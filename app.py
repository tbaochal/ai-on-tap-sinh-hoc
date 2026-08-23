import streamlit as st
import streamlit.components.v1 as components
import json
import os
import io
import pandas as pd
import uuid
import time
import re
import hashlib
import random
import urllib.parse
import urllib.request
import xml.etree.ElementTree as ET
import unicodedata
import base64
from datetime import datetime, timezone, timedelta
from google import genai
from google.genai import types

try:
    from supabase import create_client
except Exception:
    create_client = None

from muc_do_nhan_thuc import DONG_TU_MUC_DO, xac_dinh_muc_do


# ==========================================================
# CẤU HÌNH
# Bản TRẠM SINH HỌC: giữ nguyên đồng bộ ngân hàng + ảnh câu hỏi + cập nhật lời giải.
# ==========================================================
st.set_page_config(
    page_title="Trạm Sinh học",
    page_icon="🧬",
    layout="wide"
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

YCCD_PATH = os.path.join(BASE_DIR, "yccd.json")
BANK_PATH = os.path.join(BASE_DIR, "ngan_hang_cau_hoi.json")
EXAM_TEMPLATE_PATH = os.path.join(BASE_DIR, "mau_de.json")
EXAM_PATH = os.path.join(BASE_DIR, "de_da_tao.json")
MATRIX_TEST_PATH = os.path.join(BASE_DIR, "dot_kiem_tra_ma_tran.json")
HS_HISTORY_PATH = os.path.join(BASE_DIR, "lich_su_luyen_tap_hoc_sinh.json")
HS_PROFILE_PATH = os.path.join(BASE_DIR, "ho_so_nang_luc_hoc_sinh.json")
STUDENT_PATH = os.path.join(BASE_DIR, "danh_sach_hoc_sinh.json")
SEED_BANK_PATH = os.path.join(BASE_DIR, "ngan_hang_hat_giong.json")
# Dữ liệu trực quan của NGÂN HÀNG HẠT GIỐNG được lưu riêng.
# Không dùng/chỉnh thư mục dữ liệu của Ngân hàng tốt nghiệp.
SEED_MEDIA_DIR = os.path.join(BASE_DIR, "du_lieu_truc_quan_hat_giong")
SEED_SOURCE_DIR = os.path.join(BASE_DIR, "hat_giong_nguon")
GV_PROFILE_PATH = os.path.join(BASE_DIR, "ho_so_giao_vien.json")
GV_AVATAR_DIR = os.path.join(BASE_DIR, "anh_dai_dien_giao_vien")
GRAD_MEDIA_DIR = os.path.join(BASE_DIR, "du_lieu_truc_quan_tot_nghiep")
GRAD_REAL_BANK_PATH = os.path.join(BASE_DIR, "ngan_hang_tot_nghiep_thuc_te.json")
GRAD_SOURCE_DIR = os.path.join(BASE_DIR, "de_tot_nghiep_nguon")
GRAD_SCOPE_MEMORY_PATH = os.path.join(BASE_DIR, "ghi_nho_phan_loai_pham_vi_tot_nghiep.json")
os.makedirs(GRAD_MEDIA_DIR, exist_ok=True)
os.makedirs(GRAD_SOURCE_DIR, exist_ok=True)
os.makedirs(SEED_MEDIA_DIR, exist_ok=True)
os.makedirs(SEED_SOURCE_DIR, exist_ok=True)
os.makedirs(GV_AVATAR_DIR, exist_ok=True)

MODEL_AI = "gemini-3.6-flash"

# Chế độ triển khai:
# - full: trang chủ cho chọn GV/HS (dùng khi chạy nội bộ)
# - student: chỉ render giao diện học sinh, không có đường vào khu vực GV
# - teacher: chỉ render giao diện giáo viên
try:
    APP_MODE = str(st.secrets.get("APP_MODE", "full") or "full").strip().lower()
except Exception:
    APP_MODE = "full"
if APP_MODE not in {"full", "student", "teacher"}:
    APP_MODE = "full"


# ==========================================================
# DỮ LIỆU DÙNG CHUNG: SUPABASE LÀ NGUỒN CHÍNH, JSON LÀ DỰ PHÒNG
# ==========================================================
# Secret key chỉ được đọc ở phía server từ st.secrets.
# Không đưa key vào mã nguồn và không hiển thị ra giao diện.
_SUPABASE_CLIENT = None
_SUPABASE_TRIED = False


def _supabase_client():
    global _SUPABASE_CLIENT, _SUPABASE_TRIED

    if _SUPABASE_TRIED:
        return _SUPABASE_CLIENT

    _SUPABASE_TRIED = True

    if create_client is None:
        return None

    try:
        url = str(st.secrets.get("SUPABASE_URL", "") or "").strip()
        key = str(st.secrets.get("SUPABASE_SECRET_KEY", "") or "").strip()
        if not url or not key:
            return None
        _SUPABASE_CLIENT = create_client(url, key)
    except Exception:
        _SUPABASE_CLIENT = None

    return _SUPABASE_CLIENT


def _doc_json_local(path, default=None):
    try:
        if not os.path.exists(path) or os.path.getsize(path) == 0:
            return default
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default


def _luu_json_local(path, data):
    try:
        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False


def _document_key(path):
    return os.path.basename(str(path or "")).strip()


def _doc_document_shared(path, default=None):
    """Đọc app_documents trên Supabase; lỗi mạng thì quay về JSON local."""
    client_sb = _supabase_client()
    key = _document_key(path)

    if client_sb is not None and key:
        try:
            res = (
                client_sb.table("app_documents")
                .select("data")
                .eq("document_key", key)
                .limit(1)
                .execute()
            )
            rows = getattr(res, "data", None) or []
            if rows:
                data = rows[0].get("data")
                if data is not None:
                    return data
        except Exception:
            pass

    return _doc_json_local(path, default)


def _luu_document_shared(path, data):
    """Ghi Supabase và đồng thời giữ JSON local làm bản dự phòng."""
    local_ok = _luu_json_local(path, data)
    cloud_ok = False
    client_sb = _supabase_client()
    key = _document_key(path)

    if client_sb is not None and key:
        try:
            client_sb.table("app_documents").upsert(
                {
                    "document_key": key,
                    "data": data,
                    "updated_at": datetime.now(timezone.utc).isoformat(),
                },
                on_conflict="document_key",
            ).execute()
            cloud_ok = True
        except Exception:
            cloud_ok = False

    return cloud_ok or local_ok


def _doc_students_shared():
    client_sb = _supabase_client()
    if client_sb is not None:
        try:
            res = (
                client_sb.table("students")
                .select("student_id,full_name,class_name,active,data")
                .order("student_id")
                .execute()
            )
            rows = getattr(res, "data", None) or []
            if rows:
                ds = []
                for row in rows:
                    item = dict(row.get("data") or {})
                    item.setdefault("ma_hoc_sinh", row.get("student_id", ""))
                    item.setdefault("ho_ten", row.get("full_name", ""))
                    item.setdefault("lop", row.get("class_name", ""))
                    if not str(item.get("trang_thai", "")).strip():
                        item["trang_thai"] = (
                            "Đang học" if row.get("active", True) else "Tạm khóa"
                        )
                    ds.append(item)
                return ds
        except Exception:
            pass

    data = _doc_json_local(STUDENT_PATH, [])
    return data if isinstance(data, list) else []


def _luu_students_shared(ds):
    ds = list(ds or [])
    local_ok = _luu_json_local(STUDENT_PATH, ds)
    cloud_ok = False
    client_sb = _supabase_client()

    if client_sb is not None:
        try:
            rows = []
            now_iso = datetime.now(timezone.utc).isoformat()
            for hs in ds:
                if not isinstance(hs, dict):
                    continue
                sid = str(
                    hs.get("ma_hoc_sinh")
                    or hs.get("hoc_sinh_id")
                    or hs.get("student_id")
                    or ""
                ).strip()
                if not sid:
                    continue
                trang_thai = str(hs.get("trang_thai", "Đang học") or "").strip().casefold()
                active = trang_thai not in {
                    "tạm khóa", "tam khoa", "nghỉ học", "nghi hoc",
                    "đã nghỉ", "da nghi", "inactive"
                }
                rows.append({
                    "student_id": sid,
                    "full_name": str(hs.get("ho_ten", "") or "").strip(),
                    "class_name": str(hs.get("lop", "") or "").strip(),
                    "active": active,
                    "data": hs,
                    "updated_at": now_iso,
                })

            for i in range(0, len(rows), 200):
                batch = rows[i:i + 200]
                if batch:
                    client_sb.table("students").upsert(
                        batch, on_conflict="student_id"
                    ).execute()
            cloud_ok = True
        except Exception:
            cloud_ok = False

    return cloud_ok or local_ok


def _xoa_student_shared(student_id):
    """Chỉ xóa hồ sơ trong bảng students; lịch sử làm bài được giữ nguyên."""
    sid = str(student_id or "").strip()
    if not sid:
        return False
    client_sb = _supabase_client()
    if client_sb is None:
        return False
    try:
        client_sb.table("students").delete().eq("student_id", sid).execute()
        return True
    except Exception:
        return False


def _doc_attempts_shared():
    client_sb = _supabase_client()
    if client_sb is not None:
        try:
            res = (
                client_sb.table("student_attempts")
                .select("id,submitted_at,data")
                .order("submitted_at")
                .execute()
            )
            rows = getattr(res, "data", None) or []
            if rows:
                ds = []
                for row in rows:
                    item = dict(row.get("data") or {})
                    # Dữ liệu cũ khi di chuyển có thể chưa có id trong JSON gốc.
                    # Bơm id của hàng Supabase vào để lần ghi sau vẫn upsert đúng hàng.
                    if not str(item.get("id", "")).strip():
                        item["id"] = str(row.get("id", "") or "").strip()
                    if not str(item.get("thoi_gian_iso", "")).strip():
                        item["thoi_gian_iso"] = str(row.get("submitted_at", "") or "").strip()
                    ds.append(item)
                return ds
        except Exception:
            pass

    data = _doc_json_local(HS_HISTORY_PATH, [])
    return data if isinstance(data, list) else []


def _attempt_uuid(value, item):
    try:
        return str(uuid.UUID(str(value)))
    except Exception:
        raw = json.dumps(item, ensure_ascii=False, sort_keys=True, default=str)
        return str(uuid.uuid5(uuid.NAMESPACE_URL, raw))


def _luu_attempts_shared(ds):
    """Upsert từng lượt làm bài, không xóa lượt của học sinh khác."""
    ds = list(ds or [])
    local_ok = _luu_json_local(HS_HISTORY_PATH, ds)
    cloud_ok = False
    client_sb = _supabase_client()

    if client_sb is not None:
        try:
            rows = []
            for lan in ds:
                if not isinstance(lan, dict):
                    continue
                sid = str(
                    lan.get("hoc_sinh_id")
                    or lan.get("ma_hoc_sinh")
                    or ""
                ).strip()
                if not sid:
                    continue

                row_id = _attempt_uuid(lan.get("id"), lan)
                lan = dict(lan)
                lan["id"] = row_id
                pham_vi = lan.get("pham_vi", {}) or {}
                if not isinstance(pham_vi, dict):
                    pham_vi = {}

                submitted_at = str(
                    lan.get("nop_bai_iso")
                    or lan.get("thoi_gian_iso")
                    or datetime.now(timezone.utc).isoformat()
                ).strip()

                try:
                    score = lan.get("diem_chinh_thuc", lan.get("diem"))
                    score = float(score) if score is not None else None
                except Exception:
                    score = None
                try:
                    score_scale = lan.get("thang_diem", 10)
                    score_scale = float(score_scale) if score_scale is not None else None
                except Exception:
                    score_scale = None

                rows.append({
                    "id": row_id,
                    "student_id": sid,
                    "class_name": str(
                        lan.get("lop") or pham_vi.get("lop") or ""
                    ).strip(),
                    "mode": str(lan.get("che_do", "") or "").strip(),
                    "exam_id": str(
                        pham_vi.get("de_id") or pham_vi.get("mau_id") or ""
                    ).strip() or None,
                    "test_session_id": str(
                        pham_vi.get("dot_kiem_tra_id") or ""
                    ).strip() or None,
                    "submitted_at": submitted_at,
                    "score": score,
                    "score_scale": score_scale,
                    "data": lan,
                })

            for i in range(0, len(rows), 100):
                batch = rows[i:i + 100]
                if batch:
                    client_sb.table("student_attempts").upsert(
                        batch, on_conflict="id"
                    ).execute()
            cloud_ok = True
        except Exception:
            cloud_ok = False

    return cloud_ok or local_ok


# Ảnh/sơ đồ của đề thật: hiển thị vừa đủ để đọc, không kéo tràn toàn bộ màn hình.
GRAD_IMAGE_DISPLAY_WIDTH = 620
# Ảnh hạt giống dùng kích thước hiển thị tương tự, nhưng lưu/đọc độc lập.
SEED_IMAGE_DISPLAY_WIDTH = 620


# ==========================================================
# THÀNH PHẦN NĂNG LỰC SINH HỌC - TÊN CHÍNH THỨC
# ==========================================================
THANH_PHAN_NANG_LUC = [
    "Nhận thức sinh học",
    "Tìm hiểu thế giới sống",
    "Vận dụng kiến thức, kĩ năng đã học"
]


def goi_y_thanh_phan_nang_luc(yccd, muc_do=None, dang_cau=None):

    text = " ".join(
        str(yccd or "").strip().split()
    ).casefold()

    # Ưu tiên quy trình nghiên cứu / thực nghiệm / điều tra / xử lí dữ liệu
    tu_khoa_tim_hieu = [
        "thực hành", "thí nghiệm", "điều tra", "khảo sát",
        "quan sát", "thiết kế", "xây dựng giả thuyết",
        "giả thuyết", "thu thập", "xử lí dữ liệu",
        "xử lý dữ liệu", "nghiên cứu", "dự án",
        "lập kế hoạch", "báo cáo", "thực nghiệm"
    ]

    if any(tu in text for tu in tu_khoa_tim_hieu):
        return "Tìm hiểu thế giới sống"

    # Ưu tiên thực tiễn / giải pháp / đánh giá / ứng dụng
    tu_khoa_van_dung = [
        "vận dụng", "ứng dụng", "thực tiễn", "giải pháp",
        "đề xuất", "phòng", "bảo vệ", "đánh giá",
        "phản biện", "giải thích hiện tượng", "xử lí tình huống",
        "xử lý tình huống"
    ]

    if any(tu in text for tu in tu_khoa_van_dung):
        return "Vận dụng kiến thức, kĩ năng đã học"

    # Đúng/Sai tự động ưu tiên năng lực 2,3 theo yêu cầu thiết kế
    if dang_cau == "Đúng / Sai":
        if muc_do == "Vận dụng":
            return "Vận dụng kiến thức, kĩ năng đã học"
        return "Tìm hiểu thế giới sống"

    return "Nhận thức sinh học"


def cac_nang_luc_phu_hop(yccd, muc_do=None, dang_cau=None):
    """
    Trả về danh sách năng lực phù hợp theo YCCĐ.
    Không đánh số NL1/NL2/NL3, chỉ dùng tên chính thức.
    """
    goi_y = goi_y_thanh_phan_nang_luc(
        yccd,
        muc_do,
        dang_cau
    )

    if dang_cau == "Đúng / Sai":
        # Đúng/Sai ưu tiên 2 thành phần năng lực này.
        return [
            "Tìm hiểu thế giới sống",
            "Vận dụng kiến thức, kĩ năng đã học"
        ]

    # Câu thường: ưu tiên đúng năng lực gợi ý, nhưng cho phép fallback
    ds = [goi_y] + [
        x for x in THANH_PHAN_NANG_LUC
        if x != goi_y
    ]
    return ds



# ==========================================================
# GEMINI
# ==========================================================
try:
    client = genai.Client(
        api_key=st.secrets["GEMINI_API_KEY"]
    )
except Exception:
    st.error(
        "Không đọc được GEMINI_API_KEY trong "
        ".streamlit/secrets.toml"
    )
    st.stop()


# ==========================================================
# ĐỌC YCCĐ
# ==========================================================
KHO_YCCD = _doc_document_shared(YCCD_PATH, None)

if not isinstance(KHO_YCCD, dict):
    st.error(
        "Không đọc được kho YCCĐ từ Supabase hoặc file yccd.json dự phòng."
    )
    st.stop()


# ==========================================================
# NGÂN HÀNG CÂU HỎI
# ==========================================================
def doc_ngan_hang():

    data = _doc_document_shared(BANK_PATH, [])

    if isinstance(data, list):
        # Bổ sung mã chỉ báo NT/TH/VD cho câu cũ ngay khi đọc.
        # Chỉ dùng bộ quy tắc cục bộ, không gọi API.
        try:
            return gan_chi_bao_cho_ngan_hang_hien_co(data)
        except NameError:
            return data

    return []


def luu_ngan_hang(data):
    return _luu_document_shared(BANK_PATH, data)


# ==========================================================
# GIAO DIỆN & HỒ SƠ GIÁO VIÊN
# ==========================================================
DEFAULT_GV_PROFILE = {
    "ten": "Giáo viên Sinh học",
    "chuc_vu": "Giáo viên",
    "don_vi": "",
    "loi_chao": "Khơi gợi tư duy khoa học, nuôi dưỡng niềm yêu thích Sinh học.",
    "avatar_path": "",
    # Ảnh đại diện được nhúng vào hồ sơ trên Supabase để không mất khi
    # Streamlit Cloud restart/redeploy hoặc filesystem tạm bị reset.
    "avatar_base64": "",
    "avatar_mime_type": ""
}


def doc_ho_so_giao_vien():
    data = dict(DEFAULT_GV_PROFILE)
    try:
        saved = _doc_document_shared(GV_PROFILE_PATH, {})
        if isinstance(saved, dict):
            data.update({k: v for k, v in saved.items() if k in data})
            # Tương thích các tên trường từng dùng ở phiên bản cũ.
            if not str(data.get("avatar_base64", "") or "").strip():
                data["avatar_base64"] = str(
                    saved.get("avatar_data")
                    or saved.get("avatar_data_base64")
                    or saved.get("image_base64")
                    or ""
                ).strip()
            if not str(data.get("avatar_mime_type", "") or "").strip():
                data["avatar_mime_type"] = str(
                    saved.get("avatar_mime")
                    or saved.get("mime_type")
                    or ""
                ).strip()
    except Exception:
        pass
    return data


def luu_ho_so_giao_vien(profile):
    return _luu_document_shared(GV_PROFILE_PATH, profile)


def _tim_avatar_local_mac_dinh():
    """Tìm ảnh đại diện có sẵn trong thư mục dự án để phục hồi sau deploy."""
    try:
        if not os.path.isdir(GV_AVATAR_DIR):
            return ""
        uu_tien = [
            "avatar_giao_vien.png", "avatar_giao_vien.jpg", "avatar_giao_vien.jpeg",
            "avatar.png", "avatar.jpg", "avatar.jpeg"
        ]
        for ten in uu_tien:
            p = os.path.join(GV_AVATAR_DIR, ten)
            if os.path.isfile(p):
                return p
        for ten in sorted(os.listdir(GV_AVATAR_DIR)):
            if str(ten).lower().endswith((".png", ".jpg", ".jpeg")):
                p = os.path.join(GV_AVATAR_DIR, ten)
                if os.path.isfile(p):
                    return p
    except Exception:
        pass
    return ""


def _nang_cap_avatar_profile_ben_vung(profile):
    """Tự chuyển avatar local sang base64/Supabase một lần nếu có thể."""
    p = dict(profile or {})
    if str(p.get("avatar_base64", "") or "").strip():
        return p

    path = str(p.get("avatar_path", "") or "").strip()
    if not path or not os.path.exists(path):
        path = _tim_avatar_local_mac_dinh()

    if not path or not os.path.exists(path):
        return p

    try:
        with open(path, "rb") as f:
            raw = f.read()
        if not raw:
            return p
        ext = os.path.splitext(path)[1].lower()
        p["avatar_path"] = path
        p["avatar_base64"] = base64.b64encode(raw).decode("ascii")
        p["avatar_mime_type"] = "image/png" if ext == ".png" else "image/jpeg"
        # Ghi lại Supabase + local để lần restart tiếp theo không phụ thuộc filesystem.
        luu_ho_so_giao_vien(p)
    except Exception:
        pass
    return p


def _avatar_data_uri(path="", avatar_base64="", mime_type=""):
    """Ưu tiên avatar nhúng trong Supabase, sau đó mới fallback file local."""
    try:
        b64 = str(avatar_base64 or "").strip()
        if b64:
            if b64.startswith("data:"):
                return b64
            mime = str(mime_type or "").strip() or "image/jpeg"
            return f"data:{mime};base64,{b64}"

        if not path or not os.path.exists(path):
            return ""
        ext = os.path.splitext(path)[1].lower()
        mime = "image/png" if ext == ".png" else "image/jpeg"
        with open(path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("ascii")
        return f"data:{mime};base64,{b64}"
    except Exception:
        return ""


def _html_escape(value):
    import html
    return html.escape(str(value or ""))


st.markdown(
    """
    <style>
    :root { --bio-blue:#2563eb; --bio-cyan:#0891b2; --bio-green:#16a34a; --bio-ink:#172033; }
    .stApp {
        background: radial-gradient(circle at 92% 8%, rgba(14,165,233,.09), transparent 26rem),
                    radial-gradient(circle at 4% 88%, rgba(34,197,94,.06), transparent 28rem), #fbfdff;
    }
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #f8fbff 0%, #eef6ff 52%, #f7fff9 100%);
        border-right: 1px solid #dbeafe;
    }
    .stButton > button, .stDownloadButton > button {
        border-radius: 12px !important; font-weight: 700 !important; min-height: 2.8rem;
        transition: all .18s ease;
    }
    .stButton > button:hover, .stDownloadButton > button:hover {
        transform: translateY(-1px); box-shadow: 0 8px 20px rgba(37,99,235,.12);
    }
    button[kind="primary"] { border: 0 !important; background: linear-gradient(90deg,#2563eb,#0891b2) !important; }
    [data-testid="stMetric"] {
        background: rgba(255,255,255,.88); border:1px solid #e2e8f0; border-radius:15px;
        padding:.6rem .8rem; box-shadow:0 6px 22px rgba(15,23,42,.04);
    }
    div[data-testid="stExpander"] { border-radius:14px !important; border-color:#dbeafe !important; background:rgba(255,255,255,.72); }
    .bio-hero {
        position:relative; overflow:hidden; border-radius:26px; padding:2.1rem 2.2rem; margin:.2rem 0 1.25rem 0;
        color:white; background:linear-gradient(120deg,#1d4ed8 0%,#0284c7 52%,#0f9f79 100%);
        box-shadow:0 18px 50px rgba(37,99,235,.20);
    }
    .bio-hero:after { content:"🧬"; position:absolute; right:1.6rem; top:-.5rem; font-size:8rem; opacity:.12; transform:rotate(-12deg); }
    .bio-hero h1 { margin:0; font-size:2.25rem; line-height:1.12; color:white; }
    .bio-hero p { margin:.7rem 0 0 0; font-size:1.02rem; opacity:.94; max-width:70%; }
    .bio-chip { display:inline-block; padding:.35rem .75rem; margin-bottom:.75rem; border:1px solid rgba(255,255,255,.35); border-radius:999px; background:rgba(255,255,255,.13); font-weight:700; font-size:.86rem; }
    .bio-feature { min-height:118px; padding:1rem 1.05rem; border-radius:17px; background:rgba(255,255,255,.92); border:1px solid #e2e8f0; box-shadow:0 8px 28px rgba(15,23,42,.05); }
    .bio-feature b { color:#1e3a8a; }
    .teacher-card { margin:.15rem 0 .7rem 0; padding:.9rem .8rem; border-radius:18px; color:#10213b; background:linear-gradient(145deg,rgba(255,255,255,.98),rgba(239,246,255,.95)); border:1px solid #dbeafe; box-shadow:0 9px 25px rgba(37,99,235,.08); text-align:center; }
    .teacher-avatar { width:82px; height:82px; object-fit:cover; border-radius:50%; border:4px solid white; box-shadow:0 5px 18px rgba(15,23,42,.18); margin:0 auto .55rem auto; display:block; }
    .teacher-avatar-fallback { width:82px; height:82px; border-radius:50%; display:flex; align-items:center; justify-content:center; margin:0 auto .55rem auto; background:linear-gradient(145deg,#dbeafe,#dcfce7); border:4px solid white; box-shadow:0 5px 18px rgba(15,23,42,.14); font-size:2.4rem; }
    .teacher-name { font-size:1.04rem; font-weight:800; color:#172554; margin-top:.1rem; }
    .teacher-meta { font-size:.82rem; color:#64748b; margin-top:.18rem; }
    .teacher-quote { font-size:.78rem; color:#475569; margin-top:.5rem; line-height:1.35; }
    .section-title-card { padding:.82rem 1rem; border-radius:15px; background:linear-gradient(90deg,rgba(219,234,254,.75),rgba(220,252,231,.58)); border:1px solid #dbeafe; margin-bottom:1rem; }
    </style>
    """,
    unsafe_allow_html=True
)


def hien_thi_the_giao_vien_sidebar():
    profile = _nang_cap_avatar_profile_ben_vung(doc_ho_so_giao_vien())
    avatar_uri = _avatar_data_uri(
        profile.get("avatar_path"),
        profile.get("avatar_base64", ""),
        profile.get("avatar_mime_type", "")
    )
    if avatar_uri:
        avatar_html = f'<img class="teacher-avatar" src="{avatar_uri}" alt="Ảnh giáo viên">'
    else:
        avatar_html = '<div class="teacher-avatar-fallback">👩‍🏫</div>'
    meta = " • ".join([x for x in [profile.get("chuc_vu", ""), profile.get("don_vi", "")] if str(x).strip()])
    st.markdown(
        f"""
        <div class="teacher-card">
            {avatar_html}
            <div class="teacher-name">{_html_escape(profile.get("ten"))}</div>
            <div class="teacher-meta">{_html_escape(meta or "Sinh học THPT")}</div>
            <div class="teacher-quote">{_html_escape(profile.get("loi_chao"))}</div>
        </div>
        """,
        unsafe_allow_html=True
    )
    with st.expander("⚙️ Hồ sơ giáo viên", expanded=False):
        ten = st.text_input("Tên hiển thị", value=profile.get("ten", ""), key="gv_profile_name")
        chuc_vu = st.text_input("Chức vụ / vai trò", value=profile.get("chuc_vu", ""), key="gv_profile_role")
        don_vi = st.text_input("Trường / đơn vị", value=profile.get("don_vi", ""), key="gv_profile_unit")
        loi_chao = st.text_area("Dòng giới thiệu ngắn", value=profile.get("loi_chao", ""), height=80, key="gv_profile_quote")
        avatar_file = st.file_uploader("Ảnh đại diện", type=["png","jpg","jpeg"], accept_multiple_files=False, key="gv_profile_avatar")
        if st.button("💾 Lưu hồ sơ", use_container_width=True, key="gv_profile_save"):
            avatar_path = profile.get("avatar_path", "")
            avatar_b64 = str(profile.get("avatar_base64", "") or "").strip()
            avatar_mime = str(profile.get("avatar_mime_type", "") or "").strip()

            if avatar_file is not None:
                raw_avatar = bytes(avatar_file.getbuffer())
                ext = os.path.splitext(avatar_file.name)[1].lower()
                if ext not in [".png", ".jpg", ".jpeg"]:
                    ext = ".jpg"
                avatar_mime = (
                    "image/png" if ext == ".png"
                    else "image/jpeg"
                )
                avatar_b64 = base64.b64encode(raw_avatar).decode("ascii")

                # Vẫn lưu local để chạy nhanh khi còn file; Supabase/base64 mới là bản bền vững.
                avatar_path = os.path.join(GV_AVATAR_DIR, "avatar_giao_vien" + ext)
                try:
                    with open(avatar_path, "wb") as f:
                        f.write(raw_avatar)
                except Exception:
                    avatar_path = ""

            # Nếu hồ sơ cũ mới chỉ có file local, tự nâng cấp sang base64 khi bấm Lưu.
            if not avatar_b64 and avatar_path and os.path.exists(avatar_path):
                try:
                    with open(avatar_path, "rb") as f:
                        raw_avatar = f.read()
                    avatar_b64 = base64.b64encode(raw_avatar).decode("ascii")
                    ext = os.path.splitext(avatar_path)[1].lower()
                    avatar_mime = "image/png" if ext == ".png" else "image/jpeg"
                except Exception:
                    pass

            luu_ho_so_giao_vien({
                "ten": ten.strip() or DEFAULT_GV_PROFILE["ten"],
                "chuc_vu": chuc_vu.strip(),
                "don_vi": don_vi.strip(),
                "loi_chao": loi_chao.strip(),
                "avatar_path": avatar_path,
                "avatar_base64": avatar_b64,
                "avatar_mime_type": avatar_mime
            })
            st.success("Đã lưu hồ sơ giáo viên và ảnh đại diện bền vững trên Supabase.")
            st.rerun()

# ==========================================================
# SESSION STATE
# ==========================================================
if "vai_tro" not in st.session_state:
    st.session_state.vai_tro = None

if "yccd_da_chon" not in st.session_state:
    st.session_state.yccd_da_chon = []

if "cau_hinh_yccd" not in st.session_state:
    st.session_state.cau_hinh_yccd = {}

if "bang_dac_ta" not in st.session_state:
    st.session_state.bang_dac_ta = []

if "cau_hoi_ai" not in st.session_state:
    st.session_state.cau_hoi_ai = []

if "trang_thai_duyet" not in st.session_state:
    st.session_state.trang_thai_duyet = {}


if "ket_qua_kiem_dinh" not in st.session_state:
    st.session_state.ket_qua_kiem_dinh = {}


if "cache_kiem_dinh" not in st.session_state:
    st.session_state.cache_kiem_dinh = {}

if "ma_tran_de_gv" not in st.session_state:
    st.session_state.ma_tran_de_gv = []

if "de_xem_truoc" not in st.session_state:
    st.session_state.de_xem_truoc = None


# ==========================================================
# HÀM PHỤ
# ==========================================================
def tao_id_yccd(khoi, chuong, bai, yccd):

    return f"{khoi}|||{chuong}|||{bai}|||{yccd}"


def item_da_chon(item):

    return item in st.session_state.yccd_da_chon


def cau_hinh_mac_dinh(yccd=""):

    muc_do_tu_dong = xac_dinh_muc_do(yccd)

    dang_mac_dinh = "Trắc nghiệm 4 lựa chọn"

    return {
        "Số câu": 1,
        "Mức độ": muc_do_tu_dong,
        "Dạng câu hỏi": dang_mac_dinh,
        "Thành phần năng lực": goi_y_thanh_phan_nang_luc(
            yccd,
            muc_do_tu_dong,
            dang_mac_dinh
        ),
        "Nguồn tham chiếu": [
            "SGK / Chương trình"
        ]
    }


def lay_cau_hinh(item):

    yccd_id = tao_id_yccd(
        item["Khối"],
        item["Chương"],
        item["Bài"],
        item["YCCĐ"]
    )

    if yccd_id not in st.session_state.cau_hinh_yccd:
        st.session_state.cau_hinh_yccd[
    yccd_id
] = cau_hinh_mac_dinh(
    item["YCCĐ"]
)

    return (
        yccd_id,
        st.session_state.cau_hinh_yccd[yccd_id]
    )


# ==========================================================
# XÓA 1 YCCĐ
# ==========================================================
def xoa_mot_yccd(item):

    yccd_id = tao_id_yccd(
        item["Khối"],
        item["Chương"],
        item["Bài"],
        item["YCCĐ"]
    )

    if item in st.session_state.yccd_da_chon:
        st.session_state.yccd_da_chon.remove(item)

    if yccd_id in st.session_state.cau_hinh_yccd:
        del st.session_state.cau_hinh_yccd[yccd_id]

    # Xóa trạng thái checkbox
    checkbox_key = f"checkbox_{yccd_id}"

    if checkbox_key in st.session_state:
        del st.session_state[checkbox_key]

    # Xóa trạng thái widget
    cac_key = [
        f"so_cau_{yccd_id}",
        f"mucdo_{yccd_id}",
        f"dang_{yccd_id}",
        f"nangluc_{yccd_id}",
        f"nguon_{yccd_id}"
    ]

    for key in cac_key:
        if key in st.session_state:
            del st.session_state[key]

    st.session_state.bang_dac_ta = []
    st.session_state.cau_hoi_ai = []
    st.session_state.trang_thai_duyet = {}


# ==========================================================
# XÓA TOÀN BỘ
# ==========================================================
def xoa_toan_bo():

    st.session_state.yccd_da_chon = []
    st.session_state.cau_hinh_yccd = {}
    st.session_state.bang_dac_ta = []
    st.session_state.cau_hoi_ai = []
    st.session_state.trang_thai_duyet = {}

    # Xóa trạng thái widget YCCĐ cũ
    keys = list(st.session_state.keys())

    for key in keys:
        if (
            key.startswith("checkbox_")
            or key.startswith("so_cau_")
            or key.startswith("mucdo_")
            or key.startswith("dang_")
            or key.startswith("nangluc_")
            or key.startswith("nangluc_ds_")
            or key.startswith("nguon_")
        ):
            del st.session_state[key]


# ==========================================================
# BẢNG ĐẶC TẢ
# ==========================================================
def tao_bang_dac_ta():

    bang = []

    for i, item in enumerate(
        st.session_state.yccd_da_chon,
        start=1
    ):

        _, cau_hinh = lay_cau_hinh(item)

        bang.append({
            "STT": i,
            "Khối": item["Khối"],
            "Chương": item["Chương"],
            "Bài": item["Bài"],
            "YCCĐ": item["YCCĐ"],
            "Số câu": cau_hinh["Số câu"],
            "Mức độ": cau_hinh["Mức độ"],
            "Dạng câu hỏi": cau_hinh["Dạng câu hỏi"],
            "Thành phần năng lực": cau_hinh.get(
                "Thành phần năng lực",
                goi_y_thanh_phan_nang_luc(
                    item["YCCĐ"],
                    cau_hinh.get("Mức độ"),
                    cau_hinh.get("Dạng câu hỏi")
                )
            ),
            "Nguồn tham chiếu": ", ".join(
                cau_hinh["Nguồn tham chiếu"]
            )
        })

    st.session_state.bang_dac_ta = bang

    return bang

def doc_tai_lieu_giao_vien(files):

    noi_dung = []

    if not files:
        return ""

    for file_path in files:

        try:
            # file_path bây giờ là đường dẫn thật trên máy
            ten_file = os.path.basename(file_path)
            ten_file_lower = ten_file.lower()

            # ==========================================
            # TXT
            # ==========================================
            if ten_file_lower.endswith(".txt"):

                with open(
                    file_path,
                    "r",
                    encoding="utf-8",
                    errors="ignore"
                ) as f:
                    text = f.read()

            # ==========================================
            # PDF
            # ==========================================
            elif ten_file_lower.endswith(".pdf"):

                from pypdf import PdfReader

                reader = PdfReader(file_path)

                text = "\n".join(
                    page.extract_text() or ""
                    for page in reader.pages
                )

            # ==========================================
            # WORD
            # ==========================================
            elif ten_file_lower.endswith(".docx"):

                from docx import Document

                doc = Document(file_path)

                text = "\n".join(
                    p.text
                    for p in doc.paragraphs
                    if p.text.strip()
                )

            else:
                continue

            if text.strip():

                noi_dung.append(
                    f"\n===== TÀI LIỆU: {ten_file} =====\n{text}"
                )

        except Exception as e:

            st.warning(
                f"Không đọc được {os.path.basename(str(file_path))}: {e}"
            )

    return "\n".join(noi_dung)

# ==========================================================
# TRUY XUẤT NGUỒN THẬT - KHÔNG CHO AI TỰ BỊA NGUỒN
# ==========================================================
def _http_get_json(url, timeout=12):

    req = urllib.request.Request(
        url,
        headers={
            "User-Agent": (
                "SinhHocQuestionBank/1.0 "
                "(educational-use)"
            )
        }
    )

    with urllib.request.urlopen(
        req,
        timeout=timeout
    ) as response:

        return json.loads(
            response.read().decode(
                "utf-8",
                errors="ignore"
            )
        )


def _http_get_text(url, timeout=12):

    req = urllib.request.Request(
        url,
        headers={
            "User-Agent": (
                "SinhHocQuestionBank/1.0 "
                "(educational-use)"
            )
        }
    )

    with urllib.request.urlopen(
        req,
        timeout=timeout
    ) as response:

        return response.read().decode(
            "utf-8",
            errors="ignore"
        )


def _bo_dau_nguon(text):

    text = unicodedata.normalize(
        "NFD",
        str(text or "")
    )

    return "".join(
        c
        for c in text
        if unicodedata.category(c) != "Mn"
    ).lower()


def tao_truy_van_khoa_hoc(
    yccd,
    bai=""
):
    """
    Tạo truy vấn tiếng Anh bằng quy tắc cục bộ.
    Không gọi Gemini để tránh thêm quota/429.
    """

    raw = (
        str(bai or "")
        + " "
        + str(yccd or "")
    )

    raw_norm = _bo_dau_nguon(raw)

    mapping = [
        ("cap do to chuc song", "levels of biological organization"),
        ("to chuc song", "biological organization"),
        ("te bao", "cell biology"),
        ("mang sinh chat", "cell membrane"),
        ("van chuyen qua mang", "membrane transport"),
        ("enzyme", "enzyme"),
        ("enzim", "enzyme"),
        ("chuyen hoa vat chat", "metabolism"),
        ("nang luong", "bioenergetics"),
        ("quang hop", "photosynthesis"),
        ("ho hap te bao", "cellular respiration"),
        ("phan bao", "cell division"),
        ("nguyen phan", "mitosis"),
        ("giam phan", "meiosis"),
        ("nhiem sac the", "chromosome"),
        ("dna", "DNA"),
        ("rna", "RNA"),
        ("gene", "gene"),
        ("di truyen", "genetics"),
        ("dot bien", "mutation"),
        ("phien ma", "transcription"),
        ("dich ma", "translation"),
        ("tien hoa", "evolution"),
        ("chon loc tu nhien", "natural selection"),
        ("quan the", "population biology"),
        ("quan xa", "community ecology"),
        ("he sinh thai", "ecosystem"),
        ("sinh thai", "ecology"),
        ("da dang sinh hoc", "biodiversity"),
        ("vi sinh vat", "microbiology"),
        ("vi khuan", "bacteria"),
        ("virus", "virus"),
        ("mien dich", "immunology"),
        ("sinh san", "reproduction biology"),
        ("noi tiet", "endocrinology"),
        ("than kinh", "neurobiology"),
        ("tuan hoan", "circulatory physiology"),
        ("ho hap", "respiratory physiology"),
        ("tieu hoa", "digestive physiology"),
        ("bai tiet", "excretory physiology"),
        ("can bang noi moi", "homeostasis"),
        ("sinh truong", "growth biology"),
        ("phat trien", "developmental biology"),
        ("thuc vat", "plant biology"),
        ("dong vat", "animal biology"),
    ]

    terms = []

    for vi, en in mapping:
        if vi in raw_norm and en not in terms:
            terms.append(en)

    if not terms:
        terms = ["biology"]

    # Giữ truy vấn ngắn để API trả kết quả ổn định hơn.
    return " AND ".join(
        terms[:4]
    )


def _rut_gon_doan(text, max_chars=1800):

    s = " ".join(
        str(text or "").split()
    ).strip()

    if len(s) <= max_chars:
        return s

    return s[:max_chars].rsplit(
        " ",
        1
    )[0] + "…"


def _tu_khoa_lien_quan(text):

    raw = _bo_dau_nguon(text)

    stop = {
        "duoc", "va", "la", "cua", "cac", "mot",
        "trong", "ve", "cho", "tu", "den", "voi",
        "neu", "trinh", "bay", "phan", "tich",
        "giai", "thich", "xac", "dinh", "dua",
        "vao", "phat", "bieu", "mo", "ta"
    }

    words = re.findall(
        r"[a-z0-9]+",
        raw
    )

    return {
        w
        for w in words
        if len(w) >= 4
        and w not in stop
    }


def trich_doan_lien_quan(
    text,
    yccd,
    max_chars=3500
):
    """
    Lấy phần liên quan nhất trong tài liệu GV/SGK,
    thay vì nhét cả PDF vào prompt gây 429.
    """

    text = str(text or "").strip()

    if not text:
        return ""

    paragraphs = [
        " ".join(p.split())
        for p in re.split(
            r"\n\s*\n|\n",
            text
        )
        if p.strip()
    ]

    if not paragraphs:
        return _rut_gon_doan(
            text,
            max_chars
        )

    kws = _tu_khoa_lien_quan(
        yccd
    )

    scored = []

    for p in paragraphs:

        p_norm = _bo_dau_nguon(p)

        score = sum(
            1
            for kw in kws
            if kw in p_norm
        )

        scored.append(
            (
                score,
                len(p),
                p
            )
        )

    scored.sort(
        key=lambda x: (
            x[0],
            min(x[1], 800)
        ),
        reverse=True
    )

    selected = []
    total = 0

    for score, _, p in scored:

        # Nếu có từ khóa thì ưu tiên đoạn có score > 0.
        # Nếu không có đoạn khớp, vẫn lấy 1–2 đoạn đầu tốt nhất.
        if kws and score <= 0 and selected:
            continue

        room = (
            max_chars
            - total
        )

        if room <= 100:
            break

        part = _rut_gon_doan(
            p,
            min(
                1000,
                room
            )
        )

        selected.append(
            part
        )

        total += (
            len(part)
            + 2
        )

        if total >= max_chars:
            break

        if len(selected) >= 5:
            break

    return "\n\n".join(
        selected
    )


def tim_pubmed_nguon_that(
    query,
    retmax=3
):
    """
    Truy xuất thật từ NCBI PubMed bằng E-utilities.
    """

    try:

        params = urllib.parse.urlencode({
            "db": "pubmed",
            "term": query,
            "retmode": "json",
            "retmax": int(retmax),
            "sort": "relevance"
        })

        search_url = (
            "https://eutils.ncbi.nlm.nih.gov/"
            "entrez/eutils/esearch.fcgi?"
            + params
        )

        search_data = _http_get_json(
            search_url
        )

        ids = (
            search_data
            .get("esearchresult", {})
            .get("idlist", [])
        )

        if not ids:
            return []

        fetch_params = urllib.parse.urlencode({
            "db": "pubmed",
            "id": ",".join(ids),
            "retmode": "xml"
        })

        fetch_url = (
            "https://eutils.ncbi.nlm.nih.gov/"
            "entrez/eutils/efetch.fcgi?"
            + fetch_params
        )

        xml_text = _http_get_text(
            fetch_url
        )

        root = ET.fromstring(
            xml_text
        )

        ket_qua = []

        for article in root.findall(
            ".//PubmedArticle"
        ):

            pmid = article.findtext(
                ".//PMID",
                default=""
            ).strip()

            title_node = article.find(
                ".//ArticleTitle"
            )

            title = (
                "".join(
                    title_node.itertext()
                ).strip()
                if title_node is not None
                else ""
            )

            abstracts = []

            for ab in article.findall(
                ".//Abstract/AbstractText"
            ):

                ab_text = "".join(
                    ab.itertext()
                ).strip()

                if ab_text:
                    abstracts.append(
                        ab_text
                    )

            abstract = " ".join(
                abstracts
            )

            if not abstract:
                continue

            doi = ""

            for aid in article.findall(
                ".//ArticleId"
            ):
                if (
                    aid.attrib.get(
                        "IdType"
                    ) == "doi"
                ):
                    doi = (
                        aid.text
                        or ""
                    ).strip()
                    break

            url = (
                f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
                if pmid
                else ""
            )

            ket_qua.append({
                "provider": "PubMed / NCBI",
                "id": (
                    f"PMID:{pmid}"
                    if pmid
                    else "PubMed"
                ),
                "title": title,
                "url": url,
                "doi": doi,
                "excerpt": _rut_gon_doan(
                    abstract,
                    1800
                )
            })

        return ket_qua

    except Exception:
        return []


def tim_pmc_nguon_that(
    query,
    retmax=3
):
    """
    Truy xuất thật từ Europe PMC, ưu tiên bài có PMCID.
    """

    try:

        q = (
            f"({query}) "
            "AND OPEN_ACCESS:Y"
        )

        params = urllib.parse.urlencode({
            "query": q,
            "format": "json",
            "pageSize": int(retmax)
        })

        url = (
            "https://www.ebi.ac.uk/"
            "europepmc/webservices/rest/search?"
            + params
        )

        data = _http_get_json(
            url
        )

        results = (
            data.get(
                "resultList",
                {}
            )
            .get(
                "result",
                []
            )
        )

        ket_qua = []

        for r in results:

            abstract = str(
                r.get(
                    "abstractText",
                    ""
                )
                or ""
            ).strip()

            if not abstract:
                continue

            pmcid = str(
                r.get(
                    "pmcid",
                    ""
                )
                or ""
            ).strip()

            pmid = str(
                r.get(
                    "pmid",
                    ""
                )
                or ""
            ).strip()

            if pmcid:
                source_url = (
                    "https://pmc.ncbi.nlm.nih.gov/"
                    f"articles/{pmcid}/"
                )
            elif pmid:
                source_url = (
                    "https://pubmed.ncbi.nlm.nih.gov/"
                    f"{pmid}/"
                )
            else:
                source_url = ""

            ket_qua.append({
                "provider": "PubMed Central (PMC)",
                "id": (
                    pmcid
                    or (
                        f"PMID:{pmid}"
                        if pmid
                        else "PMC"
                    )
                ),
                "title": str(
                    r.get(
                        "title",
                        ""
                    )
                    or ""
                ).strip(),
                "url": source_url,
                "doi": str(
                    r.get(
                        "doi",
                        ""
                    )
                    or ""
                ).strip(),
                "excerpt": _rut_gon_doan(
                    abstract,
                    1800
                )
            })

        return ket_qua

    except Exception:
        return []


def _openalex_abstract(
    inverted
):

    if not isinstance(
        inverted,
        dict
    ):
        return ""

    positions = []

    for word, idxs in inverted.items():

        for idx in idxs or []:

            positions.append(
                (
                    int(idx),
                    str(word)
                )
            )

    positions.sort(
        key=lambda x: x[0]
    )

    return " ".join(
        word
        for _, word in positions
    )


def tim_web_khoa_hoc_uy_tin(
    query,
    retmax=3
):
    """
    Dùng OpenAlex để lấy metadata + abstract có thật.
    Không cho AI tự bịa website/DOI.
    """

    try:

        params = urllib.parse.urlencode({
            "search": query,
            "per-page": int(retmax)
        })

        url = (
            "https://api.openalex.org/works?"
            + params
        )

        data = _http_get_json(
            url
        )

        ket_qua = []

        for r in data.get(
            "results",
            []
        ):

            abstract = _openalex_abstract(
                r.get(
                    "abstract_inverted_index"
                )
            )

            if not abstract:
                continue

            doi = str(
                r.get(
                    "doi",
                    ""
                )
                or ""
            ).strip()

            source_url = (
                doi
                if doi
                else str(
                    r.get(
                        "id",
                        ""
                    )
                    or ""
                ).strip()
            )

            primary = (
                r.get(
                    "primary_location"
                )
                or {}
            )

            source = (
                primary.get(
                    "source"
                )
                or {}
            )

            provider_name = (
                source.get(
                    "display_name"
                )
                or "OpenAlex indexed source"
            )

            ket_qua.append({
                "provider": provider_name,
                "id": str(
                    r.get(
                        "id",
                        ""
                    )
                    or ""
                ).strip(),
                "title": str(
                    r.get(
                        "display_name",
                        ""
                    )
                    or ""
                ).strip(),
                "url": source_url,
                "doi": doi,
                "excerpt": _rut_gon_doan(
                    abstract,
                    1800
                )
            })

        return ket_qua

    except Exception:
        return []


def doc_file_nguon_cuc_bo(
    file_path
):

    try:

        ten_file_lower = str(
            file_path
        ).lower()

        if ten_file_lower.endswith(
            ".txt"
        ):

            with open(
                file_path,
                "r",
                encoding="utf-8",
                errors="ignore"
            ) as f:

                return f.read()

        if ten_file_lower.endswith(
            ".pdf"
        ):

            from pypdf import PdfReader

            reader = PdfReader(
                file_path
            )

            return "\n".join(
                page.extract_text()
                or ""
                for page in reader.pages
            )

        if ten_file_lower.endswith(
            ".docx"
        ):

            from docx import Document

            doc = Document(
                file_path
            )

            return "\n".join(
                p.text
                for p in doc.paragraphs
                if p.text.strip()
            )

    except Exception:
        return ""

    return ""


def tim_file_sgk_chuong_trinh(
    thu_muc="kho_tai_lieu_gv"
):

    if not os.path.isdir(
        thu_muc
    ):
        return []

    keywords = [
        "sgk",
        "sach giao khoa",
        "sinh hoc",
        "chuong trinh",
        "ctgdpt"
    ]

    ket_qua = []

    for ten_file in os.listdir(
        thu_muc
    ):

        ten_norm = _bo_dau_nguon(
            ten_file
        )

        if any(
            kw in ten_norm
            for kw in keywords
        ):

            path = os.path.join(
                thu_muc,
                ten_file
            )

            if os.path.isfile(
                path
            ):
                ket_qua.append(
                    path
                )

    return ket_qua


def thu_thap_nguon_that_cho_item(
    item,
    cau_hinh
):
    """
    Trả về DANH SÁCH NGUỒN CÓ DỮ LIỆU THẬT.
    Nếu API/tài liệu không trả dữ liệu thì không dựng nguồn giả.
    """

    nguon_da_chon = cau_hinh.get(
        "Nguồn tham chiếu",
        []
    )

    yccd = item.get(
        "YCCĐ",
        ""
    )

    bai = item.get(
        "Bài",
        ""
    )

    query = tao_truy_van_khoa_hoc(
        yccd,
        bai
    )

    records = []

    # ------------------------------------------
    # Tài liệu GV được chọn
    # ------------------------------------------
    if (
        "Tài liệu giáo viên tải lên"
        in nguon_da_chon
    ):

        for ten_file in cau_hinh.get(
            "Tài liệu giáo viên",
            []
        ):

            path = os.path.join(
                "kho_tai_lieu_gv",
                ten_file
            )

            if not os.path.isfile(
                path
            ):
                continue

            full_text = doc_file_nguon_cuc_bo(
                path
            )

            excerpt = trich_doan_lien_quan(
                full_text,
                yccd,
                3500
            )

            if excerpt:

                records.append({
                    "provider": (
                        "Tài liệu giáo viên tải lên"
                    ),
                    "id": ten_file,
                    "title": ten_file,
                    "url": (
                        "local://"
                        + ten_file
                    ),
                    "doi": "",
                    "excerpt": excerpt
                })

    # ------------------------------------------
    # SGK / Chương trình:
    # chỉ dùng khi có FILE THẬT trong kho.
    # ------------------------------------------
    if (
        "SGK / Chương trình"
        in nguon_da_chon
    ):

        for path in tim_file_sgk_chuong_trinh():

            full_text = doc_file_nguon_cuc_bo(
                path
            )

            excerpt = trich_doan_lien_quan(
                full_text,
                yccd,
                3500
            )

            if excerpt:

                ten_file = os.path.basename(
                    path
                )

                records.append({
                    "provider": (
                        "SGK / Chương trình "
                        "(tệp thật trong kho)"
                    ),
                    "id": ten_file,
                    "title": ten_file,
                    "url": (
                        "local://"
                        + ten_file
                    ),
                    "doi": "",
                    "excerpt": excerpt
                })

                # Không nhét quá nhiều SGK vào prompt.
                if sum(
                    1
                    for r in records
                    if r["provider"].startswith(
                        "SGK / Chương trình"
                    )
                ) >= 2:
                    break

    # ------------------------------------------
    # PubMed / NCBI
    # ------------------------------------------
    if (
        "PubMed" in nguon_da_chon
        or "NCBI" in nguon_da_chon
    ):

        records.extend(
            tim_pubmed_nguon_that(
                query,
                retmax=3
            )
        )

    # ------------------------------------------
    # PMC
    # ------------------------------------------
    if (
        "PubMed Central (PMC)"
        in nguon_da_chon
    ):

        records.extend(
            tim_pmc_nguon_that(
                query,
                retmax=3
            )
        )

    # ------------------------------------------
    # Web khoa học uy tín
    # ------------------------------------------
    if (
        "Nguồn web khoa học uy tín"
        in nguon_da_chon
    ):

        records.extend(
            tim_web_khoa_hoc_uy_tin(
                query,
                retmax=3
            )
        )

    # Loại trùng URL/title
    unique = []
    seen = set()

    for r in records:

        key = (
            str(
                r.get(
                    "url",
                    ""
                )
            ).strip()
            or str(
                r.get(
                    "title",
                    ""
                )
            ).strip()
        )

        if not key or key in seen:
            continue

        seen.add(
            key
        )

        unique.append(
            r
        )

    # Giới hạn để prompt không quá dài/giảm 429.
    return unique[:6]


def dinh_dang_nguon_that_cho_prompt(
    records
):

    if not records:

        return (
            "KHÔNG CÓ DỮ LIỆU NGUỒN THỰC TẾ. "
            "KHÔNG ĐƯỢC DÙNG KIẾN THỨC NGOÀI NGUỒN "
            "ĐỂ TỰ BỊA CÂU HỎI."
        )

    blocks = []

    for idx, r in enumerate(
        records,
        start=1
    ):

        blocks.append(
            "\n".join([
                f"[SOURCE_{idx}]",
                (
                    "Provider: "
                    + str(
                        r.get(
                            "provider",
                            ""
                        )
                    )
                ),
                (
                    "ID: "
                    + str(
                        r.get(
                            "id",
                            ""
                        )
                    )
                ),
                (
                    "Title: "
                    + str(
                        r.get(
                            "title",
                            ""
                        )
                    )
                ),
                (
                    "URL: "
                    + str(
                        r.get(
                            "url",
                            ""
                        )
                    )
                ),
                (
                    "DOI: "
                    + str(
                        r.get(
                            "doi",
                            ""
                        )
                    )
                ),
                "SOURCE_EXCERPT:",
                str(
                    r.get(
                        "excerpt",
                        ""
                    )
                ),
                "[/SOURCE]"
            ])
        )

    return "\n\n".join(
        blocks
    )



# ==========================================================
# PROMPT
# ==========================================================
def tao_prompt_ai():

    danh_sach_yccd = st.session_state.get(
        "yccd_da_chon",
        []
    )

    thu_muc_kho = "kho_tai_lieu_gv"

    cac_file_gv_duoc_chon = []

    for item in danh_sach_yccd:

        _, cau_hinh = lay_cau_hinh(item)

        nguon_da_chon = cau_hinh.get(
            "Nguồn tham chiếu",
            []
        )

        # Chỉ đọc tài liệu nếu YCCĐ này
        # thực sự chọn nguồn tài liệu GV
        if "Tài liệu giáo viên tải lên" in nguon_da_chon:

            ds_tai_lieu = cau_hinh.get(
                "Tài liệu giáo viên",
                []
            )

            for ten_file in ds_tai_lieu:

                duong_dan = os.path.join(
                    thu_muc_kho,
                    ten_file
                )

                if (
                    os.path.isfile(duong_dan)
                    and duong_dan not in cac_file_gv_duoc_chon
                ):
                    cac_file_gv_duoc_chon.append(
                        duong_dan
                    )

    # Không đưa toàn bộ tài liệu vào prompt.
    # Phần nguồn thật được trích theo từng YCCĐ ở vòng lặp bên dưới.
    noi_dung_tai_lieu_gv = ""


    prompt = """
Bạn là chuyên gia Sinh học THPT và chuyên gia đánh giá giáo dục.

Hãy tạo câu hỏi bám sát YCCĐ, mức độ nhận thức, dạng câu hỏi
và các nguồn tham chiếu được giáo viên lựa chọn.

QUY TẮC SỬ DỤNG NGUỒN THỰC TẾ:

- Mỗi YCCĐ sẽ có các SOURCE BLOCK chứa dữ liệu đã được hệ thống
  truy xuất thật từ tài liệu GV, tệp SGK/Chương trình trong kho,
  PubMed/NCBI, PMC hoặc nguồn khoa học đã được truy xuất.

- CHỈ được sử dụng thông tin khoa học có trong SOURCE BLOCK
  của chính YCCĐ đó để tạo nội dung câu hỏi, dữ kiện, đáp án
  và giải thích.

- KHÔNG được dùng trí nhớ chung của mô hình để bổ sung
  một dữ kiện khoa học không có trong SOURCE BLOCK.

- Không được tự bịa bài báo, tác giả, PMID, PMCID, DOI, URL,
  số liệu, kết quả thí nghiệm hoặc kết luận nghiên cứu.

- Nếu chưa truy xuất được SOURCE BLOCK phù hợp cho một YCCĐ,
  KHÔNG được bịa bài báo, số liệu nghiên cứu, PMID, DOI hoặc URL.
  Tuy nhiên vẫn được tạo câu hỏi dựa trên kiến thức chương trình phổ thông
  nếu YCCĐ thuộc kiến thức nền ổn định và câu hỏi không cần dữ liệu thực nghiệm cụ thể.
  Khi đó phải ghi "kieu_nguon": "kien_thuc_chuong_trinh".

- Tài liệu nguồn là căn cứ để PHÁT SINH CÂU HỎI MỚI.
  Không sao chép nguyên văn câu hỏi/đoạn văn từ nguồn nếu không cần thiết.

- Trường "nguon_url" phải sao chép CHÍNH XÁC URL của một SOURCE BLOCK
  thực sự được dùng. Không được tự tạo URL khác.


==================================================
QUY TẮC PHÁT SINH CÂU HỎI ĐỊNH LƯỢNG / TÍNH TOÁN
==================================================

NGUYÊN TẮC 3 TẦNG:

TẦNG 1 - ƯU TIÊN DỮ LIỆU THẬT:
- Nếu SOURCE BLOCK có bảng số liệu, số đo, tần số, tỉ lệ, kích thước,
  kết quả thí nghiệm, dữ liệu quần thể, dữ liệu sinh thái, dữ liệu gene,
  biểu đồ hoặc thông tin định lượng khác, PHẢI ưu tiên sử dụng chính
  dữ liệu thật đó để phát sinh câu hỏi.
- Không được bỏ qua dữ liệu thật đang có để tự tạo một bộ số liệu khác.
- Có thể biến đổi cách hỏi, xây dựng tình huống mới hoặc yêu cầu suy luận mới,
  nhưng các dữ kiện định lượng lấy từ nguồn phải giữ đúng bản chất và giá trị.

TẦNG 2 - CHỈ MÔ PHỎNG KHI THẬT SỰ CẦN:
- Chỉ được tạo DỮ KIỆN GIẢ ĐỊNH/MÔ PHỎNG khi:
  1) SOURCE BLOCK cung cấp cơ chế, quy luật hoặc kiến thức nền đúng;
  2) nguồn không có đủ dữ liệu định lượng phù hợp để tạo bài toán;
  3) YCCĐ có bản chất định lượng hoặc cần đánh giá năng lực tính toán/suy luận.
- Khi dùng dữ kiện mô phỏng, phải ghi nhận nội bộ:
  "kieu_du_lieu": "mo_phong"
- Số liệu mô phỏng phải:
  + hợp lí về sinh học và toán học;
  + không mâu thuẫn SOURCE BLOCK;
  + đủ dữ kiện để có một đáp án duy nhất;
  + được tự kiểm tra phép tính trước khi xuất.

TẦNG 3 - CẤM GÁN GIẢ CHO NGUỒN THẬT:
- Tuyệt đối không được tạo một số liệu giả rồi mô tả như:
  "nghiên cứu cho thấy", "tại địa điểm X", "ở quần thể Y", "theo bài báo Z"
  nếu SOURCE BLOCK không thực sự chứa dữ liệu đó.
- Không được tự bịa cỡ mẫu, tên loài, địa điểm, thời gian nghiên cứu,
  tỉ lệ phần trăm, nồng độ, khối lượng, số cá thể, kết quả thực nghiệm,
  PMID, PMCID, DOI hoặc URL.
- Nếu cần dùng tình huống mô phỏng, phải viết trung tính theo kiểu:
  "Giả sử...", "Trong một bài toán...", "Xét một quần thể giả định...",
  không được ngụy tạo thành dữ liệu nghiên cứu thật.

CHỦ ĐỀ ƯU TIÊN CÂU ĐỊNH LƯỢNG:
- Đột biến gene/NST; tái bản DNA; phiên mã - dịch mã;
  nguyên phân - giảm phân; Mendel; tương tác gene;
  liên kết gene - hoán vị gene; di truyền giới tính;
  di truyền quần thể; xác suất di truyền; phả hệ;
  sinh thái quần thể, năng lượng sinh thái và các YCCĐ định lượng khác.

THEO MỨC ĐỘ:
- VẬN DỤNG: ưu tiên xử lí dữ liệu thật, bài toán nhiều bước,
  phân tích bảng/biểu đồ/tình huống hoặc suy luận định lượng.
- THÔNG HIỂU: có thể dùng phép tính ngắn hoặc đọc dữ liệu để làm rõ quan hệ sinh học.
- NHẬN BIẾT: không ép tạo bài toán chỉ để có tính toán.

ĐA DẠNG NGÂN HÀNG:
- Không bắt buộc mọi câu trong cùng YCCĐ đều là câu tính toán.
- Cần phối hợp câu kiến thức, phân tích dữ liệu, tính toán và vận dụng tình huống.

==================================================
QUY TẮC BẮT BUỘC
==================================================

1. YCCĐ là ràng buộc cao nhất.
Không tự sửa, thay thế hoặc mở rộng YCCĐ.
Trường "yccd" trong kết quả phải sao chép NGUYÊN VĂN
đúng YCCĐ được cung cấp tương ứng với câu hỏi.

2. Mỗi YCCĐ phải tạo ĐÚNG số lượng câu được yêu cầu.

3. Câu hỏi phải đúng mức độ:
- Nhận biết
- Thông hiểu
- Vận dụng

4. Phải đúng dạng câu hỏi được yêu cầu.

4a. Phải đúng THÀNH PHẦN NĂNG LỰC được yêu cầu.

4b. Nếu là câu có tính toán/định lượng:
- lời giải phải nêu được các bước chính;
- đáp án và phương án nhiễu phải được kiểm tra bằng phép tính;
- phương án nhiễu nên phản ánh lỗi tính toán hoặc lỗi vận dụng quy luật thường gặp,
  không tạo số ngẫu nhiên vô nghĩa;
- không được có hai phương án cùng đúng do làm tròn hoặc do thiếu điều kiện.
Chỉ sử dụng đúng một trong ba tên chính thức:
- Nhận thức sinh học
- Tìm hiểu thế giới sống
- Vận dụng kiến thức, kĩ năng đã học

Không dùng nhãn NL1, NL2, NL3.

5. Không tạo hai câu giống nhau.

6. Câu hỏi phải phù hợp học sinh THPT.

==================================================
QUY TẮC THUẬT NGỮ
==================================================

Bắt buộc sử dụng:

DNA
RNA
gene
protein
amino acid
ribosome
nucleotide
enzyme
allele

Không dùng:

ADN
ARN
gen
prôtêin
axit amin
ribôxôm
nuclêôtit
enzim
alen

Các thuật ngữ tiếng Việt thông dụng như:

nhiễm sắc thể
phiên mã
dịch mã
tái bản DNA
đột biến

được giữ bằng tiếng Việt.

Toàn bộ câu hỏi, phương án, đáp án và giải thích
phải thống nhất hệ thuật ngữ này.

============================================================
QUY TẮC THÀNH PHẦN NĂNG LỰC
============================================================

1. "Nhận thức sinh học":
- Đánh giá khả năng nhận biết, trình bày, phân loại, phân tích,
  so sánh, giải thích mối quan hệ sinh học, phát hiện/chỉnh sửa sai.
- Câu hỏi trọng tâm là hiểu biết và tư duy trên kiến thức sinh học.

2. "Tìm hiểu thế giới sống":
- Đánh giá các hoạt động như đề xuất vấn đề, đặt câu hỏi,
  phán đoán/giả thuyết, lập kế hoạch, chọn phương pháp,
  quan sát, thực nghiệm, điều tra, thu thập/xử lí dữ liệu,
  đối chiếu giả thuyết, rút kết luận, báo cáo và phản biện kết quả.
- Nên dùng tình huống nghiên cứu, thí nghiệm, bảng số liệu,
  thiết kế khảo sát hoặc kết quả quan sát khi phù hợp.

3. "Vận dụng kiến thức, kĩ năng đã học":
- Đánh giá khả năng giải thích/đánh giá hiện tượng thực tiễn,
  phản biện mô hình/công nghệ, đề xuất giải pháp, lựa chọn hành vi,
  bảo vệ sức khoẻ/môi trường, thích ứng và phát triển bền vững.
- Nên có tình huống thực tiễn và yêu cầu học sinh sử dụng kiến thức
  để xử lí, đánh giá hoặc đề xuất giải pháp.

AI phải tạo nội dung thực sự đo đúng thành phần năng lực được gán,
không chỉ ghi nhãn metadata.


============================================================
BỘ ĐIỀU PHỐI PHONG CÁCH ĐỀ TỐT NGHIỆP THPT 2025–2026
============================================================

Mục tiêu không phải tạo các câu rời rạc giống nhau.
Phải xây ngân hàng đa dạng để sau này rút ra được đề có phong cách tương tự
đề tốt nghiệp THPT Sinh học 2025–2026.

Khi YCCĐ phù hợp, chủ động lựa chọn một trong các kiểu khai thác sau:
- kiến thức nền trực tiếp, ngắn gọn, chính xác;
- tình huống sinh học thực tế;
- thí nghiệm hoặc thiết kế thí nghiệm;
- bảng số liệu hoặc dữ liệu định lượng;
- sơ đồ, chuỗi quá trình, phả hệ, cấu trúc gene;
- dữ liệu quần thể hoặc sinh thái;
- bài toán di truyền, xác suất hoặc định lượng;
- phân tích kết quả nghiên cứu;
- vận dụng kiến thức để đánh giá, dự đoán hoặc giải quyết vấn đề.

KHÔNG được ép mọi YCCĐ thành câu dữ liệu.
KHÔNG được ép mọi câu thành câu tính toán.
KHÔNG được ép mọi câu dùng bài báo khoa học.

NGUYÊN TẮC CHỌN NGUỒN:
1. Truy xuất tất cả nguồn khả dụng có liên quan.
2. Chọn nguồn/đoạn phù hợp nhất với YCCĐ và kiểu câu định tạo.
3. Có thể kết hợp nhiều nguồn thật nếu cần.
4. SGK/tài liệu GV, PubMed/PMC/NCBI và nguồn web khoa học đều là các nguồn ngang hàng.
5. Kiến thức nền ổn định có thể dùng kiến thức chương trình phổ thông khi không cần dữ liệu nghiên cứu cụ thể.
6. Câu vận dụng có dữ liệu, nghiên cứu, thí nghiệm hoặc bối cảnh thực tiễn phải ưu tiên dữ liệu thật nếu truy xuất được.
7. Chỉ dùng dữ liệu mô phỏng khi bản chất bài toán cần định lượng mà nguồn không có số liệu phù hợp.

MỤC TIÊU ĐA DẠNG:
- Trong một nhóm nhiều câu của cùng bài/YCCĐ, tránh lặp cấu trúc.
- Nếu có thể, phối hợp câu kiến thức + phân tích dữ liệu + vận dụng + tính toán.
- Đúng/Sai nên dùng một tình huống/dữ kiện chung thực sự có giá trị khai thác.
- Trả lời ngắn ưu tiên định lượng, bảng số liệu, di truyền, sinh thái hoặc xử lí dữ liệu.
- Trắc nghiệm 4 lựa chọn không nhất thiết phải có dữ kiện dài; chọn cấu trúc phù hợp nhất với YCCĐ và mức độ.


============================================================
QUY TẮC CÂU HỎI HOÀN CHỈNH - TỰ ĐỦ DỮ KIỆN
============================================================

Mỗi câu hỏi phải là một đơn vị hoàn chỉnh, học sinh có thể làm chỉ bằng
những gì được nhìn thấy trong chính câu hỏi/tình huống/bảng/sơ đồ đi kèm.

BẮT BUỘC:
- Nếu câu hỏi nhắc đến "bảng dưới đây", "hình bên", "5 dòng", "3 quần thể",
  "kết quả nghiên cứu", "các mẫu", "các cá thể", "các phương án thí nghiệm",
  "số liệu sau", "sơ đồ", "biểu đồ", "phả hệ"... thì TOÀN BỘ dữ kiện cần thiết
  phải xuất hiện trong phần tình huống hoặc nội dung câu hỏi.
- Không được để dữ kiện quan trọng chỉ xuất hiện trong phần đáp án hoặc giải thích.
- Không được hỏi "bao nhiêu", "nhận định nào đúng/sai", "xác suất là bao nhiêu"
  nếu đề bài chưa cung cấp đủ dữ liệu để suy ra duy nhất đáp án.
- Không được dùng các đối tượng như Dòng 1, Dòng 2, Mẫu A, Nhóm I...
  nếu chưa mô tả rõ chúng trước khi hỏi.
- Nếu cần bảng số liệu mà giao diện hiện chưa hỗ trợ bảng riêng,
  hãy trình bày dữ liệu ngay trong "tinh_huong" hoặc "cau_hoi" theo dạng dễ đọc,
  mỗi dòng một đối tượng.
- Câu 4 lựa chọn phải có đủ A, B, C, D.
- Câu Đúng/Sai phải có tình huống chung đầy đủ và đủ 4 ý a-d.
- Câu Trả lời ngắn phải có đủ dữ kiện để tính/suy luận ra một kết quả duy nhất.

TỰ KIỂM TRA TRƯỚC KHI TRẢ JSON:
1. Đọc lại câu hỏi như một học sinh chưa biết lời giải.
2. Xóa phần đáp án/giải thích khỏi trí nhớ tạm thời.
3. Kiểm tra xem vẫn đủ dữ kiện để giải hay không.
4. Nếu chưa đủ, phải bổ sung dữ kiện vào tình huống/câu hỏi trước khi xuất.
5. Chỉ xuất câu khi hoàn chỉnh.

============================================================
QUY TẮC CÂU HỎI
============================================================

Mọi câu hỏi phải:
- Bám sát YCCĐ được chọn.
- Đúng mức độ giáo viên đã chọn.
- Nội dung chính xác về mặt khoa học.
- Có dữ kiện rõ ràng, đủ để xác định đáp án.
- Không tạo câu hỏi mơ hồ hoặc có nhiều cách hiểu.
- Không tiết lộ đáp án trong câu dẫn.
- Không tạo dữ kiện thừa vô nghĩa.
- Ưu tiên câu hỏi có ngữ cảnh sinh học thực tế, bảng số liệu,
  thí nghiệm, hình, sơ đồ hoặc tình huống khi phù hợp.
- Không tự ý nâng hoặc hạ mức độ so với cấu hình giáo viên chọn.

============================================================
I. TRẮC NGHIỆM 4 LỰA CHỌN
============================================================

Nếu dạng là "Trắc nghiệm 4 lựa chọn":

- BẮT BUỘC có đúng 4 phương án A, B, C, D.
- Chỉ có DUY NHẤT 1 phương án đúng.
- Ba phương án nhiễu phải hợp lí về mặt sinh học.
- Các phương án nhiễu nên phản ánh những lỗi học sinh
  thường mắc hoặc những cách suy luận chưa đầy đủ.
- Không tạo phương án vô lí chỉ để đủ 4 đáp án.
- Các phương án phải tương đối đồng đều về độ dài và cách diễn đạt.
- Không để đáp án đúng nổi bật do dài hơn, chi tiết hơn
  hoặc khác cấu trúc so với các phương án còn lại.
- Không sử dụng kiểu:
  "Tất cả các phương án trên".
- Không sử dụng kiểu:
  "Cả A và B đều đúng".
- Không sử dụng kiểu:
  "A và C đúng".
- Không để lộ đáp án trong câu dẫn.

CẤU TRÚC CÂU HỎI:

Câu hỏi có thể sử dụng:
- Một tình huống sinh học.
- Một thí nghiệm.
- Một bảng số liệu.
- Một sơ đồ.
- Một hình minh họa.
- Một đoạn thông tin khoa học ngắn.
- Một hiện tượng thực tiễn.
- Hoặc kiến thức sinh học trực tiếp nếu YCCĐ phù hợp.

Với mức "Nhận biết":
- Kiểm tra khả năng nhận ra, xác định, gọi tên,
  nêu đặc điểm hoặc kiến thức cơ bản.
- Không cố tình làm câu hỏi phức tạp.

Với mức "Thông hiểu":
- Yêu cầu học sinh giải thích, phân biệt, so sánh,
  xác định mối quan hệ hoặc vận dụng trực tiếp kiến thức
  vào một tình huống quen thuộc.

Với mức "Vận dụng":
- Ưu tiên dữ liệu, thí nghiệm, bảng, sơ đồ,
  tình huống thực tiễn hoặc bài toán sinh học.
- Học sinh phải xử lí thông tin hoặc thực hiện
  ít nhất một bước suy luận để chọn đáp án.

Đáp án:
- dap_an phải là một trong:
  "A", "B", "C", "D".

============================================================
II. ĐÚNG / SAI
============================================================

Nếu dạng là "Đúng / Sai":

BẮT BUỘC mỗi câu phải có MỘT TÌNH HUỐNG / DỮ KIỆN CHUNG
thực sự có nội dung để học sinh khai thác.

TUYỆT ĐỐI KHÔNG được tạo câu Đúng / Sai chỉ gồm
4 nhận định kiến thức rời rạc mà không có dữ kiện chung.

------------------------------------------------------------
1. TÌNH HUỐNG / DỮ KIỆN CHUNG
------------------------------------------------------------

Trường "tinh_huong" BẮT BUỘC:

- Không được để trống.
- Không được chỉ có một câu dẫn chung chung như:
  "Dựa vào kiến thức đã học..."
  "Cho các phát biểu sau..."
  "Xét các nhận định sau..."
- Phải chứa thông tin cụ thể để ít nhất 2 trong 4 ý
  cần khai thác trực tiếp từ tình huống đó.

Tình huống có thể là:

- Một thí nghiệm.
- Một bảng số liệu.
- Một sơ đồ.
- Một chuỗi sự kiện sinh học.
- Một đoạn mô tả hiện tượng.
- Một kết quả nghiên cứu.
- Một trường hợp thực tiễn.
- Một quần thể hoặc quần xã có số liệu.
- Một phép lai.
- Một quá trình sinh học có dữ kiện.
- Một mô hình hoặc mối quan hệ sinh học cụ thể.

Ưu tiên dữ kiện có:
- số liệu;
- biến số;
- quan hệ nhân quả;
- kết quả thí nghiệm;
- thay đổi trước / sau;
- điều kiện xử lí;
- hoặc thông tin cho phép suy luận.

Ví dụ ĐẠT:

"Một quần xã có loài A là động vật ăn thịt,
loài B là con mồi của A và loài C là thức ăn của B.
Khi số lượng A giảm 60%, số lượng B tăng gấp 2 lần,
độ phủ của C giảm từ 70% xuống 25%."

Ví dụ KHÔNG ĐẠT:

"Các mối quan hệ sinh thái rất đa dạng.
Hãy xác định các nhận định sau đúng hay sai."

------------------------------------------------------------
2. CÂU HỎI / CÂU LỆNH SAU DỮ KIỆN
------------------------------------------------------------

Trường "cau_hoi" BẮT BUỘC phải có nội dung đối với câu Đúng / Sai.

Sau phần tinh_huong, phải có MỘT câu hỏi hoặc câu lệnh rõ ràng
trước khi đưa ra 4 nhận định a, b, c, d.

Câu hỏi/câu lệnh phải gắn trực tiếp với tình huống hoặc dữ kiện,
ví dụ:
- "Dựa vào thông tin trên, hãy xác định mỗi nhận định sau là Đúng hay Sai."
- "Căn cứ vào kết quả thí nghiệm trên, hãy đánh giá tính Đúng/Sai của các nhận định sau."
- "Từ dữ liệu trên, hãy xác định các nhận định sau đúng hay sai."

Không được dùng trường "cau_hoi" để lặp lại toàn bộ tinh_huong.
Không được bỏ qua câu hỏi/câu lệnh rồi chuyển thẳng từ tinh_huong
sang các nhận định.

Cấu trúc bắt buộc của một câu Đúng / Sai:
1. tinh_huong: dữ liệu/tình huống chung.
2. cau_hoi: câu hỏi/câu lệnh yêu cầu đánh giá Đúng/Sai.
3. nhan_dinh_meta: đúng 4 nhận định a, b, c, d.

------------------------------------------------------------
YÊU CẦU VỀ CÁC NHẬN ĐỊNH a, b, c, d:

- Mỗi nhận định phải ngắn gọn, rõ nghĩa, không diễn đạt dài dòng.
- Có thể gồm 1 hoặc nhiều câu nếu cần để thể hiện đầy đủ nội dung.
- Không giới hạn cứng số câu, nhưng không được lặp lại dài dòng dữ kiện đã nêu ở tình huống.
- Độ khó phải nằm ở kiến thức và suy luận, không nằm ở việc viết câu quá dài.
- Độ dài các nhận định có thể khác nhau tự nhiên.
- Có thể sử dụng số liệu, kết quả thí nghiệm, điều kiện giả định hoặc thông tin bổ sung ngắn khi cần thiết.
- Mỗi nhận định phải có một kết luận rõ ràng để học sinh xác định Đúng/Sai.

Bốn nhận định phải đa dạng về cách khai thác dữ kiện:
- nhận diện hoặc xác định;
- giải thích;
- phân tích quan hệ nhân quả;
- dự đoán khi điều kiện thay đổi;
- xử lí hoặc đối chiếu số liệu;
- đánh giá kết luận;
- vận dụng kiến thức vào tình huống mới.

KHÔNG bắt buộc mỗi ý tương ứng cố định với một dạng trên.
AI phải lựa chọn cách hỏi phù hợp nhất với dữ kiện của từng câu.

Không được tạo 3–4 nhận định cùng một kiểu hỏi.
Tối đa 1 nhận định chỉ đơn thuần yêu cầu gọi tên/phân loại khái niệm.

Ít nhất 2 nhận định phải yêu cầu học sinh thực sự khai thác dữ kiện của tình huống.
Ít nhất 1 nhận định phải có yếu tố suy luận, dự đoán, phân tích hoặc đánh giá.

Không tạo nhận định sai bằng thủ thuật ngôn ngữ đơn giản
như chỉ thêm "luôn", "mọi", "hoàn toàn", "duy nhất".

Nhận định sai phải sai về kiến thức, dữ kiện hoặc logic suy luận.
------------------------------------------------------------
3. PHÂN BỐ MỨC ĐỘ
------------------------------------------------------------

Ưu tiên:

a) Nhận biết hoặc khai thác dữ kiện trực tiếp.
b) Thông hiểu.
c) Thông hiểu hoặc Vận dụng.
d) Vận dụng.

Nếu giáo viên đã cấu hình mức độ riêng từng ý,
phải tuân theo mức độ đó.

Ý vận dụng nên yêu cầu:
- tính toán;
- xử lí số liệu;
- dự đoán;
- xác định hệ quả;
- suy luận nhân quả;
- hoặc kết hợp nhiều dữ kiện.

------------------------------------------------------------
4. ĐÁP ÁN VÀ GIẢI THÍCH
------------------------------------------------------------

Mỗi ý phải có:

- noi_dung
- yccd
- muc_do
- thanh_phan_nang_luc
- dap_an
- giai_thich

BẮT BUỘC:
- nhan_dinh_meta phải có đúng 4 phần tử.
- Cả 4 trường noi_dung của a, b, c, d đều phải có nội dung thực sự.
- Không được chỉ tạo tình huống/câu dẫn rồi bỏ trống các nhận định.
- 4 ý a, b, c, d phải được hiểu là PHẦN HIỂN THỊ CHÍNH của câu Đúng/Sai,
  không phải metadata ẩn.

Với câu Đúng / Sai:
- Ưu tiên "Tìm hiểu thế giới sống" và
  "Vận dụng kiến thức, kĩ năng đã học".
- 4 ý a, b, c, d có thể thuộc YCCĐ khác nhau,
  mức độ khác nhau và thành phần năng lực khác nhau.
- Không bắt buộc 4 ý dùng chung một mức độ hay một YCCĐ.

dap_an của từng ý chỉ được là:

"Đúng"
hoặc
"Sai"

giai_thich phải giải thích RIÊNG từng ý
và phải dựa trên dữ kiện của tinh_huong
khi ý đó sử dụng dữ kiện.

------------------------------------------------------------
5. KIỂM TRA BẮT BUỘC TRƯỚC KHI XUẤT
------------------------------------------------------------

Trước khi trả JSON cho một câu Đúng / Sai,
AI PHẢI tự kiểm tra:

1. tinh_huong có khác rỗng không?

2. tinh_huong có dữ kiện cụ thể không?

3. cau_hoi có một câu hỏi/câu lệnh rõ ràng, gắn với tinh_huong và yêu cầu đánh giá Đúng/Sai không?

4. Có đúng 4 nhận định a, b, c, d không?

5. Ít nhất 2 nhận định có thực sự khai thác tinh_huong không?

6. Có ít nhất 1 nhận định cần suy luận không?

7. 4 nhận định có khác nhau về nhiệm vụ không?

8. Có nhận định nào có thể trả lời hoàn toàn
   mà không cần đọc tinh_huong không?
   Nếu có quá nhiều, phải tạo lại.

9. Đáp án Đúng / Sai và giải thích
   có phù hợp dữ kiện không?

Nếu MỘT trong các điều kiện trên không đạt:

KHÔNG ĐƯỢC xuất câu đó.

Phải tự tạo lại câu hỏi cho đến khi đạt.
============================================================
III. TRẢ LỜI NGẮN
============================================================

Nếu dạng là "Trả lời ngắn":

Đây là dạng câu hỏi học sinh trả lời bằng MỘT GIÁ TRỊ NGẮN có thể mã hoá bằng số,
phù hợp cấu trúc trả lời ngắn của đề tốt nghiệp.

Dạng này KHÔNG đồng nghĩa với "bắt buộc phải tính toán".
AI phải chọn cách hỏi phù hợp nhất với YCCĐ, mức độ và dữ kiện.

Có thể tạo:
- bài toán tính toán;
- xác định một giá trị từ bảng/biểu đồ/sơ đồ;
- đếm số nhận định đúng;
- xác định số bước/số đối tượng/số trường hợp;
- sắp xếp trình tự rồi trả lời bằng mã số/thứ tự theo yêu cầu;
- hoặc dạng xử lí dữ kiện khác cho ra một đáp án ngắn xác định.

Không được ép mọi câu Trả lời ngắn thành cùng một khuôn.

KHÔNG tạo đáp án bằng:
- chữ;
- thuật ngữ;
- tên sinh vật;
- tên quá trình;
- A, B, C, D;
- "Đúng" hoặc "Sai";
- câu văn.

------------------------------------------------------------
QUY TẮC 4 Ô
------------------------------------------------------------

Đáp án của học sinh phải nhập được trong TỐI ĐA 4 Ô.

Mỗi kí tự sau chiếm 1 ô:

- Mỗi chữ số từ 0 đến 9: 1 ô.
- Dấu âm "-": 1 ô.
- Dấu phẩy ",": 1 ô.

Tổng số kí tự của đáp án,
TÍNH CẢ DẤU "-" VÀ DẤU ",",
KHÔNG ĐƯỢC VƯỢT QUÁ 4.

Ví dụ HỢP LỆ:

3       → 1 ô
25      → 2 ô
125     → 3 ô
2025    → 4 ô

2,5     → 3 ô
0,25    → 4 ô

-5      → 2 ô
-25     → 3 ô
-125    → 4 ô
-2,5    → 4 ô

Ví dụ KHÔNG HỢP LỆ:

12345   → 5 ô
0,125   → 5 ô
12,50   → 5 ô
-1250   → 5 ô
-12,5   → 5 ô

------------------------------------------------------------
QUY TẮC BIỂU DIỄN ĐÁP ÁN
------------------------------------------------------------

- Dùng dấu phẩy "," làm dấu thập phân.
- KHÔNG dùng dấu chấm "." làm dấu thập phân.
- Không dùng dấu phân cách hàng nghìn.
- Không ghi đơn vị trong dap_an.
- Không ghi kí hiệu "%" trong dap_an.
- Không ghi chữ trong dap_an.
- Không ghi khoảng trắng không cần thiết.

Ví dụ:

ĐÚNG:

dap_an = "5"
dap_an = "25"
dap_an = "125"
dap_an = "2,5"
dap_an = "0,25"
dap_an = "-5"
dap_an = "-2,5"

SAI:

dap_an = "25%"
dap_an = "25 %"
dap_an = "25 cá thể"
dap_an = "0.25"
dap_an = "120 kJ"
dap_an = "A"
dap_an = "Đúng"

------------------------------------------------------------
QUY TẮC TẠO SỐ LIỆU
------------------------------------------------------------

AI phải thiết kế số liệu của câu hỏi NGAY TỪ ĐẦU
để đáp án cuối cùng nhập được trong tối đa 4 ô.

KHÔNG được tạo một bài toán có kết quả quá dài
rồi cắt bớt đáp án một cách tùy tiện.

Nếu kết quả dự kiến vượt quá 4 ô,
AI phải tự thay đổi:
- số liệu;
- đơn vị;
- cách hỏi;
- hoặc yêu cầu làm tròn;

để đáp án cuối cùng phù hợp quy tắc 4 ô.

Nếu cần làm tròn:
- Câu dẫn phải nói rõ làm tròn đến hàng nào.
- Kết quả sau khi làm tròn vẫn phải tối đa 4 ô.

------------------------------------------------------------
TRƯỜNG HỢP PHẦN TRĂM
------------------------------------------------------------

Nếu đáp án mang đơn vị %:

- Kí hiệu % phải nằm trong CÂU HỎI.
- Không đưa kí hiệu % vào dap_an.

Ví dụ:

"Tỉ lệ cá thể có kiểu hình trội là bao nhiêu %?
Chỉ ghi phần số."

Nếu kết quả là 75% thì:

dap_an = "75"

KHÔNG được:

dap_an = "75%"

------------------------------------------------------------
CÁC DẠNG TRẢ LỜI NGẮN NÊN ƯU TIÊN
------------------------------------------------------------

AI tự chọn dạng Trả lời ngắn phù hợp nhất với nội dung:

- Nếu YCCĐ có bản chất định lượng → có thể tạo bài toán tính toán.
- Nếu YCCĐ gắn bảng/biểu đồ/sơ đồ → có thể yêu cầu đọc và xác định một giá trị.
- Nếu YCCĐ phù hợp đánh giá nhiều phát biểu → có thể hỏi số nhận định đúng.
- Nếu YCCĐ là quá trình/cơ chế → có thể cho các bước rồi yêu cầu sắp xếp,
  sau đó trả lời bằng mã/thứ tự số đã quy ước trong câu hỏi.
- Nếu YCCĐ liên quan phân loại/đếm → có thể hỏi số nhóm, số loại, số trường hợp.
- Nếu có dữ liệu thực nghiệm → ưu tiên khai thác trực tiếp dữ liệu đó.

YÊU CẦU ĐA DẠNG:
- Không để phần lớn câu Trả lời ngắn cùng là "có bao nhiêu nhận định đúng".
- Không để phần lớn câu Trả lời ngắn cùng là bài tính.
- Trong một nhóm nhiều câu, phải thay đổi tự nhiên cách khai thác tùy nội dung.

------------------------------------------------------------
YÊU CẦU CHẤT LƯỢNG
------------------------------------------------------------

Mỗi câu Trả lời ngắn phải:

- Có một đáp án số xác định.
- Có đủ dữ kiện để tính.
- Bám đúng YCCĐ.
- Bám đúng mức độ được chọn.
- Có lời giải hoặc giải thích.
- Không hỏi câu chỉ cần trả lời bằng một thuật ngữ.
- Không biến câu Trả lời ngắn thành câu hỏi tự luận.

Nếu mức độ là "Vận dụng":
- Ưu tiên nhiệm vụ cần xử lí/suy luận thực sự:
  bài toán, bảng số liệu, thí nghiệm, sơ đồ, biểu đồ,
  sắp xếp trình tự, tình huống thực tiễn hoặc dữ kiện tổng hợp.
- Chọn dạng nào phù hợp YCCĐ nhất, không ép thành bài tính.

------------------------------------------------------------
KIỂM TRA BẮT BUỘC TRƯỚC KHI TRẢ KẾT QUẢ
------------------------------------------------------------

Trước khi xuất JSON cho một câu Trả lời ngắn,
AI PHẢI tự kiểm tra:

1. dap_an có phải là GIÁ TRỊ SỐ không?

2. Nếu là số thập phân,
   dap_an đã dùng dấu "," chưa?

3. Tổng số kí tự của dap_an,
   TÍNH CẢ "-" VÀ ",",
   có <= 4 không?

4. dap_an có chứa chữ không?

5. dap_an có chứa đơn vị không?

6. dap_an có chứa "%" không?

7. Dữ kiện trong câu hỏi có thực sự tính ra dap_an không?

8. Lời giải có khớp với dap_an không?

Nếu MỘT trong các điều kiện trên không đạt:

KHÔNG ĐƯỢC xuất câu hỏi đó.

Phải tự sửa số liệu hoặc tạo lại câu hỏi
cho đến khi đáp án hợp lệ.


============================================================
NGUỒN
============================================================

Nếu nguồn là "SGK / Chương trình":

- Ưu tiên kiến thức phổ thông phù hợp chương trình.
- Nội dung phải bám YCCĐ.
- Không tự bịa tên sách, trang sách hoặc đoạn trích.
- Không tự bịa URL.

Nếu nguồn là "NCBI", "PubMed" hoặc
"PubMed Central (PMC)":

- Chỉ được sử dụng các SOURCE BLOCK mà hệ thống đã truy xuất thật.
- Không có SOURCE BLOCK thì không tạo câu hỏi.
- Không được tự bịa bài báo, tác giả, PMID, PMCID, DOI hoặc URL.

Nếu nguồn là "Nguồn web khoa học uy tín":

- Không tự bịa website hoặc URL.
- Chỉ ghi URL khi thực sự có URL được cung cấp
  hoặc đã được hệ thống truy xuất.

Nếu nguồn là "Tài liệu giáo viên tải lên":

- Chỉ sử dụng thông tin có trong tài liệu được cung cấp.
- Không được tự gán nội dung ngoài tài liệu
  thành nội dung của tài liệu đó.

============================================================
QUY TẮC CUỐI CÙNG
============================================================

Trước khi trả kết quả, phải tự kiểm tra:

- Đủ số câu giáo viên yêu cầu.
- Đúng dạng câu hỏi.
- Đúng YCCĐ.
- Đúng mức độ.
- Đúng cấu trúc từng dạng.
- Đáp án và lời giải thống nhất.
- Không bịa nguồn.
- Không bịa URL.
- kieu_nguon phải phản ánh đúng cách tạo câu:
  "nguon_that" nếu có SOURCE BLOCK thật được dùng;
  "kien_thuc_chuong_trinh" nếu là kiến thức nền phổ thông ổn định;
  "ket_hop" nếu kết hợp kiến thức chương trình và nguồn thật.
- kieu_du_lieu chỉ dùng:
  "du_lieu_that", "mo_phong", hoặc "khong_dinh_luong".

Đặc biệt:

Nếu là "Trắc nghiệm 4 lựa chọn":
phải có đúng A, B, C, D và chỉ 1 đáp án đúng.

Nếu là "Đúng / Sai":
phải có 1 tình huống chung + đúng 4 ý a, b, c, d.

Nếu là "Trả lời ngắn":
dap_an bắt buộc là GIÁ TRỊ SỐ
và toàn bộ đáp án phải nhập được trong TỐI ĐA 4 Ô,
tính cả dấu "-" và dấu ",".

Nếu câu hỏi không thỏa mãn các quy tắc trên,
phải tự tạo lại trước khi xuất kết quả.

==================================================
NGUỒN
==================================================

Nếu nguồn là SGK / Chương trình:
ưu tiên kiến thức phổ thông phù hợp chương trình.

Nếu nguồn tham chiếu là NCBI, PubMed hoặc PMC:
chỉ sử dụng SOURCE BLOCK đã được hệ thống truy xuất thật.
Nếu không có SOURCE BLOCK thì không tạo câu cho YCCĐ đó.

========================================
YÊU CẦU TẠO CÂU HỎI
========================================
"""

    for item in danh_sach_yccd:

        _, cau_hinh = lay_cau_hinh(item)

        so_cau_item = int(
            cau_hinh.get("Số câu", 1)
        )

        muc_do_item = cau_hinh.get(
            "Mức độ",
            xac_dinh_muc_do(item["YCCĐ"])
        )

        dang_cau_item = cau_hinh.get(
            "Dạng câu hỏi",
            "Trắc nghiệm 4 lựa chọn"
        )

        nang_luc_item = cau_hinh.get(
            "Thành phần năng lực",
            goi_y_thanh_phan_nang_luc(
                item["YCCĐ"],
                muc_do_item,
                dang_cau_item
            )
        )

        nguon_item = cau_hinh.get(
            "Nguồn tham chiếu",
            ["SGK / Chương trình"]
        )

        tai_lieu_item = cau_hinh.get(
            "Tài liệu giáo viên",
            []
        )

        chi_tiet_4_y = cau_hinh.get(
            "Chi tiết 4 ý",
            {}
        )

        # ================================================
        # TRUY XUẤT NGUỒN THẬT CHO RIÊNG YCCĐ NÀY
        # ================================================
        nguon_that_item = thu_thap_nguon_that_cho_item(
            item,
            cau_hinh
        )

        yccd_key_nguon = " ".join(
            str(
                item["YCCĐ"]
            ).split()
        ).casefold()

        if (
            "nguon_thuc_te_theo_yccd"
            not in st.session_state
        ):
            st.session_state[
                "nguon_thuc_te_theo_yccd"
            ] = {}

        st.session_state[
            "nguon_thuc_te_theo_yccd"
        ][yccd_key_nguon] = nguon_that_item

        nguon_that_text = (
            dinh_dang_nguon_that_cho_prompt(
                nguon_that_item
            )
        )

        chi_tiet_4_y_text = ""

        if (
            dang_cau_item == "Đúng / Sai"
            and chi_tiet_4_y
        ):

            dong_4_y = []

            for ky_hieu in ["a", "b", "c", "d"]:

                meta = chi_tiet_4_y.get(
                    ky_hieu,
                    {}
                )

                dong_4_y.append(
                    f"Ý {ky_hieu}: "
                    f"YCCĐ = {meta.get('yccd', item['YCCĐ'])}; "
                    f"Mức độ = {meta.get('muc_do', muc_do_item)}; "
                    f"Thành phần năng lực = "
                    f"{meta.get('thanh_phan_nang_luc', nang_luc_item)}"
                )

            chi_tiet_4_y_text = (
                "\nCấu hình 4 ý Đúng / Sai:\n"
                + "\n".join(dong_4_y)
            )

        prompt += f"""
----------------------------------------

Khối:
{item["Khối"]}

Chương:
{item["Chương"]}

Bài:
{item["Bài"]}

YCCĐ:
{item["YCCĐ"]}

Số câu:
{so_cau_item}

Mức độ:
{muc_do_item}

Dạng câu hỏi:
{dang_cau_item}

Thành phần năng lực:
{nang_luc_item}

Nguồn tham chiếu:
Nguồn tham chiếu:
{", ".join(nguon_item) if nguon_item else "Không có"}

Tài liệu giáo viên đã chọn:
{", ".join(tai_lieu_item) if tai_lieu_item else "Không có"}

YÊU CẦU HOÀN CHỈNH:
- Câu được tạo phải tự đủ dữ kiện để học sinh làm ngay.
- Nếu sử dụng bảng/số liệu/sơ đồ/danh sách đối tượng, phải chép phần dữ kiện cần thiết
  vào tình huống hoặc câu hỏi.
- Không được để bất kỳ dữ kiện cần thiết nào chỉ xuất hiện ở phần giải thích.

==================================================
DỮ LIỆU NGUỒN THỰC TẾ CHO YCCĐ NÀY
==================================================
{nguon_that_text}

{chi_tiet_4_y_text}

"""

    prompt += """

Hãy trả về đúng dữ liệu theo schema đã yêu cầu.

Mỗi câu phải có:

- khoi
- chuong
- bai
- yccd
- muc_do
- dang_cau
- thanh_phan_nang_luc
- cau_hoi
- lua_chon
- dap_an
- giai_thich
- nguon
"""

    # ======================================================
    # BỔ SUNG QUY TẮC MỨC ĐỘ NHẬN THỨC CHO AI
    # ======================================================

    prompt += """

============================================================
QUY TẮC NHẬN DIỆN VÀ RA CÂU THEO MỨC ĐỘ NHẬN THỨC
============================================================

AI phải tuyệt đối tuân thủ MỨC ĐỘ ĐƯỢC YÊU CẦU
trong mục "Mức độ:" của từng câu hỏi.

Mức độ này là mức giáo viên đã thiết lập cho từng YCCĐ ở Kho YCCĐ.
Phải sử dụng đúng mức độ đó cho câu hỏi tương ứng.

Không được tự đổi mức độ dựa vào động từ xuất hiện trong YCCĐ.
YCCĐ chỉ xác định phạm vi nội dung và năng lực cần hướng tới.

1. NHẬN BIẾT

Các động từ thường gặp:
- nhận biết
- kể tên
- phát biểu
- nêu
- trình bày
- xác định
- mô tả

Yêu cầu khi ra câu:
- kiểm tra khả năng nhớ, nhận diện hoặc trình bày kiến thức;
- không biến câu nhận biết thành câu suy luận phức tạp;
- không yêu cầu xử lí tình huống vượt quá YCCĐ.

2. THÔNG HIỂU

Các động từ thường gặp:
- phân loại
- phân biệt
- phân tích
- so sánh
- lựa chọn
- giải thích
- kết nối thông tin
- nhận ra điểm sai
- chỉnh sửa
- thảo luận
- đưa ra nhận định

Yêu cầu khi ra câu:
- học sinh phải hiểu bản chất kiến thức;
- có thể yêu cầu giải thích, phân tích, so sánh hoặc xác định mối quan hệ;
- không chỉ hỏi nhớ máy móc.

3. VẬN DỤNG

Các động từ thường gặp:
- vận dụng
- giải thích vấn đề thực tiễn
- giải thích mô hình công nghệ
- đưa ra bằng chứng
- phản biện
- đánh giá
- đề xuất giải pháp
- đưa ra giải pháp
- thực hiện giải pháp
- xử lí tình huống
- giải quyết vấn đề

Yêu cầu khi ra câu:
- phải có dữ kiện, tình huống, số liệu, bảng, sơ đồ hoặc bối cảnh phù hợp khi cần;
- học sinh phải sử dụng kiến thức đã học để xử lí vấn đề;
- không được chỉ đổi cách diễn đạt câu nhận biết thành câu vận dụng.

QUY TẮC QUAN TRỌNG:

- Trường "muc_do" trong kết quả phải đúng với mức độ được yêu cầu.
- Nội dung câu hỏi phải thực sự tương ứng với mức độ đó.
- Không được tự nâng hoặc hạ mức độ.
- - Khi hệ thống đang xác định mức độ của YCCĐ,
  phải xét toàn bộ cụm từ và ngữ cảnh của YCCĐ,
  không chỉ dựa vào một động từ riêng lẻ.

- Tuy nhiên, khi mức độ cuối cùng đã được xác định
  hoặc giáo viên đã chủ động chọn mức độ,
  phải tuyệt đối tuân thủ mức độ đó khi tạo câu hỏi.

- Không được tự thay đổi mức độ câu hỏi chỉ vì trong YCCĐ
  có một động từ thường gắn với mức độ khác.
"""
    if noi_dung_tai_lieu_gv.strip():

        prompt += """

============================================================
TÀI LIỆU GIÁO VIÊN CUNG CẤP
============================================================

Ưu tiên sử dụng nội dung tài liệu giáo viên dưới đây khi tạo câu hỏi.

Không sao chép nguyên văn câu hỏi có sẵn nếu không cần thiết.
Có thể khai thác dữ kiện, cách hỏi, ngữ cảnh và kiến thức trong tài liệu
nhưng phải tuân thủ đúng YCCĐ, mức độ và dạng câu đã được cấu hình.

"""

        prompt += noi_dung_tai_lieu_gv
    return prompt


# ==========================================================
# SCHEMA KẾT QUẢ GEMINI
# ==========================================================
QUESTION_SCHEMA = {
    "type": "object",
    "properties": {
        "questions": {
            "type": "array",
            "items": {
                "type": "object",

                "properties": {

                    "khoi": {
                        "type": "string"
                    },

                    "chuong": {
                        "type": "string"
                    },

                    "bai": {
                        "type": "string"
                    },

                    "yccd": {
                        "type": "string"
                    },

                    "muc_do": {
                        "type": "string"
                    },

                    "dang_cau": {
                        "type": "string"
                    },

                    "thanh_phan_nang_luc": {
                        "type": "string"
                    },

                    "cau_hoi": {
                        "type": "string"
                    },

                    "lua_chon": {
                        "type": "array",
                        "items": {
                            "type": "string"
                        }
                    },

                    "dap_an": {
                        "type": "string"
                    },

                    "giai_thich": {
                        "type": "string"
                    },

                    "nguon": {
                        "type": "string"
                    },

                    "tinh_huong": {
                        "type": "string"
                    },

                    "nguon_url": {
                        "type": "string"
                    },

                    "nhan_dinh_meta": {
    "type": "array",
    "minItems": 4,
    "maxItems": 4,

    "items": {
                            "type": "object",

                            "properties": {

                                "noi_dung": {
                                    "type": "string"
                                },

                                "yccd": {
                                    "type": "string"
                                },

                                "muc_do": {
                                    "type": "string"
                                },

                                "thanh_phan_nang_luc": {
                                    "type": "string"
                                },

                                "dap_an": {
                                    "type": "string"
                                },

                                "giai_thich": {
                                    "type": "string"
                                }
                            },

                            "required": [
                                "noi_dung",
                                "yccd",
                                "muc_do",
                                "thanh_phan_nang_luc",
                                "dap_an",
                                "giai_thich"
                            ]
                        }
                    }
                },

                "required": [
                    "khoi",
                    "chuong",
                    "bai",
                    "yccd",
                    "muc_do",
                    "dang_cau",
                    "thanh_phan_nang_luc",
                    "cau_hoi",
                    "lua_chon",
                    "dap_an",
                    "giai_thich",
                    "nguon",
                    "tinh_huong",
                    "nguon_url",
                    "nhan_dinh_meta"
                ]
            }
        }
    },

    "required": [
        "questions"
    ]
}


def bo_nhan_phuong_an(text):
    s = str(text or "").strip()
    for nhan in ["A)", "A.", "A:", "B)", "B.", "B:", "C)", "C.", "C:", "D)", "D.", "D:"]:
        if s.upper().startswith(nhan.upper()):
            return s[len(nhan):].strip()
    return s



def xao_phuong_an_4_lua_chon(question, dap_an_muc_tieu=None):
    """
    Chỉ đổi THỨ TỰ 4 phương án và cập nhật lại đáp án đúng.
    Không thay đổi nội dung khoa học của câu.

    dap_an_muc_tieu: A/B/C/D nếu muốn cân bằng vị trí đáp án.
    Nếu None thì xáo ngẫu nhiên.
    """
    q = dict(question)

    if q.get("dang_cau") != "Trắc nghiệm 4 lựa chọn":
        return q

    lua_chon = list(q.get("lua_chon", []) or [])

    if len(lua_chon) != 4:
        return q

    dap_an = str(q.get("dap_an", "")).strip().upper()

    if dap_an not in ["A", "B", "C", "D"]:
        return q

    noi_dung = [
        bo_nhan_phuong_an(x)
        for x in lua_chon
    ]

    dung_idx = ["A", "B", "C", "D"].index(dap_an)

    items = [
        {
            "id_goc": i,
            "noi_dung": noi_dung[i]
        }
        for i in range(4)
    ]

    if dap_an_muc_tieu in ["A", "B", "C", "D"]:
        target_idx = ["A", "B", "C", "D"].index(
            dap_an_muc_tieu
        )

        dung_item = items[dung_idx]
        con_lai = [
            x for x in items
            if x["id_goc"] != dung_idx
        ]

        random.shuffle(con_lai)

        moi = []
        j = 0

        for pos in range(4):
            if pos == target_idx:
                moi.append(dung_item)
            else:
                moi.append(con_lai[j])
                j += 1

        items = moi

    else:
        random.shuffle(items)

    nhan = ["A", "B", "C", "D"]

    q["lua_chon"] = [
        f"{nhan[i]}. {item['noi_dung']}"
        for i, item in enumerate(items)
    ]

    vi_tri_dung_moi = next(
        i
        for i, item in enumerate(items)
        if item["id_goc"] == dung_idx
    )

    q["dap_an"] = nhan[vi_tri_dung_moi]
    q["_dap_an_da_xao"] = True

    return q


def can_bang_dap_an_4_lua_chon(ds_cau):
    """
    Cân bằng A/B/C/D gần đều trong một danh sách câu.
    Ví dụ 20 câu -> khoảng 5A, 5B, 5C, 5D.
    """
    ds = list(ds_cau)

    idx_4lc = [
        i
        for i, q in enumerate(ds)
        if q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn"
        and len(q.get("lua_chon", []) or []) == 4
        and str(q.get("dap_an", "")).strip().upper()
        in ["A", "B", "C", "D"]
    ]

    if not idx_4lc:
        return ds

    targets = []
    chu_ky = ["A", "B", "C", "D"]

    for i in range(len(idx_4lc)):
        targets.append(
            chu_ky[i % 4]
        )

    random.shuffle(targets)

    for idx, target in zip(idx_4lc, targets):
        ds[idx] = xao_phuong_an_4_lua_chon(
            ds[idx],
            target
        )

    return ds


def thong_ke_dap_an_4_lua_chon(ds_cau):
    thong_ke = {
        "A": 0,
        "B": 0,
        "C": 0,
        "D": 0
    }

    for q in ds_cau:
        if q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn":
            dap = str(
                q.get("dap_an", "")
            ).strip().upper()

            if dap in thong_ke:
                thong_ke[dap] += 1

    return thong_ke


def cau_dung_sai_thieu_hoan_thien(q):
    if q.get("dang_cau") != "Đúng / Sai":
        return []

    loi = []
    meta = list(
        q.get("nhan_dinh_meta", []) or []
    )

    if len(meta) != 4:
        return ["Câu chưa đủ 4 ý a–d"]

    for i, nd in enumerate(meta):
        ky = "abcd"[i]

        if not str(nd.get("noi_dung", "")).strip():
            loi.append(f"Ý {ky} thiếu nội dung")

        dap = str(nd.get("dap_an", "")).strip().casefold()
        if dap not in ["đúng", "dung", "true", "sai", "false"]:
            loi.append(f"Ý {ky} thiếu đáp án Đúng/Sai")

        if not str(nd.get("giai_thich", "")).strip():
            loi.append(f"Ý {ky} thiếu giải thích")

        for field, label in [
            ("yccd", "YCCĐ"),
            ("muc_do", "mức độ"),
            ("thanh_phan_nang_luc", "thành phần năng lực")
        ]:
            if not str(nd.get(field, "")).strip():
                loi.append(
                    f"Ý {ky} thiếu {label}"
                )

    return loi


def chuan_hoa_dap_an_dung_sai(value):
    s = str(value or "").strip().casefold()

    if s in ["đúng", "dung", "true", "1"]:
        return "Đúng"

    if s in ["sai", "false", "0"]:
        return "Sai"

    return str(value or "").strip()


def hoan_thien_cau_dung_sai_bang_ai(question):
    """
    Cứu câu Đ/S hay nhưng đang thiếu đáp án/giải thích/metadata.
    AI KHÔNG được đổi tình huống, câu lệnh hay nội dung 4 nhận định.
    """
    if question.get("dang_cau") != "Đúng / Sai":
        return None, "Chỉ dùng chức năng này cho câu Đúng/Sai."

    meta_cu = list(
        question.get("nhan_dinh_meta", []) or []
    )

    if len(meta_cu) != 4:
        return None, "Câu phải có đủ 4 ý a, b, c, d trước khi hoàn thiện."

    prompt = f"""
Bạn là chuyên gia kiểm định câu hỏi Sinh học THPT.

NHIỆM VỤ:
Hoàn thiện metadata cho câu Đúng/Sai dưới đây.

QUY TẮC CỨNG:
- KHÔNG thay đổi tình huống.
- KHÔNG thay đổi câu hỏi/câu lệnh.
- KHÔNG thay đổi nội dung 4 nhận định a, b, c, d.
- Giữ nguyên YCCĐ/mức độ/thành phần năng lực nào đã có sẵn và hợp lệ.
- Chỉ bổ sung/chỉnh các trường còn thiếu hoặc rõ ràng sai định dạng.
- Với mỗi ý phải có đầy đủ:
  yccd, muc_do, thanh_phan_nang_luc, dap_an, giai_thich.
- dap_an chỉ được là "Đúng" hoặc "Sai".
- Giải thích phải đủ để GV kiểm tra tính đúng/sai của chính nhận định đó.
- Không bịa nguồn, DOI, PMID, số liệu nghiên cứu.

CÂU GỐC:
{json.dumps(question, ensure_ascii=False)}

Trả đúng JSON theo QUESTION_SCHEMA với đúng 1 câu trong questions.
"""

    try:
        response = goi_gemini_co_retry(
            prompt,
            QUESTION_SCHEMA
        )

        data = json.loads(
            response.text
        )

        ds = data.get(
            "questions",
            []
        )

        if not ds:
            return None, "AI không trả về câu đã hoàn thiện."

        ai_q = dict(ds[0])
        meta_ai = list(
            ai_q.get("nhan_dinh_meta", []) or []
        )

        if len(meta_ai) != 4:
            return None, "AI chưa trả đủ 4 ý a–d."

        q_moi = dict(question)

        # Khóa tuyệt đối phần nội dung câu gốc.
        q_moi["tinh_huong"] = question.get(
            "tinh_huong",
            ""
        )
        q_moi["cau_hoi"] = question.get(
            "cau_hoi",
            ""
        )

        meta_moi = []

        for i in range(4):
            cu = dict(meta_cu[i])
            moi = dict(meta_ai[i])

            nd = dict(cu)
            nd["noi_dung"] = cu.get(
                "noi_dung",
                ""
            )

            # Metadata đã có thì giữ; thiếu mới lấy AI bổ sung.
            for field in [
                "yccd",
                "muc_do",
                "thanh_phan_nang_luc"
            ]:
                if not str(cu.get(field, "")).strip():
                    nd[field] = moi.get(
                        field,
                        ""
                    )

            nd["dap_an"] = chuan_hoa_dap_an_dung_sai(
                moi.get(
                    "dap_an",
                    cu.get("dap_an", "")
                )
            )

            nd["giai_thich"] = str(
                moi.get(
                    "giai_thich",
                    cu.get("giai_thich", "")
                )
            ).strip()

            meta_moi.append(nd)

        q_moi["nhan_dinh_meta"] = meta_moi

        # Với Đ/S không dùng đáp án chung cấp câu.
        q_moi["dap_an"] = ""
        q_moi["giai_thich"] = ""

        lich_su = list(
            question.get("lich_su_sua", []) or []
        )

        ban_cu = dict(question)
        ban_cu.pop("lich_su_sua", None)

        lich_su.append({
            "thoi_gian": datetime.now().strftime(
                "%d/%m/%Y %H:%M"
            ),
            "kieu_sua": "AI hoàn thiện đáp án/metadata Đúng-Sai",
            "cau_truoc_sua": ban_cu
        })

        q_moi["lich_su_sua"] = lich_su
        q_moi["temp_id"] = str(uuid.uuid4())
        q_moi["trang_thai"] = (
            "AI đã hoàn thiện – Chờ kiểm định lại"
        )

        q_moi = chuan_hoa_cau_truc_cau_hoi(
            q_moi
        )

        return q_moi, ""

    except Exception as e:
        if la_loi_429(e):
            return None, (
                "API đang giới hạn 429. Câu gốc vẫn được giữ nguyên; "
                "hãy thử hoàn thiện lại sau."
            )

        return None, str(e)



def chuan_hoa_cau_truc_cau_hoi(question):
    """
    Chuẩn hóa nhãn phương án và kiểm tra lỗi cấu trúc bằng Python,
    không tốn lượt gọi Gemini.
    """
    q = dict(question)

    if not str(q.get("muc_dich_su_dung", "")).strip():
        q["muc_dich_su_dung"] = "on_tap_kiem_tra"
    if "hanh_vi_nang_luc" not in q:
        q["hanh_vi_nang_luc"] = ""
    if "chi_bao" not in q:
        q["chi_bao"] = ""
    if "mo_ta_chi_bao" not in q:
        q["mo_ta_chi_bao"] = ""
    if "du_lieu_truc_quan" not in q:
        q["du_lieu_truc_quan"] = {}

    dang = q.get("dang_cau", "")
    loi = []

    if dang == "Trắc nghiệm 4 lựa chọn":
        lua_chon = list(q.get("lua_chon", []) or [])

        if len(lua_chon) != 4:
            loi.append("Trắc nghiệm 4 lựa chọn phải có đúng 4 phương án A, B, C, D.")
        else:
            noi_dung = [bo_nhan_phuong_an(x) for x in lua_chon]
            q["lua_chon"] = [
                f"A. {noi_dung[0]}",
                f"B. {noi_dung[1]}",
                f"C. {noi_dung[2]}",
                f"D. {noi_dung[3]}"
            ]

        if str(q.get("dap_an", "")).strip().upper() not in ["A", "B", "C", "D"]:
            loi.append("Đáp án trắc nghiệm phải là A, B, C hoặc D.")
        elif (
            len(q.get("lua_chon", []) or []) == 4
            and not q.get("_dap_an_da_xao")
        ):
            # Gemini thường có xu hướng để đáp án đúng ở A.
            # App tự xáo phương án và cập nhật đáp án tương ứng.
            q = xao_phuong_an_4_lua_chon(q)

    elif dang == "Đúng / Sai":
        if not str(q.get("tinh_huong", "")).strip():
            loi.append("Câu Đúng/Sai thiếu tình huống hoặc dữ kiện chung.")

        if not str(q.get("cau_hoi", "")).strip():
            loi.append("Câu Đúng/Sai thiếu câu hỏi/câu lệnh trước 4 nhận định.")

        meta = list(q.get("nhan_dinh_meta", []) or [])
        if len(meta) != 4:
            loi.append("Câu Đúng/Sai phải có đúng 4 nhận định a, b, c, d.")
        else:
            for i_nd, nd in enumerate(meta):
                if not str(nd.get("noi_dung", "")).strip():
                    loi.append(
                        f"Câu Đúng/Sai thiếu nội dung nhận định {'abcd'[i_nd]}."
                    )

                nd["dap_an"] = chuan_hoa_dap_an_dung_sai(
                    nd.get("dap_an", "")
                )

                if nd.get("dap_an") not in ["Đúng", "Sai"]:
                    loi.append(
                        f"Câu Đúng/Sai thiếu đáp án ý {'abcd'[i_nd]}."
                    )

                if not str(nd.get("giai_thich", "")).strip():
                    loi.append(
                        f"Câu Đúng/Sai thiếu giải thích ý {'abcd'[i_nd]}."
                    )

    elif dang == "Trả lời ngắn":
        dap = str(q.get("dap_an", "")).strip()
        hop_le = bool(re.fullmatch(r"-?\d+(,\d+)?", dap)) and len(dap) <= 4
        if not hop_le:
            loi.append(
                "Trả lời ngắn phải có đáp án số, dùng dấu phẩy thập phân và tối đa 4 kí tự."
            )

    q["loi_cau_truc"] = loi
    # Câu mới/câu vừa chỉnh phải có mã chỉ báo chuẩn nếu có thể suy ra.
    try:
        q = gan_chi_bao_chuan_cho_cau(q)
    except NameError:
        pass

    return q


def fingerprint_cau_hoi(q):
    du_lieu = {
        "yccd": q.get("yccd", ""),
        "muc_do": q.get("muc_do", ""),
        "dang_cau": q.get("dang_cau", ""),
        "thanh_phan_nang_luc": q.get("thanh_phan_nang_luc", ""),
        "cau_hoi": q.get("cau_hoi", ""),
        "tinh_huong": q.get("tinh_huong", ""),
        "lua_chon": q.get("lua_chon", []),
        "dap_an": q.get("dap_an", ""),
        "nhan_dinh_meta": q.get("nhan_dinh_meta", [])
    }
    raw = json.dumps(du_lieu, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def la_loi_429(exc):
    msg = str(exc).lower()
    return (
        "429" in msg
        or "resource_exhausted" in msg
        or "quota" in msg
        or "rate limit" in msg
        or "too many requests" in msg
    )


def goi_gemini_co_retry(prompt, schema, so_lan_thu=3):
    delays = [8, 20, 45]

    last_exc = None

    for lan in range(so_lan_thu):
        try:
            return client.models.generate_content(
                model=MODEL_AI,
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=schema
                )
            )
        except Exception as e:
            last_exc = e

            if not la_loi_429(e):
                raise

            if lan >= so_lan_thu - 1:
                break

            time.sleep(delays[min(lan, len(delays) - 1)])

    raise last_exc



# ==========================================================
# GỌI GEMINI
# ==========================================================
def goi_ai_tao_cau_hoi():

    prompt = tao_prompt_ai()

    # Hạt giống trước, AI sau.
    khoi_hat_giong = tao_khoi_hat_giong_cho_yccd_da_chon()
    if khoi_hat_giong:
        prompt += khoi_hat_giong

    prompt += """

============================================================
QUY TẮC ĐA DẠNG HÓA NGÂN HÀNG
============================================================
- Hạt giống là nguồn cảm hứng/chất lượng, không phải toàn bộ ngân hàng.
- Nếu tạo nhiều câu cùng YCCĐ, không để các câu cùng một mô-típ.
- Phối hợp câu trực tiếp, tình huống thực tiễn, bảng/số liệu, thí nghiệm,
  sơ đồ/hình mô tả, suy luận và tính toán khi YCCĐ cho phép.
- Tạo thêm câu mới ngoài hạt giống để mở rộng độ phủ năng lực và kĩ năng.
- Ưu tiên bổ sung các thành phần năng lực/chỉ báo còn ít câu.
- Câu mới vẫn phải bám chặt YCCĐ, mức độ và thành phần năng lực được giao.
- Không bịa dữ kiện khoa học; dữ kiện cần đủ để học sinh tự giải.
"""

    try:

        response = goi_gemini_co_retry(
            prompt,
            QUESTION_SCHEMA
        )

        data = json.loads(response.text)

        questions = data.get(
            "questions",
            []
        )

        tong_yeu_cau = 0

        for item in st.session_state.yccd_da_chon:
            _, cau_hinh = lay_cau_hinh(item)
            tong_yeu_cau += int(
                cau_hinh.get("Số câu", 1)
            )

        if len(questions) != tong_yeu_cau:
            st.warning(
                f"Gemini trả về {len(questions)} câu, "
                f"trong khi cấu hình yêu cầu {tong_yeu_cau} câu."
            )

        # Chuẩn hóa metadata theo đúng cấu hình đã chọn ở Kho YCCĐ
        for question in questions:

            yccd_cau = " ".join(
                str(question.get("yccd", "")).split()
            ).casefold()

            for item in st.session_state.yccd_da_chon:

                yccd_item = " ".join(
                    str(item["YCCĐ"]).split()
                ).casefold()

                if yccd_item == yccd_cau:

                    _, cau_hinh = lay_cau_hinh(item)

                    question["khoi"] = item["Khối"]
                    question["chuong"] = item["Chương"]
                    question["bai"] = item["Bài"]
                    question["yccd"] = item["YCCĐ"]
                    question["muc_do"] = cau_hinh.get(
                        "Mức độ",
                        xac_dinh_muc_do(item["YCCĐ"])
                    )
                    question["dang_cau"] = cau_hinh.get(
                        "Dạng câu hỏi",
                        question.get("dang_cau", "")
                    )
                    question["thanh_phan_nang_luc"] = cau_hinh.get(
                        "Thành phần năng lực",
                        goi_y_thanh_phan_nang_luc(
                            item["YCCĐ"],
                            question["muc_do"],
                            question["dang_cau"]
                        )
                    )

                    if question["dang_cau"] == "Đúng / Sai":
                        chi_tiet = cau_hinh.get("Chi tiết 4 ý", {})
                        meta = list(question.get("nhan_dinh_meta", []) or [])

                        for ky, nd in zip(["a", "b", "c", "d"], meta):
                            cfg_y = chi_tiet.get(ky, {})
                            nd["yccd"] = cfg_y.get(
                                "yccd",
                                nd.get("yccd", item["YCCĐ"])
                            )
                            nd["muc_do"] = cfg_y.get(
                                "muc_do",
                                nd.get("muc_do", question["muc_do"])
                            )
                            nd["thanh_phan_nang_luc"] = cfg_y.get(
                                "thanh_phan_nang_luc",
                                nd.get(
                                    "thanh_phan_nang_luc",
                                    goi_y_thanh_phan_nang_luc(
                                        nd["yccd"],
                                        nd["muc_do"],
                                        "Đúng / Sai"
                                    )
                                )
                            )
                    break

            # ================================================
            # KHÓA NGUỒN: chỉ chấp nhận nguồn đã truy xuất thật
            # ================================================
            yccd_key_nguon = " ".join(
                str(
                    question.get(
                        "yccd",
                        ""
                    )
                ).split()
            ).casefold()

            ds_nguon_that = (
                st.session_state
                .get(
                    "nguon_thuc_te_theo_yccd",
                    {}
                )
                .get(
                    yccd_key_nguon,
                    []
                )
            )

            if ds_nguon_that:

                url_hop_le = {
                    str(
                        r.get(
                            "url",
                            ""
                        )
                    ).strip(): r
                    for r in ds_nguon_that
                    if str(
                        r.get(
                            "url",
                            ""
                        )
                    ).strip()
                }

                url_ai = str(
                    question.get(
                        "nguon_url",
                        ""
                    )
                ).strip()

                if (
                    url_ai
                    and url_ai in url_hop_le
                ):
                    source_used = url_hop_le[
                        url_ai
                    ]
                else:
                    # Nếu AI không chỉ đúng URL đã cấp,
                    # ép về một nguồn thật thay vì để URL bịa.
                    source_used = ds_nguon_that[0]

                question["nguon"] = (
                    str(
                        source_used.get(
                            "provider",
                            ""
                        )
                    )
                    + " | "
                    + str(
                        source_used.get(
                            "title",
                            ""
                        )
                    )
                ).strip(" |")

                question["nguon_url"] = str(
                    source_used.get(
                        "url",
                        ""
                    )
                ).strip()

                question["nguon_thuc_te"] = [
                    {
                        "provider": r.get(
                            "provider",
                            ""
                        ),
                        "id": r.get(
                            "id",
                            ""
                        ),
                        "title": r.get(
                            "title",
                            ""
                        ),
                        "url": r.get(
                            "url",
                            ""
                        ),
                        "doi": r.get(
                            "doi",
                            ""
                        )
                    }
                    for r in ds_nguon_that
                ]

            else:
                # Không có SOURCE BLOCK riêng:
                # vẫn cho phép câu kiến thức chương trình ổn định,
                # nhưng tuyệt đối không được mang metadata nguồn nghiên cứu giả.
                question["kieu_nguon"] = "kien_thuc_chuong_trinh"
                question["nguon"] = (
                    question.get("nguon", "")
                    or "Kiến thức chương trình phổ thông"
                )
                question["nguon_url"] = ""
                question["nguon_thuc_te"] = []

            kieu_du_lieu = str(
                question.get(
                    "kieu_du_lieu",
                    "khong_dinh_luong"
                )
            ).strip()

            if kieu_du_lieu not in [
                "du_lieu_that",
                "mo_phong",
                "khong_dinh_luong"
            ]:
                kieu_du_lieu = "khong_dinh_luong"

            question["kieu_du_lieu"] = kieu_du_lieu

            cau_chuan = chuan_hoa_cau_truc_cau_hoi(question)
            cau_chuan["kieu_du_lieu"] = kieu_du_lieu

            # Giữ lại lỗi nguồn đã thêm trước khi chuẩn hóa.
            if question.get("loi_nguon"):
                cau_chuan["loi_nguon"] = question["loi_nguon"]
                cau_chuan["loi_cau_truc"] = list(
                    set(
                        cau_chuan.get(
                            "loi_cau_truc",
                            []
                        )
                        + [
                            "Không có dữ liệu nguồn thực tế."
                        ]
                    )
                )

            question.clear()
            question.update(cau_chuan)

            loi_du_kien = kiem_tra_cau_hoi_day_du_du_kien(
                question
            )

            if loi_du_kien:
                question["loi_cau_truc"] = list(
                    question.get("loi_cau_truc", [])
                ) + loi_du_kien

            question["temp_id"] = str(
                uuid.uuid4()
            )

        # Cân bằng đáp án 4 lựa chọn ngay sau khi tạo.
        # Không gọi AI, chỉ đổi thứ tự phương án và cập nhật đáp án.
        questions = can_bang_dap_an_4_lua_chon(
            questions
        )

        return questions

    except Exception as e:

        st.error(
            "Lỗi khi Gemini tạo câu hỏi:"
        )

        st.code(str(e))

        return []



# ==========================================================
# SCHEMA KIỂM ĐỊNH CÂU HỎI AI
# ==========================================================
VALIDATION_SCHEMA = {
    "type": "object",
    "properties": {
        "evaluations": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "temp_id": {"type": "string"},
                    "diem_yccd": {"type": "integer"},
                    "diem_muc_do": {"type": "integer"},
                    "diem_nang_luc": {"type": "integer"},
                    "diem_khoa_hoc": {"type": "integer"},
                    "diem_dang_cau": {"type": "integer"},
                    "diem_ro_rang": {"type": "integer"},
                    "diem_dap_an_giai_thich": {"type": "integer"},
                    "diem_khong_trung": {"type": "integer"},
                    "tong_diem": {"type": "integer"},
                    "muc_do_thuc_te": {"type": "string"},
                    "nang_luc_thuc_te": {"type": "string"},
                    "ket_luan": {"type": "string"},
                    "loi_nghiem_trong": {
                        "type": "array",
                        "items": {"type": "string"}
                    },
                    "canh_bao": {
                        "type": "array",
                        "items": {"type": "string"}
                    },
                    "goi_y_sua": {
                        "type": "array",
                        "items": {"type": "string"}
                    },
                    "nhan_xet_ngan": {"type": "string"},
                    "danh_gia_tung_y": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "ky_hieu": {"type": "string"},
                                "diem_yccd": {"type": "integer"},
                                "diem_muc_do": {"type": "integer"},
                                "diem_nang_luc": {"type": "integer"},
                                "muc_do_thuc_te": {"type": "string"},
                                "nang_luc_thuc_te": {"type": "string"},
                                "dap_an_hop_ly": {"type": "boolean"},
                                "loi_nghiem_trong": {
                                    "type": "array",
                                    "items": {"type": "string"}
                                },
                                "canh_bao": {
                                    "type": "array",
                                    "items": {"type": "string"}
                                },
                                "nhan_xet_ngan": {"type": "string"}
                            },
                            "required": [
                                "ky_hieu",
                                "diem_yccd",
                                "diem_muc_do",
                                "diem_nang_luc",
                                "muc_do_thuc_te",
                                "nang_luc_thuc_te",
                                "dap_an_hop_ly",
                                "loi_nghiem_trong",
                                "canh_bao",
                                "nhan_xet_ngan"
                            ]
                        }
                    }
                },
                "required": [
                    "temp_id",
                    "diem_yccd",
                    "diem_muc_do",
                    "diem_nang_luc",
                    "diem_khoa_hoc",
                    "diem_dang_cau",
                    "diem_ro_rang",
                    "diem_dap_an_giai_thich",
                    "diem_khong_trung",
                    "tong_diem",
                    "muc_do_thuc_te",
                    "nang_luc_thuc_te",
                    "ket_luan",
                    "loi_nghiem_trong",
                    "canh_bao",
                    "goi_y_sua",
                    "nhan_xet_ngan",
                    "danh_gia_tung_y"
                ]
            }
        }
    },
    "required": ["evaluations"]
}


def tao_noi_dung_cau_de_so_sanh(q):

    phan = [
        str(q.get("cau_hoi", "")),
        str(q.get("tinh_huong", ""))
    ]

    for x in q.get("lua_chon", []):
        phan.append(str(x))

    for nd in q.get("nhan_dinh_meta", []):
        phan.append(str(nd.get("noi_dung", "")))

    return " ".join(phan).strip()



def chuan_hoa_noi_dung_trung(text):
    s = str(text or "").casefold()
    s = re.sub(r"[\W_]+", " ", s, flags=re.UNICODE)
    return " ".join(s.split())


def tap_yccd_cua_cau(q):
    """
    Câu thường: 1 YCCĐ.
    Đúng/Sai: lấy toàn bộ YCCĐ của 4 ý để chống trùng đúng bản chất.
    """
    ds = set()

    if q.get("dang_cau") == "Đúng / Sai":
        for nd in q.get("nhan_dinh_meta", []) or []:
            y = chuan_hoa_noi_dung_trung(
                nd.get("yccd", "")
            )
            if y:
                ds.add(y)
    else:
        y = chuan_hoa_noi_dung_trung(
            q.get("yccd", "")
        )
        if y:
            ds.add(y)

    return ds


def do_giong_token(a, b):
    ta = set(
        chuan_hoa_noi_dung_trung(a).split()
    )
    tb = set(
        chuan_hoa_noi_dung_trung(b).split()
    )

    if not ta or not tb:
        return 0.0

    return len(ta & tb) / len(ta | tb)


def tinh_do_giong_noi_bo(a, b):
    """
    Không gọi AI/API.
    Kết hợp độ giống chuỗi + độ giống tập từ.
    """
    from difflib import SequenceMatcher

    a2 = chuan_hoa_noi_dung_trung(a)
    b2 = chuan_hoa_noi_dung_trung(b)

    if not a2 or not b2:
        return 0.0

    if a2 == b2:
        return 1.0

    seq = SequenceMatcher(
        None,
        a2,
        b2
    ).ratio()

    jac = do_giong_token(
        a2,
        b2
    )

    return max(
        seq,
        0.60 * seq + 0.40 * jac
    )


def xep_loai_trung(ti_le):
    if ti_le >= 0.92:
        return "🔴 Trùng cao"
    if ti_le >= 0.78:
        return "🟡 Gần giống"
    return "🟢 Không trùng"


def kiem_tra_trung_gan(question, bank=None):

    if bank is None:
        bank = doc_ngan_hang()

    noi_dung_moi = tao_noi_dung_cau_de_so_sanh(
        question
    )

    if not str(noi_dung_moi).strip():
        return {
            "ti_le_cao_nhat": 0.0,
            "cau_gan_nhat": "",
            "id_gan_nhat": "",
            "muc_trung": "🟢 Không trùng"
        }

    yccd_moi = tap_yccd_cua_cau(
        question
    )

    ti_le_cao_nhat = 0.0
    cau_gan_nhat = ""
    id_gan_nhat = ""

    for old in bank:

        yccd_cu = tap_yccd_cua_cau(
            old
        )

        # Nếu cả hai bên đều có metadata YCCĐ thì chỉ so sâu
        # khi có ít nhất 1 YCCĐ giao nhau.
        if (
            yccd_moi
            and yccd_cu
            and not (yccd_moi & yccd_cu)
        ):
            continue

        noi_dung_cu = tao_noi_dung_cau_de_so_sanh(
            old
        )

        if not str(noi_dung_cu).strip():
            continue

        ti_le = tinh_do_giong_noi_bo(
            noi_dung_moi,
            noi_dung_cu
        )

        if ti_le > ti_le_cao_nhat:
            ti_le_cao_nhat = ti_le
            cau_gan_nhat = old.get(
                "cau_hoi",
                ""
            )
            id_gan_nhat = old.get(
                "id",
                ""
            )

    ti_le_cao_nhat = round(
        ti_le_cao_nhat,
        3
    )

    return {
        "ti_le_cao_nhat": ti_le_cao_nhat,
        "cau_gan_nhat": cau_gan_nhat,
        "id_gan_nhat": id_gan_nhat,
        "muc_trung": xep_loai_trung(
            ti_le_cao_nhat
        )
    }


def cau_an_toan_de_duyet_lo(q, bank=None):
    """
    Điều kiện cứng cho duyệt theo lô.
    Không cần AI hoạt động.
    Nếu đã có kết quả AI kiểm định và không đạt thì không chọn tự động.
    """
    loi = kiem_tra_cau_moi_truoc_khi_luu(
        q
    )

    if loi:
        return False, "Lỗi cấu trúc"

    trung = kiem_tra_trung_gan(
        q,
        bank
    )

    if trung.get(
        "ti_le_cao_nhat",
        0
    ) >= 0.92:
        return False, "Trùng cao"

    temp_id = str(
        q.get("temp_id", "")
    )

    kd = st.session_state.ket_qua_kiem_dinh.get(
        temp_id
    )

    # Nếu đã kiểm định AI thì chỉ tự chọn khi ĐẠT.
    # Chưa kiểm định do 429 vẫn cho GV tự quyết bằng checkbox.
    if kd and kd.get("ket_luan") != "Đạt":
        return False, "AI kiểm định chưa đạt"

    return True, "An toàn"



def chuan_hoa_muc_do(value):
    text = " ".join(
        str(value or "").strip().split()
    ).casefold()

    mapping = {
        "nhận biết": "Nhận biết",
        "nhan biet": "Nhận biết",
        "thông hiểu": "Thông hiểu",
        "thong hieu": "Thông hiểu",
        "vận dụng": "Vận dụng",
        "van dung": "Vận dụng",
    }

    return mapping.get(text, str(value or "").strip())


def kiem_dinh_danh_sach_cau_hoi(questions):

    if not questions:
        return {}

    bank = doc_ngan_hang()

    # Không kiểm định lại câu không thay đổi.
    ket_qua_cache = {}
    cau_can_kiem = []

    for q in questions:
        fp = fingerprint_cau_hoi(q)
        cache_item = st.session_state.cache_kiem_dinh.get(fp)

        if cache_item:
            ket_qua_cache[str(q.get("temp_id", ""))] = dict(cache_item)
        else:
            cau_can_kiem.append(q)

    if not cau_can_kiem:
        return ket_qua_cache

    questions = cau_can_kiem

    # Chỉ đưa một lượng câu cũ vừa đủ để Gemini đối chiếu trùng.
    # Tránh prompt quá dài.
    bank_tom_tat = []

    for old in bank[-30:]:

        bank_tom_tat.append({
            "yccd": old.get("yccd", ""),
            "muc_do": old.get("muc_do", ""),
            "dang_cau": old.get("dang_cau", ""),
            "cau_hoi": old.get("cau_hoi", ""),
            "tinh_huong": old.get("tinh_huong", "")
        })

    ds_gui = []

    cau_goc_theo_temp_id = {}

    for q in questions:

        q_gui = dict(q)

        if not q_gui.get("temp_id"):
            q_gui["temp_id"] = str(
                uuid.uuid4()
            )
            q["temp_id"] = q_gui["temp_id"]

        q_gui["kiem_tra_trung_noi_bo"] = (
            kiem_tra_trung_gan(
                q,
                bank
            )
        )

        cau_goc_theo_temp_id[
            str(q_gui["temp_id"])
        ] = q

        ds_gui.append(q_gui)

    prompt = f"""
Bạn là hội đồng kiểm định câu hỏi Sinh học THPT.

NHIỆM VỤ:
Đánh giá TỪNG câu hỏi trong danh sách dưới đây.
Không được tin nhãn "muc_do" do hệ thống gắn sẵn.
Phải đọc NỘI DUNG THỰC TẾ của câu hỏi để xác định mức độ nhận thức thật.

============================================================
THANG ĐIỂM 100
============================================================

1. Bám đúng YCCĐ: 15 điểm.
2. Đúng mức độ nhận thức thực tế: 15 điểm.
3. Đúng thành phần năng lực được gán: 15 điểm.
4. Chính xác khoa học: 20 điểm.
5. Đúng cấu trúc và chất lượng dạng câu: 15 điểm.
6. Diễn đạt rõ ràng, không mơ hồ: 10 điểm.
7. Đáp án và giải thích chính xác, thống nhất: 5 điểm.
8. Không trùng hoặc quá giống câu đã có: 5 điểm.

Tổng tối đa = 100.

============================================================
QUY TẮC XÁC ĐỊNH MỨC ĐỘ THỰC TẾ
============================================================

NHẬN BIẾT:
- Chủ yếu nhớ, nhận ra, gọi tên, nêu, xác định trực tiếp.
- Không cần xử lí tình huống mới hoặc suy luận đáng kể.

THÔNG HIỂU:
- Cần giải thích, phân biệt, so sánh, phân tích mối quan hệ,
  hoặc áp dụng trực tiếp kiến thức vào tình huống quen thuộc.

VẬN DỤNG:
- Phải dùng kiến thức để xử lí dữ kiện/tình huống,
  phân tích, suy luận, dự đoán, tính toán, đánh giá
  hoặc giải quyết một vấn đề cụ thể.
- Chỉ thêm một đoạn bối cảnh dài nhưng đáp án vẫn nhận ra trực tiếp
  KHÔNG được xem là Vận dụng.

============================================================
QUY TẮC XÁC ĐỊNH THÀNH PHẦN NĂNG LỰC THỰC TẾ
============================================================

Chỉ được dùng một trong ba tên chính thức:
- Nhận thức sinh học
- Tìm hiểu thế giới sống
- Vận dụng kiến thức, kĩ năng đã học

"Nhận thức sinh học":
- trọng tâm là nhận biết, trình bày, phân loại, phân tích,
  so sánh, giải thích mối quan hệ sinh học, nhận ra/chỉnh sửa sai.

"Tìm hiểu thế giới sống":
- trọng tâm là câu hỏi nghiên cứu, giả thuyết, thiết kế/đánh giá thí nghiệm,
  điều tra, quan sát, thu thập và xử lí dữ liệu, kết luận, báo cáo nghiên cứu.

"Vận dụng kiến thức, kĩ năng đã học":
- trọng tâm là giải thích/đánh giá hiện tượng thực tiễn,
  phản biện, đề xuất/lựa chọn giải pháp, hành vi, biện pháp bảo vệ
  sức khoẻ, môi trường và phát triển bền vững.

Không được tin nhãn "thanh_phan_nang_luc" có sẵn.
Phải đọc nội dung thực tế của câu để xác định "nang_luc_thuc_te".
Nếu năng lực thực tế khác năng lực được gán,
câu không được xếp loại ĐẠT.


============================================================
QUY TẮC RIÊNG CHO CÂU ĐÚNG / SAI TÍCH HỢP
============================================================

Câu Đúng/Sai KHÔNG được xem là "1 câu = 1 YCCĐ = 1 mức độ = 1 năng lực".

Cấu trúc đúng là:
1 tình huống/dữ liệu chung
→ 4 nhận định a, b, c, d
→ mỗi nhận định có YCCĐ, mức độ và thành phần năng lực RIÊNG.

Vì vậy, nếu dang_cau = "Đúng / Sai":

A. Ở CẤP CÂU:
- Chỉ kiểm tra chất lượng tình huống/dữ liệu chung, cấu trúc 4 ý,
  độ chính xác khoa học tổng thể, độ rõ ràng, tính hợp lí của câu lệnh
  và độ trùng.
- KHÔNG được kết luận câu sai YCCĐ/mức độ/năng lực chỉ vì trường
  yccd, muc_do hoặc thanh_phan_nang_luc ở cấp câu không đại diện
  cho cả 4 nhận định.
- Các trường cấp câu đó chỉ là metadata phục vụ tổ chức/lọc dữ liệu.

B. Ở CẤP TỪNG Ý:
- Phải kiểm định RIÊNG a, b, c, d dựa vào chính metadata trong nhan_dinh_meta.
- Với mỗi ý phải xác định:
  + có bám đúng yccd của ý không;
  + muc_do_thuc_te của ý;
  + nang_luc_thuc_te của ý;
  + đáp án Đúng/Sai và giải thích có hợp lí không.
- Trả kết quả trong danh_gia_tung_y với đúng 4 phần tử theo a, b, c, d.
- diem_yccd, diem_muc_do, diem_nang_luc của mỗi ý đều theo thang 0–15.
- Nếu một ý sai mức độ hoặc sai năng lực, chỉ cảnh báo/sửa Ý ĐÓ;
  không được nói cả câu sai chỉ vì khác metadata cấp câu.

C. TỔNG HỢP:
- Với câu Đúng/Sai, ba tiêu chí YCCĐ + mức độ + năng lực ở cấp câu
  phải phản ánh TRUNG BÌNH của 4 ý.
- muc_do_thuc_te ở cấp câu ghi "Tích hợp theo từng ý".
- nang_luc_thuc_te ở cấp câu ghi "Tích hợp theo từng ý".
- Nếu cả 4 ý đều đạt metadata riêng và không có lỗi nghiêm trọng,
  câu có thể được xếp ĐẠT.
- Nếu có 1 hoặc nhiều ý lệch YCCĐ/mức độ/năng lực,
  toàn câu tối đa "Cần xem lại" cho đến khi các ý đó được sửa.
- Nếu một ý sai kiến thức hoặc đáp án Đúng/Sai sai,
  coi là lỗi nghiêm trọng của câu.

Với câu KHÔNG phải Đúng/Sai:
- danh_gia_tung_y phải trả về [].

============================================================
QUY TẮC KẾT LUẬN
============================================================

- "Đạt": tổng điểm >= 80 VÀ không có lỗi nghiêm trọng.
- "Cần xem lại": tổng điểm 60–79, hoặc có cảnh báo cần GV xem.
- "Không đạt": tổng điểm < 60, hoặc có lỗi nghiêm trọng.

LỖI NGHIÊM TRỌNG gồm:
- sai kiến thức khoa học;
- đáp án sai hoặc có nhiều đáp án đúng ở câu 4 lựa chọn;
- không đúng YCCĐ;
- dạng câu sai cấu trúc bắt buộc;
- câu Đúng/Sai thiếu dữ kiện chung hoặc thiếu 4 nhận định;
- câu Trả lời ngắn không có đáp án số hợp lệ;
- nhãn Vận dụng nhưng nội dung thực tế chỉ là Nhận biết
  mà không có xử lí/suy luận.

============================================================
YÊU CẦU ĐẶC BIỆT
============================================================

- Trường "muc_do_thuc_te" chỉ được là:
  "Nhận biết", "Thông hiểu", hoặc "Vận dụng".
- Nếu muc_do_thuc_te khác muc_do được yêu cầu,
  phải trừ điểm mục mức độ và nêu rõ trong cảnh báo.
- Không cộng điểm chỉ vì câu dài hoặc có bối cảnh.
- Kiểm tra phương án nhiễu của trắc nghiệm 4 lựa chọn.
- Kiểm tra tính hợp lí của từng đáp án/giải thích.
- Dùng thông tin "kiem_tra_trung_noi_bo" và danh sách câu ngân hàng
  để đánh giá độ trùng.
- Nếu tỉ lệ giống >= 0.82, phải cảnh báo nguy cơ trùng cao.
- Chỉ đánh giá, KHÔNG tự sửa nội dung câu hỏi.

============================================================
CÂU HỎI CẦN KIỂM ĐỊNH
============================================================

{json.dumps(ds_gui, ensure_ascii=False)}

============================================================
MỘT SỐ CÂU ĐANG CÓ TRONG NGÂN HÀNG ĐỂ ĐỐI CHIẾU
============================================================

{json.dumps(bank_tom_tat, ensure_ascii=False)}

Hãy trả đúng JSON theo schema.
"""

    try:

        response = goi_gemini_co_retry(
            prompt,
            VALIDATION_SCHEMA
        )

        data = json.loads(response.text)

        evaluations = data.get(
            "evaluations",
            []
        )

        ket_qua = {}

        for item in evaluations:

            temp_id = str(
                item.get("temp_id", "")
            )

            # --------------------------------------------------
            # CHUẨN HÓA MỨC ĐỘ
            # --------------------------------------------------
            cau_goc = cau_goc_theo_temp_id.get(
                temp_id,
                {}
            )

            la_dung_sai = (
                cau_goc.get("dang_cau") == "Đúng / Sai"
            )

            if la_dung_sai:
                # Metadata cấp câu không đại diện cho 4 ý.
                muc_do_yeu_cau = ""
                muc_do_thuc_te = "Tích hợp theo từng ý"
                nang_luc_yeu_cau = ""
                nang_luc_thuc_te = "Tích hợp theo từng ý"
                item["muc_do_thuc_te"] = muc_do_thuc_te
                item["nang_luc_thuc_te"] = nang_luc_thuc_te
                sai_nang_luc = False
            else:
                muc_do_yeu_cau = chuan_hoa_muc_do(
                    cau_goc.get("muc_do", "")
                )

                muc_do_thuc_te = chuan_hoa_muc_do(
                    item.get("muc_do_thuc_te", "")
                )

                item["muc_do_thuc_te"] = muc_do_thuc_te

                nang_luc_yeu_cau = str(
                    cau_goc.get("thanh_phan_nang_luc", "")
                ).strip()

                nang_luc_thuc_te = str(
                    item.get("nang_luc_thuc_te", "")
                ).strip()

                item["nang_luc_thuc_te"] = nang_luc_thuc_te

                sai_nang_luc = (
                    nang_luc_yeu_cau in THANH_PHAN_NANG_LUC
                    and nang_luc_thuc_te in THANH_PHAN_NANG_LUC
                    and nang_luc_yeu_cau != nang_luc_thuc_te
                )

            # --------------------------------------------------
            # CHẶN ĐIỂM NGOÀI GIỚI HẠN
            # --------------------------------------------------
            diem_yccd = max(
                0,
                min(15, int(item.get("diem_yccd", 0)))
            )

            diem_muc_do = max(
                0,
                min(15, int(item.get("diem_muc_do", 0)))
            )

            diem_nang_luc = max(
                0,
                min(15, int(item.get("diem_nang_luc", 0)))
            )

            diem_khoa_hoc = max(
                0,
                min(20, int(item.get("diem_khoa_hoc", 0)))
            )

            diem_dang_cau = max(
                0,
                min(15, int(item.get("diem_dang_cau", 0)))
            )

            diem_ro_rang = max(
                0,
                min(10, int(item.get("diem_ro_rang", 0)))
            )

            diem_dap_an = max(
                0,
                min(
                    5,
                    int(
                        item.get(
                            "diem_dap_an_giai_thich",
                            0
                        )
                    )
                )
            )

            diem_khong_trung = max(
                0,
                min(5, int(item.get("diem_khong_trung", 0)))
            )

            co_y_can_xem_lai = False
            loi_tung_y = []
            canh_bao_tung_y = []

            if la_dung_sai:
                meta_4_y = list(
                    cau_goc.get("nhan_dinh_meta", []) or []
                )
                danh_gia_4_y = list(
                    item.get("danh_gia_tung_y", []) or []
                )

                dg_map = {
                    str(dg.get("ky_hieu", "")).strip().lower(): dg
                    for dg in danh_gia_4_y
                }

                diem_yccd_y = []
                diem_muc_do_y = []
                diem_nang_luc_y = []
                danh_gia_chuan = []

                for i_y, nd in enumerate(meta_4_y[:4]):
                    ky = "abcd"[i_y]
                    dg = dict(dg_map.get(ky, {}))

                    muc_yeu_cau_y = chuan_hoa_muc_do(
                        nd.get("muc_do", "")
                    )
                    muc_thuc_te_y = chuan_hoa_muc_do(
                        dg.get("muc_do_thuc_te", "")
                    )

                    nl_yeu_cau_y = str(
                        nd.get("thanh_phan_nang_luc", "")
                    ).strip()
                    nl_thuc_te_y = str(
                        dg.get("nang_luc_thuc_te", "")
                    ).strip()

                    d_yccd = max(
                        0,
                        min(15, int(dg.get("diem_yccd", 0)))
                    )
                    d_muc = max(
                        0,
                        min(15, int(dg.get("diem_muc_do", 0)))
                    )
                    d_nl = max(
                        0,
                        min(15, int(dg.get("diem_nang_luc", 0)))
                    )

                    sai_muc_y = (
                        muc_yeu_cau_y
                        and muc_thuc_te_y
                        and muc_yeu_cau_y != muc_thuc_te_y
                    )
                    sai_nl_y = (
                        nl_yeu_cau_y in THANH_PHAN_NANG_LUC
                        and nl_thuc_te_y in THANH_PHAN_NANG_LUC
                        and nl_yeu_cau_y != nl_thuc_te_y
                    )

                    if sai_muc_y:
                        d_muc = min(d_muc, 6)
                        co_y_can_xem_lai = True
                        canh_bao_tung_y.append(
                            f"Ý {ky}: yêu cầu {muc_yeu_cau_y}, "
                            f"thực tế {muc_thuc_te_y}."
                        )

                    if sai_nl_y:
                        d_nl = min(d_nl, 5)
                        co_y_can_xem_lai = True
                        canh_bao_tung_y.append(
                            f"Ý {ky}: năng lực gán là {nl_yeu_cau_y}, "
                            f"thực tế đo {nl_thuc_te_y}."
                        )

                    if not bool(dg.get("dap_an_hop_ly", True)):
                        co_y_can_xem_lai = True
                        loi_tung_y.append(
                            f"Ý {ky}: đáp án Đúng/Sai hoặc giải thích không hợp lí."
                        )

                    for loi_y in dg.get("loi_nghiem_trong", []) or []:
                        loi_tung_y.append(
                            f"Ý {ky}: {loi_y}"
                        )

                    for cb_y in dg.get("canh_bao", []) or []:
                        canh_bao_tung_y.append(
                            f"Ý {ky}: {cb_y}"
                        )

                    diem_yccd_y.append(d_yccd)
                    diem_muc_do_y.append(d_muc)
                    diem_nang_luc_y.append(d_nl)

                    dg["ky_hieu"] = ky
                    dg["muc_do_yeu_cau"] = muc_yeu_cau_y
                    dg["muc_do_thuc_te"] = muc_thuc_te_y
                    dg["nang_luc_yeu_cau"] = nl_yeu_cau_y
                    dg["nang_luc_thuc_te"] = nl_thuc_te_y
                    dg["diem_yccd"] = d_yccd
                    dg["diem_muc_do"] = d_muc
                    dg["diem_nang_luc"] = d_nl
                    danh_gia_chuan.append(dg)

                if len(danh_gia_chuan) == 4:
                    diem_yccd = round(
                        sum(diem_yccd_y) / 4
                    )
                    diem_muc_do = round(
                        sum(diem_muc_do_y) / 4
                    )
                    diem_nang_luc = round(
                        sum(diem_nang_luc_y) / 4
                    )
                    item["danh_gia_tung_y"] = danh_gia_chuan
                else:
                    co_y_can_xem_lai = True
                    loi_tung_y.append(
                        "AI kiểm định chưa trả đủ đánh giá cho 4 ý a, b, c, d."
                    )

            # --------------------------------------------------
            # QUY TẮC CỨNG: SAI MỨC ĐỘ KHÔNG ĐƯỢC "ĐẠT"
            # --------------------------------------------------
            thu_bac = {
                "Nhận biết": 1,
                "Thông hiểu": 2,
                "Vận dụng": 3
            }

            sai_muc_do = (
                (not la_dung_sai)
                and muc_do_yeu_cau in thu_bac
                and muc_do_thuc_te in thu_bac
                and muc_do_yeu_cau != muc_do_thuc_te
            )

            do_lech = 0

            if sai_muc_do:
                do_lech = abs(
                    thu_bac[muc_do_yeu_cau]
                    - thu_bac[muc_do_thuc_te]
                )

                # Lệch 1 bậc: tiêu chí mức độ tối đa 8/20.
                # Lệch 2 bậc: tiêu chí mức độ tối đa 2/20.
                if do_lech == 1:
                    diem_muc_do = min(
                        diem_muc_do,
                        6
                    )
                else:
                    diem_muc_do = min(
                        diem_muc_do,
                        2
                    )

            if sai_nang_luc:
                diem_nang_luc = min(
                    diem_nang_luc,
                    5
                )

            # --------------------------------------------------
            # TÍNH LẠI TỔNG ĐIỂM
            # --------------------------------------------------
            tong_diem = (
                diem_yccd
                + diem_muc_do
                + diem_nang_luc
                + diem_khoa_hoc
                + diem_dang_cau
                + diem_ro_rang
                + diem_dap_an
                + diem_khong_trung
            )

            item["diem_yccd"] = diem_yccd
            item["diem_muc_do"] = diem_muc_do
            item["diem_nang_luc"] = diem_nang_luc
            item["diem_khoa_hoc"] = diem_khoa_hoc
            item["diem_dang_cau"] = diem_dang_cau
            item["diem_ro_rang"] = diem_ro_rang
            item["diem_dap_an_giai_thich"] = diem_dap_an
            item["diem_khong_trung"] = diem_khong_trung
            item["tong_diem"] = tong_diem

            loi_nghiem_trong = list(
                item.get(
                    "loi_nghiem_trong",
                    []
                )
                or []
            )

            canh_bao = list(
                item.get(
                    "canh_bao",
                    []
                )
                or []
            )

            if la_dung_sai:
                for loi_y in loi_tung_y:
                    if loi_y not in loi_nghiem_trong:
                        loi_nghiem_trong.append(loi_y)

                for cb_y in canh_bao_tung_y:
                    if cb_y not in canh_bao:
                        canh_bao.append(cb_y)

            # --------------------------------------------------
            # THÊM CẢNH BÁO / LỖI THEO MỨC ĐỘ THỰC TẾ
            # --------------------------------------------------
            if sai_muc_do:

                thong_bao_lech = (
                    f"Mức độ yêu cầu là {muc_do_yeu_cau}, "
                    f"nhưng nội dung thực tế được xác định là "
                    f"{muc_do_thuc_te}."
                )

                if thong_bao_lech not in canh_bao:
                    canh_bao.append(
                        thong_bao_lech
                    )

                # Vận dụng nhưng thực tế chỉ Nhận biết:
                # coi là sai nghiêm trọng và KHÔNG ĐẠT.
                if (
                    muc_do_yeu_cau == "Vận dụng"
                    and muc_do_thuc_te == "Nhận biết"
                ):
                    loi_vd_nb = (
                        "Câu được yêu cầu ở mức Vận dụng "
                        "nhưng nội dung thực tế chỉ ở mức Nhận biết."
                    )

                    if loi_vd_nb not in loi_nghiem_trong:
                        loi_nghiem_trong.append(
                            loi_vd_nb
                        )

            if sai_nang_luc:
                tb_nl = (
                    f"Thành phần năng lực được gán là {nang_luc_yeu_cau}, "
                    f"nhưng nội dung thực tế đo {nang_luc_thuc_te}."
                )
                if tb_nl not in canh_bao:
                    canh_bao.append(tb_nl)

            item["loi_nghiem_trong"] = loi_nghiem_trong
            item["canh_bao"] = canh_bao

            # --------------------------------------------------
            # KẾT LUẬN CUỐI CÙNG
            # --------------------------------------------------
            # Quy tắc ưu tiên:
            # 1) Vận dụng -> Nhận biết: KHÔNG ĐẠT.
            # 2) Bất kỳ sai mức độ nào khác: tối đa CẦN XEM LẠI.
            # 3) Có lỗi nghiêm trọng khác: không được ĐẠT.
            # 4) Chỉ đúng mức độ + không lỗi nghiêm trọng mới xét điểm.
            if (
                (not la_dung_sai)
                and muc_do_yeu_cau == "Vận dụng"
                and muc_do_thuc_te == "Nhận biết"
            ):
                item["ket_luan"] = "Không đạt"

            elif la_dung_sai and co_y_can_xem_lai:
                if loi_nghiem_trong and tong_diem < 60:
                    item["ket_luan"] = "Không đạt"
                else:
                    item["ket_luan"] = "Cần xem lại"

            elif sai_muc_do or sai_nang_luc:
                if tong_diem < 60:
                    item["ket_luan"] = "Không đạt"
                else:
                    item["ket_luan"] = "Cần xem lại"

            elif loi_nghiem_trong:
                if tong_diem < 60:
                    item["ket_luan"] = "Không đạt"
                else:
                    item["ket_luan"] = "Cần xem lại"

            else:
                if tong_diem >= 80:
                    item["ket_luan"] = "Đạt"
                elif tong_diem >= 60:
                    item["ket_luan"] = "Cần xem lại"
                else:
                    item["ket_luan"] = "Không đạt"

            ket_qua[temp_id] = item

            fp = fingerprint_cau_hoi(
                cau_goc
            )
            st.session_state.cache_kiem_dinh[
                fp
            ] = dict(item)

        ket_qua.update(ket_qua_cache)
        return ket_qua

    except Exception as e:

        st.error(
            "Lỗi khi AI kiểm định câu hỏi:"
        )
        st.code(str(e))

        return {}


def hien_thi_ket_qua_kiem_dinh(question):

    temp_id = str(
        question.get("temp_id", "")
    )

    kd = st.session_state.ket_qua_kiem_dinh.get(
        temp_id
    )

    if not kd:
        st.warning(
            "🛡 Chưa kiểm định câu này."
        )
        return

    tong = int(
        kd.get("tong_diem", 0)
    )

    ket_luan = kd.get(
        "ket_luan",
        "Cần xem lại"
    )

    if ket_luan == "Đạt":
        st.success(
            f"🛡 AI kiểm định: **{tong}/100 – ĐẠT**"
        )
    elif ket_luan == "Không đạt":
        st.error(
            f"🛡 AI kiểm định: **{tong}/100 – KHÔNG ĐẠT**"
        )
    else:
        st.warning(
            f"🛡 AI kiểm định: **{tong}/100 – CẦN XEM LẠI**"
        )

    if question.get("dang_cau") == "Đúng / Sai":

        st.info(
            "Câu Đúng/Sai được kiểm định **theo từng ý a–d**; "
            "không dùng một YCCĐ/mức độ/năng lực chung cho cả câu."
        )

        danh_gia_y = kd.get(
            "danh_gia_tung_y",
            []
        ) or []

        if danh_gia_y:
            rows_y = []

            meta_q = list(
                question.get("nhan_dinh_meta", []) or []
            )

            for i_y, dg in enumerate(danh_gia_y[:4]):
                nd = (
                    meta_q[i_y]
                    if i_y < len(meta_q)
                    else {}
                )

                rows_y.append({
                    "Ý": dg.get("ky_hieu", "abcd"[i_y]),
                    "YCCĐ": nd.get("yccd", ""),
                    "Mức độ gán": dg.get(
                        "muc_do_yeu_cau",
                        nd.get("muc_do", "")
                    ),
                    "Mức độ thực tế": dg.get(
                        "muc_do_thuc_te",
                        ""
                    ),
                    "Năng lực gán": dg.get(
                        "nang_luc_yeu_cau",
                        nd.get("thanh_phan_nang_luc", "")
                    ),
                    "Năng lực thực tế": dg.get(
                        "nang_luc_thuc_te",
                        ""
                    ),
                    "Đáp án hợp lí": (
                        "Có"
                        if dg.get("dap_an_hop_ly", True)
                        else "Không"
                    )
                })

            st.dataframe(
                pd.DataFrame(rows_y),
                use_container_width=True,
                hide_index=True
            )

    else:

        muc_yeu_cau = question.get(
            "muc_do",
            ""
        )

        muc_thuc_te = kd.get(
            "muc_do_thuc_te",
            ""
        )

        st.write(
            f"**Mức độ yêu cầu:** {muc_yeu_cau}"
        )
        st.write(
            f"**Mức độ AI xác định từ nội dung:** {muc_thuc_te}"
        )

        nang_luc_yeu_cau = question.get(
            "thanh_phan_nang_luc",
            ""
        )
        nang_luc_thuc_te = kd.get(
            "nang_luc_thuc_te",
            ""
        )

        st.write(
            f"**Thành phần năng lực được gán:** {nang_luc_yeu_cau}"
        )
        st.write(
            f"**Thành phần năng lực AI xác định:** {nang_luc_thuc_te}"
        )

        if (
            nang_luc_yeu_cau
            and nang_luc_thuc_te
            and nang_luc_yeu_cau != nang_luc_thuc_te
        ):
            st.error(
                "⚠️ Câu hỏi không đo đúng thành phần năng lực đã gán. "
                "Câu này không được xếp loại ĐẠT."
            )

        muc_yeu_cau_chuan = chuan_hoa_muc_do(
            muc_yeu_cau
        )
        muc_thuc_te_chuan = chuan_hoa_muc_do(
            muc_thuc_te
        )

        if (
            muc_yeu_cau_chuan
            and muc_thuc_te_chuan
            and muc_yeu_cau_chuan != muc_thuc_te_chuan
        ):
            st.error(
                "⚠️ Mức độ thực tế của câu hỏi "
                "không khớp mức độ cần tạo. "
                "Câu này không được xếp loại ĐẠT."
            )

    with st.expander(
        "📋 Xem chi tiết kiểm định"
    ):

        bang_diem = pd.DataFrame([
            {
                "Tiêu chí": "Bám đúng YCCĐ",
                "Điểm": kd.get("diem_yccd", 0),
                "Tối đa": 15
            },
            {
                "Tiêu chí": "Đúng mức độ nhận thức",
                "Điểm": kd.get("diem_muc_do", 0),
                "Tối đa": 15
            },
            {
                "Tiêu chí": "Đúng thành phần năng lực",
                "Điểm": kd.get("diem_nang_luc", 0),
                "Tối đa": 15
            },
            {
                "Tiêu chí": "Chính xác khoa học",
                "Điểm": kd.get("diem_khoa_hoc", 0),
                "Tối đa": 20
            },
            {
                "Tiêu chí": "Cấu trúc / dạng câu",
                "Điểm": kd.get("diem_dang_cau", 0),
                "Tối đa": 15
            },
            {
                "Tiêu chí": "Rõ ràng",
                "Điểm": kd.get("diem_ro_rang", 0),
                "Tối đa": 10
            },
            {
                "Tiêu chí": "Đáp án / giải thích",
                "Điểm": kd.get(
                    "diem_dap_an_giai_thich",
                    0
                ),
                "Tối đa": 5
            },
            {
                "Tiêu chí": "Không trùng",
                "Điểm": kd.get("diem_khong_trung", 0),
                "Tối đa": 5
            }
        ])

        st.dataframe(
            bang_diem,
            use_container_width=True,
            hide_index=True
        )

        st.write(
            "**Nhận xét:**",
            kd.get("nhan_xet_ngan", "")
        )

        loi = kd.get(
            "loi_nghiem_trong",
            []
        )

        if loi:
            st.markdown(
                "**🚨 Lỗi nghiêm trọng:**"
            )
            for x in loi:
                st.write(f"- {x}")

        canh_bao = kd.get(
            "canh_bao",
            []
        )

        if canh_bao:
            st.markdown(
                "**⚠️ Cảnh báo:**"
            )
            for x in canh_bao:
                st.write(f"- {x}")

        goi_y = kd.get(
            "goi_y_sua",
            []
        )

        if goi_y:
            st.markdown(
                "**✏️ Gợi ý sửa:**"
            )
            for x in goi_y:
                st.write(f"- {x}")




def thay_cau_trong_danh_sach_cho_duyet(temp_id, cau_moi):
    ds = st.session_state.get(
        "auto_cau_cho_duyet",
        []
    )

    for i, q in enumerate(ds):
        if str(q.get("temp_id", "")) == str(temp_id):
            ds[i] = cau_moi
            st.session_state.auto_cau_cho_duyet = ds
            return True

    return False


def tao_prompt_sua_cau_ai(question, ket_qua_kiem_dinh):
    """
    AI chỉ sửa đúng lỗi đã phát hiện, không được tự đổi mục tiêu đánh giá.
    Với Đúng/Sai phải giữ cấu trúc metadata riêng cho từng ý.
    """
    return f"""
Bạn là chuyên gia biên tập câu hỏi Sinh học THPT.

NHIỆM VỤ:
Sửa câu hỏi dưới đây dựa CHÍNH XÁC trên kết quả kiểm định.
Không được tự thay đổi YCCĐ, mức độ hoặc thành phần năng lực mục tiêu.
Không được bịa nguồn, dữ liệu nghiên cứu, URL, PMID hoặc DOI.

============================================================
CÂU GỐC
============================================================
{json.dumps(question, ensure_ascii=False)}

============================================================
KẾT QUẢ KIỂM ĐỊNH
============================================================
{json.dumps(ket_qua_kiem_dinh, ensure_ascii=False)}

============================================================
NGUYÊN TẮC SỬA
============================================================

1. Chỉ sửa những phần cần thiết để câu đạt chuẩn.
2. Giữ nguyên khối/chương/bài và mục tiêu đánh giá đã gán.
3. Nếu lỗi thiếu dữ kiện:
   - bổ sung toàn bộ dữ kiện cần thiết vào tình_huong hoặc cau_hoi.
4. Nếu lỗi mức độ:
   - điều chỉnh nhiệm vụ nhận thức để đúng mức độ ĐÃ GÁN,
     không đổi nhãn mức độ sang nhãn mới.
5. Nếu lỗi thành phần năng lực:
   - thiết kế lại nhiệm vụ để thực sự đo đúng thành phần năng lực ĐÃ GÁN.
6. Nếu lỗi gần trùng:
   - thay bối cảnh/dữ liệu/cách khai thác,
     nhưng không thay YCCĐ và mục tiêu đánh giá.
7. Nếu dạng 4 lựa chọn:
   - phải đủ A, B, C, D;
   - chỉ một đáp án đúng;
   - phương án nhiễu hợp lí.
8. Nếu dạng Trả lời ngắn:
   - chọn cách hỏi phù hợp nội dung: tính toán, xử lí bảng/sơ đồ,
     số ý đúng, sắp xếp, đếm trường hợp...;
   - không ép mọi câu thành một khuôn;
   - đáp án cuối vẫn tuân thủ quy tắc trả lời ngắn của hệ thống.
9. Nếu dạng Đúng/Sai:
   - 1 tình huống/dữ liệu chung;
   - đúng 4 ý a,b,c,d;
   - mỗi ý có yccd, muc_do, thanh_phan_nang_luc riêng;
   - CHỈ sửa ý bị lỗi nếu các ý còn lại đã đạt;
   - không gom 4 ý về một YCCĐ/mức độ chung.
10. Không để dữ kiện quan trọng chỉ xuất hiện ở lời giải.
11. Câu sau sửa phải hoàn chỉnh để học sinh làm ngay.

============================================================
YÊU CẦU ĐẦU RA
============================================================

Trả đúng JSON theo QUESTION_SCHEMA với đúng 1 phần tử trong "questions".
Không thêm giải thích ngoài JSON.
"""


def sua_cau_bang_ai(question):
    """
    Gọi Gemini để sửa 1 câu dựa trên lỗi kiểm định hiện có.
    Nếu đang 429 thì giữ nguyên câu cũ.
    """
    temp_id = str(
        question.get("temp_id", "")
    )

    kd = st.session_state.ket_qua_kiem_dinh.get(
        temp_id
    )

    if not kd:
        return None, (
            "Câu này chưa có kết quả kiểm định. "
            "Hãy kiểm định trước để AI biết cần sửa lỗi nào."
        )

    try:
        prompt = tao_prompt_sua_cau_ai(
            question,
            kd
        )

        response = goi_gemini_co_retry(
            prompt,
            QUESTION_SCHEMA
        )

        data = json.loads(
            response.text
        )

        ds = data.get(
            "questions",
            []
        )

        if not ds:
            return None, "AI không trả về câu đã sửa."

        cau_moi = dict(ds[0])

        # Giữ các metadata hệ thống không được phép mất.
        for field in [
            "khoi",
            "chuong",
            "bai",
            "dang_cau"
        ]:
            if not cau_moi.get(field):
                cau_moi[field] = question.get(
                    field,
                    ""
                )

        # Với câu thường, khóa metadata mục tiêu cấp câu.
        if question.get("dang_cau") != "Đúng / Sai":
            for field in [
                "yccd",
                "muc_do",
                "thanh_phan_nang_luc"
            ]:
                cau_moi[field] = question.get(
                    field,
                    cau_moi.get(field, "")
                )

        # Với Đúng/Sai, khóa metadata từng ý theo mục tiêu cũ.
        else:
            meta_cu = list(
                question.get(
                    "nhan_dinh_meta",
                    []
                )
                or []
            )
            meta_moi = list(
                cau_moi.get(
                    "nhan_dinh_meta",
                    []
                )
                or []
            )

            if len(meta_moi) == 4 and len(meta_cu) == 4:
                for i in range(4):
                    for field in [
                        "yccd",
                        "muc_do",
                        "thanh_phan_nang_luc"
                    ]:
                        meta_moi[i][field] = meta_cu[i].get(
                            field,
                            meta_moi[i].get(field, "")
                        )

                cau_moi["nhan_dinh_meta"] = meta_moi

        # Lưu lịch sử phiên bản trước.
        lich_su = list(
            question.get(
                "lich_su_sua",
                []
            )
            or []
        )

        ban_cu = dict(question)
        ban_cu.pop(
            "lich_su_sua",
            None
        )

        lich_su.append({
            "thoi_gian": datetime.now().strftime(
                "%d/%m/%Y %H:%M"
            ),
            "kieu_sua": "AI sửa theo kiểm định",
            "cau_truoc_sua": ban_cu
        })

        cau_moi["lich_su_sua"] = lich_su
        cau_moi["temp_id"] = str(
            uuid.uuid4()
        )
        cau_moi["trang_thai"] = (
            "AI đã sửa – Chờ kiểm định lại"
        )

        cau_moi = chuan_hoa_cau_truc_cau_hoi(
            cau_moi
        )

        loi_du_kien = kiem_tra_cau_hoi_day_du_du_kien(
            cau_moi
        )

        if loi_du_kien:
            cau_moi["loi_cau_truc"] = list(
                cau_moi.get(
                    "loi_cau_truc",
                    []
                )
            ) + loi_du_kien

        return cau_moi, ""

    except Exception as e:
        if la_loi_429(e):
            return None, (
                "API đang bị giới hạn 429. "
                "Câu gốc vẫn được giữ nguyên; có thể thử lại sau."
            )

        return None, str(e)


def sua_cau_thu_cong_tu_form(question, prefix):
    """
    Render form sửa thủ công ngay trong khu vực chờ duyệt.
    Trả câu mới nếu GV bấm Lưu chỉnh sửa.
    """
    dang = question.get(
        "dang_cau",
        ""
    )

    with st.form(
        key=f"manual_edit_form_{prefix}"
    ):
        tinh_huong = st.text_area(
            "Tình huống / dữ liệu",
            value=str(
                question.get(
                    "tinh_huong",
                    ""
                )
            ),
            height=140
        )

        cau_hoi = st.text_area(
            "Câu hỏi / câu lệnh",
            value=str(
                question.get(
                    "cau_hoi",
                    ""
                )
            ),
            height=110
        )

        lua_chon = list(
            question.get(
                "lua_chon",
                []
            )
            or []
        )

        meta = [
            dict(x)
            for x in (
                question.get(
                    "nhan_dinh_meta",
                    []
                )
                or []
            )
        ]

        dap_an = str(
            question.get(
                "dap_an",
                ""
            )
        )

        giai_thich = str(
            question.get(
                "giai_thich",
                ""
            )
        )

        if dang == "Trắc nghiệm 4 lựa chọn":
            while len(lua_chon) < 4:
                lua_chon.append("")

            lua_chon_moi = []
            for i, ky in enumerate(
                ["A", "B", "C", "D"]
            ):
                val = st.text_input(
                    f"Phương án {ky}",
                    value=bo_nhan_phuong_an(
                        lua_chon[i]
                    )
                )
                lua_chon_moi.append(
                    f"{ky}. {val.strip()}"
                )

            dap_an = st.selectbox(
                "Đáp án đúng",
                ["A", "B", "C", "D"],
                index=(
                    ["A", "B", "C", "D"].index(
                        dap_an.strip().upper()
                    )
                    if dap_an.strip().upper()
                    in ["A", "B", "C", "D"]
                    else 0
                )
            )

        elif dang == "Đúng / Sai":
            while len(meta) < 4:
                meta.append({
                    "noi_dung": "",
                    "yccd": "",
                    "muc_do": "",
                    "thanh_phan_nang_luc": "",
                    "dap_an": "Đúng",
                    "giai_thich": ""
                })

            for i, ky in enumerate(
                ["a", "b", "c", "d"]
            ):
                st.markdown(
                    f"**Ý {ky}**"
                )

                meta[i]["noi_dung"] = st.text_area(
                    f"Nội dung ý {ky}",
                    value=str(
                        meta[i].get(
                            "noi_dung",
                            ""
                        )
                    ),
                    key=f"{prefix}_noi_dung_{ky}",
                    height=80
                )

                # Metadata mục tiêu chỉ hiển thị, không cho sửa nhầm.
                st.caption(
                    f"YCCĐ: {meta[i].get('yccd', '')} • "
                    f"Mức độ: {meta[i].get('muc_do', '')} • "
                    f"Năng lực: {meta[i].get('thanh_phan_nang_luc', '')}"
                )

                dap_ds = str(
                    meta[i].get(
                        "dap_an",
                        "Đúng"
                    )
                ).strip()

                meta[i]["dap_an"] = st.selectbox(
                    f"Đáp án ý {ky}",
                    ["Đúng", "Sai"],
                    index=(
                        1
                        if dap_ds.casefold()
                        in ["sai", "false"]
                        else 0
                    ),
                    key=f"{prefix}_dap_an_{ky}"
                )

                meta[i]["giai_thich"] = st.text_area(
                    f"Giải thích ý {ky}",
                    value=str(
                        meta[i].get(
                            "giai_thich",
                            ""
                        )
                    ),
                    key=f"{prefix}_gt_{ky}",
                    height=70
                )

        elif dang == "Trả lời ngắn":
            dap_an = st.text_input(
                "Đáp án",
                value=dap_an
            )

        giai_thich = st.text_area(
            "Giải thích chung",
            value=giai_thich,
            height=110
        )

        submitted = st.form_submit_button(
            "💾 LƯU CHỈNH SỬA THỦ CÔNG",
            type="primary",
            use_container_width=True
        )

    if not submitted:
        return None

    cau_moi = dict(
        question
    )

    cau_moi["tinh_huong"] = (
        tinh_huong.strip()
    )
    cau_moi["cau_hoi"] = (
        cau_hoi.strip()
    )
    cau_moi["giai_thich"] = (
        giai_thich.strip()
    )

    if dang == "Trắc nghiệm 4 lựa chọn":
        cau_moi["lua_chon"] = (
            lua_chon_moi
        )
        cau_moi["dap_an"] = (
            dap_an
        )

    elif dang == "Đúng / Sai":
        cau_moi["nhan_dinh_meta"] = (
            meta[:4]
        )

    elif dang == "Trả lời ngắn":
        cau_moi["dap_an"] = (
            str(dap_an).strip()
        )

    # Lịch sử sửa
    lich_su = list(
        question.get(
            "lich_su_sua",
            []
        )
        or []
    )

    ban_cu = dict(
        question
    )
    ban_cu.pop(
        "lich_su_sua",
        None
    )

    lich_su.append({
        "thoi_gian": datetime.now().strftime(
            "%d/%m/%Y %H:%M"
        ),
        "kieu_sua": "GV sửa thủ công",
        "cau_truoc_sua": ban_cu
    })

    cau_moi["lich_su_sua"] = (
        lich_su
    )
    cau_moi["temp_id"] = str(
        uuid.uuid4()
    )
    cau_moi["trang_thai"] = (
        "GV đã sửa – Chờ kiểm định lại"
    )

    # Kết quả kiểm định cũ không còn giá trị sau sửa.
    cau_moi = chuan_hoa_cau_truc_cau_hoi(
        cau_moi
    )

    loi_du_kien = kiem_tra_cau_hoi_day_du_du_kien(
        cau_moi
    )

    if loi_du_kien:
        cau_moi["loi_cau_truc"] = list(
            cau_moi.get(
                "loi_cau_truc",
                []
            )
        ) + loi_du_kien

    return cau_moi


def hien_thi_so_sanh_phien_ban(question):
    lich_su = list(
        question.get(
            "lich_su_sua",
            []
        )
        or []
    )

    if not lich_su:
        return

    with st.expander(
        f"🕘 Lịch sử sửa ({len(lich_su)})"
    ):
        for i, ban in enumerate(
            reversed(lich_su),
            start=1
        ):
            st.markdown(
                f"**Lần {i}: {ban.get('kieu_sua', '')} "
                f"— {ban.get('thoi_gian', '')}**"
            )

            cu = ban.get(
                "cau_truoc_sua",
                {}
            )

            st.write(
                "Câu trước sửa:",
                cu.get(
                    "cau_hoi",
                    ""
                )
            )

            st.divider()



# ==========================================================
# LƯU 1 CÂU VÀO NGÂN HÀNG
# ==========================================================
def luu_cau_vao_ngan_hang(question):

    question = chuan_hoa_cau_truc_cau_hoi(
        question
    )

    if question.get("loi_cau_truc"):
        return False

    bank = doc_ngan_hang()

    # Chống trùng nội bộ trước khi lưu.
    # Chặn câu trùng/gần như trùng ở mức rất cao.
    kq_trung_luu = kiem_tra_trung_gan(
        question,
        bank
    )

    if kq_trung_luu.get(
        "ti_le_cao_nhat",
        0
    ) >= 0.92:
        return False

    cau_moi = dict(question)

    # Ngân hàng là kho dùng chung. Câu tốt nghiệp cũng mặc định được phép
    # xuất hiện trong luyện tập cá nhân nếu phù hợp YCCĐ/kiến thức.
    cau_moi.setdefault(
        "duoc_dung_luyen_hs",
        True
    )

    temp_id_kiem_dinh = str(
        question.get("temp_id", "")
    )

    if temp_id_kiem_dinh:
        kd = st.session_state.ket_qua_kiem_dinh.get(
            temp_id_kiem_dinh
        )

        if kd:
            cau_moi["kiem_dinh_ai"] = kd

    cau_moi.pop(
        "temp_id",
        None
    )

    cau_moi["id"] = str(
        uuid.uuid4()
    )

    cau_moi["ngay_tao"] = (
        datetime.now().strftime(
            "%d/%m/%Y %H:%M"
        )
    )

    cau_moi["trang_thai"] = (
        "Đã duyệt"
    )

    bank.append(
        cau_moi
    )

    luu_ngan_hang(
        bank
    )

    return True


# ==========================================================
# NHẬN DIỆN THƯƠNG HIỆU - TRẠM SINH HỌC
# ==========================================================
def hien_thi_dau_trang_tram_sinh_hoc(khu_vuc="hocsinh"):
    """Nhận diện thống nhất cho GV/HS, không để Markdown hiểu nhầm HTML thành code block."""
    la_gv = str(khu_vuc).strip().lower() == "giaovien"
    role = "KHU VỰC GIÁO VIÊN" if la_gv else "HỌC TẬP • ÔN LUYỆN • CHINH PHỤC"
    slogan = (
        "Kiến tạo học liệu • Theo dõi tiến bộ • Đồng hành cùng học sinh"
        if la_gv
        else "Học đúng trọng tâm • Luyện đúng năng lực • Tiến bộ mỗi ngày"
    )
    chips_html = "" if la_gv else (
        '<div class="tram-bio-icons">'
        '<span class="tram-bio-chip">🧬 Di truyền</span>'
        '<span class="tram-bio-chip">🔬 Khám phá</span>'
        '<span class="tram-bio-chip">🧫 Tế bào</span>'
        '<span class="tram-bio-chip">🌿 Sự sống</span>'
        '</div>'
    )

    css = """
    <style>
    .tram-bio-shell{max-width:1160px;margin:0 auto 1rem auto}
    .tram-bio-head{position:relative;overflow:hidden;border-radius:26px;padding:1.45rem 1.55rem;color:#fff;background:linear-gradient(120deg,#123f78 0%,#087c88 54%,#17835c 100%);border:1px solid rgba(255,255,255,.16);box-shadow:0 16px 38px rgba(20,67,93,.15);display:grid;grid-template-columns:minmax(0,1.8fr) minmax(150px,.45fr);gap:1rem;align-items:center}
    .tram-bio-head:before{content:"";position:absolute;width:300px;height:300px;border-radius:50%;right:-145px;top:-175px;background:rgba(255,255,255,.07)}
    .tram-bio-left,.tram-bio-right{position:relative;z-index:1}
    .tram-bio-dept{font-size:.76rem;font-weight:800;letter-spacing:.055em;opacity:.90}
    .tram-bio-school{font-size:1.01rem;font-weight:900;letter-spacing:.028em;margin-top:.15rem}
    .tram-bio-title{font-size:2.48rem;font-weight:950;line-height:1.02;margin:.68rem 0 .30rem 0}
    .tram-bio-role{display:inline-block;padding:.27rem .66rem;border-radius:999px;background:rgba(255,255,255,.12);border:1px solid rgba(255,255,255,.17);font-size:.78rem;font-weight:850;letter-spacing:.035em}
    .tram-bio-slogan{margin-top:.78rem;font-size:1rem;font-weight:700;opacity:.98}
    .tram-bio-icons{margin-top:.78rem;display:flex;gap:.46rem;flex-wrap:wrap}
    .tram-bio-chip{padding:.31rem .67rem;border-radius:999px;background:rgba(255,255,255,.10);border:1px solid rgba(255,255,255,.17);font-size:.77rem;font-weight:730}
    .tram-bio-right{display:flex;align-items:center;justify-content:center;min-height:130px}
    .tram-bio-dna{font-size:5rem;line-height:1;opacity:.16;user-select:none;transform:rotate(-8deg)}
    @media(max-width:820px){.tram-bio-head{grid-template-columns:1fr;padding:1.18rem 1.12rem}.tram-bio-right{display:none}.tram-bio-title{font-size:2rem}}
    </style>
    """

    # Thân HTML để trên một dòng liên tục để Streamlit/Markdown không biến thẻ div thành khối code.
    body = (
        '<div class="tram-bio-shell"><div class="tram-bio-head">'
        '<div class="tram-bio-left">'
        '<div class="tram-bio-dept">SỞ GIÁO DỤC VÀ ĐÀO TẠO GIA LAI</div>'
        '<div class="tram-bio-school">TRƯỜNG THPT SỐ 1 PHÙ CÁT</div>'
        '<div class="tram-bio-title">TRẠM SINH HỌC</div>'
        f'<div class="tram-bio-role">{role}</div>'
        f'<div class="tram-bio-slogan">{slogan}</div>'
        f'{chips_html}'
        '</div>'
        '<div class="tram-bio-right"><div class="tram-bio-dna">🧬</div></div>'
        '</div></div>'
    )
    st.markdown(css + body, unsafe_allow_html=True)

# ==========================================================
# TRANG CHỦ
# ==========================================================
def trang_chu():
    profile = _nang_cap_avatar_profile_ben_vung(doc_ho_so_giao_vien())
    ten = str(profile.get("ten", "") or "").strip()
    chuc_vu = str(profile.get("chuc_vu", "") or "").strip()
    don_vi = str(profile.get("don_vi", "") or "").strip()
    loi_chao = str(profile.get("loi_chao", "") or "").strip()
    avatar_uri = _avatar_data_uri(
        profile.get("avatar_path"),
        profile.get("avatar_base64", ""),
        profile.get("avatar_mime_type", "")
    )

    st.markdown(
        """
        <style>
        .home-wrap { max-width:1180px; margin:0 auto; }
        .home-hero {
            position:relative; overflow:hidden; border-radius:30px; padding:2.35rem 2.45rem;
            background:linear-gradient(125deg,#0f3d88 0%,#0b78b5 48%,#14966f 100%);
            box-shadow:0 24px 70px rgba(15,61,136,.20); color:white;
            border:1px solid rgba(255,255,255,.22);
        }
        .home-hero:before {
            content:""; position:absolute; width:330px; height:330px; border-radius:50%;
            right:-90px; top:-150px; background:rgba(255,255,255,.09);
        }
        .home-hero:after {
            content:"🧬"; position:absolute; right:2.0rem; bottom:-1.3rem; font-size:9rem;
            opacity:.13; transform:rotate(-8deg);
        }
        .home-kicker { display:inline-flex; gap:.45rem; align-items:center; padding:.38rem .78rem;
            border-radius:999px; background:rgba(255,255,255,.13); border:1px solid rgba(255,255,255,.26);
            font-size:.84rem; font-weight:800; letter-spacing:.01em; }
        .home-title { margin:.8rem 0 .35rem 0; font-size:2.55rem; line-height:1.08; font-weight:900; color:white; }
        .home-sub { max-width:760px; font-size:1.05rem; line-height:1.55; opacity:.95; margin:0; }
        .home-badges { display:flex; flex-wrap:wrap; gap:.55rem; margin-top:1.15rem; }
        .home-badge { padding:.38rem .72rem; border-radius:999px; background:rgba(255,255,255,.10);
            border:1px solid rgba(255,255,255,.20); font-size:.82rem; font-weight:700; }
        .home-teacher {
            margin-top:1rem; display:flex; align-items:center; gap:.75rem; width:max-content; max-width:100%;
            padding:.62rem .8rem; border-radius:16px; background:rgba(255,255,255,.12);
            border:1px solid rgba(255,255,255,.20);
        }
        .home-teacher img { width:50px; height:50px; border-radius:50%; object-fit:cover; border:2px solid rgba(255,255,255,.85); }
        .home-teacher-fallback { width:50px; height:50px; border-radius:50%; display:flex; align-items:center; justify-content:center;
            background:rgba(255,255,255,.18); border:2px solid rgba(255,255,255,.75); font-size:1.55rem; }
        .home-teacher-name { font-weight:850; color:white; line-height:1.1; }
        .home-teacher-meta { font-size:.80rem; opacity:.84; margin-top:.2rem; }
        .home-section-title { margin:1.55rem 0 .8rem 0; text-align:center; font-size:1.32rem; font-weight:850; color:#17315f; }
        .home-role-card {
            min-height:170px; padding:1.15rem 1.2rem; border-radius:22px; background:rgba(255,255,255,.94);
            border:1px solid #dbe7f4; box-shadow:0 12px 34px rgba(15,23,42,.07); margin-bottom:.65rem;
        }
        .home-role-icon { font-size:2.15rem; margin-bottom:.35rem; }
        .home-role-title { font-size:1.18rem; font-weight:900; color:#17315f; }
        .home-role-desc { color:#5b6b82; font-size:.90rem; line-height:1.45; margin:.38rem 0 .7rem 0; }
        .home-role-list { font-size:.83rem; color:#40536e; line-height:1.65; }
        .home-science-grid { display:grid; grid-template-columns:repeat(4,1fr); gap:.7rem; margin-top:1.2rem; }
        .home-science-item { padding:.78rem .85rem; border-radius:16px; text-align:center;
            background:linear-gradient(145deg,rgba(239,246,255,.96),rgba(240,253,244,.90));
            border:1px solid #dce9f5; color:#30435e; font-size:.82rem; font-weight:750; }
        .home-science-item span { display:block; font-size:1.45rem; margin-bottom:.2rem; }
        .home-footer-note { margin:1.0rem auto 0 auto; text-align:center; color:#718096; font-size:.80rem; }
        @media (max-width: 800px) {
            .home-title { font-size:2rem; }
            .home-science-grid { grid-template-columns:repeat(2,1fr); }
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    if avatar_uri:
        avatar_html = f'<img src="{avatar_uri}" alt="Ảnh giáo viên">'
    else:
        avatar_html = '<div class="home-teacher-fallback">👩‍🏫</div>'

    meta = " • ".join([x for x in [chuc_vu, don_vi] if x])
    teacher_html = ""
    if ten or meta or loi_chao:
        teacher_html = (
            '<div class="home-teacher">'
            + avatar_html
            + '<div>'
            + '<div class="home-teacher-name">' + _html_escape(ten or "Giáo viên Sinh học") + '</div>'
            + '<div class="home-teacher-meta">' + _html_escape(meta or "Phụ trách hệ thống ôn tập") + '</div>'
            + '</div></div>'
        )

    st.markdown(
        f"""
        <div class="home-wrap">
          <div class="home-hero">
            <div class="home-kicker">🔬 SỞ GD&ĐT GIA LAI • THPT SỐ 1 PHÙ CÁT</div>
            <div class="home-title">TRẠM SINH HỌC</div>
            <p class="home-sub">Học tập, ôn luyện và phát triển năng lực Sinh học trên một hệ thống thống nhất. Giáo viên kiến tạo học liệu; học sinh luyện tập đúng trọng tâm và theo dõi tiến bộ của chính mình.</p>
            <div class="home-badges">
              <span class="home-badge">🧬 YCCĐ làm trung tâm</span>
              <span class="home-badge">🎯 Cá nhân hóa luyện tập</span>
              <span class="home-badge">📊 Theo dõi tiến bộ</span>
              <span class="home-badge">🌱 Học từ hạt giống</span>
            </div>
            {teacher_html}
          </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.markdown('<div class="home-section-title">Chọn không gian để bắt đầu</div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2, gap="large")

    with col1:
        st.markdown(
            """
            <div class="home-role-card">
              <div class="home-role-icon">👩‍🏫</div>
              <div class="home-role-title">Không gian Giáo viên</div>
              <div class="home-role-desc">Thiết kế ngân hàng câu hỏi, tạo đề và theo dõi việc học của từng học sinh.</div>
              <div class="home-role-list">✓ Quản lý YCCĐ & năng lực<br>✓ Ngân hàng hạt giống & đề tốt nghiệp<br>✓ Phân tích lớp học & tiến bộ học sinh</div>
            </div>
            """,
            unsafe_allow_html=True
        )
        if st.button("👩‍🏫  VÀO KHÔNG GIAN GIÁO VIÊN", type="primary", use_container_width=True, key="home_teacher"):
            st.session_state.vai_tro = "giaovien"
            st.rerun()

    with col2:
        st.markdown(
            """
            <div class="home-role-card">
              <div class="home-role-icon">👨‍🎓</div>
              <div class="home-role-title">Không gian Học sinh</div>
              <div class="home-role-desc">Luyện tập theo mức độ phù hợp, nhận gợi ý kiến thức cần củng cố và xem tiến bộ của chính mình.</div>
              <div class="home-role-list">✓ Luyện theo gợi ý hôm nay<br>✓ Xem mức làm chủ & xu hướng học tập<br>✓ Chỉ xem thứ hạng cá nhân của mình</div>
            </div>
            """,
            unsafe_allow_html=True
        )
        if st.button("👨‍🎓  VÀO KHÔNG GIAN HỌC SINH", use_container_width=True, key="home_student"):
            st.session_state.vai_tro = "hocsinh"
            st.rerun()

    st.markdown(
        """
        <div class="home-science-grid">
          <div class="home-science-item"><span>🧫</span>Tế bào & cơ thể sống</div>
          <div class="home-science-item"><span>🧬</span>Di truyền & tiến hóa</div>
          <div class="home-science-item"><span>🌿</span>Sinh thái & môi trường</div>
          <div class="home-science-item"><span>📈</span>Dữ liệu & tiến bộ</div>
        </div>
        <div class="home-footer-note">A • T &nbsp;&nbsp; C • G &nbsp;&nbsp; | &nbsp;&nbsp; Học bằng hiểu biết — tiến bộ bằng dữ liệu</div>
        """,
        unsafe_allow_html=True
    )

# ==========================================================
# KHO YCCĐ
# ==========================================================
def kho_yccd():

    st.header("📚 KHO YÊU CẦU CẦN ĐẠT")

    khoi = st.selectbox(
        "Chọn khối",
        list(KHO_YCCD.keys())
    )

    chuong = st.selectbox(
        "Chọn chương / chủ đề",
        list(KHO_YCCD[khoi].keys())
    )

    bai = st.selectbox(
        "Chọn bài / nội dung",
        list(KHO_YCCD[khoi][chuong].keys())
    )

    st.divider()
    st.subheader(f"📖 {bai}")

    # Danh sách YCCĐ của bài hiện tại
    danh_sach_yccd = KHO_YCCD[khoi][chuong][bai]

    for i, yccd in enumerate(danh_sach_yccd, start=1):

        item = {
            "Khối": khoi,
            "Chương": chuong,
            "Bài": bai,
            "YCCĐ": yccd
        }

        yccd_id = tao_id_yccd(
            khoi,
            chuong,
            bai,
            yccd
        )

        checked = st.checkbox(
            f"YCCĐ {i}: {yccd}",
            value=item_da_chon(item),
            key=f"checkbox_{yccd_id}"
        )

        # =========================
        # THÊM YCCĐ
        # =========================
        if checked and item not in st.session_state.yccd_da_chon:

            st.session_state.yccd_da_chon.append(item)

        if yccd_id not in st.session_state.cau_hinh_yccd:
            st.session_state.cau_hinh_yccd[yccd_id] = cau_hinh_mac_dinh(
                item["YCCĐ"]
            )

            st.session_state.cau_hoi_ai = []

            st.rerun()
        # =========================
        # BỎ YCCĐ
        # =========================
        if not checked and item in st.session_state.yccd_da_chon:

            xoa_mot_yccd(item)

            st.rerun()

        # =====================================================
        # CẤU HÌNH YCCĐ ĐANG CHỌN
        # =====================================================
        if item in st.session_state.yccd_da_chon:

            _, cau_hinh = lay_cau_hinh(item)

            with st.container(border=True):

                st.markdown("**⚙️ Thiết lập câu hỏi**")

                col1, col2, col3 = st.columns(3)

                # -------------------------
                # SỐ CÂU
                # -------------------------
                with col1:

                    so_cau = st.number_input(
                        "Số câu",
                        min_value=1,
                        max_value=50,
                        value=cau_hinh["Số câu"],
                        key=f"so_cau_{yccd_id}"
                    )

                            # ----------------------
            # MỨC ĐỘ
            # ----------------------
            with col2:

                ds_muc_do = [
                    "Nhận biết",
                    "Thông hiểu",
                    "Vận dụng"
                ]

                muc_do_goi_y = xac_dinh_muc_do(
                    item["YCCĐ"]
                )

                st.caption(
                    f"🤖 AI đề xuất: {muc_do_goi_y}"
                )

                muc_do_cu = cau_hinh.get(
                    "Mức độ",
                    muc_do_goi_y
                )

                if muc_do_cu not in ds_muc_do:
                    muc_do_cu = muc_do_goi_y

                muc_do = st.selectbox(
                    "Mức độ",
                    ds_muc_do,
                    index=ds_muc_do.index(muc_do_cu),
                    key=f"mucdo_{yccd_id}"
                )
            
                # -------------------------
                # DẠNG CÂU
                # -------------------------
                with col3:

                    ds_dang = [
                        "Trắc nghiệm 4 lựa chọn",
                        "Đúng / Sai",
                        "Trả lời ngắn",
                    ]

                    dang_cu = cau_hinh.get(
                        "Dạng câu hỏi",
                        "Trắc nghiệm 4 lựa chọn"
                    )

                    if dang_cu not in ds_dang:
                        dang_cu = "Trắc nghiệm 4 lựa chọn"

                    dang_cau = st.selectbox(
                        "Dạng câu hỏi",
                        ds_dang,
                        index=ds_dang.index(dang_cu),
                        key=f"dang_{yccd_id}"
                    )

                # =================================================
                # THÀNH PHẦN NĂNG LỰC
                # =================================================
                nang_luc_goi_y = goi_y_thanh_phan_nang_luc(
                    item["YCCĐ"],
                    muc_do,
                    dang_cau
                )

                st.caption(
                    f"🧠 Gợi ý thành phần năng lực: {nang_luc_goi_y}"
                )

                nang_luc_cu = cau_hinh.get(
                    "Thành phần năng lực",
                    nang_luc_goi_y
                )

                if nang_luc_cu not in THANH_PHAN_NANG_LUC:
                    nang_luc_cu = nang_luc_goi_y

                nang_luc = st.selectbox(
                    "Thành phần năng lực",
                    THANH_PHAN_NANG_LUC,
                    index=THANH_PHAN_NANG_LUC.index(
                        nang_luc_cu
                    ),
                    key=f"nangluc_{yccd_id}"
                )

                # =================================================
                # ĐÚNG / SAI: CẤU HÌNH RIÊNG 4 Ý
                # =================================================
                cau_hinh_4_y = {}

            if dang_cau == "Đúng / Sai":

                st.markdown(
                    "**🎯 Cấu hình 4 ý Đúng / Sai**"
                )

                st.caption(
                    "Mỗi ý a, b, c, d có thể chọn YCCĐ và "
                    "mức độ riêng. Cả 4 ý dùng chung một "
                    "tình huống/dữ liệu."
                )

                chi_tiet_cu = cau_hinh.get(
                    "Chi tiết 4 ý",
                    {}
                )

                for ky_hieu in ["a", "b", "c", "d"]:

                    col_yccd, col_mucdo, col_nangluc = st.columns([3, 1, 2])

                    cau_hinh_y_cu = chi_tiet_cu.get(
                        ky_hieu,
                        {}
                    )

                    # YCCĐ mặc định
                    yccd_mac_dinh = cau_hinh_y_cu.get(
                        "yccd",
                        yccd
                    )

                    if yccd_mac_dinh not in danh_sach_yccd:
                        yccd_mac_dinh = yccd

                    # Mức độ mặc định
                    mucdo_mac_dinh = cau_hinh_y_cu.get(
                        "muc_do",
                        muc_do
                    )

                    if mucdo_mac_dinh not in ds_muc_do:
                        mucdo_mac_dinh = muc_do

                    # YCCĐ của từng ý
                    with col_yccd:

                        yccd_y = st.selectbox(
                            f"YCCĐ ý {ky_hieu}",
                            danh_sach_yccd,
                            index=danh_sach_yccd.index(
                                yccd_mac_dinh
                            ),
                            key=f"yccd_ds_{yccd_id}_{ky_hieu}"
                        )

                    mucdo_tu_yccd = xac_dinh_muc_do(yccd_y)
                    # Mức độ của từng ý
                    with col_mucdo:

                            mucdo_y = st.selectbox(
                                f"Mức độ ý {ky_hieu}",
                                ds_muc_do,
                                index=ds_muc_do.index(
                                    mucdo_tu_yccd
                                ),
                                key=f"mucdo_ds_{yccd_id}_{ky_hieu}_{yccd_y}"
                            )

                    nang_luc_y_goi_y = goi_y_thanh_phan_nang_luc(
                        yccd_y,
                        mucdo_y,
                        "Đúng / Sai"
                    )

                    nang_luc_y_cu = cau_hinh_y_cu.get(
                        "thanh_phan_nang_luc",
                        nang_luc_y_goi_y
                    )

                    ds_nl_ds = [
                        "Tìm hiểu thế giới sống",
                        "Vận dụng kiến thức, kĩ năng đã học",
                        "Nhận thức sinh học"
                    ]

                    if nang_luc_y_cu not in ds_nl_ds:
                        nang_luc_y_cu = nang_luc_y_goi_y

                    with col_nangluc:
                        nang_luc_y = st.selectbox(
                            f"Năng lực ý {ky_hieu}",
                            ds_nl_ds,
                            index=ds_nl_ds.index(
                                nang_luc_y_cu
                            ),
                            key=f"nangluc_ds_{yccd_id}_{ky_hieu}_{yccd_y}_{mucdo_y}"
                        )

                    cau_hinh_4_y[ky_hieu] = {
                            "yccd": yccd_y,
                            "muc_do": mucdo_y,
                            "thanh_phan_nang_luc": nang_luc_y
                        }

            # =================================================
            # NGUỒN THAM CHIẾU
            # =================================================
            st.markdown("**📚 Nguồn tham chiếu**")

            ds_nguon = [
                "SGK / Chương trình",
                "NCBI",
                "PubMed",
                "PubMed Central (PMC)",
                "Nguồn web khoa học uy tín",
                "Tài liệu giáo viên tải lên"
            ]

            nguon_mac_dinh = cau_hinh.get(
                "Nguồn tham chiếu",
                ["SGK / Chương trình"]
            )

            nguon_mac_dinh = [
                x
                for x in nguon_mac_dinh
                if x in ds_nguon
            ]

            nguon = st.multiselect(
                "Có thể chọn nhiều nguồn",
                ds_nguon,
                default=nguon_mac_dinh,
                key=f"nguon_{yccd_id}"
            )

            tai_lieu_gv = []

            if "Tài liệu giáo viên tải lên" in nguon:

                thu_muc_kho = "kho_tai_lieu_gv"

                if not os.path.exists(thu_muc_kho):
                    os.makedirs(thu_muc_kho)

                danh_sach_tai_lieu = os.listdir(thu_muc_kho)

                if danh_sach_tai_lieu:

                    tai_lieu_gv = st.multiselect(
                        "📚 Chọn tài liệu từ Kho tài liệu GV",
                        danh_sach_tai_lieu,
                        default=cau_hinh.get(
                            "Tài liệu giáo viên",
                            []
                        ),
                        key=f"tai_lieu_kho_{yccd_id}"
                    )

                else:
                    st.info(
                        "Kho tài liệu hiện đang trống. "
                        "Hãy vào 📁 Kho tài liệu GV để tải tài liệu."
                    )

            # =================================================
            # LƯU CẤU HÌNH
            # =================================================
            st.session_state.cau_hinh_yccd[yccd_id] = {
                "Số câu": so_cau,
                "Mức độ": muc_do,
                "Dạng câu hỏi": dang_cau,
                "Thành phần năng lực": nang_luc,
                "Nguồn tham chiếu": nguon,
                "Tài liệu giáo viên": tai_lieu_gv,
                "Chi tiết 4 ý": (
                    cau_hinh_4_y
                    if dang_cau == "Đúng / Sai"
                    else {}
                )
            }

    # ==========================================================
    # YCCĐ ĐÃ CHỌN
    # ==========================================================
    st.divider()

    st.subheader("✅ YCCĐ đã chọn")

    ds_hien_tai = list(
        st.session_state.yccd_da_chon
    )

    if not ds_hien_tai:

        st.warning("Chưa chọn YCCĐ nào.")

    else:

        st.success(
            f"Đã chọn {len(ds_hien_tai)} YCCĐ."
        )

        tong_so_cau = 0

        for i, item in enumerate(
            ds_hien_tai,
            start=1
        ):

            yccd_id, cau_hinh = lay_cau_hinh(item)

            tong_so_cau += cau_hinh["Số câu"]

            col1, col2 = st.columns([9, 1])

            with col1:

                with st.expander(
                    f"{i}. [{item['Khối']}] {item['Bài']}"
                ):

                    st.write(
                        "**YCCĐ:** " + item["YCCĐ"]
                    )

                    st.write(
                        f"**Số câu:** {cau_hinh['Số câu']}"
                    )

                    st.write(
                        f"**Mức độ:** {cau_hinh['Mức độ']}"
                    )

                    st.write(
                        f"**Dạng:** {cau_hinh['Dạng câu hỏi']}"
                    )

                    st.write(
                        f"**Thành phần năng lực:** "
                        f"{cau_hinh.get('Thành phần năng lực', '')}"
                    )

                    if cau_hinh["Nguồn tham chiếu"]:

                        st.write(
                            "**Nguồn:** "
                            + ", ".join(
                                cau_hinh["Nguồn tham chiếu"]
                            )
                        )

                    # Hiển thị cấu hình 4 ý
                    if (
                        cau_hinh["Dạng câu hỏi"] == "Đúng / Sai"
                        and cau_hinh.get("Chi tiết 4 ý")
                    ):

                        st.markdown(
                            "**🎯 Cấu hình 4 ý:**"
                        )

                        for ky_hieu in ["a", "b", "c", "d"]:

                            meta = cau_hinh[
                                "Chi tiết 4 ý"
                            ].get(
                                ky_hieu,
                                {}
                            )

                            st.write(
                                f"Ý {ky_hieu}: "
                                f"{meta.get('muc_do', '')} — "
                                f"{meta.get('thanh_phan_nang_luc', '')} — "
                                f"{meta.get('yccd', '')}"
                            )

            with col2:

                if st.button(
                    "❌",
                    key=f"xoa_{yccd_id}",
                    help="Xóa YCCĐ này"
                ):

                    xoa_mot_yccd(item)

                    st.rerun()

        st.info(
            f"Tổng số câu dự kiến: "
            f"**{tong_so_cau} câu**"
        )

        if st.button(
            "🗑️ XÓA TOÀN BỘ YCCĐ",
            use_container_width=True
        ):

            xoa_toan_bo()

            st.rerun()


# ==========================================================
# BẢNG ĐẶC TẢ
# ==========================================================
def bang_dac_ta():

    st.header(
        "📋 BẢNG ĐẶC TẢ"
    )

    if not st.session_state.yccd_da_chon:

        st.warning(
            "Chưa có YCCĐ nào."
        )

        return

    bang = tao_bang_dac_ta()

    df = pd.DataFrame(
        bang
    )

    st.dataframe(
        df,
        use_container_width=True,
        hide_index=True
    )

    tong = int(
        df["Số câu"].sum()
    )

    c1, c2, c3 = st.columns(3)

    with c1:
        st.metric(
            "YCCĐ",
            len(df)
        )

    with c2:
        st.metric(
            "Tổng số câu",
            tong
        )

    with c3:
        st.metric(
            "Dạng câu",
            df[
                "Dạng câu hỏi"
            ].nunique()
        )


# ==========================================================
# HIỂN THỊ 1 CÂU AI
# ==========================================================
def hien_thi_cau_ai(question, index):

    temp_id = question["temp_id"]

    trang_thai = st.session_state.trang_thai_duyet.get(
        temp_id,
        "Chờ duyệt"
    )

    with st.container(border=True):

        st.markdown(f"### Câu {index}")

        if question.get("dang_cau") == "Đúng / Sai":
            st.caption(
                f"{question.get('khoi', '')} • "
                f"{question.get('bai', '')} • "
                "Đúng / Sai • Tích hợp 4 ý"
            )
            st.caption(
                "YCCĐ, mức độ và thành phần năng lực được quản lí riêng theo từng ý a–d."
            )
        else:
            st.caption(
                f"{question.get('khoi', '')} • "
                f"{question.get('bai', '')} • "
                f"{question.get('muc_do', '')} • "
                f"{question.get('dang_cau', '')}"
            )

            st.write("**YCCĐ:**")
            st.write(question.get("yccd", ""))

            st.write(
                "**Thành phần năng lực:**",
                question.get("thanh_phan_nang_luc", "")
            )

        if question.get("nguon_thuc_te"):
            st.success(
                "📚 Câu này được phát sinh từ nguồn "
                "đã được hệ thống truy xuất thực tế."
            )

        kieu_dl_hien = question.get(
            "kieu_du_lieu",
            "khong_dinh_luong"
        )

        if kieu_dl_hien == "du_lieu_that":
            st.info(
                "📊 Dữ liệu định lượng: **DỮ LIỆU THẬT TỪ NGUỒN**"
            )
        elif kieu_dl_hien == "mo_phong":
            st.warning(
                "🧮 Dữ liệu định lượng: **MÔ PHỎNG** "
                "(chỉ dùng khi nguồn không có số liệu phù hợp)."
            )
        else:
            st.caption(
                "Dạng dữ liệu: không sử dụng dữ liệu định lượng."
            )

        if question.get("loi_nguon"):
            st.error(
                "⛔ "
                + str(
                    question.get(
                        "loi_nguon"
                    )
                )
            )

        canh_bao_gan_gia = kiem_tra_nguy_co_gan_gia_du_lieu(
            question
        )

        if canh_bao_gan_gia:
            st.error(
                "⛔ " + canh_bao_gan_gia
            )

        loi_cau_truc = question.get("loi_cau_truc", [])
        if loi_cau_truc:
            st.error(
                "🚫 Câu này có lỗi cấu trúc và không nên duyệt:"
            )
            for loi in loi_cau_truc:
                st.write(f"- {loi}")

        dang_cau = question.get("dang_cau", "")

        # ==================================================
        # ĐÚNG / SAI
        # ==================================================
        if dang_cau == "Đúng / Sai":

            tinh_huong = question.get(
                "tinh_huong",
                ""
            )

            st.markdown(
                "**Tình huống / dữ liệu:**"
            )

            if tinh_huong:
                st.write(tinh_huong)

            cau_lenh = str(
                question.get("cau_hoi", "")
            ).strip()

            st.markdown("**Câu hỏi:**")
            if cau_lenh:
                st.write(cau_lenh)
            else:
                st.write(
                    "Dựa vào thông tin trên, hãy xác định mỗi nhận định sau là Đúng hay Sai."
                )

            st.markdown(
                "**Các nhận định:**"
            )

            nhan_dinh_meta = question.get(
                "nhan_dinh_meta",
                []
            )

            # ----------------------------------------------
            # Nếu Gemini trả nhan_dinh_meta
            # ----------------------------------------------
            if nhan_dinh_meta:

                ky_hieu_ds = [
                    "a",
                    "b",
                    "c",
                    "d"
                ]

                for ky_hieu, nhan_dinh in zip(
                    ky_hieu_ds,
                    nhan_dinh_meta
                ):

                    noi_dung = nhan_dinh.get(
                        "noi_dung",
                        ""
                    ).strip()

                    # Xóa a), b), c), d) nếu Gemini
                    # đã tự ghi vào đầu nội dung
                    cac_nhan = [
                        f"{ky_hieu})",
                        f"{ky_hieu}.",
                        f"{ky_hieu}:"
                    ]

                    for nhan in cac_nhan:

                        if noi_dung.lower().startswith(
                            nhan.lower()
                        ):
                            noi_dung = noi_dung[
                                len(nhan):
                            ].strip()

                            break

                    st.write(
                        f"**{ky_hieu})** {noi_dung}"
                    )
                    st.caption(
                        f"Ý {ky_hieu}: "
                        f"{nhan_dinh.get('muc_do', '')} • "
                        f"{nhan_dinh.get('thanh_phan_nang_luc', '')} • "
                        f"YCCĐ: {nhan_dinh.get('yccd', '')}"
                    )

            # ----------------------------------------------
            # Dự phòng nếu không có nhan_dinh_meta
            # ----------------------------------------------
            else:

                lua_chon = question.get(
                    "lua_chon",
                    []
                )

                for i, lc in enumerate(
                    lua_chon[:4]
                ):

                    ky_hieu = [
                        "a",
                        "b",
                        "c",
                        "d"
                    ][i]

                    noi_dung = str(lc).strip()

                    cac_nhan = [
                        f"{ky_hieu})",
                        f"{ky_hieu}.",
                        f"{ky_hieu}:"
                    ]

                    for nhan in cac_nhan:

                        if noi_dung.lower().startswith(
                            nhan.lower()
                        ):
                            noi_dung = noi_dung[
                                len(nhan):
                            ].strip()

                            break

                    st.write(
                        f"**{ky_hieu})** {noi_dung}"
                    )

        # ==================================================
        # CÁC DẠNG CÂU KHÁC
        # ==================================================
        else:

            st.markdown(
                "**Câu hỏi:**"
            )

            st.write(
                question.get(
                    "cau_hoi",
                    ""
                )
            )

            if dang_cau == "Trắc nghiệm 4 lựa chọn":

                lua_chon = question.get(
                    "lua_chon",
                    []
                )

                for lc in lua_chon:
                    st.write(lc)

        # ==================================================
        # ĐÁP ÁN + GIẢI THÍCH
        # ==================================================
        with st.expander(
            "👁️ Xem đáp án và giải thích"
        ):

            # ----------------------------------------------
            # ĐÚNG / SAI
            # ----------------------------------------------
            if (
                dang_cau == "Đúng / Sai"
                and question.get("nhan_dinh_meta")
            ):

                nhan_dinh_meta = question[
                    "nhan_dinh_meta"
                ]

                # Tạo đáp án chuẩn từ từng ý
                dap_an_hien_thi = []

                for ky_hieu, nhan_dinh in zip(
                    ["a", "b", "c", "d"],
                    nhan_dinh_meta
                ):

                    dap_an_y = nhan_dinh.get(
                        "dap_an",
                        ""
                    )

                    dap_an_hien_thi.append(
                        f"{ky_hieu} - {dap_an_y}"
                    )

                st.write(
                    "**Đáp án:**",
                    ", ".join(dap_an_hien_thi)
                )

                st.markdown(
                    "**Giải thích từng ý:**"
                )

                for ky_hieu, nhan_dinh in zip(
                    ["a", "b", "c", "d"],
                    nhan_dinh_meta
                ):

                    st.write(
                        f"**{ky_hieu}) "
                        f"{nhan_dinh.get('dap_an', '')}:** "
                        f"{nhan_dinh.get('giai_thich', '')}"
                    )

            # ----------------------------------------------
            # CÁC DẠNG KHÁC
            # ----------------------------------------------
            else:

                st.write(
                    "**Đáp án:**",
                    question.get(
                        "dap_an",
                        ""
                    )
                )

                st.write(
                    "**Giải thích:**",
                    question.get(
                        "giai_thich",
                        ""
                    )
                )

            st.write(
                "**Nguồn:**",
                question.get(
                    "nguon",
                    ""
                )
            )

            if question.get(
                "nguon_url"
            ):
                st.caption(
                    "Nguồn truy xuất: "
                    + str(
                        question.get(
                            "nguon_url",
                            ""
                        )
                    )
                )

        # ==================================================
        # KẾT QUẢ KIỂM ĐỊNH AI
        # ==================================================
        st.markdown("---")
        hien_thi_ket_qua_kiem_dinh(
            question
        )

        # ==================================================
        # TRẠNG THÁI DUYỆT
        # ==================================================
        if trang_thai == "Đã lưu":

            st.success(
                "✅ Câu này đã được lưu "
                "vào Ngân hàng câu hỏi."
            )

        elif trang_thai == "Không lưu":

            st.warning(
                "🗑️ Câu này đã được loại."
            )

        else:

            col1, col2 = st.columns(2)

            with col1:

                if st.button(
                    "✅ DUYỆT & LƯU",
                    key=f"duyet_{temp_id}",
                    use_container_width=True,
                    type="primary",
                    disabled=bool(
                        question.get("loi_cau_truc", [])
                    )
                ):

                    luu_cau_vao_ngan_hang(
                        question
                    )

                    st.session_state.trang_thai_duyet[
                        temp_id
                    ] = "Đã lưu"

                    st.rerun()

            with col2:

                if st.button(
                    "🗑️ KHÔNG LƯU",
                    key=f"bo_{temp_id}",
                    use_container_width=True
                ):

                    st.session_state.trang_thai_duyet[
                        temp_id
                    ] = "Không lưu"

                    st.rerun()
# ==========================================================
# TẠO CÂU HỎI AI
# ==========================================================
def tao_cau_hoi_ai():

    st.header(
        "🧱 XÂY DỰNG NGÂN HÀNG ÔN TẬP / KIỂM TRA"
    )

    st.caption(
        "Mặc định nên dùng chế độ **Tự động**: app đọc độ phủ hiện có, "
        "tận dụng câu hạt giống và chỉ tạo thêm phần còn thiếu. "
        "Chế độ thủ công chỉ dùng khi GV muốn can thiệp vào một YCCĐ cụ thể."
    )

    tab_auto, tab_manual = st.tabs([
        "⚡ Xây dựng tự động",
        "✍️ Tạo theo YCCĐ (nâng cao)"
    ])

    # ======================================================
    # TAB 1 - TỰ ĐỘNG: LUỒNG CHÍNH
    # ======================================================
    with tab_auto:
        phan_tich_do_phu_va_tao_tu_dong_on_tap()

    # ======================================================
    # TAB 2 - THỦ CÔNG: CHỈ KHI GV CẦN
    # ======================================================
    with tab_manual:
        st.caption(
            "Dùng phần này khi GV đã chọn YCCĐ trong 📚 Kho YCCĐ "
            "và muốn chủ động số câu, mức độ, dạng câu hoặc nguồn."
        )

        st.caption(
            "Phần này dành cho trường hợp GV muốn chủ động chọn YCCĐ/cấu hình cụ thể. "
            "Nếu chỉ muốn app tự bù phần thiếu, dùng phần Độ phủ & Xây dựng tự động phía trên."
        )

        danh_sach_yccd = st.session_state.get(
            "yccd_da_chon",
            []
        )

        if not danh_sach_yccd:
            st.warning(
                "Hãy chọn YCCĐ trước."
            )
            return

        st.subheader(
            "🎯 Cấu hình đã chọn ở Kho YCCĐ"
        )

        st.caption(
            "Hệ thống sử dụng YCCĐ, số câu, mức độ, dạng câu hỏi và thành phần năng lực đã thiết lập. "
            "Nguồn được truy xuất linh hoạt phía sau để chọn căn cứ phù hợp nhất cho từng câu."
        )

        so_hat_giong_phu_hop = dem_hat_giong_phu_hop_yccd_da_chon()
        tong_hat_giong = len(doc_ngan_hang_hat_giong())

        if tong_hat_giong > 0:
            st.success(
                f"🌱 Có **{tong_hat_giong} câu hạt giống** trong kho; "
                f"app tìm thấy **{so_hat_giong_phu_hop} nhóm hạt giống phù hợp** "
                "với cấu hình hiện tại. Hệ thống sẽ kết hợp hạt giống + câu mới để tăng đa dạng."
            )
        else:
            st.info(
                "🌱 Chưa có câu hạt giống. Hệ thống vẫn có thể tạo câu mới, "
                "nhưng nhập hạt giống sẽ giúp ngân hàng phong phú và tiết kiệm API hơn."
            )

        tong_so_cau = 0

        for i, item in enumerate(
            danh_sach_yccd,
            start=1
        ):

            _, cau_hinh = lay_cau_hinh(item)

            so_cau = int(
                cau_hinh.get("Số câu", 1)
            )
            muc_do = cau_hinh.get(
                "Mức độ",
                xac_dinh_muc_do(item["YCCĐ"])
            )
            dang_cau = cau_hinh.get(
                "Dạng câu hỏi",
                "Trắc nghiệm 4 lựa chọn"
            )
            nang_luc = cau_hinh.get(
                "Thành phần năng lực",
                goi_y_thanh_phan_nang_luc(
                    item["YCCĐ"],
                    muc_do,
                    dang_cau
                )
            )
            nguon = cau_hinh.get(
                "Nguồn tham chiếu",
                []
            )
            tai_lieu = cau_hinh.get(
                "Tài liệu giáo viên",
                []
            )

            tong_so_cau += so_cau

            with st.container(border=True):

                st.markdown(
                    f"### YCCĐ {i}"
                )

                st.caption(
                    f"{item['Khối']} • "
                    f"{item['Chương']} • "
                    f"{item['Bài']}"
                )

                st.write(
                    f"**Yêu cầu cần đạt:** {item['YCCĐ']}"
                )

                c1, c2, c3, c4 = st.columns(4)

                with c1:
                    st.metric("Số câu", so_cau)

                with c2:
                    st.metric("Mức độ", muc_do)

                with c3:
                    st.markdown("**Dạng câu hỏi**")
                    st.write(dang_cau)

                with c4:
                    st.markdown("**Thành phần năng lực**")
                    st.write(nang_luc)

                st.markdown("**📚 Nguồn tham chiếu:**")
                if nguon:
                    for ten_nguon in nguon:
                        st.write(f"- {ten_nguon}")
                else:
                    st.write("- Không có")

                if "Tài liệu giáo viên tải lên" in nguon:
                    st.markdown("**📎 Tài liệu giáo viên đã chọn:**")
                    if tai_lieu:
                        for ten_file in tai_lieu:
                            st.write(f"- {ten_file}")
                    else:
                        st.write("- Chưa chọn tài liệu")

                if dang_cau == "Đúng / Sai":
                    chi_tiet_4_y = cau_hinh.get("Chi tiết 4 ý", {})

                    st.markdown("**🎯 Cấu hình 4 ý Đúng / Sai:**")

                    if chi_tiet_4_y:
                        for ky_hieu in ["a", "b", "c", "d"]:
                            meta = chi_tiet_4_y.get(ky_hieu, {})
                            yccd_y = meta.get("yccd", item["YCCĐ"])
                            muc_do_y = meta.get("muc_do", muc_do)

                            st.markdown(f"**Ý {ky_hieu})**")
                            st.write(f"YCCĐ: {yccd_y}")
                            nang_luc_y = meta.get(
                                "thanh_phan_nang_luc",
                                goi_y_thanh_phan_nang_luc(
                                    yccd_y,
                                    muc_do_y,
                                    "Đúng / Sai"
                                )
                            )
                            st.write(f"Mức độ: **{muc_do_y}**")
                            st.write(f"Thành phần năng lực: **{nang_luc_y}**")
                    else:
                        st.write(
                            "Chưa có cấu hình riêng 4 ý; hệ thống sẽ dùng YCCĐ "
                            "và mức độ chung của câu."
                        )

        st.success(
            f"Đã chọn **{len(danh_sach_yccd)} YCCĐ** "
            f"— tổng cộng **{tong_so_cau} câu hỏi**."
        )

        st.info(
            "Muốn đổi số câu, mức độ, dạng câu hoặc nguồn của YCCĐ nào, "
            "hãy quay lại 📚 Kho YCCĐ để chỉnh YCCĐ đó."
        )

        st.write("")

        _, cot_giua, _ = st.columns([1, 2, 1])

        with cot_giua:
            nut_tao_ai = st.button(
                "🚀 XÂY DỰNG CÂU HỎI",
                type="primary",
                use_container_width=True
            )

        if nut_tao_ai:

            with st.spinner(
                "Gemini đang tạo câu hỏi..."
            ):

                st.session_state.cau_hoi_ai = (
                    goi_ai_tao_cau_hoi()
                )

                st.session_state.trang_thai_duyet = {}

        if st.session_state.cau_hoi_ai:

            st.divider()

            st.header(
                "👩‍🏫 GIÁO VIÊN DUYỆT CÂU HỎI"
            )

            st.info(
                "Chỉ những câu bạn bấm "
                "✅ DUYỆT & LƯU "
                "mới được đưa vào Ngân hàng."
            )

            if st.button(
                "🛡 KIỂM ĐỊNH TẤT CẢ CÂU VỪA TẠO",
                type="secondary",
                use_container_width=True,
                key="kiem_dinh_cau_tao_thu_cong"
            ):
                with st.spinner(
                    "AI đang kiểm định YCCĐ, mức độ, "
                    "độ chính xác, đáp án và độ trùng..."
                ):
                    kq_kd = kiem_dinh_danh_sach_cau_hoi(
                        st.session_state.cau_hoi_ai
                    )

                    st.session_state.ket_qua_kiem_dinh.update(
                        kq_kd
                    )

                if kq_kd:
                    st.success(
                        f"Đã kiểm định **{len(kq_kd)} câu**."
                    )
                    st.rerun()

            for i, question in enumerate(
                st.session_state.cau_hoi_ai,
                start=1
            ):
                hien_thi_cau_ai(
                    question,
                    i
                )


# ==========================================================
# NGÂN HÀNG CÂU HỎI
# ==========================================================

def chuan_hoa_text_cov(value):
    return " ".join(str(value or "").strip().split())


def chuan_hoa_khoi_cov(value):
    s = chuan_hoa_text_cov(value)
    digits = "".join(ch for ch in s if ch.isdigit())
    if digits in {"10", "11", "12"}:
        return f"Khối {digits}"
    return s


def lay_khoi_cau_cov(q):
    return chuan_hoa_khoi_cov(
        q.get("khoi") or q.get("lop") or q.get("khối") or ""
    )


def lay_nang_luc_cau_cov(q):
    nl = chuan_hoa_text_cov(q.get("thanh_phan_nang_luc", ""))
    if nl:
        return nl
    return goi_y_thanh_phan_nang_luc(
        q.get("yccd", ""),
        q.get("muc_do", ""),
        q.get("dang_cau", "")
    )


def lay_nang_luc_y_ds_cov(nd):
    nl = chuan_hoa_text_cov(nd.get("thanh_phan_nang_luc", ""))
    if nl:
        return nl
    return goi_y_thanh_phan_nang_luc(
        nd.get("yccd", ""),
        nd.get("muc_do", ""),
        "Đúng / Sai"
    )











def cau_cu_khong_dat_chuan_moi(q):
    """
    Xác định câu cũ chưa đạt cấu trúc metadata của ngân hàng mới.

    Không đánh giá kiến thức bằng AI ở đây.
    Chỉ xác định các câu thiếu dữ liệu bắt buộc của thiết kế mới.
    """

    # Câu phải có các trường nền tảng.
    if not chuan_hoa_text_cov(q.get("yccd", "")):
        return True

    if not chuan_hoa_text_cov(q.get("muc_do", "")):
        return True

    if not chuan_hoa_text_cov(q.get("dang_cau", "")):
        return True

    # Thiếu thành phần năng lực cấp câu -> câu cũ không đạt chuẩn mới.
    if not chuan_hoa_text_cov(
        q.get("thanh_phan_nang_luc", "")
    ):
        return True

    dang = q.get("dang_cau", "")

    # Trắc nghiệm 4 lựa chọn phải đủ 4 phương án và đáp án A-D.
    if dang == "Trắc nghiệm 4 lựa chọn":
        lua_chon = q.get("lua_chon", []) or []

        if len(lua_chon) != 4:
            return True

        if str(q.get("dap_an", "")).strip().upper() not in [
            "A", "B", "C", "D"
        ]:
            return True

    # Đúng/Sai phải đủ 4 ý và mỗi ý có metadata mới.
    elif dang == "Đúng / Sai":
        meta = q.get("nhan_dinh_meta", []) or []

        if len(meta) != 4:
            return True

        for nd in meta:
            if not chuan_hoa_text_cov(nd.get("yccd", "")):
                return True

            if not chuan_hoa_text_cov(nd.get("muc_do", "")):
                return True

            if not chuan_hoa_text_cov(
                nd.get("thanh_phan_nang_luc", "")
            ):
                return True

            if str(nd.get("dap_an", "")).strip().lower() not in [
                "true", "false", "đúng", "sai"
            ]:
                return True

    # Trả lời ngắn: đáp án phải có.
    elif dang == "Trả lời ngắn":
        if not str(q.get("dap_an", "")).strip():
            return True

    return False




def kiem_tra_cau_moi_truoc_khi_luu(q):
    """Kiểm tra cứng trước khi câu được đưa vào ngân hàng chính thức."""
    loi = []

    for truong, ten in [
        ("yccd", "YCCĐ"),
        ("muc_do", "Mức độ"),
        ("dang_cau", "Dạng câu"),
        ("thanh_phan_nang_luc", "Thành phần năng lực"),
        ("cau_hoi", "Nội dung câu hỏi"),
    ]:
        if not str(q.get(truong, "")).strip():
            loi.append(f"Thiếu {ten}")

    dang = q.get("dang_cau", "")

    if dang == "Trắc nghiệm 4 lựa chọn":
        lc = q.get("lua_chon", []) or []
        if len(lc) != 4:
            loi.append("Trắc nghiệm phải đủ 4 lựa chọn A–D")
        if str(q.get("dap_an", "")).strip().upper() not in ["A", "B", "C", "D"]:
            loi.append("Đáp án trắc nghiệm phải là A/B/C/D")

    elif dang == "Đúng / Sai":
        meta = q.get("nhan_dinh_meta", []) or []
        if len(meta) != 4:
            loi.append("Đúng/Sai phải có đúng 4 ý a–d")
        else:
            for i, nd in enumerate(meta):
                ky = "abcd"[i]

                if not str(nd.get("noi_dung", "")).strip():
                    loi.append(
                        f"Ý {ky} thiếu nội dung nhận định"
                    )

                for field, label in [
                    ("yccd", "YCCĐ"),
                    ("muc_do", "mức độ"),
                    ("thanh_phan_nang_luc", "thành phần năng lực"),
                ]:
                    if not str(nd.get(field, "")).strip():
                        loi.append(f"Ý {ky} thiếu {label}")

                dap_y = chuan_hoa_dap_an_dung_sai(
                    nd.get("dap_an", "")
                )

                if dap_y not in ["Đúng", "Sai"]:
                    loi.append(
                        f"Ý {ky} thiếu đáp án Đúng/Sai"
                    )

                if not str(
                    nd.get("giai_thich", "")
                ).strip():
                    loi.append(
                        f"Ý {ky} thiếu giải thích"
                    )

    elif dang == "Trả lời ngắn":
        if not str(q.get("dap_an", "")).strip():
            loi.append("Trả lời ngắn thiếu đáp án")

    return loi


def thong_ke_sau_khi_luu(bank):
    """Thông tin nhanh để GV biết ngân hàng vừa được cập nhật."""
    return {
        "tong_cau": len(bank),
        "dang_4lc": sum(
            1 for q in bank
            if q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn"
        ),
        "dang_ds": sum(
            1 for q in bank
            if q.get("dang_cau") == "Đúng / Sai"
        ),
        "dang_ngan": sum(
            1 for q in bank
            if q.get("dang_cau") == "Trả lời ngắn"
        ),
    }




def kiem_tra_nguy_co_gan_gia_du_lieu(q):
    if q.get("kieu_du_lieu") != "mo_phong":
        return ""

    noi_dung = " ".join(
        str(q.get(k, ""))
        for k in ["tinh_huong", "cau_hoi", "giai_thich"]
    ).lower()

    cum_nguy_co = [
        "nghiên cứu cho thấy",
        "kết quả nghiên cứu",
        "theo nghiên cứu",
        "tại một nghiên cứu",
        "được khảo sát tại",
        "số liệu thực tế",
        "bài báo cho thấy"
    ]

    if any(cum in noi_dung for cum in cum_nguy_co):
        return (
            "Câu đang dùng dữ liệu mô phỏng nhưng diễn đạt như dữ liệu nghiên cứu thật."
        )

    return ""




def kiem_tra_cau_hoi_day_du_du_kien(q):
    """
    Kiểm tra sơ bộ câu hỏi có tự đủ dữ kiện hay không.
    Không gọi thêm AI.
    """
    cau_hoi = str(q.get("cau_hoi", "") or "")
    tinh_huong = str(q.get("tinh_huong", "") or "")
    noi_dung = (tinh_huong + "\n" + cau_hoi).strip().lower()

    loi = []

    cum_phu_thuoc_du_lieu = [
        "bảng dưới",
        "bảng sau",
        "hình bên",
        "hình dưới",
        "sơ đồ bên",
        "biểu đồ",
        "kết quả nghiên cứu",
        "dòng 1",
        "dòng 2",
        "dòng 3",
        "dòng 4",
        "dòng 5",
        "mẫu a",
        "mẫu b",
        "nhóm i",
        "nhóm ii",
        "quần thể i",
        "quần thể ii",
        "số liệu sau",
        "các số liệu"
    ]

    # Nếu câu dùng tham chiếu dữ liệu nhưng tổng độ dài quá ngắn,
    # khả năng cao là thiếu tình huống/bảng/sơ đồ.
    if any(cum in noi_dung for cum in cum_phu_thuoc_du_lieu):
        if len(noi_dung) < 280:
            loi.append(
                "Câu có tham chiếu đến dữ liệu/bảng/hình/đối tượng nhưng chưa trình bày đủ dữ kiện."
            )

    # Dòng 1..5 xuất hiện trong giải thích nhưng không xuất hiện đầy đủ trong đề
    giai_thich = str(q.get("giai_thich", "") or "").lower()
    for nhan in ["dòng 1", "dòng 2", "dòng 3", "dòng 4", "dòng 5"]:
        if nhan in giai_thich and nhan not in noi_dung:
            loi.append(
                "Dữ kiện quan trọng xuất hiện trong lời giải nhưng không có trong đề bài."
            )
            break

    dang = q.get("dang_cau", "")

    if dang == "Trắc nghiệm 4 lựa chọn":
        lc = q.get("lua_chon", []) or []
        if len(lc) != 4:
            loi.append("Câu 4 lựa chọn chưa đủ A-D.")

    if dang == "Đúng / Sai":
        nd = q.get("nhan_dinh_meta", []) or []

        if len(nd) != 4:
            loi.append("Câu Đúng/Sai chưa đủ 4 ý a, b, c, d.")
        else:
            for i_nd, item_nd in enumerate(nd):
                if not str(item_nd.get("noi_dung", "")).strip():
                    loi.append(
                        f"Câu Đúng/Sai thiếu nội dung ý {'abcd'[i_nd]}."
                    )

    return list(dict.fromkeys(loi))




AUTO_CHECKPOINT_PATH = "auto_checkpoint_ngan_hang.json"


def luu_auto_checkpoint(data):
    try:
        with open(
            AUTO_CHECKPOINT_PATH,
            "w",
            encoding="utf-8"
        ) as f:
            json.dump(
                data,
                f,
                ensure_ascii=False,
                indent=2
            )
        return True
    except Exception:
        return False


def doc_auto_checkpoint():
    try:
        if not os.path.exists(AUTO_CHECKPOINT_PATH):
            return None

        with open(
            AUTO_CHECKPOINT_PATH,
            "r",
            encoding="utf-8"
        ) as f:
            data = json.load(f)

        if isinstance(data, dict):
            return data
    except Exception:
        pass

    return None


def xoa_auto_checkpoint():
    try:
        if os.path.exists(AUTO_CHECKPOINT_PATH):
            os.remove(AUTO_CHECKPOINT_PATH)
    except Exception:
        pass


def tao_auto_checkpoint(
    cac_batch,
    batch_size,
    pham_vi,
    tong_ke_hoach
):
    return {
        "trang_thai": "dang_chay",
        "batch_tiep_theo": 0,
        "tong_batch": len(cac_batch),
        "batch_size": int(batch_size),
        "pham_vi": pham_vi,
        "tong_ke_hoach": int(tong_ke_hoach),
        "da_tao": 0,
        "cac_batch": cac_batch,
        "cau_da_tao": [],
        "loi_cuoi": ""
    }


def khoa_cau_checkpoint(q):
    return str(
        q.get("temp_id", "")
        or q.get("id", "")
        or hashlib.md5(
            (
                str(q.get("cau_hoi", ""))
                + "|"
                + str(q.get("yccd", ""))
                + "|"
                + str(q.get("dang_cau", ""))
            ).encode("utf-8")
        ).hexdigest()
    )


def khoi_phuc_cau_checkpoint_vao_cho_duyet(
    checkpoint,
    ds_cho_duyet
):
    """
    Khôi phục câu đã tạo từ checkpoint mà không tạo bản sao trùng.
    """
    da_co = {
        khoa_cau_checkpoint(q)
        for q in ds_cho_duyet
    }

    so_khoi_phuc = 0

    for q in checkpoint.get("cau_da_tao", []) or []:
        key = khoa_cau_checkpoint(q)

        if key in da_co:
            continue

        ds_cho_duyet.append(q)
        da_co.add(key)
        so_khoi_phuc += 1

    return so_khoi_phuc




def phan_tich_do_phu_va_tao_tu_dong_on_tap():
    bank = doc_ngan_hang()

    st.subheader("⚡ XÂY DỰNG TỰ ĐỘNG THEO ĐỘ PHỦ")
    st.caption(
        "Phân tích câu đã có, xác định phần còn thiếu và tự động bổ sung. "
        "Hạt giống đã kiểm tra an toàn được đưa vào ngân hàng trước; phần Xây dựng chỉ tạo phần còn thiếu."
    )

    if not bank:
        st.info(
            "Ngân hàng chuẩn hiện chưa có câu. Bạn vẫn có thể chọn phạm vi và "
            "đặt mục tiêu để hệ thống bắt đầu xây dựng từ hạt giống + câu mới."
        )

    with st.expander("📊 Xem độ phủ và phần còn thiếu", expanded=False):

        st.caption(
            "Phân tích số câu hiện có theo từng YCCĐ và mức độ "
            "Nhận biết / Thông hiểu / Vận dụng."
        )

        cov_col1, cov_col2, cov_col3 = st.columns(3)

        ds_khoi_cov = list(KHO_YCCD.keys())

        with cov_col1:
            khoi_cov = st.selectbox(
                "🎓 Khối phân tích",
                ds_khoi_cov,
                key="coverage_khoi"
            )

        ds_chuong_cov = list(
            KHO_YCCD.get(khoi_cov, {}).keys()
        )

        with cov_col2:
            chuong_cov = st.selectbox(
                "📚 Chương",
                ds_chuong_cov,
                key="coverage_chuong"
            )

        ds_bai_cov = list(
            KHO_YCCD.get(khoi_cov, {})
            .get(chuong_cov, {})
            .keys()
        )

        with cov_col3:
            bai_cov = st.selectbox(
                "📖 Bài",
                ["Tất cả"] + ds_bai_cov,
                key="coverage_bai"
            )

        muc_tieu_moi_muc = st.number_input(
            "🎯 Mục tiêu số câu cho mỗi YCCĐ theo cấu hình đã gán",
            min_value=1,
            max_value=100,
            value=10,
            step=1,
            key="coverage_target"
        )

        st.caption(
            "App kiểm tra từng YCCĐ theo đúng mức độ và thành phần năng lực đã gán, "
            "sau đó chỉ bổ sung những nhóm còn thiếu."
        )

        # --------------------------------------------------
        # LẤY TOÀN BỘ YCCĐ CẦN PHÂN TÍCH TỪ KHO YCCĐ
        # --------------------------------------------------
        danh_sach_yccd_cov = []

        if bai_cov == "Tất cả":
            for ten_bai_cov, ds_yccd_cov in (
                KHO_YCCD.get(khoi_cov, {})
                .get(chuong_cov, {})
                .items()
            ):
                for yccd_cov in ds_yccd_cov:
                    danh_sach_yccd_cov.append(
                        {
                            "bai": ten_bai_cov,
                            "yccd": yccd_cov
                        }
                    )
        else:
            for yccd_cov in (
                KHO_YCCD.get(khoi_cov, {})
                .get(chuong_cov, {})
                .get(bai_cov, [])
            ):
                danh_sach_yccd_cov.append(
                    {
                        "bai": bai_cov,
                        "yccd": yccd_cov
                    }
                )

        # Chỉ tính câu còn sử dụng trong ngân hàng
        bank_cov = [
            q for q in bank
            if lay_khoi_cau_cov(q) == chuan_hoa_khoi_cov(khoi_cov)
            and chuan_hoa_text_cov(q.get("chuong", "")) == chuan_hoa_text_cov(chuong_cov)
            and q.get("trang_thai", "Đã duyệt") != "Ngừng sử dụng"
        ]

        if bai_cov != "Tất cả":
            bank_cov = [
                q for q in bank_cov
                if chuan_hoa_text_cov(q.get("bai", ""))
                == chuan_hoa_text_cov(bai_cov)
            ]

        def dem_cau_cov(yccd_text, muc_do_text):
            dem = 0

            for q in bank_cov:

                # Câu thường: 1 câu ↔ 1 YCCĐ ↔ 1 mức độ
                if q.get("dang_cau") != "Đúng / Sai":
                    if (
                        q.get("yccd") == yccd_text
                        and q.get("muc_do") == muc_do_text
                    ):
                        dem += 1

                # Đúng/Sai: tính 1 câu tình huống nếu có ít nhất
                # một ý a-d đúng YCCĐ + mức độ đang xét.
                else:
                    nhan_dinh_meta = q.get(
                        "nhan_dinh_meta",
                        []
                    )

                    khop = any(
                        nd.get("yccd") == yccd_text
                        and nd.get("muc_do") == muc_do_text
                        for nd in nhan_dinh_meta
                    )

                    if khop:
                        dem += 1

            return dem

        def dem_cau_cov_nang_luc(
            yccd_text,
            muc_do_text,
            nang_luc_text
        ):
            dem = 0

            for q in bank_cov:

                if q.get("dang_cau") != "Đúng / Sai":
                    if (
                        chuan_hoa_text_cov(q.get("yccd")) == chuan_hoa_text_cov(yccd_text)
                        and chuan_hoa_text_cov(q.get("muc_do")) == chuan_hoa_text_cov(muc_do_text)
                        and chuan_hoa_text_cov(lay_nang_luc_cau_cov(q))
                        == chuan_hoa_text_cov(nang_luc_text)
                    ):
                        dem += 1

                else:
                    # Với Đúng/Sai, mỗi ý a-d có metadata riêng.
                    for nd in q.get("nhan_dinh_meta", []) or []:
                        if (
                            chuan_hoa_text_cov(nd.get("yccd")) == chuan_hoa_text_cov(yccd_text)
                            and chuan_hoa_text_cov(nd.get("muc_do")) == chuan_hoa_text_cov(muc_do_text)
                            and chuan_hoa_text_cov(lay_nang_luc_y_ds_cov(nd))
                            == chuan_hoa_text_cov(nang_luc_text)
                        ):
                            dem += 1

            return dem

        def nang_luc_muc_tieu_cov(
            yccd_text,
            muc_do_text
        ):
            return goi_y_thanh_phan_nang_luc(
                yccd_text,
                muc_do_text,
                "Trắc nghiệm 4 lựa chọn"
            )

        def trang_thai_cov(so_cau):
            muc_tieu = int(muc_tieu_moi_muc)

            if so_cau >= muc_tieu:
                return "🟢 Đủ"

            if so_cau >= max(1, int(muc_tieu * 0.8)):
                return "🟡 Sắp đủ"

            return "🔴 Thiếu"

        bang_do_phu = []
        tong_hien_co = 0
        tong_can_co = 0
        tong_thieu = 0

        for idx_cov, item_cov in enumerate(
            danh_sach_yccd_cov,
            start=1
        ):

            yccd_text = item_cov["yccd"]
            ten_bai_cov = item_cov["bai"]

            muc_do_duoc_gan = xac_dinh_muc_do(
                yccd_text
            )

            nang_luc_muc_tieu = nang_luc_muc_tieu_cov(
                yccd_text,
                muc_do_duoc_gan
            )

            nb = dem_cau_cov(
                yccd_text,
                "Nhận biết"
            )

            th = dem_cau_cov(
                yccd_text,
                "Thông hiểu"
            )

            vd = dem_cau_cov(
                yccd_text,
                "Vận dụng"
            )

            tong_dong = nb + th + vd

            dem_muc_duoc_gan = dem_cau_cov_nang_luc(
                yccd_text,
                muc_do_duoc_gan,
                nang_luc_muc_tieu
            )

            can_dong = int(muc_tieu_moi_muc)

            thieu_dong = max(
                0,
                int(muc_tieu_moi_muc)
                - dem_muc_duoc_gan
            )

            tong_hien_co += tong_dong
            tong_can_co += can_dong
            tong_thieu += thieu_dong

            bang_do_phu.append(
                {
                    "STT": idx_cov,
                    "Bài": ten_bai_cov,
                    "YCCĐ": yccd_text,
                    "Mức độ được gán": muc_do_duoc_gan,
                    "Thành phần năng lực mục tiêu": nang_luc_muc_tieu,
                    "Đã có đúng mục tiêu": dem_muc_duoc_gan,
                    "Nhận biết": nb,
                    "NB trạng thái": trang_thai_cov(nb),
                    "Thông hiểu": th,
                    "TH trạng thái": trang_thai_cov(th),
                    "Vận dụng": vd,
                    "VD trạng thái": trang_thai_cov(vd),
                    "Tổng": tong_dong,
                    "Còn thiếu": thieu_dong
                }
            )

        if not bang_do_phu:
            st.info(
                "Không có YCCĐ để phân tích trong phạm vi đã chọn."
            )
        else:

            so_cau_thuc_te_cov = len(bank_cov)
            da_phu_dung_cov = sum(
                row["Đã có đúng mục tiêu"]
                for row in bang_do_phu
            )

            m1, m2, m3, m4 = st.columns(4)

            with m1:
                st.metric("Câu hiện có", so_cau_thuc_te_cov)

            with m2:
                st.metric("Đã phủ đúng cấu hình", da_phu_dung_cov)

            with m3:
                st.metric("Mục tiêu", tong_can_co)

            with m4:
                st.metric("Còn thiếu", tong_thieu)

            st.caption(
                "Câu cũ chưa có trường thành phần năng lực vẫn được suy ra "
                "từ YCCĐ + mức độ để tính độ phủ."
            )

            df_do_phu = pd.DataFrame(bang_do_phu)

            if not df_do_phu.empty:
                for ten_bai_hien in list(dict.fromkeys(df_do_phu["Bài"].tolist())):
                    df_bai = df_do_phu[df_do_phu["Bài"] == ten_bai_hien].copy()

                    so_yccd_bai = len(df_bai)
                    tong_da_co_bai = int(df_bai["Đã có đúng mục tiêu"].sum())
                    tong_thieu_bai = int(df_bai["Còn thiếu"].sum())

                    with st.expander(
                        f"📖 {ten_bai_hien} — {so_yccd_bai} YCCĐ — "
                        f"đã có {tong_da_co_bai} — thiếu {tong_thieu_bai}",
                        expanded=False
                    ):
                        df_hien = df_bai[
                            [
                                "YCCĐ",
                                "Mức độ được gán",
                                "Thành phần năng lực mục tiêu",
                                "Đã có đúng mục tiêu",
                                "Còn thiếu"
                            ]
                        ].copy()

                        df_hien = df_hien.rename(
                            columns={
                                "Mức độ được gán": "Mức độ",
                                "Thành phần năng lực mục tiêu": "Thành phần năng lực",
                                "Đã có đúng mục tiêu": "Đã có",
                                "Còn thiếu": "Thiếu"
                            }
                        )

                        df_hien["Mục tiêu"] = int(muc_tieu_moi_muc)
                        df_hien = df_hien[
                            [
                                "YCCĐ",
                                "Mức độ",
                                "Thành phần năng lực",
                                "Đã có",
                                "Mục tiêu",
                                "Thiếu"
                            ]
                        ]

                        st.dataframe(
                            df_hien,
                            use_container_width=True,
                            hide_index=True,
                            column_config={
                                "YCCĐ": st.column_config.TextColumn(
                                    "YCCĐ",
                                    width="large"
                                ),
                                "Thành phần năng lực": st.column_config.TextColumn(
                                    "Thành phần năng lực",
                                    width="medium"
                                )
                            }
                        )

            df_cov_nl = pd.DataFrame(
                bang_do_phu
            )

            if not df_cov_nl.empty:
                bang_nl = (
                    df_cov_nl.groupby(
                        "Thành phần năng lực mục tiêu",
                        dropna=False
                    )[[
                        "Đã có đúng mục tiêu",
                        "Còn thiếu"
                    ]]
                    .sum()
                    .reset_index()
                )

                st.markdown(
                    "#### 🧠 ĐỘ PHỦ THEO THÀNH PHẦN NĂNG LỰC"
                )

                st.dataframe(
                    bang_nl,
                    use_container_width=True,
                    hide_index=True
                )

            if tong_thieu == 0:
                st.success(
                    "🟢 Ngân hàng đã đạt mục tiêu độ phủ "
                    "trong phạm vi đang chọn."
                )
            else:
                st.warning(
                    f"⚠️ Còn thiếu **{tong_thieu} câu** "
                    "để đạt mục tiêu độ phủ."
                )

                st.info(
                    "Bước sau sẽ dùng kết quả này để chức năng "
                    "**⚡ Xây ngân hàng tự động** chỉ tạo đúng "
                    "những ô YCCĐ × mức độ đã gán × thành phần năng lực "
                    "còn thiếu. Với Đúng/Sai, hệ thống theo dõi từng ý a–d."
                )

                st.markdown("---")
                st.subheader("⚡ XÂY NGÂN HÀNG TỰ ĐỘNG")

                ds_bai_thieu_auto = [
                    ten_bai
                    for ten_bai in list(dict.fromkeys(
                        item["bai"] for item in danh_sach_yccd_cov
                    ))
                    if any(
                        row["Bài"] == ten_bai and row["Còn thiếu"] > 0
                        for row in bang_do_phu
                    )
                ]

                pham_vi_xay_auto = st.selectbox(
                    "📖 Phạm vi xây tự động",
                    ["Tất cả bài đang thiếu"] + ds_bai_thieu_auto,
                    key="auto_build_scope_bai"
                )

                st.caption(
                    "Hệ thống chỉ tạo phần còn thiếu sau khi đã tính cả các câu hạt giống an toàn đã đồng bộ vào ngân hàng. "
                    "Với câu Đúng/Sai, độ phủ được theo dõi theo từng ý a–d "
                    "vì mỗi ý có YCCĐ, mức độ và thành phần năng lực riêng. "
                    "Câu AI tạo ra sẽ ở khu vực chờ GV duyệt, chưa tự động "
                    "đưa vào ngân hàng chính thức."
                )

                # ==========================================
                # CƠ CẤU DẠNG CÂU MỤC TIÊU
                # ==========================================
                st.markdown("**Cơ cấu dạng câu mục tiêu**")

                ty1, ty2, ty3 = st.columns(3)

                with ty1:
                    ti_le_tn = st.number_input(
                        "4 lựa chọn (%)",
                        min_value=0,
                        max_value=100,
                        value=50,
                        step=5,
                        key="auto_ratio_mcq"
                    )

                with ty2:
                    ti_le_ds = st.number_input(
                        "Đúng / Sai (%)",
                        min_value=0,
                        max_value=100,
                        value=30,
                        step=5,
                        key="auto_ratio_tf"
                    )

                with ty3:
                    ti_le_tln = st.number_input(
                        "Trả lời ngắn (%)",
                        min_value=0,
                        max_value=100,
                        value=20,
                        step=5,
                        key="auto_ratio_short"
                    )

                tong_ti_le = (
                    int(ti_le_tn)
                    + int(ti_le_ds)
                    + int(ti_le_tln)
                )

                if tong_ti_le != 100:
                    st.error(
                        f"Tổng tỉ lệ hiện là **{tong_ti_le}%**. "
                        "Cần chỉnh về đúng **100%**."
                    )

                # ==========================================
                # HÀM PHÂN BỔ SỐ NGUYÊN THEO TỈ LỆ
                # ==========================================
                def phan_bo_so_luong(tong_so, ti_le_map):

                    if tong_so <= 0:
                        return {
                            ten: 0
                            for ten in ti_le_map
                        }

                    raw = {
                        ten: tong_so * ti_le / 100
                        for ten, ti_le in ti_le_map.items()
                    }

                    ket_qua = {
                        ten: int(raw[ten])
                        for ten in raw
                    }

                    con_lai = (
                        tong_so
                        - sum(ket_qua.values())
                    )

                    thu_tu = sorted(
                        raw.keys(),
                        key=lambda ten: (
                            raw[ten] - int(raw[ten])
                        ),
                        reverse=True
                    )

                    for ten in thu_tu[:con_lai]:
                        ket_qua[ten] += 1

                    return ket_qua

                # ==========================================
                # ĐẾM SỐ CÂU HIỆN CÓ THEO DẠNG
                # ==========================================
                def dem_theo_dang(
                    yccd_text,
                    muc_do_text,
                    dang_text
                ):

                    dem = 0

                    for q in bank_cov:

                        if q.get(
                            "trang_thai",
                            "Đã duyệt"
                        ) == "Ngừng sử dụng":
                            continue

                        # Câu thường
                        if dang_text != "Đúng / Sai":

                            if (
                                q.get("dang_cau") == dang_text
                                and q.get("yccd") == yccd_text
                                and q.get("muc_do") == muc_do_text
                            ):
                                dem += 1

                        # Đúng/Sai tính theo câu tình huống
                        else:

                            if q.get("dang_cau") != "Đúng / Sai":
                                continue

                            nhan_dinh_meta = q.get(
                                "nhan_dinh_meta",
                                []
                            )

                            if any(
                                nd.get("yccd") == yccd_text
                                and nd.get("muc_do") == muc_do_text
                                for nd in nhan_dinh_meta
                            ):
                                dem += 1

                    return dem

                ti_le_map = {
                    "Trắc nghiệm 4 lựa chọn": int(ti_le_tn),
                    "Đúng / Sai": int(ti_le_ds),
                    "Trả lời ngắn": int(ti_le_tln)
                }

                ke_hoach_auto = []

                if tong_ti_le == 100:

                    for item_cov in danh_sach_yccd_cov:

                        yccd_text = item_cov["yccd"]
                        ten_bai = item_cov["bai"]

                        for muc_do_text in [
                            xac_dinh_muc_do(yccd_text)
                        ]:

                            nang_luc_text = nang_luc_muc_tieu_cov(
                                yccd_text,
                                muc_do_text
                            )

                            tong_hien_tai_o = dem_cau_cov_nang_luc(
                                yccd_text,
                                muc_do_text,
                                nang_luc_text
                            )

                            tong_can_them_o = max(
                                0,
                                int(muc_tieu_moi_muc)
                                - tong_hien_tai_o
                            )

                            if tong_can_them_o <= 0:
                                continue

                            muc_tieu_theo_dang = (
                                phan_bo_so_luong(
                                    int(muc_tieu_moi_muc),
                                    ti_le_map
                                )
                            )

                            # Ưu tiên bù đúng dạng còn thiếu.
                            thieu_theo_dang = {}

                            for dang_text in ti_le_map:

                                hien_co_dang = dem_theo_dang(
                                    yccd_text,
                                    muc_do_text,
                                    dang_text
                                )

                                thieu_theo_dang[
                                    dang_text
                                ] = max(
                                    0,
                                    muc_tieu_theo_dang[
                                        dang_text
                                    ] - hien_co_dang
                                )

                            # Do làm tròn hoặc cấu trúc dữ liệu cũ,
                            # tổng thiếu theo dạng có thể khác số còn thiếu
                            # của ô. Chốt về đúng tong_can_them_o.
                            tong_thieu_dang = sum(
                                thieu_theo_dang.values()
                            )

                            if tong_thieu_dang < tong_can_them_o:

                                # Bù lần lượt theo tỉ lệ ưu tiên cao hơn.
                                thu_tu_dang = sorted(
                                    ti_le_map.keys(),
                                    key=lambda d: ti_le_map[d],
                                    reverse=True
                                )

                                con = (
                                    tong_can_them_o
                                    - tong_thieu_dang
                                )

                                idx_bu = 0

                                while con > 0:
                                    dang_bu = thu_tu_dang[
                                        idx_bu
                                        % len(thu_tu_dang)
                                    ]
                                    thieu_theo_dang[
                                        dang_bu
                                    ] += 1
                                    idx_bu += 1
                                    con -= 1

                            elif tong_thieu_dang > tong_can_them_o:

                                # Cắt bớt từ dạng có tỉ lệ thấp hơn.
                                can_cat = (
                                    tong_thieu_dang
                                    - tong_can_them_o
                                )

                                thu_tu_cat = sorted(
                                    ti_le_map.keys(),
                                    key=lambda d: ti_le_map[d]
                                )

                                for dang_cat in thu_tu_cat:

                                    if can_cat <= 0:
                                        break

                                    cat = min(
                                        can_cat,
                                        thieu_theo_dang[
                                            dang_cat
                                        ]
                                    )

                                    thieu_theo_dang[
                                        dang_cat
                                    ] -= cat

                                    can_cat -= cat

                            for dang_text, so_can_tao in (
                                thieu_theo_dang.items()
                            ):

                                if so_can_tao <= 0:
                                    continue

                                ke_hoach_auto.append(
                                    {
                                        "Khối": khoi_cov,
                                        "Chương": chuong_cov,
                                        "Bài": ten_bai,
                                        "YCCĐ": yccd_text,
                                        "Mức độ": muc_do_text,
                                        "Thành phần năng lực": nang_luc_text,
                                        "Dạng câu hỏi": dang_text,
                                        "Số câu": int(so_can_tao)
                                    }
                                )

                # Ưu tiên tạo các ô còn thiếu nhiều trước.
                ke_hoach_auto.sort(
                    key=lambda item: item["Số câu"],
                    reverse=True
                )

                if pham_vi_xay_auto != "Tất cả bài đang thiếu":
                    ke_hoach_auto = [
                        item
                        for item in ke_hoach_auto
                        if item["Bài"] == pham_vi_xay_auto
                    ]

                tong_ke_hoach = sum(
                    item["Số câu"]
                    for item in ke_hoach_auto
                )

                if tong_ti_le == 100:

                    tong_don_vi_danh_gia_ke_hoach = sum(
                        (
                            item["Số câu"] * 4
                            if item["Dạng câu hỏi"] == "Đúng / Sai"
                            else item["Số câu"]
                        )
                        for item in ke_hoach_auto
                    )

                    st.write(
                        f"📌 Phạm vi: **{pham_vi_xay_auto}** — "
                        f"cần tạo thêm **{tong_ke_hoach} câu**."
                    )

                    st.caption(
                        f"Quy đổi tương đương "
                        f"**{tong_don_vi_danh_gia_ke_hoach} đơn vị đánh giá** "
                        "(mỗi câu Đúng/Sai = 4 nhận định)."
                    )

                    if ke_hoach_auto:

                        # ==================================================
                        # BẢNG KẾ HOẠCH GỘP THEO YCCĐ
                        # ==================================================
                        ke_hoach_gop = {}

                        for item_auto in ke_hoach_auto:

                            key_gop = (
                                item_auto["Bài"],
                                item_auto["YCCĐ"],
                                item_auto["Mức độ"],
                                item_auto["Thành phần năng lực"]
                            )

                            if key_gop not in ke_hoach_gop:
                                ke_hoach_gop[key_gop] = {
                                    "Bài": item_auto["Bài"],
                                    "YCCĐ": item_auto["YCCĐ"],
                                    "Mức độ": item_auto["Mức độ"],
                                    "Thành phần năng lực": item_auto[
                                        "Thành phần năng lực"
                                    ],
                                    "Trắc nghiệm 4 lựa chọn": 0,
                                    "Đúng / Sai": 0,
                                    "Trả lời ngắn": 0
                                }

                            dang_auto = item_auto["Dạng câu hỏi"]

                            if dang_auto in ke_hoach_gop[key_gop]:
                                ke_hoach_gop[key_gop][
                                    dang_auto
                                ] += int(
                                    item_auto["Số câu"]
                                )

                        bang_ke_hoach = []

                        ds_gop_sap_xep = sorted(
                            ke_hoach_gop.values(),
                            key=lambda x: (
                                str(x["Bài"]),
                                str(x["YCCĐ"]),
                                str(x["Mức độ"]),
                                str(x["Thành phần năng lực"])
                            )
                        )

                        for stt_auto, item_gop in enumerate(
                            ds_gop_sap_xep,
                            start=1
                        ):

                            so_tn = int(
                                item_gop[
                                    "Trắc nghiệm 4 lựa chọn"
                                ]
                            )

                            so_ds = int(
                                item_gop[
                                    "Đúng / Sai"
                                ]
                            )

                            so_ngan = int(
                                item_gop[
                                    "Trả lời ngắn"
                                ]
                            )

                            tong_cau = (
                                so_tn
                                + so_ds
                                + so_ngan
                            )

                            tong_don_vi_danh_gia = (
                                so_tn
                                + so_ngan
                                + so_ds * 4
                            )

                            bang_ke_hoach.append(
                                {
                                    "STT": stt_auto,
                                    "Bài": item_gop["Bài"],
                                    "YCCĐ": item_gop["YCCĐ"],
                                    "Mức độ": item_gop["Mức độ"],
                                    "Thành phần năng lực": item_gop[
                                        "Thành phần năng lực"
                                    ],
                                    "4 lựa chọn": so_tn,
                                    "Đúng / Sai": so_ds,
                                    "Trả lời ngắn": so_ngan,
                                    "Tổng câu": tong_cau,
                                    "Đơn vị đánh giá": tong_don_vi_danh_gia
                                }
                            )

                        with st.expander(
                            "📋 Xem kế hoạch chi tiết",
                            expanded=False
                        ):

                            st.caption(
                                "Mỗi YCCĐ chỉ hiển thị một dòng. "
                                "Một câu Đúng/Sai gồm 4 nhận định nên "
                                "được quy đổi thành 4 đơn vị đánh giá."
                            )

                            df_ke_hoach_gop = pd.DataFrame(bang_ke_hoach)

                            if not df_ke_hoach_gop.empty:
                                for ten_bai_kh in list(dict.fromkeys(
                                    df_ke_hoach_gop["Bài"].tolist()
                                )):
                                    st.markdown(f"#### 📖 {ten_bai_kh}")

                                    df_kh_bai = df_ke_hoach_gop[
                                        df_ke_hoach_gop["Bài"] == ten_bai_kh
                                    ].copy()

                                    if "Bài" in df_kh_bai.columns:
                                        df_kh_bai = df_kh_bai.drop(
                                            columns=["Bài"]
                                        )

                                    st.dataframe(
                                        df_kh_bai,
                                        use_container_width=True,
                                        hide_index=True,
                                        column_config={
                                            "YCCĐ": st.column_config.TextColumn(
                                                "YCCĐ",
                                                width="large"
                                            ),
                                            "Thành phần năng lực": st.column_config.TextColumn(
                                                "Thành phần năng lực",
                                                width="medium"
                                            )
                                        }
                                    )

                        st.markdown(
                            "**Cách gọi Gemini:** hệ thống chia nhỏ "
                            "5–10 câu/lượt để tránh prompt quá dài."
                        )

                        st.info(
                            "Hệ thống chỉ tạo đúng phần còn thiếu theo "
                            "**YCCĐ + mức độ đã gán + thành phần năng lực + dạng câu**. "
                            "Câu AI tạo ra vẫn nằm ở khu vực **chờ GV duyệt**, "
                            "chưa tự động đưa vào ngân hàng chính thức."
                        )

                        ctl1, ctl2 = st.columns(2)

                        with ctl1:
                            batch_size_auto = st.number_input(
                                "Số câu mỗi lượt",
                                min_value=5,
                                max_value=10,
                                value=5,
                                step=1,
                                key="auto_batch_size"
                            )

                        with ctl2:
                            gioi_han_auto = st.number_input(
                                "Tối đa tạo trong lần bấm này",
                                min_value=5,
                                max_value=50,
                                value=min(
                                    20,
                                    max(5, tong_ke_hoach)
                                ),
                                step=5,
                                key="auto_max_generate"
                            )

                        tu_dong_den_du = st.checkbox(
                            "⚡ Tự động tạo toàn bộ phần còn thiếu theo chỉ tiêu",
                            value=False,
                            key="auto_build_until_target"
                        )

                        if tu_dong_den_du:
                            st.caption(
                                f"Hệ thống sẽ tự chia **{int(tong_ke_hoach)} câu còn thiếu** "
                                f"thành các batch **{int(batch_size_auto)} câu/lượt** và chạy lần lượt. "
                                "Tất cả câu tạo ra vẫn chỉ vào khu vực **chờ GV duyệt**, "
                                "không tự động lưu vào ngân hàng."
                            )
                        else:
                            st.caption(
                                f"Lần bấm này tối đa tạo **{int(gioi_han_auto)} câu** "
                                f"trên tổng **{int(tong_ke_hoach)} câu còn thiếu**. "
                                "Có thể bấm nhiều lần; sau khi duyệt và lưu, độ phủ "
                                "sẽ được tính lại để lần sau chỉ bù phần còn thiếu mới."
                            )

                        if "auto_cau_cho_duyet" not in st.session_state:
                            st.session_state.auto_cau_cho_duyet = []

                        st.caption(
                            "Quy trình: **AI tạo → GV xem/duyệt → kiểm tra cấu trúc cứng "
                            "→ lưu ngân hàng → độ phủ tự tính lại từ dữ liệu đã lưu**."
                        )

                        if "auto_tien_do" not in st.session_state:
                            st.session_state.auto_tien_do = {
                                "da_tao": 0,
                                "tong_ke_hoach": 0
                            }

                        checkpoint_auto = doc_auto_checkpoint()

                        if checkpoint_auto:
                            khoi_phuc_cau_checkpoint_vao_cho_duyet(
                                checkpoint_auto,
                                st.session_state.auto_cau_cho_duyet
                            )

                            if checkpoint_auto.get("trang_thai") == "tam_dung":
                                st.warning(
                                    f"⏸ Phiên trước đã tạm dừng sau "
                                    f"**{checkpoint_auto.get('da_tao', 0)} câu** "
                                    f"(batch {checkpoint_auto.get('batch_tiep_theo', 0)}/"
                                    f"{checkpoint_auto.get('tong_batch', 0)})."
                                )

                                cp1, cp2 = st.columns(2)

                                with cp1:
                                    resume_auto = st.button(
                                        "▶ TIẾP TỤC PHIÊN TRƯỚC",
                                        use_container_width=True,
                                        key="resume_auto_checkpoint"
                                    )

                                with cp2:
                                    huy_checkpoint = st.button(
                                        "🗑 HỦY PHIÊN TRƯỚC",
                                        use_container_width=True,
                                        key="cancel_auto_checkpoint"
                                    )

                                if huy_checkpoint:
                                    xoa_auto_checkpoint()
                                    st.rerun()

                            elif checkpoint_auto.get("trang_thai") == "dang_chay":
                                st.info(
                                    f"💾 Đã có checkpoint: "
                                    f"{checkpoint_auto.get('da_tao', 0)} câu đã tạo."
                                )

                        nhan_nut_auto = (
                            "⚡ TỰ ĐỘNG TẠO TOÀN BỘ PHẦN CÒN THIẾU"
                            if tu_dong_den_du
                            else "⚡ TẠO PHẦN CÒN THIẾU"
                        )

                        bat_dau_auto_moi = st.button(
                            nhan_nut_auto,
                            type="primary",
                            use_container_width=True,
                            key="btn_auto_build_bank"
                        )

                        dang_resume_auto = (
                            "resume_auto" in locals()
                            and resume_auto
                        )

                        if bat_dau_auto_moi or dang_resume_auto:

                            # ----------------------------------
                            # Mở rộng kế hoạch thành các đơn vị
                            # rồi chia batch tối đa 5–10 câu.
                            # ----------------------------------
                            if dang_resume_auto:

                                checkpoint_run = doc_auto_checkpoint()

                                if not checkpoint_run:
                                    st.error(
                                        "Không tìm thấy checkpoint để tiếp tục."
                                    )
                                    st.stop()

                                cac_batch = checkpoint_run.get(
                                    "cac_batch",
                                    []
                                )

                                chi_so_batch_bat_dau = int(
                                    checkpoint_run.get(
                                        "batch_tiep_theo",
                                        0
                                    )
                                )

                            else:

                                if tu_dong_den_du:
                                    so_duoc_phep = int(
                                        tong_ke_hoach
                                    )
                                else:
                                    so_duoc_phep = min(
                                        int(gioi_han_auto),
                                        int(tong_ke_hoach)
                                    )

                                specs_con_lai = []
                                da_lay = 0

                                for spec in ke_hoach_auto:

                                    if da_lay >= so_duoc_phep:
                                        break

                                    so_lay = min(
                                        int(spec["Số câu"]),
                                        so_duoc_phep - da_lay
                                    )

                                    if so_lay > 0:
                                        spec_copy = dict(spec)
                                        spec_copy["Số câu"] = so_lay
                                        specs_con_lai.append(
                                            spec_copy
                                        )
                                        da_lay += so_lay

                                cac_batch = []
                                batch_hien_tai = []
                                tong_batch = 0

                                for spec in specs_con_lai:

                                    so_con = int(
                                        spec["Số câu"]
                                    )

                                    while so_con > 0:

                                        cho_trong = (
                                            int(batch_size_auto)
                                            - tong_batch
                                        )

                                        if cho_trong <= 0:
                                            cac_batch.append(
                                                batch_hien_tai
                                            )
                                            batch_hien_tai = []
                                            tong_batch = 0
                                            cho_trong = int(
                                                batch_size_auto
                                            )

                                        so_cho_batch = min(
                                            so_con,
                                            cho_trong
                                        )

                                        x = dict(spec)
                                        x["Số câu"] = so_cho_batch

                                        batch_hien_tai.append(x)
                                        tong_batch += so_cho_batch
                                        so_con -= so_cho_batch

                                if batch_hien_tai:
                                    cac_batch.append(
                                        batch_hien_tai
                                    )

                                chi_so_batch_bat_dau = 0

                                # Chỉ tạo checkpoint cho chế độ chạy toàn bộ.
                                if tu_dong_den_du:
                                    checkpoint_run = tao_auto_checkpoint(
                                        cac_batch,
                                        batch_size_auto,
                                        pham_vi_xay_auto,
                                        tong_ke_hoach
                                    )
                                    luu_auto_checkpoint(
                                        checkpoint_run
                                    )
                                else:
                                    checkpoint_run = None

                            # Backup cấu hình GV đang dùng ở Kho YCCĐ.
                            backup_yccd = list(
                                st.session_state.yccd_da_chon
                            )

                            backup_cau_hinh = dict(
                                st.session_state.cau_hinh_yccd
                            )

                            backup_ai = list(
                                st.session_state.cau_hoi_ai
                            )

                            backup_duyet = dict(
                                st.session_state.trang_thai_duyet
                            )

                            cau_moi_auto = []

                            progress = st.progress(0)
                            status = st.empty()

                            try:

                                for idx_batch in range(
                                    chi_so_batch_bat_dau,
                                    len(cac_batch)
                                ):

                                    i_batch = idx_batch + 1
                                    batch = cac_batch[idx_batch]

                                    temp_items = []
                                    temp_cfg = {}

                                    for spec in batch:

                                        item_temp = {
                                            "Khối": spec["Khối"],
                                            "Chương": spec["Chương"],
                                            "Bài": spec["Bài"],
                                            "YCCĐ": spec["YCCĐ"]
                                        }

                                        temp_items.append(
                                            item_temp
                                        )

                                        yccd_id_temp = tao_id_yccd(
                                            item_temp["Khối"],
                                            item_temp["Chương"],
                                            item_temp["Bài"],
                                            item_temp["YCCĐ"]
                                        )

                                        cfg = {
                                            "Số câu": int(
                                                spec["Số câu"]
                                            ),
                                            "Mức độ": spec[
                                                "Mức độ"
                                            ],
                                            "Dạng câu hỏi": spec[
                                                "Dạng câu hỏi"
                                            ],
                                            "Thành phần năng lực": spec[
                                                "Thành phần năng lực"
                                            ],
                                            "Nguồn tham chiếu": [
                                                "SGK / Chương trình"
                                            ]
                                        }

                                        # Đúng/Sai: ở bước auto hiện tại,
                                        # dùng cùng YCCĐ/mức độ cho 4 ý.
                                        # Sau đó có thể nâng cấp sang cụm YCCĐ.
                                        if (
                                            spec["Dạng câu hỏi"]
                                            == "Đúng / Sai"
                                        ):
                                            # Đúng/Sai tích hợp:
                                            # 4 ý có thể thuộc YCCĐ, mức độ,
                                            # thành phần năng lực khác nhau.
                                            ds_yccd_cung_bai = [
                                                x["yccd"]
                                                for x in danh_sach_yccd_cov
                                                if x["bai"] == spec["Bài"]
                                            ]

                                            ds_yccd_cum = [
                                                spec["YCCĐ"]
                                            ] + [
                                                x for x in ds_yccd_cung_bai
                                                if x != spec["YCCĐ"]
                                            ]

                                            if not ds_yccd_cum:
                                                ds_yccd_cum = [
                                                    spec["YCCĐ"]
                                                ]

                                            while len(ds_yccd_cum) < 4:
                                                ds_yccd_cum.append(
                                                    ds_yccd_cum[
                                                        len(ds_yccd_cum)
                                                        % max(1, len(ds_yccd_cum))
                                                    ]
                                                )

                                            cau_hinh_4_y_auto = {}

                                            for idx_y, ky in enumerate(
                                                ["a", "b", "c", "d"]
                                            ):
                                                yccd_y = ds_yccd_cum[
                                                    idx_y
                                                    % len(ds_yccd_cum)
                                                ]

                                                mucdo_y = xac_dinh_muc_do(
                                                    yccd_y
                                                )

                                                # Ưu tiên đúng hai thành phần
                                                # năng lực đã thống nhất.
                                                if mucdo_y == "Vận dụng":
                                                    nl_y = (
                                                        "Vận dụng kiến thức, kĩ năng đã học"
                                                    )
                                                else:
                                                    nl_y = (
                                                        "Tìm hiểu thế giới sống"
                                                    )

                                                cau_hinh_4_y_auto[ky] = {
                                                    "yccd": yccd_y,
                                                    "muc_do": mucdo_y,
                                                    "thanh_phan_nang_luc": nl_y
                                                }

                                            cfg[
                                                "Chi tiết 4 ý"
                                            ] = cau_hinh_4_y_auto

                                        temp_cfg[
                                            yccd_id_temp
                                        ] = cfg

                                    st.session_state.yccd_da_chon = (
                                        temp_items
                                    )

                                    st.session_state.cau_hinh_yccd = (
                                        temp_cfg
                                    )

                                    status.write(
                                        f"🤖 Gemini đang tạo batch "
                                        f"{i_batch}/{len(cac_batch)}..."
                                    )

                                    ket_qua_batch = (
                                        goi_ai_tao_cau_hoi()
                                    )

                                    # Đưa ngay batch vừa tạo vào vùng chờ duyệt.
                                    for q_auto_batch in ket_qua_batch:
                                        q_auto_batch[
                                            "trang_thai"
                                        ] = "AI tạo – Chờ duyệt"

                                    st.session_state.auto_cau_cho_duyet.extend(
                                        ket_qua_batch
                                    )

                                    cau_moi_auto.extend(
                                        ket_qua_batch
                                    )

                                    # Lưu checkpoint NGAY SAU mỗi batch thành công.
                                    if checkpoint_run is not None:
                                        checkpoint_run[
                                            "batch_tiep_theo"
                                        ] = i_batch

                                        checkpoint_run[
                                            "da_tao"
                                        ] = int(
                                            checkpoint_run.get(
                                                "da_tao",
                                                0
                                            )
                                        ) + len(
                                            ket_qua_batch
                                        )

                                        checkpoint_run[
                                            "cau_da_tao"
                                        ].extend(
                                            ket_qua_batch
                                        )

                                        checkpoint_run[
                                            "trang_thai"
                                        ] = "dang_chay"

                                        checkpoint_run[
                                            "loi_cuoi"
                                        ] = ""

                                        luu_auto_checkpoint(
                                            checkpoint_run
                                        )

                                    progress.progress(
                                        i_batch
                                        / len(cac_batch)
                                    )

                            except Exception as e:

                                if checkpoint_run is not None:
                                    checkpoint_run[
                                        "trang_thai"
                                    ] = "tam_dung"

                                    checkpoint_run[
                                        "loi_cuoi"
                                    ] = str(e)

                                    luu_auto_checkpoint(
                                        checkpoint_run
                                    )

                                st.error(
                                    f"❌ Quá trình tự động tạm dừng: {e}"
                                )

                                if checkpoint_run is not None:
                                    st.warning(
                                        "💾 Các batch đã hoàn thành đã được giữ lại. "
                                        "Khi API hoạt động lại, bấm "
                                        "**TIẾP TỤC PHIÊN TRƯỚC**."
                                    )

                            finally:

                                # Khôi phục nguyên trạng cấu hình GV.
                                st.session_state.yccd_da_chon = (
                                    backup_yccd
                                )

                                st.session_state.cau_hinh_yccd = (
                                    backup_cau_hinh
                                )

                                st.session_state.cau_hoi_ai = (
                                    backup_ai
                                )

                                st.session_state.trang_thai_duyet = (
                                    backup_duyet
                                )

                            if cau_moi_auto:

                                st.session_state.auto_tien_do = {
                                    "da_tao": len(
                                        cau_moi_auto
                                    ),
                                    "tong_ke_hoach": tong_ke_hoach
                                }

                                if checkpoint_run is not None:
                                    # Hoàn tất khi batch_tiep_theo đã tới cuối.
                                    if int(
                                        checkpoint_run.get(
                                            "batch_tiep_theo",
                                            0
                                        )
                                    ) >= len(cac_batch):

                                        checkpoint_run[
                                            "trang_thai"
                                        ] = "hoan_tat"

                                        luu_auto_checkpoint(
                                            checkpoint_run
                                        )

                                if (
                                    checkpoint_run is not None
                                    and checkpoint_run.get(
                                        "trang_thai"
                                    ) == "hoan_tat"
                                ):
                                    st.success(
                                        f"✅ Đã hoàn tất phiên tự động. "
                                        f"Tổng cộng đã tạo "
                                        f"**{checkpoint_run.get('da_tao', 0)} câu**. "
                                        "Các câu đang chờ GV duyệt."
                                    )

                                    # Khi đã hoàn tất, checkpoint không còn cần để resume.
                                    xoa_auto_checkpoint()

                                elif checkpoint_run is None:
                                    st.success(
                                        f"✅ Đã tạo **{len(cau_moi_auto)} câu** "
                                        "và đưa vào khu vực chờ GV duyệt."
                                    )

                                st.rerun()

                    else:
                        st.success(
                            "🟢 Không còn câu nào cần bổ sung "
                            "trong phạm vi và mục tiêu hiện tại."
                        )

                # ==========================================
                # KHU VỰC CÂU AI TẠO - CHỜ GV DUYỆT
                # ==========================================
                ds_cho_duyet_auto = st.session_state.get(
                    "auto_cau_cho_duyet",
                    []
                )

                if ds_cho_duyet_auto:

                    st.markdown("---")
                    st.subheader(
                        f"👨‍🏫 CÂU AI TẠO – CHỜ DUYỆT "
                        f"({len(ds_cho_duyet_auto)})"
                    )

                    st.caption(
                        "Quy trình: chống trùng → kiểm định → sửa câu chưa đạt "
                        "→ kiểm định lại → GV duyệt. "
                        "Chỉ câu GV duyệt mới vào ngân hàng chính thức."
                    )

                    if st.button(
                        "🛡 KIỂM ĐỊNH TẤT CẢ CÂU CHỜ DUYỆT",
                        type="secondary",
                        use_container_width=True,
                        key="auto_validate_pending"
                    ):
                        with st.spinner(
                            "AI đang kiểm định các câu hỏi..."
                        ):
                            kq_kd_auto = (
                                kiem_dinh_danh_sach_cau_hoi(
                                    ds_cho_duyet_auto
                                )
                            )

                            st.session_state.ket_qua_kiem_dinh.update(
                                kq_kd_auto
                            )

                        if kq_kd_auto:
                            st.success(
                                f"Đã kiểm định "
                                f"**{len(kq_kd_auto)} câu**."
                            )
                            st.rerun()

                    bank_hien_tai_cho_duyet = doc_ngan_hang()

                    # --------------------------------------------------
                    # CÂN BẰNG ĐÁP ÁN 4 LỰA CHỌN KHÔNG CẦN AI
                    # --------------------------------------------------
                    thong_ke_dap_an = thong_ke_dap_an_4_lua_chon(
                        ds_cho_duyet_auto
                    )

                    if sum(thong_ke_dap_an.values()) > 0:
                        st.caption(
                            "Phân bố đáp án 4 lựa chọn hiện tại: "
                            + " • ".join(
                                f"{k}: {v}"
                                for k, v in thong_ke_dap_an.items()
                            )
                        )

                        if st.button(
                            "🔀 CÂN BẰNG A / B / C / D",
                            use_container_width=True,
                            key="auto_balance_mcq_answers"
                        ):
                            st.session_state.auto_cau_cho_duyet = (
                                can_bang_dap_an_4_lua_chon(
                                    ds_cho_duyet_auto
                                )
                            )

                            # Nội dung đã đổi thứ tự phương án nên kết quả
                            # kiểm định cũ không còn giá trị.
                            for q_bal in ds_cho_duyet_auto:
                                temp_bal = str(
                                    q_bal.get("temp_id", "")
                                )
                                if temp_bal:
                                    st.session_state.ket_qua_kiem_dinh.pop(
                                        temp_bal,
                                        None
                                    )

                            st.success(
                                "Đã xáo phương án và cân bằng vị trí đáp án đúng. "
                                "Nội dung câu hỏi không thay đổi."
                            )
                            st.rerun()

                    ds_can_ai_sua = []

                    for q_fix in ds_cho_duyet_auto:
                        temp_fix = str(
                            q_fix.get(
                                "temp_id",
                                ""
                            )
                        )

                        kd_fix = st.session_state.ket_qua_kiem_dinh.get(
                            temp_fix
                        )

                        if (
                            kd_fix
                            and kd_fix.get(
                                "ket_luan"
                            ) != "Đạt"
                        ):
                            ds_can_ai_sua.append(
                                q_fix
                            )

                    if ds_can_ai_sua:
                        if st.button(
                            f"🔧 AI SỬA {len(ds_can_ai_sua)} CÂU CHƯA ĐẠT",
                            use_container_width=True,
                            key="auto_fix_failed_all"
                        ):
                            da_sua = 0
                            bi_loi_sua = 0
                            ds_moi = list(
                                ds_cho_duyet_auto
                            )

                            with st.spinner(
                                "AI đang sửa lần lượt các câu chưa đạt..."
                            ):
                                for q_fix in ds_can_ai_sua:
                                    temp_fix = str(
                                        q_fix.get(
                                            "temp_id",
                                            ""
                                        )
                                    )

                                    cau_fix_moi, loi_fix = sua_cau_bang_ai(
                                        q_fix
                                    )

                                    if cau_fix_moi is None:
                                        bi_loi_sua += 1
                                        # Nếu 429 thì dừng để tránh gọi tiếp vô ích.
                                        if "429" in str(loi_fix):
                                            break
                                        continue

                                    for i_fix, q_old_fix in enumerate(
                                        ds_moi
                                    ):
                                        if str(
                                            q_old_fix.get(
                                                "temp_id",
                                                ""
                                            )
                                        ) == temp_fix:
                                            ds_moi[i_fix] = cau_fix_moi
                                            da_sua += 1
                                            break

                                    st.session_state.ket_qua_kiem_dinh.pop(
                                        temp_fix,
                                        None
                                    )

                            st.session_state.auto_cau_cho_duyet = ds_moi

                            if da_sua:
                                st.success(
                                    f"AI đã sửa **{da_sua} câu**. "
                                    "Các bản mới cần kiểm định lại."
                                )

                            if bi_loi_sua:
                                st.warning(
                                    f"Có **{bi_loi_sua} câu** chưa sửa được "
                                    "(có thể do 429/API)."
                                )

                            st.rerun()

                    cchon1, cchon2, cchon3 = st.columns(3)

                    with cchon1:
                        if st.button(
                            "☑️ Chọn câu an toàn",
                            key="auto_select_safe"
                        ):
                            for q_auto in ds_cho_duyet_auto:
                                temp_q = str(
                                    q_auto.get(
                                        "temp_id",
                                        id(q_auto)
                                    )
                                )

                                an_toan, _ = cau_an_toan_de_duyet_lo(
                                    q_auto,
                                    bank_hien_tai_cho_duyet
                                )

                                st.session_state[
                                    "auto_pick_" + temp_q
                                ] = bool(an_toan)

                            st.rerun()

                    with cchon2:
                        if st.button(
                            "☑️ Chọn tất cả",
                            key="auto_select_all"
                        ):
                            for q_auto in ds_cho_duyet_auto:
                                st.session_state[
                                    "auto_pick_"
                                    + str(
                                        q_auto.get(
                                            "temp_id",
                                            id(q_auto)
                                        )
                                    )
                                ] = True
                            st.rerun()

                    with cchon3:
                        if st.button(
                            "⬜ Bỏ chọn",
                            key="auto_unselect_all"
                        ):
                            for key_auto in list(
                                st.session_state.keys()
                            ):
                                if key_auto.startswith(
                                    "auto_pick_"
                                ):
                                    st.session_state[
                                        key_auto
                                    ] = False
                            st.rerun()

                    ds_duoc_chon = []

                    for i_auto, q_auto in enumerate(
                        ds_cho_duyet_auto,
                        start=1
                    ):

                        temp_id_auto = str(
                            q_auto.get(
                                "temp_id",
                                f"auto_{i_auto}"
                            )
                        )

                        key_pick = (
                            "auto_pick_"
                            + temp_id_auto
                        )

                        with st.container(border=True):

                            pick = st.checkbox(
                                f"Câu {i_auto} • "
                                f"{q_auto.get('muc_do', '')} • "
                                f"{q_auto.get('dang_cau', '')}",
                                key=key_pick
                            )

                            st.caption(
                                f"{q_auto.get('khoi', '')} • "
                                f"{q_auto.get('chuong', '')} • "
                                f"{q_auto.get('bai', '')}"
                            )

                            if q_auto.get("dang_cau") == "Đúng / Sai":
                                st.caption(
                                    "Câu Đúng/Sai tích hợp: YCCĐ, mức độ và "
                                    "thành phần năng lực được quản lí riêng theo từng ý a–d."
                                )
                            else:
                                st.write(
                                    "**YCCĐ:**",
                                    q_auto.get("yccd", "")
                                )
                                st.write(
                                    "**Thành phần năng lực:**",
                                    q_auto.get(
                                        "thanh_phan_nang_luc",
                                        ""
                                    )
                                )

                            if q_auto.get("tinh_huong"):
                                st.write(
                                    "**Tình huống:**",
                                    q_auto.get(
                                        "tinh_huong",
                                        ""
                                    )
                                )

                            st.write(
                                "**Câu hỏi:**",
                                q_auto.get("cau_hoi", "")
                            )

                            if q_auto.get("dang_cau") == "Đúng / Sai":
                                meta_hien_thi_auto = list(
                                    q_auto.get("nhan_dinh_meta", []) or []
                                )

                                if len(meta_hien_thi_auto) == 4:
                                    for i_nd, nd in enumerate(
                                        meta_hien_thi_auto
                                    ):
                                        st.markdown(
                                            f"**{'abcd'[i_nd]})** "
                                            f"{nd.get('noi_dung', '')}"
                                        )
                                        st.caption(
                                            f"🎯 YCCĐ: {nd.get('yccd', '')}  |  "
                                            f"Mức độ: {nd.get('muc_do', '')}  |  "
                                            f"Năng lực: {nd.get('thanh_phan_nang_luc', '')}"
                                        )

                            kq_trung_auto = kiem_tra_trung_gan(
                                q_auto,
                                bank_hien_tai_cho_duyet
                            )

                            ti_le_trung_auto = float(
                                kq_trung_auto.get(
                                    "ti_le_cao_nhat",
                                    0
                                )
                            )

                            if ti_le_trung_auto >= 0.92:
                                st.error(
                                    f"🔴 Trùng cao: {ti_le_trung_auto:.0%} "
                                    "so với câu đã có trong ngân hàng."
                                )
                                if kq_trung_auto.get("cau_gan_nhat"):
                                    st.caption(
                                        "Câu gần nhất: "
                                        + str(
                                            kq_trung_auto.get(
                                                "cau_gan_nhat",
                                                ""
                                            )
                                        )[:260]
                                    )
                            elif ti_le_trung_auto >= 0.78:
                                st.warning(
                                    f"🟡 Gần giống: {ti_le_trung_auto:.0%} "
                                    "— GV nên xem trước khi duyệt."
                                )
                            else:
                                st.success(
                                    f"🟢 Chống trùng nội bộ: "
                                    f"{ti_le_trung_auto:.0%} — an toàn."
                                )

                            loi_cau_auto = kiem_tra_cau_moi_truoc_khi_luu(
                                q_auto
                            )

                            if loi_cau_auto:
                                st.error(
                                    "⛔ Chưa đủ chuẩn để lưu: "
                                    + "; ".join(loi_cau_auto)
                                )
                            else:
                                st.success(
                                    "✅ Đủ cấu trúc bắt buộc để GV duyệt và lưu."
                                )

                            if q_auto.get("lua_chon"):
                                for lc_auto in q_auto.get(
                                    "lua_chon",
                                    []
                                ):
                                    st.write(lc_auto)

                            st.markdown("---")
                            hien_thi_ket_qua_kiem_dinh(
                                q_auto
                            )

                            hien_thi_so_sanh_phien_ban(
                                q_auto
                            )

                            kd_hien_tai = st.session_state.ket_qua_kiem_dinh.get(
                                temp_id_auto
                            )

                            can_sua = (
                                bool(loi_cau_auto)
                                or ti_le_trung_auto >= 0.78
                                or (
                                    kd_hien_tai
                                    and kd_hien_tai.get(
                                        "ket_luan"
                                    ) != "Đạt"
                                )
                            )

                            with st.expander(
                                "✏️ SỬA CÂU HỎI",
                                expanded=False
                            ):
                                st.caption(
                                    "Có thể sửa thủ công ngay. "
                                    "AI sửa theo lỗi kiểm định chỉ dùng được khi API hoạt động."
                                )

                                cau_sua_thu_cong = sua_cau_thu_cong_tu_form(
                                    q_auto,
                                    temp_id_auto
                                )

                                if cau_sua_thu_cong is not None:
                                    temp_cu = temp_id_auto

                                    # Xóa kết quả kiểm định cũ vì nội dung đã thay đổi.
                                    st.session_state.ket_qua_kiem_dinh.pop(
                                        temp_cu,
                                        None
                                    )

                                    if thay_cau_trong_danh_sach_cho_duyet(
                                        temp_cu,
                                        cau_sua_thu_cong
                                    ):
                                        st.success(
                                            "Đã lưu bản sửa thủ công. "
                                            "Câu cần được kiểm định lại trước khi duyệt."
                                        )
                                        st.rerun()

                                if kd_hien_tai:
                                    if st.button(
                                        "🔧 AI SỬA THEO LỖI KIỂM ĐỊNH",
                                        use_container_width=True,
                                        key=f"ai_fix_{temp_id_auto}"
                                    ):
                                        with st.spinner(
                                            "AI đang sửa đúng các lỗi đã kiểm định..."
                                        ):
                                            cau_ai_sua, loi_ai_sua = sua_cau_bang_ai(
                                                q_auto
                                            )

                                        if cau_ai_sua is not None:
                                            st.session_state.ket_qua_kiem_dinh.pop(
                                                temp_id_auto,
                                                None
                                            )

                                            if thay_cau_trong_danh_sach_cho_duyet(
                                                temp_id_auto,
                                                cau_ai_sua
                                            ):
                                                st.success(
                                                    "AI đã sửa câu. "
                                                    "Bản cũ được lưu trong lịch sử; "
                                                    "hãy kiểm định lại bản mới."
                                                )
                                                st.rerun()
                                        else:
                                            st.warning(
                                                loi_ai_sua
                                            )
                                else:
                                    st.info(
                                        "Muốn dùng AI sửa theo lỗi, "
                                        "cần có kết quả kiểm định trước."
                                    )

                            with st.expander(
                                "Xem đáp án / giải thích"
                            ):
                                if q_auto.get("dang_cau") == "Đúng / Sai":
                                    meta_ds = list(
                                        q_auto.get("nhan_dinh_meta", []) or []
                                    )

                                    for i_ds, nd in enumerate(meta_ds[:4]):
                                        ky_ds = "abcd"[i_ds]
                                        dap_ds = chuan_hoa_dap_an_dung_sai(
                                            nd.get("dap_an", "")
                                        )

                                        st.markdown(
                                            f"### Ý {ky_ds}) — "
                                            f"{dap_ds if dap_ds else '⚠️ Chưa có đáp án'}"
                                        )

                                        st.write(
                                            "**Giải thích:**",
                                            nd.get("giai_thich", "")
                                            or "⚠️ Chưa có giải thích"
                                        )

                                        st.caption(
                                            f"YCCĐ: {nd.get('yccd', '') or 'Chưa có'}  |  "
                                            f"Mức độ: {nd.get('muc_do', '') or 'Chưa có'}  |  "
                                            f"Thành phần năng lực: "
                                            f"{nd.get('thanh_phan_nang_luc', '') or 'Chưa có'}"
                                        )

                                        if i_ds < len(meta_ds) - 1:
                                            st.divider()

                                    loi_thieu_ds = cau_dung_sai_thieu_hoan_thien(
                                        q_auto
                                    )

                                    if loi_thieu_ds:
                                        st.warning(
                                            "Câu này còn thiếu: "
                                            + "; ".join(loi_thieu_ds)
                                        )

                                        if st.button(
                                            "✨ HOÀN THIỆN ĐÁP ÁN / GIẢI THÍCH / METADATA",
                                            use_container_width=True,
                                            key=f"complete_tf_{temp_id_auto}"
                                        ):
                                            with st.spinner(
                                                "AI đang hoàn thiện câu nhưng giữ nguyên nội dung 4 nhận định..."
                                            ):
                                                q_hoan_thien, err_ht = (
                                                    hoan_thien_cau_dung_sai_bang_ai(
                                                        q_auto
                                                    )
                                                )

                                            if q_hoan_thien is not None:
                                                st.session_state.ket_qua_kiem_dinh.pop(
                                                    temp_id_auto,
                                                    None
                                                )

                                                if thay_cau_trong_danh_sach_cho_duyet(
                                                    temp_id_auto,
                                                    q_hoan_thien
                                                ):
                                                    st.success(
                                                        "Đã hoàn thiện câu. Nội dung tình huống và 4 nhận định được giữ nguyên. "
                                                        "Hãy kiểm định lại trước khi duyệt."
                                                    )
                                                    st.rerun()
                                            else:
                                                st.warning(err_ht)

                                else:
                                    st.write(
                                        "**Đáp án:**",
                                        q_auto.get(
                                            "dap_an",
                                            ""
                                        )
                                    )
                                    st.write(
                                        "**Giải thích:**",
                                        q_auto.get(
                                            "giai_thich",
                                            ""
                                        )
                                    )

                            if pick:
                                ds_duoc_chon.append(
                                    q_auto
                                )

                    d1, d2 = st.columns(2)

                    with d1:
                        if st.button(
                            f"✅ DUYỆT {len(ds_duoc_chon)} CÂU",
                            type="primary",
                            use_container_width=True,
                            disabled=(
                                len(ds_duoc_chon) == 0
                            ),
                            key="auto_approve_selected"
                        ):

                            da_luu = 0
                            bi_trung = 0
                            bi_loi = 0
                            kd_chua_dat = 0

                            id_da_luu = set()

                            # Đọc lại ngân hàng tại thời điểm duyệt.
                            bank_duyet_lo = doc_ngan_hang()

                            for q_auto in ds_duoc_chon:

                                loi_q = kiem_tra_cau_moi_truoc_khi_luu(
                                    q_auto
                                )

                                if loi_q:
                                    bi_loi += 1
                                    continue

                                trung_q = kiem_tra_trung_gan(
                                    q_auto,
                                    bank_duyet_lo
                                )

                                if trung_q.get(
                                    "ti_le_cao_nhat",
                                    0
                                ) >= 0.92:
                                    bi_trung += 1
                                    continue

                                temp_q = str(
                                    q_auto.get(
                                        "temp_id",
                                        ""
                                    )
                                )

                                kd_q = st.session_state.ket_qua_kiem_dinh.get(
                                    temp_q
                                )

                                # Nếu đã có kết quả kiểm định và chưa ĐẠT,
                                # không cho duyệt theo lô.
                                # Chưa kiểm định do 429: GV vẫn có thể tự duyệt.
                                if (
                                    kd_q
                                    and kd_q.get(
                                        "ket_luan"
                                    ) != "Đạt"
                                ):
                                    kd_chua_dat += 1
                                    continue

                                if luu_cau_vao_ngan_hang(
                                    q_auto
                                ):
                                    da_luu += 1
                                    id_da_luu.add(
                                        temp_q
                                    )

                                    # Cập nhật bank cục bộ để câu sau
                                    # cũng được so với câu vừa lưu.
                                    q_bank_temp = dict(q_auto)
                                    q_bank_temp["id"] = (
                                        q_bank_temp.get("id")
                                        or temp_q
                                    )
                                    bank_duyet_lo.append(
                                        q_bank_temp
                                    )
                                else:
                                    bi_trung += 1

                            # Chỉ bỏ khỏi danh sách chờ những câu đã lưu thật sự.
                            # Câu bị trùng/lỗi/chưa đạt vẫn giữ lại để GV xử lí.
                            st.session_state.auto_cau_cho_duyet = [
                                q
                                for q in ds_cho_duyet_auto
                                if str(
                                    q.get(
                                        "temp_id",
                                        ""
                                    )
                                ) not in id_da_luu
                            ]

                            st.success(
                                f"Đã duyệt và lưu **{da_luu} câu**."
                            )

                            if bi_trung:
                                st.warning(
                                    f"Giữ lại **{bi_trung} câu** vì trùng/gần trùng cao."
                                )

                            if bi_loi:
                                st.warning(
                                    f"Giữ lại **{bi_loi} câu** vì lỗi cấu trúc."
                                )

                            if kd_chua_dat:
                                st.warning(
                                    f"Giữ lại **{kd_chua_dat} câu** vì đã kiểm định "
                                    "nhưng chưa đạt."
                                )

                            st.rerun()

                    with d2:
                        if st.button(
                            f"🗑 XÓA {len(ds_duoc_chon)} CÂU ĐÃ CHỌN",
                            use_container_width=True,
                            disabled=(
                                len(ds_duoc_chon) == 0
                            ),
                            key="auto_delete_selected"
                        ):

                            id_xoa = {
                                str(
                                    q.get(
                                        "temp_id",
                                        ""
                                    )
                                )
                                for q in ds_duoc_chon
                            }

                            for temp_xoa in id_xoa:
                                if temp_xoa:
                                    st.session_state.ket_qua_kiem_dinh.pop(
                                        temp_xoa,
                                        None
                                    )

                            st.session_state.auto_cau_cho_duyet = [
                                q
                                for q in ds_cho_duyet_auto
                                if str(
                                    q.get(
                                        "temp_id",
                                        ""
                                    )
                                ) not in id_xoa
                            ]

                            for temp_xoa in id_xoa:
                                st.session_state.pop(
                                    "auto_pick_" + temp_xoa,
                                    None
                                )

                            st.rerun()

                    if st.button(
                        "🧹 XÓA TOÀN BỘ DANH SÁCH CHỜ DUYỆT",
                        use_container_width=True,
                        key="auto_clear_pending"
                    ):

                        for q_pending in ds_cho_duyet_auto:
                            temp_pending = str(
                                q_pending.get(
                                    "temp_id",
                                    ""
                                )
                            )
                            if temp_pending:
                                st.session_state.ket_qua_kiem_dinh.pop(
                                    temp_pending,
                                    None
                                )

                        st.session_state.auto_cau_cho_duyet = []

                        for key_auto in list(
                            st.session_state.keys()
                        ):
                            if key_auto.startswith(
                                "auto_pick_"
                            ):
                                del st.session_state[
                                    key_auto
                                ]

                        st.rerun()

    st.divider()

    # ======================================================
    # THỐNG KÊ GỌN
    # ======================================================
    so_khoi = len({
        q.get("khoi", "")
        for q in bank
        if q.get("khoi")
    })

    so_chuong = len({
        (q.get("khoi", ""), q.get("chuong", ""))
        for q in bank
        if q.get("chuong")
    })

    so_bai = len({
        (q.get("khoi", ""), q.get("chuong", ""), q.get("bai", ""))
        for q in bank
        if q.get("bai")
    })

    c1, c2, c3, c4 = st.columns(4)

    with c1:
        st.metric("Tổng số câu", len(bank))

    with c2:
        st.metric("Khối", so_khoi)

    with c3:
        st.metric("Chương", so_chuong)

    with c4:
        st.metric("Bài", so_bai)

    st.divider()

    # ======================================================
    # ======================================================
    # BỔ SUNG MỞ RỘNG NGÂN HÀNG
    # ======================================================
    st.markdown("---")

    with st.expander("➕ BỔ SUNG MỞ RỘNG NGÂN HÀNG", expanded=False):
        st.caption(
            "Dùng khi YCCĐ đã đủ chỉ tiêu nhưng GV vẫn muốn tạo thêm câu mới. "
            "Không làm thay đổi mục tiêu chung."
        )

        if "danh_sach_yccd_cov" in locals() and danh_sach_yccd_cov:
            ds_bai_mr = list(dict.fromkeys(x["bai"] for x in danh_sach_yccd_cov))
            bai_mr = st.selectbox("📖 Bài cần bổ sung", ds_bai_mr, key="mr_bai")

            ds_yccd_mr = [
                x["yccd"] for x in danh_sach_yccd_cov
                if x["bai"] == bai_mr
            ]
            yccd_mr = st.selectbox("🎯 YCCĐ cần bổ sung", ds_yccd_mr, key="mr_yccd")

            muc_do_mr = xac_dinh_muc_do(yccd_mr)
            dang_mr = st.selectbox(
                "🧩 Dạng câu hỏi",
                ["Trắc nghiệm 4 lựa chọn", "Đúng / Sai", "Trả lời ngắn"],
                key="mr_dang"
            )
            nl_mr = goi_y_thanh_phan_nang_luc(yccd_mr, muc_do_mr, dang_mr)

            c1_mr, c2_mr = st.columns(2)
            with c1_mr:
                st.write(f"**Mức độ đã gán:** {muc_do_mr}")
            with c2_mr:
                st.write(f"**Thành phần năng lực:** {nl_mr}")

            so_cau_mr = st.number_input(
                "➕ Số câu muốn bổ sung",
                min_value=1, max_value=50, value=5, step=1,
                key="mr_so_cau"
            )

            st.info(
                "Câu bổ sung vẫn phải qua khu vực chờ GV duyệt và kiểm định, "
                "không tự động vào ngân hàng."
            )
        else:
            st.info("Hãy chọn phạm vi ở phần Độ phủ ngân hàng trước.")



def ngan_hang_cau_hoi():

    st.header("🏦 NGÂN HÀNG CÂU HỎI")
    st.caption(
        "Chỉ hiển thị các câu đã được GV duyệt. "
        "Độ phủ và xây dựng tự động được quản lý ở mục "
        "🧱 Xây dựng NH ôn tập / kiểm tra."
    )

    bank = doc_ngan_hang()

    if not bank:
        st.info("Ngân hàng chưa có câu hỏi.")
        return

    # BỘ LỌC
    # ======================================================
    st.subheader("🔎 Chọn phạm vi câu hỏi")

    col1, col2, col3 = st.columns(3)

    ds_khoi = sorted({
        q.get("khoi", "")
        for q in bank
        if q.get("khoi")
    })

    with col1:
        loc_khoi = st.selectbox(
            "🎓 Khối",
            ["Tất cả"] + ds_khoi,
            key="bank_filter_khoi"
        )

    bank_khoi = (
        bank
        if loc_khoi == "Tất cả"
        else [
            q for q in bank
            if q.get("khoi") == loc_khoi
        ]
    )

    ds_chuong = sorted({
        q.get("chuong", "")
        for q in bank_khoi
        if q.get("chuong")
    })

    with col2:
        loc_chuong = st.selectbox(
            "📚 Chương",
            ["Tất cả"] + ds_chuong,
            key="bank_filter_chuong"
        )

    bank_chuong = (
        bank_khoi
        if loc_chuong == "Tất cả"
        else [
            q for q in bank_khoi
            if q.get("chuong") == loc_chuong
        ]
    )

    ds_bai = sorted({
        q.get("bai", "")
        for q in bank_chuong
        if q.get("bai")
    })

    with col3:
        loc_bai = st.selectbox(
            "📖 Bài",
            ["Tất cả"] + ds_bai,
            key="bank_filter_bai"
        )

    bank_bai = (
        bank_chuong
        if loc_bai == "Tất cả"
        else [
            q for q in bank_chuong
            if q.get("bai") == loc_bai
        ]
    )

    col4, col5, col6 = st.columns(3)

    ds_yccd = sorted({
        q.get("yccd", "")
        for q in bank_bai
        if q.get("yccd")
    })

    with col4:
        loc_yccd = st.selectbox(
            "🎯 YCCĐ",
            ["Tất cả"] + ds_yccd,
            key="bank_filter_yccd"
        )

    bank_yccd = (
        bank_bai
        if loc_yccd == "Tất cả"
        else [
            q for q in bank_bai
            if q.get("yccd") == loc_yccd
        ]
    )

    ds_dang = sorted({
        q.get("dang_cau", "")
        for q in bank_yccd
        if q.get("dang_cau")
    })

    with col5:
        loc_dang = st.selectbox(
            "🧩 Dạng câu",
            ["Tất cả"] + ds_dang,
            key="bank_filter_dang"
        )

    bank_dang = (
        bank_yccd
        if loc_dang == "Tất cả"
        else [
            q for q in bank_yccd
            if q.get("dang_cau") == loc_dang
        ]
    )

    ds_muc_do = sorted({
        q.get("muc_do", "")
        for q in bank_dang
        if q.get("muc_do")
    })

    with col6:
        loc_muc_do = st.selectbox(
            "📊 Mức độ",
            ["Tất cả"] + ds_muc_do,
            key="bank_filter_mucdo"
        )

    ds_nang_luc = sorted({
        q.get("thanh_phan_nang_luc", "")
        for q in bank_dang
        if q.get("thanh_phan_nang_luc")
    })

    loc_nang_luc = st.selectbox(
        "🧠 Thành phần năng lực",
        ["Tất cả"] + ds_nang_luc,
        key="bank_filter_nangluc"
    )

    tu_khoa = st.text_input(
        "🔍 Tìm theo nội dung câu hỏi / YCCĐ",
        placeholder="Nhập từ khóa nếu cần...",
        key="bank_search"
    ).strip()

    # ======================================================
    # CHỈ KHI BẤM OK MỚI CHỐT BỘ LỌC VÀ HIỆN KẾT QUẢ
    # ======================================================
    _, cot_ok, _ = st.columns([1, 2, 1])

    with cot_ok:
        bam_ok = st.button(
            "🔎 OK - XEM CÂU HỎI",
            type="primary",
            use_container_width=True
        )

    if bam_ok:
        st.session_state["bank_filter_da_chot"] = {
            "khoi": loc_khoi,
            "chuong": loc_chuong,
            "bai": loc_bai,
            "yccd": loc_yccd,
            "dang": loc_dang,
            "muc_do": loc_muc_do,
            "nang_luc": loc_nang_luc,
            "tu_khoa": tu_khoa,
        }

    cau_hinh_da_chot = st.session_state.get(
        "bank_filter_da_chot"
    )

    if not cau_hinh_da_chot:
        st.info(
            "Chọn các điều kiện phía trên rồi bấm "
            "**OK - XEM CÂU HỎI**."
        )
        return

    # ======================================================
    # LỌC THEO BỘ LỌC ĐÃ BẤM OK
    # ======================================================
    filtered = list(bank)

    if cau_hinh_da_chot["khoi"] != "Tất cả":
        filtered = [
            q for q in filtered
            if q.get("khoi") == cau_hinh_da_chot["khoi"]
        ]

    if cau_hinh_da_chot["chuong"] != "Tất cả":
        filtered = [
            q for q in filtered
            if q.get("chuong") == cau_hinh_da_chot["chuong"]
        ]

    if cau_hinh_da_chot["bai"] != "Tất cả":
        filtered = [
            q for q in filtered
            if q.get("bai") == cau_hinh_da_chot["bai"]
        ]

    if cau_hinh_da_chot["yccd"] != "Tất cả":
        filtered = [
            q for q in filtered
            if q.get("yccd") == cau_hinh_da_chot["yccd"]
        ]

    if cau_hinh_da_chot["dang"] != "Tất cả":
        filtered = [
            q for q in filtered
            if q.get("dang_cau") == cau_hinh_da_chot["dang"]
        ]

    if cau_hinh_da_chot["muc_do"] != "Tất cả":
        filtered = [
            q for q in filtered
            if q.get("muc_do") == cau_hinh_da_chot["muc_do"]
        ]

    if cau_hinh_da_chot.get("nang_luc", "Tất cả") != "Tất cả":
        filtered = [
            q for q in filtered
            if q.get("thanh_phan_nang_luc")
            == cau_hinh_da_chot["nang_luc"]
        ]

    tu_khoa_chot = str(
        cau_hinh_da_chot.get("tu_khoa", "")
    ).strip().lower()

    if tu_khoa_chot:
        filtered = [
            q for q in filtered
            if (
                tu_khoa_chot in str(q.get("cau_hoi", "")).lower()
                or tu_khoa_chot in str(q.get("tinh_huong", "")).lower()
                or tu_khoa_chot in str(q.get("yccd", "")).lower()
                or tu_khoa_chot in str(q.get("bai", "")).lower()
            )
        ]

    filtered = sorted(
        filtered,
        key=lambda q: (
            str(q.get("khoi", "")),
            str(q.get("chuong", "")),
            str(q.get("bai", "")),
            str(q.get("yccd", "")),
            str(q.get("dang_cau", "")),
            str(q.get("muc_do", ""))
        )
    )

    st.success(
        f"✅ Tìm thấy **{len(filtered)} câu hỏi** phù hợp."
    )

    if not filtered:
        return

    st.divider()

    # ======================================================
    # HIỆN THẲNG CÁC CÂU PHÙ HỢP SAU KHI BẤM OK
    # ======================================================
    for so_tt, q in enumerate(filtered, start=1):

        q_id = str(q.get("id") or f"bank_{so_tt}")
        dang_cau = q.get("dang_cau", "")
        muc_do = q.get("muc_do", "")

        if dang_cau == "Đúng / Sai":
            st.markdown(
                f"### Câu {so_tt} • Đúng / Sai • Tích hợp 4 ý"
            )
        else:
            st.markdown(
                f"### Câu {so_tt} • {dang_cau} • {muc_do}"
            )

        st.caption(
            f"{q.get('khoi', '')}"
            f" • {q.get('chuong', '')}"
            f" • {q.get('bai', '')}"
        )

        if dang_cau == "Đúng / Sai":
            st.caption(
                "YCCĐ, mức độ và thành phần năng lực được lưu riêng theo từng ý a–d."
            )
        else:
            st.markdown("**YCCĐ:**")
            st.write(q.get("yccd", ""))

            st.write(
                "**Thành phần năng lực:**",
                q.get("thanh_phan_nang_luc", "")
            )

        if q.get("trang_thai"):
            st.write(
                "**Trạng thái:**",
                q.get("trang_thai", "")
            )

        # --------------------------------------------------
        # ĐÚNG / SAI
        # --------------------------------------------------
        if dang_cau == "Đúng / Sai":

            tinh_huong = q.get("tinh_huong", "")

            if tinh_huong:
                st.markdown("**Tình huống / dữ liệu:**")
                st.write(tinh_huong)

            st.markdown("**Câu hỏi / câu lệnh:**")
            st.write(
                q.get("cau_hoi", "")
                or "Dựa vào thông tin trên, hãy xác định mỗi nhận định sau là Đúng hay Sai."
            )

            nhan_dinh_meta = q.get(
                "nhan_dinh_meta",
                []
            )

            if nhan_dinh_meta:

                st.markdown("**Các nhận định:**")

                for ky_hieu, nd in zip(
                    ["a", "b", "c", "d"],
                    nhan_dinh_meta
                ):

                    noi_dung = str(
                        nd.get("noi_dung", "")
                    ).strip()

                    for nhan_cu in [
                        f"{ky_hieu})",
                        f"{ky_hieu}.",
                        f"{ky_hieu}:"
                    ]:
                        if noi_dung.lower().startswith(
                            nhan_cu.lower()
                        ):
                            noi_dung = noi_dung[
                                len(nhan_cu):
                            ].strip()
                            break

                    st.write(
                        f"**{ky_hieu})** {noi_dung}"
                    )

                    yccd_y = nd.get("yccd", "")
                    muc_do_y = nd.get("muc_do", "")
                    nang_luc_y = nd.get(
                        "thanh_phan_nang_luc",
                        ""
                    )

                    if yccd_y or muc_do_y or nang_luc_y:
                        st.caption(
                            f"Ý {ky_hieu}: "
                            f"Mức độ: {muc_do_y or 'Chưa xác định'}"
                            + (
                                f" • Năng lực: {nang_luc_y}"
                                if nang_luc_y else ""
                            )
                            + (
                                f" • YCCĐ: {yccd_y}"
                                if yccd_y else ""
                            )
                        )

                ds_dap_an = [
                    f"{ky_hieu} - {nd.get('dap_an', '')}"
                    for ky_hieu, nd in zip(
                        ["a", "b", "c", "d"],
                        nhan_dinh_meta
                    )
                ]

                with st.expander(
                    "👁️ Xem đáp án và giải thích"
                ):
                    st.write(
                        "**Đáp án:**",
                        ", ".join(ds_dap_an)
                    )

                    for ky_hieu, nd in zip(
                        ["a", "b", "c", "d"],
                        nhan_dinh_meta
                    ):
                        st.write(
                            f"**{ky_hieu}) "
                            f"{nd.get('dap_an', '')}:** "
                            f"{nd.get('giai_thich', '')}"
                        )

            else:
                for lc in q.get("lua_chon", []):
                    st.write(lc)

                with st.expander(
                    "👁️ Xem đáp án và giải thích"
                ):
                    st.write(
                        "**Đáp án:**",
                        q.get("dap_an", "")
                    )
                    st.write(
                        "**Giải thích:**",
                        q.get("giai_thich", "")
                    )

        # --------------------------------------------------
        # TRẮC NGHIỆM / TRẢ LỜI NGẮN
        # --------------------------------------------------
        else:

            st.markdown("**Câu hỏi:**")
            st.write(q.get("cau_hoi", ""))

            lua_chon = q.get("lua_chon", [])

            if lua_chon:
                st.markdown("**Các phương án:**")
                for lc in lua_chon:
                    st.write(lc)

            with st.expander(
                "👁️ Xem đáp án và giải thích"
            ):
                st.write(
                    "**Đáp án:**",
                    q.get("dap_an", "")
                )

                st.write(
                    "**Giải thích:**",
                    q.get("giai_thich", "")
                )

        st.write(
            "**Nguồn:**",
            q.get("nguon", "")
            or "Chưa ghi nguồn"
        )

        kd_da_luu = q.get(
            "kiem_dinh_ai"
        )

        if kd_da_luu:
            ket_luan_kd = kd_da_luu.get(
                "ket_luan",
                ""
            )
            tong_kd = kd_da_luu.get(
                "tong_diem",
                0
            )

            if ket_luan_kd == "Đạt":
                st.success(
                    f"🛡 Kiểm định AI: "
                    f"**{tong_kd}/100 – Đạt**"
                )
            elif ket_luan_kd == "Không đạt":
                st.error(
                    f"🛡 Kiểm định AI: "
                    f"**{tong_kd}/100 – Không đạt**"
                )
            else:
                st.warning(
                    f"🛡 Kiểm định AI: "
                    f"**{tong_kd}/100 – Cần xem lại**"
                )

        if q.get("ngay_tao"):
            st.caption(
                "Ngày duyệt: "
                + q.get("ngay_tao", "")
            )

        # ==================================================
        # SỬA / XÓA
        # ==================================================
        edit_key = f"edit_{q_id}"

        if edit_key not in st.session_state:
            st.session_state[edit_key] = False

        if not st.session_state[edit_key]:

            col_edit, col_delete = st.columns(2)

            with col_edit:
                if st.button(
                    "✏️ SỬA CÂU HỎI",
                    key=f"btn_edit_{q_id}",
                    use_container_width=True
                ):
                    st.session_state[edit_key] = True
                    st.rerun()

            with col_delete:
                if st.button(
                    "🗑️ XÓA KHỎI NGÂN HÀNG",
                    key=f"btn_delete_{q_id}",
                    use_container_width=True
                ):
                    st.session_state[
                        f"confirm_delete_{q_id}"
                    ] = True
                    st.rerun()

            if st.session_state.get(
                f"confirm_delete_{q_id}",
                False
            ):

                st.warning(
                    "⚠️ Bạn có chắc muốn xóa câu hỏi này không?"
                )

                c_yes, c_no = st.columns(2)

                with c_yes:
                    if st.button(
                        "✅ CÓ, XÓA",
                        key=f"yes_delete_{q_id}",
                        type="primary",
                        use_container_width=True
                    ):

                        bank_moi = [
                            item
                            for item in bank
                            if item.get("id") != q.get("id")
                        ]

                        luu_ngan_hang(bank_moi)

                        st.session_state.pop(
                            f"confirm_delete_{q_id}",
                            None
                        )

                        st.session_state.pop(
                            edit_key,
                            None
                        )

                        st.rerun()

                with c_no:
                    if st.button(
                        "❌ KHÔNG",
                        key=f"no_delete_{q_id}",
                        use_container_width=True
                    ):
                        st.session_state[
                            f"confirm_delete_{q_id}"
                        ] = False
                        st.rerun()

        else:

            st.info("✏️ Đang chỉnh sửa câu hỏi")

            cau_hoi_moi = st.text_area(
                "Câu hỏi / câu lệnh",
                value=q.get("cau_hoi", ""),
                key=f"edit_question_{q_id}",
                height=110
            )

            tinh_huong_moi = q.get(
                "tinh_huong",
                ""
            )

            if dang_cau == "Đúng / Sai":
                tinh_huong_moi = st.text_area(
                    "Tình huống / dữ liệu",
                    value=q.get(
                        "tinh_huong",
                        ""
                    ),
                    key=f"edit_situation_{q_id}",
                    height=150
                )

            lua_chon_moi = list(
                q.get("lua_chon", [])
            )

            if dang_cau == "Trắc nghiệm 4 lựa chọn":

                while len(lua_chon_moi) < 4:
                    lua_chon_moi.append("")

                st.markdown("**Các phương án:**")

                ds_lua_chon_moi = []

                for j, nhan in enumerate(
                    ["A", "B", "C", "D"]
                ):

                    gia_tri = st.text_input(
                        f"Phương án {nhan}",
                        value=lua_chon_moi[j],
                        key=f"edit_option_{q_id}_{j}"
                    )

                    if gia_tri.strip():
                        ds_lua_chon_moi.append(
                            gia_tri.strip()
                        )

                lua_chon_moi = ds_lua_chon_moi

            dap_an_moi = st.text_input(
                "Đáp án đúng",
                value=q.get("dap_an", ""),
                key=f"edit_answer_{q_id}"
            )

            giai_thich_moi = st.text_area(
                "Giải thích",
                value=q.get("giai_thich", ""),
                key=f"edit_explain_{q_id}",
                height=110
            )

            nguon_moi = st.text_input(
                "Nguồn tham chiếu",
                value=q.get("nguon", ""),
                key=f"edit_source_{q_id}"
            )

            col_save, col_cancel = st.columns(2)

            with col_save:
                if st.button(
                    "💾 LƯU CHỈNH SỬA",
                    key=f"save_edit_{q_id}",
                    type="primary",
                    use_container_width=True
                ):

                    if (
                        dang_cau != "Đúng / Sai"
                        and not cau_hoi_moi.strip()
                    ):
                        st.error(
                            "Câu hỏi không được để trống."
                        )

                    else:
                        for item in bank:
                            if item.get("id") == q.get("id"):

                                item["cau_hoi"] = (
                                    cau_hoi_moi.strip()
                                )

                                if dang_cau == "Đúng / Sai":
                                    item["tinh_huong"] = (
                                        tinh_huong_moi.strip()
                                    )

                                item["lua_chon"] = (
                                    lua_chon_moi
                                )

                                item["dap_an"] = (
                                    dap_an_moi.strip()
                                )

                                item["giai_thich"] = (
                                    giai_thich_moi.strip()
                                )

                                item["nguon"] = (
                                    nguon_moi.strip()
                                )

                                item["ngay_sua"] = (
                                    datetime.now().strftime(
                                        "%d/%m/%Y %H:%M"
                                    )
                                )

                                break

                        luu_ngan_hang(bank)

                        st.session_state[
                            edit_key
                        ] = False

                        st.rerun()

            with col_cancel:
                if st.button(
                    "↩️ HỦY",
                    key=f"cancel_edit_{q_id}",
                    use_container_width=True
                ):
                    st.session_state[
                        edit_key
                    ] = False
                    st.rerun()

        st.divider()




def ket_qua_hoc_sinh():

    st.header(
        "📊 KẾT QUẢ HỌC SINH"
    )

    ds = doc_lich_su_hoc_sinh()

    if not ds:
        st.info(
            "Chưa có lượt luyện nào được lưu."
        )
        return

    ds_hs = sorted({
        str(
            x.get(
                "hoc_sinh_id",
                ""
            )
        )
        for x in ds
        if x.get(
            "hoc_sinh_id"
        )
    })

    hs = st.selectbox(
        "Chọn học sinh",
        ds_hs,
        key="gv_result_student"
    )

    lich_su = [
        x
        for x in ds
        if str(
            x.get(
                "hoc_sinh_id",
                ""
            )
        ) == hs
    ]

    profile = tao_ho_so_tu_lich_su(
        hs
    )

    r1, r2, r3 = st.columns(3)

    with r1:
        st.metric(
            "Số lượt",
            len(
                lich_su
            )
        )

    with r2:
        st.metric(
            "Tỉ lệ đúng tích lũy",
            f"{profile.get('ti_le_dung', 0) * 100:.0f}%"
        )

    with r3:
        st.metric(
            "Đơn vị đánh giá",
            profile.get(
                "tong_don_vi",
                0
            )
        )

    st.markdown(
        "### 🧠 Các điểm cần ưu tiên"
    )

    weak = tom_tat_diem_yeu(
        profile,
        10
    )

    if weak:
        st.dataframe(
            pd.DataFrame([
                {
                    "YCCĐ": x.get(
                        "yccd",
                        ""
                    ),
                    "Mức độ": x.get(
                        "muc_do",
                        ""
                    ),
                    "Năng lực": x.get(
                        "nang_luc",
                        ""
                    ),
                    "Số lần": x.get(
                        "so_lan",
                        0
                    ),
                    "Số đúng": x.get(
                        "so_dung",
                        0
                    ),
                    "Tỉ lệ đúng": (
                        f"{x.get('ti_le_dung', 0) * 100:.0f}%"
                    )
                }
                for x in weak
            ]),
            use_container_width=True,
            hide_index=True
        )

    st.markdown(
        "### 📚 Lịch sử lượt luyện"
    )

    st.dataframe(
        pd.DataFrame([
            {
                "Thời gian": x.get(
                    "thoi_gian",
                    ""
                ),
                "Chế độ": x.get(
                    "che_do",
                    ""
                ),
                "Tên lượt": x.get(
                    "ten_luot",
                    ""
                ),
                "Số câu": x.get(
                    "tong_so_cau",
                    0
                ),
                "Tỉ lệ đúng": (
                    f"{x.get('ti_le_dung_don_vi', 0):.0f}%"
                ),
                "Điểm": x.get(
                    "diem",
                    0
                )
            }
            for x in reversed(
                lich_su
            )
        ]),
        use_container_width=True,
        hide_index=True
    )




# ==========================================================
# TẠO ĐỀ TỪ NGÂN HÀNG
# ==========================================================
def doc_json_list(path):
    # Hai nhóm có bảng riêng để nhiều học sinh có thể ghi đồng thời.
    if os.path.abspath(path) == os.path.abspath(HS_HISTORY_PATH):
        return _doc_attempts_shared()
    if os.path.abspath(path) == os.path.abspath(STUDENT_PATH):
        return _doc_students_shared()

    data = _doc_document_shared(path, [])
    return data if isinstance(data, list) else []


def luu_json_list(path, data):
    if os.path.abspath(path) == os.path.abspath(HS_HISTORY_PATH):
        return _luu_attempts_shared(data)
    if os.path.abspath(path) == os.path.abspath(STUDENT_PATH):
        return _luu_students_shared(data)
    return _luu_document_shared(path, data)


# ==========================================================
# ĐỢT KIỂM TRA THEO MA TRẬN: LỊCH, 1 LẦN/HS, MỞ ĐÁP ÁN
# ==========================================================
VN_TZ = timezone(timedelta(hours=7))


def bay_gio_viet_nam():
    return datetime.now(VN_TZ)


def parse_iso_vn(value):
    value = str(value or "").strip()
    if not value:
        return None
    try:
        dt = datetime.fromisoformat(value)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=VN_TZ)
        return dt.astimezone(VN_TZ)
    except Exception:
        return None


def fmt_vn_datetime(value):
    dt = value if isinstance(value, datetime) else parse_iso_vn(value)
    if not dt:
        return "—"
    return dt.astimezone(VN_TZ).strftime("%d/%m/%Y %H:%M")


def doc_dot_kiem_tra_ma_tran():
    return doc_json_list(MATRIX_TEST_PATH)


def luu_dot_kiem_tra_ma_tran(data):
    luu_json_list(MATRIX_TEST_PATH, data)


def tim_dot_kiem_tra_theo_id(dot_id):
    dot_id = str(dot_id or "").strip()
    for item in doc_dot_kiem_tra_ma_tran():
        if str(item.get("id", "")).strip() == dot_id:
            return item
    return None


def han_cuoi_bat_dau_dot_kiem_tra(dot):
    """
    Hạn cuối HS được phép BẮT ĐẦU = thời điểm kết thúc đợt - thời gian làm bài.
    Nhờ vậy mọi HS bắt đầu hợp lệ đều có đủ đúng số phút GV quy định.
    """
    dong = parse_iso_vn(dot.get("dong_luc_iso", ""))
    thoi_gian_phut = int(dot.get("thoi_gian_lam_phut", 0) or 0)
    if not dong or thoi_gian_phut <= 0:
        return None
    return dong - timedelta(minutes=thoi_gian_phut)


def trang_thai_dot_kiem_tra(dot, now=None):
    now = now or bay_gio_viet_nam()
    mo = parse_iso_vn(dot.get("mo_tu_iso", ""))
    dong = parse_iso_vn(dot.get("dong_luc_iso", ""))
    han_bat_dau = han_cuoi_bat_dau_dot_kiem_tra(dot)

    if not mo or not dong or not han_bat_dau:
        return "Chưa cấu hình", False
    if dong <= mo or han_bat_dau < mo:
        return "Cấu hình chưa hợp lệ", False
    if now < mo:
        return "Sắp mở", False
    if now > dong:
        return "Đã kết thúc", False
    if now > han_bat_dau:
        return "Hết giờ vào", False
    return "Đang mở", True


def dot_kiem_tra_phu_hop_hoc_sinh(dot, hs_lop, hs_khoi=""):
    lop_ap_dung = str(dot.get("lop_ap_dung", "Tất cả") or "Tất cả").strip()
    khoi_dot = str(dot.get("khoi", "") or "").strip()
    if lop_ap_dung not in {"", "Tất cả"} and lop_ap_dung != str(hs_lop or "").strip():
        return False
    if khoi_dot and hs_khoi and khoi_dot != str(hs_khoi).strip():
        return False
    return True


def tim_luot_kiem_tra_da_nop(hoc_sinh_id, dot_id):
    ma = str(hoc_sinh_id or "").strip().upper()
    dot_id = str(dot_id or "").strip()
    if not ma or not dot_id:
        return None
    for lan in reversed(doc_lich_su_hoc_sinh()):
        if str(lan.get("hoc_sinh_id", "")).strip().upper() != ma:
            continue
        pham_vi = lan.get("pham_vi", {}) or {}
        if str(pham_vi.get("dot_kiem_tra_id", "")).strip() == dot_id:
            return lan
    return None


def seed_dot_kiem_tra_hoc_sinh(dot, hoc_sinh_id):
    raw = (
        f"{str(dot.get('id', ''))}|"
        f"{str(dot.get('mau_id', ''))}|"
        f"PB{int(dot.get('mau_phien_ban', 1) or 1)}|"
        f"{str(hoc_sinh_id).strip().upper()}|DOT_KIEM_TRA_V1"
    )
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def tinh_diem_ma_tran_theo_don_gia(ma_tran):
    tong = 0.0
    for dong in ma_tran or []:
        try:
            n = int(dong.get("Số câu", 0) or 0)
        except Exception:
            n = 0
        dang = str(dong.get("Dạng câu hỏi", "")).strip()
        if dang == "Trắc nghiệm 4 lựa chọn":
            tong += n * 0.25
        elif dang == "Đúng / Sai":
            tong += n * 1.0
        elif dang == "Trả lời ngắn":
            tong += n * 0.50
    return round(tong, 2)


def khoa_bai_lam_chinh_thuc(lan):
    pham_vi = lan.get("pham_vi", {}) or {}
    dot_id = str(pham_vi.get("dot_kiem_tra_id", "")).strip()
    if dot_id:
        return "DOT:" + dot_id
    de_id = str(pham_vi.get("de_id", "")).strip()
    if de_id:
        return "DE:" + de_id
    mau_id = str(pham_vi.get("mau_id", "")).strip()
    if mau_id:
        pb = str(pham_vi.get("mau_phien_ban", ""))
        return f"MAU:{mau_id}:PB{pb}"
    return "TEN:" + str(lan.get("ten_luot", "")).strip()


def nhan_bai_lam_chinh_thuc(lan):
    pham_vi = lan.get("pham_vi", {}) or {}
    if pham_vi.get("dot_kiem_tra_id"):
        return str(pham_vi.get("ten_dot_kiem_tra", lan.get("ten_luot", "Kiểm tra")))
    if pham_vi.get("de_id"):
        ma_de = str(pham_vi.get("ma_de", "")).strip()
        return str(lan.get("ten_luot", "Đề GV")) + (f" • {ma_de}" if ma_de else "")
    return str(lan.get("ten_luot", "Bài làm"))


def metadata_don_vi_cau(q):
    """
    Trả các đơn vị metadata của câu.
    Câu thường: một đơn vị.
    Đúng/Sai: 4 đơn vị a-d.
    """
    if q.get("dang_cau") == "Đúng / Sai":
        ds = []
        for nd in q.get("nhan_dinh_meta", []) or []:
            ds.append({
                "yccd": str(nd.get("yccd", "")).strip(),
                "muc_do": str(nd.get("muc_do", "")).strip(),
                "nang_luc": str(
                    nd.get("thanh_phan_nang_luc", "")
                ).strip()
            })
        return ds

    return [{
        "yccd": str(q.get("yccd", "")).strip(),
        "muc_do": str(q.get("muc_do", "")).strip(),
        "nang_luc": str(
            q.get("thanh_phan_nang_luc", "")
        ).strip()
    }]


# ==========================================================
# QUY TẮC CHẤM ĐIỂM HỌC SINH THEO TỪNG CHẾ ĐỘ
# ==========================================================
CHE_DO_DE_GV = "📝 Đề GV / ma trận"
CHE_DO_KIEM_TRA_MA_TRAN = "🧪 Kiểm tra theo ma trận GV"
CHE_DO_TOT_NGHIEP = "🎓 Luyện tốt nghiệp THPT"

# Đề tốt nghiệp vẫn ưu tiên Ngân hàng tốt nghiệp. Chỉ cho phép bổ sung
# một lượng nhỏ câu phù hợp từ Ngân hàng ôn tập để tăng độ phong phú.
GRAD_MAX_CAU_TU_NH_ON_TAP = 4
GRAD_QUOTA_ON_TAP_THEO_DANG = {
    "Trắc nghiệm 4 lựa chọn": 2,
    "Đúng / Sai": 1,
    "Trả lời ngắn": 1,
}


def la_che_do_cham_diem_co_dinh(che_do):
    return che_do in {
        CHE_DO_DE_GV,
        CHE_DO_KIEM_TRA_MA_TRAN,
        CHE_DO_TOT_NGHIEP
    }


def thang_diem_chinh_thuc_hs(che_do, diem_toi_da_thuc_te=None):
    # Ôn theo đề/ma trận GV: thang điểm đi theo đúng ma trận GV đã chọn,
    # không khóa cứng 7 điểm.
    if che_do == CHE_DO_DE_GV:
        if diem_toi_da_thuc_te is not None:
            try:
                return round(float(diem_toi_da_thuc_te), 2)
            except Exception:
                pass
        return 10.0

    if che_do in {
        CHE_DO_KIEM_TRA_MA_TRAN,
        CHE_DO_TOT_NGHIEP
    }:
        return 10.0

    return 10.0


def diem_toi_da_cua_cau_hs(q, che_do):
    dang = str(q.get("dang_cau", "")).strip()

    if che_do == CHE_DO_TOT_NGHIEP:
        if dang == "Đúng / Sai":
            return 1.0
        if dang in {"Trắc nghiệm 4 lựa chọn", "Trả lời ngắn"}:
            return 0.25
        return 0.0

    if che_do in {CHE_DO_DE_GV, CHE_DO_KIEM_TRA_MA_TRAN}:
        if dang == "Đúng / Sai":
            return 1.0  # 4 ý x 0,25
        if dang == "Trắc nghiệm 4 lựa chọn":
            return 0.25
        if dang == "Trả lời ngắn":
            return 0.50
        return 0.0

    # Các chế độ luyện cá nhân hóa cũ: giữ nguyên cách tính nội bộ
    # rồi quy đổi về thang 10 như trước.
    return 1.0


def tinh_tong_diem_toi_da_hs(ds_cau, che_do):
    return round(
        sum(
            diem_toi_da_cua_cau_hs(q, che_do)
            for q in (ds_cau or [])
        ),
        2
    )


def tinh_diem_cau_hs(
    dang,
    dung_toan_cau,
    so_y_dung,
    che_do
):
    """
    Trả điểm THỰC của câu theo đúng chế độ.
    - Đề GV/ma trận ôn tập: 0,25 MCQ; 0,25/ý Đ-S; 0,5 TLN.
    - Kiểm tra ma trận: cùng đơn giá trên, tổng chuẩn 10.
    - Tốt nghiệp: theo cách tính Bộ GDĐT.
    - Các chế độ luyện cá nhân hóa khác: giữ cách cũ.
    """
    dang = str(dang or "").strip()
    so_y_dung = int(so_y_dung or 0)

    if che_do == CHE_DO_TOT_NGHIEP:
        if dang == "Đúng / Sai":
            return {
                4: 1.00,
                3: 0.50,
                2: 0.25,
                1: 0.10,
                0: 0.00
            }.get(so_y_dung, 0.0)

        if dang in {"Trắc nghiệm 4 lựa chọn", "Trả lời ngắn"}:
            return 0.25 if dung_toan_cau else 0.0

        return 0.0

    if che_do in {CHE_DO_DE_GV, CHE_DO_KIEM_TRA_MA_TRAN}:
        if dang == "Đúng / Sai":
            return so_y_dung * 0.25

        if dang == "Trắc nghiệm 4 lựa chọn":
            return 0.25 if dung_toan_cau else 0.0

        if dang == "Trả lời ngắn":
            return 0.50 if dung_toan_cau else 0.0

        return 0.0

    # Giữ nguyên cách chấm cũ cho luyện cá nhân hóa.
    if dang == "Đúng / Sai":
        return so_y_dung * 0.25

    return 1.0 if dung_toan_cau else 0.0


def seed_kiem_tra_ma_tran_hoc_sinh(mau, hoc_sinh_id):
    """
    Mỗi học sinh nhận một mã đề ổn định theo cùng ma trận.
    Không dùng hồ sơ năng lực, điểm yếu hay lịch sử luyện tập.
    """
    mau_id = str(
        mau.get("id", "")
        or mau.get("ten_mau", "")
        or json.dumps(mau.get("ma_tran", []), ensure_ascii=False, sort_keys=True)
    )
    phien_ban = int(mau.get("phien_ban", 1) or 1)
    raw = (
        f"{mau_id}|PB{phien_ban}|{str(hoc_sinh_id).strip().upper()}|"
        "KIEM_TRA_MA_TRAN_V1"
    )
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def cau_khop_dac_ta(q, spec):
    if q.get("trang_thai") == "Ngừng sử dụng":
        return False

    if (
        spec.get("Khối")
        and q.get("khoi") != spec.get("Khối")
    ):
        return False

    if (
        spec.get("Chương")
        and spec.get("Chương") != "Tất cả"
        and q.get("chuong") != spec.get("Chương")
    ):
        return False

    if (
        spec.get("Bài")
        and spec.get("Bài") != "Tất cả"
        and q.get("bai") != spec.get("Bài")
    ):
        return False

    if (
        spec.get("Dạng câu hỏi")
        and spec.get("Dạng câu hỏi") != "Tất cả"
        and q.get("dang_cau") != spec.get("Dạng câu hỏi")
    ):
        return False

    yccd = spec.get("YCCĐ", "Tất cả")
    muc = spec.get("Mức độ", "Tất cả")
    nl = spec.get("Thành phần năng lực", "Tất cả")

    # Đúng/Sai khớp nếu có ít nhất một ý thỏa đồng thời
    # tất cả điều kiện metadata được chọn.
    for unit in metadata_don_vi_cau(q):
        # Câu đề thật chỉ bắt buộc Khối/Chương/Bài; không ép YCCĐ chi tiết.
        # Vì vậy khi GV chọn YCCĐ, câu đề thật vẫn được xét nếu đã khớp Bài/Chương
        # và metadata mức độ/năng lực; câu ngân hàng thường vẫn khớp YCCĐ như cũ.
        la_de_that = str(q.get("muc_dich_su_dung", "")) == "tot_nghiep" or bool(q.get("nguon_file"))
        y_unit = str(unit.get("yccd", "") or "")
        y_noi_bo = (not y_unit) or y_unit.startswith("Kiến thức trọng tâm")
        ok_y = (
            yccd in ["", "Tất cả"]
            or unit["yccd"] == yccd
            or (la_de_that and y_noi_bo)
        )
        ok_m = (
            muc in ["", "Tất cả"]
            or unit["muc_do"] == muc
        )
        ok_n = (
            nl in ["", "Tất cả"]
            or unit["nang_luc"] == nl
        )

        if ok_y and ok_m and ok_n:
            return True

    return False


def rut_de_theo_ma_tran(bank, specs, seed=None):
    rng = random.Random(
        seed or str(uuid.uuid4())
    )

    da_chon_id = set()
    ds_de = []
    thieu = []

    for i, spec in enumerate(specs, start=1):
        can = int(
            spec.get("Số câu", 0)
        )

        pool = [
            q
            for q in bank
            if cau_khop_dac_ta(q, spec)
            and str(
                q.get("id", "")
            ) not in da_chon_id
        ]

        rng.shuffle(pool)

        lay = pool[:can]

        for q in lay:
            ds_de.append(q)
            da_chon_id.add(
                str(
                    q.get("id", "")
                )
            )

        if len(lay) < can:
            thieu.append({
                "Dòng": i,
                "Bài": spec.get("Bài", ""),
                "YCCĐ": spec.get("YCCĐ", ""),
                "Mức độ": spec.get("Mức độ", ""),
                "Năng lực": spec.get(
                    "Thành phần năng lực",
                    ""
                ),
                "Dạng": spec.get(
                    "Dạng câu hỏi",
                    ""
                ),
                "Cần": can,
                "Có thể lấy": len(lay),
                "Thiếu": can - len(lay)
            })

    rng.shuffle(ds_de)

    return ds_de, thieu


def luu_mau_de(mau):
    ds = doc_json_list(
        EXAM_TEMPLATE_PATH
    )

    ds.append(mau)

    luu_json_list(
        EXAM_TEMPLATE_PATH,
        ds
    )


def luu_de_da_tao(de):
    ds = doc_json_list(
        EXAM_PATH
    )

    ds.append(de)

    luu_json_list(
        EXAM_PATH,
        ds
    )


def hien_thi_de_xem_truoc(de):
    if not de:
        return

    st.markdown("---")
    st.subheader(
        f"📝 {de.get('ten_de', 'Đề xem trước')}"
    )

    st.caption(
        f"Mã đề: {de.get('ma_de', '')} • "
        f"{len(de.get('cau_hoi', []))} câu • "
        f"Thời gian: {de.get('thoi_gian', '')} phút"
    )

    for i, q in enumerate(
        de.get("cau_hoi", []),
        start=1
    ):
        with st.expander(
            f"Câu {i} • {q.get('dang_cau', '')}",
            expanded=False
        ):
            if q.get("tinh_huong"):
                st.write(
                    "**Tình huống / dữ liệu:**",
                    q.get("tinh_huong", "")
                )

            if q.get("tai_nguyen_truc_quan") or q.get("du_lieu_truc_quan"):
                hien_thi_tai_nguyen_cau_tot_nghiep(q)
            else:
                hien_thi_du_lieu_truc_quan_cau(q)

            st.write(
                "**Câu hỏi:**",
                q.get("cau_hoi", "")
            )

            if q.get(
                "dang_cau"
            ) == "Trắc nghiệm 4 lựa chọn":
                for lc in q.get(
                    "lua_chon",
                    []
                ):
                    st.write(lc)

            elif q.get(
                "dang_cau"
            ) == "Đúng / Sai":
                for ky, nd in zip(
                    ["a", "b", "c", "d"],
                    q.get(
                        "nhan_dinh_meta",
                        []
                    )
                ):
                    st.write(
                        f"**{ky})** "
                        f"{nd.get('noi_dung', '')}"
                    )

            st.caption(
                f"{q.get('bai', '')}"
            )


def tao_de_giao_vien():
    st.header(
        "📝 TẠO ĐỀ TỪ NGÂN HÀNG"
    )

    # Nguồn dùng chung khi GV ra đề: ngân hàng ôn tập + câu đề thật đã đủ điều kiện.
    # Câu đề thật vẫn giữ nguyên hình/bảng và metadata Khối → Chương → Bài.
    bank_main = [
        q for q in doc_ngan_hang()
        if q.get("trang_thai", "Đã duyệt") != "Ngừng sử dụng"
    ]
    bank_grad = [
        q for q in doc_ngan_hang_tot_nghiep_thuc_te()
        if cau_tot_nghiep_du_dieu_kien_su_dung(q)
    ]
    bank = []
    seen_ids = set()
    for q in bank_main + bank_grad:
        key = str(q.get("id", "") or fingerprint_cau_hoi(q))
        if key in seen_ids:
            continue
        seen_ids.add(key)
        bank.append(q)

    if not bank:
        st.info(
            "Ngân hàng chưa có câu đã duyệt để tạo đề."
        )
        return

    st.caption(
        "Chỉ rút câu đã duyệt trong ngân hàng. "
        "Nếu một ô ma trận thiếu câu, hệ thống báo thiếu thay vì lấy sai mục tiêu."
    )

    tab1, tab2, tab3 = st.tabs([
        "① Theo ma trận / đặc tả GV",
        "② Luyện đề tốt nghiệp THPT",
        "③ Mẫu đề đã lưu"
    ])

    # ======================================================
    # TAB 1 - MA TRẬN GV
    # ======================================================
    with tab1:
        st.subheader(
            "🎯 Ma trận / bản đặc tả của giáo viên"
        )

        ten_de = st.text_input(
            "Tên đề",
            value="Đề ôn tập / kiểm tra",
            key="exam_custom_name"
        )

        c0, c1, c2 = st.columns(3)

        ds_khoi = sorted({
            q.get("khoi", "")
            for q in bank
            if q.get("khoi")
        })

        with c0:
            khoi = st.selectbox(
                "Khối",
                ds_khoi,
                key="exam_custom_grade"
            )

        with c1:
            loai_kiem_tra = st.selectbox(
                "Mục đích",
                [
                    "Ôn tập",
                    "Kiểm tra thường xuyên",
                    "Kiểm tra giữa kì",
                    "Kiểm tra cuối kì"
                ],
                key="exam_custom_type"
            )

        with c2:
            thoi_gian = st.number_input(
                "Thời gian (phút)",
                min_value=5,
                max_value=180,
                value=45,
                step=5,
                key="exam_custom_time"
            )

        bank_khoi = [
            q for q in bank
            if q.get("khoi") == khoi
        ]

        ds_chuong = sorted({
            q.get("chuong", "")
            for q in bank_khoi
            if q.get("chuong")
        })

        ch = st.selectbox(
            "Chương",
            ["Tất cả"] + ds_chuong,
            key="exam_spec_chuong"
        )

        bank_ch = (
            bank_khoi
            if ch == "Tất cả"
            else [
                q for q in bank_khoi
                if q.get("chuong") == ch
            ]
        )

        ds_bai = sorted({
            q.get("bai", "")
            for q in bank_ch
            if q.get("bai")
        })

        bai = st.selectbox(
            "Bài",
            ["Tất cả"] + ds_bai,
            key="exam_spec_bai"
        )

        # YCCĐ lấy từ ngân hàng + metadata từng ý Đ/S
        ds_yccd = set()

        for q in bank_ch:
            if (
                bai != "Tất cả"
                and q.get("bai") != bai
            ):
                continue

            for unit in metadata_don_vi_cau(q):
                if unit["yccd"]:
                    ds_yccd.add(
                        unit["yccd"]
                    )

        c3, c4, c5 = st.columns(3)

        with c3:
            yccd = st.selectbox(
                "YCCĐ",
                ["Tất cả"] + sorted(ds_yccd),
                key="exam_spec_yccd"
            )

        with c4:
            muc = st.selectbox(
                "Mức độ",
                [
                    "Tất cả",
                    "Nhận biết",
                    "Thông hiểu",
                    "Vận dụng"
                ],
                key="exam_spec_level"
            )

        with c5:
            nl = st.selectbox(
                "Thành phần năng lực",
                ["Tất cả"] + THANH_PHAN_NANG_LUC,
                key="exam_spec_comp"
            )

        c6, c7 = st.columns(2)

        with c6:
            dang = st.selectbox(
                "Dạng câu hỏi",
                [
                    "Trắc nghiệm 4 lựa chọn",
                    "Đúng / Sai",
                    "Trả lời ngắn"
                ],
                key="exam_spec_type"
            )

        with c7:
            so_cau = st.number_input(
                "Số câu ở dòng này",
                min_value=1,
                max_value=100,
                value=1,
                step=1,
                key="exam_spec_count"
            )

        if st.button(
            "➕ THÊM VÀO MA TRẬN",
            type="secondary",
            use_container_width=True,
            key="exam_add_spec"
        ):
            st.session_state.ma_tran_de_gv.append({
                "Khối": khoi,
                "Chương": ch,
                "Bài": bai,
                "YCCĐ": yccd,
                "Mức độ": muc,
                "Thành phần năng lực": nl,
                "Dạng câu hỏi": dang,
                "Số câu": int(so_cau)
            })
            st.rerun()

        if st.session_state.ma_tran_de_gv:
            st.markdown(
                "#### 📋 Ma trận hiện tại"
            )

            df_mt = pd.DataFrame(
                st.session_state.ma_tran_de_gv
            )

            st.dataframe(
                df_mt,
                use_container_width=True,
                hide_index=True
            )

            tong = sum(
                int(x["Số câu"])
                for x in st.session_state.ma_tran_de_gv
            )

            st.info(
                f"Tổng số câu dự kiến: **{tong}**"
            )

            # Tính trước điểm tối đa theo đúng đơn giá đề GV/ma trận:
            # 0,25/câu 4LC; 0,25/ý Đ-S; 0,50/câu TLN.
            diem_mt = 0.0
            for dong_mt in st.session_state.ma_tran_de_gv:
                n_mt = int(dong_mt.get("Số câu", 0) or 0)
                dang_mt = dong_mt.get("Dạng câu hỏi", "")
                if dang_mt == "Trắc nghiệm 4 lựa chọn":
                    diem_mt += n_mt * 0.25
                elif dang_mt == "Đúng / Sai":
                    diem_mt += n_mt * 1.00
                elif dang_mt == "Trả lời ngắn":
                    diem_mt += n_mt * 0.50

            st.caption(
                f"Tổng điểm của ma trận hiện tại theo đơn giá đã quy ước: **{diem_mt:.2f} điểm**. "
                "Khi HS **ôn theo đề/ma trận GV**, app chấm đúng trên chính tổng điểm này; "
                "không khóa cứng 7 điểm. Với **kiểm tra theo ma trận GV**, ma trận chính thức vẫn cần đủ 10,0 điểm."
            )

            st.info(
                "ℹ️ **Rút 1 mã đề để xem/lưu**: ma trận chỉ là khung yêu cầu. "
                "Nút này lấy ngẫu nhiên một bộ câu cụ thể đúng từng ô ma trận để GV xem trước, lưu hoặc in. "
                "Khi **kiểm tra theo ma trận**, GV không cần rút sẵn cho từng HS; app sẽ tự rút mã đề khác nhau cho từng em."
            )

            a1, a2, a3 = st.columns(3)

            with a1:
                if st.button(
                    "🎲 RÚT 1 MÃ ĐỀ ĐỂ XEM / LƯU",
                    type="primary",
                    use_container_width=True,
                    key="exam_draw_custom"
                ):
                    seed = str(
                        uuid.uuid4()
                    )

                    ds_cau, thieu = (
                        rut_de_theo_ma_tran(
                            bank,
                            st.session_state.ma_tran_de_gv,
                            seed=seed
                        )
                    )

                    de = {
                        "id": str(uuid.uuid4()),
                        "ma_de": seed[:8].upper(),
                        "ten_de": ten_de,
                        "loai": loai_kiem_tra,
                        "khoi": khoi,
                        "thoi_gian": int(thoi_gian),
                        "ngay_tao": datetime.now().strftime(
                            "%d/%m/%Y %H:%M"
                        ),
                        "ma_tran": list(
                            st.session_state.ma_tran_de_gv
                        ),
                        "cau_hoi": ds_cau,
                        "thieu": thieu
                    }

                    st.session_state.de_xem_truoc = de

                    if thieu:
                        st.warning(
                            "Ngân hàng chưa đủ một số ô ma trận. "
                            "Xem bảng thiếu phía dưới."
                        )
                    else:
                        st.success(
                            "Đã rút đủ đề theo ma trận."
                        )

            with a2:
                if st.button(
                    "💾 LƯU MẪU MA TRẬN",
                    use_container_width=True,
                    key="exam_save_custom_template"
                ):
                    luu_mau_de({
                        "id": str(uuid.uuid4()),
                        "ten_mau": ten_de,
                        "che_do": "Ma trận GV",
                        "phien_ban": 1,
                        "khoi": khoi,
                        "thoi_gian": int(thoi_gian),
                        "ma_tran": list(
                            st.session_state.ma_tran_de_gv
                        ),
                        "ngay_tao": datetime.now().strftime(
                            "%d/%m/%Y %H:%M"
                        )
                    })
                    st.success(
                        "Đã lưu mẫu ma trận."
                    )

            with a3:
                if st.button(
                    "🗑 XÓA MA TRẬN",
                    use_container_width=True,
                    key="exam_clear_matrix"
                ):
                    st.session_state.ma_tran_de_gv = []
                    st.session_state.de_xem_truoc = None
                    st.rerun()

            de_custom = st.session_state.de_xem_truoc

            if (
                de_custom
                and de_custom.get("loai")
                != "Luyện tốt nghiệp THPT"
            ):
                if de_custom.get("thieu"):
                    st.dataframe(
                        pd.DataFrame(
                            de_custom["thieu"]
                        ),
                        use_container_width=True,
                        hide_index=True
                    )
                else:
                    if st.button(
                        "✅ LƯU ĐỀ NÀY",
                        use_container_width=True,
                        key="exam_save_custom_exam"
                    ):
                        luu_de_da_tao(
                            de_custom
                        )
                        st.success(
                            "Đã lưu đề."
                        )

                hien_thi_de_xem_truoc(
                    de_custom
                )

        else:
            st.info(
                "Thêm ít nhất một dòng vào ma trận để bắt đầu."
            )

    # ======================================================
    # TAB 2 - TỐT NGHIỆP
    # ======================================================
    with tab2:
        st.subheader("🎓 Luyện đề tốt nghiệp THPT từ đề thật")
        st.caption(
            "Form được khóa đúng **18 câu 4 lựa chọn + 4 câu Đúng/Sai + 6 câu Trả lời ngắn**, "
            "thời gian 50 phút. Nguồn chính là Ngân hàng tốt nghiệp GV đã nhập; app có thể chen tối đa một ít câu phù hợp "
            "từ Ngân hàng ôn tập Khối 12 và không gọi AI để viết lại câu khi rút đề."
        )

        bank12 = [
            q for q in doc_ngan_hang_tot_nghiep_thuc_te()
            if cau_tot_nghiep_du_dieu_kien_su_dung(q)
        ]
        bank_on_tap_12 = [
            q for q in doc_ngan_hang()
            if cau_on_tap_bo_sung_du_dieu_kien_tot_nghiep(q)
        ]

        if not bank12:
            st.warning(
                "Ngân hàng tốt nghiệp từ đề thật chưa có câu dùng được. "
                "Hãy vào **🎓 Xây dựng NH tốt nghiệp** để nhập đề trước."
            )
        else:
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("4 lựa chọn", sum(q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn" for q in bank12))
            c2.metric("Đúng/Sai", sum(q.get("dang_cau") == "Đúng / Sai" for q in bank12))
            c3.metric("Trả lời ngắn", sum(q.get("dang_cau") == "Trả lời ngắn" for q in bank12))
            c4.metric("Nguồn đề", len({q.get("nguon_file", "") for q in bank12 if q.get("nguon_file")}))

            st.info(
                "Mỗi mã đề mới vẫn **ưu tiên câu từ Ngân hàng tốt nghiệp**, trộn nhiều file nguồn và không lặp câu. "
                f"App có thể bổ sung tối đa **{GRAD_MAX_CAU_TU_NH_ON_TAP}/28 câu** phù hợp từ Ngân hàng ôn tập Khối 12 "
                "để tăng độ phong phú; phần còn lại vẫn lấy từ kho tốt nghiệp."
            )

            specs_grad = [
                {"Dạng câu hỏi": "Trắc nghiệm 4 lựa chọn", "Số câu": 18},
                {"Dạng câu hỏi": "Đúng / Sai", "Số câu": 4},
                {"Dạng câu hỏi": "Trả lời ngắn", "Số câu": 6},
            ]
            tg_grad = 50

            if st.button(
                "🎲 TẠO MÃ ĐỀ MỚI 18/4/6",
                type="primary",
                use_container_width=True,
                key="grad_generate",
            ):
                seed = str(uuid.uuid4())
                ds_cau, thieu = rut_de_tot_nghiep_tu_de_that(
                    bank12,
                    seed=seed,
                    bank_on_tap=bank_on_tap_12
                )
                de = {
                    "id": str(uuid.uuid4()),
                    "ma_de": seed[:8].upper(),
                    "ten_de": "Luyện đề tốt nghiệp THPT – Sinh học",
                    "loai": "Luyện tốt nghiệp THPT",
                    "khoi": "Khối 12",
                    "thoi_gian": tg_grad,
                    "ngay_tao": datetime.now().strftime("%d/%m/%Y %H:%M"),
                    "ma_tran": specs_grad,
                    "cau_hoi": ds_cau,
                    "thieu": thieu,
                    "nguon_de_that": True,
                }
                st.session_state.de_xem_truoc = de
                if thieu:
                    st.warning("Ngân hàng chưa đủ câu không trùng để tạo trọn form 18/4/6.")
                else:
                    st.success("Đã tạo một mã đề mới từ các câu gốc trong ngân hàng tốt nghiệp.")

            de_grad = st.session_state.de_xem_truoc
            if de_grad and de_grad.get("loai") == "Luyện tốt nghiệp THPT":
                if de_grad.get("thieu"):
                    st.dataframe(pd.DataFrame(de_grad["thieu"]), use_container_width=True, hide_index=True)
                else:
                    if st.button(
                        "✅ LƯU MÃ ĐỀ NÀY",
                        use_container_width=True,
                        key="grad_save_exam",
                    ):
                        luu_de_da_tao(de_grad)
                        st.success("Đã lưu mã đề.")
                hien_thi_de_xem_truoc(de_grad)

    # ======================================================
    # TAB 3 - QUẢN LÝ MẪU / ĐỀ ĐÃ LƯU
    # ======================================================
    with tab3:
        st.subheader("🗂️ Quản lý mẫu ma trận và đề đã lưu")
        st.caption(
            "GV có thể **xem – sửa – xóa** mẫu ma trận sau khi học sinh đã ôn/kiểm tra xong. "
            "Xóa mẫu/đề chỉ làm nó biến mất khỏi danh sách giao cho học sinh; "
            "**lịch sử và điểm học sinh đã nộp vẫn được giữ nguyên**."
        )

        ql_mau_tab, ql_de_tab, ql_kt_tab = st.tabs([
            "📋 Mẫu ma trận",
            "📝 Đề đã lưu",
            "🧪 Đợt kiểm tra"
        ])

        # --------------------------------------------------
        # QUẢN LÝ MẪU MA TRẬN
        # --------------------------------------------------
        with ql_mau_tab:
            ds_mau = doc_json_list(EXAM_TEMPLATE_PATH)

            if not ds_mau:
                st.info("Chưa có mẫu ma trận nào được lưu.")
            else:
                st.dataframe(
                    pd.DataFrame([
                        {
                            "Tên mẫu": x.get("ten_mau", ""),
                            "Khối": x.get("khoi", ""),
                            "Thời gian (phút)": x.get("thoi_gian", ""),
                            "Số dòng ma trận": len(x.get("ma_tran", []) or []),
                            "Ngày tạo": x.get("ngay_tao", ""),
                            "Cập nhật": x.get("ngay_cap_nhat", "")
                        }
                        for x in ds_mau
                    ]),
                    use_container_width=True,
                    hide_index=True
                )

                idx_mau = st.selectbox(
                    "Chọn mẫu để xem / sửa / xóa",
                    options=list(range(len(ds_mau))),
                    format_func=lambda i: (
                        f"{ds_mau[i].get('ten_mau', 'Không tên')} • "
                        f"{ds_mau[i].get('khoi', '')} • "
                        f"{len(ds_mau[i].get('ma_tran', []) or [])} dòng"
                    ),
                    key="exam_manage_template_pick"
                )

                mau_chon = dict(ds_mau[idx_mau])
                mau_id = str(mau_chon.get("id", idx_mau))

                with st.container(border=True):
                    st.markdown("### 👁️ Xem / chỉnh sửa mẫu")

                    ten_mau_sua = st.text_input(
                        "Tên mẫu",
                        value=str(mau_chon.get("ten_mau", "")),
                        key=f"exam_manage_template_name_{mau_id}"
                    )

                    ec1, ec2, ec3 = st.columns(3)
                    with ec1:
                        khoi_mau_sua = st.text_input(
                            "Khối",
                            value=str(mau_chon.get("khoi", "")),
                            key=f"exam_manage_template_grade_{mau_id}"
                        )
                    with ec2:
                        thoi_gian_mau_sua = st.number_input(
                            "Thời gian (phút)",
                            min_value=1,
                            max_value=300,
                            value=max(1, int(mau_chon.get("thoi_gian", 45) or 45)),
                            step=5,
                            key=f"exam_manage_template_time_{mau_id}"
                        )
                    with ec3:
                        st.metric(
                            "Số dòng ma trận",
                            len(mau_chon.get("ma_tran", []) or [])
                        )

                    st.markdown("**Ma trận / bản đặc tả**")
                    df_mau_sua = pd.DataFrame(
                        mau_chon.get("ma_tran", []) or [],
                        columns=[
                            "Khối", "Chương", "Bài", "YCCĐ", "Mức độ",
                            "Thành phần năng lực", "Dạng câu hỏi", "Số câu"
                        ]
                    )

                    df_mau_sua = st.data_editor(
                        df_mau_sua,
                        num_rows="dynamic",
                        use_container_width=True,
                        hide_index=True,
                        key=f"exam_manage_template_editor_{mau_id}",
                        column_config={
                            "Số câu": st.column_config.NumberColumn(
                                "Số câu", min_value=1, step=1
                            )
                        }
                    )

                    # Hiển thị điểm tối đa dự kiến ngay khi sửa.
                    diem_du_kien = 0.0
                    for _, r in df_mau_sua.iterrows():
                        try:
                            n = int(r.get("Số câu", 0) or 0)
                        except Exception:
                            n = 0
                        d = str(r.get("Dạng câu hỏi", "")).strip()
                        if d == "Trắc nghiệm 4 lựa chọn":
                            diem_du_kien += n * 0.25
                        elif d == "Đúng / Sai":
                            diem_du_kien += n * 1.00
                        elif d == "Trả lời ngắn":
                            diem_du_kien += n * 0.50

                    st.info(
                        f"Điểm tối đa theo đơn giá hiện hành: **{diem_du_kien:.2f} điểm**. "
                        "Ôn theo đề GV thường hướng đến 7 điểm tự động; kiểm tra theo ma trận hướng đến 10 điểm."
                    )

                    sm1, sm2 = st.columns(2)
                    with sm1:
                        if st.button(
                            "💾 LƯU THAY ĐỔI MẪU",
                            type="primary",
                            use_container_width=True,
                            key=f"exam_manage_template_save_{mau_id}"
                        ):
                            ma_tran_moi = []
                            for rec in df_mau_sua.to_dict("records"):
                                # Bỏ dòng hoàn toàn rỗng.
                                if not any(str(v).strip() for v in rec.values() if v is not None):
                                    continue
                                try:
                                    rec["Số câu"] = max(1, int(rec.get("Số câu", 1) or 1))
                                except Exception:
                                    rec["Số câu"] = 1
                                for field in [
                                    "Khối", "Chương", "Bài", "YCCĐ", "Mức độ",
                                    "Thành phần năng lực", "Dạng câu hỏi"
                                ]:
                                    rec[field] = str(rec.get(field, "") or "").strip()
                                ma_tran_moi.append(rec)

                            if not ma_tran_moi:
                                st.error("Mẫu phải có ít nhất một dòng ma trận.")
                            else:
                                mau_moi = dict(mau_chon)
                                mau_moi["ten_mau"] = ten_mau_sua.strip() or "Ma trận không tên"
                                mau_moi["khoi"] = khoi_mau_sua.strip()
                                mau_moi["thoi_gian"] = int(thoi_gian_mau_sua)
                                mau_moi["ma_tran"] = ma_tran_moi
                                mau_moi["phien_ban"] = int(mau_chon.get("phien_ban", 1) or 1) + 1
                                mau_moi["ngay_cap_nhat"] = datetime.now().strftime("%d/%m/%Y %H:%M")
                                ds_mau[idx_mau] = mau_moi
                                luu_json_list(EXAM_TEMPLATE_PATH, ds_mau)
                                st.success("Đã cập nhật mẫu ma trận. Học sinh sẽ thấy phiên bản mới ở lượt tiếp theo.")
                                st.rerun()

                    with sm2:
                        xac_nhan_xoa_mau = st.checkbox(
                            "Tôi xác nhận xóa mẫu này khỏi danh sách giao cho học sinh.",
                            key=f"exam_manage_template_delete_confirm_{mau_id}"
                        )
                        if st.button(
                            "🗑 XÓA MẪU MA TRẬN",
                            use_container_width=True,
                            disabled=not xac_nhan_xoa_mau,
                            key=f"exam_manage_template_delete_{mau_id}"
                        ):
                            mau_xoa_id = str(mau_chon.get("id", ""))
                            ds_mau.pop(idx_mau)
                            luu_json_list(EXAM_TEMPLATE_PATH, ds_mau)

                            # Xóa luôn các đợt kiểm tra còn đang tham chiếu mẫu này
                            # để HS không thể tiếp tục vào một đợt đã bị GV thu hồi.
                            if mau_xoa_id:
                                ds_dot_lien_quan = [
                                    d for d in doc_dot_kiem_tra_ma_tran()
                                    if str(d.get("mau_id", "")) != mau_xoa_id
                                ]
                                luu_dot_kiem_tra_ma_tran(ds_dot_lien_quan)

                            st.success(
                                "Đã xóa mẫu ma trận và thu hồi các đợt kiểm tra đang tham chiếu mẫu này. "
                                "Lịch sử/điểm học sinh đã nộp vẫn được giữ."
                            )
                            st.rerun()

        # --------------------------------------------------
        # QUẢN LÝ ĐỀ ĐÃ LƯU
        # --------------------------------------------------
        with ql_de_tab:
            ds_de = doc_json_list(EXAM_PATH)

            if not ds_de:
                st.info("Chưa có đề cụ thể nào được lưu.")
            else:
                st.dataframe(
                    pd.DataFrame([
                        {
                            "Tên đề": x.get("ten_de", ""),
                            "Mã đề": x.get("ma_de", ""),
                            "Loại": x.get("loai", ""),
                            "Khối": x.get("khoi", ""),
                            "Thời gian": x.get("thoi_gian", ""),
                            "Số câu": len(x.get("cau_hoi", []) or []),
                            "Ngày tạo": x.get("ngay_tao", "")
                        }
                        for x in ds_de
                    ]),
                    use_container_width=True,
                    hide_index=True
                )

                idx_de = st.selectbox(
                    "Chọn đề để xem / sửa thông tin / xóa",
                    options=list(range(len(ds_de))),
                    format_func=lambda i: (
                        f"{ds_de[i].get('ten_de', ds_de[i].get('ma_de', 'Không tên'))} • "
                        f"{ds_de[i].get('ma_de', '')}"
                    ),
                    key="exam_manage_exam_pick"
                )

                de_chon = dict(ds_de[idx_de])
                de_id = str(de_chon.get("id", idx_de))

                with st.container(border=True):
                    st.markdown("### 👁️ Xem / chỉnh sửa thông tin đề")

                    dc1, dc2, dc3 = st.columns(3)
                    with dc1:
                        ten_de_sua = st.text_input(
                            "Tên đề",
                            value=str(de_chon.get("ten_de", "")),
                            key=f"exam_manage_exam_name_{de_id}"
                        )
                    with dc2:
                        loai_de_sua = st.text_input(
                            "Loại / mục đích",
                            value=str(de_chon.get("loai", "")),
                            key=f"exam_manage_exam_type_{de_id}"
                        )
                    with dc3:
                        thoi_gian_de_sua = st.number_input(
                            "Thời gian (phút)",
                            min_value=1,
                            max_value=300,
                            value=max(1, int(de_chon.get("thoi_gian", 45) or 45)),
                            step=5,
                            key=f"exam_manage_exam_time_{de_id}"
                        )

                    st.caption(
                        "Để bảo toàn đáp án và dữ liệu bài đã giao, phần sửa đề đã lưu chỉ thay đổi **tên / loại / thời gian**. "
                        "Muốn thay câu hỏi, GV nên sửa ma trận rồi rút một mã đề mới."
                    )

                    sd1, sd2 = st.columns(2)
                    with sd1:
                        if st.button(
                            "💾 LƯU THÔNG TIN ĐỀ",
                            type="primary",
                            use_container_width=True,
                            key=f"exam_manage_exam_save_{de_id}"
                        ):
                            de_moi = dict(de_chon)
                            de_moi["ten_de"] = ten_de_sua.strip() or de_chon.get("ten_de", "Đề không tên")
                            de_moi["loai"] = loai_de_sua.strip()
                            de_moi["thoi_gian"] = int(thoi_gian_de_sua)
                            de_moi["ngay_cap_nhat"] = datetime.now().strftime("%d/%m/%Y %H:%M")
                            ds_de[idx_de] = de_moi
                            luu_json_list(EXAM_PATH, ds_de)
                            st.success("Đã cập nhật thông tin đề.")
                            st.rerun()

                    with sd2:
                        xac_nhan_xoa_de = st.checkbox(
                            "Tôi xác nhận xóa đề này khỏi danh sách học sinh.",
                            key=f"exam_manage_exam_delete_confirm_{de_id}"
                        )
                        if st.button(
                            "🗑 XÓA ĐỀ ĐÃ LƯU",
                            use_container_width=True,
                            disabled=not xac_nhan_xoa_de,
                            key=f"exam_manage_exam_delete_{de_id}"
                        ):
                            ds_de.pop(idx_de)
                            luu_json_list(EXAM_PATH, ds_de)
                            st.success(
                                "Đã xóa đề khỏi danh sách. Kết quả học sinh đã nộp trước đó không bị xóa."
                            )
                            st.rerun()

                with st.expander("📄 Xem nội dung đề", expanded=False):
                    hien_thi_de_xem_truoc(de_chon)


        # --------------------------------------------------
        # QUẢN LÝ ĐỢT KIỂM TRA THEO MA TRẬN
        # --------------------------------------------------
        with ql_kt_tab:
            st.subheader("🧪 Giao đợt kiểm tra theo ma trận")
            st.caption(
                "Mỗi **đợt kiểm tra** có mã riêng, lớp áp dụng, giờ mở, giờ kết thúc đợt, thời gian làm và quyền xem đáp án. "
                "Một học sinh chỉ được nộp **1 lần trong mỗi đợt**. Ma trận được chụp lại tại lúc giao để đề của cả đợt không thay đổi giữa chừng."
            )

            ds_mau_kt = [
                x for x in doc_json_list(EXAM_TEMPLATE_PATH)
                if x.get("ma_tran")
            ]
            ds_dot = doc_dot_kiem_tra_ma_tran()

            if not ds_mau_kt:
                st.info("Chưa có mẫu ma trận. Hãy lưu mẫu ma trận trước khi tạo đợt kiểm tra.")
            else:
                st.markdown("### ➕ Tạo đợt kiểm tra mới")

                idx_mau_kt = st.selectbox(
                    "Chọn ma trận dùng để kiểm tra",
                    options=list(range(len(ds_mau_kt))),
                    format_func=lambda i: (
                        f"{ds_mau_kt[i].get('ten_mau', 'Không tên')} • "
                        f"{ds_mau_kt[i].get('khoi', '')} • PB{int(ds_mau_kt[i].get('phien_ban', 1) or 1)}"
                    ),
                    key="exam_assignment_template_pick"
                )
                mau_giao = ds_mau_kt[idx_mau_kt]
                diem_mau_giao = tinh_diem_ma_tran_theo_don_gia(
                    mau_giao.get("ma_tran", []) or []
                )

                st.info(
                    f"Ma trận này có tổng điểm tự động **{diem_mau_giao:.2f}/10,00** theo đơn giá kiểm tra."
                )
                if abs(diem_mau_giao - 10.0) > 1e-9:
                    st.error(
                        "Chưa thể giao kiểm tra chính thức: ma trận phải đạt đúng **10,00 điểm**. "
                        "Hãy sửa ma trận trước."
                    )

                now_vn = bay_gio_viet_nam()
                start_default = now_vn + timedelta(minutes=5)
                end_default = start_default + timedelta(hours=2)

                ten_dot_moi = st.text_input(
                    "Tên đợt kiểm tra",
                    value=f"Kiểm tra – {mau_giao.get('ten_mau', 'Ma trận')}",
                    key="exam_assignment_name_new"
                )

                lop_options = ["Tất cả"] + lay_danh_sach_lop_tu_hoc_sinh()
                lop_dot_moi = st.selectbox(
                    "Lớp được làm bài",
                    lop_options,
                    key="exam_assignment_class_new"
                )

                t1, t2 = st.columns(2)
                with t1:
                    ngay_mo = st.date_input(
                        "Ngày mở",
                        value=start_default.date(),
                        key="exam_assignment_start_date_new"
                    )
                    gio_mo = st.time_input(
                        "Giờ mở",
                        value=start_default.time().replace(second=0, microsecond=0),
                        key="exam_assignment_start_time_new"
                    )
                with t2:
                    ngay_dong = st.date_input(
                        "Ngày kết thúc đợt",
                        value=end_default.date(),
                        key="exam_assignment_end_date_new"
                    )
                    gio_dong = st.time_input(
                        "Giờ kết thúc đợt",
                        value=end_default.time().replace(second=0, microsecond=0),
                        key="exam_assignment_end_time_new"
                    )

                thoi_gian_lam_moi = st.number_input(
                    "Thời gian làm bài của mỗi học sinh (phút)",
                    min_value=5,
                    max_value=300,
                    value=max(5, int(mau_giao.get("thoi_gian", 45) or 45)),
                    step=5,
                    key="exam_assignment_duration_new"
                )

                mo_preview = datetime.combine(ngay_mo, gio_mo).replace(tzinfo=VN_TZ)
                dong_preview = datetime.combine(ngay_dong, gio_dong).replace(tzinfo=VN_TZ)
                han_bat_dau_preview = dong_preview - timedelta(minutes=int(thoi_gian_lam_moi))
                cua_so_hop_le_moi = (
                    dong_preview > mo_preview
                    and han_bat_dau_preview >= mo_preview
                )

                if cua_so_hop_le_moi:
                    st.info(
                        "⏰ **Cách tính thời gian:** HS được bắt đầu từ "
                        f"**{mo_preview.strftime('%d/%m/%Y %H:%M')}** đến chậm nhất "
                        f"**{han_bat_dau_preview.strftime('%d/%m/%Y %H:%M')}**. "
                        f"Mỗi HS có đúng **{int(thoi_gian_lam_moi)} phút** kể từ lúc bấm Bắt đầu; "
                        f"toàn bộ đợt kết thúc lúc **{dong_preview.strftime('%d/%m/%Y %H:%M')}**."
                    )
                else:
                    st.error(
                        "Khoảng từ giờ mở đến giờ kết thúc đợt phải ít nhất bằng "
                        f"**{int(thoi_gian_lam_moi)} phút** để HS có đủ thời gian làm bài."
                    )

                if st.button(
                    "📤 GIAO ĐỢT KIỂM TRA",
                    type="primary",
                    use_container_width=True,
                    disabled=(
                        abs(diem_mau_giao - 10.0) > 1e-9
                        or not cua_so_hop_le_moi
                    ),
                    key="exam_assignment_create"
                ):
                    mo_dt = mo_preview
                    dong_dt = dong_preview
                    han_bat_dau_dt = han_bat_dau_preview
                    if dong_dt <= mo_dt:
                        st.error("Giờ kết thúc đợt phải sau giờ mở.")
                    elif han_bat_dau_dt < mo_dt:
                        st.error("Khung giờ kiểm tra quá ngắn so với thời gian làm bài của mỗi học sinh.")
                    else:
                        dot_moi = {
                            "id": str(uuid.uuid4()),
                            "ten_dot": ten_dot_moi.strip() or f"Kiểm tra – {mau_giao.get('ten_mau', 'Ma trận')}",
                            "mau_id": str(mau_giao.get("id", "")),
                            "ten_mau": str(mau_giao.get("ten_mau", "")),
                            "mau_phien_ban": int(mau_giao.get("phien_ban", 1) or 1),
                            "khoi": str(mau_giao.get("khoi", "")),
                            "lop_ap_dung": lop_dot_moi,
                            "ma_tran_snapshot": list(mau_giao.get("ma_tran", []) or []),
                            "thoi_gian_lam_phut": int(thoi_gian_lam_moi),
                            "mo_tu_iso": mo_dt.isoformat(),
                            "dong_luc_iso": dong_dt.isoformat(),
                            "han_cuoi_bat_dau_iso": han_bat_dau_dt.isoformat(),
                            "mo_dap_an": False,
                            "ngay_tao": bay_gio_viet_nam().strftime("%d/%m/%Y %H:%M"),
                            "ngay_cap_nhat": ""
                        }
                        ds_dot.append(dot_moi)
                        luu_dot_kiem_tra_ma_tran(ds_dot)
                        st.success(
                            "Đã giao đợt kiểm tra. Học sinh đúng lớp chỉ thấy đợt này trong thời gian được phép. "
                            "Đáp án đang khóa mặc định."
                        )
                        st.rerun()

            st.markdown("---")
            st.markdown("### 🗂️ Các đợt kiểm tra đã giao")
            ds_dot = doc_dot_kiem_tra_ma_tran()

            if not ds_dot:
                st.info("Chưa có đợt kiểm tra nào.")
            else:
                lich_su_all_dot = doc_lich_su_hoc_sinh()
                rows_dot = []
                for d in ds_dot:
                    status, _ = trang_thai_dot_kiem_tra(d)
                    dot_id = str(d.get("id", ""))
                    so_nop = sum(
                        1 for lan in lich_su_all_dot
                        if str((lan.get("pham_vi", {}) or {}).get("dot_kiem_tra_id", "")) == dot_id
                    )
                    rows_dot.append({
                        "Đợt kiểm tra": d.get("ten_dot", ""),
                        "Ma trận": d.get("ten_mau", ""),
                        "Lớp": d.get("lop_ap_dung", "Tất cả"),
                        "Mở từ": fmt_vn_datetime(d.get("mo_tu_iso", "")),
                        "Hạn cuối bắt đầu": fmt_vn_datetime(han_cuoi_bat_dau_dot_kiem_tra(d)),
                        "Kết thúc đợt": fmt_vn_datetime(d.get("dong_luc_iso", "")),
                        "Thời gian làm": f"{int(d.get('thoi_gian_lam_phut', 0) or 0)} phút",
                        "Trạng thái": status,
                        "Đáp án": "Đã mở" if d.get("mo_dap_an") else "Đang khóa",
                        "Số bài đã nộp": so_nop
                    })

                st.dataframe(
                    pd.DataFrame(rows_dot),
                    use_container_width=True,
                    hide_index=True
                )

                idx_dot = st.selectbox(
                    "Chọn đợt để sửa / mở đáp án / xóa",
                    options=list(range(len(ds_dot))),
                    format_func=lambda i: (
                        f"{ds_dot[i].get('ten_dot', 'Không tên')} • "
                        f"{ds_dot[i].get('lop_ap_dung', 'Tất cả')} • "
                        f"{fmt_vn_datetime(ds_dot[i].get('mo_tu_iso', ''))}"
                    ),
                    key="exam_assignment_manage_pick"
                )
                dot_chon = dict(ds_dot[idx_dot])
                dot_id = str(dot_chon.get("id", idx_dot))
                mo_old = parse_iso_vn(dot_chon.get("mo_tu_iso", "")) or bay_gio_viet_nam()
                dong_old = parse_iso_vn(dot_chon.get("dong_luc_iso", "")) or (mo_old + timedelta(hours=2))

                with st.container(border=True):
                    st.markdown("### ⚙️ Cài đặt đợt kiểm tra")
                    ten_dot_sua = st.text_input(
                        "Tên đợt",
                        value=str(dot_chon.get("ten_dot", "")),
                        key=f"exam_assignment_edit_name_{dot_id}"
                    )
                    ds_lop_sua = ["Tất cả"] + lay_danh_sach_lop_tu_hoc_sinh()
                    lop_cu = str(dot_chon.get("lop_ap_dung", "Tất cả") or "Tất cả")
                    if lop_cu not in ds_lop_sua:
                        ds_lop_sua.append(lop_cu)
                    lop_dot_sua = st.selectbox(
                        "Lớp áp dụng",
                        ds_lop_sua,
                        index=ds_lop_sua.index(lop_cu),
                        key=f"exam_assignment_edit_class_{dot_id}"
                    )

                    e1, e2 = st.columns(2)
                    with e1:
                        ngay_mo_sua = st.date_input(
                            "Ngày mở",
                            value=mo_old.date(),
                            key=f"exam_assignment_edit_start_date_{dot_id}"
                        )
                        gio_mo_sua = st.time_input(
                            "Giờ mở",
                            value=mo_old.time().replace(tzinfo=None, second=0, microsecond=0),
                            key=f"exam_assignment_edit_start_time_{dot_id}"
                        )
                    with e2:
                        ngay_dong_sua = st.date_input(
                            "Ngày kết thúc đợt",
                            value=dong_old.date(),
                            key=f"exam_assignment_edit_end_date_{dot_id}"
                        )
                        gio_dong_sua = st.time_input(
                            "Giờ kết thúc đợt",
                            value=dong_old.time().replace(tzinfo=None, second=0, microsecond=0),
                            key=f"exam_assignment_edit_end_time_{dot_id}"
                        )

                    tg_sua = st.number_input(
                        "Thời gian làm bài (phút)",
                        min_value=5,
                        max_value=300,
                        value=max(5, int(dot_chon.get("thoi_gian_lam_phut", 45) or 45)),
                        step=5,
                        key=f"exam_assignment_edit_duration_{dot_id}"
                    )

                    mo_preview_sua = datetime.combine(ngay_mo_sua, gio_mo_sua).replace(tzinfo=VN_TZ)
                    dong_preview_sua = datetime.combine(ngay_dong_sua, gio_dong_sua).replace(tzinfo=VN_TZ)
                    han_bat_dau_preview_sua = dong_preview_sua - timedelta(minutes=int(tg_sua))
                    cua_so_hop_le_sua = (
                        dong_preview_sua > mo_preview_sua
                        and han_bat_dau_preview_sua >= mo_preview_sua
                    )
                    if cua_so_hop_le_sua:
                        st.info(
                            "⏰ Hạn cuối HS được bắt đầu: "
                            f"**{han_bat_dau_preview_sua.strftime('%d/%m/%Y %H:%M')}** • "
                            f"mỗi HS có đúng **{int(tg_sua)} phút** • "
                            f"kết thúc toàn bộ đợt lúc **{dong_preview_sua.strftime('%d/%m/%Y %H:%M')}**."
                        )
                    else:
                        st.error(
                            "Khung giờ hiện tại không đủ cho mỗi HS làm trọn "
                            f"**{int(tg_sua)} phút**."
                        )

                    mo_dap_an_sua = st.toggle(
                        "🔓 Cho học sinh xem đáp án và lời giải sau khi đã nộp",
                        value=bool(dot_chon.get("mo_dap_an", False)),
                        key=f"exam_assignment_edit_release_{dot_id}"
                    )
                    st.caption(
                        "Khi tắt: HS sau khi nộp chỉ thấy **điểm và thời gian**, không thấy câu đúng/sai, đáp án, lời giải hay phản hồi có thể làm lộ đáp án."
                    )

                    em1, em2 = st.columns(2)
                    with em1:
                        if st.button(
                            "💾 LƯU CÀI ĐẶT ĐỢT",
                            type="primary",
                            use_container_width=True,
                            key=f"exam_assignment_edit_save_{dot_id}"
                        ):
                            mo_new = mo_preview_sua
                            dong_new = dong_preview_sua
                            han_bat_dau_new = han_bat_dau_preview_sua
                            if dong_new <= mo_new:
                                st.error("Giờ kết thúc đợt phải sau giờ mở.")
                            elif han_bat_dau_new < mo_new:
                                st.error("Khung giờ kiểm tra quá ngắn so với thời gian làm bài của mỗi học sinh.")
                            else:
                                dot_new = dict(dot_chon)
                                dot_new["ten_dot"] = ten_dot_sua.strip() or dot_chon.get("ten_dot", "Đợt kiểm tra")
                                dot_new["lop_ap_dung"] = lop_dot_sua
                                dot_new["mo_tu_iso"] = mo_new.isoformat()
                                dot_new["dong_luc_iso"] = dong_new.isoformat()
                                dot_new["han_cuoi_bat_dau_iso"] = han_bat_dau_new.isoformat()
                                dot_new["thoi_gian_lam_phut"] = int(tg_sua)
                                dot_new["mo_dap_an"] = bool(mo_dap_an_sua)
                                dot_new["ngay_cap_nhat"] = bay_gio_viet_nam().strftime("%d/%m/%Y %H:%M")
                                ds_dot[idx_dot] = dot_new
                                luu_dot_kiem_tra_ma_tran(ds_dot)
                                st.success("Đã cập nhật đợt kiểm tra.")
                                st.rerun()

                    with em2:
                        xoa_dot = st.checkbox(
                            "Tôi xác nhận xóa đợt kiểm tra này khỏi danh sách HS.",
                            key=f"exam_assignment_delete_confirm_{dot_id}"
                        )
                        if st.button(
                            "🗑 XÓA ĐỢT KIỂM TRA",
                            use_container_width=True,
                            disabled=not xoa_dot,
                            key=f"exam_assignment_delete_{dot_id}"
                        ):
                            ds_dot.pop(idx_dot)
                            luu_dot_kiem_tra_ma_tran(ds_dot)
                            st.success("Đã xóa đợt kiểm tra. Điểm/lịch sử đã nộp vẫn được giữ.")
                            st.rerun()



# ==========================================================
# GIÁO VIÊN
# ==========================================================
def kho_tai_lieu_gv():
    st.header("📁 KHO TÀI LIỆU GIÁO VIÊN")

    st.info(
        "Tải tài liệu vào kho để sử dụng khi AI tạo câu hỏi."
    )

    thu_muc_kho = "kho_tai_lieu_gv"

    # Tạo thư mục kho nếu chưa có
    if not os.path.exists(thu_muc_kho):
        os.makedirs(thu_muc_kho)

    # Upload nhiều tài liệu
    files = st.file_uploader(
        "📎 Chọn tài liệu",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="upload_kho_tai_lieu_gv"
    )

    # Lưu tài liệu vào kho
    if files:
        so_file_moi = 0

        for file in files:
            duong_dan = os.path.join(
                thu_muc_kho,
                file.name
            )

            if not os.path.exists(duong_dan):
                with open(duong_dan, "wb") as f:
                    f.write(file.getbuffer())

                so_file_moi += 1

        if so_file_moi > 0:
            st.success(
                f"Đã thêm {so_file_moi} tài liệu mới vào kho."
            )
        else:
            st.info(
                "Các tài liệu này đã có trong kho."
            )

    # Hiển thị danh sách tài liệu hiện có
    danh_sach_file = os.listdir(thu_muc_kho)

    st.subheader(
        f"📚 Tài liệu trong kho ({len(danh_sach_file)})"
    )

    if not danh_sach_file:
        st.caption("Kho tài liệu hiện đang trống.")

    else:
        for ten_file in danh_sach_file:
            st.write("📄", ten_file)

# ==========================================================
# QUẢN LÝ HỌC SINH
# ==========================================================
def doc_danh_sach_hoc_sinh():
    return _doc_students_shared()


def luu_danh_sach_hoc_sinh(ds):
    return _luu_students_shared(ds)


def chuan_hoa_ten_lop(value):
    s = re.sub(
        r"[^0-9A-Za-z_-]+",
        "",
        str(value or "").strip().upper()
    )

    return s or "LOP"


def tao_ma_hoc_sinh_tu_dong(lop, ds_hien_co):
    lop_code = chuan_hoa_ten_lop(
        lop
    )

    pattern = re.compile(
        rf"^{re.escape(lop_code)}-(\d+)$"
    )

    nums = []

    for hs in ds_hien_co:
        ma = str(
            hs.get(
                "ma_hoc_sinh",
                ""
            )
        ).strip().upper()

        m = pattern.match(
            ma
        )

        if m:
            nums.append(
                int(
                    m.group(1)
                )
            )

    so_moi = (
        max(nums) + 1
        if nums
        else 1
    )

    return (
        f"{lop_code}-{so_moi:03d}"
    )


def tim_hoc_sinh_theo_ma(ma):
    ma_chuan = str(
        ma or ""
    ).strip().upper()

    for hs in doc_danh_sach_hoc_sinh():
        if str(
            hs.get(
                "ma_hoc_sinh",
                ""
            )
        ).strip().upper() == ma_chuan:
            return hs

    return None


def them_hoc_sinh_moi(
    ho_ten,
    lop,
    khoi="",
    ghi_chu=""
):
    ds = doc_danh_sach_hoc_sinh()

    ho_ten = " ".join(
        str(
            ho_ten or ""
        ).strip().split()
    )

    lop = str(
        lop or ""
    ).strip()

    if not ho_ten or not lop:
        return None

    ma = tao_ma_hoc_sinh_tu_dong(
        lop,
        ds
    )

    hs = {
        "id": str(
            uuid.uuid4()
        ),
        "ma_hoc_sinh": ma,
        "ho_ten": ho_ten,
        "lop": lop,
        "khoi": str(
            khoi or ""
        ).strip(),
        "trang_thai": "Đang học",
        "ghi_chu": str(
            ghi_chu or ""
        ).strip(),
        "ngay_tao": datetime.now().strftime(
            "%d/%m/%Y %H:%M"
        )
    }

    ds.append(
        hs
    )

    luu_danh_sach_hoc_sinh(
        ds
    )

    return hs



def doc_file_danh_sach_hoc_sinh_upload(file):
    """
    Chấp nhận XLSX/XLS/CSV.
    Ưu tiên cột có tên Họ và tên / Họ tên / Tên học sinh.
    Nếu không có thì lấy cột dữ liệu đầu tiên.
    """
    if file is None:
        return []

    ten = str(
        getattr(file, "name", "")
    ).lower()

    try:
        if ten.endswith(
            (".xlsx", ".xls")
        ):
            df = pd.read_excel(
                file
            )
        elif ten.endswith(".csv"):
            try:
                df = pd.read_csv(
                    file,
                    encoding="utf-8-sig"
                )
            except Exception:
                file.seek(0)
                df = pd.read_csv(
                    file
                )
        else:
            return []

    except Exception:
        return []

    if df is None or df.empty:
        return []

    # Bỏ các cột rỗng hoàn toàn.
    df = df.dropna(
        axis=1,
        how="all"
    )

    if df.empty:
        return []

    aliases = {
        "họ và tên",
        "họ tên",
        "ho va ten",
        "ho ten",
        "tên học sinh",
        "ten hoc sinh",
        "học sinh",
        "hoc sinh",
        "student",
        "name"
    }

    cot_ten = None

    for col in df.columns:
        chuan = (
            unicodedata.normalize(
                "NFKC",
                str(col)
            )
            .strip()
            .casefold()
        )

        if chuan in aliases:
            cot_ten = col
            break

    if cot_ten is None:
        # Ưu tiên bỏ cột STT nếu file mẫu có STT.
        cot_kha_dung = [
            c
            for c in df.columns
            if str(c).strip().casefold()
            not in {
                "stt",
                "số thứ tự",
                "so thu tu",
                "no",
                "id"
            }
        ]

        cot_ten = (
            cot_kha_dung[0]
            if cot_kha_dung
            else df.columns[0]
        )

    ds_ten = []

    for value in df[
        cot_ten
    ].tolist():
        if pd.isna(
            value
        ):
            continue

        ten_hs = " ".join(
            str(
                value
            ).strip().split()
        )

        if not ten_hs:
            continue

        # Tránh đọc nhầm header lặp lại thành dữ liệu.
        if ten_hs.casefold() in aliases:
            continue

        ds_ten.append(
            ten_hs
        )

    return ds_ten


def tao_excel_danh_sach_ma_hoc_sinh(ds_hs):
    """
    Tạo file Excel để GV tải về và phát mã cho HS.
    """
    df = pd.DataFrame([
        {
            "STT": i,
            "Mã học sinh": hs.get(
                "ma_hoc_sinh",
                ""
            ),
            "Họ và tên": hs.get(
                "ho_ten",
                ""
            ),
            "Lớp": hs.get(
                "lop",
                ""
            ),
            "Khối": hs.get(
                "khoi",
                ""
            )
        }
        for i, hs in enumerate(
            ds_hs,
            start=1
        )
    ])

    output = io.BytesIO()

    with pd.ExcelWriter(
        output,
        engine="openpyxl"
    ) as writer:
        df.to_excel(
            writer,
            index=False,
            sheet_name="Danh sách mã HS"
        )

        ws = writer.book[
            "Danh sách mã HS"
        ]

        # Căn chỉnh gọn để in/phát cho học sinh.
        widths = {
            "A": 8,
            "B": 18,
            "C": 32,
            "D": 14,
            "E": 14
        }

        for col, width in widths.items():
            ws.column_dimensions[
                col
            ].width = width

        for cell in ws[1]:
            cell.font = cell.font.copy(
                bold=True
            )

        ws.freeze_panes = "A2"

    output.seek(0)

    return output.getvalue()


def tao_excel_mau_import_hoc_sinh():
    """
    File mẫu tối giản: chỉ cần STT + Họ và tên.
    Lớp và Khối chọn trực tiếp trên app.
    """
    df = pd.DataFrame({
        "STT": [
            1,
            2,
            3
        ],
        "Họ và tên": [
            "Nguyễn Văn A",
            "Trần Thị B",
            "Lê Văn C"
        ]
    })

    output = io.BytesIO()

    with pd.ExcelWriter(
        output,
        engine="openpyxl"
    ) as writer:
        df.to_excel(
            writer,
            index=False,
            sheet_name="Danh sách học sinh"
        )

        ws = writer.book[
            "Danh sách học sinh"
        ]

        ws.column_dimensions[
            "A"
        ].width = 8

        ws.column_dimensions[
            "B"
        ].width = 32

        for cell in ws[1]:
            cell.font = cell.font.copy(
                bold=True
            )

    output.seek(0)

    return output.getvalue()



def quan_ly_hoc_sinh():
    st.header(
        "👥 QUẢN LÝ HỌC SINH"
    )

    st.caption(
        "GV tạo lớp và danh sách học sinh. "
        "App tự cấp mã duy nhất để dùng cho đăng nhập, lưu lịch sử và cá nhân hóa."
    )

    ds = doc_danh_sach_hoc_sinh()

    tab1, tab2, tab3 = st.tabs([
        "➕ Thêm học sinh",
        "📋 Danh sách lớp",
        "🧹 Quản lý"
    ])

    # ------------------------------------------------------
    # TAB 1: thêm từng HS / dán danh sách
    # ------------------------------------------------------
    with tab1:
        st.subheader(
            "Thêm học sinh"
        )

        c1, c2, c3 = st.columns(3)

        with c1:
            khoi_moi = st.selectbox(
                "Khối",
                [
                    "Khối 10",
                    "Khối 11",
                    "Khối 12"
                ],
                key="student_add_grade"
            )

        with c2:
            lop_moi = st.text_input(
                "Lớp",
                placeholder="VD: 12A2",
                key="student_add_class"
            )

        with c3:
            ten_moi = st.text_input(
                "Họ và tên",
                placeholder="Nguyễn Văn A",
                key="student_add_name"
            )

        if st.button(
            "➕ THÊM 1 HỌC SINH",
            type="primary",
            use_container_width=True,
            key="student_add_one"
        ):
            hs_moi = them_hoc_sinh_moi(
                ten_moi,
                lop_moi,
                khoi_moi
            )

            if hs_moi:
                st.success(
                    f"Đã thêm **{hs_moi['ho_ten']}** – "
                    f"Mã: **{hs_moi['ma_hoc_sinh']}**"
                )
                st.rerun()
            else:
                st.warning(
                    "Cần nhập đủ họ tên và lớp."
                )

        st.markdown("---")

        st.markdown(
            "### ⚡ Tạo danh sách học sinh"
        )

        st.caption(
            "Chọn khối và lớp, sau đó tải file Excel/CSV có cột **Họ và tên**."
        )

        c_bulk1, c_bulk2 = st.columns(2)

        with c_bulk1:
            khoi_hang_loat = st.selectbox(
                "Khối",
                [
                    "Khối 10",
                    "Khối 11",
                    "Khối 12"
                ],
                key="student_bulk_grade"
            )

        with c_bulk2:
            lop_hang_loat = st.text_input(
                "Lớp",
                placeholder="VD: 12A2",
                key="student_bulk_class"
            )

        upload_col, sample_col = st.columns(
            [4, 1]
        )

        with upload_col:
            file_hs = st.file_uploader(
                "📤 Danh sách học sinh",
                type=[
                    "xlsx",
                    "xls",
                    "csv"
                ],
                help=(
                    "File nên có cột 'Họ và tên'. "
                    "App cũng có thể nhận file chỉ gồm STT và Họ và tên."
                ),
                key="student_bulk_file"
            )

        with sample_col:
            st.write("")
            st.write("")
            mau_bytes = tao_excel_mau_import_hoc_sinh()

            st.download_button(
                "📄 File mẫu",
                data=mau_bytes,
                file_name="mau_danh_sach_hoc_sinh.xlsx",
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "spreadsheetml.sheet"
                ),
                use_container_width=True,
                key="student_download_template"
            )

        ds_ten_upload = doc_file_danh_sach_hoc_sinh_upload(
            file_hs
        )

        if file_hs is not None:
            if ds_ten_upload:
                st.success(
                    f"✅ Đã đọc được **{len(ds_ten_upload)} học sinh**."
                )

                with st.expander(
                    "👁️ Xem trước danh sách",
                    expanded=True
                ):
                    st.dataframe(
                        pd.DataFrame({
                            "STT": range(
                                1,
                                len(ds_ten_upload) + 1
                            ),
                            "Họ và tên": ds_ten_upload
                        }),
                        use_container_width=True,
                        hide_index=True
                    )
            else:
                st.warning(
                    "Không đọc được danh sách. "
                    "Bạn có thể bấm **File mẫu** để xem định dạng."
                )

        if "student_bulk_created" not in st.session_state:
            st.session_state.student_bulk_created = []

        can_create = bool(
            ds_ten_upload
        ) and bool(
            str(lop_hang_loat).strip()
        )

        if st.button(
            "⚡ TẠO MÃ HỌC SINH",
            type="primary",
            use_container_width=True,
            disabled=not can_create,
            key="student_bulk_add"
        ):
            da_them = []

            for ten in ds_ten_upload:
                hs_moi = them_hoc_sinh_moi(
                    ten,
                    lop_hang_loat,
                    khoi_hang_loat
                )

                if hs_moi:
                    da_them.append(
                        hs_moi
                    )

            st.session_state.student_bulk_created = da_them

            st.success(
                f"🎉 Đã tạo mã cho **{len(da_them)} học sinh**."
            )

        da_tao_vua_xong = st.session_state.get(
            "student_bulk_created",
            []
        )

        if da_tao_vua_xong:
            st.markdown("---")
            st.markdown(
                "### ✅ Danh sách đã cấp mã"
            )

            st.dataframe(
                pd.DataFrame([
                    {
                        "STT": i,
                        "Mã học sinh": x.get(
                            "ma_hoc_sinh",
                            ""
                        ),
                        "Họ và tên": x.get(
                            "ho_ten",
                            ""
                        ),
                        "Lớp": x.get(
                            "lop",
                            ""
                        ),
                        "Khối": x.get(
                            "khoi",
                            ""
                        )
                    }
                    for i, x in enumerate(
                        da_tao_vua_xong,
                        start=1
                    )
                ]),
                use_container_width=True,
                hide_index=True
            )

            file_xuat = tao_excel_danh_sach_ma_hoc_sinh(
                da_tao_vua_xong
            )

            ten_lop_file = chuan_hoa_ten_lop(
                da_tao_vua_xong[0].get(
                    "lop",
                    "lop"
                )
            )

            st.download_button(
                "⬇️ TẢI DANH SÁCH MÃ HỌC SINH",
                data=file_xuat,
                file_name=(
                    f"danh_sach_ma_hoc_sinh_{ten_lop_file}.xlsx"
                ),
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "spreadsheetml.sheet"
                ),
                type="primary",
                use_container_width=True,
                key="student_download_codes"
            )

            st.caption(
                "File gồm mã học sinh, họ tên, lớp và khối để GV lưu, in hoặc gửi cho học sinh."
            )


    # ------------------------------------------------------
    # TAB 2: xem danh sách
    # ------------------------------------------------------
    with tab2:
        if not ds:
            st.info(
                "Chưa có học sinh."
            )
        else:
            ds_lop = sorted({
                str(
                    x.get(
                        "lop",
                        ""
                    )
                )
                for x in ds
                if x.get(
                    "lop"
                )
            })

            lop_loc = st.selectbox(
                "Chọn lớp",
                ["Tất cả"] + ds_lop,
                key="student_filter_class"
            )

            ds_xem = (
                ds
                if lop_loc == "Tất cả"
                else [
                    x
                    for x in ds
                    if x.get(
                        "lop"
                    ) == lop_loc
                ]
            )

            st.metric(
                "Số học sinh",
                len(
                    ds_xem
                )
            )

            st.dataframe(
                pd.DataFrame([
                    {
                        "Mã học sinh": x.get(
                            "ma_hoc_sinh",
                            ""
                        ),
                        "Họ và tên": x.get(
                            "ho_ten",
                            ""
                        ),
                        "Lớp": x.get(
                            "lop",
                            ""
                        ),
                        "Khối": x.get(
                            "khoi",
                            ""
                        ),
                        "Trạng thái": x.get(
                            "trang_thai",
                            ""
                        )
                    }
                    for x in ds_xem
                ]),
                use_container_width=True,
                hide_index=True
            )

            st.info(
                "Học sinh chỉ cần nhập **Mã học sinh** ở khu vực HS. "
                "Họ tên và lớp sẽ được app nhận diện tự động."
            )

    # ------------------------------------------------------
    # TAB 3: khóa / xóa HS
    # ------------------------------------------------------
    with tab3:
        if not ds:
            st.info(
                "Chưa có học sinh để quản lý."
            )
        else:
            options = [
                f"{x.get('ma_hoc_sinh', '')} – "
                f"{x.get('ho_ten', '')} – "
                f"{x.get('lop', '')}"
                for x in ds
            ]

            selected_label = st.selectbox(
                "Chọn học sinh",
                options,
                key="student_manage_pick"
            )

            idx = options.index(
                selected_label
            )

            hs_chon = dict(
                ds[idx]
            )

            st.write(
                f"**Mã:** {hs_chon.get('ma_hoc_sinh', '')}"
            )
            st.write(
                f"**Họ tên:** {hs_chon.get('ho_ten', '')}"
            )
            st.write(
                f"**Lớp:** {hs_chon.get('lop', '')}"
            )

            trang_thai_moi = st.selectbox(
                "Trạng thái",
                [
                    "Đang học",
                    "Tạm khóa"
                ],
                index=(
                    1
                    if hs_chon.get(
                        "trang_thai"
                    ) == "Tạm khóa"
                    else 0
                ),
                key="student_status"
            )

            if st.button(
                "💾 CẬP NHẬT TRẠNG THÁI",
                use_container_width=True,
                key="student_update_status"
            ):
                ds[idx][
                    "trang_thai"
                ] = trang_thai_moi

                luu_danh_sach_hoc_sinh(
                    ds
                )

                st.success(
                    "Đã cập nhật."
                )
                st.rerun()

            st.markdown("---")

            xac_nhan = st.checkbox(
                "Tôi xác nhận muốn xóa học sinh này khỏi danh sách.",
                key="student_delete_confirm"
            )

            if st.button(
                "🗑 XÓA HỌC SINH",
                use_container_width=True,
                disabled=not xac_nhan,
                key="student_delete"
            ):
                ma_xoa = hs_chon.get(
                    "ma_hoc_sinh",
                    ""
                )

                ds_moi = [
                    x
                    for x in ds
                    if x.get(
                        "ma_hoc_sinh"
                    ) != ma_xoa
                ]

                luu_danh_sach_hoc_sinh(
                    ds_moi
                )
                _xoa_student_shared(
                    ma_xoa
                )

                st.success(
                    "Đã xóa khỏi danh sách học sinh. "
                    "Lịch sử luyện tập cũ vẫn được giữ để tránh mất dữ liệu."
                )
                st.rerun()




# ==========================================================
# PHÂN TÍCH LỚP HỌC - KHÔNG DÙNG AI/API
# ==========================================================
def lay_danh_sach_lop_tu_hoc_sinh():
    ds_hs = doc_danh_sach_hoc_sinh()
    return sorted({
        str(x.get("lop", "")).strip()
        for x in ds_hs
        if str(x.get("lop", "")).strip()
    })


def tong_hop_du_lieu_lop(lop_chon=None, che_do_filter=None, ten_luot_filter=None, bai_key_filter=None):
    """
    Tổng hợp trực tiếp từ lịch sử làm bài.
    Không gọi AI.
    Đúng/Sai đã được lưu theo từng ý trong don_vi_danh_gia.
    """
    lich_su = doc_lich_su_hoc_sinh()
    ds_hs = doc_danh_sach_hoc_sinh()

    map_hs = {
        str(x.get("ma_hoc_sinh", "")).strip().upper(): x
        for x in ds_hs
    }

    if lop_chon and lop_chon != "Tất cả":
        ma_lop = {
            ma
            for ma, hs in map_hs.items()
            if str(hs.get("lop", "")).strip() == lop_chon
        }
        lich_su = [
            x for x in lich_su
            if str(x.get("hoc_sinh_id", "")).strip().upper() in ma_lop
        ]

    if che_do_filter:
        if isinstance(che_do_filter, (list, tuple, set)):
            che_do_set = {str(x).strip() for x in che_do_filter if str(x).strip()}
        else:
            che_do_set = {str(che_do_filter).strip()}
        lich_su = [
            x for x in lich_su
            if str(x.get("che_do", "")).strip() in che_do_set
        ]

    if ten_luot_filter:
        ten_loc = str(ten_luot_filter).strip()
        lich_su = [
            x for x in lich_su
            if str(x.get("ten_luot", "")).strip() == ten_loc
        ]

    if bai_key_filter:
        key_loc = str(bai_key_filter).strip()
        lich_su = [
            x for x in lich_su
            if khoa_bai_lam_chinh_thuc(x) == key_loc
        ]

    tong_hop = {
        "so_luot": len(lich_su),
        "_lich_su": lich_su,
        "hoc_sinh": {},
        "yccd": {},
        "muc_do": {},
        "nang_luc": {},
        "dang_cau": {},
        "bai": {},
        "chuong": {},
    }

    for lan in lich_su:
        ma = str(
            lan.get("hoc_sinh_id", "")
        ).strip().upper()

        hs_info = map_hs.get(
            ma,
            {}
        )

        hs_stat = tong_hop["hoc_sinh"].setdefault(
            ma,
            {
                "ma": ma,
                "ho_ten": hs_info.get(
                    "ho_ten",
                    lan.get("ho_ten", "")
                ),
                "lop": hs_info.get(
                    "lop",
                    ""
                ),
                "so_luot": 0,
                "tong_don_vi": 0,
                "dung_don_vi": 0,
                "diem": [],
            }
        )

        hs_stat["so_luot"] += 1
        hs_stat["diem"].append(
            float(lan.get("diem", 0) or 0)
        )

        for item in lan.get(
            "chi_tiet",
            []
        ) or []:
            q = item.get(
                "cau_snapshot",
                {}
            ) or {}

            dang = str(
                item.get(
                    "dang_cau",
                    q.get("dang_cau", "")
                )
            ).strip()

            bai = str(
                q.get("bai", "")
            ).strip()

            chuong = str(
                q.get("chuong", "")
            ).strip()

            units = item.get(
                "don_vi_danh_gia",
                []
            ) or []

            if not units:
                units = [{
                    "yccd": q.get("yccd", ""),
                    "muc_do": q.get("muc_do", ""),
                    "nang_luc": q.get(
                        "thanh_phan_nang_luc",
                        ""
                    ),
                    "dung": bool(
                        item.get(
                            "dung_toan_cau"
                        )
                    )
                }]

            for unit in units:
                dung = bool(
                    unit.get("dung")
                )

                hs_stat["tong_don_vi"] += 1
                hs_stat["dung_don_vi"] += int(
                    dung
                )

                def cap_nhat(bucket, key):
                    key = str(
                        key or ""
                    ).strip()

                    if not key:
                        return

                    s = bucket.setdefault(
                        key,
                        {
                            "tong": 0,
                            "dung": 0,
                            "sai": 0,
                            "hoc_sinh_sai": set()
                        }
                    )

                    s["tong"] += 1
                    s["dung"] += int(
                        dung
                    )
                    s["sai"] += int(
                        not dung
                    )

                    if not dung and ma:
                        s["hoc_sinh_sai"].add(
                            ma
                        )

                cap_nhat(
                    tong_hop["yccd"],
                    unit.get("yccd", "")
                )
                cap_nhat(
                    tong_hop["muc_do"],
                    unit.get("muc_do", "")
                )
                cap_nhat(
                    tong_hop["nang_luc"],
                    unit.get("nang_luc", "")
                )
                cap_nhat(
                    tong_hop["dang_cau"],
                    dang
                )
                cap_nhat(
                    tong_hop["bai"],
                    bai
                )
                cap_nhat(
                    tong_hop["chuong"],
                    chuong
                )

    return tong_hop


def bang_thong_ke_bucket(bucket, ten_cot):
    rows = []

    for key, s in bucket.items():
        tong = int(
            s.get("tong", 0)
        )
        dung = int(
            s.get("dung", 0)
        )

        rows.append({
            ten_cot: key,
            "Số lượt đánh giá": tong,
            "Số đúng": dung,
            "Số sai": int(
                s.get("sai", 0)
            ),
            "Tỉ lệ đúng": (
                round(
                    dung / tong * 100,
                    1
                )
                if tong
                else 0
            ),
            "Số HS từng sai": len(
                s.get(
                    "hoc_sinh_sai",
                    set()
                )
            )
        })

    rows.sort(
        key=lambda x: (
            x["Tỉ lệ đúng"],
            -x["Số lượt đánh giá"]
        )
    )

    return rows


def phan_tich_lop_hoc():
    st.header(
        "📊 PHÂN TÍCH LỚP HỌC"
    )

    st.caption(
        "Dữ liệu được tổng hợp trực tiếp từ bài làm của học sinh. "
        "Phần này **không gọi AI và không tốn API**."
    )

    ds_lop = lay_danh_sach_lop_tu_hoc_sinh()

    if not ds_lop:
        st.info(
            "Chưa có lớp/học sinh trong hệ thống."
        )
        return

    lop = st.selectbox(
        "Chọn lớp",
        ["Tất cả"] + ds_lop,
        key="gv_class_analysis_class"
    )

    pham_vi_phan_tich = st.selectbox(
        "Nguồn dữ liệu phân tích",
        [
            "Tất cả hoạt động",
            "Chỉ các bài có điểm chính thức",
            CHE_DO_DE_GV,
            CHE_DO_KIEM_TRA_MA_TRAN,
            CHE_DO_TOT_NGHIEP
        ],
        key="gv_class_analysis_scope"
    )

    if pham_vi_phan_tich == "Tất cả hoạt động":
        che_do_loc = None
    elif pham_vi_phan_tich == "Chỉ các bài có điểm chính thức":
        che_do_loc = [
            CHE_DO_DE_GV,
            CHE_DO_KIEM_TRA_MA_TRAN,
            CHE_DO_TOT_NGHIEP
        ]
    else:
        che_do_loc = pham_vi_phan_tich

    bai_key_phan_tich = None

    # Khi GV đang xem dữ liệu chấm điểm chính thức, cho chọn chính xác từng đề/đợt.
    if pham_vi_phan_tich != "Tất cả hoạt động":
        ds_luot_pt = loc_luot_co_diem_chinh_thuc(
            lop=lop,
            che_do=(
                None
                if isinstance(che_do_loc, (list, tuple, set))
                else che_do_loc
            )
        )
        if isinstance(che_do_loc, (list, tuple, set)):
            ds_luot_pt = [
                x for x in ds_luot_pt
                if str(x.get("che_do", "")).strip() in set(che_do_loc)
            ]

        key_label_pt = {}
        for lan in ds_luot_pt:
            k = khoa_bai_lam_chinh_thuc(lan)
            if k and k not in key_label_pt:
                key_label_pt[k] = nhan_bai_lam_chinh_thuc(lan)

        if key_label_pt:
            ds_key_pt = sorted(key_label_pt, key=lambda k: key_label_pt[k])
            key_chon_pt = st.selectbox(
                "Chọn riêng từng đề / đợt để phân tích",
                ["Tất cả"] + ds_key_pt,
                format_func=lambda k: "Tất cả" if k == "Tất cả" else key_label_pt.get(k, k),
                key="gv_class_analysis_exact_exam"
            )
            if key_chon_pt != "Tất cả":
                bai_key_phan_tich = key_chon_pt

    data = tong_hop_du_lieu_lop(
        lop,
        che_do_filter=che_do_loc,
        bai_key_filter=bai_key_phan_tich
    )

    hs_stats = list(
        data.get(
            "hoc_sinh",
            {}
        ).values()
    )

    if not hs_stats:
        st.info(
            "Lớp này chưa có dữ liệu làm bài."
        )
        return

    tong_don_vi = sum(
        x.get("tong_don_vi", 0)
        for x in hs_stats
    )

    tong_dung = sum(
        x.get("dung_don_vi", 0)
        for x in hs_stats
    )

    c1, c2, c3, c4 = st.columns(4)

    with c1:
        st.metric(
            "Học sinh có dữ liệu",
            len(hs_stats)
        )

    with c2:
        st.metric(
            "Tổng lượt làm bài",
            data.get(
                "so_luot",
                0
            )
        )

    with c3:
        st.metric(
            "Tỉ lệ đúng toàn lớp",
            (
                f"{tong_dung / tong_don_vi * 100:.1f}%"
                if tong_don_vi
                else "0%"
            )
        )

    with c4:
        diem_all = [
            d
            for hs in hs_stats
            for d in hs.get(
                "diem",
                []
            )
        ]

        st.metric(
            "Điểm TB các lượt",
            (
                f"{sum(diem_all) / len(diem_all):.1f}"
                if diem_all
                else "—"
            )
        )

    tab1, tab2, tab3, tab4 = st.tabs([
        "🎯 Nội dung lớp còn yếu",
        "👥 Học sinh cần hỗ trợ",
        "📈 Mức độ & năng lực",
        "📚 Bài / chương"
    ])

    # ------------------------------------------------------
    # TAB 1
    # ------------------------------------------------------
    with tab1:
        st.subheader(
            "🎯 YCCĐ lớp đang gặp khó khăn"
        )

        rows = bang_thong_ke_bucket(
            data.get("yccd", {}),
            "YCCĐ"
        )

        if rows:
            st.dataframe(
                pd.DataFrame(
                    rows[:20]
                ),
                use_container_width=True,
                hide_index=True
            )

            can_day_lai = [
                x for x in rows
                if x["Số lượt đánh giá"] >= 3
                and x["Tỉ lệ đúng"] < 60
            ]

            if can_day_lai:
                st.warning(
                    f"⚠️ Có **{len(can_day_lai)} YCCĐ** có tỉ lệ đúng dưới 60%. "
                    "GV nên cân nhắc ôn lại, đổi cách giải thích hoặc tăng bài luyện."
                )

                for i, x in enumerate(
                    can_day_lai[:5],
                    start=1
                ):
                    st.write(
                        f"**{i}. {x['YCCĐ']}**"
                    )
                    st.caption(
                        f"Tỉ lệ đúng: {x['Tỉ lệ đúng']}% • "
                        f"{x['Số HS từng sai']} học sinh từng sai."
                    )

    # ------------------------------------------------------
    # TAB 2
    # ------------------------------------------------------
    with tab2:
        st.caption(
            "App ưu tiên HS cần GV quan tâm dựa trên kết quả chung, "
            "số YCCĐ còn yếu và xu hướng các lượt gần đây."
        )

        rows_hs = []

        for hs in hs_stats:
            tong = hs.get("tong_don_vi", 0)
            dung = hs.get("dung_don_vi", 0)

            ti_le = (
                dung / tong * 100
                if tong
                else 0
            )

            diem = hs.get("diem", [])
            so_luot = int(hs.get("so_luot", 0) or 0)

            # Chỉ phân tích đúng phạm vi dữ liệu GV đang chọn ở đầu trang.
            ma_hs = hs.get("ma", "")
            ls_hs_scope = [
                lan
                for lan in data.get("_lich_su", []) or []
                if str(lan.get("hoc_sinh_id", "")).strip().upper()
                == str(ma_hs).strip().upper()
            ]

            weak_scope = tong_hop_diem_yeu_tu_luot(
                ls_hs_scope,
                100
            )

            # Đếm YCCĐ yếu có đủ tối thiểu 2 lần đánh giá.
            yccd_yeu = set()
            nl_tong = {}

            for s in weak_scope:
                sl = int(s.get("Số lượt đánh giá", 0) or 0)
                sd = int(s.get("Số đúng", 0) or 0)
                tl = sd / sl if sl else 0

                if sl >= 2 and tl < 0.60:
                    yccd = str(s.get("YCCĐ", "")).strip()
                    if yccd:
                        yccd_yeu.add(yccd)

                nl = str(s.get("Năng lực", "")).strip()
                if nl:
                    bucket = nl_tong.setdefault(
                        nl,
                        {"tong": 0, "dung": 0}
                    )
                    bucket["tong"] += sl
                    bucket["dung"] += sd

            nl_yeu = []
            for nl, s_nl in nl_tong.items():
                if s_nl["tong"] >= 3:
                    tl_nl = s_nl["dung"] / s_nl["tong"]
                    if tl_nl < 0.60:
                        nl_yeu.append(nl)

            # Xu hướng điểm: so TB nửa đầu và nửa sau khi có >= 4 lượt.
            xu_huong = "Chưa đủ dữ liệu"
            if len(diem) >= 4:
                mid = len(diem) // 2
                dau = diem[:mid]
                sau = diem[mid:]
                tb_dau = sum(dau) / len(dau)
                tb_sau = sum(sau) / len(sau)
                chenhlech = tb_sau - tb_dau

                if chenhlech >= 0.5:
                    xu_huong = "Đang tiến bộ"
                elif chenhlech <= -0.5:
                    xu_huong = "Có dấu hiệu giảm"
                else:
                    xu_huong = "Ổn định"
            elif len(diem) >= 2:
                chenhlech = diem[-1] - diem[0]
                if chenhlech >= 0.5:
                    xu_huong = "Đang tiến bộ"
                elif chenhlech <= -0.5:
                    xu_huong = "Có dấu hiệu giảm"
                else:
                    xu_huong = "Ổn định"

            ly_do = []
            muc_uu_tien = 0

            if so_luot >= 1 and ti_le < 60:
                ly_do.append("Kết quả chung thấp")
                muc_uu_tien += 2

            if len(yccd_yeu) >= 1:
                ly_do.append(f"{len(yccd_yeu)} YCCĐ cần củng cố")
                muc_uu_tien += min(len(yccd_yeu), 3)

            if nl_yeu:
                ly_do.append(
                    "Yếu " + ", ".join(nl_yeu[:2])
                )
                muc_uu_tien += 2

            if xu_huong == "Có dấu hiệu giảm":
                ly_do.append("Kết quả gần đây giảm")
                muc_uu_tien += 2

            # Chỉ xếp vào nhóm cần hỗ trợ khi có ít nhất một dấu hiệu rõ.
            can_ho_tro = bool(ly_do)

            rows_hs.append({
                "Mã HS": ma_hs,
                "Họ và tên": hs.get("ho_ten", ""),
                "Lớp": hs.get("lop", ""),
                "Số lượt": so_luot,
                "Tỉ lệ đúng": round(ti_le, 1),
                "Điểm TB": round(
                    sum(diem) / len(diem),
                    1
                ) if diem else 0,
                "Xu hướng": xu_huong,
                "Cần hỗ trợ về": " • ".join(ly_do) if ly_do else "Chưa có dấu hiệu đáng chú ý",
                "_uu_tien": muc_uu_tien,
                "_can_ho_tro": can_ho_tro
            })

        nhom_can_ho_tro = [
            x for x in rows_hs
            if x["_can_ho_tro"]
        ]

        nhom_can_ho_tro.sort(
            key=lambda x: (
                -x["_uu_tien"],
                x["Tỉ lệ đúng"],
                x["Điểm TB"]
            )
        )

        if nhom_can_ho_tro:
            st.warning(
                f"👥 Có **{len(nhom_can_ho_tro)} học sinh** đang có dấu hiệu cần GV quan tâm."
            )

            bang_ho_tro = pd.DataFrame([
                {
                    "Mã HS": x["Mã HS"],
                    "Họ và tên": x["Họ và tên"],
                    "Lớp": x["Lớp"],
                    "Số lượt": x["Số lượt"],
                    "Tỉ lệ đúng": x["Tỉ lệ đúng"],
                    "Điểm TB": x["Điểm TB"],
                    "Xu hướng": x["Xu hướng"],
                    "Cần hỗ trợ về": x["Cần hỗ trợ về"]
                }
                for x in nhom_can_ho_tro
            ])

            st.dataframe(
                bang_ho_tro,
                use_container_width=True,
                hide_index=True
            )

            st.caption(
                "Danh sách được xếp ưu tiên theo mức độ cần hỗ trợ; "
                "không chỉ dựa vào ngưỡng điểm dưới 60%."
            )
        else:
            st.success(
                "Hiện chưa phát hiện học sinh có dấu hiệu cần hỗ trợ rõ rệt."
            )

        with st.expander(
            "Xem toàn bộ học sinh có dữ liệu",
            expanded=False
        ):
            tat_ca = pd.DataFrame([
                {
                    "Mã HS": x["Mã HS"],
                    "Họ và tên": x["Họ và tên"],
                    "Lớp": x["Lớp"],
                    "Số lượt": x["Số lượt"],
                    "Tỉ lệ đúng": x["Tỉ lệ đúng"],
                    "Điểm TB": x["Điểm TB"],
                    "Xu hướng": x["Xu hướng"]
                }
                for x in sorted(
                    rows_hs,
                    key=lambda z: (
                        z["Lớp"],
                        z["Họ và tên"]
                    )
                )
            ])

            st.dataframe(
                tat_ca,
                use_container_width=True,
                hide_index=True
            )

    # ------------------------------------------------------
    # TAB 3
    # ------------------------------------------------------
    with tab3:
        col_a, col_b = st.columns(2)

        with col_a:
            st.markdown(
                "### Theo mức độ"
            )

            rows_muc = bang_thong_ke_bucket(
                data.get("muc_do", {}),
                "Mức độ"
            )

            if rows_muc:
                st.dataframe(
                    pd.DataFrame(
                        rows_muc
                    ),
                    use_container_width=True,
                    hide_index=True
                )

        with col_b:
            st.markdown(
                "### Theo thành phần năng lực"
            )

            rows_nl = bang_thong_ke_bucket(
                data.get("nang_luc", {}),
                "Thành phần năng lực"
            )

            if rows_nl:
                st.dataframe(
                    pd.DataFrame(
                        rows_nl
                    ),
                    use_container_width=True,
                    hide_index=True
                )

        st.markdown(
            "### Theo dạng câu"
        )

        rows_dang = bang_thong_ke_bucket(
            data.get("dang_cau", {}),
            "Dạng câu"
        )

        if rows_dang:
            st.dataframe(
                pd.DataFrame(
                    rows_dang
                ),
                use_container_width=True,
                hide_index=True
            )

    # ------------------------------------------------------
    # TAB 4
    # ------------------------------------------------------
    with tab4:
        st.markdown(
            "### Theo bài"
        )

        rows_bai = bang_thong_ke_bucket(
            data.get("bai", {}),
            "Bài"
        )

        if rows_bai:
            st.dataframe(
                pd.DataFrame(
                    rows_bai[:30]
                ),
                use_container_width=True,
                hide_index=True
            )

        st.markdown(
            "### Theo chương"
        )

        rows_chuong = bang_thong_ke_bucket(
            data.get("chuong", {}),
            "Chương"
        )

        if rows_chuong:
            st.dataframe(
                pd.DataFrame(
                    rows_chuong
                ),
                use_container_width=True,
                hide_index=True
            )




# ==========================================================
# QUẢN LÝ DỮ LIỆU & TIẾN BỘ HỌC SINH
# ==========================================================
def sap_xep_lich_su_hs(lich_su):
    def key_time(x):
        iso = str(
            x.get(
                "thoi_gian_iso",
                ""
            )
        ).strip()

        if iso:
            return iso

        return str(
            x.get(
                "thoi_gian",
                ""
            )
        )

    return sorted(
        lich_su,
        key=key_time
    )


def tinh_tien_bo_hoc_sinh(lich_su):
    """
    Đánh giá xu hướng bằng dữ liệu thực tế, không dùng AI.
    So sánh trung bình tối đa 3 lượt đầu và 3 lượt gần nhất.
    """
    ls = sap_xep_lich_su_hs(
        lich_su
    )

    if not ls:
        return {
            "so_luot": 0,
            "diem_tb": 0,
            "ti_le_tb": 0,
            "diem_dau": 0,
            "diem_gan_day": 0,
            "chenh_lech": 0,
            "xu_huong": "Chưa có dữ liệu"
        }

    diem = [
        float(
            x.get(
                "diem",
                0
            )
            or 0
        )
        for x in ls
    ]

    ti_le = [
        float(
            x.get(
                "ti_le_dung_don_vi",
                0
            )
            or 0
        )
        for x in ls
    ]

    n = min(
        3,
        len(
            diem
        )
    )

    diem_dau = (
        sum(
            diem[:n]
        ) / n
    )

    diem_gan_day = (
        sum(
            diem[-n:]
        ) / n
    )

    chenh = (
        diem_gan_day
        - diem_dau
    )

    if len(diem) < 2:
        xu_huong = "Mới bắt đầu"
    elif chenh >= 1.5:
        xu_huong = "Tiến bộ rõ"
    elif chenh >= 0.5:
        xu_huong = "Có tiến bộ"
    elif chenh <= -1.5:
        xu_huong = "Giảm rõ"
    elif chenh <= -0.5:
        xu_huong = "Có dấu hiệu giảm"
    else:
        xu_huong = "Ổn định"

    return {
        "so_luot": len(
            ls
        ),
        "diem_tb": round(
            sum(
                diem
            ) / len(
                diem
            ),
            2
        ),
        "ti_le_tb": round(
            sum(
                ti_le
            ) / len(
                ti_le
            ),
            1
        ),
        "diem_dau": round(
            diem_dau,
            2
        ),
        "diem_gan_day": round(
            diem_gan_day,
            2
        ),
        "chenh_lech": round(
            chenh,
            2
        ),
        "xu_huong": xu_huong
    }




# ==========================================================
# XẾP HẠNG CÁ NHÂN HỌC SINH
# ==========================================================
def tinh_diem_xep_hang_hoc_sinh(lich_su):
    """
    Điểm xếp hạng phản ánh MỨC LÀM CHỦ HIỆN TẠI, không phải tốc độ tiến bộ.

    - 75%: kết quả tích lũy trên toàn bộ đơn vị đánh giá đã làm.
    - 25%: kết quả tối đa 3 lượt gần nhất.

    Số lượt luyện không được cộng điểm trực tiếp. Xu hướng tiến bộ được tính
    riêng bằng tinh_tien_bo_hoc_sinh().
    """
    ls = sap_xep_lich_su_hs(list(lich_su or []))

    if not ls:
        return {
            "diem_xep_hang": 0.0,
            "tich_luy": 0.0,
            "gan_day": 0.0,
            "so_luot": 0,
            "tong_don_vi": 0,
        }

    tong = 0
    dung = 0
    for lan in ls:
        t = int(lan.get("tong_don_vi", 0) or 0)
        d = int(lan.get("dung_don_vi", 0) or 0)
        if t > 0:
            tong += t
            dung += max(0, min(d, t))

    # Dữ liệu cũ có thể chưa lưu tổng đơn vị; khi đó dùng tỉ lệ từng lượt.
    if tong > 0:
        tich_luy = dung / tong * 100.0
    else:
        vals = [
            float(x.get("ti_le_dung_don_vi", 0) or 0)
            for x in ls
        ]
        tich_luy = sum(vals) / len(vals) if vals else 0.0

    gan = ls[-min(3, len(ls)):]
    tong_gan = sum(int(x.get("tong_don_vi", 0) or 0) for x in gan)
    dung_gan = sum(int(x.get("dung_don_vi", 0) or 0) for x in gan)

    if tong_gan > 0:
        gan_day = dung_gan / tong_gan * 100.0
    else:
        vals_gan = [
            float(x.get("ti_le_dung_don_vi", 0) or 0)
            for x in gan
        ]
        gan_day = sum(vals_gan) / len(vals_gan) if vals_gan else tich_luy

    diem_xep_hang = 0.75 * tich_luy + 0.25 * gan_day

    return {
        "diem_xep_hang": round(diem_xep_hang, 2),
        "tich_luy": round(tich_luy, 2),
        "gan_day": round(gan_day, 2),
        "so_luot": len(ls),
        "tong_don_vi": tong,
    }


def tinh_bang_xep_hang_lop(lop_chon):
    """
    Tính bảng xếp hạng nội bộ của MỘT LỚP.
    Giao diện HS chỉ lấy đúng dòng của chính HS; không lộ danh sách bạn khác.
    """
    lop_chon = str(lop_chon or "").strip()
    ds_hs = doc_danh_sach_hoc_sinh()
    lich_su_all = doc_lich_su_hoc_sinh()

    if lop_chon and lop_chon != "Tất cả":
        ds_hs = [
            hs for hs in ds_hs
            if str(hs.get("lop", "")).strip() == lop_chon
        ]

    rows = []
    for hs in ds_hs:
        ma = str(hs.get("ma_hoc_sinh", "")).strip().upper()
        if not ma:
            continue

        ls_hs = [
            x for x in lich_su_all
            if str(x.get("hoc_sinh_id", "")).strip().upper() == ma
        ]
        if not ls_hs:
            continue

        diem = tinh_diem_xep_hang_hoc_sinh(ls_hs)
        tien_bo = tinh_tien_bo_hoc_sinh(ls_hs)

        rows.append({
            "ma_hoc_sinh": ma,
            "ho_ten": hs.get("ho_ten", ""),
            "lop": hs.get("lop", ""),
            "diem_xep_hang": diem["diem_xep_hang"],
            "tich_luy": diem["tich_luy"],
            "gan_day": diem["gan_day"],
            "so_luot": diem["so_luot"],
            "tong_don_vi": diem["tong_don_vi"],
            "xu_huong": tien_bo["xu_huong"],
            "muc_thay_doi": tien_bo["chenh_lech"],
        })

    # Nếu không chỉ định lớp thì vẫn không dùng bảng này để công bố xếp hạng
    # toàn trường. Khi GV xem "Tất cả", hàm tổng hợp sẽ tính hạng theo từng lớp.
    rows.sort(
        key=lambda r: (
            -float(r.get("diem_xep_hang", 0) or 0),
            -float(r.get("gan_day", 0) or 0),
            -float(r.get("tich_luy", 0) or 0),
            -int(r.get("tong_don_vi", 0) or 0),
            str(r.get("ho_ten", "")),
        )
    )

    for idx, row in enumerate(rows, start=1):
        row["hang"] = idx
        row["si_so_co_du_lieu"] = len(rows)

    return rows


def lay_xep_hang_ca_nhan_hs(ma_hoc_sinh, lop):
    ma = str(ma_hoc_sinh or "").strip().upper()
    for row in tinh_bang_xep_hang_lop(lop):
        if str(row.get("ma_hoc_sinh", "")).strip().upper() == ma:
            return row
    return None


def hien_thi_xep_hang_ca_nhan_hs(ma_hoc_sinh, lop):
    """Chỉ hiển thị thứ hạng của chính HS, không lộ tên/điểm của bạn khác."""
    row = lay_xep_hang_ca_nhan_hs(ma_hoc_sinh, lop)
    if not row:
        return

    st.markdown("### 🏆 Xếp hạng của em")
    c1, c2, c3 = st.columns(3)

    with c1:
        st.metric(
            "Hạng hiện tại",
            f"{row['hang']}/{row['si_so_co_du_lieu']}"
        )

    with c2:
        st.metric(
            "Mức làm chủ tích lũy",
            f"{row['tich_luy']:.0f}%"
        )

    with c3:
        xu_huong = str(row.get("xu_huong", ""))
        if xu_huong in ["Tiến bộ rõ", "Có tiến bộ"]:
            nhan = f"📈 {xu_huong}"
        elif xu_huong in ["Giảm rõ", "Có dấu hiệu giảm"]:
            nhan = f"📉 {xu_huong}"
        elif xu_huong == "Mới bắt đầu":
            nhan = "🌱 Mới bắt đầu"
        else:
            nhan = f"➡️ {xu_huong or 'Ổn định'}"
        st.metric("Xu hướng học tập", nhan)

    st.caption(
        "Em chỉ xem được vị trí của chính mình. Hạng được tính trong số học sinh "
        "cùng lớp đã có dữ liệu làm bài: 75% mức làm chủ tích lũy + 25% kết quả "
        "tối đa 3 lượt gần nhất. Xu hướng tiến bộ được tính riêng."
    )


def phan_tich_tien_bo_thuc_theo_don_vi(lich_su):
    """
    Đo tiến bộ trên CÙNG một đơn vị đánh giá:
    YCCĐ × mức độ × thành phần năng lực.

    Nhờ vậy không kết luận sai kiểu:
    lần đầu làm Vận dụng 40%, lần sau làm Nhận biết 90% => tiến bộ.
    """
    ls = sap_xep_lich_su_hs(
        lich_su
    )

    groups = {}

    for thu_tu_luot, lan in enumerate(
        ls,
        start=1
    ):
        for item in lan.get(
            "chi_tiet",
            []
        ) or []:
            for unit in item.get(
                "don_vi_danh_gia",
                []
            ) or []:
                yccd = str(
                    unit.get(
                        "yccd",
                        ""
                    )
                ).strip()

                muc_do = str(
                    unit.get(
                        "muc_do",
                        ""
                    )
                ).strip()

                nang_luc = str(
                    unit.get(
                        "nang_luc",
                        ""
                    )
                ).strip()

                if not yccd:
                    continue

                key = khoa_nang_luc(
                    yccd,
                    muc_do,
                    nang_luc
                )

                g = groups.setdefault(
                    key,
                    {
                        "yccd": yccd,
                        "muc_do": muc_do,
                        "nang_luc": nang_luc,
                        "ket_qua": []
                    }
                )

                g["ket_qua"].append({
                    "luot": thu_tu_luot,
                    "dung": 1 if unit.get("dung") else 0
                })

    rows = []

    for g in groups.values():
        values = [
            x["dung"]
            for x in g["ket_qua"]
        ]

        n = len(values)

        # Cần ít nhất 3 lần đánh giá cùng đơn vị mới bắt đầu kết luận.
        if n < 3:
            ti_le_dau = (
                sum(values) / n * 100
                if n
                else 0
            )
            ti_le_gan = ti_le_dau
            chenh = 0
            trang_thai = "Chưa đủ dữ liệu"

        else:
            window = min(
                3,
                max(
                    1,
                    n // 2
                )
            )

            dau = values[:window]
            gan = values[-window:]

            ti_le_dau = (
                sum(dau) / len(dau) * 100
            )

            ti_le_gan = (
                sum(gan) / len(gan) * 100
            )

            chenh = (
                ti_le_gan
                - ti_le_dau
            )

            # Trạng thái vừa phản ánh mức hiện tại,
            # vừa phản ánh xu hướng trên cùng mức độ/năng lực.
            if n >= 6 and ti_le_gan >= 85 and chenh >= 0:
                trang_thai = "Thành thạo"
            elif ti_le_gan >= 75 and n >= 4:
                trang_thai = "Đạt"
            elif chenh >= 15:
                trang_thai = "Đang tiến bộ"
            elif ti_le_gan < 60:
                trang_thai = "Cần củng cố"
            else:
                trang_thai = "Ổn định"

        rows.append({
            "YCCĐ": g["yccd"],
            "Mức độ": g["muc_do"],
            "Năng lực": g["nang_luc"],
            "Số lần cùng chuẩn": n,
            "Tỉ lệ đầu (%)": round(
                ti_le_dau,
                1
            ),
            "Tỉ lệ gần đây (%)": round(
                ti_le_gan,
                1
            ),
            "Thay đổi (điểm %)": round(
                chenh,
                1
            ),
            "Trạng thái": trang_thai
        })

    # Ưu tiên hiển thị phần cần củng cố / đang tiến bộ trước.
    thu_tu = {
        "Cần củng cố": 0,
        "Đang tiến bộ": 1,
        "Ổn định": 2,
        "Đạt": 3,
        "Thành thạo": 4,
        "Chưa đủ dữ liệu": 5
    }

    rows.sort(
        key=lambda x: (
            thu_tu.get(
                x["Trạng thái"],
                9
            ),
            x["Tỉ lệ gần đây (%)"]
        )
    )

    return rows


def tom_tat_tien_bo_thuc(lich_su):
    rows = phan_tich_tien_bo_thuc_theo_don_vi(
        lich_su
    )

    hop_le = [
        r
        for r in rows
        if r["Trạng thái"]
        != "Chưa đủ dữ liệu"
    ]

    if not hop_le:
        return {
            "so_don_vi": 0,
            "dang_tien_bo": 0,
            "dat": 0,
            "thanh_thao": 0,
            "can_cung_co": 0
        }

    return {
        "so_don_vi": len(
            hop_le
        ),
        "dang_tien_bo": sum(
            1
            for r in hop_le
            if r["Trạng thái"]
            == "Đang tiến bộ"
        ),
        "dat": sum(
            1
            for r in hop_le
            if r["Trạng thái"]
            == "Đạt"
        ),
        "thanh_thao": sum(
            1
            for r in hop_le
            if r["Trạng thái"]
            == "Thành thạo"
        ),
        "can_cung_co": sum(
            1
            for r in hop_le
            if r["Trạng thái"]
            == "Cần củng cố"
        )
    }


def hien_thi_the_tien_bo_day_du(row, index=None):
    """
    Dùng container thay cho ô bảng đối với nội dung dài,
    để YCCĐ không bị cắt bằng dấu ...
    """
    trang_thai = row.get(
        "Trạng thái",
        ""
    )

    with st.container(
        border=True
    ):
        if index is not None:
            st.markdown(
                f"### {index}. {trang_thai}"
            )
        else:
            st.markdown(
                f"### {trang_thai}"
            )

        st.markdown(
            "**Yêu cầu cần đạt**"
        )
        st.write(
            row.get(
                "YCCĐ",
                ""
            )
        )

        c1, c2, c3 = st.columns(
            3
        )

        with c1:
            st.write(
                "**Mức độ**"
            )
            st.write(
                row.get(
                    "Mức độ",
                    ""
                )
            )

        with c2:
            st.write(
                "**Thành phần năng lực**"
            )
            st.write(
                row.get(
                    "Năng lực",
                    ""
                )
            )

        with c3:
            st.write(
                "**Số lần cùng chuẩn**"
            )
            st.write(
                row.get(
                    "Số lần cùng chuẩn",
                    0
                )
            )

        c4, c5, c6 = st.columns(
            3
        )

        with c4:
            st.metric(
                "Ban đầu",
                f"{row.get('Tỉ lệ đầu (%)', 0):.0f}%"
            )

        with c5:
            st.metric(
                "Gần đây",
                f"{row.get('Tỉ lệ gần đây (%)', 0):.0f}%"
            )

        with c6:
            st.metric(
                "Thay đổi",
                f"{row.get('Thay đổi (điểm %)', 0):+.0f} điểm %"
            )



def tong_hop_hoc_sinh_theo_lop(lop_chon):
    ds_hs = doc_danh_sach_hoc_sinh()
    lich_su = doc_lich_su_hoc_sinh()

    if lop_chon and lop_chon != "Tất cả":
        ds_hs = [
            hs
            for hs in ds_hs
            if str(hs.get("lop", "")).strip() == lop_chon
        ]

    # Hạng luôn được tính TRONG TỪNG LỚP, kể cả khi GV chọn "Tất cả".
    cac_lop = sorted({
        str(hs.get("lop", "")).strip()
        for hs in ds_hs
        if str(hs.get("lop", "")).strip()
    })
    xep_hang_map = {}
    for lop in cac_lop:
        for r in tinh_bang_xep_hang_lop(lop):
            xep_hang_map[str(r.get("ma_hoc_sinh", "")).strip().upper()] = r

    rows = []

    for hs in ds_hs:
        ma = str(
            hs.get(
                "ma_hoc_sinh",
                ""
            )
        ).strip().upper()

        ls_hs = [
            x
            for x in lich_su
            if str(
                x.get(
                    "hoc_sinh_id",
                    ""
                )
            ).strip().upper() == ma
        ]

        tb = tinh_tien_bo_hoc_sinh(
            ls_hs
        )

        tb_thuc = tom_tat_tien_bo_thuc(
            ls_hs
        )

        profile = tao_ho_so_tu_lich_su(
            ma
        )

        weak = tom_tat_diem_yeu(
            profile,
            3
        )

        noi_dung_yeu = " | ".join(
            [
                str(
                    x.get(
                        "yccd",
                        ""
                    )
                ).strip()
                for x in weak
                if str(
                    x.get(
                        "yccd",
                        ""
                    )
                ).strip()
            ]
        )

        xh_row = xep_hang_map.get(ma, {})

        rows.append({
            "Mã học sinh": ma,
            "Họ và tên": hs.get(
                "ho_ten",
                ""
            ),
            "Lớp": hs.get(
                "lop",
                ""
            ),
            "Khối": hs.get(
                "khoi",
                ""
            ),
            "Hạng lớp": xh_row.get("hang", "—"),
            "Điểm xếp hạng": xh_row.get("diem_xep_hang", 0),
            "Mức làm chủ tích lũy (%)": xh_row.get("tich_luy", 0),
            "Kết quả 3 lượt gần (%)": xh_row.get("gan_day", 0),
            "Số lượt": tb["so_luot"],
            "Điểm TB": tb["diem_tb"],
            "Tỉ lệ đúng TB (%)": tb["ti_le_tb"],
            "Điểm giai đoạn đầu": tb["diem_dau"],
            "Điểm gần đây": tb["diem_gan_day"],
            "Mức thay đổi": tb["chenh_lech"],
            "Xu hướng": tb["xu_huong"],
            "Đơn vị đang tiến bộ": tb_thuc["dang_tien_bo"],
            "Đơn vị đạt": tb_thuc["dat"],
            "Đơn vị thành thạo": tb_thuc["thanh_thao"],
            "Đơn vị cần củng cố": tb_thuc["can_cung_co"],
            "Nội dung cần ưu tiên": noi_dung_yeu
        })

    rows.sort(
        key=lambda x: (
            x["Lớp"],
            999999 if x.get("Hạng lớp") == "—" else int(x.get("Hạng lớp", 999999)),
            x["Họ và tên"]
        )
    )

    return rows


def tao_excel_tong_hop_hoc_sinh_lop(lop_chon):
    """
    Workbook 1: theo dõi từng HS của lớp.
    """
    import io

    rows = tong_hop_hoc_sinh_theo_lop(
        lop_chon
    )

    lich_su = doc_lich_su_hoc_sinh()

    ma_hs = {
        r["Mã học sinh"]
        for r in rows
    }

    chi_tiet = []

    for lan in lich_su:
        ma = str(
            lan.get(
                "hoc_sinh_id",
                ""
            )
        ).strip().upper()

        if ma not in ma_hs:
            continue

        chi_tiet.append({
            "Mã học sinh": ma,
            "Họ và tên": lan.get(
                "ho_ten",
                ""
            ),
            "Bắt đầu": lan.get(
                "bat_dau_luc",
                ""
            ),
            "Nộp bài": lan.get(
                "nop_bai_luc",
                lan.get("thoi_gian", "")
            ),
            "Thời lượng làm (phút)": lan.get(
                "thoi_luong_lam_phut",
                0
            ),
            "Thời gian quy định (phút)": lan.get(
                "thoi_gian_quy_dinh_phut",
                0
            ),
            "Chế độ": lan.get(
                "che_do",
                ""
            ),
            "Bài luyện": lan.get(
                "ten_luot",
                ""
            ),
            "Số câu": lan.get(
                "tong_so_cau",
                0
            ),
            "Tỉ lệ đúng (%)": lan.get(
                "ti_le_dung_don_vi",
                0
            ),
            "Điểm chính thức": lan.get(
                "diem_chinh_thuc",
                lan.get("diem", 0)
            ),
            "Thang điểm": lan.get(
                "thang_diem",
                10
            ),
            "Điểm quy đổi /10": lan.get(
                "diem",
                0
            )
        })

    output = io.BytesIO()

    with pd.ExcelWriter(
        output,
        engine="openpyxl"
    ) as writer:
        pd.DataFrame(
            rows
        ).to_excel(
            writer,
            index=False,
            sheet_name="Tong hop tung HS"
        )

        pd.DataFrame(
            chi_tiet
        ).to_excel(
            writer,
            index=False,
            sheet_name="Lich su luot lam"
        )

        # Sheet điểm yếu chi tiết theo HS
        weak_rows = []

        for r in rows:
            ma = r[
                "Mã học sinh"
            ]

            profile = tao_ho_so_tu_lich_su(
                ma
            )

            for x in tom_tat_diem_yeu(
                profile,
                10
            ):
                weak_rows.append({
                    "Mã học sinh": ma,
                    "Họ và tên": r[
                        "Họ và tên"
                    ],
                    "YCCĐ": x.get(
                        "yccd",
                        ""
                    ),
                    "Mức độ": x.get(
                        "muc_do",
                        ""
                    ),
                    "Năng lực": x.get(
                        "nang_luc",
                        ""
                    ),
                    "Số lần": x.get(
                        "so_lan",
                        0
                    ),
                    "Số đúng": x.get(
                        "so_dung",
                        0
                    ),
                    "Tỉ lệ đúng (%)": round(
                        x.get(
                            "ti_le_dung",
                            0
                        )
                        * 100,
                        1
                    )
                })

        pd.DataFrame(
            weak_rows
        ).to_excel(
            writer,
            index=False,
            sheet_name="Diem yeu tung HS"
        )

        tien_bo_rows = []

        for r in rows:
            ma = r[
                "Mã học sinh"
            ]

            ls_hs = lay_lich_su_cua_hoc_sinh(
                ma
            )

            for tb_row in phan_tich_tien_bo_thuc_theo_don_vi(
                ls_hs
            ):
                item_tb = {
                    "Mã học sinh": ma,
                    "Họ và tên": r[
                        "Họ và tên"
                    ]
                }
                item_tb.update(
                    tb_row
                )
                tien_bo_rows.append(
                    item_tb
                )

        pd.DataFrame(
            tien_bo_rows
        ).to_excel(
            writer,
            index=False,
            sheet_name="Tien bo cung chuan"
        )

        # Căn cột cơ bản
        for ws in writer.book.worksheets:
            ws.freeze_panes = "A2"

            for col_cells in ws.columns:
                max_len = 0
                letter = col_cells[
                    0
                ].column_letter

                for cell in col_cells:
                    value = (
                        ""
                        if cell.value is None
                        else str(
                            cell.value
                        )
                    )

                    max_len = max(
                        max_len,
                        len(
                            value
                        )
                    )

                ws.column_dimensions[
                    letter
                ].width = min(
                    max(
                        max_len + 2,
                        10
                    ),
                    45
                )

            for cell in ws[
                1
            ]:
                cell.font = cell.font.copy(
                    bold=True
                )

    output.seek(
        0
    )

    return output.getvalue()


def tao_excel_tong_hop_chung_lop(lop_chon):
    """
    Workbook 2: bức tranh chung của lớp.
    """
    import io

    data = tong_hop_du_lieu_lop(
        lop_chon
    )

    output = io.BytesIO()

    rows_hs = tong_hop_hoc_sinh_theo_lop(
        lop_chon
    )

    with pd.ExcelWriter(
        output,
        engine="openpyxl"
    ) as writer:
        # Tổng quan
        tong_don_vi = sum(
            x.get(
                "tong_don_vi",
                0
            )
            for x in data.get(
                "hoc_sinh",
                {}
            ).values()
        )

        tong_dung = sum(
            x.get(
                "dung_don_vi",
                0
            )
            for x in data.get(
                "hoc_sinh",
                {}
            ).values()
        )

        diem_all = [
            d
            for hs in data.get(
                "hoc_sinh",
                {}
            ).values()
            for d in hs.get(
                "diem",
                []
            )
        ]

        tong_quan = [{
            "Lớp": lop_chon,
            "Số học sinh có dữ liệu": len(
                data.get(
                    "hoc_sinh",
                    {}
                )
            ),
            "Tổng lượt luyện": data.get(
                "so_luot",
                0
            ),
            "Tỉ lệ đúng toàn lớp (%)": round(
                (
                    tong_dung
                    / tong_don_vi
                    * 100
                )
                if tong_don_vi
                else 0,
                1
            ),
            "Điểm TB các lượt": round(
                (
                    sum(
                        diem_all
                    )
                    / len(
                        diem_all
                    )
                )
                if diem_all
                else 0,
                2
            )
        }]

        pd.DataFrame(
            tong_quan
        ).to_excel(
            writer,
            index=False,
            sheet_name="Tong quan lop"
        )

        pd.DataFrame(
            rows_hs
        ).to_excel(
            writer,
            index=False,
            sheet_name="Tong hop HS"
        )

        for bucket, sheet_name, col_name in [
            (
                "yccd",
                "Theo YCCD",
                "YCCĐ"
            ),
            (
                "muc_do",
                "Theo muc do",
                "Mức độ"
            ),
            (
                "nang_luc",
                "Theo nang luc",
                "Thành phần năng lực"
            ),
            (
                "dang_cau",
                "Theo dang cau",
                "Dạng câu"
            ),
            (
                "bai",
                "Theo bai",
                "Bài"
            ),
            (
                "chuong",
                "Theo chuong",
                "Chương"
            )
        ]:
            rows_bucket = bang_thong_ke_bucket(
                data.get(
                    bucket,
                    {}
                ),
                col_name
            )

            pd.DataFrame(
                rows_bucket
            ).to_excel(
                writer,
                index=False,
                sheet_name=sheet_name[:31]
            )

        for ws in writer.book.worksheets:
            ws.freeze_panes = "A2"

            for col_cells in ws.columns:
                max_len = 0
                letter = col_cells[
                    0
                ].column_letter

                for cell in col_cells:
                    value = (
                        ""
                        if cell.value is None
                        else str(
                            cell.value
                        )
                    )

                    max_len = max(
                        max_len,
                        len(
                            value
                        )
                    )

                ws.column_dimensions[
                    letter
                ].width = min(
                    max(
                        max_len + 2,
                        10
                    ),
                    45
                )

            for cell in ws[
                1
            ]:
                cell.font = cell.font.copy(
                    bold=True
                )

    output.seek(
        0
    )

    return output.getvalue()


def du_lieu_va_tien_bo_hoc_sinh():
    st.header(
        "🗂️ DỮ LIỆU & TIẾN BỘ HỌC SINH"
    )

    st.caption(
        "Theo dõi sự tiến bộ từng học sinh và xuất báo cáo theo lớp. "
        "Toàn bộ phần này không dùng API."
    )

    ds_lop = lay_danh_sach_lop_tu_hoc_sinh()

    if not ds_lop:
        st.info(
            "Chưa có lớp/học sinh."
        )
        return

    lop = st.selectbox(
        "Chọn lớp",
        ds_lop,
        key="gv_data_progress_class"
    )

    rows = tong_hop_hoc_sinh_theo_lop(
        lop
    )

    if not rows:
        st.info(
            "Lớp này chưa có dữ liệu."
        )
        return

    tab1, tab2, tab3 = st.tabs([
        "👤 Tiến bộ từng học sinh",
        "👥 Tổng hợp cả lớp",
        "⬇️ Xuất dữ liệu"
    ])

    # ------------------------------------------------------
    # TAB 1
    # ------------------------------------------------------
    with tab1:
        labels = [
            f"{r['Mã học sinh']} – "
            f"{r['Họ và tên']}"
            for r in rows
        ]

        selected = st.selectbox(
            "Chọn học sinh",
            labels,
            key="gv_progress_student"
        )

        idx = labels.index(
            selected
        )

        hs_row = rows[
            idx
        ]

        ma = hs_row[
            "Mã học sinh"
        ]

        lich_su_hs = lay_lich_su_cua_hoc_sinh(
            ma
        )

        tb = tinh_tien_bo_hoc_sinh(
            lich_su_hs
        )

        c1, c2, c3, c4 = st.columns(
            4
        )

        with c1:
            st.metric(
                "Số lượt",
                tb[
                    "so_luot"
                ]
            )

        with c2:
            st.metric(
                "Điểm TB",
                f"{tb['diem_tb']:.1f}"
            )

        with c3:
            st.metric(
                "Điểm gần đây",
                f"{tb['diem_gan_day']:.1f}"
            )

        with c4:
            st.metric(
                "Mức thay đổi",
                f"{tb['chenh_lech']:+.1f}"
            )

        if tb[
            "xu_huong"
        ] in [
            "Tiến bộ rõ",
            "Có tiến bộ"
        ]:
            st.success(
                f"📈 **Xu hướng: {tb['xu_huong']}**"
            )
        elif tb[
            "xu_huong"
        ] in [
            "Giảm rõ",
            "Có dấu hiệu giảm"
        ]:
            st.warning(
                f"📉 **Xu hướng: {tb['xu_huong']}**"
            )
        else:
            st.info(
                f"➡️ **Xu hướng: {tb['xu_huong']}**"
            )

        if lich_su_hs:
            lich_su_sort = sap_xep_lich_su_hs(
                lich_su_hs
            )

            df_tien_bo = pd.DataFrame([
                {
                    "Lượt": i,
                    "Thời gian": x.get(
                        "thoi_gian",
                        ""
                    ),
                    "Điểm": float(
                        x.get(
                            "diem",
                            0
                        )
                        or 0
                    ),
                    "Tỉ lệ đúng": float(
                        x.get(
                            "ti_le_dung_don_vi",
                            0
                        )
                        or 0
                    )
                }
                for i, x in enumerate(
                    lich_su_sort,
                    start=1
                )
            ])

            st.markdown(
                "### 📈 Đường tiến bộ"
            )

            if not df_tien_bo.empty:
                st.line_chart(
                    df_tien_bo.set_index(
                        "Lượt"
                    )[
                        [
                            "Điểm"
                        ]
                    ]
                )

            st.dataframe(
                df_tien_bo,
                use_container_width=True,
                hide_index=True
            )

        profile = tao_ho_so_tu_lich_su(
            ma
        )

        # ==================================================
        # PHẦN DƯỚI: GỌN, TẬP TRUNG VÀO 3 THÀNH PHẦN NĂNG LỰC
        # ==================================================
        st.markdown("---")
        st.markdown(
            "## 🧠 NĂNG LỰC SINH HỌC"
        )

        st.caption(
            "Tổng hợp theo 3 thành phần năng lực để GV nhanh chóng nhận ra "
            "điểm mạnh, điểm cần hỗ trợ và điều chỉnh cách dạy/cách học."
        )

        # Tổng hợp trực tiếp từ hồ sơ động đã có.
        stats_profile = list(
            profile.get(
                "stats",
                {}
            ).values()
        )

        tong_hop_nl = {}

        for s in stats_profile:
            nl = str(
                s.get(
                    "nang_luc",
                    ""
                )
            ).strip()

            if not nl:
                continue

            item_nl = tong_hop_nl.setdefault(
                nl,
                {
                    "so_lan": 0,
                    "so_dung": 0
                }
            )

            item_nl["so_lan"] += int(
                s.get(
                    "so_lan",
                    0
                )
                or 0
            )

            item_nl["so_dung"] += int(
                s.get(
                    "so_dung",
                    0
                )
                or 0
            )

        def _trang_thai_nang_luc(ti_le, so_lan):
            if so_lan < 3:
                return "Chưa đủ dữ liệu"

            if ti_le >= 0.80:
                return "Đạt tốt"

            if ti_le >= 0.65:
                return "Đạt"

            if ti_le >= 0.50:
                return "Đang củng cố"

            return "Cần hỗ trợ"


        nl_rows = []

        for nl_name in THANH_PHAN_NANG_LUC:
            s = tong_hop_nl.get(
                nl_name,
                {
                    "so_lan": 0,
                    "so_dung": 0
                }
            )

            so_lan_nl = int(
                s.get(
                    "so_lan",
                    0
                )
                or 0
            )

            so_dung_nl = int(
                s.get(
                    "so_dung",
                    0
                )
                or 0
            )

            ti_le_nl = (
                so_dung_nl / so_lan_nl
                if so_lan_nl > 0
                else 0
            )

            nl_rows.append({
                "Năng lực": nl_name,
                "Số lần đánh giá": so_lan_nl,
                "Tỉ lệ đúng": ti_le_nl,
                "Trạng thái": _trang_thai_nang_luc(
                    ti_le_nl,
                    so_lan_nl
                )
            })

        # Hiển thị 3 thẻ ngang, không tạo bảng dài.
        cols_nl = st.columns(
            3
        )

        for i_nl, row_nl in enumerate(
            nl_rows
        ):
            with cols_nl[
                i_nl
            ]:
                with st.container(
                    border=True
                ):
                    st.markdown(
                        f"### {row_nl['Năng lực']}"
                    )

                    st.metric(
                        "Tỉ lệ đúng",
                        f"{row_nl['Tỉ lệ đúng'] * 100:.0f}%"
                    )

                    st.caption(
                        f"Số lần đánh giá: {row_nl['Số lần đánh giá']}"
                    )

                    trang_thai_nl = row_nl[
                        "Trạng thái"
                    ]

                    if trang_thai_nl == "Đạt tốt":
                        st.success(
                            "⭐ Đạt tốt"
                        )
                    elif trang_thai_nl == "Đạt":
                        st.success(
                            "🟢 Đạt"
                        )
                    elif trang_thai_nl == "Đang củng cố":
                        st.info(
                            "🟡 Đang củng cố"
                        )
                    elif trang_thai_nl == "Cần hỗ trợ":
                        st.warning(
                            "🔴 Cần hỗ trợ"
                        )
                    else:
                        st.info(
                            "⚪ Chưa đủ dữ liệu"
                        )

        # ==================================================
        # NỘI DUNG CẦN CHÚ Ý - CHỈ TOP 5
        # ==================================================
        st.markdown(
            "## 🎯 NỘI DUNG GV CẦN CHÚ Ý"
        )

        weak = tom_tat_diem_yeu(
            profile,
            5
        )

        if weak:
            for i_w, x in enumerate(
                weak,
                start=1
            ):
                with st.container(
                    border=True
                ):
                    st.markdown(
                        f"**{i_w}. {x.get('yccd', '')}**"
                    )

                    c_w1, c_w2, c_w3 = st.columns(
                        3
                    )

                    with c_w1:
                        st.caption(
                            "Mức độ"
                        )
                        st.write(
                            x.get(
                                "muc_do",
                                ""
                            )
                        )

                    with c_w2:
                        st.caption(
                            "Năng lực"
                        )
                        st.write(
                            x.get(
                                "nang_luc",
                                ""
                            )
                        )

                    with c_w3:
                        st.metric(
                            "Tỉ lệ đúng",
                            f"{x.get('ti_le_dung', 0) * 100:.0f}%"
                        )
        else:
            st.success(
                "Chưa phát hiện nội dung yếu nổi bật."
            )

        # ==================================================
        # GỢI Ý HÀNH ĐỘNG CHO GV - KHÔNG DÙNG API
        # ==================================================
        st.markdown(
            "## 🧭 GỢI Ý ĐIỀU CHỈNH DẠY – HỌC"
        )

        nl_can_ho_tro = [
            r
            for r in nl_rows
            if r["Trạng thái"] in [
                "Cần hỗ trợ",
                "Đang củng cố"
            ]
        ]

        if nl_can_ho_tro:
            ten_nl = ", ".join(
                r["Năng lực"]
                for r in nl_can_ho_tro
            )

            st.warning(
                f"Ưu tiên hỗ trợ các thành phần năng lực: **{ten_nl}**. "
                "Nên tăng câu luyện đúng các YCCĐ liên quan, giữ mức độ phù hợp "
                "và theo dõi lại sau vài lượt."
            )
        else:
            st.success(
                "Các thành phần năng lực hiện chưa có dấu hiệu yếu rõ. "
                "Có thể tiếp tục duy trì và tăng dần độ khó ở nội dung phù hợp."
            )

        if weak:
            st.info(
                "Trong lượt luyện tiếp theo, app nên ưu tiên một phần câu ở "
                "**3–5 YCCĐ yếu nhất**, nhưng vẫn xen kẽ nội dung đã làm tốt "
                "để kiểm tra độ bền kiến thức."
            )

    # ------------------------------------------------------
    # TAB 2
    # ------------------------------------------------------
    with tab2:
        st.caption(
            "Bảng dưới chỉ giữ các cột ngắn để dễ nhìn. "
            "Nội dung YCCĐ cần ưu tiên được hiển thị đầy đủ ở các thẻ bên dưới."
        )

        cot_ngan = [
            "Mã học sinh",
            "Họ và tên",
            "Lớp",
            "Số lượt",
            "Điểm TB",
            "Tỉ lệ đúng TB (%)",
            "Xu hướng",
            "Đơn vị đang tiến bộ",
            "Đơn vị đạt",
            "Đơn vị thành thạo",
            "Đơn vị cần củng cố"
        ]

        df_lop_ngan = pd.DataFrame(
            rows
        )

        cot_co = [
            c
            for c in cot_ngan
            if c in df_lop_ngan.columns
        ]

        st.dataframe(
            df_lop_ngan[
                cot_co
            ],
            use_container_width=True,
            hide_index=True,
            height=420
        )

        st.markdown(
            "### 👤 Xem đầy đủ từng học sinh"
        )

        for r_hs in rows:
            with st.expander(
                f"{r_hs.get('Mã học sinh', '')} – "
                f"{r_hs.get('Họ và tên', '')} – "
                f"{r_hs.get('Xu hướng', '')}",
                expanded=False
            ):
                s1, s2, s3, s4 = st.columns(
                    4
                )

                with s1:
                    st.metric(
                        "Số lượt",
                        r_hs.get(
                            "Số lượt",
                            0
                        )
                    )

                with s2:
                    st.metric(
                        "Điểm TB",
                        r_hs.get(
                            "Điểm TB",
                            0
                        )
                    )

                with s3:
                    st.metric(
                        "Đang tiến bộ",
                        r_hs.get(
                            "Đơn vị đang tiến bộ",
                            0
                        )
                    )

                with s4:
                    st.metric(
                        "Cần củng cố",
                        r_hs.get(
                            "Đơn vị cần củng cố",
                            0
                        )
                    )

                st.markdown(
                    "**Nội dung cần ưu tiên:**"
                )

                noi_dung_day_du = str(
                    r_hs.get(
                        "Nội dung cần ưu tiên",
                        ""
                    )
                ).strip()

                if noi_dung_day_du:
                    for nd_uu_tien in noi_dung_day_du.split(
                        " | "
                    ):
                        with st.container(
                            border=True
                        ):
                            st.write(
                                nd_uu_tien
                            )
                else:
                    st.write(
                        "Chưa có nội dung cần ưu tiên."
                    )

        tien_bo = sum(
            1
            for r in rows
            if r[
                "Xu hướng"
            ] in [
                "Tiến bộ rõ",
                "Có tiến bộ"
            ]
        )

        can_ho_tro = sum(
            1
            for r in rows
            if r[
                "Xu hướng"
            ] in [
                "Giảm rõ",
                "Có dấu hiệu giảm"
            ]
        )

        d1, d2, d3 = st.columns(
            3
        )

        with d1:
            st.metric(
                "Sĩ số",
                len(
                    rows
                )
            )

        with d2:
            st.metric(
                "Đang tiến bộ",
                tien_bo
            )

        with d3:
            st.metric(
                "Cần chú ý",
                can_ho_tro
            )

    # ------------------------------------------------------
    # TAB 3
    # ------------------------------------------------------
    with tab3:
        st.markdown(
            "### ⬇️ Xuất báo cáo"
        )

        st.caption(
            "Có 2 loại file: báo cáo từng học sinh của lớp và báo cáo bức tranh chung của lớp."
        )

        file_hs = tao_excel_tong_hop_hoc_sinh_lop(
            lop
        )

        file_lop = tao_excel_tong_hop_chung_lop(
            lop
        )

        f1, f2 = st.columns(
            2
        )

        with f1:
            st.download_button(
                "⬇️ TẢI TỔNG HỢP TỪNG HỌC SINH",
                data=file_hs,
                file_name=(
                    f"tong_hop_tung_hoc_sinh_{chuan_hoa_ten_lop(lop)}.xlsx"
                ),
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "spreadsheetml.sheet"
                ),
                use_container_width=True,
                key="gv_download_student_progress"
            )

            st.caption(
                "Gồm: tổng hợp từng HS, lịch sử lượt làm, điểm yếu và tiến bộ theo cùng YCCĐ/mức độ/năng lực."
            )

        with f2:
            st.download_button(
                "⬇️ TẢI TỔNG HỢP CHUNG CỦA LỚP",
                data=file_lop,
                file_name=(
                    f"tong_hop_chung_lop_{chuan_hoa_ten_lop(lop)}.xlsx"
                ),
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "spreadsheetml.sheet"
                ),
                use_container_width=True,
                key="gv_download_class_summary"
            )

            st.caption(
                "Gồm: tổng quan lớp, YCCĐ, mức độ, năng lực, dạng câu, bài và chương."
            )




# ==========================================================
# NGÂN HÀNG HẠT GIỐNG + NGÂN HÀNG TỐT NGHIỆP
# ==========================================================
def doc_ngan_hang_hat_giong():
    return doc_json_list(SEED_BANK_PATH)


def luu_ngan_hang_hat_giong(ds):
    luu_json_list(SEED_BANK_PATH, ds)


def _anh_base64_bytes(value):
    """Giải mã ảnh nhúng base64; chấp nhận cả data URI."""
    s = str(value or "").strip()
    if not s:
        return b""
    if s.startswith("data:") and "," in s:
        s = s.split(",", 1)[1]
    try:
        return base64.b64decode(s, validate=False)
    except Exception:
        return b""


def _hien_thi_anh_tu_resource(data, width=620):
    """
    Hiển thị ảnh theo thứ tự ưu tiên:
    1) file local nếu còn tồn tại;
    2) URL nếu có;
    3) bytes base64 đã nhúng trong JSON/Supabase.

    Nhờ bước (3), app học sinh vẫn thấy ảnh khi chạy ở deployment/instance
    khác app giáo viên hoặc sau khi filesystem tạm của Streamlit bị reset.
    """
    if not isinstance(data, dict):
        return False

    for key in ("duong_dan", "duong_dan_anh"):
        path = str(data.get(key, "") or "").strip()
        if path and os.path.exists(path):
            st.image(path, width=width)
            return True

    for key in ("url_anh", "url", "public_url"):
        url = str(data.get(key, "") or "").strip()
        if url.startswith(("http://", "https://")):
            try:
                st.image(url, width=width)
                return True
            except Exception:
                pass

    blob = _anh_base64_bytes(
        data.get("du_lieu_base64")
        or data.get("image_base64")
        or data.get("base64")
        or ""
    )
    if blob:
        try:
            st.image(blob, width=width)
            return True
        except Exception:
            pass

    return False


def _seed_safe_stem(ten_file):
    stem = os.path.splitext(os.path.basename(str(ten_file or "hat_giong")))[0]
    stem = re.sub(r"[^0-9A-Za-zÀ-ỹ_-]+", "_", stem).strip("_")
    return (stem or "hat_giong")[:80]


def _seed_luu_file_nguon(raw_bytes, ten_file):
    """Lưu bản nguồn hạt giống để ảnh/bảng có nguồn ổn định sau khi Streamlit rerun."""
    raw_bytes = bytes(raw_bytes or b"")
    safe_stem = _seed_safe_stem(ten_file)
    ext = os.path.splitext(str(ten_file or ""))[1].lower()
    source_hash = hashlib.sha1(
        str(ten_file or "").encode("utf-8", errors="ignore") + raw_bytes
    ).hexdigest()[:10]
    source_name = f"{safe_stem}_{source_hash}{ext}"
    source_path = os.path.join(SEED_SOURCE_DIR, source_name)
    try:
        with open(source_path, "wb") as f:
            f.write(raw_bytes)
    except Exception:
        source_path = ""

    media_dir = os.path.join(SEED_MEDIA_DIR, f"{safe_stem}_{source_hash}")
    os.makedirs(media_dir, exist_ok=True)
    return source_path, media_dir, safe_stem


def _seed_compat_du_lieu_truc_quan(resources, ten_file=""):
    """Tạo trường tương thích renderer cũ từ tài nguyên đầu tiên, nhưng vẫn giữ toàn bộ list."""
    for res in resources or []:
        if res.get("loai") == "anh" and (
            str(res.get("duong_dan", "") or "").strip()
            or str(res.get("du_lieu_base64", "") or "").strip()
        ):
            return {
                "loai": "hinh_tu_tai_lieu",
                "duong_dan_anh": str(res.get("duong_dan", "") or "").strip(),
                "du_lieu_base64": str(res.get("du_lieu_base64", "") or "").strip(),
                "mime_type": str(res.get("mime_type", "") or "").strip(),
                "nguon": str(ten_file or res.get("nguon", "")),
                "mo_ta": "Hình trích nguyên từ file hạt giống",
            }
        if res.get("loai") == "bang" and (res.get("du_lieu") or []):
            rows = list(res.get("du_lieu") or [])
            if rows:
                return {
                    "loai": "bang_so_lieu",
                    "cot": list(rows[0]),
                    "du_lieu": [list(r) for r in rows[1:]],
                    "nguon": str(ten_file or res.get("nguon", "")),
                    "mo_ta": "Bảng trích nguyên từ file hạt giống",
                }
    return {}


def _seed_docx_text_va_tai_nguyen(raw_bytes, ten_file, media_dir, safe_stem):
    """
    Đọc DOCX theo thứ tự paragraph/bảng và GẮN marker tài nguyên vào đúng vùng câu.
    Ảnh được lưu thành file thật; bảng giữ nguyên hàng/cột.
    """
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(raw_bytes))
    blocks = []
    resources = {}
    resource_counter = 0
    image_counter = 0

    def add_resource(res):
        nonlocal resource_counter
        resource_counter += 1
        rid = f"R{resource_counter:04d}"
        resources[rid] = res
        blocks.append(f"[[SEED_RESOURCE:{rid}]]")

    def extract_images_from_para(para):
        nonlocal image_counter
        try:
            blips = para._p.xpath(".//a:blip")
        except Exception:
            blips = []

        for blip in blips:
            rel_id = blip.get(qn("r:embed"))
            if not rel_id:
                continue
            try:
                part = doc.part.related_parts[rel_id]
                blob = bytes(part.blob)
                # Không khử trùng theo hash toàn tài liệu: cùng một hình có thể được
                # tác giả chủ ý dùng lại ở nhiều câu khác nhau. Mỗi lần xuất hiện
                # trong luồng Word phải được gắn với đúng câu tương ứng.
                content_type = str(getattr(part, "content_type", "") or "").lower()
                ext = ".png"
                if "jpeg" in content_type or "jpg" in content_type:
                    ext = ".jpg"
                elif "gif" in content_type:
                    ext = ".gif"
                elif "webp" in content_type:
                    ext = ".webp"
                elif "bmp" in content_type:
                    ext = ".bmp"

                image_counter += 1
                img_name = f"{safe_stem}_img_{image_counter:03d}{ext}"
                img_path = os.path.join(media_dir, img_name)
                with open(img_path, "wb") as f:
                    f.write(blob)

                add_resource({
                    "loai": "anh",
                    "duong_dan": img_path,
                    # Lưu thêm bản base64 vào JSON/Supabase để app HS không phụ
                    # thuộc đường dẫn local của app GV.
                    "du_lieu_base64": base64.b64encode(blob).decode("ascii"),
                    "mime_type": content_type or (
                        "image/jpeg" if ext == ".jpg" else f"image/{ext.lstrip('.')}"
                    ),
                    "nguon": ten_file,
                    "mo_ta": "Ảnh/sơ đồ/biểu đồ trích nguyên từ file hạt giống",
                })
            except Exception:
                continue

    for child in doc.element.body.iterchildren():
        tag = str(child.tag)

        if tag.endswith("}p"):
            para = Paragraph(child, doc)
            txt = str(para.text or "").strip()
            if txt:
                blocks.append(txt)
            # Ảnh có thể nằm trong paragraph không có chữ.
            extract_images_from_para(para)

        elif tag.endswith("}tbl"):
            table = Table(child, doc)
            rows = []
            for row in table.rows:
                vals = []
                for cell in row.cells:
                    txt = " ".join(
                        " ".join(str(p.text or "").split())
                        for p in cell.paragraphs
                        if str(p.text or "").strip()
                    ).strip()
                    vals.append(txt)
                if any(str(x).strip() for x in vals):
                    rows.append(vals)
                    # Giữ text bảng trong luồng parser để A/B/C/D, metadata vẫn đọc được.
                    blocks.append("\t".join(vals))

            if rows:
                add_resource({
                    "loai": "bang",
                    "du_lieu": rows,
                    "nguon": ten_file,
                    "mo_ta": "Bảng Word trích nguyên từ file hạt giống",
                })

            # Nếu trong ô bảng có ảnh, vẫn trích riêng ảnh đó.
            seen_para_ids = set()
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        pid = id(para._p)
                        if pid in seen_para_ids:
                            continue
                        seen_para_ids.add(pid)
                        extract_images_from_para(para)

    return "\n".join(blocks), resources


def doc_goi_tu_file_hat_giong(file):
    """
    Trả về cả text + tài nguyên trực quan.
    - DOCX: giữ ảnh và bảng.
    - PDF/TXT: giữ luồng đọc text hiện có; file nguồn vẫn được lưu để không mất nguồn.
    """
    if file is None:
        return {
            "text": "",
            "resources": {},
            "source_path": "",
            "media_dir": "",
        }

    ten_file = str(getattr(file, "name", "") or "hat_giong")
    try:
        raw = bytes(file.getvalue())
    except Exception:
        file.seek(0)
        raw = bytes(file.read())

    source_path, media_dir, safe_stem = _seed_luu_file_nguon(raw, ten_file)
    ten_lower = ten_file.lower()

    if ten_lower.endswith(".docx"):
        try:
            txt, resources = _seed_docx_text_va_tai_nguyen(
                raw, ten_file, media_dir, safe_stem
            )
            return {
                "text": txt,
                "resources": resources,
                "source_path": source_path,
                "media_dir": media_dir,
            }
        except Exception:
            pass

    # Fallback cho PDF/TXT hoặc DOCX lỗi media: vẫn đọc text như trước.
    try:
        temp = io.BytesIO(raw)
        temp.name = ten_file
        txt = doc_text_tu_file_hat_giong(temp)
    except Exception:
        txt = ""

    return {
        "text": txt,
        "resources": {},
        "source_path": source_path,
        "media_dir": media_dir,
    }


def _seed_tai_nguyen_tu_chunk(chunk, resources_map):
    ids = re.findall(r"\[\[SEED_RESOURCE:(R\d{4})\]\]", str(chunk or ""))
    out = []
    for rid in ids:
        res = (resources_map or {}).get(rid)
        if isinstance(res, dict):
            out.append(dict(res))
    return out


def _seed_xoa_marker_tai_nguyen(chunk):
    return re.sub(
        r"(?m)^\s*\[\[SEED_RESOURCE:R\d{4}\]\]\s*$",
        "",
        str(chunk or ""),
    ).strip()


def _seed_hien_thi_tai_nguyen(seed):
    """Hiển thị ảnh/bảng hạt giống mà không phụ thuộc code Ngân hàng tốt nghiệp."""
    resources = list(seed.get("tai_nguyen_truc_quan", []) or [])
    if not resources:
        data = seed.get("du_lieu_truc_quan", {}) or {}
        if data:
            _hien_thi_anh_tu_resource(
                data,
                width=SEED_IMAGE_DISPLAY_WIDTH
            )
            rows = list(data.get("du_lieu", []) or [])
            cols = list(data.get("cot", []) or [])
            if rows or cols:
                try:
                    st.dataframe(
                        pd.DataFrame(rows, columns=cols or None),
                        use_container_width=True,
                        hide_index=True,
                    )
                except Exception:
                    st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)
        return

    for res in resources:
        loai = str(res.get("loai", "") or "").strip()
        if loai == "anh":
            _hien_thi_anh_tu_resource(
                res,
                width=SEED_IMAGE_DISPLAY_WIDTH
            )
        elif loai == "bang":
            rows = [list(r) for r in (res.get("du_lieu") or [])]
            if not rows:
                continue
            width = max(len(r) for r in rows)
            padded = [r + [""] * (width - len(r)) for r in rows]
            headers = padded[0]
            seen = {}
            cols = []
            for j, h in enumerate(headers):
                base = str(h or f"Cột {j+1}")
                seen[base] = seen.get(base, 0) + 1
                cols.append(base if seen[base] == 1 else f"{base} ({seen[base]})")
            df = (
                pd.DataFrame(padded[1:], columns=cols)
                if len(padded) > 1
                else pd.DataFrame(columns=cols)
            )
            st.dataframe(df, use_container_width=True, hide_index=True)


def _seed_duong_dan_media_dang_duoc_main_su_dung():
    """Không xóa media hạt giống nếu câu đã đồng bộ sang ngân hàng chính còn tham chiếu."""
    used = set()
    for q in doc_ngan_hang() or []:
        for res in q.get("tai_nguyen_truc_quan", []) or []:
            p = str(res.get("duong_dan", "") or "").strip()
            if p:
                used.add(os.path.abspath(p))
        data = q.get("du_lieu_truc_quan", {}) or {}
        p = str(data.get("duong_dan_anh", "") or "").strip()
        if p:
            used.add(os.path.abspath(p))
    return used


def _seed_don_media_khi_xoa_nguon(seed_items):
    """
    Dọn file nguồn/media của các seed bị xóa, nhưng KHÔNG xóa ảnh đang được
    ngân hàng chính sử dụng.
    """
    import shutil

    protected = _seed_duong_dan_media_dang_duoc_main_su_dung()
    source_paths = set()
    media_dirs = set()

    for q in seed_items or []:
        sp = str(q.get("nguon_file_path", "") or "").strip()
        md = str(q.get("thu_muc_media", "") or "").strip()
        if sp:
            source_paths.add(os.path.abspath(sp))
        if md:
            media_dirs.add(os.path.abspath(md))

    seed_source_root = os.path.abspath(SEED_SOURCE_DIR)
    seed_media_root = os.path.abspath(SEED_MEDIA_DIR)

    for sp in source_paths:
        try:
            if os.path.commonpath([sp, seed_source_root]) == seed_source_root and os.path.isfile(sp):
                os.remove(sp)
        except Exception:
            pass

    for md in media_dirs:
        try:
            if os.path.commonpath([md, seed_media_root]) != seed_media_root or not os.path.isdir(md):
                continue
            # Nếu có bất kỳ file nào trong thư mục đang được main dùng thì giữ cả thư mục.
            keep = any(p == md or p.startswith(md + os.sep) for p in protected)
            if not keep:
                shutil.rmtree(md, ignore_errors=True)
        except Exception:
            pass


def _seed_docx_blocks_in_order(doc):
    """Đọc paragraph + bảng Word theo gần đúng thứ tự xuất hiện trong tài liệu."""
    blocks = []
    try:
        from docx.text.paragraph import Paragraph
        from docx.table import Table

        for child in doc.element.body.iterchildren():
            tag = str(child.tag)
            if tag.endswith('}p'):
                text = Paragraph(child, doc).text
                if str(text).strip():
                    blocks.append(str(text).strip())
            elif tag.endswith('}tbl'):
                table = Table(child, doc)
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        txt = " ".join(
                            " ".join(str(p.text or "").split())
                            for p in cell.paragraphs
                            if str(p.text or "").strip()
                        ).strip()
                        if txt:
                            cells.append(txt)
                    if cells:
                        # Tab giúp giữ cấu trúc A/B/C/D nằm trên cùng một hàng.
                        blocks.append("\t".join(cells))
    except Exception:
        # Fallback an toàn nếu phiên bản python-docx khác.
        blocks.extend(
            str(p.text).strip()
            for p in getattr(doc, 'paragraphs', [])
            if str(getattr(p, 'text', '')).strip()
        )
        for table in getattr(doc, 'tables', []):
            for row in table.rows:
                vals = [" ".join(str(c.text or "").split()) for c in row.cells]
                if any(vals):
                    blocks.append("\t".join(vals))
    return blocks


def doc_text_tu_file_hat_giong(file):
    """
    Đọc file hạt giống theo hướng ưu tiên thu hồi tối đa nội dung.
    DOCX: đọc cả paragraph và bảng Word, thay vì chỉ doc.paragraphs.
    """
    if file is None:
        return ""
    ten = str(getattr(file, "name", "")).lower()
    try:
        if ten.endswith(".txt"):
            raw = file.getvalue()
            try:
                return raw.decode("utf-8")
            except Exception:
                return raw.decode("utf-8", errors="ignore")

        if ten.endswith(".docx"):
            from docx import Document
            file.seek(0)
            doc = Document(file)
            return "\n".join(_seed_docx_blocks_in_order(doc))

        if ten.endswith(".pdf"):
            from pypdf import PdfReader
            file.seek(0)
            reader = PdfReader(file)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return ""
    return ""


def tao_file_mau_hat_giong_docx_bytes():
    """Tạo file Word mẫu hạt giống ngay trong RAM để GV tải xuống."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("MẪU NHẬP NGÂN HÀNG HẠT GIỐNG")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph(
        "Có thể dùng Câu 1., Câu 1:, Câu 1), Câu 1 -, Câu hỏi 1 hoặc 1. "
        "Phương án A–D / nhận định a–d có thể nằm cùng dòng hoặc khác dòng. "
        "File nên được kiểm tra và hoàn thiện trước khi nhập. Dòng Đáp án là bắt buộc để đưa câu vào ngân hàng chính. "
        "YCCĐ, Mức độ, Thành phần năng lực, Chỉ báo và Kiến thức nên được điền sẵn; app không gọi AI để suy đoán khi nhập hạt giống."
    )

    doc.add_heading("Ví dụ 1 – Trắc nghiệm 4 lựa chọn", level=2)
    doc.add_paragraph("Câu 1. Thành phần nào sau đây trực tiếp mang thông tin di truyền ở sinh vật?")
    doc.add_paragraph("A. DNA")
    doc.add_paragraph("B. Lipid")
    doc.add_paragraph("C. Glucose")
    doc.add_paragraph("D. Nước")
    doc.add_paragraph("Đáp án: A")
    doc.add_paragraph("Mức độ: Nhận biết")
    doc.add_paragraph("Kiến thức: Vật chất di truyền")
    doc.add_paragraph("YCCĐ: Nêu được khái niệm gene và mô tả được cấu trúc DNA phù hợp nội dung câu hỏi")
    doc.add_paragraph("Thành phần năng lực: Nhận thức sinh học")
    doc.add_paragraph("Chỉ báo: NT1")

    doc.add_heading("Ví dụ 2 – A/B/C/D cùng một dòng", level=2)
    doc.add_paragraph(
        "Câu 2: Ở sinh vật nhân thực, quá trình dịch mã diễn ra chủ yếu ở đâu? "
        "A. Nhân tế bào    B. Ribosome    C. Lysosome    D. Trung thể"
    )
    doc.add_paragraph("Đáp án: B")

    doc.add_heading("Ví dụ 3 – Đúng/Sai", level=2)
    doc.add_paragraph("Câu 3) Một quần thể có tần số allele A = 0,6 và a = 0,4.")
    doc.add_paragraph("Dựa vào dữ kiện trên, hãy xác định các nhận định sau đúng hay sai.")
    doc.add_paragraph("a) Tần số allele A lớn hơn allele a.")
    doc.add_paragraph("b) Tổng tần số hai allele bằng 1.")
    doc.add_paragraph("c) Nếu quần thể cân bằng Hardy–Weinberg thì tần số AA là 0,36.")
    doc.add_paragraph("d) Nếu không có tác động tiến hóa thì tần số allele luôn bằng 0,5.")
    doc.add_paragraph("Đáp án: ĐĐĐS")

    doc.add_heading("Ví dụ 4 – Trả lời ngắn", level=2)
    doc.add_paragraph("Câu hỏi 4. Một tế bào nguyên phân 3 lần liên tiếp tạo ra bao nhiêu tế bào con?")
    doc.add_paragraph("Đáp án: 8")

    doc.add_heading("Ví dụ 5 – Câu trong bảng Word", level=2)
    table = doc.add_table(rows=5, cols=2)
    table.cell(0, 0).text = "5."
    table.cell(0, 1).text = "Trong giảm phân, sự tiếp hợp NST tương đồng xảy ra ở kì nào?"
    table.cell(1, 0).text = "A."
    table.cell(1, 1).text = "Kì đầu I"
    table.cell(2, 0).text = "B."
    table.cell(2, 1).text = "Kì giữa I"
    table.cell(3, 0).text = "C."
    table.cell(3, 1).text = "Kì sau I"
    table.cell(4, 0).text = "D."
    table.cell(4, 1).text = "Kì cuối I"
    doc.add_paragraph("Đáp án: A")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()



# ==========================================================
# CHUẨN PHÂN LOẠI NĂNG LỰC SINH HỌC
# Dựa trên biểu hiện của 3 thành phần năng lực trong CT môn Sinh học.
# Nhãn có sẵn trong hạt giống chỉ là THAM KHẢO.
# ==========================================================
CHI_BAO_NANG_LUC_CHUAN = {
    "Nhận thức sinh học": [
        ("NT1", "Nhận biết, kể tên, phát biểu, nêu được đối tượng, khái niệm, quy luật, quá trình sống"),
        ("NT2", "Trình bày đặc điểm, vai trò của đối tượng và quá trình sống"),
        ("NT3", "Phân loại đối tượng, hiện tượng sống theo các tiêu chí"),
        ("NT4", "Phân tích đặc điểm của đối tượng, sự vật, quá trình theo một logic nhất định"),
        ("NT5", "So sánh, lựa chọn đối tượng, khái niệm, cơ chế, quá trình theo tiêu chí"),
        ("NT6", "Giải thích mối quan hệ giữa các sự vật và hiện tượng: nguyên nhân - kết quả, cấu tạo - chức năng"),
        ("NT7", "Nhận ra và chỉnh sửa điểm sai, đưa ra nhận định phản biện"),
        ("NT8", "Tìm từ khóa, sử dụng thuật ngữ khoa học, kết nối thông tin, lập sơ đồ/hình biểu đạt")
    ],
    "Tìm hiểu thế giới sống": [
        ("TH1", "Đề xuất vấn đề liên quan đến thế giới sống, đặt câu hỏi nghiên cứu"),
        ("TH2", "Đưa ra phán đoán, xây dựng giả thuyết nghiên cứu"),
        ("TH3", "Lập kế hoạch nghiên cứu, lựa chọn phương pháp quan sát, thực nghiệm, điều tra, hồi cứu tư liệu"),
        ("TH4", "Thực hiện kế hoạch, thu thập và lưu giữ dữ liệu"),
        ("TH5", "Xử lí, phân tích dữ liệu và trình bày kết quả"),
        ("TH6", "So sánh kết quả với giả thuyết, giải thích, rút kết luận và điều chỉnh"),
        ("TH7", "Viết, trình bày báo cáo, thảo luận và bảo vệ kết quả nghiên cứu")
    ],
    "Vận dụng kiến thức, kĩ năng đã học": [
        ("VD1", "Giải thích, đánh giá hiện tượng thường gặp trong tự nhiên và đời sống"),
        ("VD2", "Giải quyết vấn đề thực tiễn liên quan sức khỏe, vệ sinh, môi trường, sản xuất"),
        ("VD3", "Đề xuất hoặc thực hiện giải pháp bảo vệ sức khỏe, thiên nhiên, môi trường, thích ứng biến đổi khí hậu")
    ]
}


def xac_dinh_nang_luc_chi_bao_tu_noi_dung(noi_dung):
    """
    Phân loại cục bộ, không gọi API.
    Trọng tâm là HÀNH VI học sinh phải thực hiện, không chỉ dựa vào chủ đề kiến thức.
    """
    raw = " ".join(str(noi_dung or "").split())
    s = raw.casefold()

    # Điểm theo từng chỉ báo. Từ khóa có tính quyết định được cho trọng số cao hơn.
    rules = {
        "TH1": ["đề xuất vấn đề", "đặt câu hỏi nghiên cứu", "câu hỏi nghiên cứu"],
        "TH2": ["giả thuyết", "phán đoán", "dự đoán nghiên cứu"],
        "TH3": ["lập kế hoạch", "thiết kế thí nghiệm", "bố trí thí nghiệm", "chọn phương pháp", "phương pháp nghiên cứu"],
        "TH4": ["thu thập dữ liệu", "tiến hành thí nghiệm", "thực hiện thí nghiệm", "điều tra", "khảo sát"],
        "TH5": ["bảng số liệu", "biểu đồ", "đồ thị", "xử lí số liệu", "xử lý số liệu", "phân tích số liệu", "kết quả thí nghiệm"],
        "TH6": ["rút ra kết luận", "kết luận nào", "so sánh với giả thuyết", "giải thích kết quả thí nghiệm", "đánh giá kết luận"],
        "TH7": ["báo cáo kết quả", "trình bày kết quả nghiên cứu", "thảo luận kết quả"],

        "VD1": ["trong thực tiễn", "trong đời sống", "hiện tượng thực tế", "giải thích hiện tượng", "đánh giá hiện tượng"],
        "VD2": ["giải quyết", "xử lí tình huống", "xử lý tình huống", "sức khỏe", "vệ sinh", "sản xuất", "ô nhiễm"],
        "VD3": ["đề xuất giải pháp", "biện pháp", "bảo vệ môi trường", "bảo vệ thiên nhiên", "biến đổi khí hậu", "phòng tránh"],

        "NT1": ["nêu", "kể tên", "nhận biết", "phát biểu", "khái niệm", "là gì"],
        "NT2": ["trình bày", "đặc điểm", "vai trò", "chức năng"],
        "NT3": ["phân loại", "xếp vào nhóm", "thuộc nhóm"],
        "NT4": ["phân tích đặc điểm", "phân tích quá trình", "phân tích cơ chế"],
        "NT5": ["so sánh", "phân biệt", "lựa chọn", "đối chiếu"],
        "NT6": ["giải thích mối quan hệ", "nguyên nhân", "kết quả", "cấu tạo", "chức năng", "cơ chế"],
        "NT7": ["nhận định sai", "chỉnh sửa", "phản biện", "phát biểu nào sai"],
        "NT8": ["từ khóa", "thuật ngữ", "sơ đồ", "sắp xếp", "kết nối thông tin"]
    }

    scores = {}
    for code, keywords in rules.items():
        score = 0
        for kw in keywords:
            if kw in s:
                score += 2 if len(kw.split()) >= 2 else 1
        scores[code] = score

    # Đọc bảng/biểu đồ không tự động là Tìm hiểu TGS nếu chỉ hỏi nhận biết đơn giản,
    # nhưng khi có động từ phân tích/xử lí/kết luận thì tăng mạnh.
    if any(x in s for x in ["bảng số liệu", "biểu đồ", "đồ thị"]):
        if any(x in s for x in ["phân tích", "xử lí", "xử lý", "rút", "kết luận", "nhận xét"]):
            scores["TH5"] += 4

    # Tình huống thực tiễn + yêu cầu giải pháp => Vận dụng rõ.
    if any(x in s for x in ["thực tiễn", "đời sống", "sức khỏe", "môi trường", "sản xuất"]):
        if any(x in s for x in ["đề xuất", "giải pháp", "biện pháp", "xử lí", "xử lý", "đánh giá"]):
            scores["VD2"] += 4

    best_code = max(scores, key=scores.get)
    best_score = scores.get(best_code, 0)
    sorted_scores = sorted(scores.values(), reverse=True)
    second = sorted_scores[1] if len(sorted_scores) > 1 else 0

    if best_score <= 0:
        return {
            "thanh_phan_nang_luc": "Chưa xác định",
            "chi_bao": "",
            "mo_ta_chi_bao": "",
            "do_tin_cay": 0.0,
            "trang_thai": "Cần GV kiểm tra"
        }

    if best_code.startswith("TH"):
        nl = "Tìm hiểu thế giới sống"
    elif best_code.startswith("VD"):
        nl = "Vận dụng kiến thức, kĩ năng đã học"
    else:
        nl = "Nhận thức sinh học"

    mo_ta = ""
    for code, desc in CHI_BAO_NANG_LUC_CHUAN[nl]:
        if code == best_code:
            mo_ta = desc
            break

    # Confidence tương đối: cao khi điểm nổi trội so với lựa chọn thứ hai.
    confidence = min(
        0.98,
        0.50 + 0.08 * best_score + 0.05 * max(0, best_score - second)
    )

    return {
        "thanh_phan_nang_luc": nl,
        "chi_bao": best_code,
        "mo_ta_chi_bao": mo_ta,
        "do_tin_cay": round(confidence, 2),
        "trang_thai": (
            "Tự xác định"
            if confidence >= 0.72
            else "Cần GV kiểm tra"
        )
    }


def mo_ta_chi_bao_theo_ma(ma_chi_bao):
    ma = str(ma_chi_bao or "").strip().upper()
    for nl, ds in CHI_BAO_NANG_LUC_CHUAN.items():
        for code, desc in ds:
            if code == ma:
                return desc
    return ""


def nang_luc_theo_ma_chi_bao(ma_chi_bao):
    ma = str(ma_chi_bao or "").strip().upper()
    if ma.startswith("NT"):
        return "Nhận thức sinh học"
    if ma.startswith("TH"):
        return "Tìm hiểu thế giới sống"
    if ma.startswith("VD"):
        return "Vận dụng kiến thức, kĩ năng đã học"
    return ""


def noi_dung_de_phan_loai_cau(q):
    """Ghép phần thể hiện nhiệm vụ HS phải làm để suy chỉ báo."""
    parts = [
        q.get("tinh_huong", ""),
        q.get("cau_hoi", ""),
        q.get("yccd", ""),
        q.get("hanh_vi_nang_luc", "")
    ]
    return " ".join(str(x or "") for x in parts if str(x or "").strip())


def gan_chi_bao_chuan_cho_cau(q, force=False):
    """
    Gắn mã NT/TH/VD cho câu hỏi bằng bộ quy tắc cục bộ, không gọi API.
    - Nếu đã có mã chuẩn thì giữ.
    - Nếu thiếu/sai dạng thì suy từ hành vi của câu.
    - Đ/S: gắn riêng cho từng ý nếu có thể.
    """
    q2 = dict(q)
    ma_cu = str(q2.get("chi_bao", "")).strip().upper()
    ma_hop_le = bool(re.fullmatch(r"(NT[1-8]|TH[1-7]|VD[1-3])", ma_cu))

    if force or not ma_hop_le:
        kq = xac_dinh_nang_luc_chi_bao_tu_noi_dung(
            noi_dung_de_phan_loai_cau(q2)
        )
        if kq.get("chi_bao"):
            q2["chi_bao"] = kq["chi_bao"]
            q2["mo_ta_chi_bao"] = kq["mo_ta_chi_bao"]
            q2["thanh_phan_nang_luc"] = kq["thanh_phan_nang_luc"]
            q2["do_tin_cay_phan_loai"] = kq["do_tin_cay"]
            q2["trang_thai_phan_loai"] = kq["trang_thai"]

    if q2.get("dang_cau") == "Đúng / Sai":
        meta_moi = []
        for nd in list(q2.get("nhan_dinh_meta", []) or []):
            nd2 = dict(nd)
            ma_nd = str(nd2.get("chi_bao", "")).strip().upper()
            if force or not re.fullmatch(r"(NT[1-8]|TH[1-7]|VD[1-3])", ma_nd):
                noi_dung_nd = " ".join([
                    str(q2.get("tinh_huong", "") or ""),
                    str(q2.get("cau_hoi", "") or ""),
                    str(nd2.get("noi_dung", "") or ""),
                    str(nd2.get("yccd", "") or "")
                ])
                kq_nd = xac_dinh_nang_luc_chi_bao_tu_noi_dung(noi_dung_nd)
                if kq_nd.get("chi_bao"):
                    nd2["chi_bao"] = kq_nd["chi_bao"]
                    nd2["mo_ta_chi_bao"] = kq_nd["mo_ta_chi_bao"]
                    nd2["thanh_phan_nang_luc"] = kq_nd["thanh_phan_nang_luc"]
                    nd2["do_tin_cay_phan_loai"] = kq_nd["do_tin_cay"]
                    nd2["trang_thai_phan_loai"] = kq_nd["trang_thai"]
            meta_moi.append(nd2)
        q2["nhan_dinh_meta"] = meta_moi

    return q2


def gan_chi_bao_cho_ngan_hang_hien_co(bank, force=False):
    """Chuẩn hóa metadata khi đọc ngân hàng; không tốn API."""
    return [
        gan_chi_bao_chuan_cho_cau(q, force=force)
        for q in (bank or [])
        if isinstance(q, dict)
    ]


def ra_soat_nang_luc_hat_giong(seed):
    seed_moi = dict(seed)
    nhan_goc = str(seed_moi.get("thanh_phan_nang_luc", "")).strip()

    kq = xac_dinh_nang_luc_chi_bao_tu_noi_dung(
        seed_moi.get("noi_dung_goc", "")
    )

    seed_moi["nang_luc_goc"] = nhan_goc
    seed_moi["thanh_phan_nang_luc"] = kq["thanh_phan_nang_luc"]
    seed_moi["chi_bao"] = kq["chi_bao"]
    seed_moi["mo_ta_chi_bao"] = kq["mo_ta_chi_bao"]
    seed_moi["do_tin_cay_phan_loai"] = kq["do_tin_cay"]
    seed_moi["trang_thai_phan_loai"] = kq["trang_thai"]

    if (
        nhan_goc
        and kq["thanh_phan_nang_luc"] != "Chưa xác định"
        and chuan_hoa_ten_nang_luc_seed(nhan_goc)
        != kq["thanh_phan_nang_luc"]
    ):
        seed_moi["canh_bao_nhan_goc"] = (
            f"Nhãn gốc '{nhan_goc}' khác với phân loại lại "
            f"'{kq['thanh_phan_nang_luc']}'."
        )
    else:
        seed_moi["canh_bao_nhan_goc"] = ""

    return seed_moi


def dem_hat_giong_phu_hop_yccd_da_chon():
    bank_seed = doc_ngan_hang_hat_giong()
    if not bank_seed:
        return 0

    ids = set()
    for item in st.session_state.get("yccd_da_chon", []):
        _, cfg = lay_cau_hinh(item)
        seed = chon_hat_giong_phu_hop(
            bank_seed,
            cfg.get("Thành phần năng lực", ""),
            cfg.get("Mức độ", ""),
            cfg.get("Dạng câu hỏi", ""),
            item.get("YCCĐ", "")
        )
        if seed:
            ids.add(seed.get("id") or str(seed.get("noi_dung_goc", ""))[:80])

    return len(ids)


def _seed_chuan_hoa_text_de_tach(text_goc):
    """Chuẩn hóa nhẹ để nhận A-D/a-d nhưng không phá nội dung phương án.

    Quan trọng: A-D và a-d là hai hệ nhãn khác nhau. Không dùng IGNORECASE.
    Chỉ tách nhãn nằm sau tab/nhiều khoảng trắng (kiểu Word) hoặc sau dấu chấm
    của phương án trước; không tách các chữ a/b/c/d đang nằm trong nội dung.
    """
    s = unicodedata.normalize("NFKC", str(text_goc or ""))
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = s.replace("\u00a0", " ")

    # Tách phương án/nhận định cùng một dòng TRƯỚC khi gom khoảng trắng.
    s = re.sub(r"(?:\t+| {2,})([A-D])\s*[\.\)\:]\s*", r"\n\1. ", s)
    s = re.sub(r"(?:\t+| {2,})([a-d])\s*[\.\)]\s+", r"\n\1) ", s)
    # Trường hợp chỉ cách một khoảng trắng nhưng phương án trước kết thúc bằng dấu chấm.
    s = re.sub(r"(?<=\.)\s+([A-D])\s*[\.\)\:]\s*", r"\n\1. ", s)

    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

def _seed_tim_moc_cau(s):
    """
    Ưu tiên mốc mạnh có chữ Câu/Câu hỏi.
    Nếu không có, mới dùng dạng số thứ tự 1., 2), 3: ở đầu dòng.
    """
    patterns_strong = [
        r"(?im)^\s*Câu\s+hỏi\s+(\d+)\s*(?:[\.\:\)\-–—]|$)\s*",
        r"(?im)^\s*Câu\s+(\d+)\s*(?:[\.\:\)\-–—]|$)\s*",
    ]
    matches = []
    for pat in patterns_strong:
        matches.extend(list(re.finditer(pat, s)))

    # Khử trùng vị trí nếu 2 pattern cùng chạm một mốc.
    by_start = {}
    for m in matches:
        by_start.setdefault(m.start(), m)
    matches = sorted(by_start.values(), key=lambda x: x.start())
    if matches:
        return matches

    # Dạng tài liệu chỉ đánh số: 1. / 2) / 3:
    bare = list(re.finditer(r"(?im)^\s*(\d{1,3})\s*[\.\:\)]\s+", s))
    if len(bare) < 2:
        return []

    # Chỉ chấp nhận nếu dãy số có tính chất câu hỏi tương đối liên tục,
    # tránh nhầm danh sách 1., 2., 3. nằm trong một câu.
    nums = [int(m.group(1)) for m in bare]
    continuity = sum(1 for a, b in zip(nums, nums[1:]) if b == a + 1)
    if continuity < max(1, len(nums) // 3):
        return []
    return bare


def _seed_doan_co_4_phuong_an(chunk):
    labels = re.findall(r"(?im)^\s*([A-D])\s*[\.\)\:]\s+", chunk)
    return len(set(x.upper() for x in labels)) >= 3


def _seed_doan_co_4_nhan_dinh(chunk):
    labels = re.findall(r"(?m)^\s*([a-d])\s*[\.\)\:]\s+", chunk)
    return len(set(x.lower() for x in labels)) >= 3


def _seed_chuan_hoa_muc_do(value):
    raw = " ".join(str(value or "").strip().split())
    cf = raw.casefold()
    if "vận dụng" in cf:
        return "Vận dụng"
    if "hiểu" in cf or "thông hiểu" in cf:
        return "Thông hiểu"
    if "biết" in cf or "nhận biết" in cf:
        return "Nhận biết"
    return raw


def _seed_lay_meta(chunk, nhan):
    """Đọc metadata hạt giống theo cả mẫu ``[[NHÃN]]`` và mẫu Word thông thường.

    Hỗ trợ: ``[[ĐÁP ÁN]] B``, ``Đáp án: B``, nhãn ở một dòng rồi giá trị
    ở dòng kế tiếp, và nhãn/giá trị trong bảng Word.
    """
    text = unicodedata.normalize("NFKC", str(chunk or ""))
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    label = re.escape(str(nhan or "").strip())

    mm = re.search(
        rf"(?is)\[\[\s*{label}\s*\]\]\s*[:：]?\s*(.*?)"
        rf"(?=\s*\[\[|\n\s*(?:Câu(?:\s+hỏi)?\s+\d+|\d{{1,3}}\s*[\.\:\)])|$)",
        text,
    )
    if mm:
        value = " ".join(str(mm.group(1)).strip().split())
        if value:
            return value

    lines = text.splitlines()
    known_labels = re.compile(
        r"^(?:ĐÁP\s*ÁN|DAP\s*AN|YCCĐ|YÊU\s*CẦU\s*CẦN\s*ĐẠT|"
        r"ĐƠN\s*VỊ\s*KIẾN\s*THỨC|KIẾN\s*THỨC|MỨC\s*ĐỘ(?:\s*TƯ\s*DUY)?|"
        r"CẤP\s*ĐỘ\s*TƯ\s*DUY|THÀNH\s*PHẦN\s*NĂNG\s*LỰC|CHỈ\s*BÁO|"
        r"HƯỚNG\s*DẪN\s*GIẢI|HUONG\s*DAN\s*GIAI)\b",
        flags=re.I,
    )
    target_re = re.compile(rf"^\s*{label}\s*(?:[:：]\s*)?(.*)$", flags=re.I)

    for i, line in enumerate(lines):
        raw_line = str(line or "").strip()
        if not raw_line:
            continue
        m = target_re.match(raw_line)
        if not m:
            continue
        tail = re.sub(r"^[:：]\s*", "", str(m.group(1) or "").strip()).strip()
        if tail:
            return " ".join(tail.split())
        for j in range(i + 1, min(len(lines), i + 4)):
            nxt = str(lines[j] or "").strip()
            if not nxt:
                continue
            if known_labels.match(nxt):
                break
            if re.match(r"^(?:Câu(?:\s+hỏi)?\s+\d+|\d{1,3}\s*[\.\:\)]|[A-Da-d]\s*[\.\)\:])", nxt, flags=re.I):
                break
            return " ".join(nxt.split())
        break
    return ""



def _seed_lay_huong_dan_giai(chunk):
    """Đọc lời giải/hướng dẫn giải từ file hạt giống.

    Hỗ trợ cả dạng Word thông thường ``Hướng dẫn giải: ...`` / ``Lời giải: ...``
    và mẫu ``[[HƯỚNG DẪN GIẢI]] ...``. Chỉ lấy nội dung lời giải của chính
    câu hiện tại; dừng trước tiêu đề phần/chủ đề nếu có.
    """
    text = unicodedata.normalize("NFKC", str(chunk or ""))
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.splitlines()

    start_idx = None
    first_value = ""
    label_re = re.compile(
        r"^\s*(?:\[\[\s*)?(?:HƯỚNG\s*DẪN\s*GIẢI|HUONG\s*DAN\s*GIAI|"
        r"HƯỚNG\s*GIẢI|HUONG\s*GIAI|LỜI\s*GIẢI|LOI\s*GIAI)"
        r"(?:\s*\]\])?\s*[:：]?\s*(.*)$",
        flags=re.I,
    )

    for i, line in enumerate(lines):
        m = label_re.match(str(line or "").strip())
        if m:
            start_idx = i
            first_value = str(m.group(1) or "").strip()
            break

    if start_idx is None:
        return ""

    parts = []
    if first_value:
        parts.append(first_value)

    stop_re = re.compile(
        r"^\s*(?:DẠNG\s*\d+|DANG\s*\d+|CHỦ\s*ĐỀ\b|CHU\s*DE\b|"
        r"PHẦN\s+[IVXLC]+\b|PHAN\s+[IVXLC]+\b|NGÂN\s*HÀNG\s*CÂU\s*HỎI\b)",
        flags=re.I,
    )
    metadata_re = re.compile(
        r"^\s*(?:ĐÁP\s*ÁN|DAP\s*AN|YCCĐ|YÊU\s*CẦU\s*CẦN\s*ĐẠT|"
        r"KIẾN\s*THỨC|MỨC\s*ĐỘ|THÀNH\s*PHẦN\s*NĂNG\s*LỰC|CHỈ\s*BÁO)\b",
        flags=re.I,
    )

    for line in lines[start_idx + 1:]:
        stp = str(line or "").strip()
        if not stp:
            if parts and parts[-1] != "":
                parts.append("")
            continue
        if stop_re.match(stp):
            break
        # Nếu file đặt lời giải trước metadata, không nuốt metadata vào lời giải.
        if metadata_re.match(stp):
            break
        if re.match(r"^\[\[\s*(?:ĐÁP\s*ÁN|DAP\s*AN|YCCĐ|META\s+[A-Da-d])", stp, flags=re.I):
            break
        parts.append(stp)

    # Chuẩn hoá khoảng trắng nhưng giữ ngắt dòng khi lời giải có nhiều bước.
    out = []
    for x in parts:
        if x == "":
            if out and out[-1] != "":
                out.append("")
            continue
        out.append(" ".join(x.split()))
    return "\n".join(out).strip()


def _seed_tach_nhan_dinh_ds(chunk):
    """Đọc 4 nhận định a-d và metadata từng ý.

    Hỗ trợ cả ``[[META A]]`` và bảng Word kiểu:
    ``a | Đ | NT1 | Biết | Nhận thức sinh học``.
    Tuyệt đối không coi A/B/C/D của câu nhiều lựa chọn là a/b/c/d.
    """
    noi_dung = {}
    for mm in re.finditer(r"(?m)^\s*([a-d])\s*[\.\)\:]\s*(.+?)\s*$", str(chunk or "")):
        ky = mm.group(1)
        if ky not in noi_dung:
            noi_dung[ky] = " ".join(mm.group(2).strip().split())

    meta_by_key = {}

    # Dạng [[META A]] ...
    meta_matches = list(re.finditer(
        r"(?is)\[\[\s*META\s+([A-Da-d])\s*\]\]\s*(.*?)(?=\s*\[\[\s*META\s+[A-Da-d]\s*\]\]|[\n\r]|$)",
        str(chunk or "")
    ))
    for mm in meta_matches:
        ky = mm.group(1).lower()
        body = mm.group(2).strip()

        def field(name):
            m2 = re.search(rf"(?i)(?:^|\|)\s*{re.escape(name)}\s*=\s*([^|]+)", body)
            return " ".join(m2.group(1).strip().split()) if m2 else ""

        dap = field("Đáp án")
        dap_cf = dap.casefold()
        if dap_cf in {"đ", "đúng", "dung", "true"}:
            dap = "Đúng"
        elif dap_cf in {"s", "sai", "false"}:
            dap = "Sai"
        meta_by_key[ky] = {
            "ky_hieu": ky,
            "noi_dung": noi_dung.get(ky, ""),
            "dap_an": dap,
            "chi_bao": re.sub(r"\s+", "", field("Chỉ báo").upper()),
            "muc_do": _seed_chuan_hoa_muc_do(field("Mức độ")),
            "thanh_phan_nang_luc": field("Thành phần năng lực"),
            "yccd": "", "giai_thich": ""
        }

    # Dạng bảng Word. Sau khi trích DOCX, mỗi hàng trở thành một dòng text.
    # Ví dụ: a Đ NT1 Biết Nhận thức sinh học
    for line in str(chunk or "").splitlines():
        stp = " ".join(str(line or "").strip().split())
        mrow = re.match(
            r"^([a-d])\s*[\.\)]?\s+(Đ|S|Đúng|Sai)\s+((?:NT|TH|VD)\s*\d+)\s+"
            r"(Nhận\s*biết|Biết|Thông\s*hiểu|Hiểu|Vận\s*dụng)\s+(.+)$",
            stp, flags=re.I
        )
        if not mrow:
            continue
        ky = mrow.group(1).lower()
        if ky in meta_by_key:
            continue
        dap_raw = mrow.group(2).casefold()
        dap = "Đúng" if dap_raw in {"đ", "đúng", "dung"} else "Sai"
        meta_by_key[ky] = {
            "ky_hieu": ky,
            "noi_dung": noi_dung.get(ky, ""),
            "dap_an": dap,
            "chi_bao": re.sub(r"\s+", "", mrow.group(3).upper()),
            "muc_do": _seed_chuan_hoa_muc_do(mrow.group(4)),
            "thanh_phan_nang_luc": " ".join(mrow.group(5).split()),
            "yccd": "", "giai_thich": ""
        }

    # Nếu bảng không có đủ metadata thì vẫn tạo 4 ý từ nội dung để giữ cấu trúc câu.
    for ky in "abcd":
        if ky in noi_dung and ky not in meta_by_key:
            meta_by_key[ky] = {
                "ky_hieu": ky, "noi_dung": noi_dung.get(ky, ""), "dap_an": "",
                "chi_bao": "", "muc_do": "", "thanh_phan_nang_luc": "",
                "yccd": "", "giai_thich": ""
            }

    return [meta_by_key[k] for k in "abcd" if k in meta_by_key]

def _seed_tom_tat_meta_ds(meta):
    """Tạo metadata cấp câu để các bộ lọc cũ vẫn hoạt động, nhưng giữ chi tiết ở nhan_dinh_meta."""
    if not meta:
        return "", "", ""

    mucs = [x.get("muc_do", "") for x in meta if str(x.get("muc_do", "")).strip()]
    nls = [x.get("thanh_phan_nang_luc", "") for x in meta if str(x.get("thanh_phan_nang_luc", "")).strip()]
    cbs = [x.get("chi_bao", "") for x in meta if str(x.get("chi_bao", "")).strip()]

    # Với mức độ, dùng mức cao nhất để câu không bị đánh giá thấp hơn yêu cầu của các ý.
    rank = {"Nhận biết": 1, "Thông hiểu": 2, "Vận dụng": 3}
    muc = max(mucs, key=lambda x: rank.get(x, 0)) if mucs else ""

    # Nếu mọi ý cùng năng lực/chỉ báo thì dùng trực tiếp; nếu khác nhau, để chuỗi tổng hợp dễ nhìn.
    uniq_nl = list(dict.fromkeys(nls))
    uniq_cb = list(dict.fromkeys(cbs))
    nl = uniq_nl[0] if len(uniq_nl) == 1 else " / ".join(uniq_nl)
    cb = uniq_cb[0] if len(uniq_cb) == 1 else ", ".join(uniq_cb)
    return muc, nl, cb



def _seed_doc_dap_an_nguon(chunk, dang, nhan_dinh_meta=None):
    """Đọc đáp án GV ghi trong file hạt giống; không tự suy đáp án."""
    raw = _seed_lay_meta(chunk, "ĐÁP ÁN") or _seed_lay_meta(chunk, "DAP AN")
    raw = " ".join(str(raw or "").strip().split())

    if dang == "Đúng / Sai":
        # Ưu tiên đáp án từng ý trong [[META A]]... nếu đủ 4 ý.
        meta = list(nhan_dinh_meta or [])
        if len(meta) == 4 and all(x.get("dap_an") in {"Đúng", "Sai"} for x in meta):
            return [x.get("dap_an") for x in meta]

        compact = re.sub(r"[^DdSsĐđ]", "", raw)
        if len(compact) >= 4:
            compact = compact[:4].upper().replace("Đ", "D")
            return ["Đúng" if ch == "D" else "Sai" for ch in compact]

        vals = re.findall(r"(?i)\b(Đúng|Sai)\b", raw)
        if len(vals) >= 4:
            return ["Đúng" if v.casefold().startswith("đ") else "Sai" for v in vals[:4]]
        return []

    if dang == "Trắc nghiệm 4 lựa chọn":
        m = re.search(r"(?i)\b([ABCD])\b", raw)
        return m.group(1).upper() if m else ""

    return raw


def _seed_co_dap_an_nguon(q):
    if q.get("dang_cau_goi_y") == "Đúng / Sai":
        ans = q.get("dap_an_nguon", []) or []
        return len(ans) == 4 and all(x in {"Đúng", "Sai"} for x in ans)
    return bool(str(q.get("dap_an_nguon", "") or "").strip())


SEED_REVIEW_BATCH_SCHEMA = {
    "type": "object",
    "properties": {
        "items": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "id": {"type": "string"},
                    "ket_luan": {"type": "string"},
                    "dap_an_de_xuat": {"type": "string"},
                    "dap_an_4_y": {"type": "array", "items": {"type": "string"}},
                    "do_tin_cay": {"type": "number"},
                    "ly_do": {"type": "string"}
                },
                "required": ["id", "ket_luan", "dap_an_de_xuat", "dap_an_4_y", "do_tin_cay", "ly_do"]
            }
        }
    },
    "required": ["items"]
}


def _seed_text_review(q):
    dang = q.get("dang_cau_goi_y", "")
    ans = q.get("dap_an_nguon", "")
    if isinstance(ans, list):
        ans = "; ".join(f"{chr(97+i)}) {v}" for i, v in enumerate(ans))

    # Không đưa đáp án/metadata lẫn trong thân câu cho AI; đáp án nguồn được truyền riêng bên dưới.
    noi_dung_sach = (
        str(q.get("noi_dung_hien_thi", "")).strip()
        or _seed_xoa_metadata_de_hien_thi(q.get("noi_dung_goc", ""))
    )
    return (
        f"ID: {q.get('id','')}\n"
        f"Dạng: {dang}\n"
        f"Câu hỏi và dữ kiện:\n{noi_dung_sach[:7000]}\n"
        f"ĐÁP ÁN NGUỒN CỦA GV: {ans}\n"
    )


def ra_soat_hat_giong_bang_ai(ds_cau, batch_size=8):
    """
    Kiểm tra đáp án hạt giống theo lô để giảm số lần gọi API.
    KHÔNG tự sửa đáp án nguồn. Chỉ gắn trạng thái và đề xuất để GV duyệt.
    """
    ds = [dict(q) for q in (ds_cau or [])]

    # Câu thiếu đáp án nguồn: vẫn lưu nhưng khóa sử dụng.
    can_review = []
    for q in ds:
        if not _seed_co_dap_an_nguon(q):
            q["kiem_tra_dap_an"] = "Chưa đủ dữ kiện"
            q["trang_thai_kiem_tra_dap_an"] = "Chưa đủ dữ kiện"
            q["do_tin_cay_dap_an"] = 0.0
            q["canh_bao_dap_an"] = "Không đọc được đáp án nguồn đầy đủ để đối chiếu."
            q["duoc_dung_lam_hat_giong"] = False
            q["gv_da_duyet_dap_an"] = False
        else:
            can_review.append(q)

    by_id = {q.get("id"): q for q in ds}

    for start in range(0, len(can_review), max(1, int(batch_size))):
        batch = can_review[start:start + max(1, int(batch_size))]
        blocks = "\n\n====================\n\n".join(_seed_text_review(q) for q in batch)
        prompt = f"""
Bạn là chuyên gia thẩm định câu hỏi Sinh học THPT Việt Nam.
Hãy KIỂM TRA ĐÁP ÁN của các câu hạt giống dưới đây. Tuyệt đối không viết lại câu hỏi và không tự sửa dữ liệu của giáo viên.

QUY TẮC:
1. Tự giải độc lập bằng kiến thức Sinh học THPT và các dữ kiện có trong câu.
2. So sánh với ĐÁP ÁN NGUỒN CỦA GV.
3. ket_luan chỉ dùng đúng một trong ba giá trị: "Khớp", "Không khớp", "Chưa đủ dữ kiện".
4. do_tin_cay từ 0 đến 1.
5. Trắc nghiệm 4 lựa chọn: dap_an_de_xuat chỉ A/B/C/D.
6. Đúng/Sai: dap_an_4_y phải có đúng 4 phần tử "Đúng"/"Sai" nếu đủ dữ kiện; nếu không đủ thì để mảng rỗng.
7. Trả lời ngắn: dap_an_de_xuat là kết quả ngắn.
8. Nếu câu phụ thuộc hình/bảng mà phần văn bản không đủ để kết luận, bắt buộc chọn "Chưa đủ dữ kiện".
9. Không vì đáp án nguồn đã có mà mặc định cho là đúng.
10. ly_do ngắn gọn, nêu đúng điểm mâu thuẫn hoặc lý do chưa thể xác minh.

CÁC CÂU CẦN KIỂM TRA:
{blocks}
"""
        try:
            response = goi_gemini_co_retry(prompt, SEED_REVIEW_BATCH_SCHEMA, so_lan_thu=3)
            data = json.loads(response.text)
            results = data.get("items", []) or []
        except Exception as e:
            results = []
            err = str(e)[:300]
            for q in batch:
                qq = by_id.get(q.get("id"))
                if qq is not None:
                    qq["kiem_tra_dap_an"] = "Chưa đủ dữ kiện"
                    qq["trang_thai_kiem_tra_dap_an"] = "Chưa đủ dữ kiện"
                    qq["do_tin_cay_dap_an"] = 0.0
                    qq["canh_bao_dap_an"] = "Chưa kiểm tra được bằng AI: " + err
                    qq["duoc_dung_lam_hat_giong"] = False
                    qq["gv_da_duyet_dap_an"] = False
            continue

        result_map = {str(x.get("id", "")): x for x in results if isinstance(x, dict)}
        for q in batch:
            qq = by_id.get(q.get("id"))
            if qq is None:
                continue
            r = result_map.get(str(q.get("id", "")), {})
            ket = str(r.get("ket_luan", "")).strip()
            if ket not in {"Khớp", "Không khớp", "Chưa đủ dữ kiện"}:
                ket = "Chưa đủ dữ kiện"
            conf = float(r.get("do_tin_cay", 0) or 0)
            qq["kiem_tra_dap_an"] = ket
            qq["trang_thai_kiem_tra_dap_an"] = ket
            qq["do_tin_cay_dap_an"] = conf
            qq["canh_bao_dap_an"] = str(r.get("ly_do", "") or "").strip()
            qq["dap_an_ai_de_xuat"] = (
                list(r.get("dap_an_4_y", []) or [])
                if qq.get("dang_cau_goi_y") == "Đúng / Sai"
                else str(r.get("dap_an_de_xuat", "") or "").strip()
            )
            qq["gv_da_duyet_dap_an"] = False

            # Chỉ câu Khớp với độ tin cậy đủ cao mới tự động được dùng.
            # Không khớp/không đủ dữ kiện phải chờ GV duyệt.
            if ket == "Khớp" and conf >= 0.80:
                qq["duoc_dung_lam_hat_giong"] = True
            else:
                qq["duoc_dung_lam_hat_giong"] = False

    return ds


def _seed_fingerprint_noi_dung(seed_or_text):
    """Dấu vân tay nội dung hạt giống, bỏ metadata/đáp án/số thứ tự; không gọi AI."""
    if isinstance(seed_or_text, dict):
        raw = str(seed_or_text.get("noi_dung_hien_thi", "") or seed_or_text.get("noi_dung_goc", "") or "")
    else:
        raw = str(seed_or_text or "")
    try:
        clean = _seed_xoa_metadata_de_hien_thi(raw)
    except Exception:
        clean = raw
    clean = re.sub(r"(?im)^\s*(?:Câu(?:\s+hỏi)?\s*)?\d+\s*[\.\:\)\-]\s*", "", clean, count=1)
    clean = chuan_hoa_noi_dung_trung(clean)
    return hashlib.sha256(clean.encode("utf-8")).hexdigest() if clean else ""


def _seed_text_key(seed_or_text):
    if isinstance(seed_or_text, dict):
        raw = str(seed_or_text.get("noi_dung_hien_thi", "") or seed_or_text.get("noi_dung_goc", "") or "")
    else:
        raw = str(seed_or_text or "")
    try:
        raw = _seed_xoa_metadata_de_hien_thi(raw)
    except Exception:
        pass
    raw = re.sub(r"(?im)^\s*(?:Câu(?:\s+hỏi)?\s*)?\d+\s*[\.\:\)\-]\s*", "", raw, count=1)
    return chuan_hoa_noi_dung_trung(raw)


def _seed_trung_voi_hat_giong(seed, seed_bank, nguong=0.92):
    """Phát hiện trùng/gần trùng trong kho hạt giống bằng Python cục bộ."""
    moi = _seed_text_key(seed)
    if not moi:
        return False, 0.0
    cao = 0.0
    for old in seed_bank or []:
        cu = _seed_text_key(old)
        if not cu:
            continue
        if moi == cu:
            return True, 1.0
        score = tinh_do_giong_noi_bo(moi, cu)
        if score > cao:
            cao = score
    return cao >= float(nguong), round(cao, 3)


def _seed_du_dieu_kien_dong_bo(seed):
    """Kiểm tra trường bắt buộc bằng Python; không thẩm định đáp án bằng AI."""
    if not str(seed.get("yccd", "") or "").strip():
        return False, "Thiếu YCCĐ"
    if not _seed_co_dap_an_nguon(seed):
        return False, "Thiếu đáp án nguồn"
    dang = str(seed.get("dang_cau", seed.get("dang_cau_goi_y", "")) or "").strip()
    if dang not in {"Trắc nghiệm 4 lựa chọn", "Đúng / Sai", "Trả lời ngắn"}:
        return False, "Chưa xác định dạng câu"
    return True, "Sẵn sàng"


def _seed_noi_dung_bank_key(q):
    return chuan_hoa_noi_dung_trung(tao_noi_dung_cau_de_so_sanh(q))


def _seed_trung_voi_ngan_hang(q, bank, nguong=0.92):
    """So trùng với ngân hàng chính bằng thuật toán cục bộ, không gọi API."""
    noi_moi = _seed_noi_dung_bank_key(q)
    if not noi_moi:
        return False, 0.0, ""
    cao_nhat = 0.0; id_gan = ""
    for old in bank or []:
        noi_cu = _seed_noi_dung_bank_key(old)
        if not noi_cu:
            continue
        if noi_moi == noi_cu:
            return True, 1.0, str(old.get("id", ""))
        score = tinh_do_giong_noi_bo(noi_moi, noi_cu)
        if score > cao_nhat:
            cao_nhat = score; id_gan = str(old.get("id", ""))
    return cao_nhat >= float(nguong), round(cao_nhat, 3), id_gan


def seed_duoc_phep_su_dung(seed):
    """Nguồn đã được GV chuẩn hóa trước; app không gọi AI để duyệt đáp án hạt giống."""
    return bool(seed.get("duoc_dung_lam_hat_giong", True))


def _seed_strip_meta_lines(text):
    """Bỏ các dòng metadata nhưng không làm mất phương án/nhận định nằm phía sau."""
    label_re = re.compile(
        r"^(?:ĐÁP\s*ÁN|DAP\s*AN|Đ/A|YCCĐ|YÊU\s*CẦU\s*CẦN\s*ĐẠT|"
        r"ĐƠN\s*VỊ\s*KIẾN\s*THỨC|KIẾN\s*THỨC|MỨC\s*ĐỘ(?:\s*TƯ\s*DUY)?|"
        r"CẤP\s*ĐỘ\s*TƯ\s*DUY|THÀNH\s*PHẦN\s*NĂNG\s*LỰC|CHỈ\s*BÁO)\b",
        flags=re.I,
    )
    out=[]
    skip_value=False
    for line in str(text or "").splitlines():
        stp=str(line or "").strip()
        if not stp:
            continue
        if re.match(r"^\[\[\s*(?:ĐÁP\s*ÁN|DAP\s*AN|HƯỚNG\s*DẪN\s*GIẢI|HUONG\s*DAN\s*GIAI|META\s+[A-Da-d])", stp, flags=re.I):
            continue
        if skip_value:
            # Giá trị của metadata dạng nhãn một dòng, giá trị ở dòng kế.
            if not re.match(r"^(?:Câu(?:\s+hỏi)?\s+\d+|[A-D]\s*[\.\)\:]|[a-d]\s*[\.\)\:])", stp):
                skip_value=False
                continue
            skip_value=False
        mm=label_re.match(stp)
        if mm:
            # nếu chỉ có nhãn, bỏ luôn giá trị ở dòng kế
            tail=re.sub(r"^[^:：]*[:：]?\s*", "", stp).strip()
            if not tail or tail.casefold() == stp.casefold():
                if ":" not in stp and "：" not in stp:
                    skip_value=True
            continue
        # Bỏ header và các hàng metadata bảng Word.
        if re.match(r"^Ý\s+Đáp\s*án\s+Chỉ\s*(?:báo|bảo)\s+Mức\s*độ\s*tư\s*duy", stp, flags=re.I):
            continue
        if re.match(r"^[a-d]\s+(?:Đ|S|Đúng|Sai)\s+(?:NT|TH|VD)\s*\d+\s+(?:Biết|Hiểu|Nhận\s*biết|Thông\s*hiểu|Vận\s*dụng)\b", stp, flags=re.I):
            continue
        out.append(line.rstrip())
    return "\n".join(out).strip()


def _seed_noi_dung_sach_theo_dang(chunk, dang):
    """Lấy đúng phần đề để hiển thị/so trùng/chuyển ngân hàng.

    File nguồn có thể đặt ``Đáp án:`` ngay sau thân câu rồi mới tới A-D hoặc bảng dữ liệu,
    vì vậy tuyệt đối không được cắt tài liệu tại dòng Đáp án đầu tiên.
    """
    raw = unicodedata.normalize("NFKC", str(chunk or "")).replace("\r\n", "\n").replace("\r", "\n")

    if dang == "Trắc nghiệm 4 lựa chọn":
        opts=[]
        first_pos=None
        for mm in re.finditer(r"(?m)^\s*([A-D])\s*[\.\)\:]\s*(.+?)\s*$", raw):
            lab=mm.group(1).upper()
            if lab in {x[0] for x in opts}:
                continue
            if first_pos is None:
                first_pos=mm.start()
            opts.append((lab, " ".join(mm.group(2).strip().split())))
            if len(opts) == 4:
                break
        if len(opts) >= 4 and first_pos is not None:
            stem=_seed_strip_meta_lines(raw[:first_pos])
            return (stem + "\n" + "\n".join(f"{a}. {b}" for a,b in opts[:4])).strip()
        # Một số file Word cũ bị mất riêng nhãn "A." nhưng vẫn còn B/C/D.
        # Nếu đúng 3 nhãn B,C,D và ngay trước B có một dòng ngắn chưa gắn nhãn,
        # coi dòng đó là phương án A thay vì làm mất cả câu.
        labs=[x[0] for x in opts]
        if labs == ["B","C","D"] and first_pos is not None:
            prefix=_seed_strip_meta_lines(raw[:first_pos])
            plines=[x.strip() for x in prefix.splitlines() if x.strip()]
            if len(plines) >= 2:
                a_text=plines[-1]
                stem="\n".join(plines[:-1]).strip()
                if len(a_text) <= 220 and not re.match(r"^(?:Câu|Câu hỏi)\b", a_text, flags=re.I):
                    all_opts=[("A",a_text)]+opts
                    return (stem + "\n" + "\n".join(f"{a}. {b}" for a,b in all_opts[:4])).strip()

    if dang == "Đúng / Sai":
        nds=[]; first_pos=None
        for mm in re.finditer(r"(?m)^\s*([a-d])\s*[\.\)\:]\s*(.+?)\s*$", raw):
            lab=mm.group(1)
            if lab in {x[0] for x in nds}:
                continue
            if first_pos is None:
                first_pos=mm.start()
            nds.append((lab, " ".join(mm.group(2).strip().split())))
            if len(nds) == 4:
                break
        if len(nds) >= 4 and first_pos is not None:
            stem=_seed_strip_meta_lines(raw[:first_pos])
            return (stem + "\n" + "\n".join(f"{a}) {b}" for a,b in nds[:4])).strip()

    # Trả lời ngắn/chưa xác định: file chuẩn hoá thường có 2 dòng đáp án:
    # dòng đầu là metadata mới, dòng sau là đáp án/lời giải cũ. Giữ toàn bộ dữ kiện
    # ở giữa hai dòng và cắt trước đáp án cũ để không kéo lời giải/tiêu đề phần sau vào câu.
    ans_marks = list(re.finditer(r"(?im)^\s*(?:Đ/A|ĐÁP\s*ÁN)\s*[:：]", raw))
    raw_question = raw[:ans_marks[1].start()] if len(ans_marks) >= 2 else raw
    clean=_seed_strip_meta_lines(raw_question)
    clean=re.split(
        r"(?im)^\s*(?:\[\[\s*)?(?:HƯỚNG\s*DẪN(?:\s*GIẢI)?|HUONG\s*DAN(?:\s*GIAI)?|HƯỚNG\s*GIẢI|GIẢI)(?:\s*\]\])?\s*[:：]?",
        clean, maxsplit=1
    )[0].strip()
    return clean


def _seed_xoa_metadata_de_hien_thi(text):
    """Bản làm sạch chung, không bao giờ cắt mất A-D/a-d chỉ vì gặp dòng Đáp án."""
    return _seed_strip_meta_lines(text)

def _seed_chuan_hoa_yccd_de_doi_chieu(text):
    """Chuẩn hóa YCCĐ chỉ để đối chiếu an toàn với KHO_YCCD.

    Không dùng AI/so khớp ngữ nghĩa. Chỉ bỏ khác biệt trình bày như Unicode,
    khoảng trắng, dấu đầu dòng và dấu câu; nội dung chữ của YCCĐ vẫn phải trùng.
    """
    s = unicodedata.normalize("NFKC", str(text or ""))
    s = s.replace("\u00a0", " ").strip().casefold()
    s = re.sub(r"^\s*(?:[-–—•·▪◦]+|yccđ\s*[:：])\s*", "", s, flags=re.I)
    s = re.sub(r"[^0-9a-zà-ỹđ]+", " ", s, flags=re.I)
    return " ".join(s.split())


def _seed_tim_pham_vi_tu_yccd(yccd):
    """Tìm Khối → Chương → Bài từ YCCĐ theo KHO_YCCD, không đoán.

    Chỉ trả kết quả khi YCCĐ chuẩn hóa khớp DUY NHẤT một phạm vi. Nếu cùng một
    YCCĐ xuất hiện ở nhiều bài khác nhau thì để trống để tránh gắn sai.
    """
    target = _seed_chuan_hoa_yccd_de_doi_chieu(yccd)
    if not target:
        return {"khoi": "", "chuong": "", "bai": ""}

    matches = []
    for khoi, ds_chuong in (KHO_YCCD or {}).items():
        for chuong, ds_bai in (ds_chuong or {}).items():
            for bai, ds_yccd in (ds_bai or {}).items():
                for yc in ds_yccd or []:
                    if isinstance(yc, dict):
                        yc_text = str(yc.get("noi_dung", yc.get("YCCĐ", yc.get("yccd", ""))))
                    else:
                        yc_text = str(yc)
                    if _seed_chuan_hoa_yccd_de_doi_chieu(yc_text) == target:
                        item = {
                            "khoi": str(khoi),
                            "chuong": str(chuong),
                            "bai": str(bai),
                            "yccd_chuan": yc_text.strip(),
                        }
                        if not any(
                            x["khoi"] == item["khoi"]
                            and x["chuong"] == item["chuong"]
                            and x["bai"] == item["bai"]
                            for x in matches
                        ):
                            matches.append(item)

    if len(matches) == 1:
        return matches[0]
    return {"khoi": "", "chuong": "", "bai": ""}


def _seed_tach_yccd_tung_y_tu_chuoi(value):
    """Tách YCCĐ dạng: a) ...; b) ...; c) ...; d) ... mà không sửa parser câu."""
    raw = unicodedata.normalize("NFKC", str(value or "")).strip()
    if not raw:
        return []

    marks = list(re.finditer(r"(?i)(?:^|[;\n])\s*([a-d])\s*[\)\.\:]\s*", raw))
    if len(marks) < 2:
        return [raw]

    values = []
    for i, m in enumerate(marks):
        end = marks[i + 1].start() if i + 1 < len(marks) else len(raw)
        part = raw[m.end():end].strip(" \t\r\n;,.：:")
        if part:
            values.append(part)
    return values


def _seed_lay_cac_yccd_cho_pham_vi(seed):
    """Lấy các YCCĐ của một seed chỉ để xác định phạm vi luyện HS.

    Câu thường thường có 1 YCCĐ. Câu Đúng/Sai mẫu File 4 có thể ghi 4 YCCĐ
    a-d trên cùng một dòng; hàm này đọc được kiểu đó mà không đụng tới các
    metadata khác của parser hiện có.
    """
    ds = []

    for nd in list(seed.get("nhan_dinh_meta", []) or []):
        if isinstance(nd, dict):
            y = str(nd.get("yccd", "") or "").strip()
            for item in _seed_tach_yccd_tung_y_tu_chuoi(y):
                if item:
                    ds.append(item)

    yccd_seed = str(seed.get("yccd", "") or "").strip()
    for item in _seed_tach_yccd_tung_y_tu_chuoi(yccd_seed):
        if item:
            ds.append(item)

    if not ds:
        raw = str(seed.get("noi_dung_goc", "") or "")
        if raw:
            y_raw = _seed_lay_meta(raw, "YCCĐ") or _seed_lay_meta(raw, "YÊU CẦU CẦN ĐẠT")
            for item in _seed_tach_yccd_tung_y_tu_chuoi(y_raw):
                if item:
                    ds.append(item)

    out = []
    seen = set()
    for y in ds:
        key = _seed_chuan_hoa_yccd_de_doi_chieu(y)
        if key and key not in seen:
            seen.add(key)
            out.append(y)
    return out


def _seed_gan_pham_vi_luyen_hs_tu_yccd(seed, q):
    """Gắn phạm vi luyện HS từ YCCĐ chuẩn trong KHO_YCCD.

    Không dùng Gemini, không đoán bằng từ khóa; chỉ gắn khi YCCĐ khớp duy nhất.
    Với Đúng/Sai nhiều YCCĐ, lưu toàn bộ phạm vi để lọc đúng từng bài/chương.
    """
    scopes = []
    seen_scope = set()

    for y in _seed_lay_cac_yccd_cho_pham_vi(seed):
        pv = _seed_tim_pham_vi_tu_yccd(y)
        if not pv.get("khoi") or not pv.get("chuong") or not pv.get("bai"):
            continue
        key = (pv.get("khoi", ""), pv.get("chuong", ""), pv.get("bai", ""))
        if key in seen_scope:
            continue
        seen_scope.add(key)
        scopes.append({
            "khoi": pv.get("khoi", ""),
            "chuong": pv.get("chuong", ""),
            "bai": pv.get("bai", ""),
            "yccd": pv.get("yccd_chuan", y),
        })

    if scopes:
        q["pham_vi_hat_giong"] = scopes
        if len(scopes) == 1:
            q["khoi"] = scopes[0]["khoi"]
            q["chuong"] = scopes[0]["chuong"]
            q["bai"] = scopes[0]["bai"]
            if str(q.get("dang_cau", "")) != "Đúng / Sai":
                q["yccd"] = scopes[0]["yccd"]
        else:
            khoi_set = {x["khoi"] for x in scopes}
            chuong_set = {(x["khoi"], x["chuong"]) for x in scopes}
            if len(khoi_set) == 1:
                q["khoi"] = next(iter(khoi_set))
            if len(chuong_set) == 1:
                q["chuong"] = scopes[0]["chuong"]
            q["bai"] = q.get("bai", "") if any(q.get("bai") == x["bai"] for x in scopes) else ""

        q["pham_vi_hat_giong_xac_dinh_bang"] = "YCCĐ chuẩn trong yccd.json"
        q["pham_vi_hat_giong_da_xac_dinh"] = True
        return True

    q["pham_vi_hat_giong"] = []
    q["pham_vi_hat_giong_da_xac_dinh"] = False
    return False


def _seed_goi_y_yccd_tu_kho(seed):
    """Gợi ý YCCĐ khi file nguồn không ghi YCCĐ; không ghi đè metadata GV."""
    try:
        text = " ".join([
            str(seed.get("kien_thuc_chu_de", "") or ""),
            str(seed.get("noi_dung_hien_thi", "") or seed.get("noi_dung_goc", "") or ""),
        ]).strip()
        if not text:
            return {"yccd": "", "khoi": "", "chuong": "", "bai": "", "do_tin_cay": 0.0}
        q_tmp = {"cau_hoi": text, "tinh_huong": "", "yccd": "", "dang_cau": seed.get("dang_cau", seed.get("dang_cau_goi_y", ""))}
        ranked_scope = _grad_xep_hang_ung_vien_pham_vi(q_tmp, top_n=5)
        allowed = None
        if ranked_scope:
            b = ranked_scope[0]
            allowed = (str(b.get("khoi", "")), str(b.get("chuong", "")), str(b.get("bai", "")))
        query_tokens = _grad_tokens(text)
        if not query_tokens:
            return {"yccd": "", "khoi": "", "chuong": "", "bai": "", "do_tin_cay": 0.0}
        scored=[]
        for khoi, ds_chuong in (KHO_YCCD or {}).items():
            for chuong, ds_bai in (ds_chuong or {}).items():
                for bai, ds_yccd in (ds_bai or {}).items():
                    if allowed and (str(khoi), str(chuong), str(bai)) != allowed:
                        continue
                    for yc in ds_yccd or []:
                        yc_text = str(yc.get("noi_dung", yc.get("YCCĐ", yc.get("yccd", ""))) if isinstance(yc, dict) else yc).strip()
                        if not yc_text:
                            continue
                        yt=_grad_tokens(yc_text)
                        inter=query_tokens & yt
                        if not inter:
                            continue
                        jac=len(inter)/max(1,len(query_tokens|yt))
                        coverage=len(inter)/max(1,len(yt))
                        score=0.55*coverage+0.45*jac
                        scored.append((score,yc_text,str(khoi),str(chuong),str(bai)))
        if not scored:
            return {"yccd": "", "khoi": "", "chuong": "", "bai": "", "do_tin_cay": 0.0}
        scored.sort(key=lambda x:x[0], reverse=True)
        best=scored[0]; second=scored[1][0] if len(scored)>1 else 0.0
        conf=min(0.94,0.45+0.55*best[0]+0.20*max(0.0,best[0]-second))
        return {"yccd":best[1],"khoi":best[2],"chuong":best[3],"bai":best[4],"do_tin_cay":round(conf,2)}
    except Exception:
        return {"yccd": "", "khoi": "", "chuong": "", "bai": "", "do_tin_cay": 0.0}


def _seed_chuyen_thanh_cau_ngan_hang(seed):
    """Chuyển câu hạt giống đã an toàn thành câu dùng được trong ngân hàng chính, không gọi AI."""
    if not seed_duoc_phep_su_dung(seed):
        return None

    dang = str(seed.get("dang_cau", seed.get("dang_cau_goi_y", ""))).strip()
    raw = str(seed.get("noi_dung_hien_thi", "") or _seed_noi_dung_sach_theo_dang(seed.get("noi_dung_goc", ""), dang)).strip()
    raw = re.sub(r"(?im)^\s*Câu\s+\d+\s*[\.\:\)\-]\s*", "", raw, count=1).strip()

    q = {
        "id": str(uuid.uuid4()),
        "temp_id": str(uuid.uuid4()),
        "nguon_seed_id": seed.get("id", ""),
        "nguon_tao": "Hạt giống đã kiểm tra an toàn",
        "nguon": str(seed.get("nguon_file", "")),
        "nguon_file": str(seed.get("nguon_file", "")),
        "so_cau_goc": str(seed.get("so_cau_goc", "")),
        "dang_cau": dang,
        "yccd": str(seed.get("yccd", "")).strip(),
        "muc_do": str(seed.get("muc_do", "")).strip(),
        "thanh_phan_nang_luc": chuan_hoa_ten_nang_luc_seed(seed.get("thanh_phan_nang_luc", "")),
        "chi_bao": str(seed.get("chi_bao", "")).strip(),
        "hanh_vi_nang_luc": str(seed.get("chi_bao", "")).strip(),
        "kien_thuc_chu_de": str(seed.get("kien_thuc_chu_de", "")).strip(),
        "muc_dich_su_dung": "on_tap_kiem_tra",
        "trang_thai": "Hạt giống an toàn",
        "duoc_dung_luyen_hs": True,
        "ngay_tao": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "giai_thich": str(
            seed.get("giai_thich_nguon", seed.get("giai_thich", "")) or ""
        ).strip(),
        "nguon_giai_thich": (
            str(seed.get("nguon_giai_thich", "") or "").strip()
            or (
                "Lời giải từ file nguồn"
                if str(seed.get("giai_thich_nguon", seed.get("giai_thich", "")) or "").strip()
                else ""
            )
        ),
        "lua_chon": [],
        "dap_an": "",
        "tinh_huong": "",
        "cau_hoi": "",
        "nhan_dinh_meta": [],
        # Giữ nguyên ảnh/bảng từ hạt giống khi chuyển sang ngân hàng ôn tập.
        "tai_nguyen_truc_quan": [
            dict(x)
            for x in (seed.get("tai_nguyen_truc_quan", []) or [])
            if isinstance(x, dict)
        ],
        "du_lieu_truc_quan": dict(
            seed.get("du_lieu_truc_quan", {}) or {}
        ),
    }

    # Chỉ bổ sung phạm vi luyện HS từ YCCĐ chuẩn; không thay đổi nội dung câu.
    pham_vi_tu_yccd_ok = _seed_gan_pham_vi_luyen_hs_tu_yccd(seed, q)

    if dang == "Trắc nghiệm 4 lựa chọn":
        ms = list(re.finditer(r"(?m)^\s*([A-D])[\.\)]\s*", raw))
        if len(ms) >= 4:
            q["cau_hoi"] = raw[:ms[0].start()].strip()
            choices = []
            for i, m in enumerate(ms[:4]):
                end = ms[i+1].start() if i+1 < 4 else len(raw)
                choices.append(f"{m.group(1).upper()}. {raw[m.end():end].strip()}")
            q["lua_chon"] = choices
        else:
            q["cau_hoi"] = raw
        dap = seed.get("dap_an_nguon", "")
        if isinstance(dap, list):
            dap = dap[0] if dap else ""
        q["dap_an"] = str(dap).strip().upper()

    elif dang == "Đúng / Sai":
        meta = []
        for i, nd in enumerate(list(seed.get("nhan_dinh_meta", []) or [])[:4]):
            nd2 = dict(nd)
            nd2.setdefault("ky_hieu", "abcd"[i])
            nd2["yccd"] = nd2.get("yccd") or q["yccd"]
            nd2["muc_do"] = nd2.get("muc_do") or q["muc_do"]
            nd2["thanh_phan_nang_luc"] = chuan_hoa_ten_nang_luc_seed(nd2.get("thanh_phan_nang_luc") or q["thanh_phan_nang_luc"])
            nd2["chi_bao"] = nd2.get("chi_bao") or q["chi_bao"]
            nd2.setdefault("giai_thich", "")
            meta.append(nd2)
        # Nếu file chỉ có đáp án chung ĐĐSĐ mà bảng metadata không đọc đủ,
        # vẫn tạo đủ 4 nhận định để không mất đáp án khi chuyển sang ngân hàng.
        ans_list = list(seed.get("dap_an_nguon", []) or []) if isinstance(seed.get("dap_an_nguon", []), list) else []
        if len(meta) < 4:
            nd_text = {}
            for mm_nd in re.finditer(r"(?m)^\s*([a-d])[\.\)]\s*(.+?)\s*$", raw):
                nd_text.setdefault(mm_nd.group(1), " ".join(mm_nd.group(2).strip().split()))
            meta = []
            for i_nd, ky in enumerate("abcd"):
                if ky not in nd_text:
                    continue
                meta.append({
                    "ky_hieu": ky, "noi_dung": nd_text[ky],
                    "dap_an": ans_list[i_nd] if i_nd < len(ans_list) else "",
                    "yccd": q["yccd"], "muc_do": q["muc_do"],
                    "thanh_phan_nang_luc": q["thanh_phan_nang_luc"],
                    "chi_bao": q["chi_bao"], "giai_thich": ""
                })
        else:
            for i_nd, nd2 in enumerate(meta[:4]):
                if not str(nd2.get("dap_an", "")).strip() and i_nd < len(ans_list):
                    nd2["dap_an"] = ans_list[i_nd]
        q["nhan_dinh_meta"] = meta
        first_nd = re.search(r"(?m)^\s*[a-d][\.\)]\s*", raw)
        common = raw[:first_nd.start()].strip() if first_nd else raw
        q["tinh_huong"] = common
        q["cau_hoi"] = "Hãy xác định mỗi nhận định sau là Đúng hay Sai."
        q["dap_an"] = "".join("Đ" if chuan_hoa_dap_an_dung_sai(x.get("dap_an", "")) == "Đúng" else "S" for x in meta)

    else:
        q["cau_hoi"] = raw
        dap = seed.get("dap_an_nguon", "")
        if isinstance(dap, list):
            dap = dap[0] if dap else ""
        q["dap_an"] = str(dap).strip()

    # Giữ cơ chế cũ chỉ khi seed hoàn toàn không có YCCĐ để đối chiếu.
    # Có YCCĐ nhưng không khớp kho thì KHÔNG đoán, tránh gắn sai bài/chương/khối.
    if (
        not pham_vi_tu_yccd_ok
        and not _seed_lay_cac_yccd_cho_pham_vi(seed)
        and (not q.get("khoi") or not q.get("chuong"))
    ):
        try:
            pv2 = xac_dinh_pham_vi_bai_chuong_tot_nghiep(q, dung_ai=False)
            q["khoi"] = q.get("khoi") or pv2.get("khoi", "")
            q["chuong"] = q.get("chuong") or pv2.get("chuong", "")
            q["bai"] = q.get("bai") or pv2.get("bai", "")
        except Exception:
            pass

    return q


def _seed_tim_trung_chinh_xac_ngan_hang(q, bank):
    """Chỉ coi là trùng tự động khi nội dung chuẩn hoá bằng nhau.

    Có thêm đường nâng cấp cho dữ liệu cũ bị lỗi parser: cùng thân câu nhưng bản cũ
    thiếu lựa chọn/nhận định thì xem là cùng câu để bổ sung metadata, không tạo bản sao.
    """
    key_new=_seed_noi_dung_bank_key(q)
    stem_new=chuan_hoa_noi_dung_trung(str(q.get("cau_hoi", "") or q.get("tinh_huong", "")))
    src_new = str(q.get("nguon_file", q.get("nguon", "")) or "").strip().casefold()
    num_new = str(q.get("so_cau_goc", "") or "").strip()
    for old in bank or []:
        # Cùng file + cùng số câu CHỈ được coi là cùng câu khi CÙNG DẠNG.
        # Trong file hạt giống, số câu thường được đánh lại từ 1 ở mỗi phần
        # (MCQ / Đúng-Sai / Trả lời ngắn). Nếu bỏ điều kiện cùng dạng,
        # Câu 1 Đ/S hoặc TL ngắn sẽ bị nhầm với Câu 1 MCQ.
        src_old = str(old.get("nguon_file", old.get("nguon", "")) or "").strip().casefold()
        num_old = str(old.get("so_cau_goc", "") or "").strip()
        dang_old = str(old.get("dang_cau", "") or "").strip()
        dang_new = str(q.get("dang_cau", "") or "").strip()
        if (
            src_new and num_new
            and src_old == src_new
            and num_old == num_new
            and dang_old == dang_new
        ):
            return old, 1.0
        key_old=_seed_noi_dung_bank_key(old)
        if key_new and key_old and key_new == key_old:
            return old, 1.0
        # Sửa dữ liệu cũ đã nhập từ hạt giống nhưng bị mất A-D/a-d.
        stem_old=chuan_hoa_noi_dung_trung(str(old.get("cau_hoi", "") or old.get("tinh_huong", "")))
        old_incomplete = (
            (q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn" and not list(old.get("lua_chon", []) or []))
            or (q.get("dang_cau") == "Đúng / Sai" and not list(old.get("nhan_dinh_meta", []) or []))
        )
        if old_incomplete and stem_new and stem_old and stem_new == stem_old:
            return old, 1.0
    return None, 0.0


def _seed_bo_sung_main_tu_nguon(old, new):
    """Bổ sung/sửa bản main từ file nguồn đã chuẩn hoá, không tạo bản sao.

    Nếu câu main vốn được sinh từ hạt giống cùng file/số câu, cho phép sửa cấu trúc
    bị parser cũ làm hỏng (mất A-D, nhầm MCQ thành Đ/S, thiếu bảng dữ kiện...).
    Các câu main không có nguồn hạt giống chỉ được điền trường còn trống.
    """
    changed=False
    src_old=str(old.get("nguon_file", old.get("nguon", "")) or "").strip().casefold()
    src_new=str(new.get("nguon_file", new.get("nguon", "")) or "").strip().casefold()
    num_old=str(old.get("so_cau_goc", "") or "").strip()
    num_new=str(new.get("so_cau_goc", "") or "").strip()
    from_seed = bool(old.get("nguon_seed_id")) or "hạt giống" in str(old.get("nguon_tao", "")).casefold()

    # Nếu đúng cùng seed_id thì cho phép sửa lại bản main đã từng bị parser cũ
    # ghi nhầm dạng câu. Đây cũng giúp tự phục hồi dữ liệu đã bị lỗi trước bản sửa này.
    old_seed_id = str(old.get("nguon_seed_id", "") or "").strip()
    new_seed_id = str(new.get("nguon_seed_id", "") or "").strip()
    same_seed_id = bool(
        old_seed_id and new_seed_id and old_seed_id == new_seed_id
    )

    # Nếu chỉ dựa vào file+số câu thì BẮT BUỘC cùng dạng câu,
    # vì mỗi phần MCQ / Đúng-Sai / TL ngắn có thể đánh số lại từ Câu 1.
    dang_old = str(old.get("dang_cau", "") or "").strip()
    dang_new = str(new.get("dang_cau", "") or "").strip()
    same_source_number_type = bool(
        src_old and src_new
        and src_old == src_new
        and num_old and num_old == num_new
        and dang_old == dang_new
    )

    same_source = same_seed_id or same_source_number_type

    # Chỉ cập nhật lại phạm vi của đúng bản main sinh từ seed này.
    if from_seed and same_source and list(new.get("pham_vi_hat_giong", []) or []):
        for f in ["khoi", "chuong", "bai", "pham_vi_hat_giong",
                  "pham_vi_hat_giong_xac_dinh_bang", "pham_vi_hat_giong_da_xac_dinh"]:
            nv = new.get(f)
            if old.get(f) != nv:
                old[f] = nv
                changed = True

    if from_seed and same_source:
        # Nguồn đã được kiểm tra trước khi nhập -> dùng lại cấu trúc chuẩn để sửa bản cũ.
        for f in [
            "dang_cau", "cau_hoi", "tinh_huong", "lua_chon",
            "nhan_dinh_meta", "giai_thich", "nguon_giai_thich",
            "tai_nguyen_truc_quan", "du_lieu_truc_quan"
        ]:
            nv=new.get(f)
            if nv not in (None, "", [], {}) and old.get(f) != nv:
                old[f]=nv; changed=True
        if str(new.get("dap_an", "") or "").strip() and old.get("dap_an") != new.get("dap_an"):
            old["dap_an"]=new.get("dap_an"); changed=True

    fill_fields=["yccd","muc_do","thanh_phan_nang_luc","chi_bao","hanh_vi_nang_luc","kien_thuc_chu_de","khoi","chuong","bai"]
    for f in fill_fields:
        if not str(old.get(f, "") or "").strip() and str(new.get(f, "") or "").strip():
            old[f]=new.get(f); changed=True
    if not str(old.get("dap_an", "") or "").strip() and str(new.get("dap_an", "") or "").strip():
        old["dap_an"]=new.get("dap_an"); changed=True
    if not str(old.get("giai_thich", "") or "").strip() and str(new.get("giai_thich", "") or "").strip():
        old["giai_thich"] = new.get("giai_thich", "")
        old["nguon_giai_thich"] = new.get("nguon_giai_thich", "") or "Lời giải từ file nguồn"
        changed = True
    if new.get("dang_cau") == "Trắc nghiệm 4 lựa chọn" and not list(old.get("lua_chon", []) or []) and list(new.get("lua_chon", []) or []):
        old["lua_chon"]=new.get("lua_chon"); changed=True
    if new.get("dang_cau") == "Đúng / Sai" and not list(old.get("nhan_dinh_meta", []) or []) and list(new.get("nhan_dinh_meta", []) or []):
        old["nhan_dinh_meta"]=new.get("nhan_dinh_meta"); changed=True
        old["tinh_huong"]=new.get("tinh_huong", old.get("tinh_huong", "")); changed=True
    if changed:
        old["ngay_cap_nhat_tu_hat_giong"] = datetime.now().strftime("%d/%m/%Y %H:%M")
        old["trang_thai"] = "Đã cập nhật từ nguồn hạt giống chuẩn hóa"
    return changed


def dong_bo_hat_giong_an_toan_sang_ngan_hang(seed_bank=None):
    """Đồng bộ câu đủ metadata sang ngân hàng chính, chỉ dùng chống trùng cục bộ."""
    seed_bank = seed_bank if seed_bank is not None else doc_ngan_hang_hat_giong()
    bank = doc_ngan_hang()
    main_theo_seed = {
        str(q.get("nguon_seed_id", "")): q
        for q in bank
        if str(q.get("nguon_seed_id", "")).strip()
    }
    da_co_seed = set(main_theo_seed)
    them = 0; bo_qua = 0
    for seed in seed_bank or []:
        sid = str(seed.get("id", ""))
        if sid and sid in da_co_seed:
            # Seed đã từng đồng bộ: vẫn nâng cấp ảnh/bảng cho bản main cũ.
            q_cap_nhat = _seed_chuyen_thanh_cau_ngan_hang(seed)
            old_main = main_theo_seed.get(sid)
            if q_cap_nhat is not None and old_main is not None:
                _seed_bo_sung_main_tu_nguon(old_main, q_cap_nhat)
            seed["da_chuyen_sang_ngan_hang"] = True
            seed["trang_thai_dong_bo"] = "Đã có trong ngân hàng chính"
            bo_qua += 1
            continue
        ok, ly_do = _seed_du_dieu_kien_dong_bo(seed)
        if not ok:
            seed["trang_thai_dong_bo"] = ly_do; bo_qua += 1; continue
        q = _seed_chuyen_thanh_cau_ngan_hang(seed)
        if not q or not str(q.get("cau_hoi", "")).strip():
            seed["trang_thai_dong_bo"] = "Không chuyển được cấu trúc câu"; bo_qua += 1; continue
        old_exact, ti_le = _seed_tim_trung_chinh_xac_ngan_hang(q, bank)
        if old_exact is not None:
            _seed_bo_sung_main_tu_nguon(old_exact, q)
            seed["da_chuyen_sang_ngan_hang"] = True
            seed["trang_thai_dong_bo"] = "Đã có trong ngân hàng chính"
            seed["id_cau_trung_ngan_hang"] = str(old_exact.get("id", ""))
            if sid: da_co_seed.add(sid)
            bo_qua += 1
            continue
        # Gần giống chỉ là cảnh báo, KHÔNG tự xoá/bỏ câu.
        near, near_score, near_id = _seed_trung_voi_ngan_hang(q, bank, nguong=0.92)
        if near and near_score < 1.0:
            seed["nghi_trung_ngan_hang"] = True
            seed["do_giong_ngan_hang"] = near_score
            seed["id_cau_gan_giong_ngan_hang"] = near_id
        q["nguon_tao"] = "Hạt giống đã chuẩn hóa trước khi nhập"
        q["trang_thai"] = "Đã duyệt từ nguồn chuẩn hóa"
        bank.append(q)
        if sid:
            da_co_seed.add(sid)
            main_theo_seed[sid] = q
        seed["da_chuyen_sang_ngan_hang"] = True
        seed["trang_thai_dong_bo"] = "Đã chuyển"
        seed["ngay_chuyen_sang_ngan_hang"] = datetime.now().strftime("%d/%m/%Y %H:%M")
        them += 1
    luu_ngan_hang_hat_giong(seed_bank)
    luu_ngan_hang(bank)
    return them, bo_qua


def tach_cau_hoi_hat_giong(
    text_goc,
    ten_file,
    resources_map=None,
    source_path="",
    media_dir=""
):
    s = _seed_chuan_hoa_text_de_tach(text_goc)
    matches = _seed_tim_moc_cau(s)
    if not matches:
        return []

    ds = []

    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(s)
        chunk_with_resources = s[start:end].strip()
        tai_nguyen_cau = _seed_tai_nguyen_tu_chunk(
            chunk_with_resources,
            resources_map or {}
        )
        chunk = _seed_xoa_marker_tai_nguyen(
            chunk_with_resources
        )

        if len(chunk) < 12:
            continue

        dang = "Chưa xác định"
        co_abcd = _seed_doan_co_4_phuong_an(chunk)
        co_tf = _seed_doan_co_4_nhan_dinh(chunk)
        low = chunk.casefold()

        # Phân dạng dựa vào cấu trúc câu trước. Không dùng tiêu đề của PHẦN KẾ TIẾP
        # (ví dụ "TRẮC NGHIỆM ĐÚNG/SAI") để gán nhầm câu cuối phần MCQ.
        if co_tf:
            dang = "Đúng / Sai"
        elif co_abcd:
            dang = "Trắc nghiệm 4 lựa chọn"
        elif any(x in low for x in [
            "bao nhiêu", "hãy viết", "tính giá trị", "xác định số",
            "kết quả bằng", "chỉ ghi phần số", "trả lời ngắn"
        ]) or _seed_lay_meta(chunk, "ĐÁP ÁN") or _seed_lay_meta(chunk, "DAP AN"):
            dang = "Trả lời ngắn"
        elif "đúng hay sai" in low or "đúng/sai" in low or "đánh giá tính đúng" in low:
            dang = "Đúng / Sai"

        # Metadata cấp câu: hỗ trợ cả file cũ lẫn mẫu mới có [[...]].
        nl_seed = _seed_lay_meta(chunk, "THÀNH PHẦN NĂNG LỰC")
        chi_bao_seed = _seed_lay_meta(chunk, "CHỈ BÁO").upper()
        muc_do_seed = _seed_chuan_hoa_muc_do(
            _seed_lay_meta(chunk, "MỨC ĐỘ")
            or _seed_lay_meta(chunk, "CẤP ĐỘ TƯ DUY")
            or _seed_lay_meta(chunk, "MỨC ĐỘ TƯ DUY")
        )
        kien_thuc_seed = (
            _seed_lay_meta(chunk, "ĐƠN VỊ KIẾN THỨC")
            or _seed_lay_meta(chunk, "KIẾN THỨC")
        )
        yccd_seed = (
            _seed_lay_meta(chunk, "YCCĐ")
            or _seed_lay_meta(chunk, "YÊU CẦU CẦN ĐẠT")
        )

        nhan_dinh_meta = []
        if dang == "Đúng / Sai":
            nhan_dinh_meta = _seed_tach_nhan_dinh_ds(chunk)
            if nhan_dinh_meta:
                md_ds, nl_ds, cb_ds = _seed_tom_tat_meta_ds(nhan_dinh_meta)
                # Metadata từng ý là dữ liệu GV cụ thể hơn -> ưu tiên hơn metadata chung.
                muc_do_seed = md_ds or muc_do_seed
                nl_seed = nl_ds or nl_seed
                chi_bao_seed = cb_ds or chi_bao_seed
                for nd in nhan_dinh_meta:
                    nd["yccd"] = yccd_seed

        dap_an_nguon = _seed_doc_dap_an_nguon(chunk, dang, nhan_dinh_meta)
        giai_thich_nguon = _seed_lay_huong_dan_giai(chunk)
        # Nếu Đ/S có đáp án chung nhưng META chưa ghi đáp án, phân bổ lại cho từng ý.
        if dang == "Đúng / Sai" and isinstance(dap_an_nguon, list) and len(dap_an_nguon) == 4 and nhan_dinh_meta:
            for idx_nd, nd in enumerate(nhan_dinh_meta[:4]):
                if not str(nd.get("dap_an", "")).strip():
                    nd["dap_an"] = dap_an_nguon[idx_nd]

        so_cau = str(m.group(1)).strip()

        ds.append({
            "id": str(uuid.uuid4()),
            "nguon_file": ten_file,
            "so_cau_goc": so_cau,
            "noi_dung_goc": chunk,
            "noi_dung_hien_thi": _seed_noi_dung_sach_theo_dang(chunk, dang),
            "dang_cau_goi_y": dang,
            "dang_cau": dang,
            "trang_thai": "Đã đọc metadata" if (
                nl_seed or chi_bao_seed or muc_do_seed or kien_thuc_seed
                or yccd_seed or nhan_dinh_meta
            ) else "Chờ phân loại",
            "yccd": yccd_seed,
            "yccd_goi_y": "",
            "do_tin_cay_yccd_goi_y": 0.0,
            "trang_thai_yccd": "GV/file cung cấp" if yccd_seed else "Chưa xác định",
            "muc_do": muc_do_seed,
            "thanh_phan_nang_luc": nl_seed,
            "chi_bao": chi_bao_seed,
            "hanh_vi_nang_luc": chi_bao_seed,
            "kien_thuc_chu_de": kien_thuc_seed,
            "nhan_dinh_meta": nhan_dinh_meta,
            "tai_nguyen_truc_quan": tai_nguyen_cau,
            "du_lieu_truc_quan": _seed_compat_du_lieu_truc_quan(
                tai_nguyen_cau,
                ten_file
            ),
            "nguon_file_path": source_path,
            "thu_muc_media": media_dir,
            "dap_an_nguon": dap_an_nguon,
            "giai_thich_nguon": giai_thich_nguon,
            "nguon_giai_thich": (
                "Lời giải từ file nguồn" if str(giai_thich_nguon or "").strip() else ""
            ),
            "kiem_tra_dap_an": "Nguồn đã chuẩn hóa trước khi nhập",
            "trang_thai_kiem_tra_dap_an": "Không kiểm tra AI",
            "do_tin_cay_dap_an": None,
            "canh_bao_dap_an": "",
            "dap_an_ai_de_xuat": [],
            "gv_da_duyet_dap_an": True,
            "duoc_dung_lam_hat_giong": True,
            "nguon_da_kiem_tra_truoc": True,
            "muc_dich_su_dung": "",
            "ngay_nhap": datetime.now().strftime("%d/%m/%Y %H:%M")
        })

    for q in ds:
        if not str(q.get("yccd", "") or "").strip():
            gy = _seed_goi_y_yccd_tu_kho(q)
            if gy.get("yccd"):
                q["yccd_goi_y"] = gy.get("yccd", "")
                q["do_tin_cay_yccd_goi_y"] = float(gy.get("do_tin_cay", 0) or 0)
                q["trang_thai_yccd"] = "Cần GV kiểm tra"
                q["khoi_goi_y"] = gy.get("khoi", "")
                q["chuong_goi_y"] = gy.get("chuong", "")
                q["bai_goi_y"] = gy.get("bai", "")

    return ds


def doc_lai_metadata_hat_giong_hien_co(ds_cau):
    """Đọc lại metadata/đáp án từ ``noi_dung_goc`` cho dữ liệu hạt giống đã nhập trước bản sửa.

    Không gọi AI và không tự thay đổi đáp án của GV. Hàm chỉ sửa lỗi parser cũ,
    đồng thời tạo ``noi_dung_hien_thi`` sạch để giao diện không trộn metadata vào đề.
    """
    ds_moi = []
    da_sua = 0
    for old in (ds_cau or []):
        q = dict(old)
        chunk = str(q.get("noi_dung_goc", "") or "")
        if not chunk.strip():
            ds_moi.append(q)
            continue

        dang = str(q.get("dang_cau", q.get("dang_cau_goi_y", "")) or "").strip()
        if not dang or dang == "Chưa xác định":
            co_abcd = _seed_doan_co_4_phuong_an(chunk)
            co_tf = _seed_doan_co_4_nhan_dinh(chunk)
            low = chunk.casefold()
            if "đúng hay sai" in low or "đúng/sai" in low or co_tf:
                dang = "Đúng / Sai"
            elif co_abcd:
                dang = "Trắc nghiệm 4 lựa chọn"
            else:
                dang = "Trả lời ngắn"

        nl = _seed_lay_meta(chunk, "THÀNH PHẦN NĂNG LỰC")
        cb = _seed_lay_meta(chunk, "CHỈ BÁO").upper()
        md = _seed_chuan_hoa_muc_do(
            _seed_lay_meta(chunk, "MỨC ĐỘ")
            or _seed_lay_meta(chunk, "CẤP ĐỘ TƯ DUY")
            or _seed_lay_meta(chunk, "MỨC ĐỘ TƯ DUY")
        )
        kt = _seed_lay_meta(chunk, "ĐƠN VỊ KIẾN THỨC") or _seed_lay_meta(chunk, "KIẾN THỨC")
        yc = _seed_lay_meta(chunk, "YCCĐ") or _seed_lay_meta(chunk, "YÊU CẦU CẦN ĐẠT")

        meta = []
        if dang == "Đúng / Sai":
            meta = _seed_tach_nhan_dinh_ds(chunk)
            if meta:
                md_ds, nl_ds, cb_ds = _seed_tom_tat_meta_ds(meta)
                md = md_ds or md
                nl = nl_ds or nl
                cb = cb_ds or cb
                for nd in meta:
                    nd["yccd"] = yc or str(q.get("yccd", ""))

        ans = _seed_doc_dap_an_nguon(chunk, dang, meta)
        if dang == "Đúng / Sai" and isinstance(ans, list) and len(ans) == 4 and meta:
            for i_nd, nd in enumerate(meta[:4]):
                if not str(nd.get("dap_an", "")).strip():
                    nd["dap_an"] = ans[i_nd]

        before = (
            q.get("dap_an_nguon"), q.get("yccd"), q.get("muc_do"),
            q.get("thanh_phan_nang_luc"), q.get("chi_bao"), q.get("kien_thuc_chu_de")
        )
        q["dang_cau"] = dang
        q["dang_cau_goi_y"] = dang
        q["noi_dung_hien_thi"] = _seed_noi_dung_sach_theo_dang(chunk, dang)
        if yc:
            q["yccd"] = yc
        if md:
            q["muc_do"] = md
        if nl:
            q["thanh_phan_nang_luc"] = nl
        if cb:
            q["chi_bao"] = cb
            q["hanh_vi_nang_luc"] = cb
        if kt:
            q["kien_thuc_chu_de"] = kt
        if meta:
            q["nhan_dinh_meta"] = meta
        if ans:
            q["dap_an_nguon"] = ans

        if not str(q.get("yccd", "") or "").strip():
            gy = _seed_goi_y_yccd_tu_kho(q)
            if gy.get("yccd"):
                q["yccd_goi_y"] = gy.get("yccd", "")
                q["do_tin_cay_yccd_goi_y"] = float(gy.get("do_tin_cay", 0) or 0)
                q["trang_thai_yccd"] = "Cần GV kiểm tra"
                q["khoi_goi_y"] = gy.get("khoi", "")
                q["chuong_goi_y"] = gy.get("chuong", "")
                q["bai_goi_y"] = gy.get("bai", "")

        after = (
            q.get("dap_an_nguon"), q.get("yccd"), q.get("muc_do"),
            q.get("thanh_phan_nang_luc"), q.get("chi_bao"), q.get("kien_thuc_chu_de")
        )
        if before != after or not old.get("noi_dung_hien_thi"):
            da_sua += 1
        ds_moi.append(q)
    return ds_moi, da_sua


def danh_sach_yccd_khoi_12():
    ds = []
    for khoi_key, ds_chuong in KHO_YCCD.items():
        if str(khoi_key).strip() != "Khối 12" or not isinstance(ds_chuong, dict):
            continue
        for chuong, ds_bai in ds_chuong.items():
            if not isinstance(ds_bai, dict):
                continue
            for bai, ds_yccd in ds_bai.items():
                for yc in ds_yccd or []:
                    if isinstance(yc, dict):
                        noi_dung = str(yc.get("noi_dung", yc.get("YCCĐ", yc.get("yccd", "")))).strip()
                    else:
                        noi_dung = str(yc).strip()
                    if noi_dung:
                        ds.append({"chuong": str(chuong), "bai": str(bai), "yccd": noi_dung})
    return ds


HANH_VI_NANG_LUC_TOT_NGHIEP = [
    "Nhận biết / tái hiện kiến thức sinh học", "Giải thích cơ chế hoặc hiện tượng sinh học",
    "Đọc và khai thác bảng số liệu", "Phân tích biểu đồ / đồ thị", "Phân tích hình / sơ đồ sinh học",
    "Phân tích kết quả thí nghiệm", "Đánh giá giả thuyết / kết luận nghiên cứu",
    "Suy luận di truyền từ dữ kiện", "Tính toán và xử lí số liệu sinh học",
    "Vận dụng kiến thức vào tình huống thực tiễn", "Đề xuất / lựa chọn giải pháp sinh học"
]


def cau_phu_hop_tot_nghiep(q):
    muc_dich = str(q.get("muc_dich_su_dung", "")).strip()
    if muc_dich in {"tot_nghiep", "ca_hai"}:
        return True
    if muc_dich == "on_tap_kiem_tra":
        return False
    if str(q.get("hanh_vi_nang_luc", "")).strip() or str(q.get("tinh_huong", "")).strip():
        return True
    if q.get("du_lieu_truc_quan"):
        return True
    if q.get("dang_cau") in {"Đúng / Sai", "Trả lời ngắn"}:
        return True
    if q.get("thanh_phan_nang_luc") in {"Tìm hiểu thế giới sống", "Vận dụng kiến thức, kĩ năng đã học"}:
        return True
    return False


def hien_thi_du_lieu_truc_quan_cau(q):
    data = q.get("du_lieu_truc_quan", {}) or {}
    if not data:
        return
    loai = str(data.get("loai", "")).strip()
    if not loai or loai == "khong_co":
        return
    tieu_de = str(data.get("tieu_de", "")).strip()
    if tieu_de:
        st.markdown(f"**📊 {tieu_de}**")
    if _hien_thi_anh_tu_resource(
        data,
        width=GRAD_IMAGE_DISPLAY_WIDTH
    ):
        return
    cot = data.get("cot", []) or []
    rows = data.get("du_lieu", []) or []
    if cot and rows:
        try:
            df = pd.DataFrame(rows, columns=cot)
        except Exception:
            df = pd.DataFrame(rows)
        if loai == "bang_so_lieu":
            st.dataframe(df, use_container_width=True, hide_index=True)
        elif loai in {"bieu_do_cot", "bieu_do_duong"}:
            if len(df.columns) >= 2:
                x_col = df.columns[0]
                chart_df = df.copy()
                for c in chart_df.columns[1:]:
                    chart_df[c] = pd.to_numeric(chart_df[c], errors="coerce")
                chart_df = chart_df.set_index(x_col)
                if loai == "bieu_do_cot":
                    st.bar_chart(chart_df)
                else:
                    st.line_chart(chart_df)
            else:
                st.dataframe(df, use_container_width=True, hide_index=True)
        else:
            st.dataframe(df, use_container_width=True, hide_index=True)
    mo_ta = str(data.get("mo_ta", "")).strip()
    if mo_ta:
        st.caption(mo_ta)
    nguon = str(data.get("nguon", "")).strip()
    if nguon:
        st.caption("Nguồn dữ liệu trực quan: " + nguon)


GRAD_QUESTION_SCHEMA = {
    "type": "object", "properties": {"questions": {"type": "array", "minItems": 1, "maxItems": 1,
    "items": {"type": "object", "properties": {
        "khoi": {"type": "string"}, "chuong": {"type": "string"}, "bai": {"type": "string"},
        "yccd": {"type": "string"}, "muc_do": {"type": "string"}, "dang_cau": {"type": "string"},
        "thanh_phan_nang_luc": {"type": "string"}, "hanh_vi_nang_luc": {"type": "string"},
        "cau_hoi": {"type": "string"}, "lua_chon": {"type": "array", "items": {"type": "string"}},
        "dap_an": {"type": "string"}, "giai_thich": {"type": "string"}, "nguon": {"type": "string"},
        "tinh_huong": {"type": "string"}, "nguon_url": {"type": "string"},
        "nhan_dinh_meta": {"type": "array", "items": {"type": "object", "properties": {
            "noi_dung": {"type": "string"}, "yccd": {"type": "string"}, "muc_do": {"type": "string"},
            "thanh_phan_nang_luc": {"type": "string"}, "dap_an": {"type": "string"}, "giai_thich": {"type": "string"}
        }, "required": ["noi_dung", "yccd", "muc_do", "thanh_phan_nang_luc", "dap_an", "giai_thich"]}},
        "du_lieu_truc_quan": {"type": "object", "properties": {
            "loai": {"type": "string"}, "tieu_de": {"type": "string"},
            "cot": {"type": "array", "items": {"type": "string"}},
            "du_lieu": {"type": "array", "items": {"type": "array", "items": {"type": "string"}}},
            "mo_ta": {"type": "string"}, "nguon": {"type": "string"}
        }, "required": ["loai", "tieu_de", "cot", "du_lieu", "mo_ta", "nguon"]}
    }, "required": ["khoi", "chuong", "bai", "yccd", "muc_do", "dang_cau", "thanh_phan_nang_luc",
    "hanh_vi_nang_luc", "cau_hoi", "lua_chon", "dap_an", "giai_thich", "nguon", "tinh_huong",
    "nguon_url", "nhan_dinh_meta", "du_lieu_truc_quan"]}}}, "required": ["questions"]
}


def tao_cau_tot_nghiep_bang_ai(yccd_meta, dang_cau, muc_do, nang_luc, hanh_vi, kieu_truc_quan, nguon_truc_quan, noi_dung_hat_giong=""):
    seed_block = ""
    if str(noi_dung_hat_giong or "").strip():
        seed_block = f"""\nCÂU HỎI HẠT GIỐNG THAM KHẢO:\n{str(noi_dung_hat_giong)[:6000]}\nChỉ tham khảo kiến thức/cấu trúc, không sao chép nguyên văn nếu không cần.\n"""
    prompt = f"""
Bạn là chuyên gia xây dựng câu hỏi luyện thi tốt nghiệp THPT môn Sinh học theo định hướng đánh giá năng lực.
Khối 12; Chương: {yccd_meta.get('chuong','')}; Bài: {yccd_meta.get('bai','')};
YCCĐ kiến thức nền: {yccd_meta.get('yccd','')}; Dạng: {dang_cau}; Mức độ: {muc_do};
Năng lực chính: {nang_luc}; Hành vi năng lực: {hanh_vi}; Dữ liệu trực quan: {kieu_truc_quan}; Nguồn: {nguon_truc_quan}.
{seed_block}
QUY TẮC:
- Đánh giá việc huy động kiến thức để xử lí nhiệm vụ, không chỉ hỏi nhớ máy móc.
- Đúng/Sai: một dữ kiện chung + đúng 4 ý, mỗi ý có đáp án và giải thích.
- 4 lựa chọn: đúng 4 phương án, chỉ một đúng.
- Trả lời ngắn: đáp án số/kí hiệu ngắn.
- Dữ kiện phải đủ để giải, khoa học nhất quán.
- Không bịa nghiên cứu thật. Dữ liệu mô phỏng phải ghi 'Dữ liệu mô phỏng do app tạo'.
- bang_so_lieu / bieu_do_cot / bieu_do_duong: tạo cot + du_lieu để app tự dựng.
- hinh_tu_tai_lieu: không bịa hình, chỉ mô tả hình cần dùng và để dữ liệu rỗng.
"""
    response = goi_gemini_co_retry(prompt, GRAD_QUESTION_SCHEMA)
    data = json.loads(response.text)
    questions = data.get("questions", [])
    if not questions:
        return None
    q = questions[0]
    q["id"] = str(uuid.uuid4()); q["temp_id"] = str(uuid.uuid4())
    q["muc_dich_su_dung"] = "tot_nghiep"
    q["duoc_dung_luyen_hs"] = True
    q["nguon_tao"] = "AI tốt nghiệp"
    q["ngay_tao"] = datetime.now().strftime("%d/%m/%Y %H:%M")
    return chuan_hoa_cau_truc_cau_hoi(q)



def chuan_hoa_ten_nang_luc_seed(value):
    s = " ".join(str(value or "").strip().split()).casefold()
    if "nhận thức" in s:
        return "Nhận thức sinh học"
    if "tìm hiểu" in s or "thế giới sống" in s:
        return "Tìm hiểu thế giới sống"
    if "vận dụng" in s:
        return "Vận dụng kiến thức, kĩ năng đã học"
    return str(value or "").strip()


def chon_hat_giong_phu_hop(seed_bank, nang_luc="", muc_do="", dang_cau="", tu_khoa=""):
    nl_target = chuan_hoa_ten_nang_luc_seed(nang_luc)
    tokens = re.findall(r"[0-9A-Za-zÀ-ỹ]{4,}", str(tu_khoa or "").casefold())[:8]
    scored = []

    for seed in seed_bank or []:
        if not seed_duoc_phep_su_dung(seed):
            continue
        score = 0
        nl_seed = chuan_hoa_ten_nang_luc_seed(seed.get("thanh_phan_nang_luc", ""))
        meta_ds = list(seed.get("nhan_dinh_meta", []) or [])
        nl_ds = {
            chuan_hoa_ten_nang_luc_seed(x.get("thanh_phan_nang_luc", ""))
            for x in meta_ds if isinstance(x, dict)
        }
        md_ds = {
            str(x.get("muc_do", "")).strip()
            for x in meta_ds if isinstance(x, dict)
        }

        if nl_target and (nl_seed == nl_target or nl_target in nl_ds):
            score += 5

        if muc_do and (
            str(seed.get("muc_do", "")).strip() == str(muc_do).strip()
            or str(muc_do).strip() in md_ds
        ):
            score += 3

        if dang_cau and str(seed.get("dang_cau_goi_y", "")).strip() == str(dang_cau).strip():
            score += 3

        haystack = (
            str(seed.get("kien_thuc_chu_de", ""))
            + " "
            + str(seed.get("noi_dung_goc", ""))
        ).casefold()

        if tokens and any(t in haystack for t in tokens):
            score += 2

        if score > 0:
            scored.append((score, seed))

    if not scored:
        return None

    scored.sort(key=lambda x: -x[0])
    return scored[0][1]


def tao_khoi_hat_giong_cho_yccd_da_chon():
    seed_bank = doc_ngan_hang_hat_giong()
    if not seed_bank:
        return ""

    blocks = []
    used_ids = set()

    for item in st.session_state.get("yccd_da_chon", []):
        _, cfg = lay_cau_hinh(item)

        seed = chon_hat_giong_phu_hop(
            seed_bank,
            cfg.get("Thành phần năng lực", ""),
            cfg.get("Mức độ", ""),
            cfg.get("Dạng câu hỏi", ""),
            item.get("YCCĐ", "")
        )

        if not seed or seed.get("id") in used_ids:
            continue

        used_ids.add(seed.get("id"))

        blocks.append({
            "YCCĐ_mục_tiêu": item.get("YCCĐ", ""),
            "Mức_độ_mục_tiêu": cfg.get("Mức độ", ""),
            "Năng_lực_mục_tiêu": cfg.get("Thành phần năng lực", ""),
            "Dạng_câu": cfg.get("Dạng câu hỏi", ""),
            "Hạt_giống": str(seed.get("noi_dung_goc", ""))[:3500]
        })

        if len(blocks) >= 8:
            break

    if not blocks:
        return ""

    return (
        "\n\n============================================================\n"
        "NGÂN HÀNG HẠT GIỐNG + YÊU CẦU ĐA DẠNG HÓA\n"
        "============================================================\n"
        "Các câu dưới đây là HẠT GIỐNG để định hướng chất lượng, cấu trúc, "
        "năng lực và cách khai thác kiến thức; KHÔNG phải khuôn để sao chép.\n"
        "BẮT BUỘC đa dạng hóa: thay đổi dữ kiện, đối tượng, ngữ cảnh, cách hỏi, "
        "hướng suy luận, phương án nhiễu và kiểu xử lí thông tin khi phù hợp.\n"
        "Trong một nhóm câu, phải phối hợp cả: "
        "(1) câu phát triển từ hạt giống; "
        "(2) câu biến thể xa từ hạt giống; "
        "(3) câu mới độc lập nhưng vẫn đúng YCCĐ/mức độ/năng lực.\n"
        "Không được chỉ đổi tên, đổi số hoặc đảo phương án rồi coi là câu mới.\n"
        + json.dumps(blocks, ensure_ascii=False)
    )


def thong_ke_hat_giong_nang_luc_chi_bao(seed_bank):
    dem = {}
    for seed in seed_bank or []:
        meta_ds = list(seed.get("nhan_dinh_meta", []) or [])
        if meta_ds:
            # Đúng/Sai: mỗi ý là một nhiệm vụ đánh giá riêng, nên thống kê theo từng ý.
            for nd in meta_ds:
                if not isinstance(nd, dict):
                    continue
                nl = chuan_hoa_ten_nang_luc_seed(
                    nd.get("thanh_phan_nang_luc", "")
                ) or "Chưa xác định"
                cb = str(nd.get("chi_bao", "")).strip() or "Chưa xác định"
                dem[(nl, cb)] = dem.get((nl, cb), 0) + 1
            continue

        nl = chuan_hoa_ten_nang_luc_seed(
            seed.get("thanh_phan_nang_luc", "")
        ) or "Chưa xác định"
        cb = str(seed.get("chi_bao", "")).strip() or "Chưa xác định"
        dem[(nl, cb)] = dem.get((nl, cb), 0) + 1

    rows = [
        {
            "Thành phần năng lực": k[0],
            "Chỉ báo / kĩ năng": k[1],
            "Số câu/ý hạt giống": v
        }
        for k, v in dem.items()
    ]
    rows.sort(
        key=lambda x: (
            x["Thành phần năng lực"],
            x["Số câu/ý hạt giống"]
        )
    )
    return rows


def tom_tat_nang_luc_chi_bao_hoc_sinh(profile):
    nl_map = {}

    for s in profile.get("stats", {}).values():
        nl = str(s.get("nang_luc", "")).strip() or "Chưa xác định"
        cb_name = str(s.get("chi_bao", "")).strip()

        item = nl_map.setdefault(
            nl,
            {"so_lan": 0, "so_dung": 0, "chi_bao": {}}
        )

        sl = int(s.get("so_lan", 0) or 0)
        sd = int(s.get("so_dung", 0) or 0)
        item["so_lan"] += sl
        item["so_dung"] += sd

        # Câu cũ chưa có chỉ báo vẫn được tính vào thành phần năng lực,
        # nhưng không hiển thị "Chưa gắn chỉ báo" như một chỉ báo yếu.
        if cb_name:
            cb = item["chi_bao"].setdefault(
                cb_name,
                {"so_lan": 0, "so_dung": 0}
            )
            cb["so_lan"] += sl
            cb["so_dung"] += sd

    result = []

    for nl in THANH_PHAN_NANG_LUC:
        item = nl_map.get(nl, {"so_lan": 0, "so_dung": 0, "chi_bao": {}})
        so_lan = item["so_lan"]
        ti_le = item["so_dung"] / so_lan if so_lan else 0

        cb_rows = []
        for cb_name, cb in item["chi_bao"].items():
            sl = cb["so_lan"]
            tl = cb["so_dung"] / sl if sl else 0
            cb_rows.append(
                {"chi_bao": cb_name, "so_lan": sl, "ti_le_dung": tl}
            )

        cb_rows.sort(key=lambda x: (x["ti_le_dung"], -x["so_lan"]))

        if so_lan < 3:
            status = "Chưa đủ dữ liệu"
        elif ti_le < 0.50:
            status = "Cần hỗ trợ"
        elif ti_le < 0.65:
            status = "Đang củng cố"
        elif ti_le < 0.80:
            status = "Đạt"
        else:
            status = "Đạt tốt"

        result.append({
            "nang_luc": nl,
            "so_lan": so_lan,
            "ti_le_dung": ti_le,
            "trang_thai": status,
            "chi_bao_yeu": [
                x for x in cb_rows
                if x["so_lan"] >= 2 and x["ti_le_dung"] < 0.65
            ][:5]
        })

    return result


def lap_ke_hoach_ngan_hang_tot_nghiep(
    ds_yccd,
    so_mcq,
    so_ds,
    so_tln
):
    """
    GV chỉ chọn phạm vi + số lượng.
    App tự cân đối:
    - YCCĐ nền
    - mức độ
    - 3 thành phần năng lực
    - hành vi năng lực
    - nhu cầu dữ liệu trực quan
    """
    ds_yccd = list(
        ds_yccd or []
    )

    if not ds_yccd:
        return []

    # Trục chính là năng lực, không phải YCCĐ.
    nl_nt = "Nhận thức sinh học"
    nl_th = "Tìm hiểu thế giới sống"
    nl_vd = "Vận dụng kiến thức, kĩ năng đã học"

    # Tỉ trọng mục tiêu toàn ngân hàng.
    # Gần cân bằng nhưng vẫn dành đủ câu cho kiến thức nền.
    competency_cycle = [
        nl_nt,
        nl_th,
        nl_vd,
        nl_nt,
        nl_th,
        nl_vd,
        nl_nt,
        nl_vd,
        nl_th,
        nl_nt
    ]

    hanh_vi_theo_nl = {
        nl_nt: [
            "Nhận biết / tái hiện kiến thức sinh học",
            "Giải thích cơ chế hoặc hiện tượng sinh học",
            "Suy luận di truyền từ dữ kiện"
        ],
        nl_th: [
            "Đọc và khai thác bảng số liệu",
            "Phân tích biểu đồ / đồ thị",
            "Phân tích hình / sơ đồ sinh học",
            "Phân tích kết quả thí nghiệm",
            "Đánh giá giả thuyết / kết luận nghiên cứu"
        ],
        nl_vd: [
            "Tính toán và xử lí số liệu sinh học",
            "Vận dụng kiến thức vào tình huống thực tiễn",
            "Đề xuất / lựa chọn giải pháp sinh học",
            "Suy luận di truyền từ dữ kiện"
        ]
    }

    muc_do_theo_dang = {
        "Trắc nghiệm 4 lựa chọn": [
            "Nhận biết",
            "Thông hiểu",
            "Vận dụng",
            "Thông hiểu",
            "Vận dụng"
        ],
        "Đúng / Sai": [
            "Thông hiểu",
            "Vận dụng",
            "Vận dụng",
            "Thông hiểu"
        ],
        "Trả lời ngắn": [
            "Vận dụng",
            "Thông hiểu",
            "Vận dụng"
        ]
    }

    def chon_visual(
        dang,
        nang_luc,
        index
    ):
        # Không lạm dụng hình/bảng. Ưu tiên khi nhiệm vụ thực sự cần dữ liệu.
        if dang == "Đúng / Sai":
            options = [
                "bang_so_lieu",
                "bieu_do_duong",
                "khong_co",
                "bieu_do_cot"
            ]
            return options[
                index % len(options)
            ]

        if dang == "Trả lời ngắn":
            options = [
                "bang_so_lieu",
                "khong_co",
                "bieu_do_cot"
            ]
            return options[
                index % len(options)
            ]

        if nang_luc == nl_th:
            options = [
                "khong_co",
                "bang_so_lieu",
                "bieu_do_cot"
            ]
            return options[
                index % len(options)
            ]

        return "khong_co"

    danh_sach_dang = (
        ["Trắc nghiệm 4 lựa chọn"] * int(so_mcq)
        + ["Đúng / Sai"] * int(so_ds)
        + ["Trả lời ngắn"] * int(so_tln)
    )

    ke_hoach = []

    for i, dang in enumerate(
        danh_sach_dang
    ):
        yccd_meta = ds_yccd[
            i % len(ds_yccd)
        ]

        nang_luc = competency_cycle[
            i % len(competency_cycle)
        ]

        ds_hanh_vi = hanh_vi_theo_nl[
            nang_luc
        ]

        hanh_vi = ds_hanh_vi[
            (i // 2) % len(ds_hanh_vi)
        ]

        ds_muc_do = muc_do_theo_dang[
            dang
        ]

        muc_do = ds_muc_do[
            i % len(ds_muc_do)
        ]

        # Tránh mâu thuẫn: hành vi tái hiện không đi với Vận dụng.
        if (
            hanh_vi
            == "Nhận biết / tái hiện kiến thức sinh học"
            and muc_do == "Vận dụng"
        ):
            muc_do = "Thông hiểu"

        kieu_truc_quan = chon_visual(
            dang,
            nang_luc,
            i
        )

        ke_hoach.append({
            "stt": i + 1,
            "dang_cau": dang,
            "yccd_meta": yccd_meta,
            "muc_do": muc_do,
            "nang_luc": nang_luc,
            "hanh_vi": hanh_vi,
            "kieu_truc_quan": kieu_truc_quan
        })

    return ke_hoach


def tom_tat_ke_hoach_tot_nghiep(
    ke_hoach
):
    result = {
        "tong": len(
            ke_hoach
        ),
        "dang": {},
        "nang_luc": {},
        "muc_do": {},
        "co_truc_quan": 0,
        "yccd": set()
    }

    for item in ke_hoach:
        dang = item.get(
            "dang_cau",
            ""
        )
        nl = item.get(
            "nang_luc",
            ""
        )
        md = item.get(
            "muc_do",
            ""
        )

        result[
            "dang"
        ][dang] = (
            result[
                "dang"
            ].get(
                dang,
                0
            )
            + 1
        )

        result[
            "nang_luc"
        ][nl] = (
            result[
                "nang_luc"
            ].get(
                nl,
                0
            )
            + 1
        )

        result[
            "muc_do"
        ][md] = (
            result[
                "muc_do"
            ].get(
                md,
                0
            )
            + 1
        )

        if item.get(
            "kieu_truc_quan"
        ) != "khong_co":
            result[
                "co_truc_quan"
            ] += 1

        yccd = str(
            item.get(
                "yccd_meta",
                {}
            ).get(
                "yccd",
                ""
            )
        ).strip()

        if yccd:
            result[
                "yccd"
            ].add(
                yccd
            )

    result[
        "so_yccd"
    ] = len(
        result[
            "yccd"
        ]
    )

    return result


def chon_hat_giong_tu_dong(seed_bank, index, plan=None):
    """
    Kết hợp HẠT GIỐNG + TẠO MỚI để tránh ngân hàng nghèo nàn.

    Chu kỳ 5 câu:
    - 3 câu có hạt giống phù hợp để phát triển/biến thể.
    - 2 câu không truyền hạt giống, buộc AI tạo hướng mới độc lập.

    Như vậy hạt giống vẫn giữ vai trò neo chất lượng nhưng không khóa
    sự đa dạng của ngân hàng.
    """
    if not seed_bank:
        return ""

    plan = plan or {}

    # 40% câu chủ động tạo mới hoàn toàn.
    if index % 5 in (3, 4):
        return ""

    yccd_meta = plan.get("yccd_meta", {}) or {}

    item = chon_hat_giong_phu_hop(
        seed_bank,
        plan.get("nang_luc", ""),
        plan.get("muc_do", ""),
        plan.get("dang_cau", ""),
        str(yccd_meta.get("bai", ""))
        + " "
        + str(yccd_meta.get("yccd", ""))
    )

    if not item:
        return ""

    return str(item.get("noi_dung_goc", "")).strip()


def _seed_kiem_tra_cau_truc_truoc_khi_nhap(q):
    """Kiểm tra cục bộ, không AI. Trả về danh sách lỗi cấu trúc nghiêm trọng."""
    errs=[]
    dang=str(q.get("dang_cau", q.get("dang_cau_goi_y", "")) or "").strip()
    clean=str(q.get("noi_dung_hien_thi", "") or "").strip()
    if dang not in {"Trắc nghiệm 4 lựa chọn","Đúng / Sai","Trả lời ngắn"}:
        errs.append("chưa xác định dạng câu")
    if not _seed_co_dap_an_nguon(q):
        errs.append("không đọc được đáp án nguồn")
    if dang == "Trắc nghiệm 4 lựa chọn":
        labs=re.findall(r"(?m)^\s*([A-D])\s*[\.\)\:]\s+", clean)
        if len(set(labs)) < 4:
            errs.append("không đọc đủ 4 phương án A-D")
    elif dang == "Đúng / Sai":
        labs=re.findall(r"(?m)^\s*([a-d])\s*[\.\)\:]\s+", clean)
        if len(set(labs)) < 4:
            errs.append("không đọc đủ 4 nhận định a-d")
        ans=q.get("dap_an_nguon", []) or []
        if not isinstance(ans,list) or len(ans) != 4:
            errs.append("đáp án Đ/S không đủ 4 ý")
    elif dang == "Trả lời ngắn":
        if len(clean) < 12:
            errs.append("nội dung trả lời ngắn quá ngắn/thiếu dữ kiện")
    return errs


def _seed_kiem_tra_file_truoc_khi_nhap(ds_tach):
    loi=[]
    for q in ds_tach or []:
        es=_seed_kiem_tra_cau_truc_truoc_khi_nhap(q)
        if es:
            loi.append({"Câu": str(q.get("so_cau_goc", "?")), "Lỗi": "; ".join(es)})
    return loi



def _main_la_cau_tu_hat_giong(q):
    """Nhận diện câu trong ngân hàng chính có nguồn gốc trực tiếp từ hạt giống."""
    if not isinstance(q, dict):
        return False
    if str(q.get("nguon_seed_id", "") or "").strip():
        return True
    nguon_tao = str(q.get("nguon_tao", "") or "").strip().casefold()
    return "hạt giống" in nguon_tao or "hat giong" in nguon_tao


def _main_ten_nguon_hat_giong(q):
    ten = str(q.get("nguon_file", "") or q.get("nguon", "") or "").strip()
    return ten or "(Không rõ file nguồn)"


def _sao_luu_ngan_hang_chinh_truoc_khi_don(bank, ly_do="don_hat_giong"):
    """Tạo bản sao JSON trước khi xóa hàng loạt câu hạt giống khỏi ngân hàng chính."""
    thu_muc = os.path.join(BASE_DIR, "sao_luu_ngan_hang")
    os.makedirs(thu_muc, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    path = os.path.join(thu_muc, f"ngan_hang_cau_hoi_{ly_do}_{stamp}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(bank, f, ensure_ascii=False, indent=2)
    return path

def ngan_hang_hat_giong():
    st.header("🌱 NGÂN HÀNG HẠT GIỐNG")
    st.caption(
        "Kho nguồn đã được GV chuẩn hóa trước khi tải lên. App KHÔNG gọi AI để kiểm tra lại đáp án khi nhập; "
        "chỉ đọc metadata, chống trùng cục bộ và đồng bộ câu đủ YCCĐ + đáp án vào ngân hàng chính."
    )

    st.subheader("🌱 Nhập câu hỏi có sẵn từ Word / PDF / TXT")
    st.info(
        "Chỉ tải các file đã được kiểm tra đáp án, dữ kiện, YCCĐ, mức độ và năng lực trước. "
        "Khi nhập, app chỉ tách câu, đọc metadata và kiểm tra trùng bằng Python cục bộ — không dùng Gemini, không phát sinh lỗi 429 ở bước này."
    )

    try:
        st.download_button(
            "📄 TẢI FILE WORD MẪU NGÂN HÀNG HẠT GIỐNG",
            data=tao_file_mau_hat_giong_docx_bytes(),
            file_name="MAU_NGAN_HANG_HAT_GIONG.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="seed_download_template"
        )
    except Exception as e:
        st.caption(f"Chưa tạo được file mẫu Word: {e}")

    with st.expander("ℹ️ Các kiểu trình bày app nhận được", expanded=False):
        st.markdown(
            "- Mốc câu: `Câu 1.`, `Câu 1:`, `Câu 1)`, `Câu 1 -`, `Câu hỏi 1`, hoặc `1.` khi cả tài liệu đánh số liên tục.\n"
            "- Trắc nghiệm: A/B/C/D có thể **cùng một dòng hoặc khác dòng**.\n"
            "- Đúng/Sai: a/b/c/d có thể **cùng một dòng hoặc khác dòng**.\n"
            "- DOCX: đọc cả **paragraph và bảng Word**, đồng thời **giữ ảnh/sơ đồ/biểu đồ và cấu trúc bảng** gắn với câu.\n"
            "- PDF/TXT vẫn đọc nội dung chữ; nếu câu có ảnh/bảng cần giữ nguyên, nên dùng **DOCX**.\n"
            "- `Đáp án:` và `YCCĐ:` cần có nếu muốn câu được đưa thẳng vào ngân hàng chính. Mức độ, Thành phần năng lực, Chỉ báo, Kiến thức nên hoàn thiện trước khi tải lên. App **không gọi AI để kiểm tra hoặc gán lại** khi nhập."
        )

    files = st.file_uploader(
        "Chọn một hoặc nhiều file câu hỏi",
        type=["docx", "pdf", "txt"],
        accept_multiple_files=True,
        key="seed_upload_files"
    )

    if st.button("📥 TÁCH & NHẬP HẠT GIỐNG", type="primary", use_container_width=True, key="seed_import_btn"):
        if not files:
            st.warning("Hãy chọn ít nhất một file.")
        else:
            bank_seed = doc_ngan_hang_hat_giong()
            bank_main = doc_ngan_hang()
            fp_old = {_seed_fingerprint_noi_dung(x) for x in bank_seed if _seed_fingerprint_noi_dung(x)}
            them = trung_seed = trung_main = 0
            bao_cao = []
            for f in files:
                goi_nguon = doc_goi_tu_file_hat_giong(f)
                txt = goi_nguon.get("text", "")
                ds_tach = tach_cau_hoi_hat_giong(
                    txt,
                    f.name,
                    resources_map=goi_nguon.get("resources", {}),
                    source_path=goi_nguon.get("source_path", ""),
                    media_dir=goi_nguon.get("media_dir", "")
                )
                them_file = trung_seed_file = trung_main_file = thieu_yccd_file = thieu_dap_an_file = nghi_trung_file = 0

                # KIỂM TRA TRƯỚC KHI NHẬP: lỗi parser cấu trúc -> không nhập nửa chừng.
                loi_file = _seed_kiem_tra_file_truoc_khi_nhap(ds_tach)
                if loi_file:
                    bao_cao.append({
                        "File": f.name, "Phát hiện": len(ds_tach), "Nhập mới": 0,
                        "MCQ": sum(1 for x in ds_tach if x.get("dang_cau") == "Trắc nghiệm 4 lựa chọn"),
                        "Đ/S": sum(1 for x in ds_tach if x.get("dang_cau") == "Đúng / Sai"),
                        "TL ngắn": sum(1 for x in ds_tach if x.get("dang_cau") == "Trả lời ngắn"),
                        "Trùng hạt giống (chính xác)": 0, "Đã có ngân hàng": 0, "Nghi gần trùng": 0,
                        "Thiếu YCCĐ": sum(1 for x in ds_tach if not str(x.get("yccd", "") or "").strip()),
                        "Thiếu đáp án": sum(1 for x in ds_tach if not _seed_co_dap_an_nguon(x)),
                        "Ghi chú": "KHÔNG NHẬP - lỗi cấu trúc: " + ", ".join(f"Câu {x['Câu']}: {x['Lỗi']}" for x in loi_file[:8])
                                   + (f" ... (+{len(loi_file)-8} câu)" if len(loi_file) > 8 else "")
                    })
                    st.error(
                        f"File {f.name} chưa được nhập vì parser phát hiện {len(loi_file)} câu lỗi cấu trúc. "
                        "Không có câu nào của file này được ghi vào hạt giống. "
                        + " | ".join(f"Câu {x['Câu']}: {x['Lỗi']}" for x in loi_file[:6])
                    )
                    continue

                for q in ds_tach:
                    fp = _seed_fingerprint_noi_dung(q)
                    # Chỉ trùng CHÍNH XÁC mới bỏ bản sao trong hạt giống.
                    old_exact = None
                    if fp:
                        for old_seed in bank_seed:
                            if _seed_fingerprint_noi_dung(old_seed) == fp:
                                old_exact = old_seed; break
                    if old_exact is not None:
                        # NHẬP LẠI FILE CHUẨN: làm mới dữ liệu parser của bản hạt giống cũ,
                        # nhưng GIỮ NGUYÊN id để không mất liên kết với Ngân hàng câu hỏi.
                        #
                        # Trước đây chỉ điền trường còn trống. Vì vậy nếu bản cũ đã bị parser
                        # nhận sai dạng/đáp án/metadata nhưng trường đó không rỗng, nhập lại file
                        # vẫn không sửa được và câu tiếp tục bị giữ ở hạt giống.
                        #
                        # Từ bản này: nội dung vừa đọc từ file là nguồn chuẩn mới nhất cho các
                        # trường do parser tạo ra.
                        refresh_fields = [
                            "nguon_file", "so_cau_goc",
                            "noi_dung_goc", "noi_dung_hien_thi",
                            "dang_cau_goi_y", "dang_cau",
                            "yccd", "muc_do", "thanh_phan_nang_luc",
                            "chi_bao", "hanh_vi_nang_luc", "kien_thuc_chu_de",
                            "dap_an_nguon", "giai_thich_nguon", "nguon_giai_thich",
                            "nhan_dinh_meta",
                            "kiem_tra_dap_an", "trang_thai_kiem_tra_dap_an",
                            "gv_da_duyet_dap_an", "duoc_dung_lam_hat_giong",
                            "nguon_da_kiem_tra_truoc"
                        ]

                        for fld in refresh_fields:
                            if fld in q:
                                old_exact[fld] = q.get(fld)

                        # Ảnh/bảng chỉ thay khi lần đọc mới thực sự có dữ liệu,
                        # tránh vô tình làm mất tài nguyên trực quan của bản cũ.
                        for fld in [
                            "tai_nguyen_truc_quan", "du_lieu_truc_quan",
                            "nguon_file_path", "thu_muc_media"
                        ]:
                            val = q.get(fld)
                            if val not in (None, "", [], {}):
                                old_exact[fld] = val

                        old_exact["ngay_cap_nhat_tu_file"] = datetime.now().strftime(
                            "%d/%m/%Y %H:%M"
                        )

                        # Không xóa id/nguon_seed_id liên kết cũ.
                        # Ngay sau vòng nhập, hàm đồng bộ sẽ quét lại toàn bộ seed bank;
                        # câu đủ YCCĐ + đáp án sẽ được đưa/cập nhật vào ngân hàng chính.
                        trung_seed += 1; trung_seed_file += 1
                        continue

                    # Gần trùng chỉ cảnh báo, không tự loại.
                    near_seed, score_seed = _seed_trung_voi_hat_giong(q, bank_seed, nguong=0.92)
                    if near_seed and score_seed < 1.0:
                        q["nghi_trung_hat_giong"] = True
                        q["do_giong_hat_giong"] = score_seed
                        nghi_trung_file += 1

                    q["kiem_tra_dap_an"] = "Nguồn đã chuẩn hóa trước khi nhập"
                    q["trang_thai_kiem_tra_dap_an"] = "Không kiểm tra AI"
                    q["gv_da_duyet_dap_an"] = True
                    q["duoc_dung_lam_hat_giong"] = True
                    q["nguon_da_kiem_tra_truoc"] = True
                    if not str(q.get("yccd", "") or "").strip(): thieu_yccd_file += 1
                    if not _seed_co_dap_an_nguon(q): thieu_dap_an_file += 1

                    # Câu đã có trong ngân hàng chính vẫn được giữ ở hạt giống; chỉ ghi nhận/link, không bỏ nguồn.
                    ok_sync, _ = _seed_du_dieu_kien_dong_bo(q)
                    if ok_sync:
                        q_bank = _seed_chuyen_thanh_cau_ngan_hang(q)
                        if q_bank:
                            old_main, _ = _seed_tim_trung_chinh_xac_ngan_hang(q_bank, bank_main)
                            if old_main is not None:
                                trung_main += 1; trung_main_file += 1
                                q["id_cau_trung_ngan_hang"] = str(old_main.get("id", ""))
                    bank_seed.append(q)
                    if fp: fp_old.add(fp)
                    them += 1; them_file += 1
                loai_counts = {k: sum(1 for x in ds_tach if x.get("dang_cau") == k) for k in ["Trắc nghiệm 4 lựa chọn","Đúng / Sai","Trả lời ngắn","Chưa xác định"]}
                bao_cao.append({
                    "File": f.name, "Phát hiện": len(ds_tach), "Nhập mới": them_file,
                    "MCQ": loai_counts["Trắc nghiệm 4 lựa chọn"], "Đ/S": loai_counts["Đúng / Sai"], "TL ngắn": loai_counts["Trả lời ngắn"],
                    "Trùng hạt giống (chính xác)": trung_seed_file, "Đã có ngân hàng": trung_main_file,
                    "Nghi gần trùng": nghi_trung_file,
                    "Thiếu YCCĐ": thieu_yccd_file, "Thiếu đáp án": thieu_dap_an_file,
                    "Ghi chú": ("Không tìm thấy mốc câu phù hợp" if txt.strip() and not ds_tach else ("Không đọc được nội dung" if not txt.strip() else ""))
                })
            luu_ngan_hang_hat_giong(bank_seed)
            da_chuyen, bo_qua_sync = dong_bo_hat_giong_an_toan_sang_ngan_hang(bank_seed)
            st.success(
                f"Đã nhập {them} câu vào hạt giống; {trung_seed} câu trùng chính xác đã được gộp/cập nhật "
                f"(đáp án, hướng dẫn giải, ảnh/bảng và metadata nếu file mới có); "
                f"{trung_main} câu đã có trong ngân hàng chính vẫn được giữ nguồn. "
                f"📚 Đã đồng bộ thêm {da_chuyen} câu mới đủ YCCĐ + đáp án vào ngân hàng chính. Không gọi Gemini."
            )
            if bo_qua_sync:
                st.caption(f"Có {bo_qua_sync} câu đang được giữ ở hạt giống nhưng chưa đồng bộ hoặc bị bỏ qua khi đồng bộ (thiếu YCCĐ/đáp án hoặc trùng).")
            if bao_cao:
                st.dataframe(pd.DataFrame(bao_cao), use_container_width=True, hide_index=True)

    bank_seed = doc_ngan_hang_hat_giong()
    st.metric("Số câu hạt giống hiện có", len(bank_seed))

    # ======================================================
    # DỌN CÁC CÂU ĐÃ ĐỒNG BỘ SANG NGÂN HÀNG CHÍNH
    # Kể cả khi file/câu tương ứng đã bị xóa khỏi Ngân hàng hạt giống.
    # Dựa trên dấu vết bền vững: nguon_seed_id / nguon_tao / nguon_file.
    # ======================================================
    bank_main_hien_tai = doc_ngan_hang()
    ds_main_tu_seed = [
        q for q in bank_main_hien_tai
        if _main_la_cau_tu_hat_giong(q)
    ]

    if ds_main_tu_seed:
        with st.expander(
            "🧹 Dọn câu hạt giống đã vào Ngân hàng câu hỏi",
            expanded=False
        ):
            st.warning(
                "Các câu này vẫn nhận diện được ngay cả khi file hạt giống đã bị xóa khỏi app, "
                "vì bản trong Ngân hàng câu hỏi còn lưu mã hạt giống và tên file nguồn. "
                "Chức năng dưới đây KHÔNG đụng đến Ngân hàng tốt nghiệp."
            )

            dem_main_theo_file = {}
            for q in ds_main_tu_seed:
                ten = _main_ten_nguon_hat_giong(q)
                dem_main_theo_file[ten] = dem_main_theo_file.get(ten, 0) + 1

            rows_main_seed = []
            for ten in sorted(dem_main_theo_file):
                ds_file = [
                    q for q in ds_main_tu_seed
                    if _main_ten_nguon_hat_giong(q) == ten
                ]
                co_truc_quan = sum(
                    1 for q in ds_file
                    if (q.get("tai_nguyen_truc_quan") or q.get("du_lieu_truc_quan"))
                )
                rows_main_seed.append({
                    "File nguồn còn lưu trong NH chung": ten,
                    "Số câu đã đồng bộ": len(ds_file),
                    "Có ảnh/bảng đang lưu": co_truc_quan,
                    "Không có dữ liệu trực quan": len(ds_file) - co_truc_quan,
                })

            st.dataframe(
                pd.DataFrame(rows_main_seed),
                use_container_width=True,
                hide_index=True,
                height=min(360, 70 + 35 * len(rows_main_seed))
            )

            st.caption(
                "Nếu một file cũ từng bị mất ảnh/đồ thị, hãy chọn đúng tên file đó và xóa toàn bộ "
                "các bản đã đồng bộ; sau đó nhập lại DOCX gốc bằng parser mới."
            )

            files_main_xoa = st.multiselect(
                "Chọn file nguồn có các câu cần xóa khỏi Ngân hàng câu hỏi",
                sorted(dem_main_theo_file),
                key="seed_cleanup_main_sources"
            )

            if files_main_xoa:
                tap_main_xoa = set(files_main_xoa)
                ds_se_xoa_main = [
                    q for q in ds_main_tu_seed
                    if _main_ten_nguon_hat_giong(q) in tap_main_xoa
                ]
                st.error(
                    f"Sẽ xóa **{len(ds_se_xoa_main)} câu** đã đồng bộ từ "
                    f"**{len(files_main_xoa)} file hạt giống** khỏi Ngân hàng câu hỏi."
                )
                xac_nhan_main = st.checkbox(
                    "Tôi xác nhận xóa các câu này khỏi Ngân hàng câu hỏi. App sẽ tự sao lưu trước khi xóa.",
                    key="seed_cleanup_main_confirm"
                )

                if st.button(
                    "🗑️ XÓA CÁC CÂU ĐÃ ĐỒNG BỘ TỪ FILE ĐÃ CHỌN",
                    type="primary",
                    use_container_width=True,
                    key="seed_cleanup_main_btn",
                    disabled=not xac_nhan_main
                ):
                    ids_xoa = {str(q.get("id", "")) for q in ds_se_xoa_main}
                    backup_path = _sao_luu_ngan_hang_chinh_truoc_khi_don(
                        bank_main_hien_tai,
                        "truoc_xoa_cau_hat_giong"
                    )
                    bank_main_moi = [
                        q for q in bank_main_hien_tai
                        if str(q.get("id", "")) not in ids_xoa
                    ]
                    luu_ngan_hang(bank_main_moi)
                    st.success(
                        f"Đã xóa {len(ids_xoa)} câu hạt giống khỏi Ngân hàng câu hỏi. "
                        f"Đã sao lưu trước khi xóa tại: {os.path.basename(backup_path)}"
                    )
                    st.rerun()

    if bank_seed:
        with st.expander("🛠️ Đọc lại metadata nguồn hiện có", expanded=False):
            st.caption("Chỉ đọc lại metadata bằng parser cục bộ; không kiểm tra đáp án bằng AI và không gọi Gemini.")
            if st.button("🔄 ĐỌC LẠI METADATA HIỆN CÓ", use_container_width=True, key="seed_reparse_existing"):
                bank_seed, nfix = doc_lai_metadata_hat_giong_hien_co(bank_seed)
                for q in bank_seed:
                    q["kiem_tra_dap_an"] = "Nguồn đã chuẩn hóa trước khi nhập"
                    q["trang_thai_kiem_tra_dap_an"] = "Không kiểm tra AI"
                    q["gv_da_duyet_dap_an"] = True
                    q["duoc_dung_lam_hat_giong"] = True
                    q["nguon_da_kiem_tra_truoc"] = True
                luu_ngan_hang_hat_giong(bank_seed)
                da_chuyen, _ = dong_bo_hat_giong_an_toan_sang_ngan_hang(bank_seed)
                st.success(f"Đã đọc lại {len(bank_seed)} câu; cập nhật/làm sạch {nfix} câu và đồng bộ thêm {da_chuyen} câu không trùng.")
                st.rerun()

        san_sang = thieu_yccd = thieu_dap_an = 0
        for q in bank_seed:
            ok, ly_do = _seed_du_dieu_kien_dong_bo(q)
            if ok: san_sang += 1
            elif ly_do == "Thiếu YCCĐ": thieu_yccd += 1
            elif ly_do == "Thiếu đáp án nguồn": thieu_dap_an += 1
        c1, c2, c3 = st.columns(3)
        c1.metric("✅ Đủ metadata để dùng", san_sang)
        c2.metric("⚠️ Thiếu YCCĐ", thieu_yccd)
        c3.metric("⚠️ Thiếu đáp án", thieu_dap_an)
        if thieu_yccd or thieu_dap_an:
            st.info("Các câu thiếu YCCĐ/đáp án vẫn được giữ trong Ngân hàng hạt giống nhưng chưa tự chuyển sang Ngân hàng câu hỏi. Nên hoàn thiện file nguồn bên ngoài app rồi nhập lại; app không dùng AI để đoán bổ sung.")

    # ======================================================
    # QUẢN LÝ / XÓA MỘT HOẶC NHIỀU FILE NGUỒN HẠT GIỐNG
    # ======================================================
    if bank_seed:
        nguon_seed_all = sorted({
            str(x.get("nguon_file", "")).strip()
            for x in bank_seed
            if str(x.get("nguon_file", "")).strip()
        })

        if nguon_seed_all:
            with st.expander("🗑️ Quản lý / xóa file nguồn hạt giống", expanded=False):
                st.caption(
                    "Chọn một hoặc nhiều file nguồn để xóa toàn bộ câu hạt giống "
                    "được nhập từ các file đó. Thao tác này chỉ ảnh hưởng Ngân hàng hạt giống."
                )

                dem_theo_file = {
                    ten: sum(
                        1 for x in bank_seed
                        if str(x.get("nguon_file", "")).strip() == ten
                    )
                    for ten in nguon_seed_all
                }

                st.dataframe(
                    pd.DataFrame([
                        {"File nguồn": ten, "Số câu hạt giống": dem_theo_file.get(ten, 0)}
                        for ten in nguon_seed_all
                    ]),
                    use_container_width=True,
                    hide_index=True,
                    height=min(320, 42 + 35 * len(nguon_seed_all))
                )

                files_xoa = st.multiselect(
                    "Chọn file muốn xóa",
                    nguon_seed_all,
                    key="seed_delete_source_files"
                )

                so_cau_se_xoa = sum(
                    dem_theo_file.get(ten, 0)
                    for ten in files_xoa
                )

                if files_xoa:
                    st.warning(
                        f"Đã chọn **{len(files_xoa)} file**; sẽ xóa "
                        f"**{so_cau_se_xoa} câu hạt giống** thuộc các file này."
                    )

                    xac_nhan_xoa = st.checkbox(
                        "Tôi xác nhận muốn xóa các câu hạt giống thuộc những file đã chọn.",
                        key="seed_delete_confirm"
                    )

                    if st.button(
                        "🗑️ XÓA CÁC FILE ĐÃ CHỌN",
                        type="primary",
                        use_container_width=True,
                        key="seed_delete_files_btn",
                        disabled=not xac_nhan_xoa
                    ):
                        tap_xoa = set(files_xoa)
                        seed_bi_xoa = [
                            x for x in bank_seed
                            if str(x.get("nguon_file", "")).strip() in tap_xoa
                        ]
                        bank_moi = [
                            x for x in bank_seed
                            if str(x.get("nguon_file", "")).strip() not in tap_xoa
                        ]
                        # Chỉ dọn media không còn được ngân hàng chính tham chiếu.
                        _seed_don_media_khi_xoa_nguon(seed_bi_xoa)
                        luu_ngan_hang_hat_giong(bank_moi)
                        st.success(
                            f"Đã xóa {so_cau_se_xoa} câu hạt giống từ "
                            f"{len(files_xoa)} file nguồn."
                        )
                        st.rerun()

    if bank_seed:
        with st.expander("🧭 Độ phủ & danh sách câu hạt giống", expanded=False):
            st.markdown("### 🧭 Độ phủ hạt giống theo năng lực – chỉ báo")
            rows_seed_cov = thong_ke_hat_giong_nang_luc_chi_bao(bank_seed)
    
            if rows_seed_cov:
                st.dataframe(
                    pd.DataFrame(rows_seed_cov),
                    use_container_width=True,
                    hide_index=True,
                    height=300
                )
    
            chua_xd = sum(
                1 for x in bank_seed
                if not str(x.get("thanh_phan_nang_luc", "")).strip()
            )
    
            if chua_xd:
                st.warning(
                    f"Còn {chua_xd} câu chưa xác định được năng lực. "
                    "App vẫn giữ lại; nên hoàn thiện metadata ở file nguồn rồi nhập lại khi cần."
                )
    
            st.info(
                "Nguyên tắc: **hạt giống + tạo mới song song**. "
                "App dùng hạt giống làm mẫu định hướng, đồng thời tạo thêm câu mới "
                "và ưu tiên bù các năng lực/chỉ báo còn thiếu."
            )
            nguon_seed = sorted({str(x.get("nguon_file","")) for x in bank_seed if x.get("nguon_file")})
            nguon_chon = st.selectbox("Lọc theo file nguồn", ["Tất cả"] + nguon_seed, key="seed_source_filter")
            ds_seed_hien = [x for x in bank_seed if nguon_chon == "Tất cả" or x.get("nguon_file") == nguon_chon]
            st.dataframe(pd.DataFrame([{"Câu gốc":x.get("so_cau_goc",""),"Dạng gợi ý":x.get("dang_cau_goi_y",""),"Nguồn":x.get("nguon_file",""),"Trạng thái nguồn":x.get("kiem_tra_dap_an","Nguồn đã chuẩn hóa"),"Trạng thái":x.get("trang_thai","")} for x in ds_seed_hien[:300]]), use_container_width=True, hide_index=True, height=360)
            with st.expander("👁️ Xem nội dung một câu hạt giống", expanded=False):
                labels=[f"{i+1}. Câu {x.get('so_cau_goc','')} • {x.get('dang_cau_goi_y','')}" for i,x in enumerate(ds_seed_hien)]
                if labels:
                    label = st.selectbox(
                        "Chọn câu",
                        labels,
                        key="seed_preview_select"
                    )
                    idx = labels.index(label)
                    seed_xem = ds_seed_hien[idx]
                    st.write(
                        seed_xem.get("noi_dung_hien_thi", "")
                        or seed_xem.get("noi_dung_goc", "")
                    )
                    _seed_hien_thi_tai_nguyen(seed_xem)
    
    
    
# ==========================================================
# NGÂN HÀNG TỐT NGHIỆP TỪ ĐỀ THẬT / ĐỀ THI THỬ
# ==========================================================
def doc_ngan_hang_tot_nghiep_thuc_te():
    """
    Đọc ngân hàng đề thật và tự nâng cấp dữ liệu cũ.

    Câu đã có đáp án nguồn rõ ràng được phép dùng ngay để tạo đề/luyện tập,
    dù vẫn còn trạng thái "Chờ rà soát". AI rà soát chỉ là lớp kiểm tra bổ sung;
    nếu AI phát hiện nghi vấn thì câu mới bị chuyển sang "Cần GV xem".
    """
    ds = doc_json_list(GRAD_REAL_BANK_PATH)
    if not ds:
        return []

    changed = False
    upgraded = []
    for q in ds:
        q2 = _grad_chuan_hoa_cau_da_nhap(q)
        if q2 != q:
            changed = True
        upgraded.append(q2)

    if changed:
        try:
            luu_json_list(GRAD_REAL_BANK_PATH, upgraded)
        except Exception:
            pass

    return upgraded


def luu_ngan_hang_tot_nghiep_thuc_te(ds):
    luu_json_list(GRAD_REAL_BANK_PATH, ds)


def _grad_norm_text(value):
    s = unicodedata.normalize("NFKC", str(value or ""))
    s = " ".join(s.split()).casefold()
    return s




def _grad_co_dap_an_day_du(q):
    """Kiểm tra câu đề thật đã có đáp án nguồn đủ để dùng hay chưa."""
    dang = str(q.get("dang_cau", "")).strip()
    if dang == "Đúng / Sai":
        meta = list(q.get("nhan_dinh_meta", []) or [])
        return len(meta) == 4 and all(
            str(x.get("dap_an", "")).strip() in {"Đúng", "Sai"}
            for x in meta[:4]
        )
    if dang == "Trắc nghiệm 4 lựa chọn":
        return (
            str(q.get("dap_an", "")).strip().upper() in {"A", "B", "C", "D"}
            and len(list(q.get("lua_chon", []) or [])) >= 4
        )
    if dang == "Trả lời ngắn":
        return bool(str(q.get("dap_an", "")).strip())
    return False


def cau_tot_nghiep_du_dieu_kien_su_dung(q):
    """
    Câu đề thật được dùng ngay nếu đã có đáp án nguồn đầy đủ.
    Không bắt buộc phải chờ AI rà soát; nhưng câu đã bị AI/GV cảnh báo thì khóa.
    """
    if not isinstance(q, dict):
        return False
    if q.get("trang_thai") in {"Thiếu đáp án", "Cần GV xem", "Ngừng sử dụng"}:
        return False
    if not q.get("duoc_dung_luyen_hs", False):
        return False
    return _grad_co_dap_an_day_du(q)


def _grad_noi_dung_phan_loai(q, nd=None):
    parts = [
        q.get("tinh_huong", ""),
        q.get("cau_hoi", ""),
        " ".join(str(x or "") for x in (q.get("lua_chon", []) or [])),
    ]
    if nd is not None:
        parts.append(nd.get("noi_dung", ""))
    else:
        parts.extend(
            str(x.get("noi_dung", ""))
            for x in (q.get("nhan_dinh_meta", []) or [])
        )
    return " ".join(str(x or "") for x in parts if str(x or "").strip())


def _grad_tokens(value):
    raw = _bo_dau_nguon(value)
    words = re.findall(r"[a-z0-9]+", raw)
    stop = {
        "cau", "hoi", "sau", "day", "dua", "vao", "thong", "tin", "tren",
        "cho", "biet", "xac", "dinh", "nhan", "dinh", "phat", "bieu", "dung",
        "sai", "nao", "nhat", "mot", "cac", "cua", "la", "va", "trong", "voi",
        "duoc", "co", "the", "khi", "neu", "tu", "den", "theo", "sinh", "hoc",
        "phan", "bai", "chuong", "noi", "dung", "kien", "thuc", "yeu", "cau",
        "can", "dat", "hay", "chon", "dap", "an", "loai", "trinh", "bay",
    }
    return {
        w for w in words
        if len(w) >= 3 and w not in stop
    }


def _grad_ds_pham_vi_toan_bo():
    """
    Lấy toàn bộ Khối → Chương → Bài → YCCĐ trong kho chương trình để
    đối chiếu câu đề thật. Không mặc định câu tốt nghiệp là kiến thức Khối 12:
    một câu có thể thuộc Khối 10, 11 hoặc 12 tùy nội dung thực tế.

    Hàm chỉ dùng dữ liệu KHO_YCCD hiện có và chạy cục bộ, không gọi API.
    """
    out = []

    for khoi, root in (KHO_YCCD or {}).items():
        if not isinstance(root, dict):
            continue

        for chuong, lessons in root.items():
            if not isinstance(lessons, dict):
                continue

            for bai, yccds in lessons.items():
                yccd_texts = []

                if isinstance(yccds, dict):
                    for _, val in yccds.items():
                        if isinstance(val, list):
                            for x in val:
                                if isinstance(x, dict):
                                    yccd_texts.append(str(
                                        x.get("noi_dung", x.get("YCCĐ", x.get("yccd", "")))
                                    ))
                                else:
                                    yccd_texts.append(str(x))
                        elif isinstance(val, dict):
                            yccd_texts.append(str(
                                val.get("noi_dung", val.get("YCCĐ", val.get("yccd", "")))
                            ))
                        else:
                            yccd_texts.append(str(val))

                elif isinstance(yccds, list):
                    for yc in yccds:
                        if isinstance(yc, dict):
                            yccd_texts.append(str(
                                yc.get("noi_dung", yc.get("YCCĐ", yc.get("yccd", "")))
                            ))
                        else:
                            yccd_texts.append(str(yc))
                else:
                    yccd_texts = [str(yccds or "")]

                candidate_text = " ".join([
                    str(khoi or ""),
                    str(chuong or ""),
                    str(bai or ""),
                    " ".join(x for x in yccd_texts if str(x).strip()),
                ])

                out.append({
                    "khoi": str(khoi or "").strip(),
                    "chuong": str(chuong or "").strip(),
                    "bai": str(bai or "").strip(),
                    "text": candidate_text,
                    "tokens": _grad_tokens(candidate_text),
                    "bai_tokens": _grad_tokens(bai),
                    "chuong_tokens": _grad_tokens(chuong),
                    "yccd_tokens": _grad_tokens(" ".join(yccd_texts)),
                })

    return out


# Giữ tên cũ để không làm hỏng dữ liệu/code đã tham chiếu ở bản trước.
def _grad_ds_pham_vi_khoi_12():
    return _grad_ds_pham_vi_toan_bo()


def _grad_doc_ghi_nho_pham_vi():
    """Các ví dụ Khối → Chương → Bài do GV đã xác nhận."""
    try:
        data = doc_json_list(GRAD_SCOPE_MEMORY_PATH)
        return [x for x in data if isinstance(x, dict)]
    except Exception:
        return []


def _grad_luu_ghi_nho_pham_vi(data):
    try:
        luu_json_list(GRAD_SCOPE_MEMORY_PATH, data)
    except Exception:
        pass


def _grad_ghi_nho_pham_vi_gv(q, khoi, chuong, bai):
    """Học từ lần GV sửa để các câu tương tự sau này được ưu tiên đúng phạm vi."""
    text = _grad_noi_dung_phan_loai(q)
    norm = _grad_norm_text(text)
    if not norm or not str(khoi or '').strip() or not str(chuong or '').strip():
        return
    fp = hashlib.sha1(norm.encode('utf-8', errors='ignore')).hexdigest()
    mem = _grad_doc_ghi_nho_pham_vi()
    item = {
        'fingerprint': fp,
        'text': text[:6000],
        'tokens': sorted(_grad_tokens(text)),
        'khoi': str(khoi or '').strip(),
        'chuong': str(chuong or '').strip(),
        'bai': str(bai or '').strip(),
        'ngay_xac_nhan': datetime.now().strftime('%d/%m/%Y %H:%M'),
    }
    found = False
    for i, old in enumerate(mem):
        if old.get('fingerprint') == fp:
            mem[i] = item
            found = True
            break
    if not found:
        mem.append(item)
    # Giữ kho học gọn, các ví dụ mới nhất ở cuối.
    _grad_luu_ghi_nho_pham_vi(mem[-1000:])


def _grad_do_giong_bo_tu(tokens_a, tokens_b):
    a, b = set(tokens_a or []), set(tokens_b or [])
    if not a or not b:
        return 0.0
    return len(a & b) / max(1, len(a | b))


def _grad_xep_hang_ung_vien_pham_vi(q, nd=None, top_n=5):
    """Tầng 1: tìm tối đa 5 Bài phù hợp nhất trong KHO_YCCD, có học từ GV."""
    text = _grad_noi_dung_phan_loai(q, nd=nd)
    q_tokens = _grad_tokens(text)
    candidates = _grad_ds_pham_vi_toan_bo()
    if not q_tokens or not candidates:
        return []

    text_norm = _bo_dau_nguon(text)
    memories = _grad_doc_ghi_nho_pham_vi()
    scored = []

    for c in candidates:
        all_inter = q_tokens & c['tokens']
        bai_inter = q_tokens & c['bai_tokens']
        chuong_inter = q_tokens & c['chuong_tokens']
        yccd_inter = q_tokens & c.get('yccd_tokens', set())
        score = (
            1.0 * len(all_inter)
            + 3.2 * len(bai_inter)
            + 2.0 * len(chuong_inter)
            + 1.35 * len(yccd_inter)
        )

        bai_words_all = re.findall(r'[a-z0-9]+', _bo_dau_nguon(c['bai']))
        chuong_words_all = re.findall(r'[a-z0-9]+', _bo_dau_nguon(c['chuong']))
        bai_words = [w for w in bai_words_all if w in c['bai_tokens'] and len(w) >= 4]
        chuong_words = [w for w in chuong_words_all if w in c['chuong_tokens'] and len(w) >= 4]
        if len(bai_words) >= 2 and ' '.join(bai_words) in text_norm:
            score += 7.0
        if len(chuong_words) >= 2 and ' '.join(chuong_words) in text_norm:
            score += 3.5

        distinctive = {
            w for w in all_inter
            if len(w) >= 5 and w not in {
                'protein', 'enzyme', 'tebao', 'sinhhoc', 'cautruc',
                'hoatdong', 'quatrinh', 'ketqua', 'dulieu'
            }
        }
        score += 0.8 * len(distinctive)

        # Học từ ví dụ GV đã sửa: chỉ thưởng cho đúng phạm vi của ví dụ gần giống.
        best_mem = 0.0
        for m in memories[-500:]:
            if (
                str(m.get('khoi', '')) == str(c.get('khoi', ''))
                and str(m.get('chuong', '')) == str(c.get('chuong', ''))
                and (
                    not str(m.get('bai', '')).strip()
                    or str(m.get('bai', '')) == str(c.get('bai', ''))
                )
            ):
                sim = _grad_do_giong_bo_tu(q_tokens, m.get('tokens') or _grad_tokens(m.get('text', '')))
                best_mem = max(best_mem, sim)
        score += 14.0 * best_mem

        scored.append({
            'score': round(float(score), 4),
            'memory_similarity': round(float(best_mem), 3),
            'khoi': c.get('khoi', ''),
            'chuong': c.get('chuong', ''),
            'bai': c.get('bai', ''),
        })

    scored.sort(key=lambda x: x['score'], reverse=True)
    return scored[:max(1, int(top_n))]


GRAD_SCOPE_SCHEMA = {
    'type': 'object',
    'properties': {
        'choice': {'type': 'integer'},
        'confidence': {'type': 'integer'},
        'reason': {'type': 'string'},
    },
    'required': ['choice', 'confidence', 'reason'],
}


def _grad_ai_chon_pham_vi(text, ung_vien):
    """Tầng 2: Gemini chỉ được CHỌN trong các ứng viên từ KHO_YCCD, không tự bịa."""
    if not ung_vien:
        return None
    ds = [
        {
            'choice': i + 1,
            'khoi': c.get('khoi', ''),
            'chuong': c.get('chuong', ''),
            'bai': c.get('bai', ''),
        }
        for i, c in enumerate(ung_vien)
    ]
    prompt = f"""
Bạn là chuyên gia chương trình Sinh học THPT.
Hãy phân loại CÂU HỎI vào đúng Khối → Chương → Bài.

QUY TẮC CỨNG:
- Chỉ được chọn MỘT trong các ứng viên đã cung cấp.
- Không được sáng tác tên Khối/Chương/Bài mới.
- Đọc toàn bộ câu hỏi, tình huống, phương án, đáp án/lời giải nếu có.
- Chọn theo KIẾN THỨC cốt lõi thực sự cần để giải câu, không chọn theo vài từ khóa bề mặt.
- confidence là số nguyên 0–100.

CÂU HỎI:
{text[:9000]}

ỨNG VIÊN:
{json.dumps(ds, ensure_ascii=False)}
"""
    try:
        response = goi_gemini_co_retry(prompt, GRAD_SCOPE_SCHEMA, so_lan_thu=2)
        data = json.loads(response.text)
        idx = int(data.get('choice', 0)) - 1
        if idx < 0 or idx >= len(ung_vien):
            return None
        c = dict(ung_vien[idx])
        c['ai_confidence'] = max(0, min(100, int(data.get('confidence', 0))))
        c['ai_reason'] = str(data.get('reason', '')).strip()
        return c
    except Exception:
        return None


def xac_dinh_pham_vi_bai_chuong_tot_nghiep(q, nd=None, dung_ai=False):
    """
    Phân loại Khối → Chương → Bài theo 2 tầng:
    1) KHO_YCCD + các ví dụ GV đã xác nhận chọn 5 ứng viên tốt nhất.
    2) Khi dung_ai=True, Gemini chỉ chọn trong 5 ứng viên đó.

    YCCĐ chi tiết không bị ép gán cho câu đề thật.
    """
    text = _grad_noi_dung_phan_loai(q, nd=nd)
    ranked = _grad_xep_hang_ung_vien_pham_vi(q, nd=nd, top_n=5)
    if not ranked:
        return {'khoi': '', 'chuong': '', 'bai': '', 'do_tin_cay': 0.0, 'nguon_phan_loai': 'khong_xac_dinh'}

    best = ranked[0]
    second_score = ranked[1]['score'] if len(ranked) > 1 else 0.0
    best_score = best['score']
    margin = max(0.0, best_score - second_score)
    local_conf = min(0.96, 0.45 + 0.03 * best_score + 0.025 * margin + 0.12 * best.get('memory_similarity', 0))
    if best_score < 2.8:
        local_conf = min(local_conf, 0.55)

    chosen = None
    if dung_ai:
        chosen = _grad_ai_chon_pham_vi(text, ranked)

    if chosen:
        ai_conf = chosen.get('ai_confidence', 0) / 100.0
        conf = round(max(0.0, min(0.99, 0.25 * local_conf + 0.75 * ai_conf)), 2)
        # <65%: chỉ giữ tới Chương; 65–84%: gán Bài nhưng đánh dấu cần kiểm tra; >=85%: tin cậy cao.
        bai = chosen.get('bai', '') if conf >= 0.65 else ''
        return {
            'khoi': chosen.get('khoi', ''),
            'chuong': chosen.get('chuong', ''),
            'bai': bai,
            'do_tin_cay': conf,
            'nguon_phan_loai': 'AI chọn trong ứng viên KHO_YCCD',
            'can_kiem_tra': conf < 0.85,
            'ly_do': chosen.get('ai_reason', ''),
            'ung_vien': ranked,
        }

    # Không gọi AI: chỉ tự gắn Bài khi thuật toán + ký ức GV đủ chắc.
    bai = best.get('bai', '') if local_conf >= 0.72 else ''
    return {
        'khoi': best.get('khoi', ''),
        'chuong': best.get('chuong', ''),
        'bai': bai,
        'do_tin_cay': round(local_conf, 2),
        'nguon_phan_loai': 'KHO_YCCD + ghi nhớ GV',
        'can_kiem_tra': local_conf < 0.85,
        'ung_vien': ranked,
    }


def phan_loai_pham_vi_tot_nghiep_bang_ai(q):
    """Phân loại AI cho 1 câu, giữ nguyên mọi nội dung/đáp án khác."""
    q2 = dict(q)
    if (q2.get('phan_loai_on_tap') or {}).get('gv_xac_nhan'):
        return q2
    pv = xac_dinh_pham_vi_bai_chuong_tot_nghiep(q2, dung_ai=True)
    if pv.get('khoi'):
        q2['khoi'] = pv.get('khoi', '')
    if pv.get('chuong'):
        q2['chuong'] = pv.get('chuong', '')
    q2['bai'] = pv.get('bai', '') or ''
    q2['do_tin_cay_pham_vi_kien_thuc'] = float(pv.get('do_tin_cay', 0) or 0)
    q2['phan_loai_on_tap'] = {
        'khoi': q2.get('khoi', ''),
        'chuong': q2.get('chuong', ''),
        'bai': q2.get('bai', ''),
        'muc': 'Bài' if str(q2.get('bai', '')).strip() else 'Chương',
        'do_tin_cay': float(pv.get('do_tin_cay', 0) or 0),
        'ghi_chu': pv.get('ly_do', ''),
        'phien_ban': 'khoi_chuong_bai_v3_ai_candidate',
        'gv_xac_nhan': False,
        'can_kiem_tra': bool(pv.get('can_kiem_tra', False)),
        'nguon_phan_loai': pv.get('nguon_phan_loai', ''),
    }
    return q2

def xac_dinh_muc_do_cau_tot_nghiep(noi_dung, co_du_lieu=False):
    """Phân mức Nhận biết/Thông hiểu/Vận dụng bằng quy tắc cục bộ."""
    s = _bo_dau_nguon(noi_dung)

    van_dung = [
        "tinh", "xac suat", "ti le", "ty le", "bao nhieu", "suy ra", "du doan",
        "neu thay doi", "neu xay ra", "ket qua thi nghiem", "bang so lieu", "bieu do",
        "do thi", "pha he", "lai", "quan the", "tan so allele", "tan so alen",
        "phan tich du lieu", "xu li du lieu", "xu ly du lieu", "de xuat giai phap",
        "thuc tien", "trong san xuat", "trong doi song", "giai quyet",
    ]
    thong_hieu = [
        "giai thich", "vi sao", "phan biet", "so sanh", "moi quan he", "nguyen nhan",
        "co che", "nhan xet", "phan tich", "ket luan", "chung minh", "lien he",
    ]

    score_vd = sum(1 for kw in van_dung if kw in s)
    score_th = sum(1 for kw in thong_hieu if kw in s)

    # Có bảng/hình/số liệu chỉ nâng lên Vận dụng khi câu thật sự đòi khai thác dữ kiện.
    if co_du_lieu and any(kw in s for kw in [
        "dua vao", "can cu", "tu du lieu", "ket qua", "bao nhieu", "nhan dinh",
        "suy ra", "xac dinh", "du doan", "phan tich"
    ]):
        score_vd += 2

    if score_vd >= 2:
        return "Vận dụng"
    if score_th >= 1 or score_vd == 1:
        return "Thông hiểu"
    return "Nhận biết"


def _grad_fallback_nang_luc(noi_dung, muc_do, co_du_lieu=False):
    kq = xac_dinh_nang_luc_chi_bao_tu_noi_dung(noi_dung)
    nl = str(kq.get("thanh_phan_nang_luc", "")).strip()
    if nl in THANH_PHAN_NANG_LUC:
        return kq

    s = _bo_dau_nguon(noi_dung)
    if any(x in s for x in ["de xuat giai phap", "thuc tien", "doi song", "suc khoe", "moi truong", "san xuat"]):
        nl = "Vận dụng kiến thức, kĩ năng đã học"
        ma = "VD2"
    elif co_du_lieu and any(x in s for x in ["bang", "bieu do", "do thi", "thi nghiem", "du lieu", "ket qua"]):
        nl = "Tìm hiểu thế giới sống"
        ma = "TH5"
    else:
        nl = "Nhận thức sinh học"
        ma = "NT6" if muc_do in {"Thông hiểu", "Vận dụng"} else "NT1"

    return {
        "thanh_phan_nang_luc": nl,
        "chi_bao": ma,
        "mo_ta_chi_bao": mo_ta_chi_bao_theo_ma(ma),
        "do_tin_cay": 0.62,
        "trang_thai": "Tự xác định",
    }


def gan_phan_loai_on_tap_cho_cau_tot_nghiep(q, force=False):
    """
    Gắn phạm vi Khối → Chương → Bài + mức độ + năng lực để câu đề thật
    có thể dùng cho ôn tập thường.

    - Không mặc định Khối 12.
    - Không ép gán YCCĐ chi tiết.
    - Câu Đúng/Sai được phân loại mức độ/năng lực và phạm vi kiến thức
      riêng cho từng ý a–d, đồng thời vẫn có phạm vi chung cấp câu để lọc.
    """
    q2 = dict(q)
    resources = list(q2.get("tai_nguyen_truc_quan", []) or [])
    co_du_lieu = bool(resources or q2.get("du_lieu_truc_quan"))

    phan_loai_cu = q2.get("phan_loai_on_tap", {}) or {}
    phien_ban_cu = str(phan_loai_cu.get("phien_ban", ""))
    gv_xac_nhan = bool(phan_loai_cu.get("gv_xac_nhan", False))

    # 1) Phạm vi kiến thức: Khối → Chương → Bài.
    # Dữ liệu từ bản cũ (mặc định Khối 12) được tự phân loại lại một lần.
    can_phan_loai_lai = (
        force
        or phien_ban_cu not in {"khoi_chuong_bai_v2", "khoi_chuong_bai_v3_ai_candidate"}
        or not str(q2.get("khoi", "")).strip()
        or not str(q2.get("chuong", "")).strip()
    )

    if can_phan_loai_lai and not gv_xac_nhan:
        pham_vi = xac_dinh_pham_vi_bai_chuong_tot_nghiep(q2)
        if pham_vi.get("khoi"):
            q2["khoi"] = pham_vi.get("khoi", "")
        if pham_vi.get("chuong"):
            q2["chuong"] = pham_vi.get("chuong", "")
        # Nếu chưa đủ chắc để gắn bài, xóa bài mặc định/cũ để tránh lọc sai.
        q2["bai"] = pham_vi.get("bai", "") or ""
        q2["do_tin_cay_pham_vi_kien_thuc"] = pham_vi.get("do_tin_cay", 0.0)

    # Nhãn theo dõi nội bộ cho hồ sơ HS, không giả là YCCĐ chính thức.
    yccd_hien_tai = str(q2.get("yccd", "") or "").strip()
    if (not yccd_hien_tai) or yccd_hien_tai.startswith("Kiến thức trọng tâm"):
        khoi_txt = str(q2.get("khoi", "") or "").strip()
        bai_txt = str(q2.get("bai", "") or "").strip()
        chuong_txt = str(q2.get("chuong", "") or "").strip()
        phan = bai_txt or chuong_txt
        if phan:
            q2["yccd"] = "Kiến thức trọng tâm • " + " • ".join(
                x for x in [khoi_txt, phan] if x
            )

    q2["phan_loai_on_tap"] = {
        "khoi": q2.get("khoi", ""),
        "chuong": q2.get("chuong", ""),
        "bai": q2.get("bai", ""),
        "muc": "Bài" if str(q2.get("bai", "")).strip() else "Chương",
        "do_tin_cay": float(q2.get("do_tin_cay_pham_vi_kien_thuc", 0) or 0),
        "ghi_chu": (
            "Phân loại thô Khối → Chương → Bài để ôn tập; "
            "không ép gán YCCĐ chi tiết."
        ),
        "phien_ban": "khoi_chuong_bai_v3_ai_candidate",
        "gv_xac_nhan": gv_xac_nhan,
    }

    # 2) Mức độ + năng lực. Đúng/Sai làm riêng từng ý.
    if q2.get("dang_cau") == "Đúng / Sai":
        meta_moi = []
        muc_rank = {"Nhận biết": 1, "Thông hiểu": 2, "Vận dụng": 3}
        muc_max = "Nhận biết"

        for nd in list(q2.get("nhan_dinh_meta", []) or []):
            nd2 = dict(nd)
            noi_dung_nd = " ".join([
                str(q2.get("tinh_huong", "") or ""),
                str(q2.get("cau_hoi", "") or ""),
                str(nd2.get("noi_dung", "") or ""),
            ])

            # Phạm vi riêng từng ý để sau này phân tích HS chính xác hơn.
            # Nếu GV đã sửa/xác nhận phạm vi thì tuyệt đối không cho bộ phân loại tự động ghi đè.
            nd_pv_cu = dict(nd2.get("phan_loai_on_tap", {}) or {})
            nd_gv_xac_nhan = bool(nd_pv_cu.get("gv_xac_nhan", False))

            if not nd_gv_xac_nhan:
                pham_vi_nd = xac_dinh_pham_vi_bai_chuong_tot_nghiep(q2, nd=nd2)
                if pham_vi_nd.get("khoi"):
                    nd2["khoi"] = pham_vi_nd.get("khoi", "")
                if pham_vi_nd.get("chuong"):
                    nd2["chuong"] = pham_vi_nd.get("chuong", "")
                nd2["bai"] = pham_vi_nd.get("bai", "") or ""
                nd2["do_tin_cay_pham_vi_kien_thuc"] = pham_vi_nd.get("do_tin_cay", 0.0)
                nd2["phan_loai_on_tap"] = {
                    "khoi": nd2.get("khoi", q2.get("khoi", "")),
                    "chuong": nd2.get("chuong", q2.get("chuong", "")),
                    "bai": nd2.get("bai", q2.get("bai", "")),
                    "muc": "Bài" if str(nd2.get("bai", "")).strip() else "Chương",
                    "do_tin_cay": float(nd2.get("do_tin_cay_pham_vi_kien_thuc", 0) or 0),
                    "phien_ban": "khoi_chuong_bai_v3_ai_candidate",
                    "gv_xac_nhan": False,
                }
            else:
                # Đồng bộ lại các trường phẳng từ phạm vi GV đã xác nhận.
                nd2["khoi"] = nd_pv_cu.get("khoi", nd2.get("khoi", q2.get("khoi", "")))
                nd2["chuong"] = nd_pv_cu.get("chuong", nd2.get("chuong", q2.get("chuong", "")))
                nd2["bai"] = nd_pv_cu.get("bai", nd2.get("bai", q2.get("bai", "")))
                nd2["do_tin_cay_pham_vi_kien_thuc"] = 1.0

            if force or str(nd2.get("muc_do", "")).strip() not in {"Nhận biết", "Thông hiểu", "Vận dụng"}:
                nd2["muc_do"] = xac_dinh_muc_do_cau_tot_nghiep(
                    noi_dung_nd,
                    co_du_lieu=co_du_lieu
                )

            kq = _grad_fallback_nang_luc(
                noi_dung_nd,
                nd2.get("muc_do", ""),
                co_du_lieu=co_du_lieu
            )

            if force or str(nd2.get("thanh_phan_nang_luc", "")).strip() not in THANH_PHAN_NANG_LUC:
                nd2["thanh_phan_nang_luc"] = kq.get(
                    "thanh_phan_nang_luc",
                    "Nhận thức sinh học"
                )
                nd2["chi_bao"] = kq.get("chi_bao", "")

            if not str(nd2.get("yccd", "")).strip() or str(nd2.get("yccd", "")).startswith("Kiến thức trọng tâm"):
                khoi_nd = str(nd2.get("khoi", q2.get("khoi", "")) or "").strip()
                phan_nd = str(nd2.get("bai", "") or nd2.get("chuong", "") or q2.get("bai", "") or q2.get("chuong", "")).strip()
                nd2["yccd"] = "Kiến thức trọng tâm • " + " • ".join(
                    x for x in [khoi_nd, phan_nd] if x
                )

            if muc_rank.get(nd2.get("muc_do"), 0) > muc_rank.get(muc_max, 0):
                muc_max = nd2.get("muc_do")
            meta_moi.append(nd2)

        q2["nhan_dinh_meta"] = meta_moi
        q2["muc_do"] = muc_max

        # Năng lực cấp câu chỉ dùng hiển thị tổng quát; phân tích HS dùng từng ý.
        nls = [
            x.get("thanh_phan_nang_luc")
            for x in meta_moi
            if x.get("thanh_phan_nang_luc") in THANH_PHAN_NANG_LUC
        ]
        if nls:
            q2["thanh_phan_nang_luc"] = max(set(nls), key=nls.count)

    else:
        noi_dung = _grad_noi_dung_phan_loai(q2)

        if force or str(q2.get("muc_do", "")).strip() not in {"Nhận biết", "Thông hiểu", "Vận dụng"}:
            q2["muc_do"] = xac_dinh_muc_do_cau_tot_nghiep(
                noi_dung,
                co_du_lieu=co_du_lieu
            )

        kq = _grad_fallback_nang_luc(
            noi_dung,
            q2.get("muc_do", ""),
            co_du_lieu=co_du_lieu
        )

        if force or str(q2.get("thanh_phan_nang_luc", "")).strip() not in THANH_PHAN_NANG_LUC:
            q2["thanh_phan_nang_luc"] = kq.get(
                "thanh_phan_nang_luc",
                "Nhận thức sinh học"
            )
            q2["chi_bao"] = kq.get("chi_bao", "")
            q2["mo_ta_chi_bao"] = kq.get("mo_ta_chi_bao", "")

    return q2


def _grad_chuan_hoa_cau_da_nhap(q):
    """Nâng cấp câu cũ theo cơ chế mới mà không làm mất dữ liệu nguồn."""
    q2 = gan_phan_loai_on_tap_cho_cau_tot_nghiep(q, force=False)
    status = str(q2.get("trang_thai", "") or "").strip()

    if not _grad_co_dap_an_day_du(q2):
        q2["trang_thai"] = "Thiếu đáp án"
        q2["duoc_dung_luyen_hs"] = False
        q2["san_sang_tu_dap_an_nguon"] = False
        if not str(q2.get("canh_bao_dap_an", "") or "").strip():
            q2["canh_bao_dap_an"] = "Chưa đọc được đáp án đầy đủ từ file nguồn."
        return q2

    if status in {"Cần GV xem", "Ngừng sử dụng"}:
        q2["duoc_dung_luyen_hs"] = False
        q2["san_sang_tu_dap_an_nguon"] = False
        return q2

    if status == "Thiếu đáp án":
        # Dữ liệu cũ có thể đã từng bị đánh dấu thiếu nhưng hiện đã có đáp án sau khi sửa.
        q2["trang_thai"] = "Chờ rà soát"

    if not str(q2.get("trang_thai", "") or "").strip():
        q2["trang_thai"] = "Chờ rà soát"

    q2["duoc_dung_luyen_hs"] = True
    q2["san_sang_tu_dap_an_nguon"] = True
    return q2

def fingerprint_cau_tot_nghiep_thuc_te(q):
    raw = {
        "dang": q.get("dang_cau", ""),
        "cau_hoi": _grad_norm_text(q.get("cau_hoi", "")),
        "lua_chon": [_grad_norm_text(x) for x in (q.get("lua_chon", []) or [])],
        "nhan_dinh": [
            _grad_norm_text(x.get("noi_dung", ""))
            for x in (q.get("nhan_dinh_meta", []) or [])
        ],
    }
    return hashlib.sha256(
        json.dumps(raw, ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()


def _grad_part_key(value):
    s = str(value or "").strip().upper()
    if s in {"1", "I"}:
        return "I"
    if s in {"2", "II"}:
        return "II"
    if s in {"3", "III"}:
        return "III"
    return ""




def _grad_run_format_score(run):
    """Chấm tín hiệu định dạng của một run Word dùng để đánh dấu đáp án."""
    from docx.oxml.ns import qn

    strong = 0
    bold = 0
    italic = 0

    try:
        if run.font.highlight_color is not None:
            strong += 6
    except Exception:
        pass

    try:
        if bool(run.underline):
            strong += 5
    except Exception:
        pass

    # Màu chữ: ưu tiên mạnh màu đỏ, vẫn nhận các màu nổi bật khác.
    try:
        rgb = run.font.color.rgb
        if rgb is not None:
            rgb_s = str(rgb).upper().replace('#', '')
            if len(rgb_s) >= 6:
                r = int(rgb_s[0:2], 16)
                g = int(rgb_s[2:4], 16)
                b = int(rgb_s[4:6], 16)
                if r >= 150 and r >= g * 1.35 and r >= b * 1.35:
                    strong += 7
                elif rgb_s not in {"000000", "FFFFFF", "AUTO"}:
                    strong += 4
    except Exception:
        pass

    # Một số file dùng Theme Color thay vì RGB trực tiếp.
    try:
        if run.font.color.theme_color is not None and run.font.color.rgb is None:
            strong += 2
    except Exception:
        pass

    # Tô nền/shading ở cấp run (khác với text highlight).
    try:
        rpr = run._r.get_or_add_rPr()
        shd = rpr.find(qn('w:shd'))
        if shd is not None:
            fill = str(shd.get(qn('w:fill')) or '').upper()
            if fill and fill not in {"AUTO", "FFFFFF", "000000", "NIL", "CLEAR"}:
                strong += 6
    except Exception:
        pass

    try:
        if run.bold is True:
            bold += 1
    except Exception:
        pass

    try:
        if run.italic is True:
            italic += 1
    except Exception:
        pass

    return {"strong": strong, "bold": bold, "italic": italic}


def _grad_merge_format_score(a, b):
    a = dict(a or {})
    b = dict(b or {})
    return {
        "strong": int(a.get("strong", 0) or 0) + int(b.get("strong", 0) or 0),
        "bold": int(a.get("bold", 0) or 0) + int(b.get("bold", 0) or 0),
        "italic": int(a.get("italic", 0) or 0) + int(b.get("italic", 0) or 0),
    }


def _grad_option_format_signals_from_paragraph(para):
    """
    Đọc tín hiệu định dạng cho TẤT CẢ phương án A/B/C/D trong một paragraph.

    Hỗ trợ cả hai kiểu Word phổ biến:
    - mỗi phương án nằm trên một dòng riêng;
    - A/B/C/D nằm chung một dòng, ngăn bằng tab/khoảng trắng.

    Các dấu hiệu đáp án được nhận diện: tô nền/highlight, màu chữ (kể cả đỏ),
    gạch chân, in đậm và in nghiêng. Chỉ chấm phần nội dung phương án,
    không lấy định dạng riêng của nhãn A./B./C./D.
    """
    full_text = str(getattr(para, "text", "") or "")
    if not full_text.strip():
        return {}

    pat = re.compile(r"(?i)(?<![A-Za-z0-9À-ỹ])([A-D])\s*[\.\)]\s*")
    matches = list(pat.finditer(full_text))
    if not matches:
        return {}

    # Vị trí ký tự của từng run trong paragraph.
    run_spans = []
    pos = 0
    for run in getattr(para, "runs", []) or []:
        rt = str(getattr(run, "text", "") or "")
        start = pos
        end = pos + len(rt)
        run_spans.append((start, end, run))
        pos = end

    result = {}
    for i, m in enumerate(matches):
        letter = m.group(1).upper()
        body_start = m.end()
        body_end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        score = {"strong": 0, "bold": 0, "italic": 0}

        for rs, re_, run in run_spans:
            # Chỉ lấy run có giao với PHẦN NỘI DUNG, bỏ nhãn phương án.
            if re_ <= body_start or rs >= body_end:
                continue
            overlap_start = max(rs, body_start)
            overlap_end = min(re_, body_end)
            if overlap_end <= overlap_start:
                continue
            run_text = str(getattr(run, "text", "") or "")
            local_a = max(0, overlap_start - rs)
            local_b = max(local_a, overlap_end - rs)
            if not run_text[local_a:local_b].strip():
                continue
            score = _grad_merge_format_score(score, _grad_run_format_score(run))

        # Tô nền cả paragraph: chỉ có ý nghĩa nếu paragraph chứa đúng 1 phương án.
        if len(matches) == 1:
            try:
                from docx.oxml.ns import qn
                ppr = para._p.get_or_add_pPr()
                shd = ppr.find(qn('w:shd'))
                if shd is not None:
                    fill = str(shd.get(qn('w:fill')) or '').upper()
                    if fill and fill not in {"AUTO", "FFFFFF", "000000", "NIL", "CLEAR"}:
                        score["strong"] += 5
            except Exception:
                pass

        result[letter] = _grad_merge_format_score(result.get(letter), score)

    return result


def _grad_option_format_signal(para):
    """Giữ tương thích với code cũ: trả về tín hiệu của phương án đầu tiên."""
    signals = _grad_option_format_signals_from_paragraph(para)
    if not signals:
        return None, {"strong": 0, "bold": 0, "italic": 0}
    letter = next(iter(signals.keys()))
    return letter, signals[letter]


def _grad_chon_dap_an_tu_dinh_dang(option_scores):
    """Chỉ suy đáp án khi một phương án có định dạng nổi bật duy nhất."""
    if not isinstance(option_scores, dict) or len(option_scores) < 3:
        return ""

    rows = []
    for letter in ["A", "B", "C", "D"]:
        sc = option_scores.get(letter, {}) or {}
        rows.append((
            letter,
            int(sc.get("strong", 0) or 0),
            int(sc.get("bold", 0) or 0),
            int(sc.get("italic", 0) or 0),
        ))

    # Ưu tiên tín hiệu mạnh: highlight/tô nền/màu chữ/gạch chân.
    by_strong = sorted(rows, key=lambda x: x[1], reverse=True)
    if by_strong[0][1] > 0 and by_strong[0][1] >= by_strong[1][1] + 2:
        return by_strong[0][0]

    # Nếu chỉ dùng in đậm: phải có đúng một phương án nổi bật.
    by_bold = sorted(rows, key=lambda x: x[2], reverse=True)
    if by_bold[0][2] > 0 and by_bold[0][2] > by_bold[1][2]:
        return by_bold[0][0]

    # In nghiêng chỉ dùng như fallback rất thận trọng.
    by_italic = sorted(rows, key=lambda x: x[3], reverse=True)
    if by_italic[0][3] > 0 and by_italic[0][3] > by_italic[1][3]:
        return by_italic[0][0]

    return ""


def _grad_split_mcq_line(line):
    """
    Tách một dòng Word có thể chứa 1 hoặc nhiều phương án A/B/C/D.
    Trả về (phần trước phương án đầu tiên, [(A,nội dung), ...]).
    """
    text = str(line or "")
    pat = re.compile(r"(?i)(?<![A-Za-z0-9À-ỹ])([A-D])\s*[\.\)]\s*")
    matches = list(pat.finditer(text))
    if not matches:
        return text.strip(), []

    prefix = text[:matches[0].start()].strip()
    options = []
    for i, m in enumerate(matches):
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[m.end():end].strip(" \t;|")
        if body:
            options.append((m.group(1).upper(), body))
    return prefix, options


def _grad_extract_images_from_paragraph(para, doc, media_dir, stem, counter_start=0):
    """Trích ảnh nổi/inline gắn với một paragraph Word và lưu thành file thật."""
    from docx.oxml.ns import qn

    paths = []
    counter = counter_start
    try:
        blips = para._p.xpath('.//a:blip')
    except Exception:
        blips = []

    for blip in blips:
        rid = blip.get(qn('r:embed'))
        if not rid:
            continue
        try:
            part = doc.part.related_parts[rid]
            blob = part.blob
            ext = ".png"
            content_type = str(getattr(part, "content_type", "") or "").lower()
            if "jpeg" in content_type or "jpg" in content_type:
                ext = ".jpg"
            elif "gif" in content_type:
                ext = ".gif"
            elif "webp" in content_type:
                ext = ".webp"
            counter += 1
            name = f"{stem}_img_{counter:03d}{ext}"
            path = os.path.join(media_dir, name)
            with open(path, "wb") as f:
                f.write(blob)
            paths.append(path)
        except Exception:
            continue
    return paths, counter


def _grad_parse_answer_tables(doc):
    """Đọc bảng đáp án dạng phổ biến của đề tốt nghiệp 18/4/6."""
    ans = {"I": {}, "II": {}, "III": {}}

    for tbl in doc.tables:
        rows = [
            [" ".join(str(c.text or "").split()) for c in row.cells]
            for row in tbl.rows
        ]
        if len(rows) < 2:
            continue

        for r in range(len(rows) - 1):
            row1 = rows[r]
            row2 = rows[r + 1]
            if not row1 or not row2:
                continue
            if _grad_norm_text(row1[0]) not in {"câu", "cau"}:
                continue

            label = _grad_norm_text(row2[0])
            nums = []
            vals = []
            for x in row1[1:]:
                m = re.search(r"\d+", str(x))
                nums.append(int(m.group()) if m else None)
            vals = [str(x or "").strip() for x in row2[1:]]
            pairs = [(n, v) for n, v in zip(nums, vals) if n is not None and v]
            if not pairs:
                continue

            # Phần I: A/B/C/D
            if label in {"đa", "da", "đáp án", "dap an"} and all(
                str(v).strip().upper() in {"A", "B", "C", "D"}
                for _, v in pairs
            ):
                for n, v in pairs:
                    ans["I"][n] = str(v).strip().upper()
                continue

            # Phần II: chuỗi 4 ký tự D/S, ví dụ SDDD.
            if label in {"đáp án", "dap an"} and all(
                re.fullmatch(r"[DdSsĐđ]{4}", str(v).replace(" ", ""))
                for _, v in pairs
            ):
                for n, v in pairs:
                    seq = str(v).replace(" ", "").upper()
                    seq = seq.replace("Đ", "D")
                    ans["II"][n] = [
                        "Đúng" if ch == "D" else "Sai"
                        for ch in seq
                    ]
                continue

            # Phần III: các giá trị ngắn, thường có 6 câu.
            if label in {"đáp án", "dap an"}:
                for n, v in pairs:
                    # Tránh nuốt nhầm bảng Phần II đã xử lý ở trên.
                    if not re.fullmatch(r"[DdSsĐđ]{4}", str(v).replace(" ", "")):
                        ans["III"][n] = str(v).strip()

    return ans


def _grad_extract_inline_meta(lines):
    """Tách metadata từ file mẫu: [[ĐÁP ÁN]], [[HƯỚNG DẪN GIẢI]]."""
    answer = ""
    solution = ""
    clean = []
    collecting_solution = False

    for line in lines:
        s = str(line or "").strip()
        m_ans = re.match(r"^\[\[?\s*(?:ĐÁP\s*ÁN|DAP\s*AN)\s*\]?\]\s*[:：]?\s*(.*)$", s, flags=re.I)
        m_sol = re.match(r"^\[\[?\s*(?:HƯỚNG\s*DẪN\s*GIẢI|HUONG\s*DAN\s*GIAI|LỜI\s*GIẢI|LOI\s*GIAI)\s*\]?\]\s*[:：]?\s*(.*)$", s, flags=re.I)
        if m_ans:
            answer = m_ans.group(1).strip()
            collecting_solution = False
            continue
        if m_sol:
            solution = m_sol.group(1).strip()
            collecting_solution = True
            continue
        if collecting_solution:
            if s.casefold().startswith("lưu ý quan trọng"):
                collecting_solution = False
                continue
            if s.startswith("[["):
                collecting_solution = False
                clean.append(s)
            else:
                solution = (solution + "\n" + s).strip()
            continue
        clean.append(line)

    return clean, answer, solution


def _grad_parse_question_structure(part, qnum, lines, resources, answer_map, source_name):
    lines = [" ".join(str(x or "").split()) for x in lines if str(x or "").strip()]
    if not lines:
        return None

    # Bỏ nhãn Câu n ở dòng đầu nhưng giữ nguyên phần nội dung phía sau.
    lines[0] = re.sub(r"(?i)^\s*Câu\s+\d+\s*[\.:\)]\s*", "", lines[0]).strip()
    lines, inline_answer, inline_solution = _grad_extract_inline_meta(lines)

    part = _grad_part_key(part)
    if part == "I":
        dang = "Trắc nghiệm 4 lựa chọn"
    elif part == "II":
        dang = "Đúng / Sai"
    elif part == "III":
        dang = "Trả lời ngắn"
    else:
        # fallback theo dấu hiệu nội dung
        if sum(bool(re.match(r"^[A-D][\.\)]\s+", x)) for x in lines) >= 3:
            dang = "Trắc nghiệm 4 lựa chọn"
        elif sum(bool(re.match(r"^[abcd][\.\)]\s+", x)) for x in lines) >= 3:
            dang = "Đúng / Sai"
        else:
            dang = "Trả lời ngắn"

    q = {
        "id": str(uuid.uuid4()),
        "temp_id": str(uuid.uuid4()),
        "khoi": "Khối 12",
        "chuong": "",
        "bai": "",
        "yccd": "",
        "muc_do": "",
        "dang_cau": dang,
        "thanh_phan_nang_luc": "",
        "chi_bao": "",
        "cau_hoi": "",
        "tinh_huong": "",
        "lua_chon": [],
        "dap_an": "",
        "giai_thich": inline_solution,
        "nguon_giai_thich": "Lời giải từ file nguồn" if str(inline_solution or "").strip() else "",
        "giai_thich_ai_de_xuat": "",
        "nhan_dinh_meta": [],
        "tai_nguyen_truc_quan": resources or [],
        "nguon_file": source_name,
        "so_cau_goc": int(qnum),
        "phan_goc": part,
        "nguon": source_name,
        "nguon_tao": "Đề thật / đề thi thử GV tải lên",
        "muc_dich_su_dung": "tot_nghiep",
        # Nếu file nguồn đã có đáp án đầy đủ, câu được phép dùng ngay.
        # AI rà soát là lớp kiểm tra bổ sung; câu đáng nghi mới bị khóa lại.
        "duoc_dung_luyen_hs": False,
        "ngay_nhap": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "trang_thai": "Chờ rà soát",
        "canh_bao_dap_an": "",
        "do_tin_cay_dap_an": None,
    }

    if dang == "Trắc nghiệm 4 lựa chọn":
        opts_map = {}
        body = []
        for line in lines:
            prefix, found_opts = _grad_split_mcq_line(line)
            if found_opts:
                if prefix:
                    body.append(prefix)
                for letter, content in found_opts:
                    # Giữ phương án đầu tiên nếu file lặp nhãn ngoài ý muốn.
                    if letter not in opts_map and str(content).strip():
                        opts_map[letter] = str(content).strip()
            else:
                body.append(line)
        q["cau_hoi"] = "\n".join(body).strip()
        q["lua_chon"] = [
            f"{letter}. {opts_map[letter]}"
            for letter in ["A", "B", "C", "D"]
            if letter in opts_map
        ]
        q["dap_an"] = inline_answer.strip().upper() if inline_answer else str(answer_map.get("I", {}).get(int(qnum), "")).strip().upper()

    elif dang == "Đúng / Sai":
        statements = []
        body = []
        for line in lines:
            m = re.match(r"^([abcd])[\)\.]\s*(.*)$", line)
            if m:
                statements.append((m.group(1), m.group(2).strip()))
            else:
                body.append(line)
        q["cau_hoi"] = "\n".join(body).strip()

        ans4 = answer_map.get("II", {}).get(int(qnum), [])
        if inline_answer:
            raw = inline_answer.replace(";", " ").replace(",", " ")
            # Hỗ trợ SDDD / D S D D / a) Sai; b) Đúng...
            compact = re.sub(r"[^DdSsĐđ]", "", raw)
            if len(compact) >= 4:
                compact = compact[:4].upper().replace("Đ", "D")
                ans4 = ["Đúng" if c == "D" else "Sai" for c in compact]
            else:
                vals = re.findall(r"(?i)\b(Đúng|Sai)\b", raw)
                if len(vals) >= 4:
                    ans4 = ["Đúng" if v.casefold().startswith("đ") else "Sai" for v in vals[:4]]

        for idx, (ky, nd_text) in enumerate(statements[:4]):
            q["nhan_dinh_meta"].append({
                "ky_hieu": ky,
                "noi_dung": nd_text,
                "yccd": "",
                "muc_do": "",
                "thanh_phan_nang_luc": "",
                "chi_bao": "",
                "dap_an": ans4[idx] if idx < len(ans4) else "",
                "giai_thich": "",
            })

    else:
        q["cau_hoi"] = "\n".join(lines).strip()
        q["dap_an"] = inline_answer.strip() if inline_answer else str(answer_map.get("III", {}).get(int(qnum), "")).strip()

    # Tạo trường tương thích renderer cũ từ tài nguyên đầu tiên.
    for res in q["tai_nguyen_truc_quan"]:
        if res.get("loai") == "anh" and res.get("duong_dan"):
            q["du_lieu_truc_quan"] = {
                "loai": "hinh_tu_tai_lieu",
                "duong_dan_anh": res.get("duong_dan"),
                "nguon": source_name,
                "mo_ta": "Hình trích nguyên từ đề nguồn",
            }
            break
        if res.get("loai") == "bang" and res.get("du_lieu"):
            rows = res.get("du_lieu") or []
            if rows:
                q["du_lieu_truc_quan"] = {
                    "loai": "bang_so_lieu",
                    "cot": rows[0],
                    "du_lieu": rows[1:],
                    "nguon": source_name,
                    "mo_ta": "Bảng trích nguyên từ đề nguồn",
                }
                break
    if "du_lieu_truc_quan" not in q:
        q["du_lieu_truc_quan"] = {}

    # Nếu thiếu đáp án thì bắt buộc GV xử lý trước khi dùng.
    if dang == "Đúng / Sai":
        co_dap_an = len(q["nhan_dinh_meta"]) == 4 and all(
            x.get("dap_an") in {"Đúng", "Sai"} for x in q["nhan_dinh_meta"]
        )
    else:
        co_dap_an = bool(str(q.get("dap_an", "")).strip())

    if not co_dap_an:
        q["trang_thai"] = "Thiếu đáp án"
        q["duoc_dung_luyen_hs"] = False
        q["san_sang_tu_dap_an_nguon"] = False
        q["canh_bao_dap_an"] = "Chưa đọc được đáp án từ file nguồn."
    else:
        # Đề thật có đáp án nguồn: dùng được ngay, không phải chờ AI rà soát.
        q["trang_thai"] = "Chờ rà soát"
        q["duoc_dung_luyen_hs"] = True
        q["san_sang_tu_dap_an_nguon"] = True
        q["canh_bao_dap_an"] = ""

    # Phân loại cục bộ để câu đề thật dùng được cả trong ôn theo bài/chương.
    # Không ép gán YCCĐ chi tiết; chỉ gắn phạm vi bài/chương, mức độ và năng lực.
    try:
        q = gan_phan_loai_on_tap_cho_cau_tot_nghiep(q, force=True)
    except Exception:
        try:
            q = gan_chi_bao_chuan_cho_cau(q, force=True)
        except Exception:
            pass

    return q


def tach_de_tot_nghiep_docx(file_path, source_name=None):
    """
    Tách đề Word đúng form tốt nghiệp, giữ câu + bảng + ảnh.
    Hỗ trợ numbering tự động Word và bảng đáp án kiểu 18/4/6.
    """
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    source_name = source_name or os.path.basename(file_path)
    doc = Document(file_path)
    answer_map = _grad_parse_answer_tables(doc)

    safe_stem = re.sub(r"[^0-9A-Za-zÀ-ỹ_-]+", "_", os.path.splitext(source_name)[0])[:80]
    source_hash = hashlib.sha1(source_name.encode("utf-8", errors="ignore")).hexdigest()[:8]
    media_dir = os.path.join(GRAD_MEDIA_DIR, f"{safe_stem}_{source_hash}")
    os.makedirs(media_dir, exist_ok=True)

    questions_raw = []
    current = None
    current_part = ""
    question_counter = 0
    question_num_id = None
    image_counter = 0
    in_answer_section = False

    def finalize():
        nonlocal current
        if current and current.get("lines"):
            questions_raw.append(current)
        current = None

    def num_info(para):
        try:
            ppr = para._p.pPr
            if ppr is None or ppr.numPr is None:
                return None, None
            num_id = ppr.numPr.numId.val if ppr.numPr.numId is not None else None
            ilvl = ppr.numPr.ilvl.val if ppr.numPr.ilvl is not None else 0
            return num_id, ilvl
        except Exception:
            return None, None

    for child in doc.element.body.iterchildren():
        tag = str(child.tag)
        if tag.endswith("}p"):
            para = Paragraph(child, doc)
            txt = str(para.text or "").strip()

            # Ảnh có thể nằm trong paragraph không có text.
            img_paths, image_counter = _grad_extract_images_from_paragraph(
                para, doc, media_dir, safe_stem, image_counter
            )

            if txt:
                # Đáp án / hướng dẫn giải ở cuối file: dừng gom nội dung đề.
                if re.match(r"(?i)^\s*(ĐÁP\s*ÁN|DAP\s*AN)\s*$", txt):
                    finalize()
                    in_answer_section = True
                    continue

                m_part = re.match(r"(?i)^\s*PHẦN\s+(I{1,3}|[1-3])\b", txt)
                if m_part and not in_answer_section:
                    finalize()
                    current_part = _grad_part_key(m_part.group(1))
                    question_counter = 0
                    question_num_id = None
                    continue

                if in_answer_section:
                    continue

                # Khôi phục nhãn Câu n khi Word dùng numbering tự động.
                num_id, ilvl = num_info(para)
                if current_part in {"I", "II", "III"} and num_id is not None and int(ilvl or 0) == 0:
                    if question_num_id is None:
                        question_num_id = num_id
                    if num_id == question_num_id and not re.match(r"(?i)^\s*Câu\s+\d+", txt):
                        question_counter += 1
                        txt = f"Câu {question_counter}: {txt}"

                m_q = re.match(r"(?i)^\s*Câu\s+(\d+)\s*[\.:\)]\s*(.*)$", txt)
                if m_q:
                    finalize()
                    qnum = int(m_q.group(1))
                    current = {
                        "part": current_part,
                        "qnum": qnum,
                        "lines": [txt],
                        "resources": [],
                        "option_format_scores": {},
                    }
                    # Có file để cả câu hỏi + A/B/C/D trên cùng một paragraph.
                    # Vì vậy phải đọc định dạng ngay cả ở paragraph mở đầu câu.
                    try:
                        signals = _grad_option_format_signals_from_paragraph(para)
                        for letter_fmt, score_fmt in signals.items():
                            current["option_format_scores"][letter_fmt] = _grad_merge_format_score(
                                current["option_format_scores"].get(letter_fmt),
                                score_fmt,
                            )
                    except Exception:
                        pass
                elif current is not None:
                    # Ghi tín hiệu định dạng của TẤT CẢ phương án A/B/C/D.
                    # Hỗ trợ tô màu/highlight, màu chữ (kể cả đỏ), gạch chân, in đậm
                    # và cả trường hợp nhiều phương án nằm chung một dòng Word.
                    try:
                        signals = _grad_option_format_signals_from_paragraph(para)
                        for letter_fmt, score_fmt in signals.items():
                            current.setdefault("option_format_scores", {})[letter_fmt] = _grad_merge_format_score(
                                current.setdefault("option_format_scores", {}).get(letter_fmt),
                                score_fmt,
                            )
                    except Exception:
                        pass

                    # Không đưa dòng HẾT vào câu cuối.
                    if "HẾT" not in txt.upper() or len(txt) > 120:
                        current["lines"].append(txt)

            if current is not None and img_paths and not in_answer_section:
                for path in img_paths:
                    current["resources"].append({
                        "loai": "anh",
                        "duong_dan": path,
                        "nguon": source_name,
                    })

        elif tag.endswith("}tbl"):
            tbl = Table(child, doc)
            rows = [
                [" ".join(str(c.text or "").split()) for c in row.cells]
                for row in tbl.rows
            ]
            if current is not None and not in_answer_section and rows:
                current["resources"].append({
                    "loai": "bang",
                    "du_lieu": rows,
                    "nguon": source_name,
                })

    finalize()

    questions = []
    for raw in questions_raw:
        q = _grad_parse_question_structure(
            raw.get("part"), raw.get("qnum"), raw.get("lines", []),
            raw.get("resources", []), answer_map, source_name
        )
        if q and str(q.get("cau_hoi", "")).strip():
            # Nếu chưa có đáp án từ bảng/dòng [[ĐÁP ÁN]], thử đọc từ định dạng Word.
            if q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn" and not str(q.get("dap_an", "")).strip():
                dap_fmt = _grad_chon_dap_an_tu_dinh_dang(
                    raw.get("option_format_scores", {})
                )
                if dap_fmt:
                    q["dap_an"] = dap_fmt
                    q["dap_an_doc_tu_dinh_dang"] = True
                    q["trang_thai"] = "Chờ rà soát"
                    q["duoc_dung_luyen_hs"] = True
                    q["san_sang_tu_dap_an_nguon"] = True
                    q["canh_bao_dap_an"] = ""
                    try:
                        q = gan_phan_loai_on_tap_cho_cau_tot_nghiep(q, force=False)
                    except Exception:
                        pass
            questions.append(q)

    return questions


def _grad_question_to_review_text(q):
    parts = [
        f"Dạng: {q.get('dang_cau','')}",
        f"Câu hỏi/ngữ liệu: {q.get('cau_hoi','')}",
    ]
    if q.get("lua_chon"):
        parts.append("Phương án:\n" + "\n".join(q.get("lua_chon") or []))
    if q.get("nhan_dinh_meta"):
        parts.append("Nhận định:\n" + "\n".join(
            f"{x.get('ky_hieu','')}) {x.get('noi_dung','')}"
            for x in (q.get("nhan_dinh_meta") or [])
        ))
        parts.append("Đáp án nguồn: " + "; ".join(
            f"{x.get('ky_hieu','')}) {x.get('dap_an','')}"
            for x in (q.get("nhan_dinh_meta") or [])
        ))
    else:
        parts.append(f"Đáp án nguồn: {q.get('dap_an','')}")
    if q.get("giai_thich"):
        parts.append("Hướng dẫn giải nguồn: " + str(q.get("giai_thich", "")))
    for res in q.get("tai_nguyen_truc_quan", []) or []:
        if res.get("loai") == "bang":
            parts.append("Bảng dữ liệu:\n" + "\n".join(
                " | ".join(map(str, row)) for row in (res.get("du_lieu") or [])
            ))
    return "\n\n".join(parts)


GRAD_REVIEW_SCHEMA = {
    "type": "object",
    "properties": {
        "ket_luan": {"type": "string"},
        "dap_an_de_xuat": {"type": "string"},
        "dap_an_4_y": {"type": "array", "items": {"type": "string"}},
        "do_tin_cay": {"type": "number"},
        "ly_do": {"type": "string"},
        "thanh_phan_nang_luc": {"type": "string"},
        "nang_luc_4_y": {"type": "array", "items": {"type": "string"}},
        "muc_do": {"type": "string"},
        "muc_do_4_y": {"type": "array", "items": {"type": "string"}},
        "huong_dan_giai_de_xuat": {"type": "string"},
        "giai_thich_4_y": {"type": "array", "items": {"type": "string"}},
    },
    "required": [
        "ket_luan", "dap_an_de_xuat", "dap_an_4_y", "do_tin_cay",
        "ly_do", "thanh_phan_nang_luc", "nang_luc_4_y",
        "muc_do", "muc_do_4_y",
        "huong_dan_giai_de_xuat", "giai_thich_4_y"
    ]
}



def ra_soat_cau_tot_nghiep_bang_ai(q):
    """
    AI chỉ thẩm định câu đề thật: kiểm tra đáp án, xác nhận năng lực và
    bổ sung lời giải khi FILE NGUỒN KHÔNG CÓ lời giải. Không viết lại câu.

    Quy tắc lưu:
    - Có lời giải nguồn: giữ nguyên 100%, AI không ghi đè.
    - Không có lời giải nguồn + đáp án khớp, độ tin cậy >= 0.80:
      lưu lời giải AI đã rà soát để HS dùng về sau (không gọi API lúc HS làm).
    - Không khớp/chưa chắc: chỉ lưu bản đề xuất trong ai_ra_soat và đánh dấu
      Cần GV xem; không dùng cho HS cho tới khi GV xác nhận.
    """
    co_loi_giai_nguon = bool(str(q.get("giai_thich", "") or "").strip()) and (
        str(q.get("nguon_giai_thich", "") or "").strip() != "AI đã rà soát"
    )

    prompt = f"""
Bạn là chuyên gia thẩm định đề tốt nghiệp THPT môn Sinh học Việt Nam.
Nhiệm vụ: KIỂM TRA câu hỏi có sẵn, tuyệt đối không viết lại câu.

{_grad_question_to_review_text(q)}

Yêu cầu:
1. Tự giải độc lập theo kiến thức Sinh học THPT và CHỈ các dữ kiện đã cho trong câu/hình/bảng.
2. So sánh với đáp án nguồn. ket_luan chỉ dùng: "Khớp", "Không khớp", "Chưa đủ dữ kiện".
3. do_tin_cay từ 0 đến 1.
4. Với câu 4 lựa chọn: dap_an_de_xuat là A/B/C/D.
5. Với Trả lời ngắn: dap_an_de_xuat là kết quả ngắn.
6. Với Đúng/Sai: dap_an_4_y phải đúng 4 phần tử "Đúng"/"Sai".
7. Xác định thành phần năng lực chính theo đúng 3 tên:
   - Nhận thức sinh học
   - Tìm hiểu thế giới sống
   - Vận dụng kiến thức, kĩ năng đã học
   Câu Đúng/Sai phải trả nang_luc_4_y đúng 4 phần tử, từng ý có thể khác nhau.
8. Xác định mức độ nhận thức thực tế theo đúng 3 mức: "Nhận biết", "Thông hiểu", "Vận dụng".
   - Câu thường: trả trong muc_do.
   - Câu Đúng/Sai: trả muc_do_4_y đúng 4 phần tử, từng ý có thể khác nhau.
   Mức độ phải dựa vào thao tác học sinh thực sự phải làm, không dựa vào độ dài câu.
9. Nếu câu phụ thuộc hình, hãy dùng chính hình được gửi kèm; nếu hình không đủ rõ thì chọn "Chưa đủ dữ kiện".
10. Không tự sửa đáp án nguồn. Chỉ báo chỗ cần GV xem.
11. HƯỚNG DẪN GIẢI:
   - Nếu câu nguồn CHƯA CÓ hướng dẫn giải, hãy tạo lời giải NGẮN GỌN, CHÍNH XÁC, đúng trình độ THPT,
     đủ để HS hiểu vì sao đáp án đúng. Không đưa kiến thức ngoài chương trình nếu câu không cung cấp.
   - Với Đúng/Sai, giai_thich_4_y phải có đúng 4 giải thích riêng a,b,c,d.
   - Với câu 4 lựa chọn/Trả lời ngắn, huong_dan_giai_de_xuat là lời giải hoàn chỉnh ngắn gọn.
   - Nếu câu nguồn ĐÃ CÓ hướng dẫn giải, không chỉnh sửa/viết lại lời giải nguồn; vẫn có thể trả trường lời giải đề xuất
     nhưng hệ thống sẽ KHÔNG ghi đè lời giải nguồn.
"""

    contents = [prompt]
    for res in q.get("tai_nguyen_truc_quan", []) or []:
        if res.get("loai") != "anh":
            continue
        path = str(res.get("duong_dan", ""))
        if not path or not os.path.exists(path):
            continue
        try:
            ext = os.path.splitext(path)[1].lower()
            mime = "image/png"
            if ext in {".jpg", ".jpeg"}:
                mime = "image/jpeg"
            elif ext == ".webp":
                mime = "image/webp"
            with open(path, "rb") as f:
                contents.append(types.Part.from_bytes(data=f.read(), mime_type=mime))
        except Exception:
            pass

    response = client.models.generate_content(
        model=MODEL_AI,
        contents=contents,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=GRAD_REVIEW_SCHEMA,
        ),
    )
    data = json.loads(response.text)

    q2 = dict(q)
    q2["ai_ra_soat"] = data
    q2["do_tin_cay_dap_an"] = float(data.get("do_tin_cay", 0) or 0)

    # Cập nhật nhãn năng lực + mức độ để phục vụ nhận xét HS, không đổi nội dung/đáp án.
    valid_nl = set(THANH_PHAN_NANG_LUC)
    valid_md = {"Nhận biết", "Thông hiểu", "Vận dụng"}
    if q2.get("dang_cau") == "Đúng / Sai":
        meta = [dict(x) for x in (q2.get("nhan_dinh_meta", []) or [])]
        nl4 = list(data.get("nang_luc_4_y", []) or [])
        md4 = list(data.get("muc_do_4_y", []) or [])
        rank_md = {"Nhận biết": 1, "Thông hiểu": 2, "Vận dụng": 3}
        md_max = "Nhận biết"
        for i, nd in enumerate(meta):
            if i < len(nl4) and nl4[i] in valid_nl:
                nd["thanh_phan_nang_luc"] = nl4[i]
            if i < len(md4) and md4[i] in valid_md:
                nd["muc_do"] = md4[i]
            if rank_md.get(nd.get("muc_do"), 0) > rank_md.get(md_max, 0):
                md_max = nd.get("muc_do")
        q2["nhan_dinh_meta"] = meta
        q2["muc_do"] = md_max
    else:
        nl = str(data.get("thanh_phan_nang_luc", "")).strip()
        if nl in valid_nl:
            q2["thanh_phan_nang_luc"] = nl
        md = str(data.get("muc_do", "")).strip()
        if md in valid_md:
            q2["muc_do"] = md

    ket_luan = str(data.get("ket_luan", "")).strip()
    confidence = q2["do_tin_cay_dap_an"]

    # Luôn lưu bản lời giải AI đề xuất ở vùng thẩm định để GV có thể xem,
    # nhưng chỉ đưa thành lời giải chuẩn khi đáp án khớp và đủ tin cậy.
    if q2.get("dang_cau") == "Đúng / Sai":
        q2["giai_thich_ai_de_xuat"] = list(data.get("giai_thich_4_y", []) or [])
    else:
        q2["giai_thich_ai_de_xuat"] = str(data.get("huong_dan_giai_de_xuat", "") or "").strip()

    if ket_luan == "Khớp" and confidence >= 0.80:
        q2["trang_thai"] = "Đã rà soát"
        q2["canh_bao_dap_an"] = ""
        q2["duoc_dung_luyen_hs"] = True

        # Chỉ bổ sung lời giải khi file nguồn không có. Không bao giờ ghi đè lời giải nguồn.
        if not co_loi_giai_nguon:
            if q2.get("dang_cau") == "Đúng / Sai":
                meta = [dict(x) for x in (q2.get("nhan_dinh_meta", []) or [])]
                gt4 = list(data.get("giai_thich_4_y", []) or [])
                for i, nd in enumerate(meta):
                    if i < len(gt4) and str(gt4[i] or "").strip() and not str(nd.get("giai_thich", "") or "").strip():
                        nd["giai_thich"] = str(gt4[i]).strip()
                        nd["nguon_giai_thich"] = "AI đã rà soát"
                q2["nhan_dinh_meta"] = meta
                # Tạo bản tóm tắt chung để AI cố vấn/hệ thống cũ vẫn có nội dung giải thích.
                if not str(q2.get("giai_thich", "") or "").strip():
                    tong = []
                    for nd in meta:
                        if str(nd.get("giai_thich", "") or "").strip():
                            tong.append(f"{nd.get('ky_hieu','')}) {nd.get('giai_thich','')}")
                    q2["giai_thich"] = "\n".join(tong).strip()
            else:
                solution = str(data.get("huong_dan_giai_de_xuat", "") or "").strip()
                if solution and not str(q2.get("giai_thich", "") or "").strip():
                    q2["giai_thich"] = solution
            if str(q2.get("giai_thich", "") or "").strip() or any(
                str(x.get("giai_thich", "") or "").strip()
                for x in (q2.get("nhan_dinh_meta", []) or [])
            ):
                q2["nguon_giai_thich"] = "AI đã rà soát"
        else:
            q2["nguon_giai_thich"] = q2.get("nguon_giai_thich") or "Lời giải từ file nguồn"
    else:
        q2["trang_thai"] = "Cần GV xem"
        q2["canh_bao_dap_an"] = str(data.get("ly_do", "")).strip() or ket_luan
        q2["duoc_dung_luyen_hs"] = False

    # Giữ/hoàn thiện phân loại bài-chương sau rà soát; không ép YCCĐ chi tiết.
    q2 = gan_phan_loai_on_tap_cho_cau_tot_nghiep(q2, force=False)
    q2["ngay_ra_soat_ai"] = datetime.now().strftime("%d/%m/%Y %H:%M")
    return q2



def hien_thi_tai_nguyen_cau_tot_nghiep(q):
    resources = q.get("tai_nguyen_truc_quan", []) or []
    if not resources:
        # Tương thích dữ liệu cũ.
        data = q.get("du_lieu_truc_quan", {}) or {}
        if data:
            hien_thi_du_lieu_truc_quan_cau(q)
        return

    for res in resources:
        if res.get("loai") == "anh":
            _hien_thi_anh_tu_resource(
                res,
                width=GRAD_IMAGE_DISPLAY_WIDTH
            )
        elif res.get("loai") == "bang":
            rows = res.get("du_lieu") or []
            if rows:
                width = max(len(r) for r in rows)
                padded = [list(r) + [""] * (width - len(r)) for r in rows]
                headers = padded[0]
                # Tránh tên cột trùng nhau làm pandas lỗi/khó hiển thị.
                seen = {}
                cols = []
                for j, h in enumerate(headers):
                    base = str(h or f"Cột {j+1}")
                    seen[base] = seen.get(base, 0) + 1
                    cols.append(base if seen[base] == 1 else f"{base} ({seen[base]})")
                df = pd.DataFrame(padded[1:], columns=cols) if len(padded) > 1 else pd.DataFrame(columns=cols)
                st.dataframe(df, use_container_width=True, hide_index=True)


def cau_on_tap_bo_sung_du_dieu_kien_tot_nghiep(q):
    """
    Câu từ Ngân hàng ôn tập chỉ được dùng như nguồn BỔ SUNG nhỏ cho đề tốt nghiệp.
    Ưu tiên an toàn: câu đã duyệt, không ngừng sử dụng, thuộc Khối 12, đúng 1 trong 3 dạng
    và có đáp án đầy đủ. Không thay thế vai trò chính của Ngân hàng tốt nghiệp.
    """
    if not isinstance(q, dict):
        return False
    if q.get("trang_thai", "Đã duyệt") == "Ngừng sử dụng":
        return False
    if q.get("duoc_dung_luyen_hs", True) is False:
        return False
    if str(q.get("khoi", "")).strip() not in {"Khối 12", "12", "Lớp 12"}:
        return False
    if q.get("dang_cau") not in {"Trắc nghiệm 4 lựa chọn", "Đúng / Sai", "Trả lời ngắn"}:
        return False
    return _grad_co_dap_an_day_du(q)


def _grad_fingerprint_noi_dung_chung(q):
    raw = {
        "dang": q.get("dang_cau", ""),
        "cau_hoi": _grad_norm_text(q.get("cau_hoi", "")),
        "tinh_huong": _grad_norm_text(q.get("tinh_huong", "")),
        "lua_chon": [_grad_norm_text(x) for x in (q.get("lua_chon", []) or [])],
        "nhan_dinh": [
            _grad_norm_text(x.get("noi_dung", ""))
            for x in (q.get("nhan_dinh_meta", []) or [])
        ],
    }
    return hashlib.sha256(
        json.dumps(raw, ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()


def _grad_round_robin_sample(items, n, rng):
    # Khử trùng trước khi rút.
    unique = {}
    for q in items:
        unique.setdefault(fingerprint_cau_tot_nghiep_thuc_te(q), q)
    items = list(unique.values())

    by_source = {}
    for q in items:
        by_source.setdefault(str(q.get("nguon_file", "Không rõ")), []).append(q)
    for arr in by_source.values():
        rng.shuffle(arr)
    sources = list(by_source.keys())
    rng.shuffle(sources)

    out = []
    while sources and len(out) < n:
        next_sources = []
        for src in sources:
            arr = by_source[src]
            if arr and len(out) < n:
                out.append(arr.pop())
            if arr:
                next_sources.append(src)
        sources = next_sources
    return out[:n]


def rut_de_tot_nghiep_tu_de_that(
    bank,
    seed=None,
    bank_on_tap=None,
    toi_da_on_tap=GRAD_MAX_CAU_TU_NH_ON_TAP
):
    """
    Rút cứng form 18/4/6.
    - Nguồn chính: Ngân hàng tốt nghiệp.
    - Nguồn bổ sung: tối đa một ít câu Khối 12 từ Ngân hàng ôn tập.
    - Không trùng nội dung trong cùng mã đề.
    - Vẫn ưu tiên trộn nhiều file nguồn ở phần Ngân hàng tốt nghiệp.
    """
    rng = random.Random(seed)
    specs = {
        "Trắc nghiệm 4 lựa chọn": 18,
        "Đúng / Sai": 4,
        "Trả lời ngắn": 6,
    }
    quota_on_tap = dict(GRAD_QUOTA_ON_TAP_THEO_DANG)
    toi_da_on_tap = max(0, min(int(toi_da_on_tap or 0), GRAD_MAX_CAU_TU_NH_ON_TAP))

    selected = []
    missing = []
    used_fp = set()
    so_on_tap_da_dung = 0

    for dang, n in specs.items():
        picked_all = []

        # 1) Chỉ chen một lượng nhỏ câu phù hợp từ ngân hàng ôn tập.
        quota_dang = min(
            int(quota_on_tap.get(dang, 0) or 0),
            max(0, toi_da_on_tap - so_on_tap_da_dung),
            n,
        )
        pool_on_tap = [
            q for q in (bank_on_tap or [])
            if q.get("dang_cau") == dang
            and cau_on_tap_bo_sung_du_dieu_kien_tot_nghiep(q)
        ]
        rng.shuffle(pool_on_tap)
        for q in pool_on_tap:
            fp = _grad_fingerprint_noi_dung_chung(q)
            if fp in used_fp:
                continue
            picked_all.append(q)
            used_fp.add(fp)
            so_on_tap_da_dung += 1
            if len(picked_all) >= quota_dang:
                break

        # 2) Phần còn lại lấy từ ngân hàng tốt nghiệp.
        can_grad = n - len(picked_all)
        pool_grad = [
            q for q in (bank or [])
            if q.get("dang_cau") == dang
            and cau_tot_nghiep_du_dieu_kien_su_dung(q)
            and _grad_fingerprint_noi_dung_chung(q) not in used_fp
        ]
        picked_grad = _grad_round_robin_sample(pool_grad, can_grad, rng)
        for q in picked_grad:
            fp = _grad_fingerprint_noi_dung_chung(q)
            if fp in used_fp:
                continue
            picked_all.append(q)
            used_fp.add(fp)
            if len(picked_all) >= n:
                break

        selected.extend(picked_all[:n])
        if len(picked_all) < n:
            missing.append({
                "Dạng": dang,
                "Cần": n,
                "Có thể rút": len(picked_all),
                "Thiếu": n - len(picked_all),
            })

    # Theo đúng thứ tự 3 phần.
    order = {"Trắc nghiệm 4 lựa chọn": 0, "Đúng / Sai": 1, "Trả lời ngắn": 2}
    selected.sort(key=lambda q: order.get(q.get("dang_cau"), 9))
    return selected, missing


def tao_file_mau_ngan_hang_tot_nghiep():
    """Tạo Word mẫu tối giản để GV nhập đề có đáp án/hướng dẫn giải."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MẪU NHẬP NGÂN HÀNG TỐT NGHIỆP – SINH HỌC")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph(
        "Hướng dẫn: Giữ đúng 3 phần. Ảnh/sơ đồ/bảng có thể chèn ngay trong câu. "
        "Đáp án 4 lựa chọn có thể ghi bằng [[ĐÁP ÁN]] hoặc đánh dấu DUY NHẤT phương án đúng "
        "bằng tô nền/highlight, màu chữ (kể cả đỏ), gạch chân hoặc in đậm. "
        "[[HƯỚNG DẪN GIẢI]] là tùy chọn. Khi quản lý ngân hàng, GV chỉ cần xác nhận "
        "Khối – Chương – Mức độ – Thành phần năng lực; không bắt buộc gán Bài/YCCĐ chi tiết."
    )

    doc.add_heading("PHẦN I. Trắc nghiệm 4 lựa chọn", level=1)
    doc.add_paragraph("Câu 1: Nội dung câu hỏi mẫu. Có thể chèn hình/bảng ngay dưới đoạn này.")
    doc.add_paragraph("A. Phương án A")
    p_b = doc.add_paragraph()
    r_b = p_b.add_run("B. Phương án B (ví dụ đáp án được in đậm/màu đỏ)")
    r_b.bold = True
    try:
        from docx.shared import RGBColor
        r_b.font.color.rgb = RGBColor(220, 38, 38)
    except Exception:
        pass
    doc.add_paragraph("C. Phương án C")
    doc.add_paragraph("D. Phương án D")
    doc.add_paragraph("[[ĐÁP ÁN]] B  (có thể bỏ dòng này nếu đã đánh dấu B bằng định dạng nổi bật)")
    doc.add_paragraph("[[HƯỚNG DẪN GIẢI]] Viết lời giải/nguyên nhân chọn đáp án B.")

    doc.add_heading("PHẦN II. Đúng / Sai", level=1)
    doc.add_paragraph("Câu 1: Ngữ liệu/tình huống chung của câu Đúng/Sai.")
    doc.add_paragraph("a) Nhận định a.")
    doc.add_paragraph("b) Nhận định b.")
    doc.add_paragraph("c) Nhận định c.")
    doc.add_paragraph("d) Nhận định d.")
    doc.add_paragraph("[[ĐÁP ÁN]] SDDD")
    doc.add_paragraph(
        "[[HƯỚNG DẪN GIẢI]] a) Giải thích ý a.\n"
        "b) Giải thích ý b.\n"
        "c) Giải thích ý c.\n"
        "d) Giải thích ý d."
    )

    doc.add_heading("PHẦN III. Trả lời ngắn", level=1)
    doc.add_paragraph("Câu 1: Nội dung câu hỏi trả lời ngắn; đáp án tối đa 4 kí tự theo form app.")
    doc.add_paragraph("[[ĐÁP ÁN]] 0,25")
    doc.add_paragraph("[[HƯỚNG DẪN GIẢI]] Trình bày phép tính/lập luận dẫn tới 0,25.")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()




def _grad_sources_of_question(q):
    sources = list(q.get("nguon_files", []) or [])
    primary = str(q.get("nguon_file", "") or "").strip()
    if primary and primary not in sources:
        sources.insert(0, primary)
    return [str(x).strip() for x in sources if str(x).strip()]


def _grad_status_group_match(q, status_filter):
    """Bộ lọc trạng thái dùng chung cho Kho câu và khu vực rà soát."""
    status = str(q.get("trang_thai", "") or "").strip()
    if status_filter in {"", "Tất cả", None}:
        return True
    # Nút "Cần GV xem" phải khớp đúng với con số tổng quan:
    # gồm cả câu AI cảnh báo và câu chưa đọc được đáp án.
    if status_filter in {"Cần GV xem", "⚠️ Cần GV xem"}:
        return status in {"Cần GV xem", "Thiếu đáp án"}
    return status == status_filter


def _grad_label_question(q):
    src = str(q.get("nguon_file", "") or "Không rõ nguồn")
    part = str(q.get("phan_goc", "") or "")
    num = str(q.get("so_cau_goc", "") or "")
    dang = str(q.get("dang_cau", "") or "")
    khoi = str(q.get("khoi", "") or "")
    bai = str(q.get("bai", "") or q.get("chuong", "") or "Chưa phân loại")
    return f"{src} • P.{part} Câu {num} • {dang} • {khoi} • {bai}"


def _grad_save_managed_question(original, edited):
    bank2 = doc_ngan_hang_tot_nghiep_thuc_te()
    q2 = dict(edited)

    # GV đã chủ động chỉnh phạm vi nên khóa phân loại này, tránh lần đọc sau tự ghi đè.
    pv = dict(q2.get("phan_loai_on_tap", {}) or {})
    pv.update({
        "khoi": q2.get("khoi", ""),
        "chuong": q2.get("chuong", ""),
        "bai": q2.get("bai", ""),
        "muc": "Bài" if str(q2.get("bai", "")).strip() else "Chương",
        "do_tin_cay": 1.0,
        "ghi_chu": "GV xác nhận phạm vi Khối → Chương → Bài.",
        "phien_ban": "khoi_chuong_bai_v3_ai_candidate",
        "gv_xac_nhan": True,
    })
    q2["phan_loai_on_tap"] = pv
    q2["do_tin_cay_pham_vi_kien_thuc"] = 1.0

    # Nhãn nội bộ phục vụ hồ sơ HS; không giả là YCCĐ chi tiết chính thức.
    scope_name = str(q2.get("bai", "") or q2.get("chuong", "") or "").strip()
    if scope_name:
        q2["yccd"] = "Kiến thức trọng tâm • " + " • ".join(
            x for x in [str(q2.get("khoi", "") or "").strip(), scope_name] if x
        )

    if _grad_co_dap_an_day_du(q2):
        q2["trang_thai"] = "GV đã duyệt"
        q2["duoc_dung_luyen_hs"] = True
        q2["san_sang_tu_dap_an_nguon"] = True
        q2["canh_bao_dap_an"] = ""
    else:
        q2["trang_thai"] = "Thiếu đáp án"
        q2["duoc_dung_luyen_hs"] = False
        q2["san_sang_tu_dap_an_nguon"] = False
        q2["canh_bao_dap_an"] = "GV chưa bổ sung đủ đáp án."

    q2["ngay_gv_duyet"] = datetime.now().strftime("%d/%m/%Y %H:%M")

    _grad_ghi_nho_pham_vi_gv(q2, q2.get("khoi", ""), q2.get("chuong", ""), q2.get("bai", ""))

    found = False
    for i, item in enumerate(bank2):
        if item.get("id") == original.get("id"):
            bank2[i] = q2
            found = True
            break
    if not found:
        return False

    luu_ngan_hang_tot_nghiep_thuc_te(bank2)
    return True


def _grad_cap_nhat_pham_vi_gv(original, grade, chapter, lesson, ap_dung_4_y=True):
    """
    Chỉ cập nhật phạm vi ôn tập Khối → Chương → Bài do GV xác nhận.
    Không đổi đáp án, lời giải, trạng thái rà soát hay nội dung câu.
    """
    bank2 = doc_ngan_hang_tot_nghiep_thuc_te()
    q2 = dict(original)

    grade = str(grade or "").strip()
    chapter = str(chapter or "").strip()
    lesson = str(lesson or "").strip()

    q2["khoi"] = grade
    q2["chuong"] = chapter
    q2["bai"] = lesson
    q2["do_tin_cay_pham_vi_kien_thuc"] = 1.0
    q2["phan_loai_on_tap"] = {
        "khoi": grade,
        "chuong": chapter,
        "bai": lesson,
        "muc": "Bài" if lesson else "Chương",
        "do_tin_cay": 1.0,
        "ghi_chu": "GV sửa và xác nhận phạm vi Khối → Chương → Bài.",
        "phien_ban": "khoi_chuong_bai_v3_ai_candidate",
        "gv_xac_nhan": True,
    }
    q2["ngay_gv_xac_nhan_pham_vi"] = datetime.now().strftime("%d/%m/%Y %H:%M")

    # Chỉ tạo nhãn theo dõi nội bộ; không giả đây là YCCĐ chi tiết chính thức.
    scope_name = lesson or chapter
    yccd_old = str(q2.get("yccd", "") or "").strip()
    if scope_name and (not yccd_old or yccd_old.startswith("Kiến thức trọng tâm")):
        q2["yccd"] = "Kiến thức trọng tâm • " + " • ".join(
            x for x in [grade, scope_name] if x
        )

    # Câu Đúng/Sai có metadata từng ý. Khi GV muốn, áp dụng cùng phạm vi
    # cho cả 4 ý nhưng vẫn giữ nguyên mức độ/năng lực/đáp án riêng từng ý.
    if q2.get("dang_cau") == "Đúng / Sai" and ap_dung_4_y:
        meta = []
        for nd in list(q2.get("nhan_dinh_meta", []) or []):
            nd2 = dict(nd)
            nd2["khoi"] = grade
            nd2["chuong"] = chapter
            nd2["bai"] = lesson
            nd2["do_tin_cay_pham_vi_kien_thuc"] = 1.0
            nd2["phan_loai_on_tap"] = {
                "khoi": grade,
                "chuong": chapter,
                "bai": lesson,
                "muc": "Bài" if lesson else "Chương",
                "do_tin_cay": 1.0,
                "phien_ban": "khoi_chuong_bai_v3_ai_candidate",
                "gv_xac_nhan": True,
            }
            nd_scope = lesson or chapter
            nd_yccd_old = str(nd2.get("yccd", "") or "").strip()
            if nd_scope and (not nd_yccd_old or nd_yccd_old.startswith("Kiến thức trọng tâm")):
                nd2["yccd"] = "Kiến thức trọng tâm • " + " • ".join(
                    x for x in [grade, nd_scope] if x
                )
            meta.append(nd2)
        q2["nhan_dinh_meta"] = meta

    _grad_ghi_nho_pham_vi_gv(q2, grade, chapter, lesson)

    found = False
    for i, item in enumerate(bank2):
        if item.get("id") == original.get("id"):
            bank2[i] = q2
            found = True
            break

    if not found:
        return False

    luu_ngan_hang_tot_nghiep_thuc_te(bank2)
    return True


def hien_thi_kho_cau_tot_nghiep_da_nhap(bank):
    """
    Kho quản trị riêng cho các câu GV đã nhập từ đề thật.
    Chỉ thao tác trên GRAD_REAL_BANK_PATH; không đụng ngân hàng khác.
    """
    st.markdown("### 📚 Kho câu đã nhập")
    st.caption(
        "Các câu đã tách từ đề thật được **lưu lâu dài trong ngân hàng tốt nghiệp**. "
        "GV có thể lọc → chọn một câu → xem đầy đủ hình/bảng → sửa → xóa. "
        "GV chỉ cần xác nhận **Khối – Chương – Mức độ – Thành phần năng lực**; app tự giữ dạng câu, đáp án, nguồn và tài nguyên trực quan."
    )

    if not bank:
        st.info("Kho câu đã nhập đang trống.")
        return

    sources = sorted({src for q in bank for src in _grad_sources_of_question(q)})
    c1, c2, c3, c4 = st.columns([1.5, 1.2, 1.2, 1.8])
    with c1:
        src_filter = st.selectbox(
            "Nguồn",
            ["Tất cả"] + sources,
            key="grad_bank_manage_source",
        )
    with c2:
        type_filter = st.selectbox(
            "Dạng câu",
            ["Tất cả", "Trắc nghiệm 4 lựa chọn", "Đúng / Sai", "Trả lời ngắn"],
            key="grad_bank_manage_type",
        )
    with c3:
        status_filter = st.selectbox(
            "Trạng thái",
            ["Tất cả", "Cần GV xem", "Chờ rà soát", "Đã rà soát", "GV đã duyệt", "Thiếu đáp án"],
            key="grad_bank_manage_status",
        )
    with c4:
        search_text = st.text_input(
            "Tìm câu / chương",
            value="",
            placeholder="Nhập từ khóa...",
            key="grad_bank_manage_search",
        )

    search_norm = _grad_norm_text(search_text)
    ds = []
    for q in bank:
        if src_filter != "Tất cả" and src_filter not in _grad_sources_of_question(q):
            continue
        if type_filter != "Tất cả" and q.get("dang_cau") != type_filter:
            continue
        if not _grad_status_group_match(q, status_filter):
            continue
        if search_norm:
            hay = _grad_norm_text(" ".join([
                str(q.get("cau_hoi", "")),
                str(q.get("khoi", "")),
                str(q.get("chuong", "")),
                str(q.get("nguon_file", "")),
            ]))
            if search_norm not in hay:
                continue
        ds.append(q)

    rows = []
    for q in ds[:500]:
        pv = q.get("phan_loai_on_tap", {}) or {}
        rows.append({
            "Nguồn": q.get("nguon_file", ""),
            "Phần/Câu": f"{q.get('phan_goc','')}/{q.get('so_cau_goc','')}",
            "Dạng": q.get("dang_cau", ""),
            "Khối": q.get("khoi", ""),
            "Chương": q.get("chuong", ""),
            "Mức độ": q.get("muc_do", ""),
            "Năng lực": q.get("thanh_phan_nang_luc", ""),
            "Đáp án": "Đủ" if _grad_co_dap_an_day_du(q) else "Thiếu",
            "Trạng thái": q.get("trang_thai", ""),
            "Ôn HS": "Có" if q.get("duoc_dung_luyen_hs", False) else "Không",
        })

    st.caption(f"Tìm thấy **{len(ds)}** câu.")
    if rows:
        st.dataframe(
            pd.DataFrame(rows),
            use_container_width=True,
            hide_index=True,
            height=min(420, 80 + 35 * min(len(rows), 10)),
        )

    if not ds:
        st.info("Không có câu phù hợp bộ lọc này.")
        return

    ids = [str(q.get("id", "")) for q in ds]
    by_id = {str(q.get("id", "")): q for q in ds}
    selected_id = st.selectbox(
        "Chọn câu để xem / chỉnh sửa",
        ids,
        format_func=lambda x: _grad_label_question(by_id[x]),
        key="grad_bank_manage_select",
    )
    q = by_id[selected_id]
    qkey = re.sub(r"[^A-Za-z0-9]", "", selected_id)[:10] or "q"

    with st.container(border=True):
        st.write(f"**Nguồn:** {q.get('nguon_file','')} • **Phần {q.get('phan_goc','')} – Câu {q.get('so_cau_goc','')}**")
        st.caption(
            "Phạm vi: " + str(q.get("khoi", "") or "Chưa xác định")
            + (" → " + str(q.get("chuong", "")) if str(q.get("chuong", "")).strip() else "")
            + " • " + str(q.get("muc_do", "") or "Chưa xác định mức độ")
            + " • " + str(q.get("thanh_phan_nang_luc", "") or "Chưa xác định năng lực")
        )
        st.markdown("**Nội dung:**")
        st.write(q.get("cau_hoi", ""))
        hien_thi_tai_nguyen_cau_tot_nghiep(q)
        if q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn":
            for lc in q.get("lua_chon", []) or []:
                st.write(lc)
        elif q.get("dang_cau") == "Đúng / Sai":
            for nd in q.get("nhan_dinh_meta", []) or []:
                st.write(f"**{nd.get('ky_hieu','')})** {nd.get('noi_dung','')}")
        if q.get("dang_cau") == "Đúng / Sai":
            st.caption("Đáp án: " + "; ".join(
                f"{nd.get('ky_hieu','')}) {nd.get('dap_an','—')}"
                for nd in (q.get("nhan_dinh_meta", []) or [])
            ))
        else:
            st.caption(f"Đáp án: {q.get('dap_an','—')}")

    with st.expander("✏️ Sửa câu đang chọn", expanded=False):
        q_edit = dict(q)

        # Phạm vi Khối → Chương → Bài chọn từ chính KHO_YCCD.
        grades = list((KHO_YCCD or {}).keys())
        if not grades:
            grades = [str(q.get("khoi", "") or "")]
        grade0 = q.get("khoi", "") if q.get("khoi", "") in grades else grades[0]
        grade = st.selectbox(
            "Khối",
            grades,
            index=grades.index(grade0),
            key=f"grad_manage_grade_{qkey}",
        )
        chapters = list((KHO_YCCD.get(grade, {}) or {}).keys())
        if not chapters:
            chapters = [str(q.get("chuong", "") or "")]
        chapter0 = q.get("chuong", "") if q.get("chuong", "") in chapters else chapters[0]
        chapter = st.selectbox(
            "Chương",
            chapters,
            index=chapters.index(chapter0),
            key=f"grad_manage_chapter_{qkey}_{grade}",
        )
        # Ngân hàng tốt nghiệp chỉ yêu cầu đến mức Chương để GV thao tác nhanh.
        lesson_value = ""

        cmeta1, cmeta2 = st.columns(2)
        with cmeta1:
            muc_options = ["Nhận biết", "Thông hiểu", "Vận dụng"]
            muc0 = q.get("muc_do", "") if q.get("muc_do", "") in muc_options else "Nhận biết"
            muc_do = st.selectbox(
                "Mức độ",
                muc_options,
                index=muc_options.index(muc0),
                key=f"grad_manage_level_{qkey}",
            )
        with cmeta2:
            nl0 = q.get("thanh_phan_nang_luc", "") if q.get("thanh_phan_nang_luc", "") in THANH_PHAN_NANG_LUC else THANH_PHAN_NANG_LUC[0]
            nang_luc = st.selectbox(
                "Thành phần năng lực",
                THANH_PHAN_NANG_LUC,
                index=THANH_PHAN_NANG_LUC.index(nl0),
                key=f"grad_manage_comp_{qkey}",
            )

        cau_hoi = st.text_area(
            "Nội dung câu hỏi / ngữ liệu",
            value=str(q.get("cau_hoi", "")),
            height=150,
            key=f"grad_manage_question_{qkey}",
        )

        if q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn":
            opts_old = list(q.get("lua_chon", []) or [])
            while len(opts_old) < 4:
                opts_old.append("")
            opts_new = []
            for i, letter in enumerate(["A", "B", "C", "D"]):
                txt = st.text_input(
                    f"Phương án {letter}",
                    value=bo_nhan_phuong_an(opts_old[i]),
                    key=f"grad_manage_opt_{qkey}_{letter}",
                )
                opts_new.append(f"{letter}. {txt.strip()}")
            ans0 = str(q.get("dap_an", "") or "A").upper()
            if ans0 not in ["A", "B", "C", "D"]:
                ans0 = "A"
            dap_an = st.selectbox(
                "Đáp án đúng",
                ["A", "B", "C", "D"],
                index=["A", "B", "C", "D"].index(ans0),
                key=f"grad_manage_answer_{qkey}",
            )
            q_edit["lua_chon"] = opts_new
            q_edit["dap_an"] = dap_an

        elif q.get("dang_cau") == "Đúng / Sai":
            meta = [dict(x) for x in (q.get("nhan_dinh_meta", []) or [])]
            while len(meta) < 4:
                meta.append({"ky_hieu": "abcd"[len(meta)], "noi_dung": "", "dap_an": "", "muc_do": muc_do, "thanh_phan_nang_luc": nang_luc})
            for i, nd in enumerate(meta[:4]):
                ky = "abcd"[i]
                nd["ky_hieu"] = ky
                nd["noi_dung"] = st.text_area(
                    f"Ý {ky}",
                    value=str(nd.get("noi_dung", "")),
                    height=80,
                    key=f"grad_manage_tf_text_{qkey}_{ky}",
                )
                ans_old = str(nd.get("dap_an", "") or "Đúng")
                nd["dap_an"] = st.selectbox(
                    f"Đáp án ý {ky}",
                    ["Đúng", "Sai"],
                    index=0 if ans_old == "Đúng" else 1,
                    key=f"grad_manage_tf_ans_{qkey}_{ky}",
                )
                # Cho phép sửa riêng mức độ/năng lực từng ý vì Đ/S là câu tích hợp.
                cta, ctb = st.columns(2)
                with cta:
                    md0 = nd.get("muc_do", "") if nd.get("muc_do", "") in muc_options else muc_do
                    nd["muc_do"] = st.selectbox(
                        f"Mức độ ý {ky}",
                        muc_options,
                        index=muc_options.index(md0),
                        key=f"grad_manage_tf_level_{qkey}_{ky}",
                    )
                with ctb:
                    n0 = nd.get("thanh_phan_nang_luc", "") if nd.get("thanh_phan_nang_luc", "") in THANH_PHAN_NANG_LUC else nang_luc
                    nd["thanh_phan_nang_luc"] = st.selectbox(
                        f"Năng lực ý {ky}",
                        THANH_PHAN_NANG_LUC,
                        index=THANH_PHAN_NANG_LUC.index(n0),
                        key=f"grad_manage_tf_comp_{qkey}_{ky}",
                    )
                nd["khoi"] = grade
                nd["chuong"] = chapter
                nd["bai"] = lesson_value
                nd_scope = str(chapter or "").strip()
                if nd_scope:
                    nd["yccd"] = "Kiến thức trọng tâm • " + " • ".join(
                        x for x in [grade, nd_scope] if str(x).strip()
                    )
                nd["phan_loai_on_tap"] = {
                    "khoi": grade,
                    "chuong": chapter,
                    "bai": lesson_value,
                    "muc": "Chương",
                    "do_tin_cay": 1.0,
                    "phien_ban": "khoi_chuong_bai_v3_ai_candidate",
                    "gv_xac_nhan": True,
                }
            q_edit["nhan_dinh_meta"] = meta[:4]

        else:
            q_edit["dap_an"] = st.text_input(
                "Đáp án",
                value=str(q.get("dap_an", "")),
                key=f"grad_manage_short_ans_{qkey}",
            )

        giai_thich = st.text_area(
            "Hướng dẫn giải / ghi chú",
            value=str(q.get("giai_thich", "")),
            height=130,
            key=f"grad_manage_solution_{qkey}",
        )

        q_edit["khoi"] = grade
        q_edit["chuong"] = chapter
        q_edit["bai"] = ""
        q_edit["yccd"] = "Kiến thức trọng tâm • " + " • ".join(
            x for x in [grade, chapter] if str(x).strip()
        )
        q_edit["muc_do"] = muc_do
        q_edit["thanh_phan_nang_luc"] = nang_luc
        q_edit["cau_hoi"] = cau_hoi.strip()
        q_edit["giai_thich"] = giai_thich.strip()
        if q_edit["giai_thich"] and q_edit["giai_thich"] != str(q.get("giai_thich", "")):
            q_edit["nguon_giai_thich"] = "GV xác nhận / bổ sung"

        if st.button(
            "💾 LƯU CHỈNH SỬA CÂU",
            type="primary",
            use_container_width=True,
            key=f"grad_manage_save_{qkey}",
        ):
            if _grad_save_managed_question(q, q_edit):
                st.success("Đã lưu chỉnh sửa vào Kho câu tốt nghiệp.")
                st.rerun()
            else:
                st.error("Không tìm thấy câu để cập nhật.")

    with st.expander("🗑️ Xóa câu đang chọn", expanded=False):
        st.warning(
            "Thao tác này chỉ xóa **câu đang chọn** khỏi ngân hàng; file nguồn và các câu khác không bị xóa."
        )
        confirm = st.checkbox(
            "Tôi xác nhận xóa câu này",
            key=f"grad_manage_delete_confirm_{qkey}",
        )
        if st.button(
            "🗑️ XÓA CÂU KHỎI NGÂN HÀNG",
            use_container_width=True,
            disabled=not confirm,
            key=f"grad_manage_delete_{qkey}",
        ):
            bank2 = [item for item in doc_ngan_hang_tot_nghiep_thuc_te() if item.get("id") != q.get("id")]
            luu_ngan_hang_tot_nghiep_thuc_te(bank2)
            st.success("Đã xóa câu khỏi ngân hàng tốt nghiệp.")
            st.rerun()


def xay_dung_ngan_hang_tot_nghiep():
    st.header("🎓 NGÂN HÀNG TỐT NGHIỆP – ĐỀ THẬT / ĐỀ THI THỬ")
    st.caption(
        "Phần này **không dùng AI để sáng tác câu mới**. GV đưa vào các đề thật/đề thi thử "
        "đúng form, **có đáp án** (hướng dẫn giải có thể có hoặc chưa có); app giữ nguyên câu, bảng, hình/sơ đồ, "
        "trộn câu từ nhiều đề để tạo mã đề mới. GV chỉ cần quản lý **Khối – Chương – Mức độ – Thành phần năng lực**; "
        "AI có thể gợi ý/kiểm tra các trường này và bổ sung lời giải khi thiếu."
    )

    st.info(
        "Form chuẩn dùng để luyện: **18 câu 4 lựa chọn + 4 câu Đúng/Sai + 6 câu Trả lời ngắn**. "
        "Các câu trùng nội dung chỉ lưu một lần."
    )

    # ------------------------------------------------------
    # MẪU FILE
    # ------------------------------------------------------
    ctpl1, ctpl2 = st.columns([2, 3])
    with ctpl1:
        st.download_button(
            "⬇️ TẢI FILE WORD MẪU",
            data=tao_file_mau_ngan_hang_tot_nghiep(),
            file_name="Mau_nhap_ngan_hang_tot_nghiep_Sinh_hoc.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="grad_real_download_template",
        )
    with ctpl2:
        st.caption(
            "Có thể nhập trực tiếp đề Word hiện có. Chỉ cần có đáp án; hướng dẫn giải không bắt buộc. "
            "Đáp án 4 lựa chọn có thể nằm trong bảng/dòng đáp án hoặc được đánh dấu DUY NHẤT bằng "
            "**tô nền/highlight, màu chữ (kể cả đỏ), gạch chân, in đậm**. "
            "Nếu file Word phức tạp hoặc muốn app đọc chắc chắn nhất, dùng file mẫu. "
            "Dòng [[HƯỚNG DẪN GIẢI]] có thể bỏ nếu đề nguồn chưa có lời giải."
        )

    st.divider()
    st.subheader("① Nhập đề vào ngân hàng")
    files = st.file_uploader(
        "Chọn một hoặc nhiều đề Word",
        type=["docx"],
        accept_multiple_files=True,
        key="grad_real_upload",
        help="Word cho phép app giữ ảnh/bảng gắn với từng câu chính xác hơn PDF.",
    )

    if files:
        st.caption(f"Đã chọn **{len(files)} file**. Chưa nhập vào ngân hàng cho đến khi bấm nút bên dưới.")

    if st.button(
        "📥 TÁCH & NHẬP ĐỀ VÀO NGÂN HÀNG TỐT NGHIỆP",
        type="primary",
        use_container_width=True,
        disabled=not bool(files),
        key="grad_real_import",
    ):
        bank = doc_ngan_hang_tot_nghiep_thuc_te()
        by_fp = {fingerprint_cau_tot_nghiep_thuc_te(q): i for i, q in enumerate(bank)}
        total_new = 0
        total_dup = 0
        reports = []

        for uploaded in files or []:
            try:
                source_name = os.path.basename(uploaded.name)
                safe_name = re.sub(r"[^0-9A-Za-zÀ-ỹ_.() -]+", "_", source_name)
                save_path = os.path.join(GRAD_SOURCE_DIR, safe_name)
                with open(save_path, "wb") as f:
                    f.write(uploaded.getbuffer())

                qs = tach_de_tot_nghiep_docx(save_path, source_name)
                counts = {
                    "Trắc nghiệm 4 lựa chọn": 0,
                    "Đúng / Sai": 0,
                    "Trả lời ngắn": 0,
                }
                new_here = 0
                dup_here = 0

                for q in qs:
                    counts[q.get("dang_cau", "")] = counts.get(q.get("dang_cau", ""), 0) + 1
                    fp = fingerprint_cau_tot_nghiep_thuc_te(q)
                    if fp in by_fp:
                        idx = by_fp[fp]
                        item = dict(bank[idx])
                        sources = list(item.get("nguon_files", []) or [])
                        primary = str(item.get("nguon_file", "")).strip()
                        if primary and primary not in sources:
                            sources.append(primary)
                        if source_name not in sources:
                            sources.append(source_name)
                        item["nguon_files"] = sources

                        # Nếu câu đã có từ lần nhập trước nhưng khi đó app chưa đọc được
                        # đáp án/định dạng, lần nhập lại sẽ nâng cấp dữ liệu thay vì bỏ qua.
                        if (not _grad_co_dap_an_day_du(item)) and _grad_co_dap_an_day_du(q):
                            if q.get("dang_cau") == "Đúng / Sai":
                                item["nhan_dinh_meta"] = q.get("nhan_dinh_meta", [])
                            else:
                                item["dap_an"] = q.get("dap_an", "")
                            item["trang_thai"] = "Chờ rà soát"
                            item["duoc_dung_luyen_hs"] = True
                            item["san_sang_tu_dap_an_nguon"] = True
                            item["canh_bao_dap_an"] = ""
                            if q.get("dap_an_doc_tu_dinh_dang"):
                                item["dap_an_doc_tu_dinh_dang"] = True

                        # Bổ sung lời giải/tài nguyên nếu bản cũ còn thiếu.
                        if not str(item.get("giai_thich", "") or "").strip() and str(q.get("giai_thich", "") or "").strip():
                            item["giai_thich"] = q.get("giai_thich", "")
                            item["nguon_giai_thich"] = q.get("nguon_giai_thich", "")
                        if not (item.get("tai_nguyen_truc_quan") or []) and (q.get("tai_nguyen_truc_quan") or []):
                            item["tai_nguyen_truc_quan"] = q.get("tai_nguyen_truc_quan", [])
                            item["du_lieu_truc_quan"] = q.get("du_lieu_truc_quan", {})

                        # Luôn cập nhật phân loại ôn tập thô theo cơ chế mới.
                        item = gan_phan_loai_on_tap_cho_cau_tot_nghiep(item, force=False)
                        bank[idx] = item
                        dup_here += 1
                        continue

                    q["nguon_files"] = [source_name]
                    by_fp[fp] = len(bank)
                    bank.append(q)
                    new_here += 1

                total_new += new_here
                total_dup += dup_here
                reports.append({
                    "File": source_name,
                    "Tách được": len(qs),
                    "4 lựa chọn": counts.get("Trắc nghiệm 4 lựa chọn", 0),
                    "Đúng/Sai": counts.get("Đúng / Sai", 0),
                    "Trả lời ngắn": counts.get("Trả lời ngắn", 0),
                    "Câu mới": new_here,
                    "Trùng": dup_here,
                })
            except Exception as e:
                reports.append({
                    "File": getattr(uploaded, "name", ""),
                    "Tách được": 0,
                    "4 lựa chọn": 0,
                    "Đúng/Sai": 0,
                    "Trả lời ngắn": 0,
                    "Câu mới": 0,
                    "Trùng": 0,
                    "Lỗi": str(e),
                })

        luu_ngan_hang_tot_nghiep_thuc_te(bank)
        st.success(f"Đã thêm **{total_new} câu mới**; bỏ qua **{total_dup} câu trùng**.")
        st.dataframe(pd.DataFrame(reports), use_container_width=True, hide_index=True)
        for r in reports:
            if r.get("Tách được") and (
                r.get("4 lựa chọn") != 18 or r.get("Đúng/Sai") != 4 or r.get("Trả lời ngắn") != 6
            ):
                st.warning(
                    f"File **{r.get('File')}** chưa được đọc đúng form 18/4/6. "
                    "Hãy xem lại file hoặc dùng file mẫu để nhập ổn định hơn."
                )

    # ------------------------------------------------------
    # TỔNG QUAN BANK
    # ------------------------------------------------------
    bank = doc_ngan_hang_tot_nghiep_thuc_te()
    st.divider()
    st.subheader("② Kiểm tra & duyệt ngân hàng")

    if not bank:
        st.info("Chưa có câu nào trong ngân hàng tốt nghiệp từ đề thật.")
        return

    count_mcq = sum(q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn" for q in bank)
    count_tf = sum(q.get("dang_cau") == "Đúng / Sai" for q in bank)
    count_short = sum(q.get("dang_cau") == "Trả lời ngắn" for q in bank)
    count_warn = sum(q.get("trang_thai") in {"Thiếu đáp án", "Cần GV xem"} for q in bank)

    m1, m2, m3, m4, m5 = st.columns(5)
    m1.metric("Tổng câu", len(bank))
    m2.metric("4 lựa chọn", count_mcq)
    m3.metric("Đúng/Sai", count_tf)
    m4.metric("Trả lời ngắn", count_short)
    m5.metric("Cần GV xem", count_warn)

    with st.expander("🔎 Mở quản lý, bộ lọc và chi tiết ngân hàng tốt nghiệp", expanded=False):
        all_sources = sorted({
            src
            for q in bank
            for src in (q.get("nguon_files", []) or [q.get("nguon_file", "")])
            if str(src or "").strip()
        })
    
        # Kho quản trị rõ ràng: xem → chọn → sửa → xóa từng câu đã nhập.
        hien_thi_kho_cau_tot_nghiep_da_nhap(bank)
    
        st.markdown("### 🧭 Phân loại nhanh cho Ngân hàng tốt nghiệp")
        st.info(
            "Để GV thao tác nhẹ, ngân hàng tốt nghiệp chỉ cần **Khối – Chương – Mức độ – Thành phần năng lực**. "
            "Không bắt buộc gán Bài, YCCĐ hay chỉ báo. Có thể sửa trực tiếp 4 trường này trong **Kho câu đã nhập**."
        )

        st.markdown("### 🔎 Rà soát / kiểm tra nhanh")
        st.caption(
            "Bộ lọc dưới đây dùng cho khâu rà soát. Mục **Cần GV xem** bao gồm cả "
            "câu AI cảnh báo và câu **Thiếu đáp án**, nên số hiển thị luôn khớp với ô tổng quan phía trên."
        )
    
        # ------------------------------------------------------
        # XÓA FILE NGUỒN ĐÃ NHẬP
        # ------------------------------------------------------
        if all_sources:
            with st.expander("🗑️ Xóa file nguồn đã nhập", expanded=False):
                st.caption(
                    "Xóa dữ liệu theo đúng file nguồn. Câu chỉ thuộc file đó sẽ bị xóa; "
                    "câu trùng còn có nguồn khác sẽ được giữ lại."
                )
                src_del = st.selectbox(
                    "Chọn file nguồn cần xóa",
                    all_sources,
                    key="grad_delete_source_select",
                )
                confirm_del = st.checkbox(
                    "Tôi xác nhận muốn xóa dữ liệu của file nguồn này",
                    key="grad_delete_source_confirm",
                )
                if st.button(
                    "🗑️ XÓA FILE NGUỒN & DỮ LIỆU LIÊN QUAN",
                    use_container_width=True,
                    disabled=not confirm_del,
                    key="grad_delete_source_button",
                ):
                    bank_new = []
                    removed_questions = 0
                    kept_shared = 0
                    for item in bank:
                        item2 = dict(item)
                        sources = list(item2.get("nguon_files", []) or [])
                        primary = str(item2.get("nguon_file", "") or "").strip()
                        if primary and primary not in sources:
                            sources.append(primary)
    
                        if src_del not in sources:
                            bank_new.append(item2)
                            continue
    
                        remain = [x for x in sources if x != src_del]
                        if remain:
                            item2["nguon_files"] = remain
                            if primary == src_del or not primary:
                                item2["nguon_file"] = remain[0]
                                if item2.get("nguon_tao") == "Đề thật / đề thi thử GV tải lên":
                                    item2["nguon"] = remain[0]
                            bank_new.append(item2)
                            kept_shared += 1
                        else:
                            removed_questions += 1
    
                    luu_ngan_hang_tot_nghiep_thuc_te(bank_new)
    
                    # Xóa bản Word nguồn đã lưu. Ảnh/bảng chỉ xóa khi chắc chắn không còn câu dùng;
                    # ở đây giữ media để tránh làm mất hình của câu trùng còn nguồn khác.
                    safe_name = re.sub(r"[^0-9A-Za-zÀ-ỹ_.() -]+", "_", src_del)
                    source_path = os.path.join(GRAD_SOURCE_DIR, safe_name)
                    try:
                        if os.path.exists(source_path):
                            os.remove(source_path)
                    except Exception:
                        pass
    
                    st.success(
                        f"Đã xóa file nguồn **{src_del}**: xóa {removed_questions} câu riêng; "
                        f"giữ {kept_shared} câu còn nguồn khác."
                    )
                    st.rerun()
    
        f1, f2, f3 = st.columns(3)
        with f1:
            source_filter = st.selectbox("Lọc theo file nguồn", ["Tất cả"] + all_sources, key="grad_real_filter_source")
        with f2:
            type_filter = st.selectbox(
                "Lọc theo dạng",
                ["Tất cả", "Trắc nghiệm 4 lựa chọn", "Đúng / Sai", "Trả lời ngắn"],
                key="grad_real_filter_type",
            )
        with f3:
            status_filter = st.selectbox(
                "Lọc theo trạng thái",
                ["Tất cả", "Chờ rà soát", "Đã rà soát", "GV đã duyệt", "Cần GV xem", "Thiếu đáp án"],
                key="grad_real_filter_status",
            )
    
        filtered = []
        for q in bank:
            sources = q.get("nguon_files", []) or [q.get("nguon_file", "")]
            if source_filter != "Tất cả" and source_filter not in sources:
                continue
            if type_filter != "Tất cả" and q.get("dang_cau") != type_filter:
                continue
            if not _grad_status_group_match(q, status_filter):
                continue
            filtered.append(q)
    
        st.caption(f"Đang hiển thị **{len(filtered)}/{len(bank)} câu**.")
    
        # AI rà soát theo lô nhỏ, không tự chạy khi nhập để tiết kiệm quota.
        b1, b2 = st.columns([1, 2])
        with b1:
            review_n = st.number_input(
                "Số câu rà soát/lượt",
                min_value=1,
                max_value=10,
                value=3,
                step=1,
                key="grad_real_review_n",
            )
        with b2:
            st.caption(
                "AI chỉ **kiểm tra đáp án, xác nhận mức độ và thành phần năng lực**, không sáng tác/sửa câu. "
                "Câu đã có đáp án nguồn vẫn dùng được trước khi rà soát; nếu AI không chắc hoặc không khớp, "
                "app mới khóa câu và đưa sang Cần GV xem."
            )
    
        if st.button(
            "🔎 AI RÀ SOÁT CÁC CÂU CHƯA KIỂM TRA",
            use_container_width=True,
            key="grad_real_review_batch",
        ):
            targets = [
                q for q in bank
                if q.get("trang_thai") == "Chờ rà soát"
                and not q.get("ai_ra_soat")
            ][: int(review_n)]
            if not targets:
                st.info("Không còn câu ở trạng thái Chờ rà soát.")
            else:
                progress = st.progress(0)
                status_box = st.empty()
                updated = list(bank)
                by_id = {q.get("id"): i for i, q in enumerate(updated)}
                done = 0
                for k, q in enumerate(targets, start=1):
                    status_box.write(f"Đang rà soát {k}/{len(targets)} • {q.get('nguon_file','')} • Câu {q.get('so_cau_goc','')}")
                    try:
                        q2 = ra_soat_cau_tot_nghiep_bang_ai(q)
                        idx = by_id.get(q.get("id"))
                        if idx is not None:
                            updated[idx] = q2
                        done += 1
                    except Exception as e:
                        st.warning(f"Dừng rà soát tại câu {q.get('so_cau_goc','')}: {e}")
                        break
                    progress.progress(k / max(len(targets), 1))
                luu_ngan_hang_tot_nghiep_thuc_te(updated)
                status_box.success(f"Đã rà soát {done} câu. Câu chưa chắc được giữ lại để GV xem.")
                st.rerun()
    
        if not filtered:
            st.info("Không có câu phù hợp bộ lọc.")
            return
    
        labels = [
            f"{q.get('nguon_file','')} • Phần {q.get('phan_goc','')} • Câu {q.get('so_cau_goc','')} • {q.get('dang_cau','')} • {q.get('trang_thai','')}"
            for q in filtered
        ]
        selected_label = st.selectbox("Chọn câu để xem / sửa", labels, key="grad_real_select_question")
        q = filtered[labels.index(selected_label)]
    
        with st.container(border=True):
            st.write(f"**Nguồn:** {q.get('nguon_file','')} • **Phần {q.get('phan_goc','')} – Câu {q.get('so_cau_goc','')}**")
            if q.get("trang_thai") == "Cần GV xem":
                st.error("⚠️ AI chưa chắc đáp án này: " + str(q.get("canh_bao_dap_an", "")))
            elif q.get("trang_thai") == "Thiếu đáp án":
                st.warning("⚠️ File nguồn chưa đọc được đáp án. GV cần bổ sung trước khi dùng.")
            elif q.get("trang_thai") == "Đã rà soát":
                st.success(f"✅ AI đã rà soát • độ tin cậy {float(q.get('do_tin_cay_dap_an',0) or 0)*100:.0f}%")
            elif q.get("trang_thai") == "Chờ rà soát" and cau_tot_nghiep_du_dieu_kien_su_dung(q):
                st.info(
                    "✅ Câu đã có đáp án nguồn đầy đủ nên **được dùng ngay** để tạo đề và ôn tập. "
                    "AI rà soát là bước kiểm tra bổ sung; nếu phát hiện nghi vấn, app sẽ tự khóa câu để GV xem."
                )
    
            st.caption(
                "📚 Phân loại tốt nghiệp: "
                + (str(q.get("khoi", "")) or "Chưa xác định khối")
                + (" → " + str(q.get("chuong", "")) if str(q.get("chuong", "")).strip() else " → Chưa xác định chương")
                + " • " + str(q.get("muc_do", "Chưa xác định mức độ"))
                + " • " + str(q.get("thanh_phan_nang_luc", "Chưa xác định năng lực"))
            )
    
            # GV chỉ cần sửa Khối/Chương ở đây; mức độ/năng lực sửa tại phần xác nhận câu.
            with st.expander("🛠️ Sửa Khối / Chương nếu cần", expanded=False):
                scope_key = re.sub(r"[^A-Za-z0-9]", "", str(q.get("id", "")))[:12] or "scope"
                grades_scope = list((KHO_YCCD or {}).keys())
                if not grades_scope:
                    grades_scope = [str(q.get("khoi", "") or "")]
                grade0_scope = q.get("khoi", "") if q.get("khoi", "") in grades_scope else grades_scope[0]
                grade_scope = st.selectbox(
                    "Khối đúng",
                    grades_scope,
                    index=grades_scope.index(grade0_scope),
                    key=f"grad_review_scope_grade_{scope_key}",
                )
    
                chapters_scope = list((KHO_YCCD.get(grade_scope, {}) or {}).keys())
                if not chapters_scope:
                    chapters_scope = [str(q.get("chuong", "") or "")]
                chapter0_scope = q.get("chuong", "") if q.get("chuong", "") in chapters_scope else chapters_scope[0]
                chapter_scope = st.selectbox(
                    "Chương đúng",
                    chapters_scope,
                    index=chapters_scope.index(chapter0_scope),
                    key=f"grad_review_scope_chapter_{scope_key}_{grade_scope}",
                )
    
                # Ngân hàng tốt nghiệp không yêu cầu GV gán tới Bài.
                lesson_value_scope = ""
    
                apply_tf_scope = True
                if q.get("dang_cau") == "Đúng / Sai":
                    apply_tf_scope = st.checkbox(
                        "Áp dụng phạm vi này cho cả 4 ý a–d",
                        value=True,
                        key=f"grad_review_scope_apply_tf_{scope_key}",
                        help="Mức độ, năng lực và đáp án của từng ý vẫn được giữ riêng; chỉ đồng bộ Khối/Chương.",
                    )
    
                if st.button(
                    "💾 LƯU KHỐI / CHƯƠNG",
                    type="primary",
                    use_container_width=True,
                    key=f"grad_review_scope_save_{scope_key}",
                ):
                    if _grad_cap_nhat_pham_vi_gv(
                        q, grade_scope, chapter_scope, lesson_value_scope, ap_dung_4_y=apply_tf_scope
                    ):
                        st.success("Đã lưu Khối / Chương do GV xác nhận.")
                        st.rerun()
                    else:
                        st.error("Không tìm thấy câu để cập nhật phạm vi.")
    
            st.markdown("### Nội dung nguyên câu")
            st.write(q.get("cau_hoi", ""))
            hien_thi_tai_nguyen_cau_tot_nghiep(q)
    
            if q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn":
                for lc in q.get("lua_chon", []) or []:
                    st.write(lc)
            elif q.get("dang_cau") == "Đúng / Sai":
                for nd in q.get("nhan_dinh_meta", []) or []:
                    ky = nd.get("ky_hieu", "")
                    st.write(f"**{ky})** {nd.get('noi_dung','')}")
                    st.caption(
                        f"Phạm vi: {nd.get('khoi', q.get('khoi',''))}"
                        + (f" → {nd.get('chuong','')}" if str(nd.get('chuong','')).strip() else "")
                        + f" • Mức độ: {nd.get('muc_do','Chưa xác định')}"
                        + f" • Năng lực: {nd.get('thanh_phan_nang_luc','Chưa xác định')}"
                    )
    
            st.markdown("#### Đáp án / hướng dẫn giải")
            if q.get("dang_cau") == "Đúng / Sai":
                for nd in q.get("nhan_dinh_meta", []) or []:
                    st.write(f"**{nd.get('ky_hieu','')}) {nd.get('dap_an','—')}**")
                    if str(nd.get("giai_thich", "") or "").strip():
                        st.caption(f"Giải thích: {nd.get('giai_thich','')}")
            else:
                st.write("**Đáp án:**", q.get("dap_an", "—"))
            if q.get("giai_thich"):
                nguon_gt = str(q.get("nguon_giai_thich", "") or "").strip()
                if nguon_gt:
                    st.caption(f"Nguồn lời giải: {nguon_gt}")
                st.write("**Hướng dẫn giải:**", q.get("giai_thich", ""))
            elif q.get("dang_cau") != "Đúng / Sai":
                st.caption("Chưa có hướng dẫn giải. Khi AI rà soát đáp án khớp và đủ tin cậy, app sẽ tự tạo lời giải để lưu dùng lâu dài.")
    
            if q.get("dang_cau") != "Đúng / Sai":
                st.caption(
                    "Mức độ: " + str(q.get("muc_do", "Chưa xác định"))
                    + " • Thành phần năng lực: "
                    + str(q.get("thanh_phan_nang_luc", "Chưa xác định"))
                )
    
            if q.get("ai_ra_soat"):
                with st.expander("Xem kết quả AI rà soát", expanded=False):
                    data_rs = q.get("ai_ra_soat") or {}
                    st.write("**Kết luận:**", data_rs.get("ket_luan", ""))
                    st.write("**Độ tin cậy:**", f"{float(data_rs.get('do_tin_cay',0) or 0)*100:.0f}%")
                    if q.get("dang_cau") == "Đúng / Sai":
                        gt4 = list(data_rs.get("giai_thich_4_y", []) or [])
                        if gt4:
                            st.write("**Giải thích AI đề xuất:**")
                            for i, gt in enumerate(gt4[:4]):
                                st.write(f"{'abcd'[i]}) {gt}")
                    else:
                        gt = str(data_rs.get("huong_dan_giai_de_xuat", "") or "").strip()
                        if gt:
                            st.write("**Hướng dẫn giải AI đề xuất:**", gt)
                    st.caption("Bản AI đề xuất chỉ trở thành lời giải dùng cho HS khi đáp án nguồn khớp và AI đủ tin cậy; nếu không, GV phải duyệt.")
    
        # ------------------------------------------------------
        # SỬA / DUYỆT CÂU ĐANG CHỌN
        # ------------------------------------------------------
        with st.expander("✏️ GV sửa đáp án / năng lực và xác nhận", expanded=q.get("trang_thai") in {"Cần GV xem", "Thiếu đáp án"}):
            q_edit = dict(q)
            if q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn":
                current = str(q.get("dap_an", "") or "A").upper()
                if current not in ["A", "B", "C", "D"]:
                    current = "A"
                new_answer = st.selectbox("Đáp án đúng", ["A", "B", "C", "D"], index=["A", "B", "C", "D"].index(current), key="grad_edit_mcq_ans")
                current_nl = q.get("thanh_phan_nang_luc", "")
                new_nl = st.selectbox(
                    "Thành phần năng lực",
                    THANH_PHAN_NANG_LUC,
                    index=THANH_PHAN_NANG_LUC.index(current_nl) if current_nl in THANH_PHAN_NANG_LUC else 0,
                    key="grad_edit_mcq_nl",
                )
                q_edit["dap_an"] = new_answer
                q_edit["thanh_phan_nang_luc"] = new_nl
    
            elif q.get("dang_cau") == "Trả lời ngắn":
                q_edit["dap_an"] = st.text_input("Đáp án đúng", value=str(q.get("dap_an", "")), key="grad_edit_short_ans")
                current_nl = q.get("thanh_phan_nang_luc", "")
                q_edit["thanh_phan_nang_luc"] = st.selectbox(
                    "Thành phần năng lực",
                    THANH_PHAN_NANG_LUC,
                    index=THANH_PHAN_NANG_LUC.index(current_nl) if current_nl in THANH_PHAN_NANG_LUC else 0,
                    key="grad_edit_short_nl",
                )
    
            else:
                meta = [dict(x) for x in (q.get("nhan_dinh_meta", []) or [])]
                for i, nd in enumerate(meta):
                    st.markdown(f"**Ý {nd.get('ky_hieu','abcd'[i] if i < 4 else i+1)}**")
                    ca, cb = st.columns(2)
                    with ca:
                        ans = str(nd.get("dap_an", "") or "Đúng")
                        nd["dap_an"] = st.selectbox(
                            "Đáp án",
                            ["Đúng", "Sai"],
                            index=0 if ans == "Đúng" else 1,
                            key=f"grad_edit_tf_ans_{i}",
                        )
                    with cb:
                        nl0 = nd.get("thanh_phan_nang_luc", "")
                        nd["thanh_phan_nang_luc"] = st.selectbox(
                            "Thành phần năng lực",
                            THANH_PHAN_NANG_LUC,
                            index=THANH_PHAN_NANG_LUC.index(nl0) if nl0 in THANH_PHAN_NANG_LUC else 0,
                            key=f"grad_edit_tf_nl_{i}",
                        )
                q_edit["nhan_dinh_meta"] = meta
    
            q_edit["giai_thich"] = st.text_area(
                "Hướng dẫn giải / ghi chú của GV",
                value=str(q.get("giai_thich", "")),
                height=130,
                key="grad_edit_solution",
            )
    
            if st.button("✅ GV XÁC NHẬN & LƯU CÂU NÀY", type="primary", use_container_width=True, key="grad_edit_confirm"):
                q_edit["trang_thai"] = "GV đã duyệt"
                q_edit["duoc_dung_luyen_hs"] = True
                q_edit["canh_bao_dap_an"] = ""
                if str(q_edit.get("giai_thich", "") or "").strip():
                    # Nếu GV có sửa/nhập lời giải thì GV là nguồn xác nhận cuối cùng.
                    if str(q_edit.get("giai_thich", "")) != str(q.get("giai_thich", "")) or not q_edit.get("nguon_giai_thich"):
                        q_edit["nguon_giai_thich"] = "GV xác nhận / bổ sung"
                q_edit["ngay_gv_duyet"] = datetime.now().strftime("%d/%m/%Y %H:%M")
                bank2 = doc_ngan_hang_tot_nghiep_thuc_te()
                for i, item in enumerate(bank2):
                    if item.get("id") == q.get("id"):
                        bank2[i] = q_edit
                        break
                luu_ngan_hang_tot_nghiep_thuc_te(bank2)
                st.success("Đã lưu xác nhận của GV.")
                st.rerun()
    
        st.divider()
        st.subheader("③ Sẵn sàng tạo đề cho học sinh")
        st.caption(
            "Câu có **đáp án nguồn đầy đủ** được tính là dùng được ngay, kể cả đang ở trạng thái Chờ rà soát. "
            "Chỉ câu Thiếu đáp án / Cần GV xem / Ngừng sử dụng mới bị loại khỏi pool."
        )
        usable = [
            q for q in bank
            if cau_tot_nghiep_du_dieu_kien_su_dung(q)
        ]
        u1, u2, u3 = st.columns(3)
        u1.metric("4 lựa chọn dùng được", sum(q.get("dang_cau") == "Trắc nghiệm 4 lựa chọn" for q in usable))
        u2.metric("Đúng/Sai dùng được", sum(q.get("dang_cau") == "Đúng / Sai" for q in usable))
        u3.metric("Trả lời ngắn dùng được", sum(q.get("dang_cau") == "Trả lời ngắn" for q in usable))
    
        _, missing = rut_de_tot_nghiep_tu_de_that(usable, seed="kiem-tra-form")
        if missing:
            st.warning("Ngân hàng chưa đủ để tạo một đề 18/4/6 không trùng câu.")
            st.dataframe(pd.DataFrame(missing), use_container_width=True, hide_index=True)
        else:
            st.success(
                "Ngân hàng đã đủ form. Khi tạo đề, app sẽ trộn câu từ các file nguồn khác nhau, "
                "không lấy trùng một câu trong cùng mã đề và giữ nguyên ảnh/bảng của câu gốc."
            )


def gan_muc_dich_cau_hien_co():
    with st.expander("🏷️ Gắn mục đích sử dụng cho câu hiện có", expanded=False):
        st.subheader("🏷️ Gắn mục đích sử dụng cho câu đã có")
        st.caption("Một câu có thể dùng cho kiểm tra ở trường, luyện tốt nghiệp hoặc cả hai.")
        bank=doc_ngan_hang()
        if not bank: st.info("Ngân hàng chuẩn chưa có câu.")
        else:
            loc=st.selectbox("Lọc dạng câu",["Tất cả","Trắc nghiệm 4 lựa chọn","Đúng / Sai","Trả lời ngắn"],key="tag_bank_type")
            bank_hien=[q for q in bank if loc=="Tất cả" or q.get("dang_cau")==loc]
            labels=[f"{i+1}. {q.get('dang_cau','')} • {str(q.get('cau_hoi',''))[:100]}" for i,q in enumerate(bank_hien)]
            if labels:
                lb=st.selectbox("Chọn câu",labels,key="tag_bank_select"); q_chon=bank_hien[labels.index(lb)]; st.write(q_chon.get("cau_hoi",""))
                muc_cu=q_chon.get("muc_dich_su_dung",""); keys=["","on_tap_kiem_tra","tot_nghiep","ca_hai"]
                map_label={"":"Chưa gắn","on_tap_kiem_tra":"Ôn tập / kiểm tra ở trường","tot_nghiep":"Luyện tốt nghiệp","ca_hai":"Cả hai"}
                new_purpose=st.selectbox("Mục đích sử dụng",keys,index=keys.index(muc_cu) if muc_cu in keys else 0,format_func=lambda x:map_label[x],key="tag_bank_purpose")
                new_behavior=st.selectbox("Hành vi năng lực",[""]+HANH_VI_NANG_LUC_TOT_NGHIEP,index=0,key="tag_bank_behavior")
                if st.button("💾 LƯU NHÃN",type="primary",use_container_width=True,key="tag_bank_save"):
                    q_id=q_chon.get("id") or q_chon.get("temp_id")
                    for item in bank:
                        item_id=item.get("id") or item.get("temp_id")
                        if q_id and item_id==q_id:
                            item["muc_dich_su_dung"]=new_purpose
                            if new_behavior: item["hanh_vi_nang_luc"]=new_behavior
                            break
                    luu_ngan_hang(bank); st.success("Đã cập nhật mục đích sử dụng.")




# ==========================================================
# BẢNG ĐIỂM ÔN / KIỂM TRA GỬI VỀ GIÁO VIÊN
# ==========================================================
def _thoi_gian_sort_lich_su(lan):
    return str(
        lan.get("nop_bai_iso", "")
        or lan.get("thoi_gian_iso", "")
        or lan.get("thoi_gian", "")
    )


def loc_luot_co_diem_chinh_thuc(lop=None, che_do=None, ten_luot=None):
    ds_hs = doc_danh_sach_hoc_sinh()
    map_hs = {
        str(x.get("ma_hoc_sinh", "")).strip().upper(): x
        for x in ds_hs
    }

    ket_qua = []
    for lan in doc_lich_su_hoc_sinh():
        mode = str(lan.get("che_do", "")).strip()
        if mode not in {
            CHE_DO_DE_GV,
            CHE_DO_KIEM_TRA_MA_TRAN,
            CHE_DO_TOT_NGHIEP
        }:
            continue

        ma = str(lan.get("hoc_sinh_id", "")).strip().upper()
        hs = map_hs.get(ma, {})
        lop_hs = str(hs.get("lop", "")).strip()

        if lop and lop != "Tất cả" and lop_hs != lop:
            continue
        if che_do and che_do != "Tất cả" and mode != che_do:
            continue
        if ten_luot and ten_luot != "Tất cả" and str(lan.get("ten_luot", "")).strip() != ten_luot:
            continue

        item = dict(lan)
        item["_lop"] = lop_hs
        item["_ho_ten"] = hs.get("ho_ten", lan.get("ho_ten", ""))
        ket_qua.append(item)

    return sorted(ket_qua, key=_thoi_gian_sort_lich_su)


def chi_giu_luot_gan_nhat_moi_hs(ds_luot):
    latest = {}
    for lan in sorted(ds_luot, key=_thoi_gian_sort_lich_su):
        ma = str(lan.get("hoc_sinh_id", "")).strip().upper()
        if ma:
            latest[ma] = lan
    return list(latest.values())


def tong_hop_diem_yeu_tu_luot(ds_luot, limit=10):
    bucket = {}

    for lan in ds_luot or []:
        for item in lan.get("chi_tiet", []) or []:
            q = item.get("cau_snapshot", {}) or {}
            units = item.get("don_vi_danh_gia", []) or []
            if not units:
                units = [{
                    "yccd": q.get("yccd", ""),
                    "muc_do": q.get("muc_do", ""),
                    "nang_luc": q.get("thanh_phan_nang_luc", ""),
                    "dung": bool(item.get("dung_toan_cau"))
                }]

            for unit in units:
                yccd = str(unit.get("yccd", "")).strip()
                if not yccd:
                    continue
                muc = str(unit.get("muc_do", "")).strip()
                nl = str(unit.get("nang_luc", "")).strip()
                key = (yccd, muc, nl)
                s = bucket.setdefault(key, {"tong": 0, "dung": 0})
                s["tong"] += 1
                s["dung"] += int(bool(unit.get("dung")))

    rows = []
    for (yccd, muc, nl), stat in bucket.items():
        tong = int(stat["tong"])
        dung = int(stat["dung"])
        rows.append({
            "YCCĐ": yccd,
            "Mức độ": muc,
            "Năng lực": nl,
            "Số lượt đánh giá": tong,
            "Số đúng": dung,
            "Số sai": tong - dung,
            "Tỉ lệ đúng (%)": round(dung / tong * 100, 1) if tong else 0
        })

    rows.sort(key=lambda r: (r["Tỉ lệ đúng (%)"], -r["Số lượt đánh giá"]))
    return rows[:limit]


def bang_diem_tu_lich_su(ds_luot):
    rows = []
    for lan in ds_luot:
        pham_vi = lan.get("pham_vi", {}) or {}
        diem = float(lan.get("diem_chinh_thuc", lan.get("diem", 0)) or 0)
        thang = float(lan.get("thang_diem", 10) or 10)
        rows.append({
            "Mã học sinh": str(lan.get("hoc_sinh_id", "")).strip().upper(),
            "Họ và tên": lan.get("_ho_ten", lan.get("ho_ten", "")),
            "Lớp": lan.get("_lop", ""),
            "Chế độ": lan.get("che_do", ""),
            "Bài / ma trận": lan.get("ten_luot", ""),
            "Mã đợt kiểm tra": pham_vi.get("dot_kiem_tra_id", ""),
            "Mã đề": pham_vi.get("ma_de_hs", pham_vi.get("ma_de", "")),
            "Phiên bản ma trận": pham_vi.get("mau_phien_ban", ""),
            "Bắt đầu": lan.get("bat_dau_luc", ""),
            "Nộp bài": lan.get("nop_bai_luc", lan.get("thoi_gian", "")),
            "Thời lượng làm (phút)": round(float(lan.get("thoi_luong_lam_phut", 0) or 0), 1),
            "Thời gian quy định (phút)": int(lan.get("thoi_gian_quy_dinh_phut", 0) or 0),
            "Điểm": round(diem, 2),
            "Thang điểm": round(thang, 2),
            "Tỉ lệ điểm (%)": round((diem / thang * 100) if thang else 0, 1),
            "Tỉ lệ đúng (%)": float(lan.get("ti_le_dung_don_vi", 0) or 0),
            "Số câu": int(lan.get("tong_so_cau", 0) or 0),
            "Đơn vị đúng": int(lan.get("dung_don_vi", 0) or 0),
            "Tổng đơn vị": int(lan.get("tong_don_vi", 0) or 0),
        })
    return rows


def tao_excel_bang_diem_lop(ds_luot, lop_chon=""):
    output = io.BytesIO()
    rows = bang_diem_tu_lich_su(ds_luot)
    weak_class = tong_hop_diem_yeu_tu_luot(ds_luot, 50)

    weak_student = []
    ma_list = sorted({str(x.get("hoc_sinh_id", "")).strip().upper() for x in ds_luot})
    for ma in ma_list:
        ls_hs = [x for x in ds_luot if str(x.get("hoc_sinh_id", "")).strip().upper() == ma]
        ho_ten = ls_hs[-1].get("_ho_ten", ls_hs[-1].get("ho_ten", "")) if ls_hs else ""
        for r in tong_hop_diem_yeu_tu_luot(ls_hs, 20):
            item = {"Mã học sinh": ma, "Họ và tên": ho_ten}
            item.update(r)
            weak_student.append(item)

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame(rows).to_excel(writer, index=False, sheet_name="Bang diem")
        pd.DataFrame(weak_class).to_excel(writer, index=False, sheet_name="Diem yeu lop")
        pd.DataFrame(weak_student).to_excel(writer, index=False, sheet_name="Diem yeu tung HS")

        for ws in writer.book.worksheets:
            ws.freeze_panes = "A2"
            for cell in ws[1]:
                cell.font = cell.font.copy(bold=True)
            for col_cells in ws.columns:
                letter = col_cells[0].column_letter
                max_len = max((len(str(c.value)) if c.value is not None else 0) for c in col_cells)
                ws.column_dimensions[letter].width = min(max(max_len + 2, 10), 48)

    output.seek(0)
    return output.getvalue()


def quan_ly_diem_on_kiem_tra():
    st.header("🧾 ĐIỂM ÔN / KIỂM TRA THEO LỚP")
    st.caption(
        "Kết quả học sinh nộp bài được lưu tự động vào cùng kho dữ liệu tiến bộ. "
        "GV có thể lọc theo lớp, loại bài và đề/ma trận cụ thể, sau đó xuất Excel."
    )

    ds_lop = lay_danh_sach_lop_tu_hoc_sinh()
    if not ds_lop:
        st.info("Chưa có lớp/học sinh trong hệ thống.")
        return

    f1, f2 = st.columns(2)
    with f1:
        lop = st.selectbox(
            "Chọn lớp",
            ds_lop,
            key="gv_score_class"
        )
    with f2:
        che_do = st.selectbox(
            "Loại bài",
            [
                "Tất cả",
                CHE_DO_DE_GV,
                CHE_DO_KIEM_TRA_MA_TRAN,
                CHE_DO_TOT_NGHIEP
            ],
            key="gv_score_mode"
        )

    ds_co_so = loc_luot_co_diem_chinh_thuc(
        lop=lop,
        che_do=None if che_do == "Tất cả" else che_do
    )

    if not ds_co_so:
        st.info("Chưa có học sinh nào nộp bài phù hợp bộ lọc này.")
        return

    # Lọc theo ID thật của từng đề/ma trận/đợt kiểm tra, không chỉ theo tên.
    # Nhờ vậy hai đề trùng tên hoặc hai phiên bản ma trận không bị trộn điểm.
    key_to_label = {}
    for x in ds_co_so:
        k = khoa_bai_lam_chinh_thuc(x)
        if k and k not in key_to_label:
            key_to_label[k] = nhan_bai_lam_chinh_thuc(x)

    bai_keys = sorted(
        key_to_label.keys(),
        key=lambda k: key_to_label.get(k, k)
    )

    bai_key = st.selectbox(
        "Chọn đúng đề / ma trận / đợt kiểm tra cần xem",
        ["Tất cả"] + bai_keys,
        format_func=lambda k: (
            "Tất cả"
            if k == "Tất cả"
            else key_to_label.get(k, k)
        ),
        key="gv_score_exam_exact"
    )

    ds_loc = [
        x for x in ds_co_so
        if bai_key == "Tất cả"
        or khoa_bai_lam_chinh_thuc(x) == bai_key
    ]

    cach_lay = st.radio(
        "Cách lấy điểm khi một học sinh đã làm nhiều lần",
        ["Lượt gần nhất mỗi học sinh", "Tất cả lượt nộp"],
        horizontal=True,
        key="gv_score_attempt_policy"
    )

    ds_hien = (
        chi_giu_luot_gan_nhat_moi_hs(ds_loc)
        if cach_lay.startswith("Lượt gần nhất")
        else ds_loc
    )

    rows = bang_diem_tu_lich_su(ds_hien)
    if not rows:
        st.info("Không có dữ liệu sau khi lọc.")
        return

    diem_pct = [float(r.get("Tỉ lệ điểm (%)", 0) or 0) for r in rows]
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Học sinh có điểm", len({r["Mã học sinh"] for r in rows}))
    c2.metric("Số lượt nộp", len(rows))
    c3.metric("TB quy đổi", f"{sum(diem_pct) / len(diem_pct):.1f}%" if diem_pct else "—")
    c4.metric(
        "Nộp dưới 50%",
        sum(1 for x in diem_pct if x < 50)
    )

    st.dataframe(
        pd.DataFrame(rows),
        use_container_width=True,
        hide_index=True,
        height=430
    )

    st.download_button(
        "⬇️ XUẤT BẢNG ĐIỂM EXCEL THEO LỚP",
        data=tao_excel_bang_diem_lop(ds_hien, lop),
        file_name=f"bang_diem_{chuan_hoa_ten_lop(lop)}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
        key="gv_score_export"
    )

    st.markdown("---")
    st.subheader("🎯 Phản hồi nhanh về điểm yếu")

    weak_class = tong_hop_diem_yeu_tu_luot(ds_hien, 10)
    if weak_class:
        st.markdown("**Điểm yếu nổi bật của lớp trong các bài đang lọc:**")
        st.dataframe(
            pd.DataFrame(weak_class),
            use_container_width=True,
            hide_index=True
        )
    else:
        st.success("Chưa phát hiện điểm yếu nổi bật trong dữ liệu đang chọn.")

    ma_options = sorted({
        (str(x.get("hoc_sinh_id", "")).strip().upper(), str(x.get("_ho_ten", x.get("ho_ten", ""))))
        for x in ds_hien
        if str(x.get("hoc_sinh_id", "")).strip()
    })

    if ma_options:
        idx_hs = st.selectbox(
            "Xem điểm yếu của một học sinh",
            options=list(range(len(ma_options))),
            format_func=lambda i: f"{ma_options[i][0]} – {ma_options[i][1]}",
            key="gv_score_student_weak"
        )
        ma_chon = ma_options[idx_hs][0]
        ls_hs = [
            x for x in ds_hien
            if str(x.get("hoc_sinh_id", "")).strip().upper() == ma_chon
        ]
        weak_hs = tong_hop_diem_yeu_tu_luot(ls_hs, 8)
        if weak_hs:
            st.dataframe(
                pd.DataFrame(weak_hs),
                use_container_width=True,
                hide_index=True
            )
        else:
            st.success("Học sinh này chưa có điểm yếu nổi bật trong các bài đang lọc.")

def giao_vien():
    with st.sidebar:
        hien_thi_the_giao_vien_sidebar()
        st.divider()

        menu = st.radio(
    "Chức năng",
    [
        "📚 Kho YCCĐ",
        "📁 Kho tài liệu GV",
        "🌱 Ngân hàng hạt giống",
        "🧱 Xây dựng NH ôn tập / kiểm tra",
        "🎓 Xây dựng NH tốt nghiệp",
        "🏦 Ngân hàng câu hỏi",
        "📝 Tạo đề",
        "👥 Quản lý học sinh",
        "🧾 Điểm ôn / kiểm tra",
        "🗂️ Dữ liệu & tiến bộ HS",
        "📈 Phân tích lớp học"
    ]
)

        st.divider()

        so_yccd = len(
            st.session_state.yccd_da_chon
        )

        tong_so_cau = 0

        for item in st.session_state.yccd_da_chon:
            _, cau_hinh = lay_cau_hinh(item)
            tong_so_cau += int(
                cau_hinh.get("Số câu", 1)
            )

        bank = doc_ngan_hang()

        st.write(
            f"✅ YCCĐ đã chọn: "
            f"**{so_yccd}**"
        )

        st.write(
            f"📝 Tổng số câu: "
            f"**{tong_so_cau}**"
        )

        st.write(
            f"🏦 Ngân hàng: "
            f"**{len(bank)} câu**"
        )

        if so_yccd > 0:

            if st.button(
                "🆕 TẠO BỘ MỚI",
                use_container_width=True
            ):

                xoa_toan_bo()

                st.rerun()

        st.divider()

        if st.button(
            "⬅️ Về trang chủ",
            use_container_width=True
        ):

            st.session_state.vai_tro = None

            st.rerun()

    hien_thi_dau_trang_tram_sinh_hoc("giaovien")

    profile = _nang_cap_avatar_profile_ben_vung(doc_ho_so_giao_vien())
    st.markdown(
        f"""
        <div style="max-width:1180px;margin:0 auto .95rem auto;padding:.66rem .88rem;
                    border:1px solid #e2e8f0;border-radius:14px;background:rgba(255,255,255,.88);
                    box-shadow:0 5px 18px rgba(15,23,42,.035);color:#334155;font-weight:760;">
            {_html_escape(menu)}
        </div>
        """,
        unsafe_allow_html=True
    )

    if menu == "📚 Kho YCCĐ":

        kho_yccd()

    elif menu == "📁 Kho tài liệu GV":

        kho_tai_lieu_gv()

    elif menu == "🌱 Ngân hàng hạt giống":

        ngan_hang_hat_giong()

    elif menu == "🧱 Xây dựng NH ôn tập / kiểm tra":

        tao_cau_hoi_ai()

    elif menu == "🎓 Xây dựng NH tốt nghiệp":

        xay_dung_ngan_hang_tot_nghiep()

    elif menu == "🏦 Ngân hàng câu hỏi":

        ngan_hang_cau_hoi()

    elif menu == "📝 Tạo đề":

        tao_de_giao_vien()

    elif menu == "👥 Quản lý học sinh":

        quan_ly_hoc_sinh()

    elif menu == "🧾 Điểm ôn / kiểm tra":

        quan_ly_diem_on_kiem_tra()

    elif menu == "🗂️ Dữ liệu & tiến bộ HS":

        du_lieu_va_tien_bo_hoc_sinh()

    elif menu == "📈 Phân tích lớp học":

        phan_tich_lop_hoc()

# ==========================================================
# HỌC SINH
# ==========================================================

# ==========================================================
# HỒ SƠ HỌC TẬP & CÁ NHÂN HÓA HỌC SINH
# ==========================================================
def chuan_hoa_ma_hoc_sinh(value):
    s = re.sub(
        r"[^0-9A-Za-zÀ-ỹ_-]+",
        "_",
        str(value or "").strip()
    )
    return s[:80] or "hoc_sinh"


def doc_lich_su_hoc_sinh():
    return _doc_attempts_shared()


def luu_lich_su_hoc_sinh(ds):
    return _luu_attempts_shared(ds)


def lay_lich_su_cua_hoc_sinh(hoc_sinh_id):
    hid = chuan_hoa_ma_hoc_sinh(
        hoc_sinh_id
    )
    return [
        x
        for x in doc_lich_su_hoc_sinh()
        if chuan_hoa_ma_hoc_sinh(
            x.get("hoc_sinh_id", "")
        ) == hid
    ]


def khoa_nang_luc(yccd, muc_do, nang_luc, chi_bao=""):
    return "|||".join([
        str(yccd or "").strip(),
        str(muc_do or "").strip(),
        str(nang_luc or "").strip(),
        str(chi_bao or "").strip()
    ])


def tao_ho_so_tu_lich_su(hoc_sinh_id):
    """
    Hồ sơ động theo từng đơn vị đánh giá:
    YCCĐ × mức độ × thành phần năng lực.
    """
    lich_su = lay_lich_su_cua_hoc_sinh(
        hoc_sinh_id
    )

    stats = {}
    cau_da_gap = {}
    tong_don_vi = 0
    tong_dung = 0

    # Dùng metadata hiện tại của ngân hàng để bù chỉ báo cho lịch sử cũ.
    # Nhờ vậy các lượt HS đã làm trước đây cũng được thống kê theo NT/TH/VD.
    try:
        bank_hien_tai = gan_chi_bao_cho_ngan_hang_hien_co(doc_ngan_hang())
    except Exception:
        bank_hien_tai = doc_ngan_hang()

    cau_theo_id = {
        str(q.get("id", "")): q
        for q in bank_hien_tai
        if str(q.get("id", "")).strip()
    }

    for lan in lich_su:
        for item in lan.get(
            "chi_tiet",
            []
        ) or []:

            cau_id = str(
                item.get("cau_id", "")
            )

            if cau_id:
                info = cau_da_gap.setdefault(
                    cau_id,
                    {
                        "so_lan": 0,
                        "so_lan_dung": 0,
                        "lan_cuoi": ""
                    }
                )
                info["so_lan"] += 1
                if item.get("dung_toan_cau"):
                    info["so_lan_dung"] += 1
                info["lan_cuoi"] = lan.get(
                    "thoi_gian_iso",
                    lan.get("thoi_gian", "")
                )

            for unit in item.get(
                "don_vi_danh_gia",
                []
            ) or []:
                unit = dict(unit)
                ma_cb = str(unit.get("chi_bao", "")).strip().upper()

                # Lịch sử cũ có thể chưa lưu chỉ báo. Suy lại từ câu hiện tại.
                if not re.fullmatch(r"(NT[1-8]|TH[1-7]|VD[1-3])", ma_cb):
                    q_ref = cau_theo_id.get(cau_id, {})
                    q_ref = gan_chi_bao_chuan_cho_cau(q_ref) if q_ref else {}
                    ma_cb = str(q_ref.get("chi_bao", "")).strip().upper()

                    if ma_cb:
                        unit["chi_bao"] = ma_cb
                        unit["nang_luc"] = (
                            nang_luc_theo_ma_chi_bao(ma_cb)
                            or unit.get("nang_luc", "")
                        )

                key = khoa_nang_luc(
                    unit.get("yccd", ""),
                    unit.get("muc_do", ""),
                    unit.get("nang_luc", ""),
                    unit.get("chi_bao", "")
                )

                s = stats.setdefault(
                    key,
                    {
                        "yccd": unit.get("yccd", ""),
                        "muc_do": unit.get("muc_do", ""),
                        "nang_luc": unit.get("nang_luc", ""),
                        "chi_bao": unit.get("chi_bao", ""),
                        "so_lan": 0,
                        "so_dung": 0
                    }
                )

                s["so_lan"] += 1
                tong_don_vi += 1

                if unit.get("dung"):
                    s["so_dung"] += 1
                    tong_dung += 1

    for s in stats.values():
        s["ti_le_dung"] = (
            s["so_dung"] / s["so_lan"]
            if s["so_lan"] > 0
            else 0
        )

        # Bayesian smoothing nhẹ để tránh 1 câu sai = yếu tuyệt đối.
        s["mastery"] = (
            s["so_dung"] + 1
        ) / (
            s["so_lan"] + 2
        )

    return {
        "hoc_sinh_id": chuan_hoa_ma_hoc_sinh(
            hoc_sinh_id
        ),
        "so_luot_lam": len(lich_su),
        "tong_don_vi": tong_don_vi,
        "tong_dung": tong_dung,
        "ti_le_dung": (
            tong_dung / tong_don_vi
            if tong_don_vi > 0
            else 0
        ),
        "stats": stats,
        "cau_da_gap": cau_da_gap
    }


def mastery_cua_unit(profile, unit):
    key = khoa_nang_luc(
        unit.get("yccd", ""),
        unit.get("muc_do", ""),
        unit.get("nang_luc", ""),
        unit.get("chi_bao", "")
    )

    s = profile.get(
        "stats",
        {}
    ).get(
        key
    )

    if not s:
        # Chưa học: xem như mức trung bình thấp để ưu tiên vừa phải.
        return 0.45

    return float(
        s.get(
            "mastery",
            0.5
        )
    )


def mastery_cua_cau(profile, q):
    units = metadata_don_vi_cau(
        q
    )

    if not units:
        return 0.5

    return sum(
        mastery_cua_unit(
            profile,
            u
        )
        for u in units
    ) / len(units)


def diem_uu_tien_cau_ca_nhan(profile, q, rng):
    """
    Điểm càng cao càng được ưu tiên.
    Quy tắc:
    - chưa gặp > từng sai > đã đúng nhiều;
    - điểm yếu > trung bình > điểm mạnh;
    - vẫn có nhiễu ngẫu nhiên để đề không lặp cứng.
    """
    cau_id = str(
        q.get("id", "")
    )
    gap = profile.get(
        "cau_da_gap",
        {}
    ).get(
        cau_id
    )

    mastery = mastery_cua_cau(
        profile,
        q
    )

    # 60% trọng tâm điểm yếu, 25% trung bình, 15% duy trì phần mạnh.
    if mastery < 0.50:
        base = 100.0
    elif mastery < 0.75:
        base = 60.0
    else:
        base = 30.0

    if not gap:
        base += 55.0
    else:
        so_lan = int(
            gap.get(
                "so_lan",
                0
            )
        )
        so_dung = int(
            gap.get(
                "so_lan_dung",
                0
            )
        )

        if so_lan > 0 and so_dung < so_lan:
            # Câu từng sai: ôn lại có chủ đích.
            base += 35.0
        else:
            # Đã làm đúng nhiều lần: giảm ưu tiên.
            base -= min(
                25.0,
                so_lan * 6.0
            )

    return base + rng.random() * 18.0


def rut_cau_ca_nhan_hoa(
    pool,
    so_cau,
    profile,
    dang_counts=None,
    seed=None
):
    rng = random.Random(
        seed or str(uuid.uuid4())
    )

    pool = list(pool)

    def rank(ds):
        return sorted(
            ds,
            key=lambda q: diem_uu_tien_cau_ca_nhan(
                profile,
                q,
                rng
            ),
            reverse=True
        )

    da_chon = []
    ids = set()

    if dang_counts:
        for dang, can in dang_counts.items():
            if can <= 0:
                continue

            ds_dang = [
                q
                for q in pool
                if q.get("dang_cau") == dang
                and str(
                    q.get("id", "")
                ) not in ids
            ]

            for q in rank(ds_dang)[:int(can)]:
                da_chon.append(q)
                ids.add(
                    str(
                        q.get("id", "")
                    )
                )

    con = max(
        0,
        int(so_cau) - len(da_chon)
    )

    if con > 0:
        ds_con = [
            q
            for q in pool
            if str(
                q.get("id", "")
            ) not in ids
        ]

        for q in rank(ds_con)[:con]:
            da_chon.append(q)
            ids.add(
                str(
                    q.get("id", "")
                )
            )

    rng.shuffle(
        da_chon
    )

    return da_chon[:int(so_cau)]


def phan_bo_3_dang_tu_dong(so_cau, pool):
    """
    Mặc định gần 60% 4LC, 20% Đ/S, 20% TLN.
    Tự co giãn nếu pool của một dạng không đủ.
    """
    n = int(so_cau)

    target = {
        "Trắc nghiệm 4 lựa chọn": round(
            n * 0.60
        ),
        "Đúng / Sai": round(
            n * 0.20
        ),
        "Trả lời ngắn": 0
    }

    target[
        "Trả lời ngắn"
    ] = max(
        0,
        n
        - target["Trắc nghiệm 4 lựa chọn"]
        - target["Đúng / Sai"]
    )

    available = {
        dang: sum(
            1
            for q in pool
            if q.get("dang_cau") == dang
        )
        for dang in target
    }

    actual = {
        dang: min(
            target[dang],
            available[dang]
        )
        for dang in target
    }

    con = n - sum(
        actual.values()
    )

    if con > 0:
        for dang in sorted(
            actual,
            key=lambda d: (
                available[d] - actual[d]
            ),
            reverse=True
        ):
            if con <= 0:
                break

            them = min(
                con,
                max(
                    0,
                    available[dang]
                    - actual[dang]
                )
            )

            actual[dang] += them
            con -= them

    return actual


def tao_don_vi_ket_qua_cau(
    q,
    ket_qua_tung_y=None,
    dung_toan_cau=False
):
    """
    Chuyển kết quả câu thành các đơn vị để cập nhật hồ sơ.
    Đ/S: mỗi ý là 1 đơn vị.
    Câu thường: 1 đơn vị.
    """
    if q.get("dang_cau") == "Đúng / Sai":
        units = []

        meta = list(
            q.get(
                "nhan_dinh_meta",
                []
            )
            or []
        )

        kq = list(
            ket_qua_tung_y
            or []
        )

        for i, nd in enumerate(
            meta[:4]
        ):
            dung_y = (
                bool(kq[i])
                if i < len(kq)
                else False
            )

            units.append({
                "yccd": nd.get(
                    "yccd",
                    ""
                ),
                "muc_do": nd.get(
                    "muc_do",
                    ""
                ),
                "nang_luc": nd.get(
                    "thanh_phan_nang_luc",
                    ""
                ),
                "chi_bao": (
                    nd.get("chi_bao", "")
                    or nd.get("hanh_vi_nang_luc", "")
                    or q.get("hanh_vi_nang_luc", "")
                ),
                "dung": dung_y
            })

        return units

    return [{
        "yccd": q.get(
            "yccd",
            ""
        ),
        "muc_do": q.get(
            "muc_do",
            ""
        ),
        "nang_luc": q.get(
            "thanh_phan_nang_luc",
            ""
        ),
        "chi_bao": (
            q.get("chi_bao", "")
            or q.get("hanh_vi_nang_luc", "")
        ),
        "dung": bool(
            dung_toan_cau
        )
    }]


def tom_tat_diem_yeu(profile, limit=5):
    ds = list(
        profile.get(
            "stats",
            {}
        ).values()
    )

    ds = [
        x
        for x in ds
        if x.get(
            "so_lan",
            0
        ) >= 1
    ]

    ds.sort(
        key=lambda x: (
            x.get(
                "mastery",
                0.5
            ),
            -x.get(
                "so_lan",
                0
            )
        )
    )

    return ds[:limit]


def tao_nhan_xet_quy_tac(profile):
    weak = tom_tat_diem_yeu(
        profile,
        3
    )

    if not weak:
        return (
            "Hệ thống chưa có đủ lịch sử để xác định điểm yếu ổn định. "
            "Hãy làm thêm vài lượt để app cá nhân hóa chính xác hơn."
        )

    parts = []

    for x in weak:
        ti_le = round(
            x.get(
                "ti_le_dung",
                0
            ) * 100
        )

        parts.append(
            f"{x.get('yccd', '')} "
            f"({x.get('muc_do', '')}, {ti_le}% đúng)"
        )

    return (
        "Các nội dung cần ưu tiên hiện tại: "
        + "; ".join(parts)
        + ". Lượt luyện tiếp theo sẽ tăng ưu tiên cho các nhóm này, "
          "đồng thời vẫn giữ một phần câu ở nội dung đã làm tốt."
    )


AI_STUDENT_FEEDBACK_SCHEMA = {
    "type": "object",
    "properties": {
        "nhan_xet_tong_quan": {
            "type": "string"
        },
        "loi_co_ban": {
            "type": "array",
            "items": {
                "type": "string"
            }
        },
        "can_on_lai": {
            "type": "array",
            "items": {
                "type": "string"
            }
        },
        "dan_do": {
            "type": "array",
            "items": {
                "type": "string"
            }
        },
        "ke_hoach_luot_tiep": {
            "type": "string"
        }
    },
    "required": [
        "nhan_xet_tong_quan",
        "loi_co_ban",
        "can_on_lai",
        "dan_do",
        "ke_hoach_luot_tiep"
    ]
}


def ai_co_van_sau_bai_lam(
    hoc_sinh_id,
    ten_hoc_sinh,
    ban_ghi,
    profile
):
    """
    AI KHÔNG chấm lại.
    Chỉ dùng đáp án/lời giải đã được ngân hàng duyệt để giải thích
    và phân tích lỗi.
    """
    sai = [
        x
        for x in ban_ghi.get(
            "chi_tiet",
            []
        )
        if not x.get(
            "dung_toan_cau"
        )
    ]

    ds_gui = []

    for x in sai[:12]:
        q = x.get(
            "cau_snapshot",
            {}
        )

        ds_gui.append({
            "dang_cau": q.get(
                "dang_cau",
                ""
            ),
            "cau_hoi": q.get(
                "cau_hoi",
                ""
            ),
            "tinh_huong": q.get(
                "tinh_huong",
                ""
            ),
            "hoc_sinh_tra_loi": x.get(
                "hoc_sinh_tra_loi",
                ""
            ),
            "dap_an_chuan": x.get(
                "dap_an_chuan",
                ""
            ),
            "giai_thich_chuan": q.get(
                "giai_thich",
                ""
            ),
            "don_vi_danh_gia": x.get(
                "don_vi_danh_gia",
                []
            ),
            "nhan_dinh_meta": q.get(
                "nhan_dinh_meta",
                []
            )
        })

    weak = tom_tat_diem_yeu(
        profile,
        5
    )

    prompt = f"""
Bạn là cố vấn học tập Sinh học THPT.

QUY TẮC QUAN TRỌNG:
- KHÔNG chấm lại bài.
- Kết quả đúng/sai, đáp án chuẩn và lời giải chuẩn đã được hệ thống xác định.
- Không được thay đổi đáp án chuẩn.
- Chỉ giải thích dễ hiểu, phân tích lỗi tư duy/kiến thức của học sinh,
  và đưa ra lời khuyên cụ thể.
- Không được bịa thêm dữ kiện ngoài câu hỏi/lời giải chuẩn.
- Nếu dữ liệu chưa đủ để kết luận một loại lỗi, hãy nói thận trọng.

Học sinh: {ten_hoc_sinh or hoc_sinh_id}

KẾT QUẢ LƯỢT VỪA LÀM:
{json.dumps({
    "diem": ban_ghi.get("diem_chinh_thuc", ban_ghi.get("diem")),
    "thang_diem": ban_ghi.get("thang_diem", 10),
    "ti_le_dung_don_vi": ban_ghi.get("ti_le_dung_don_vi"),
    "che_do": ban_ghi.get("che_do"),
    "pham_vi": ban_ghi.get("pham_vi"),
    "cac_cau_sai": ds_gui
}, ensure_ascii=False)}

HỒ SƠ ĐIỂM YẾU TÍCH LŨY:
{json.dumps(weak, ensure_ascii=False)}

PHÂN TÍCH NĂNG LỰC / CHỈ BÁO TỪ HỆ THỐNG:
{json.dumps(tom_tat_nang_luc_chi_bao_hoc_sinh(profile), ensure_ascii=False)}

Hãy:
1. Nhận xét trước hết theo 3 thành phần năng lực.
2. Chỉ rõ thành phần năng lực nào còn yếu hoặc đang tiến bộ.
3. Nếu có dữ liệu, nêu chính xác chỉ báo/hành vi năng lực còn yếu.
4. Sau đó mới nêu YCCĐ/kiến thức nền cần ôn.
5. Dặn học sinh cần làm gì cụ thể để cải thiện năng lực/chỉ báo đó.
6. Đề xuất lượt luyện tiếp theo phù hợp.
"""

    try:
        response = goi_gemini_co_retry(
            prompt,
            AI_STUDENT_FEEDBACK_SCHEMA
        )

        return json.loads(
            response.text
        ), ""

    except Exception as e:
        if la_loi_429(e):
            return None, (
                "AI cố vấn đang tạm giới hạn 429. "
                "Kết quả và hồ sơ cá nhân hóa vẫn đã được lưu; "
                "có thể bấm phân tích lại sau."
            )

        return None, str(e)


def tao_de_tu_mau_cho_hoc_sinh(
    bank,
    mau,
    profile,
    seed=None,
    ca_nhan_hoa=True
):
    specs = list(
        mau.get(
            "ma_tran",
            []
        )
        or []
    )

    # Với từng ô ma trận, rút theo ưu tiên cá nhân nhưng vẫn giữ đặc tả.
    rng = random.Random(
        seed or str(uuid.uuid4())
    )

    ids = set()
    ds_de = []
    thieu = []

    for i, spec in enumerate(
        specs,
        start=1
    ):
        can = int(
            spec.get(
                "Số câu",
                0
            )
        )

        pool = [
            q
            for q in bank
            if cau_khop_dac_ta(
                q,
                spec
            )
            and str(
                q.get(
                    "id",
                    ""
                )
            ) not in ids
        ]

        if ca_nhan_hoa:
            # Ôn theo form: chỉ cá nhân hóa BÊN TRONG chính ô ma trận.
            # Không đổi số câu, dạng câu, YCCĐ, mức độ hay phạm vi của ô.
            lay = rut_cau_ca_nhan_hoa(
                pool,
                can,
                profile,
                seed=str(
                    rng.random()
                )
            )
        else:
            # Mô phỏng kiểm tra: vẫn đúng ô ma trận nhưng không dùng hồ sơ cá nhân.
            pool_random = list(pool)
            rng.shuffle(pool_random)
            lay = pool_random[:can]

        for q in lay:
            ds_de.append(q)
            ids.add(
                str(
                    q.get(
                        "id",
                        ""
                    )
                )
            )

        if len(lay) < can:
            thieu.append({
                "Dòng": i,
                "Cần": can,
                "Có": len(lay),
                "Thiếu": can - len(lay),
                "Dạng": spec.get(
                    "Dạng câu hỏi",
                    ""
                ),
                "YCCĐ": spec.get(
                    "YCCĐ",
                    ""
                )
            })

    rng.shuffle(
        ds_de
    )

    return ds_de, thieu




def trang_thai_hoc_tap_unit(stat):
    """
    Trạng thái đơn giản, ổn định và không dùng AI.
    Không coi 1 lần đúng là đã thành thạo.
    """
    so_lan = int(stat.get("so_lan", 0) or 0)
    ti_le = float(stat.get("ti_le_dung", 0) or 0)

    if so_lan < 2:
        return "🔴 Cần củng cố"
    if so_lan >= 4 and ti_le >= 0.80:
        return "🟢 Đã đạt"
    if ti_le >= 0.60:
        return "🟡 Đang tiến bộ"
    return "🔴 Cần củng cố"


def diem_khop_cau_voi_muc_tieu_hom_nay(q, muc_tieu):
    """
    Điểm khớp giữa một câu và một mục tiêu củng cố.

    YCCĐ là điều kiện chính. Mức độ được ưu tiên tiếp theo.
    Thành phần năng lực/chỉ báo chỉ là tín hiệu ĐIỀU TIẾT, tuyệt đối
    không dùng như điều kiện loại cứng. Nhờ đó câu được xây dựng cho
    tốt nghiệp THPT vẫn có thể xuất hiện trong lượt ôn nếu đúng YCCĐ.
    """
    yccd = str(muc_tieu.get("yccd", "")).strip()
    muc_do = str(muc_tieu.get("muc_do", "")).strip()
    nl = str(muc_tieu.get("nang_luc", "")).strip()
    cb = str(muc_tieu.get("chi_bao", "")).strip()

    if not yccd:
        return 0.0

    yccd_norm = " ".join(yccd.split()).casefold()
    best = 0.0

    units = list(metadata_don_vi_cau(q) or [])

    # Một số câu Đúng/Sai cũ có thể chưa đủ metadata từng ý;
    # vẫn cho phép dùng YCCĐ cấp câu làm phương án dự phòng.
    if not any(str(u.get("yccd", "")).strip() for u in units):
        units = [{
            "yccd": q.get("yccd", ""),
            "muc_do": q.get("muc_do", ""),
            "nang_luc": q.get("thanh_phan_nang_luc", "")
        }]

    for unit in units:
        unit_yccd_norm = " ".join(
            str(unit.get("yccd", "")).strip().split()
        ).casefold()

        if unit_yccd_norm != yccd_norm:
            continue

        score = 100.0

        # Mức độ là ưu tiên thứ hai sau YCCĐ.
        if muc_do and str(unit.get("muc_do", "")).strip() == muc_do:
            score += 24.0

        # Năng lực chỉ cộng điểm mềm, KHÔNG loại câu khác năng lực.
        if nl and str(unit.get("nang_luc", "")).strip() == nl:
            score += 6.0

        # Chỉ báo cũng chỉ dùng để ưu tiên nhẹ.
        q_cb = str(
            q.get("chi_bao", "")
            or q.get("hanh_vi_nang_luc", "")
        ).strip()
        if cb and q_cb and q_cb == cb:
            score += 10.0

        best = max(best, score)

    return best


def goi_y_hoc_tap_hom_nay(profile, bank, limit=3):
    """
    Chọn các mục tiêu học tập ưu tiên từ yếu -> mạnh hơn.

    Một YCCĐ chỉ xuất hiện một lần trong danh sách ưu tiên để học sinh
    không nhìn thấy nhiều mục giống nhau chỉ vì khác thành phần năng lực.

    Khi kiểm tra xem ngân hàng có câu phù hợp hay không, app ưu tiên YCCĐ
    và mức độ; KHÔNG bắt buộc trùng thành phần năng lực. Đây là chủ đích
    để toàn bộ ngân hàng chung, kể cả câu tốt nghiệp, vẫn phục vụ ôn tập.
    """
    stats = [
        dict(x)
        for x in profile.get("stats", {}).values()
    ]

    for s in stats:
        s["trang_thai_hoc_tap"] = trang_thai_hoc_tap_unit(s)

    # Điểm ưu tiên càng nhỏ càng cần học trước.
    stats.sort(
        key=lambda s: (
            0 if s["trang_thai_hoc_tap"].startswith("🔴") else
            1 if s["trang_thai_hoc_tap"].startswith("🟡") else 2,
            float(s.get("mastery", 0.5)),
            -int(s.get("so_lan", 0) or 0)
        )
    )

    goi_y = []
    yccd_da_co = set()

    for s in stats:
        yccd = str(s.get("yccd", "")).strip()
        if not yccd or yccd in yccd_da_co:
            continue

        pool = [
            q
            for q in bank
            if (
                q.get("trang_thai", "Đã duyệt") != "Ngừng sử dụng"
                and q.get("duoc_dung_luyen_hs", True)
                and diem_khop_cau_voi_muc_tieu_hom_nay(q, s) > 0
            )
        ]

        if pool:
            item = dict(s)
            item["so_cau_kha_dung"] = len(pool)
            goi_y.append(item)
            yccd_da_co.add(yccd)

        if len(goi_y) >= int(limit):
            break

    return goi_y


def tao_pool_luyen_goi_y_hom_nay(profile, bank, cac_muc_tieu):
    """
    Tạo pool CHUNG cho một lượt luyện gợi ý hôm nay.

    Không còn cơ chế "mỗi lượt chỉ lấy Ưu tiên 1". Mọi YCCĐ đang yếu
    (Ưu tiên 1 -> 2 -> 3 -> ...) đều có thể đóng góp câu vào cùng lượt.
    Câu tốt nghiệp nằm trong ngân hàng chung vẫn được dùng nếu khớp YCCĐ.
    """
    cac_muc_tieu = list(cac_muc_tieu or [])
    pool = []

    for q in bank:
        if q.get("trang_thai", "Đã duyệt") == "Ngừng sử dụng":
            continue
        if not q.get("duoc_dung_luyen_hs", True):
            continue

        best_rank = None
        best_score = 0.0

        for rank, muc_tieu in enumerate(cac_muc_tieu, start=1):
            score = diem_khop_cau_voi_muc_tieu_hom_nay(
                q,
                muc_tieu
            )

            if score > best_score:
                best_score = score
                best_rank = rank

        if best_rank is None or best_score <= 0:
            continue

        q2 = dict(q)
        q2["_hs_uu_tien_rank"] = int(best_rank)
        q2["_hs_diem_khop_muc_tieu"] = float(best_score)
        pool.append(q2)

    return pool


def tao_pool_luyen_bu_hom_nay(profile, bank, muc_tieu):
    """Tương thích với lời gọi cũ: một mục tiêu vẫn tạo được pool."""
    return tao_pool_luyen_goi_y_hom_nay(
        profile,
        bank,
        [muc_tieu] if muc_tieu else []
    )


def _trong_so_uu_tien_hom_nay(muc_tieu, rank):
    """Trọng số mềm: càng yếu và càng đứng đầu thì càng nhiều câu."""
    ti_le = float(muc_tieu.get("ti_le_dung", 0) or 0)
    ti_le = min(1.0, max(0.0, ti_le))

    # Có một phần nền để mục tiêu 2,3... vẫn xuất hiện trong cùng lượt.
    do_can_cung_co = 0.25 + 0.75 * (1.0 - ti_le)
    he_so_thu_tu = 0.72 ** max(0, int(rank) - 1)

    return do_can_cung_co * he_so_thu_tu


def _phan_bo_so_cau_uu_tien_hom_nay(cac_muc_tieu, so_cau, kha_dung):
    """
    Phân bổ số câu theo nhiều ưu tiên.
    Ví dụ 10 câu thường xấp xỉ 5-3-2 nếu ba ưu tiên đầu yếu rõ ràng.
    Nếu một ưu tiên thiếu câu, phần thiếu tự động dồn sang ưu tiên tiếp theo.
    """
    n = max(0, int(so_cau))
    if n <= 0:
        return {}

    active = []
    for rank, muc_tieu in enumerate(cac_muc_tieu, start=1):
        co = int(kha_dung.get(rank, 0) or 0)
        if co <= 0:
            continue
        active.append((rank, muc_tieu, co))

    if not active:
        return {}

    weights = {
        rank: _trong_so_uu_tien_hom_nay(muc_tieu, rank)
        for rank, muc_tieu, _ in active
    }
    tong_w = sum(weights.values()) or 1.0

    raw = {
        rank: n * weights[rank] / tong_w
        for rank, _, _ in active
    }
    quota = {
        rank: min(int(raw[rank]), co)
        for rank, _, co in active
    }

    # Nếu đủ số câu, đảm bảo vài ưu tiên đầu cùng xuất hiện trong lượt.
    if n >= 3:
        for rank, _, co in active[:min(3, len(active))]:
            if co > 0 and quota.get(rank, 0) == 0:
                quota[rank] = 1

    # Nếu vừa đảm bảo tối thiểu làm vượt tổng, rút từ ưu tiên thấp hơn trước.
    while sum(quota.values()) > n:
        for rank, _, _ in reversed(active):
            if sum(quota.values()) <= n:
                break
            min_keep = 1 if (n >= 3 and rank in [x[0] for x in active[:3]]) else 0
            if quota.get(rank, 0) > min_keep:
                quota[rank] -= 1

    # Chia phần còn lại theo mức thiếu so với tỷ trọng lý tưởng.
    # Nếu một ưu tiên đã chạm trần số câu khả dụng, phần thiếu sẽ tự
    # phân phối sang các ưu tiên còn lại theo trọng số, không dồn cứng.
    while sum(quota.values()) < n:
        candidates = []

        for rank, _, co in active:
            q_now = quota.get(rank, 0)
            if q_now >= co:
                continue

            deficit = raw.get(rank, 0) - q_now
            fair_score = weights.get(rank, 0) / max(1, q_now + 1)
            candidates.append((
                1 if deficit > 0 else 0,
                deficit,
                fair_score,
                -rank,
                rank
            ))

        if not candidates:
            break

        candidates.sort(reverse=True)
        *_, rank = candidates[0]
        quota[rank] = quota.get(rank, 0) + 1

    return quota


def rut_cau_goi_y_hom_nay(
    pool,
    so_cau,
    profile,
    cac_muc_tieu,
    dang_counts=None,
    seed=None
):
    """
    Rút một lượt luyện gồm NHIỀU ưu tiên trong cùng bộ câu.

    Thứ tự ưu tiên quyết định tỷ trọng, nhưng bên trong từng ưu tiên vẫn
    ưu tiên câu chưa gặp, câu từng sai và câu phù hợp mức độ. Thành phần
    năng lực chỉ là bonus mềm, không phải bộ lọc cứng.
    """
    rng = random.Random(seed or str(uuid.uuid4()))
    pool = list(pool or [])
    cac_muc_tieu = list(cac_muc_tieu or [])
    n = min(max(0, int(so_cau)), len(pool))

    if n <= 0:
        return []

    by_rank = {}
    for q in pool:
        rank = int(q.get("_hs_uu_tien_rank", 999) or 999)
        by_rank.setdefault(rank, []).append(q)

    kha_dung = {
        rank: len({str(q.get("id", "")) or str(id(q)) for q in ds})
        for rank, ds in by_rank.items()
    }

    quota = _phan_bo_so_cau_uu_tien_hom_nay(
        cac_muc_tieu,
        n,
        kha_dung
    )

    con_dang = dict(dang_counts or {})
    da_chon = []
    ids = set()

    def qid(q):
        return str(q.get("id", "")) or str(id(q))

    def score(q):
        return (
            diem_uu_tien_cau_ca_nhan(profile, q, rng)
            + float(q.get("_hs_diem_khop_muc_tieu", 0) or 0) * 0.35
            - int(q.get("_hs_uu_tien_rank", 99) or 99) * 1.5
        )

    def chon_tu_ds(ds, can):
        nonlocal da_chon, ids, con_dang
        ds = [q for q in ds if qid(q) not in ids]
        ds.sort(key=score, reverse=True)

        for _ in range(max(0, int(can))):
            if not ds or len(da_chon) >= n:
                break

            idx = None
            if con_dang:
                for j, q in enumerate(ds):
                    dang = q.get("dang_cau", "")
                    if int(con_dang.get(dang, 0) or 0) > 0:
                        idx = j
                        break

            if idx is None:
                idx = 0

            q = ds.pop(idx)
            q_id = qid(q)
            if q_id in ids:
                continue

            da_chon.append(q)
            ids.add(q_id)

            dang = q.get("dang_cau", "")
            if dang in con_dang and con_dang[dang] > 0:
                con_dang[dang] -= 1

    # Lấy theo quota Ưu tiên 1 -> 2 -> 3 -> ...
    for rank in sorted(quota):
        chon_tu_ds(
            by_rank.get(rank, []),
            quota.get(rank, 0)
        )

    # Nếu một ưu tiên thiếu dạng câu/câu khả dụng, bù từ mọi ưu tiên tiếp theo.
    if len(da_chon) < n:
        con_lai = [q for q in pool if qid(q) not in ids]
        con_lai.sort(
            key=lambda q: (
                -int(q.get("_hs_uu_tien_rank", 999) or 999),
                score(q)
            ),
            reverse=True
        )
        chon_tu_ds(
            con_lai,
            n - len(da_chon)
        )

    rng.shuffle(da_chon)

    # Xóa metadata tạm trước khi đưa câu vào phiên làm bài.
    ket_qua = []
    for q in da_chon[:n]:
        q2 = dict(q)
        q2.pop("_hs_uu_tien_rank", None)
        q2.pop("_hs_diem_khop_muc_tieu", None)
        ket_qua.append(q2)

    return ket_qua


def hien_thi_ke_hoach_hom_nay(profile, bank):
    goi_y_all = goi_y_hoc_tap_hom_nay(profile, bank, 10)

    st.markdown("---")
    st.markdown("## 🎯 HÔM NAY EM NÊN LUYỆN GÌ?")

    if not profile.get("tong_don_vi", 0):
        st.info(
            "App chưa có đủ dữ liệu cá nhân của em. "
            "Hãy làm một lượt ôn theo bài hoặc theo chương; "
            "sau đó hệ thống sẽ tự đề xuất nội dung phù hợp."
        )
        return []

    if not goi_y_all:
        st.success(
            "🌟 Hiện chưa có nội dung yếu rõ ràng có đủ câu trong ngân hàng. "
            "Em có thể tiếp tục ôn theo bài/chương hoặc luyện đề."
        )
        return []

    st.caption(
        "App chỉ hiển thị nội dung cần ưu tiên nhất. "
        "Các nội dung khác được thu gọn ở phía dưới."
    )

    uu_tien = goi_y_all[0]

    with st.container(border=True):
        st.markdown("### 📌 Nội dung ưu tiên")

        yccd = str(uu_tien.get("yccd", "")).strip()
        if yccd:
            st.write(f"**{yccd}**")

        st.write(
            f"**Năng lực:** "
            f"{uu_tien.get('nang_luc') or 'Chưa xác định'}"
        )

        cb = str(uu_tien.get("chi_bao", "")).strip()
        if cb and cb not in ["Chưa gắn chỉ báo", "Chưa xác định"]:
            st.write(f"**Chỉ báo/kĩ năng:** {cb}")

        st.write(
            f"**Mức độ:** {uu_tien.get('muc_do') or '—'}"
            f"  •  **Hiện tại:** "
            f"{uu_tien.get('ti_le_dung', 0) * 100:.0f}% đúng"
        )

    con_lai = goi_y_all[1:]

    if con_lai:
        with st.expander(
            f"📚 Xem thêm {len(con_lai)} nội dung cần luyện",
            expanded=False
        ):
            for j, x in enumerate(con_lai, start=2):
                st.markdown(f"**{j}. {x.get('yccd', '')}**")
                st.caption(
                    f"{x.get('nang_luc', '')} • "
                    f"{x.get('muc_do', '')} • "
                    f"{x.get('ti_le_dung', 0) * 100:.0f}% đúng"
                )

                cb = str(x.get("chi_bao", "")).strip()
                if cb and cb not in ["Chưa gắn chỉ báo", "Chưa xác định"]:
                    st.caption(f"Chỉ báo/kĩ năng: {cb}")

                st.markdown("---")

    # Trả toàn bộ danh sách ưu tiên cho bộ chọn câu phía dưới.
    # Giao diện vẫn có thể chỉ hiển thị phần đầu, nhưng lượt luyện sẽ
    # tự bù tiếp Ưu tiên 4, 5... nếu các ưu tiên đầu thiếu câu.
    return goi_y_all


def sap_xep_cau_hoi_theo_dang_hs(ds_cau):
    """
    Thứ tự cố định khi HS làm bài:
    1. Trắc nghiệm 4 lựa chọn
    2. Đúng / Sai
    3. Trả lời ngắn

    Stable sort: thứ tự ngẫu nhiên/cá nhân hóa bên trong từng dạng vẫn được giữ.
    """
    thu_tu = {
        "Trắc nghiệm 4 lựa chọn": 0,
        "Đúng / Sai": 1,
        "Trả lời ngắn": 2
    }

    return sorted(
        list(
            ds_cau or []
        ),
        key=lambda q: thu_tu.get(
            q.get(
                "dang_cau",
                ""
            ),
            99
        )
    )



# ==========================================================
# PHẠM VI LUYỆN HS CHO CÂU HẠT GIỐNG
# ==========================================================
def _cac_pham_vi_cua_cau_luyen_hs(q):
    """Trả các phạm vi Khối/Chương/Bài của câu, tương thích dữ liệu cũ."""
    ds = []
    seen = set()
    direct = {
        "khoi": str(q.get("khoi", "") or ""),
        "chuong": str(q.get("chuong", "") or ""),
        "bai": str(q.get("bai", "") or ""),
    }
    if direct["khoi"] or direct["chuong"] or direct["bai"]:
        key = (direct["khoi"], direct["chuong"], direct["bai"])
        seen.add(key)
        ds.append(direct)

    for pv in list(q.get("pham_vi_hat_giong", []) or []):
        if not isinstance(pv, dict):
            continue
        item = {
            "khoi": str(pv.get("khoi", "") or ""),
            "chuong": str(pv.get("chuong", "") or ""),
            "bai": str(pv.get("bai", "") or ""),
        }
        key = (item["khoi"], item["chuong"], item["bai"])
        if key not in seen and any(key):
            seen.add(key)
            ds.append(item)
    return ds


def _cau_thuoc_pham_vi_luyen_hs(q, khoi="", chuong="", bai=""):
    """Kiểm tra câu có thuộc phạm vi luyện HS; câu cũ giữ nguyên hành vi."""
    khoi = str(khoi or "")
    chuong = str(chuong or "")
    bai = str(bai or "")
    for pv in _cac_pham_vi_cua_cau_luyen_hs(q):
        if khoi and pv.get("khoi") != khoi:
            continue
        if chuong and pv.get("chuong") != chuong:
            continue
        if bai and pv.get("bai") != bai:
            continue
        return True
    return False


def _gia_tri_pham_vi_co_trong_bank(bank, field, khoi="", chuong=""):
    """Lấy danh sách Khối/Chương/Bài, kể cả câu Đ/S hạt giống đa YCCĐ."""
    values = set()
    for q in bank or []:
        for pv in _cac_pham_vi_cua_cau_luyen_hs(q):
            if khoi and pv.get("khoi") != str(khoi):
                continue
            if chuong and pv.get("chuong") != str(chuong):
                continue
            val = str(pv.get(field, "") or "").strip()
            if val:
                values.add(val)
    return sorted(values)


def hoc_sinh():

    # Chữ khu vực học sinh: lớn hơn, thoáng hơn để dễ đọc trên màn hình lớp học.
    st.markdown(
        """
        <style>
        /* Nội dung văn bản chính */
        [data-testid="stMainBlockContainer"] p,
        [data-testid="stMainBlockContainer"] li,
        [data-testid="stMainBlockContainer"] label {
            font-size: 1.08rem !important;
            line-height: 1.65 !important;
        }

        /* Tiêu đề */
        [data-testid="stMainBlockContainer"] h1 {
            font-size: 2.15rem !important;
            line-height: 1.25 !important;
        }
        [data-testid="stMainBlockContainer"] h2 {
            font-size: 1.75rem !important;
            line-height: 1.3 !important;
        }
        [data-testid="stMainBlockContainer"] h3 {
            font-size: 1.42rem !important;
            line-height: 1.4 !important;
        }
        [data-testid="stMainBlockContainer"] h4 {
            font-size: 1.20rem !important;
            line-height: 1.45 !important;
        }

        /* Radio, checkbox và lựa chọn đáp án */
        [data-testid="stRadio"] label,
        [data-testid="stCheckbox"] label {
            font-size: 1.08rem !important;
            line-height: 1.6 !important;
        }

        /* Ô nhập câu trả lời */
        [data-testid="stTextInput"] input,
        [data-testid="stTextArea"] textarea {
            font-size: 1.08rem !important;
        }

        /* Nút */
        [data-testid="stButton"] button,
        [data-testid="stDownloadButton"] button {
            font-size: 1.05rem !important;
            font-weight: 600 !important;
            min-height: 2.8rem !important;
        }

        /* Alert: success / warning / info / error */
        [data-testid="stAlert"] p,
        [data-testid="stAlert"] div {
            font-size: 1.08rem !important;
            line-height: 1.65 !important;
        }

        /* Caption vẫn nhỏ hơn nội dung nhưng không quá nhỏ */
        [data-testid="stCaptionContainer"] p {
            font-size: 0.98rem !important;
            line-height: 1.55 !important;
        }

        /* Metric ở trang kết quả */
        [data-testid="stMetricLabel"] {
            font-size: 1.02rem !important;
        }
        [data-testid="stMetricValue"] {
            font-size: 2rem !important;
        }

        /* Bố cục khu vực học sinh */
        .tram-page-shell {
            max-width:1160px; margin:0 auto;
        }
        .tram-guide-line {
            max-width:1040px; margin:-.12rem auto 1.05rem auto; text-align:center;
            color:#5b6b7f; font-weight:650; font-size:.98rem;
        }
        .hs-login-shell { max-width:860px; margin:0 auto .85rem auto; }
        .hs-login-banner {
            margin:.1rem 0 .68rem 0; padding:1rem 1.12rem; border-radius:18px;
            background:linear-gradient(135deg,#f4f9ff,#f2fbf6);
            border:1px solid #9fd6c5; box-shadow:0 9px 24px rgba(15,90,100,.075);
            text-align:left;
        }
        .hs-login-title { font-size:1.22rem; font-weight:900; color:#0b5f69; }
        .hs-login-note { margin-top:.25rem; color:#5a6d7e; font-size:.96rem; line-height:1.55; }
        [data-testid="stTextInput"] input {
            border:2px solid #54a895 !important; border-radius:14px !important;
            background:#ffffff !important; min-height:3.15rem !important;
        }
        [data-testid="stTextInput"] input:focus {
            border-color:#087f8c !important; box-shadow:0 0 0 .18rem rgba(8,127,140,.12) !important;
        }
        </style>
        """,
        unsafe_allow_html=True
    )


    hien_thi_dau_trang_tram_sinh_hoc("hocsinh")
    st.markdown(
        '<div class="tram-guide-line">'
        'Luyện theo bài • theo chương • đề giáo viên • tốt nghiệp THPT • cá nhân hóa theo kết quả của chính em'
        '</div>',
        unsafe_allow_html=True
    )

    # ======================================================
    # HÀM PHỤ HIỂN THỊ / CHẤM
    # ======================================================
    def bo_nhan_dinh(noi_dung, ky_hieu):
        noi_dung = str(
            noi_dung
        ).strip()

        for nhan in [
            f"{ky_hieu})",
            f"{ky_hieu}.",
            f"{ky_hieu}:"
        ]:
            if noi_dung.lower().startswith(
                nhan.lower()
            ):
                return noi_dung[
                    len(nhan):
                ].strip()

        return noi_dung


    def lay_dap_an_dung_sai(q):
        dap = {}

        for ky, nd in zip(
            ["a", "b", "c", "d"],
            q.get(
                "nhan_dinh_meta",
                []
            )
            or []
        ):
            dap[ky] = str(
                nd.get(
                    "dap_an",
                    ""
                )
            ).strip()

        return dap


    def reset_bai_hs():
        st.session_state.hs_de_thi = []
        st.session_state.hs_dang_lam = False
        st.session_state.hs_da_nop = False
        st.session_state.hs_da_luu_ket_qua = False
        st.session_state.hs_ban_ghi_hien_tai = None
        st.session_state.hs_ai_feedback = None
        st.session_state.hs_ai_feedback_attempted = False
        st.session_state.hs_ai_feedback_error = ""
        st.session_state.hs_bat_dau_epoch = None
        st.session_state.hs_bat_dau_iso = ""
        st.session_state.hs_bat_dau_hien_thi = ""
        st.session_state.hs_nop_bai_epoch = None
        st.session_state.hs_nop_bai_iso = ""
        st.session_state.hs_nop_bai_hien_thi = ""
        st.session_state.hs_thoi_gian_quy_dinh_phut = 0
        st.session_state.hs_han_nop_epoch = None
        st.session_state.hs_xem_lai_kiem_tra = False

        for key in list(
            st.session_state.keys()
        ):
            if (
                key.startswith("hs_answer_")
                or key.startswith("hs_short_answer_")
            ):
                del st.session_state[key]


    # ======================================================
    # SESSION
    # ======================================================
    defaults = {
        "hs_de_thi": [],
        "hs_dang_lam": False,
        "hs_da_nop": False,
        "hs_da_luu_ket_qua": False,
        "hs_ban_ghi_hien_tai": None,
        "hs_ai_feedback": None,
        "hs_ai_feedback_attempted": False,
        "hs_ai_feedback_error": "",
        "hs_pham_vi_hien_tai": {},
        "hs_che_do_hien_tai": "",
        "hs_bat_dau_epoch": None,
        "hs_bat_dau_iso": "",
        "hs_bat_dau_hien_thi": "",
        "hs_nop_bai_epoch": None,
        "hs_nop_bai_iso": "",
        "hs_nop_bai_hien_thi": "",
        "hs_thoi_gian_quy_dinh_phut": 0,
        "hs_han_nop_epoch": None,
        "hs_xem_lai_kiem_tra": False
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    # ======================================================
    # NHẬN DIỆN HỌC SINH
    # ======================================================
    st.markdown(
        """
        <div class="hs-login-shell">
            <div class="hs-login-banner">
                <div class="hs-login-title">🔐 ĐĂNG NHẬP HỌC SINH</div>
                <div class="hs-login-note">Nhập mã học sinh do giáo viên cấp để vào đúng hồ sơ học tập, bài luyện và tiến độ của em.</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    _, col_hs_login, _ = st.columns([1, 2.4, 1])
    with col_hs_login:
        hs_id = st.text_input(
            "Mã học sinh",
            value=st.session_state.get(
                "hs_id",
                ""
            ),
            placeholder="VD: 12A2-001",
            key="hs_id"
        )

    if not str(
        hs_id
    ).strip():
        st.info(
            "Nhập mã học sinh do giáo viên cấp."
        )
        return

    hs_info = tim_hoc_sinh_theo_ma(
        hs_id
    )

    if not hs_info:
        st.error(
            "Không tìm thấy mã học sinh này. "
            "Hãy kiểm tra lại hoặc liên hệ giáo viên."
        )
        return

    if hs_info.get(
        "trang_thai"
    ) == "Tạm khóa":
        st.warning(
            "Tài khoản học sinh này đang tạm khóa."
        )
        return

    hs_id_chuan = str(
        hs_info.get(
            "ma_hoc_sinh",
            ""
        )
    ).strip().upper()

    hs_ten = str(
        hs_info.get(
            "ho_ten",
            ""
        )
    ).strip()

    hs_lop = str(
        hs_info.get(
            "lop",
            ""
        )
    ).strip()

    hs_khoi_ds = str(
        hs_info.get(
            "khoi",
            ""
        )
    ).strip()

    st.success(
        f"👋 Xin chào **{hs_ten}**"
    )

    info1, info2, info3 = st.columns(3)

    with info1:
        st.write(
            f"**Lớp:** {hs_lop or 'Chưa xác định'}"
        )

    with info2:
        st.write(
            f"**Khối:** {hs_khoi_ds or 'Chưa xác định'}"
        )

    with info3:
        st.write(
            f"**Mã học sinh:** {hs_id_chuan}"
        )

    profile = tao_ho_so_tu_lich_su(
        hs_id_chuan
    )

    # Hồ sơ cá nhân hóa vẫn được tính và lưu ngầm để chọn câu phù hợp.
    # Không hiển thị bảng YCCĐ/mức độ/năng lực cho học sinh ở màn hình chính.

    # ======================================================
    # NGÂN HÀNG
    # ======================================================
    # Ngân hàng dùng cho học sinh là NGÂN HÀNG CHUNG:
    # gồm câu ôn tập/kiểm tra và câu được xây dựng cho tốt nghiệp THPT.
    # Mặc định mọi câu đã duyệt đều được dùng luyện HS, trừ khi GV tắt cờ này.
    bank = [
        q
        for q in (doc_ngan_hang() + doc_ngan_hang_tot_nghiep_thuc_te())
        if (
            q.get(
                "trang_thai",
                "Đã duyệt"
            ) not in {"Ngừng sử dụng", "Thiếu đáp án", "Cần GV xem"}
            and q.get(
                "duoc_dung_luyen_hs",
                True
            )
        )
    ]

    if not bank:
        st.warning(
            "Ngân hàng chưa có câu đã duyệt."
        )
        return

    # Xếp hạng cá nhân chỉ hiện ở màn hình đầu; HS không xem danh sách bạn khác.
    if (
        not st.session_state.hs_dang_lam
        and not st.session_state.hs_da_nop
    ):
        hien_thi_xep_hang_ca_nhan_hs(
            hs_id_chuan,
            hs_lop
        )

    # ======================================================
    # DASHBOARD CÁ NHÂN - CHỈ HIỆN Ở MÀN HÌNH ĐẦU
    # ======================================================
    hs_goi_y_hom_nay = goi_y_hoc_tap_hom_nay(
        profile,
        bank,
        3
    )

    if (
        not st.session_state.hs_dang_lam
        and not st.session_state.hs_da_nop
    ):
        hs_goi_y_hom_nay = hien_thi_ke_hoach_hom_nay(
            profile,
            bank
        )

    # ======================================================
    # CHỌN CHẾ ĐỘ LUYỆN
    # ======================================================
    if not st.session_state.hs_dang_lam:

        st.markdown("---")
        st.subheader(
            "🎯 Chọn cách luyện"
        )

        st.caption(
            "Ôn theo bài/chương được cá nhân hóa mạnh. "
            "Ôn theo ma trận và tốt nghiệp vẫn giữ nguyên form đề."
        )

        che_do = st.radio(
            "Hình thức luyện",
            [
                "🎯 Luyện theo gợi ý hôm nay",
                "📖 Ôn theo bài",
                "📚 Ôn theo chương",
                "📝 Đề GV / ma trận",
                "🧪 Kiểm tra theo ma trận GV",
                "🎓 Luyện tốt nghiệp THPT",
                "📈 Lịch sử & tiến bộ"
            ],
            horizontal=True,
            key="hs_mode"
        )

        ds_khoi = _gia_tri_pham_vi_co_trong_bank(
            bank,
            "khoi"
        )

        pool = []
        pham_vi = {}
        so_cau = 10
        dang_counts = None
        ten_luot = che_do

        # --------------------------------------------------
        # LUYỆN THEO GỢI Ý HÔM NAY
        # --------------------------------------------------
        if che_do == "🎯 Luyện theo gợi ý hôm nay":
            if not hs_goi_y_hom_nay:
                st.info(
                    "Chưa có mục tiêu cá nhân đủ dữ liệu. "
                    "Hãy làm một lượt ôn theo bài/chương trước."
                )
                pool = []
            else:
                st.info(
                    "🎯 App sẽ tự tạo **một lượt luyện gồm nhiều câu hỏi**, "
                    "ưu tiên nhiều hơn cho phần yếu nhất. Các **Ưu tiên 1 → 2 → 3…** "
                    "được phối hợp trong cùng lượt; sau mỗi lượt, thứ tự và tỷ trọng "
                    "sẽ được tính lại tự động."
                )

                # Pool chung của nhiều ưu tiên, không còn khóa vào riêng Ưu tiên 1.
                pool = tao_pool_luyen_goi_y_hom_nay(
                    profile,
                    bank,
                    hs_goi_y_hom_nay
                )

                # Luôn giữ thanh chọn số câu. Nếu Ưu tiên 1 thiếu câu,
                # app tự bù từ Ưu tiên 2, 3, 4... trong chính pool này.
                if pool:
                    so_toi_da = min(
                        30,
                        len(pool)
                    )

                    so_cau = st.slider(
                        "Số câu luyện",
                        min_value=1,
                        max_value=so_toi_da,
                        value=min(10, so_toi_da),
                        step=1,
                        key="hs_today_n"
                    )

                    if len(pool) < 10:
                        st.caption(
                            f"Hiện có {len(pool)} câu phù hợp trong toàn bộ các ưu tiên. "
                            "App đã tự mở rộng sang các ưu tiên tiếp theo; không giới hạn "
                            "lượt luyện ở riêng Ưu tiên 1."
                        )
                else:
                    so_cau = 0

                dang_counts = (
                    phan_bo_3_dang_tu_dong(
                        int(so_cau),
                        pool
                    )
                    if pool
                    else None
                )

                pham_vi = {
                    "kieu": "goi_y_hom_nay",
                    "cac_uu_tien": [
                        {
                            "thu_tu": i,
                            "yccd": x.get("yccd", ""),
                            "muc_do": x.get("muc_do", ""),
                            "nang_luc": x.get("nang_luc", ""),
                            "chi_bao": x.get("chi_bao", ""),
                            "ti_le_dung": x.get("ti_le_dung", 0)
                        }
                        for i, x in enumerate(
                            hs_goi_y_hom_nay,
                            start=1
                        )
                    ]
                }

                ten_luot = "Luyện theo gợi ý hôm nay"

                st.success(
                    "App ưu tiên **YCCĐ/kiến thức đang yếu**, câu chưa gặp và câu từng sai. "
                    "Thành phần năng lực chỉ dùng để điều tiết nhẹ, **không dùng để loại cứng câu**. "
                    "Vì vậy câu trong ngân hàng tốt nghiệp vẫn được dùng để ôn nếu phù hợp nội dung."
                )

        # --------------------------------------------------
        # ÔN THEO BÀI / CHƯƠNG
        # --------------------------------------------------
        elif che_do in [
            "📖 Ôn theo bài",
            "📚 Ôn theo chương"
        ]:

            if (
                hs_khoi_ds
                and hs_khoi_ds in ds_khoi
            ):
                khoi = hs_khoi_ds
                st.write(
                    f"**Khối:** {khoi}"
                )
            else:
                khoi = st.selectbox(
                    "Khối",
                    ds_khoi,
                    key="hs_scope_grade"
                )

            bank_khoi = [
                q
                for q in bank
                if _cau_thuoc_pham_vi_luyen_hs(
                    q,
                    khoi=khoi
                )
            ]

            ds_chuong = _gia_tri_pham_vi_co_trong_bank(
                bank_khoi,
                "chuong",
                khoi=khoi
            )

            chuong = st.selectbox(
                "Chương",
                ds_chuong,
                key="hs_scope_chapter"
            )

            bank_chuong = [
                q
                for q in bank_khoi
                if _cau_thuoc_pham_vi_luyen_hs(
                    q,
                    khoi=khoi,
                    chuong=chuong
                )
            ]

            if che_do == "📖 Ôn theo bài":
                ds_bai = _gia_tri_pham_vi_co_trong_bank(
                    bank_chuong,
                    "bai",
                    khoi=khoi,
                    chuong=chuong
                )

                bai = st.selectbox(
                    "Bài",
                    ds_bai,
                    key="hs_scope_lesson"
                )

                pool = [
                    q
                    for q in bank_chuong
                    if _cau_thuoc_pham_vi_luyen_hs(
                        q,
                        khoi=khoi,
                        chuong=chuong,
                        bai=bai
                    )
                ]

                pham_vi = {
                    "khoi": khoi,
                    "chuong": chuong,
                    "bai": bai
                }

                ten_luot = (
                    "Ôn theo bài: "
                    + bai
                )

            else:
                pool = bank_chuong

                pham_vi = {
                    "khoi": khoi,
                    "chuong": chuong
                }

                ten_luot = (
                    "Ôn theo chương: "
                    + chuong
                )

            if pool:
                so_cau = st.select_slider(
                    "Số câu mỗi lượt",
                    options=[
                        x
                        for x in [
                            10,
                            20,
                            30
                        ]
                        if x <= len(pool)
                    ]
                    or [
                        min(
                            10,
                            len(pool)
                        )
                    ],
                    value=(
                        10
                        if len(pool) >= 10
                        else min(
                            10,
                            len(pool)
                        )
                    ),
                    key="hs_scope_n"
                )

                dang_counts = phan_bo_3_dang_tu_dong(
                    int(
                        so_cau
                    ),
                    pool
                )

                st.info(
                    "App tự phối hợp 3 dạng câu theo ngân hàng hiện có. "
                    "Mỗi lượt ưu tiên câu chưa làm và các YCCĐ/mức độ học sinh đang yếu."
                )

                st.write(
                    "**Cơ cấu lượt này:** "
                    + " • ".join(
                        f"{dang}: {n}"
                        for dang, n in dang_counts.items()
                    )
                )

        # --------------------------------------------------
        # ĐỀ GV / MA TRẬN
        # --------------------------------------------------
        elif che_do == "📝 Đề GV / ma trận":
            ds_mau = doc_json_list(EXAM_TEMPLATE_PATH)
            ds_de = doc_json_list(EXAM_PATH)

            nhom_on = []
            if ds_mau:
                nhom_on.append("📋 Ôn theo mẫu ma trận")
            if ds_de:
                nhom_on.append("📝 Ôn theo đề GV đã lưu")

            lua_chon = []
            if nhom_on:
                loai_nguon_on = st.radio(
                    "Chọn loại tài liệu ôn tập",
                    nhom_on,
                    horizontal=True,
                    key="hs_teacher_exam_source_type"
                )

                if loai_nguon_on.startswith("📋"):
                    for x in ds_mau:
                        lua_chon.append({
                            "label": (
                                "📋 " + x.get("ten_mau", "Không tên")
                                + " • " + str(x.get("khoi", ""))
                                + f" • PB{int(x.get('phien_ban', 1) or 1)}"
                                + (f" • {x.get('ngay_cap_nhat', x.get('ngay_tao', ''))}" if x.get('ngay_cap_nhat') or x.get('ngay_tao') else "")
                            ),
                            "type": "template",
                            "data": x
                        })
                else:
                    for x in ds_de:
                        lua_chon.append({
                            "label": (
                                "📝 " + x.get("ten_de", x.get("ma_de", "Không tên"))
                                + (f" • Mã {x.get('ma_de', '')}" if x.get("ma_de") else "")
                                + (f" • {x.get('ngay_tao', '')}" if x.get("ngay_tao") else "")
                            ),
                            "type": "exam",
                            "data": x
                        })

            if not lua_chon:
                st.info(
                    "GV chưa lưu mẫu ma trận hoặc mã đề nào."
                )
                pool = []
            else:
                labels = [
                    x["label"]
                    for x in lua_chon
                ]

                label = st.selectbox(
                    "Chọn đề / mẫu",
                    labels,
                    key="hs_teacher_exam"
                )

                selected = next(
                    x
                    for x in lua_chon
                    if x["label"] == label
                )

                if selected["type"] == "template":
                    cach_su_dung_mau = st.radio(
                        "Cách luyện với mẫu đề",
                        [
                            "Ôn tập theo form – cá nhân hóa trong từng ô ma trận",
                            "Mô phỏng kiểm tra – không cá nhân hóa"
                        ],
                        horizontal=True,
                        key="hs_template_mode"
                    )

                    ca_nhan_hoa_mau = cach_su_dung_mau.startswith(
                        "Ôn tập theo form"
                    )

                    if ca_nhan_hoa_mau:
                        st.info(
                            "Giữ nguyên toàn bộ ma trận. Hệ thống chỉ ưu tiên "
                            "câu phù hợp với điểm yếu của em **trong từng ô ma trận**."
                        )
                    else:
                        st.info(
                            "Mô phỏng kiểm tra: giữ nguyên ma trận và rút câu "
                            "không dựa vào hồ sơ cá nhân."
                        )

                    ds_tmp, thieu = tao_de_tu_mau_cho_hoc_sinh(
                        bank,
                        selected["data"],
                        profile,
                        seed=str(
                            uuid.uuid4()
                        ),
                        ca_nhan_hoa=ca_nhan_hoa_mau
                    )

                    if thieu:
                        st.warning(
                            "Ngân hàng đang thiếu một số ô của mẫu này."
                        )
                        st.dataframe(
                            pd.DataFrame(
                                thieu
                            ),
                            use_container_width=True,
                            hide_index=True
                        )
                        pool = []
                    else:
                        pool = ds_tmp

                else:
                    # Đề cụ thể đã lưu: giữ nguyên nội dung đề.
                    pool = list(
                        selected["data"].get(
                            "cau_hoi",
                            []
                        )
                        or []
                    )

                so_cau = len(
                    pool
                )

                du_lieu_de_chon = selected.get("data", {}) or {}
                pham_vi = {
                    "mau_de": label,
                    "loai_nguon": selected.get("type", ""),
                    "mau_id": (
                        du_lieu_de_chon.get("id", "")
                        if selected.get("type") == "template"
                        else ""
                    ),
                    "de_id": (
                        du_lieu_de_chon.get("id", "")
                        if selected.get("type") == "exam"
                        else ""
                    ),
                    "ma_de": du_lieu_de_chon.get("ma_de", ""),
                    "ten_mau": du_lieu_de_chon.get("ten_mau", ""),
                    "mau_phien_ban": int(du_lieu_de_chon.get("phien_ban", 1) or 1),
                    "thoi_gian_phut": int(du_lieu_de_chon.get("thoi_gian", 0) or 0)
                }

                ten_luot = label

                if pool:
                    diem_form_gv = tinh_tong_diem_toi_da_hs(
                        pool,
                        CHE_DO_DE_GV
                    )
                    st.info(
                        f"Điểm tối đa của đề/form này: **{diem_form_gv:.2f} điểm** "
                        "(4 lựa chọn 0,25/câu; Đúng/Sai 0,25/ý; Trả lời ngắn 0,50/câu). "
                        "HS được chấm đúng trên tổng điểm mà GV đã thiết kế trong ma trận."
                    )

        # --------------------------------------------------
        # KIỂM TRA THEO MA TRẬN GV - KHÔNG CÁ NHÂN HÓA
        # Mỗi HS chỉ 1 lần / đợt, có giờ mở - hạn cuối bắt đầu - giờ kết thúc đợt.
        # --------------------------------------------------
        elif che_do == CHE_DO_KIEM_TRA_MA_TRAN:
            ds_dot_all = [
                d for d in doc_dot_kiem_tra_ma_tran()
                if dot_kiem_tra_phu_hop_hoc_sinh(
                    d,
                    hs_lop,
                    hs_khoi_ds
                )
            ]

            if not ds_dot_all:
                st.info(
                    "Hiện chưa có đợt kiểm tra nào được giao cho lớp của em. "
                    "GV cần tạo **Đợt kiểm tra** từ một mẫu ma trận trước."
                )
                pool = []
            else:
                # Giữ các đợt đang/sắp mở hoặc vừa hết giờ vào; đợt đã kết thúc
                # chỉ hiện nếu chính HS đã nộp để em còn xem điểm/đáp án khi GV mở.
                ds_dot_hien = []
                for d in ds_dot_all:
                    trang_thai_d, _ = trang_thai_dot_kiem_tra(d)
                    da_nop_d = tim_luot_kiem_tra_da_nop(
                        hs_id_chuan,
                        d.get("id", "")
                    )
                    if trang_thai_d != "Đã kết thúc" or da_nop_d:
                        ds_dot_hien.append(d)

                if not ds_dot_hien:
                    st.info(
                        "Các đợt kiểm tra của lớp đã kết thúc và em chưa có bài đã nộp để xem lại."
                    )
                    pool = []
                else:
                    ds_dot_hien.sort(
                        key=lambda d: str(d.get("mo_tu_iso", "")),
                        reverse=True
                    )

                    labels_dot = []
                    for d in ds_dot_hien:
                        tt, _ = trang_thai_dot_kiem_tra(d)
                        labels_dot.append(
                            f"{d.get('ten_dot', 'Kiểm tra')} • {tt} • "
                            f"Mở {fmt_vn_datetime(d.get('mo_tu_iso', ''))} → "
                            f"Kết thúc {fmt_vn_datetime(d.get('dong_luc_iso', ''))}"
                        )

                    label_dot = st.selectbox(
                        "Chọn đợt kiểm tra",
                        labels_dot,
                        key="hs_matrix_test_assignment"
                    )
                    idx_dot_hs = labels_dot.index(label_dot)
                    dot_mt = ds_dot_hien[idx_dot_hs]
                    dot_id = str(dot_mt.get("id", ""))
                    trang_thai_dot, dang_mo_dot = trang_thai_dot_kiem_tra(dot_mt)
                    luot_da_nop = tim_luot_kiem_tra_da_nop(
                        hs_id_chuan,
                        dot_id
                    )

                    st.info(
                        f"**{dot_mt.get('ten_dot', 'Đợt kiểm tra')}**  \n"
                        f"Ma trận: {dot_mt.get('ten_mau', '')} • "
                        f"Lớp: {dot_mt.get('lop_ap_dung', 'Tất cả')} • "
                        f"Thời gian làm: **{int(dot_mt.get('thoi_gian_lam_phut', 0) or 0)} phút**"
                    )
                    han_bat_dau_hs = han_cuoi_bat_dau_dot_kiem_tra(dot_mt)
                    st.caption(
                        f"Mở từ **{fmt_vn_datetime(dot_mt.get('mo_tu_iso', ''))}** • "
                        f"Hạn cuối bắt đầu **{fmt_vn_datetime(han_bat_dau_hs)}** • "
                        f"Kết thúc đợt **{fmt_vn_datetime(dot_mt.get('dong_luc_iso', ''))}** • "
                        f"Trạng thái: **{trang_thai_dot}**"
                    )

                    if luot_da_nop:
                        diem_cu = float(
                            luot_da_nop.get(
                                "diem_chinh_thuc",
                                luot_da_nop.get("diem", 0)
                            ) or 0
                        )
                        thang_cu = float(luot_da_nop.get("thang_diem", 10) or 10)
                        st.warning(
                            "✅ Em **đã nộp bài ở đợt kiểm tra này**. "
                            "Mỗi học sinh chỉ được làm 1 lần nên hệ thống không cho bắt đầu lại."
                        )
                        sc1, sc2, sc3 = st.columns(3)
                        sc1.metric("Điểm", f"{diem_cu:.2f}/{thang_cu:.0f}")
                        sc2.metric(
                            "Thời gian làm",
                            f"{float(luot_da_nop.get('thoi_luong_lam_phut', 0) or 0):.1f} phút"
                        )
                        sc3.metric(
                            "Nộp bài",
                            str(luot_da_nop.get("nop_bai_luc", luot_da_nop.get("thoi_gian", ""))) or "—"
                        )

                        if bool(dot_mt.get("mo_dap_an", False)):
                            st.success(
                                "🔓 Giáo viên đã mở đáp án. Em có thể xem lại bài đã nộp."
                            )
                            if st.button(
                                "📋 XEM LẠI BÀI & ĐÁP ÁN",
                                type="primary",
                                use_container_width=True,
                                key=f"hs_matrix_review_{dot_id}"
                            ):
                                de_da_lam = [
                                    (x.get("cau_snapshot", {}) or {})
                                    for x in (luot_da_nop.get("chi_tiet", []) or [])
                                    if x.get("cau_snapshot")
                                ]
                                st.session_state.hs_de_thi = de_da_lam
                                st.session_state.hs_dang_lam = True
                                st.session_state.hs_da_nop = True
                                st.session_state.hs_da_luu_ket_qua = True
                                st.session_state.hs_ban_ghi_hien_tai = luot_da_nop
                                st.session_state.hs_che_do_hien_tai = CHE_DO_KIEM_TRA_MA_TRAN
                                st.session_state.hs_pham_vi_hien_tai = luot_da_nop.get("pham_vi", {}) or {}
                                st.session_state.hs_ten_luot = luot_da_nop.get("ten_luot", "Kiểm tra")
                                st.session_state.hs_xem_lai_kiem_tra = True
                                st.rerun()
                        else:
                            st.info(
                                "🔒 Giáo viên **chưa mở đáp án**. Hiện em chỉ được xem điểm và thời gian làm bài."
                            )
                        pool = []

                    elif not dang_mo_dot:
                        if trang_thai_dot == "Sắp mở":
                            st.warning(
                                "⏳ Đợt kiểm tra chưa đến giờ mở. Em chưa thể bắt đầu bài."
                            )
                        elif trang_thai_dot == "Hết giờ vào":
                            st.error(
                                "⛔ Đã quá hạn bắt đầu bài. Vì mỗi học sinh phải được làm đủ thời gian quy định, "
                                "hệ thống không nhận lượt bắt đầu mới trong phần thời gian còn lại của đợt."
                            )
                        elif trang_thai_dot == "Đã kết thúc":
                            st.error(
                                "⛔ Đợt kiểm tra đã kết thúc. Hệ thống không cho bắt đầu bài mới."
                            )
                        else:
                            st.warning("Đợt kiểm tra chưa được cấu hình thời gian hợp lệ.")
                        pool = []

                    else:
                        mau_mt = {
                            "id": dot_mt.get("mau_id", dot_id),
                            "ten_mau": dot_mt.get("ten_mau", "Ma trận"),
                            "khoi": dot_mt.get("khoi", ""),
                            "phien_ban": int(dot_mt.get("mau_phien_ban", 1) or 1),
                            "thoi_gian": int(dot_mt.get("thoi_gian_lam_phut", 0) or 0),
                            "ma_tran": list(dot_mt.get("ma_tran_snapshot", []) or [])
                        }
                        seed_mt = seed_dot_kiem_tra_hoc_sinh(
                            dot_mt,
                            hs_id_chuan
                        )

                        ds_tmp, thieu_mt = tao_de_tu_mau_cho_hoc_sinh(
                            bank,
                            mau_mt,
                            profile={},
                            seed=seed_mt,
                            ca_nhan_hoa=False
                        )

                        st.info(
                            "🧪 **Kiểm tra chính thức:** câu của các học sinh có thể khác nhau nhưng "
                            "cùng ma trận, cùng nội dung/YCCĐ, mức độ, dạng câu và cơ cấu điểm. "
                            "Hệ thống **không dùng hồ sơ cá nhân hóa** để chọn câu."
                        )
                        st.caption(
                            f"Mã đề của em: **{seed_mt[:8].upper()}** • "
                            "Mã này cố định cho chính em trong đợt kiểm tra này."
                        )

                        if thieu_mt:
                            st.warning(
                                "Ngân hàng đang thiếu câu ở một số ô ma trận. "
                                "Hệ thống không lấy câu sai nội dung/mức độ để bù nên tạm chưa cho kiểm tra."
                            )
                            st.dataframe(
                                pd.DataFrame(thieu_mt),
                                use_container_width=True,
                                hide_index=True
                            )
                            pool = []
                        else:
                            pool = ds_tmp

                        so_cau = len(pool)
                        pham_vi = {
                            "kieu": "kiem_tra_ma_tran_gv",
                            "dot_kiem_tra_id": dot_id,
                            "ten_dot_kiem_tra": dot_mt.get("ten_dot", ""),
                            "mau_id": dot_mt.get("mau_id", ""),
                            "ten_mau": dot_mt.get("ten_mau", ""),
                            "mau_phien_ban": int(dot_mt.get("mau_phien_ban", 1) or 1),
                            "ma_de_hs": seed_mt[:8].upper(),
                            "khong_ca_nhan_hoa": True,
                            "thoi_gian_phut": int(dot_mt.get("thoi_gian_lam_phut", 0) or 0),
                            "mo_tu_iso": dot_mt.get("mo_tu_iso", ""),
                            "dong_luc_iso": dot_mt.get("dong_luc_iso", ""),
                            "ma_tran": list(dot_mt.get("ma_tran_snapshot", []) or [])
                        }
                        ten_luot = (
                            "Kiểm tra: "
                            + str(dot_mt.get("ten_dot", dot_mt.get("ten_mau", "Ma trận")))
                        )

                        if pool:
                            diem_form_mt = tinh_tong_diem_toi_da_hs(
                                pool,
                                CHE_DO_KIEM_TRA_MA_TRAN
                            )
                            if abs(diem_form_mt - 10.0) <= 1e-9:
                                st.success(
                                    f"Cơ cấu điểm của mã đề này: **{diem_form_mt:.2f}/10,00 điểm**."
                                )
                            else:
                                st.error(
                                    f"Ma trận hiện cho **{diem_form_mt:.2f}/10,00 điểm**; "
                                    "GV cần sửa ma trận trước khi kiểm tra."
                                )
                                pool = []

        # --------------------------------------------------
        # LỊCH SỬ & TIẾN BỘ
        # --------------------------------------------------
        elif che_do == "📈 Lịch sử & tiến bộ":
            pool = []

            lich_su_hs = lay_lich_su_cua_hoc_sinh(
                hs_id_chuan
            )

            st.subheader(
                "📈 Tiến bộ của em"
            )

            if not lich_su_hs:
                st.info(
                    "Em chưa có lượt luyện nào. "
                    "Hãy bắt đầu bằng Ôn theo bài hoặc Ôn theo chương."
                )
            else:
                diem_tb = sum(
                    float(x.get("diem", 0))
                    for x in lich_su_hs
                ) / len(lich_su_hs)

                xh_ca_nhan = lay_xep_hang_ca_nhan_hs(
                    hs_id_chuan,
                    hs_lop
                )

                p1, p2, p3, p4 = st.columns(4)

                with p1:
                    st.metric(
                        "Số lượt đã làm",
                        len(lich_su_hs)
                    )

                with p2:
                    st.metric(
                        "Điểm trung bình",
                        f"{diem_tb:.1f}/10"
                    )

                with p3:
                    st.metric(
                        "Mức làm chủ tích lũy",
                        f"{(xh_ca_nhan or {}).get('tich_luy', profile.get('ti_le_dung', 0) * 100):.0f}%"
                    )

                with p4:
                    if xh_ca_nhan:
                        st.metric(
                            "Hạng hiện tại",
                            f"{xh_ca_nhan['hang']}/{xh_ca_nhan['si_so_co_du_lieu']}"
                        )
                    else:
                        st.metric("Hạng hiện tại", "—")

                tb_xu_huong_hs = tinh_tien_bo_hoc_sinh(lich_su_hs)
                xh_text = tb_xu_huong_hs.get("xu_huong", "Chưa có dữ liệu")
                if xh_text in ["Tiến bộ rõ", "Có tiến bộ"]:
                    st.success(f"📈 **Xu hướng cá nhân: {xh_text}**")
                elif xh_text in ["Giảm rõ", "Có dấu hiệu giảm"]:
                    st.warning(f"📉 **Xu hướng cá nhân: {xh_text}**")
                else:
                    st.info(f"➡️ **Xu hướng cá nhân: {xh_text}**")

                st.markdown(
                    "#### 🧠 Nội dung cần ưu tiên"
                )

                weak = tom_tat_diem_yeu(
                    profile,
                    5
                )

                if weak:
                    for idx_w, x in enumerate(
                        weak,
                        start=1
                    ):
                        st.write(
                            f"**{idx_w}.** {x.get('yccd', '')}"
                        )
                        st.caption(
                            f"Mức độ: {x.get('muc_do', '')} • "
                            f"Tỉ lệ đúng: {x.get('ti_le_dung', 0) * 100:.0f}%"
                        )
                else:
                    st.success(
                        "Chưa phát hiện điểm yếu ổn định."
                    )

                st.markdown(
                    "#### 🕘 Các lượt gần đây"
                )

                st.dataframe(
                    pd.DataFrame([
                        {
                            "Thời gian": x.get("thoi_gian", ""),
                            "Hình thức": x.get("che_do", ""),
                            "Bài luyện": x.get("ten_luot", ""),
                            "Số câu": x.get("tong_so_cau", 0),
                            "Thời lượng": (
                                f"{float(x.get('thoi_luong_lam_phut', 0) or 0):.1f} phút"
                                if x.get("thoi_luong_lam_phut") is not None
                                else ""
                            ),
                            "Tỉ lệ đúng": (
                                f"{x.get('ti_le_dung_don_vi', 0):.0f}%"
                            ),
                            "Điểm": (
                                f"{float(x.get('diem_chinh_thuc', x.get('diem', 0)) or 0):.2f}/"
                                f"{float(x.get('thang_diem', 10) or 10):.0f}"
                            )
                        }
                        for x in reversed(
                            lich_su_hs[-10:]
                        )
                    ]),
                    use_container_width=True,
                    hide_index=True
                )

        # --------------------------------------------------
        # LUYỆN TỐT NGHIỆP
        # --------------------------------------------------
        else:
            pool12 = [
                q for q in doc_ngan_hang_tot_nghiep_thuc_te()
                if cau_tot_nghiep_du_dieu_kien_su_dung(q)
            ]
            pool_on_tap_12 = [
                q for q in doc_ngan_hang()
                if cau_on_tap_bo_sung_du_dieu_kien_tot_nghiep(q)
            ]

            dang_counts = {
                "Trắc nghiệm 4 lựa chọn": 18,
                "Đúng / Sai": 4,
                "Trả lời ngắn": 6
            }
            so_cau = 28

            _, thieu_form = rut_de_tot_nghiep_tu_de_that(
                pool12,
                seed="kiem-tra-hs",
                bank_on_tap=pool_on_tap_12
            )
            if thieu_form:
                for item in thieu_form:
                    st.warning(
                        f"Thiếu {item.get('Thiếu', 0)} câu dạng {item.get('Dạng','')} "
                        "để tạo đủ một đề tốt nghiệp không trùng câu."
                    )
            else:
                pool = pool12

            pham_vi = {
                "khoi": "Khối 12",
                "form": "Tốt nghiệp THPT – đề thật",
                "thoi_gian_phut": 50
            }
            ten_luot = "Luyện tốt nghiệp THPT"

            st.info(
                "Mỗi lượt là một mã đề mới **18/4/6**. Nguồn chính vẫn là Ngân hàng tốt nghiệp; "
                f"app chỉ bổ sung tối đa **{GRAD_MAX_CAU_TU_NH_ON_TAP}/28 câu** phù hợp từ Ngân hàng ôn tập Khối 12. "
                "Không lặp câu trong cùng mã đề và vẫn giữ nguyên ảnh/bảng của câu gốc."
            )

        if pool:
            st.caption(
                f"Ngân hàng phù hợp: {len(pool)} câu."
            )

            nut_bat_dau_label = (
                "🧪 BẮT ĐẦU KIỂM TRA"
                if che_do == CHE_DO_KIEM_TRA_MA_TRAN
                else "🚀 BẮT ĐẦU LƯỢT LUYỆN"
            )

            if st.button(
                nut_bat_dau_label,
                type="primary",
                use_container_width=True
            ):

                # Kiểm tra chính thức: chặn lại ở phía máy chủ trước khi mở đề.
                if che_do == CHE_DO_KIEM_TRA_MA_TRAN:
                    dot_id_start = str((pham_vi or {}).get("dot_kiem_tra_id", ""))
                    dot_start = tim_dot_kiem_tra_theo_id(dot_id_start)
                    if not dot_start:
                        st.error("Đợt kiểm tra đã bị GV thu hồi hoặc không còn tồn tại.")
                        return
                    if tim_luot_kiem_tra_da_nop(hs_id_chuan, dot_id_start):
                        st.error("Em đã nộp bài ở đợt này. Mỗi học sinh chỉ được làm 1 lần.")
                        return
                    trang_thai_start, dot_dang_mo = trang_thai_dot_kiem_tra(dot_start)
                    if not dot_dang_mo:
                        if trang_thai_start == "Hết giờ vào":
                            st.error(
                                "Đã quá hạn bắt đầu bài. Hệ thống chặn lượt mới để bảo đảm mỗi học sinh "
                                "được làm đủ thời gian quy định trước khi đợt kiểm tra kết thúc."
                            )
                        elif trang_thai_start == "Sắp mở":
                            st.error("Đợt kiểm tra chưa đến giờ mở.")
                        elif trang_thai_start == "Đã kết thúc":
                            st.error("Đợt kiểm tra đã kết thúc.")
                        else:
                            st.error("Đợt kiểm tra hiện không nằm trong khung giờ được phép bắt đầu.")
                        return

                if che_do in {CHE_DO_DE_GV, CHE_DO_KIEM_TRA_MA_TRAN}:
                    # Đề GV/ma trận: pool đã được xử lý đúng từng ô.
                    # Riêng kiểm tra ma trận đã rút KHÔNG cá nhân hóa bằng seed học sinh ổn định.
                    de_thi = list(
                        pool
                    )

                elif che_do == CHE_DO_TOT_NGHIEP:
                    # Rút nguyên bản 18/4/6 từ ngân hàng đề thật, không trùng câu.
                    de_thi, thieu_luot = rut_de_tot_nghiep_tu_de_that(
                        pool,
                        seed=str(uuid.uuid4()),
                        bank_on_tap=pool_on_tap_12
                    )
                    if thieu_luot:
                        st.error("Ngân hàng vừa thay đổi và không còn đủ form 18/4/6.")
                        st.stop()

                elif che_do == "🎯 Luyện theo gợi ý hôm nay":
                    # Một lượt phối hợp nhiều ưu tiên. Ưu tiên 1 nhiều câu nhất,
                    # nhưng vẫn có câu của Ưu tiên 2, 3... và tự bù khi thiếu.
                    de_thi = rut_cau_goi_y_hom_nay(
                        pool,
                        int(
                            so_cau
                        ),
                        profile,
                        hs_goi_y_hom_nay,
                        dang_counts=dang_counts,
                        seed=str(
                            uuid.uuid4()
                        )
                    )

                else:
                    # Ôn theo bài/chương: cá nhân hóa mạnh theo điểm yếu,
                    # câu từng sai và câu chưa gặp.
                    de_thi = rut_cau_ca_nhan_hoa(
                        pool,
                        int(
                            so_cau
                        ),
                        profile,
                        dang_counts=dang_counts,
                        seed=str(
                            uuid.uuid4()
                        )
                    )

                # Học sinh luôn làm theo đúng 3 phần:
                # 4 lựa chọn -> Đúng/Sai -> Trả lời ngắn.
                de_thi = sap_xep_cau_hoi_theo_dang_hs(
                    de_thi
                )

                st.session_state.hs_de_thi = de_thi
                st.session_state.hs_dang_lam = True
                st.session_state.hs_da_nop = False
                st.session_state.hs_da_luu_ket_qua = False
                st.session_state.hs_ai_feedback = None
                st.session_state.hs_ai_feedback_attempted = False
                st.session_state.hs_ai_feedback_error = ""
                st.session_state.hs_che_do_hien_tai = (
                    che_do
                )
                st.session_state.hs_pham_vi_hien_tai = (
                    pham_vi
                )
                st.session_state.hs_ten_luot = (
                    ten_luot
                )

                # Ghi nhận thời điểm bắt đầu theo giờ Việt Nam để hiển thị và chuyển về GV.
                bat_dau_dt = bay_gio_viet_nam()
                st.session_state.hs_bat_dau_epoch = time.time()
                st.session_state.hs_bat_dau_iso = bat_dau_dt.isoformat()
                st.session_state.hs_bat_dau_hien_thi = bat_dau_dt.strftime(
                    "%d/%m/%Y %H:%M:%S"
                )
                st.session_state.hs_nop_bai_epoch = None
                st.session_state.hs_nop_bai_iso = ""
                st.session_state.hs_nop_bai_hien_thi = ""
                st.session_state.hs_thoi_gian_quy_dinh_phut = int(
                    (pham_vi or {}).get("thoi_gian_phut", 0) or 0
                )
                st.session_state.hs_xem_lai_kiem_tra = False

                # Hạn nộp cá nhân = thời điểm bắt đầu + đúng thời gian làm bài.
                # Kiểm tra chính thức chỉ cho phép bắt đầu trước hạn cuối, nên HS hợp lệ
                # luôn có đủ thời gian và vẫn kết thúc không muộn hơn giờ kết thúc đợt.
                han_nop = None
                if st.session_state.hs_thoi_gian_quy_dinh_phut > 0:
                    han_nop = (
                        st.session_state.hs_bat_dau_epoch
                        + st.session_state.hs_thoi_gian_quy_dinh_phut * 60
                    )
                if che_do == CHE_DO_KIEM_TRA_MA_TRAN:
                    dong_dt = parse_iso_vn((pham_vi or {}).get("dong_luc_iso", ""))
                    if dong_dt:
                        dong_epoch = dong_dt.timestamp()
                        han_nop = min(han_nop, dong_epoch) if han_nop else dong_epoch
                st.session_state.hs_han_nop_epoch = han_nop

                for key in list(
                    st.session_state.keys()
                ):
                    if (
                        key.startswith(
                            "hs_answer_"
                        )
                        or key.startswith(
                            "hs_short_answer_"
                        )
                    ):
                        del st.session_state[
                            key
                        ]

                st.rerun()

    # ======================================================
    # ĐANG LÀM
    # ======================================================
    if (
        st.session_state.hs_dang_lam
        and not st.session_state.hs_da_nop
    ):
        de_thi = st.session_state.hs_de_thi

        st.markdown("---")
        st.header(
            "📝 "
            + str(
                st.session_state.get(
                    "hs_ten_luot",
                    "Bài luyện"
                )
            )
        )

        st.caption(
            f"{len(de_thi)} câu"
        )

        # Đồng hồ hiển thị trực tiếp trong lúc làm bài.
        bat_dau_epoch = st.session_state.get("hs_bat_dau_epoch")
        thoi_gian_quy_dinh = int(
            st.session_state.get("hs_thoi_gian_quy_dinh_phut", 0) or 0
        )
        han_nop_epoch = st.session_state.get("hs_han_nop_epoch")

        # Hết giờ thì khóa bài ở phía server. Radio/text input của Streamlit tạo rerun,
        # nên ngay thao tác đầu tiên sau hạn nộp hệ thống sẽ tự chuyển sang chấm.
        if (
            st.session_state.get("hs_che_do_hien_tai") == CHE_DO_KIEM_TRA_MA_TRAN
            and han_nop_epoch
            and time.time() >= float(han_nop_epoch)
        ):
            nop_dt = bay_gio_viet_nam()
            st.session_state.hs_nop_bai_epoch = time.time()
            st.session_state.hs_nop_bai_iso = nop_dt.isoformat()
            st.session_state.hs_nop_bai_hien_thi = nop_dt.strftime("%d/%m/%Y %H:%M:%S")
            st.session_state.hs_da_nop = True
            st.warning("⏰ Đã hết thời gian. Hệ thống tự khóa và nộp bài.")
            st.rerun()

        if bat_dau_epoch:
            st.caption(
                "Bắt đầu: **"
                + str(st.session_state.get("hs_bat_dau_hien_thi", ""))
                + "**"
                + (
                    f" • Thời gian quy định: **{thoi_gian_quy_dinh} phút**"
                    if thoi_gian_quy_dinh > 0
                    else " • Không giới hạn thời gian"
                )
            )

            limit_seconds = thoi_gian_quy_dinh * 60
            if han_nop_epoch:
                limit_seconds = max(0, int(float(han_nop_epoch) - float(bat_dau_epoch)))
            timer_title = (
                "⏳ Thời gian còn lại"
                if limit_seconds > 0
                else "⏱️ Thời gian đã làm"
            )

            components.html(
                f"""
                <div style="font-family:Arial,sans-serif;border:1px solid #dbeafe;border-radius:14px;
                            padding:12px 16px;background:#f8fbff;display:flex;justify-content:space-between;
                            align-items:center;gap:12px;">
                  <div style="font-weight:700;color:#1e3a8a;">{timer_title}</div>
                  <div id="bio_timer" style="font-size:1.55rem;font-weight:800;color:#0f172a;">00:00</div>
                </div>
                <script>
                const start = {float(bat_dau_epoch) * 1000.0};
                const limit = {int(limit_seconds) * 1000};
                function fmt(sec) {{
                    sec = Math.max(0, Math.floor(sec));
                    const h = Math.floor(sec / 3600);
                    const m = Math.floor((sec % 3600) / 60);
                    const s = sec % 60;
                    return (h > 0 ? String(h).padStart(2,'0') + ':' : '')
                         + String(m).padStart(2,'0') + ':' + String(s).padStart(2,'0');
                }}
                function tick() {{
                    const elapsed = (Date.now() - start) / 1000;
                    const value = limit > 0 ? Math.max(0, limit / 1000 - elapsed) : elapsed;
                    const el = document.getElementById('bio_timer');
                    if (el) {{
                        el.textContent = fmt(value);
                        if (limit > 0 && value <= 300) el.style.color = '#b91c1c';
                    }}
                }}
                tick();
                setInterval(tick, 1000);
                </script>
                """,
                height=75
            )

        # Theo dõi nhanh số câu đã có thao tác trả lời.
        da_tra_loi = 0

        for idx_q, q_check in enumerate(
            de_thi,
            start=1
        ):
            dang_check = q_check.get(
                "dang_cau",
                ""
            )

            if dang_check == "Đúng / Sai":
                co_tra_loi = any(
                    str(
                        st.session_state.get(
                            f"hs_answer_{idx_q}_{j}",
                            ""
                        )
                    ).strip()
                    for j in range(
                        1,
                        5
                    )
                )
            elif dang_check == "Trả lời ngắn":
                co_tra_loi = bool(
                    str(
                        st.session_state.get(
                            f"hs_short_answer_{idx_q}",
                            ""
                        )
                    ).strip()
                )
            else:
                co_tra_loi = bool(
                    str(
                        st.session_state.get(
                            f"hs_answer_{idx_q}",
                            ""
                        )
                    ).strip()
                )

            if co_tra_loi:
                da_tra_loi += 1

        st.progress(
            (
                da_tra_loi / len(de_thi)
                if de_thi
                else 0
            )
        )

        st.caption(
            f"Đã trả lời **{da_tra_loi}/{len(de_thi)} câu**"
        )

        for i, q in enumerate(
            de_thi,
            start=1
        ):
            dang = q.get(
                "dang_cau",
                ""
            )

            st.markdown(
                f"### Câu {i}"
            )

            st.caption(
                f"Loại câu: {dang}"
            )

            if q.get(
                "tinh_huong"
            ):
                st.write(
                    "**Tình huống / dữ liệu:**"
                )
                st.write(
                    q.get(
                        "tinh_huong",
                        ""
                    )
                )

            st.write(
                q.get(
                    "cau_hoi",
                    ""
                )
            )

            if q.get("tai_nguyen_truc_quan") or q.get("du_lieu_truc_quan"):
                hien_thi_tai_nguyen_cau_tot_nghiep(q)

            if dang == "Trắc nghiệm 4 lựa chọn":
                st.radio(
                    "Chọn đáp án:",
                    q.get(
                        "lua_chon",
                        []
                    ),
                    index=None,
                    key=f"hs_answer_{i}"
                )

            elif dang == "Đúng / Sai":
                for j, (
                    ky,
                    nd
                ) in enumerate(
                    zip(
                        ["a", "b", "c", "d"],
                        q.get(
                            "nhan_dinh_meta",
                            []
                        )
                        or []
                    ),
                    start=1
                ):
                    st.markdown(
                        f"**{ky})** "
                        + bo_nhan_dinh(
                            nd.get(
                                "noi_dung",
                                ""
                            ),
                            ky
                        )
                    )

                    st.radio(
                        f"Ý {ky}:",
                        [
                            "Đúng",
                            "Sai"
                        ],
                        index=None,
                        horizontal=True,
                        key=f"hs_answer_{i}_{j}"
                    )

            elif dang == "Trả lời ngắn":
                st.text_input(
                    "Trả lời:",
                    key=f"hs_short_answer_{i}",
                    max_chars=4,
                    placeholder="VD: 3 | 25 | 0,25 | -2,5"
                )

            st.divider()

        con_trong = max(
            0,
            len(de_thi) - da_tra_loi
        )

        if con_trong:
            st.warning(
                f"Em còn **{con_trong} câu** chưa trả lời đầy đủ."
            )
        else:
            st.success(
                "✅ Em đã trả lời tất cả các câu."
            )

        if st.button(
            "✅ NỘP BÀI",
            type="primary",
            use_container_width=True
        ):
            nop_dt = datetime.now()
            st.session_state.hs_nop_bai_epoch = time.time()
            st.session_state.hs_nop_bai_iso = nop_dt.isoformat()
            st.session_state.hs_nop_bai_hien_thi = nop_dt.strftime(
                "%d/%m/%Y %H:%M:%S"
            )
            st.session_state.hs_da_nop = True
            st.rerun()

    # ======================================================
    # CHẤM & LƯU KẾT QUẢ
    # ======================================================
    if (
        st.session_state.hs_dang_lam
        and st.session_state.hs_da_nop
    ):
        de_thi = st.session_state.hs_de_thi

        chi_tiet = []
        tong_don_vi = 0
        dung_don_vi = 0
        tong_diem = 0.0
        tong_toi_da = 0.0

        che_do_hien_tai = st.session_state.get(
            "hs_che_do_hien_tai",
            ""
        )

        cham_tot_nghiep = (
            che_do_hien_tai
            == CHE_DO_TOT_NGHIEP
        )
        cham_de_gv = (
            che_do_hien_tai
            == CHE_DO_DE_GV
        )
        cham_kiem_tra_ma_tran = (
            che_do_hien_tai
            == CHE_DO_KIEM_TRA_MA_TRAN
        )
        cham_diem_co_dinh = la_che_do_cham_diem_co_dinh(
            che_do_hien_tai
        )

        for i, q in enumerate(
            de_thi,
            start=1
        ):
            dang = q.get(
                "dang_cau",
                ""
            )

            hs_text = ""
            dap_chuan = ""
            dung_toan_cau = False
            kq_tung_y = None
            diem_cau = 0.0

            if dang == "Đúng / Sai":
                meta = list(
                    q.get(
                        "nhan_dinh_meta",
                        []
                    )
                    or []
                )

                hs_parts = []
                dap_parts = []
                kq_tung_y = []

                for j, (
                    ky,
                    nd
                ) in enumerate(
                    zip(
                        ["a", "b", "c", "d"],
                        meta
                    ),
                    start=1
                ):
                    hs = str(
                        st.session_state.get(
                            f"hs_answer_{i}_{j}",
                            ""
                        )
                    ).strip()

                    dap = str(
                        nd.get(
                            "dap_an",
                            ""
                        )
                    ).strip()

                    ok = (
                        hs.casefold()
                        == dap.casefold()
                        and bool(
                            hs
                        )
                    )

                    hs_parts.append(
                        f"{ky}: {hs or 'Chưa trả lời'}"
                    )
                    dap_parts.append(
                        f"{ky}: {dap}"
                    )
                    kq_tung_y.append(
                        ok
                    )

                so_y_dung = sum(
                    1
                    for x in kq_tung_y
                    if x
                )

                hs_text = ", ".join(
                    hs_parts
                )
                dap_chuan = ", ".join(
                    dap_parts
                )
                dung_toan_cau = (
                    so_y_dung == 4
                )

                tong_don_vi += len(
                    kq_tung_y
                )
                dung_don_vi += so_y_dung

                diem_cau = tinh_diem_cau_hs(
                    dang,
                    dung_toan_cau,
                    so_y_dung,
                    che_do_hien_tai
                )

                tong_toi_da += diem_toi_da_cua_cau_hs(
                    q,
                    che_do_hien_tai
                )

            elif dang == "Trả lời ngắn":
                hs_text = str(
                    st.session_state.get(
                        f"hs_short_answer_{i}",
                        ""
                    )
                ).strip()

                dap_chuan = str(
                    q.get(
                        "dap_an",
                        ""
                    )
                ).strip()

                dung_toan_cau = (
                    hs_text == dap_chuan
                    and bool(
                        hs_text
                    )
                )

                tong_don_vi += 1
                dung_don_vi += int(
                    dung_toan_cau
                )
                diem_cau = tinh_diem_cau_hs(
                    dang,
                    dung_toan_cau,
                    0,
                    che_do_hien_tai
                )
                tong_toi_da += diem_toi_da_cua_cau_hs(
                    q,
                    che_do_hien_tai
                )

            else:
                raw = str(
                    st.session_state.get(
                        f"hs_answer_{i}",
                        ""
                    )
                ).strip()

                hs_text = raw

                dap_chuan = str(
                    q.get(
                        "dap_an",
                        ""
                    )
                ).strip()

                def lay_letter(s):
                    s = str(
                        s
                    ).strip()

                    return (
                        s[0].upper()
                        if s
                        and s[0].upper()
                        in [
                            "A",
                            "B",
                            "C",
                            "D"
                        ]
                        else s.casefold()
                    )

                dung_toan_cau = (
                    lay_letter(
                        hs_text
                    )
                    == lay_letter(
                        dap_chuan
                    )
                    and bool(
                        hs_text
                    )
                )

                tong_don_vi += 1
                dung_don_vi += int(
                    dung_toan_cau
                )
                diem_cau = tinh_diem_cau_hs(
                    dang,
                    dung_toan_cau,
                    0,
                    che_do_hien_tai
                )
                tong_toi_da += diem_toi_da_cua_cau_hs(
                    q,
                    che_do_hien_tai
                )

            tong_diem += diem_cau

            units = tao_don_vi_ket_qua_cau(
                q,
                ket_qua_tung_y=kq_tung_y,
                dung_toan_cau=dung_toan_cau
            )

            chi_tiet.append({
                "stt": i,
                "cau_id": q.get(
                    "id",
                    ""
                ),
                "dang_cau": dang,
                "hoc_sinh_tra_loi": hs_text,
                "dap_an_chuan": dap_chuan,
                "dung_toan_cau": dung_toan_cau,
                "diem_cau": diem_cau,
                "don_vi_danh_gia": units,
                "cau_snapshot": q
            })

        # Điểm nội bộ /10 vẫn được giữ để tương thích lịch sử,
        # xếp hạng và hồ sơ năng lực cũ.
        diem_10 = (
            tong_diem
            / tong_toi_da
            * 10
            if tong_toi_da > 0
            else 0
        )

        if cham_diem_co_dinh:
            diem_chinh_thuc = round(tong_diem, 2)
            thang_diem = thang_diem_chinh_thuc_hs(
                che_do_hien_tai,
                diem_toi_da_thuc_te=tong_toi_da
            )
        else:
            diem_chinh_thuc = round(diem_10, 2)
            thang_diem = 10.0

        # Thời gian làm bài được ghi cùng kết quả để GV theo dõi/xuất điểm.
        if not st.session_state.get("hs_nop_bai_epoch"):
            nop_dt_fallback = datetime.now()
            st.session_state.hs_nop_bai_epoch = time.time()
            st.session_state.hs_nop_bai_iso = nop_dt_fallback.isoformat()
            st.session_state.hs_nop_bai_hien_thi = nop_dt_fallback.strftime(
                "%d/%m/%Y %H:%M:%S"
            )

        bat_epoch = st.session_state.get("hs_bat_dau_epoch")
        nop_epoch = st.session_state.get("hs_nop_bai_epoch")
        thoi_luong_giay = (
            max(0.0, float(nop_epoch) - float(bat_epoch))
            if bat_epoch and nop_epoch
            else 0.0
        )

        ban_ghi = {
            "id": str(
                uuid.uuid4()
            ),
            "hoc_sinh_id": hs_id_chuan,
            "ho_ten": hs_ten,
            "bat_dau_luc": st.session_state.get("hs_bat_dau_hien_thi", ""),
            "bat_dau_iso": st.session_state.get("hs_bat_dau_iso", ""),
            "nop_bai_luc": st.session_state.get("hs_nop_bai_hien_thi", ""),
            "nop_bai_iso": st.session_state.get("hs_nop_bai_iso", ""),
            "thoi_luong_lam_giay": round(thoi_luong_giay, 1),
            "thoi_luong_lam_phut": round(thoi_luong_giay / 60.0, 2),
            "thoi_gian_quy_dinh_phut": int(
                st.session_state.get("hs_thoi_gian_quy_dinh_phut", 0) or 0
            ),
            "thoi_gian": datetime.now().strftime(
                "%d/%m/%Y %H:%M"
            ),
            "thoi_gian_iso": datetime.now().isoformat(),
            "che_do": che_do_hien_tai,
            "pham_vi": st.session_state.get(
                "hs_pham_vi_hien_tai",
                {}
            ),
            "ten_luot": st.session_state.get(
                "hs_ten_luot",
                ""
            ),
            "tong_so_cau": len(
                de_thi
            ),
            "tong_don_vi": tong_don_vi,
            "dung_don_vi": dung_don_vi,
            "ti_le_dung_don_vi": round(
                (
                    dung_don_vi
                    / tong_don_vi
                    * 100
                )
                if tong_don_vi > 0
                else 0,
                1
            ),
            # "diem" giữ chuẩn /10 cho tương thích các thống kê cũ.
            "diem": round(
                diem_10,
                2
            ),
            # Điểm học sinh/GV nhìn thấy theo đúng chế độ.
            "diem_chinh_thuc": diem_chinh_thuc,
            "thang_diem": thang_diem,
            "diem_toi_da_cau_hoi": round(tong_toi_da, 2),
            "chi_tiet": chi_tiet
        }

        if not st.session_state.hs_da_luu_ket_qua:
            # Chặn trùng lần cuối ở phía dữ liệu đối với kiểm tra chính thức.
            existing_kiem_tra = None
            if che_do_hien_tai == CHE_DO_KIEM_TRA_MA_TRAN:
                dot_id_save = str(
                    (st.session_state.get("hs_pham_vi_hien_tai", {}) or {}).get(
                        "dot_kiem_tra_id",
                        ""
                    )
                )
                existing_kiem_tra = tim_luot_kiem_tra_da_nop(
                    hs_id_chuan,
                    dot_id_save
                )

            if existing_kiem_tra:
                ban_ghi = existing_kiem_tra
            else:
                ds_ls = doc_lich_su_hoc_sinh()
                ds_ls.append(
                    ban_ghi
                )
                luu_lich_su_hoc_sinh(
                    ds_ls
                )

            st.session_state.hs_ban_ghi_hien_tai = (
                ban_ghi
            )
            st.session_state.hs_da_luu_ket_qua = True

        else:
            ban_ghi = (
                st.session_state.hs_ban_ghi_hien_tai
                or ban_ghi
            )

        # Hồ sơ SAU lượt vừa làm
        profile_moi = tao_ho_so_tu_lich_su(
            hs_id_chuan
        )

        st.markdown("---")
        st.header(
            "🏁 KẾT QUẢ BÀI LÀM"
        )

        if che_do_hien_tai in {
            CHE_DO_DE_GV,
            CHE_DO_KIEM_TRA_MA_TRAN,
            CHE_DO_TOT_NGHIEP
        }:
            st.success(
                "📤 Điểm, thời gian làm bài và chi tiết đúng/sai đã được lưu vào dữ liệu lớp để giáo viên theo dõi."
            )

        st.caption(
            "Tiến bộ không được xác nhận chỉ bằng một câu đúng. "
            "App theo dõi lặp lại cùng năng lực/chỉ báo qua nhiều câu và nhiều lượt."
        )

        st.caption(
            "Kết quả được chấm bằng đáp án chuẩn đã có trong ngân hàng. "
            "AI chỉ hỗ trợ giải thích và tư vấn sau khi chấm."
        )

        if che_do_hien_tai == CHE_DO_DE_GV:
            st.info(
                "📝 Cách tính đề GV/ma trận: **0,25 điểm/câu 4 lựa chọn; "
                "0,25 điểm/ý Đúng-Sai; 0,50 điểm/câu Trả lời ngắn**. "
                "Phần app chấm tự động tính trên **7 điểm**; 3 điểm tự luận do GV chấm riêng."
            )
        elif che_do_hien_tai == CHE_DO_KIEM_TRA_MA_TRAN:
            st.info(
                "🧪 Kiểm tra theo ma trận GV: **0,25 điểm/câu 4 lựa chọn; "
                "0,25 điểm/ý Đúng-Sai; 0,50 điểm/câu Trả lời ngắn**, tổng **10 điểm**."
            )
        elif che_do_hien_tai == CHE_DO_TOT_NGHIEP:
            st.info(
                "🎓 Đề tốt nghiệp: 4 lựa chọn **0,25 điểm/câu**; Trả lời ngắn **0,25 điểm/câu**; "
                "Đúng-Sai theo số ý đúng: **1 ý = 0,10; 2 ý = 0,25; 3 ý = 0,50; 4 ý = 1,00 điểm**."
            )

        # Kiểm tra chính thức: sau khi nộp chỉ hiện ĐIỂM + THỜI GIAN
        # cho đến khi GV mở đáp án ở đúng đợt kiểm tra.
        if che_do_hien_tai == CHE_DO_KIEM_TRA_MA_TRAN:
            pham_vi_kq = ban_ghi.get("pham_vi", {}) or {}
            dot_id_kq = str(pham_vi_kq.get("dot_kiem_tra_id", ""))
            dot_kq = tim_dot_kiem_tra_theo_id(dot_id_kq)
            mo_dap_an_kq = bool(dot_kq and dot_kq.get("mo_dap_an", False))

            if not mo_dap_an_kq:
                diem_khoa = float(
                    ban_ghi.get(
                        "diem_chinh_thuc",
                        ban_ghi.get("diem", 0)
                    ) or 0
                )
                thang_khoa = float(ban_ghi.get("thang_diem", 10) or 10)
                k1, k2 = st.columns(2)
                with k1:
                    st.metric("Điểm", f"{diem_khoa:.2f}/{thang_khoa:.0f}")
                with k2:
                    st.metric(
                        "Thời gian làm",
                        f"{float(ban_ghi.get('thoi_luong_lam_phut', 0) or 0):.1f} phút"
                    )

                st.caption(
                    f"Bắt đầu: {ban_ghi.get('bat_dau_luc', '') or '—'} • "
                    f"Nộp bài: {ban_ghi.get('nop_bai_luc', ban_ghi.get('thoi_gian', '')) or '—'}"
                )
                st.info(
                    "🔒 **Đáp án đang được khóa.** Theo cài đặt của giáo viên, em hiện chỉ được xem điểm. "
                    "Khi GV mở đáp án, quay lại đúng đợt kiểm tra này để xem bài và lời giải."
                )

                kb1, kb2 = st.columns(2)
                with kb1:
                    if st.button(
                        "🧪 VỀ DANH SÁCH KIỂM TRA",
                        type="primary",
                        use_container_width=True,
                        key="hs_locked_back_tests"
                    ):
                        reset_bai_hs()
                        st.session_state.hs_mode = CHE_DO_KIEM_TRA_MA_TRAN
                        st.rerun()
                with kb2:
                    if st.button(
                        "🏠 VỀ MÀN HÌNH HỌC SINH",
                        use_container_width=True,
                        key="hs_locked_back_home"
                    ):
                        reset_bai_hs()
                        st.rerun()
                return

        r1, r2, r3, r4 = st.columns(4)

        with r1:
            diem_hien = float(
                ban_ghi.get(
                    "diem_chinh_thuc",
                    ban_ghi.get("diem", 0)
                )
                or 0
            )
            thang_hien = float(
                ban_ghi.get("thang_diem", 10)
                or 10
            )
            st.metric(
                "Điểm",
                f"{diem_hien:.2f}/{thang_hien:.0f}"
            )

        with r2:
            st.metric(
                "Tỉ lệ đúng",
                f"{ban_ghi.get('ti_le_dung_don_vi', 0):.0f}%"
            )

        with r3:
            st.metric(
                "Đơn vị đúng",
                f"{ban_ghi.get('dung_don_vi', 0)}/"
                f"{ban_ghi.get('tong_don_vi', 0)}"
            )

        with r4:
            phut_lam = float(
                ban_ghi.get("thoi_luong_lam_phut", 0) or 0
            )
            st.metric(
                "Thời gian làm",
                f"{phut_lam:.1f} phút"
            )

        st.caption(
            f"Bắt đầu: {ban_ghi.get('bat_dau_luc', '') or '—'} • "
            f"Nộp bài: {ban_ghi.get('nop_bai_luc', ban_ghi.get('thoi_gian', '')) or '—'}"
        )

        # --------------------------------------------------
        # 1. NĂNG LỰC CỦA EM
        # --------------------------------------------------
        weak_now = tom_tat_diem_yeu(
            profile_moi,
            5
        )

        nl_summary = tom_tat_nang_luc_chi_bao_hoc_sinh(
            profile_moi
        )

        st.markdown("---")
        st.header(
            "🧠 NĂNG LỰC CỦA EM"
        )

        cols_nl_hs = st.columns(
            3
        )

        for idx_nl, nl_item in enumerate(
            nl_summary
        ):
            with cols_nl_hs[
                idx_nl
            ]:
                with st.container(
                    border=True
                ):
                    st.markdown(
                        f"### {nl_item['nang_luc']}"
                    )

                    st.metric(
                        "Tỉ lệ đúng",
                        (
                            f"{nl_item['ti_le_dung'] * 100:.0f}%"
                            if nl_item[
                                "so_lan"
                            ]
                            else "—"
                        )
                    )

                    st.caption(
                        f"{nl_item['trang_thai']} • "
                        f"{nl_item['so_lan']} lần đánh giá"
                    )

        nl_can_ho_tro = [
            x
            for x in nl_summary
            if x[
                "trang_thai"
            ] in [
                "Cần hỗ trợ",
                "Đang củng cố"
            ]
        ]

        # --------------------------------------------------
        # 2. NHẬN XÉT
        # --------------------------------------------------
        st.markdown("---")
        st.header(
            "💬 NHẬN XÉT"
        )

        ti_le_luot = float(
            ban_ghi.get(
                "ti_le_dung_don_vi",
                0
            )
            or 0
        )

        if ti_le_luot >= 80:
            st.success(
                "🌟 **Bài làm tốt.** Em đã nắm khá chắc phần lớn nhiệm vụ trong lượt này."
            )

        elif ti_le_luot >= 60:
            st.info(
                "👍 **Bài làm khá.** Em đã có nền tảng nhưng vẫn còn một số năng lực/chỉ báo cần củng cố."
            )

        else:
            st.warning(
                "⚠️ **Em chưa đạt ở một số nội dung quan trọng.** "
                "Nên củng cố trước khi tăng độ khó."
            )

        if nl_can_ho_tro:
            st.write(
                "**Năng lực cần chú ý:**"
            )

            for nl_item in nl_can_ho_tro:
                st.markdown(
                    f"**• {nl_item['nang_luc']}**"
                )

                cb_yeu = nl_item.get(
                    "chi_bao_yeu",
                    []
                )

                if cb_yeu:
                    for cb in cb_yeu[:3]:
                        st.write(
                            f"   - **{cb['chi_bao']}**: "
                            f"{round(cb['ti_le_dung'] * cb['so_lan'])} lần đúng "
                            f"qua {cb['so_lan']} lần "
                            f"({cb['ti_le_dung'] * 100:.0f}%)."
                        )
                else:
                    st.caption(
                        "Chưa đủ dữ liệu để kết luận chính xác chỉ báo yếu."
                    )
        else:
            st.success(
                "Chưa phát hiện thành phần năng lực nào yếu rõ ràng."
            )

        # --------------------------------------------------
        # 3. GỢI Ý LÀM BÀI
        # --------------------------------------------------
        st.markdown("---")
        st.header(
            "🧭 GỢI Ý LÀM BÀI"
        )

        # Chỉ lấy kiến thức nền chưa đạt / đang yếu.
        weak_chua_dat = [
            x
            for x in weak_now
            if (
                int(
                    x.get(
                        "so_lan",
                        0
                    )
                    or 0
                ) < 4
                or float(
                    x.get(
                        "ti_le_dung",
                        0
                    )
                    or 0
                ) < 0.80
            )
        ]

        if weak_chua_dat:
            st.markdown(
                "### 📚 Kiến thức nền cần củng cố"
            )

            st.caption(
                "Ba kiến thức cần ưu tiên nhất được hiển thị trước."
            )

            yccd_map = {}

            for x in weak_chua_dat:
                ten = str(x.get("yccd", "")).strip()

                if not ten:
                    continue

                ti_le = float(x.get("ti_le_dung", 0) or 0)

                if (
                    ten not in yccd_map
                    or ti_le < float(
                        yccd_map[ten].get("ti_le_dung", 1) or 1
                    )
                ):
                    yccd_map[ten] = dict(x)

            ds_uu_tien = sorted(
                yccd_map.values(),
                key=lambda x: (
                    float(x.get("ti_le_dung", 0) or 0),
                    -int(x.get("so_lan", 0) or 0)
                )
            )

            # Hiện 3 kiến thức ưu tiên nhất.
            for i_w, x in enumerate(
                ds_uu_tien[:3],
                start=1
            ):
                st.markdown(
                    f"**{i_w}. {x.get('yccd', '')}**"
                )

                st.caption(
                    f"Mức độ: {x.get('muc_do', '')} • "
                    f"Hiện tại: {x.get('ti_le_dung', 0) * 100:.0f}% đúng"
                )

            # Sau mục số 3 vẫn có một thanh chọn để xem các kiến thức tiếp theo.
            con_lai = ds_uu_tien[3:]

            if con_lai:
                options_more = [
                    "— Chọn để xem kiến thức nền tiếp theo —"
                ] + [
                    f"{i + 4}. {x.get('yccd', '')}"
                    for i, x in enumerate(con_lai)
                ]

                st.caption(
                    f"Còn {len(con_lai)} kiến thức nền khác. Em có thể bấm chọn để xem tiếp."
                )

                chon_more = st.selectbox(
                    "⌄ Xem thêm các kiến thức nền cần củng cố",
                    options_more,
                    key="hs_weak_more_select"
                )

                if chon_more != options_more[0]:
                    idx_more = options_more.index(chon_more) - 1
                    x_more = con_lai[idx_more]

                    with st.container(border=True):
                        st.markdown(
                            f"**{x_more.get('yccd', '')}**"
                        )

                        st.write(
                            f"**Mức độ:** {x_more.get('muc_do', '')}"
                            f"  •  **Hiện tại:** "
                            f"{x_more.get('ti_le_dung', 0) * 100:.0f}% đúng"
                        )

                        if x_more.get("nang_luc"):
                            st.caption(
                                f"Năng lực liên quan: {x_more.get('nang_luc', '')}"
                            )

                        cb_more = str(
                            x_more.get("chi_bao", "")
                        ).strip()

                        if cb_more:
                            st.caption(
                                f"Chỉ báo: {cb_more}"
                            )

        else:
            st.success(
                "Không còn kiến thức nền yếu rõ ràng trong dữ liệu hiện có."
            )

        st.markdown(
            "### ✅ Em nên làm gì tiếp theo?"
        )

        if nl_can_ho_tro:
            ten_nl = ", ".join(
                x[
                    "nang_luc"
                ]
                for x in nl_can_ho_tro
            )

            st.info(
                f"Ưu tiên luyện thêm câu thuộc **{ten_nl}**; "
                "làm câu mới cùng chỉ báo/kĩ năng, đọc kĩ dữ kiện và "
                "đối chiếu lại kiến thức nền trước khi chọn đáp án."
            )
        else:
            st.info(
                "Tiếp tục luyện câu mới ở mức Thông hiểu/Vận dụng "
                "để kiểm tra độ bền kiến thức và khả năng vận dụng."
            )

        # AI chỉ gọi khi học sinh thực sự muốn giải thích sâu hơn.
        st.markdown(
            "### 🤖 Muốn AI giải thích sâu hơn?"
        )
        st.caption(
            "AI chỉ được gọi khi em bấm nút dưới đây. "
            "Điểm số, hồ sơ tiến bộ và cá nhân hóa đề không phụ thuộc vào AI."
        )

        if st.session_state.hs_ai_feedback is None:
            if st.button(
                "✨ NHỜ AI PHÂN TÍCH LỖI & DẶN DÒ",
                type="secondary",
                use_container_width=True,
                key="hs_ai_analyze_optional"
            ):
                with st.spinner(
                    "AI đang phân tích các lỗi sai của em..."
                ):
                    fb_moi, err_ai = ai_co_van_sau_bai_lam(
                        hs_id_chuan,
                        hs_ten,
                        ban_ghi,
                        profile_moi
                    )

                if fb_moi is not None:
                    st.session_state.hs_ai_feedback = fb_moi
                    st.session_state.hs_ai_feedback_error = ""
                    st.rerun()
                else:
                    st.session_state.hs_ai_feedback_error = str(
                        err_ai or ""
                    )

        fb = st.session_state.hs_ai_feedback

        if fb:
            st.markdown(
                "## 🤖 CỐ VẤN AI"
            )

            nhan_xet = str(
                fb.get(
                    "nhan_xet_tong_quan",
                    ""
                )
            ).strip()

            if nhan_xet:
                st.success(
                    "🌟 **NHẬN XÉT CHUNG**\n\n"
                    + nhan_xet
                )

            loi = fb.get(
                "loi_co_ban",
                []
            ) or []

            if loi:
                st.markdown(
                    "### ⚠️ Em đang vướng ở đâu?"
                )

                for i_fb, x in enumerate(
                    loi,
                    start=1
                ):
                    st.warning(
                        f"**{i_fb}. {str(x).strip()}**"
                    )

            can_on = fb.get(
                "can_on_lai",
                []
            ) or []

            if can_on:
                st.markdown(
                    "### 📚 Kiến thức cần ôn lại"
                )

                for i_fb, x in enumerate(
                    can_on,
                    start=1
                ):
                    st.info(
                        f"**{i_fb}. {str(x).strip()}**"
                    )

            dan_do = fb.get(
                "dan_do",
                []
            ) or []

            if dan_do:
                st.markdown(
                    "### 💡 Cách làm tốt hơn"
                )

                for i_fb, x in enumerate(
                    dan_do,
                    start=1
                ):
                    st.write(
                        f"#### {i_fb}. {str(x).strip()}"
                    )

            ke_hoach = str(
                fb.get(
                    "ke_hoach_luot_tiep",
                    ""
                )
            ).strip()

            if ke_hoach:
                st.markdown(
                    "### 🎯 Kế hoạch lượt tiếp"
                )

                st.info(
                    "**" + ke_hoach + "**"
                )

        elif st.session_state.get(
            "hs_ai_feedback_error"
        ):
            st.warning(
                "AI tạm thời chưa phản hồi được. "
                "Kết quả học tập và hồ sơ tiến bộ của em vẫn đã được lưu."
            )

        # --------------------------------------------------
        # CHỮA TỪNG CÂU
        # --------------------------------------------------
        st.markdown("---")
        st.header(
            "📋 XEM LẠI BÀI LÀM"
        )
        st.caption(
            "Xem câu đúng, câu cần sửa, đáp án chuẩn và giải thích để rút kinh nghiệm."
        )

        for item in ban_ghi.get(
            "chi_tiet",
            []
        ):
            i = item.get(
                "stt",
                ""
            )
            q = item.get(
                "cau_snapshot",
                {}
            )

            if item.get(
                "dung_toan_cau"
            ):
                st.success(
                    f"✅ Câu {i}: Đúng"
                )
            else:
                st.error(
                    f"❌ Câu {i}: Cần xem lại"
                )

            if q.get(
                "tinh_huong"
            ):
                st.write(
                    "**Tình huống / dữ liệu:**",
                    q.get(
                        "tinh_huong",
                        ""
                    )
                )

            st.write(
                "**Câu hỏi:**",
                q.get(
                    "cau_hoi",
                    ""
                )
            )

            if q.get(
                "dang_cau"
            ) == "Đúng / Sai":
                for ky, nd in zip(
                    ["a", "b", "c", "d"],
                    q.get(
                        "nhan_dinh_meta",
                        []
                    )
                    or []
                ):
                    st.write(
                        f"**{ky})** "
                        + bo_nhan_dinh(
                            nd.get(
                                "noi_dung",
                                ""
                            ),
                            ky
                        )
                    )

                    st.caption(
                        f"Đáp án: {nd.get('dap_an', '')} • "
                        f"Giải thích: {nd.get('giai_thich', '')}"
                    )

                # Một số file nguồn có lời giải chung cho cả câu Đúng/Sai thay vì tách từng ý.
                # Nếu từng ý chưa có giải thích, vẫn hiển thị nguyên lời giải nguồn/đã lưu.
                if str(q.get("giai_thich", "") or "").strip() and not any(
                    str(nd.get("giai_thich", "") or "").strip()
                    for nd in (q.get("nhan_dinh_meta", []) or [])
                ):
                    st.write("**Hướng dẫn giải:**", q.get("giai_thich", ""))

            else:
                if q.get(
                    "lua_chon"
                ):
                    for lc in q.get(
                        "lua_chon",
                        []
                    ):
                        st.write(
                            lc
                        )

                st.write(
                    "**Bạn trả lời:**",
                    item.get(
                        "hoc_sinh_tra_loi",
                        ""
                    )
                    or "Chưa trả lời"
                )

                st.write(
                    "**Đáp án đúng:**",
                    item.get(
                        "dap_an_chuan",
                        ""
                    )
                )

                st.write(
                    "**Giải thích chuẩn:**",
                    q.get(
                        "giai_thich",
                        ""
                    )
                )

            st.divider()

        # --------------------------------------------------
        # LƯỢT TIẾP THEO
        # --------------------------------------------------
        cnext1, cnext2 = st.columns(2)

        with cnext1:
            if st.button(
                "🔄 LUYỆN LƯỢT MỚI",
                type="primary",
                use_container_width=True
            ):
                reset_bai_hs()
                st.rerun()

        with cnext2:
            if st.button(
                "🏠 ĐỔI CHẾ ĐỘ LUYỆN",
                use_container_width=True
            ):
                reset_bai_hs()
                st.rerun()

    st.divider()

    if st.button(
        "⬅️ Về trang chủ",
        use_container_width=True
    ):
        reset_bai_hs()
        st.session_state.vai_tro = None
        st.rerun()


# ==========================================================
# KHÓA KHU VỰC GIÁO VIÊN KHI TRIỂN KHAI
# ==========================================================
def cho_phep_vao_khu_vuc_giao_vien():
    # Khi chạy local/full mà chưa cấu hình mã, giữ hành vi cũ để GV phát triển app.
    try:
        ma_bao_ve = str(st.secrets.get("TEACHER_ACCESS_CODE", "") or "").strip()
    except Exception:
        ma_bao_ve = ""

    if not ma_bao_ve:
        return True

    if st.session_state.get("gv_da_xac_thuc_trien_khai"):
        return True

    hien_thi_dau_trang_tram_sinh_hoc("giaovien")
    st.subheader("🔐 ĐĂNG NHẬP GIÁO VIÊN")
    st.caption("Khu vực này được tách khỏi giao diện học sinh và yêu cầu mã bảo vệ phía máy chủ.")
    ma_nhap = st.text_input(
        "Mã truy cập giáo viên",
        type="password",
        key="gv_deploy_access_code"
    )
    if st.button(
        "ĐĂNG NHẬP GIÁO VIÊN",
        type="primary",
        use_container_width=True,
        key="gv_deploy_login"
    ):
        if ma_nhap == ma_bao_ve:
            st.session_state.gv_da_xac_thuc_trien_khai = True
            st.rerun()
        else:
            st.error("Mã truy cập không đúng.")
    return False


# ĐIỀU HƯỚNG
# ==========================================================
if APP_MODE == "student":
    # Bản gửi học sinh: server không render trang chủ chọn vai trò và không có đường vào menu GV.
    st.session_state.vai_tro = "hocsinh"
    hoc_sinh()

elif APP_MODE == "teacher":
    # Bản giáo viên: nên đặt TEACHER_ACCESS_CODE trong secrets hoặc dùng OIDC ở lớp triển khai.
    st.session_state.vai_tro = "giaovien"
    if cho_phep_vao_khu_vuc_giao_vien():
        giao_vien()
else:
    # Chế độ full chỉ dùng local/nội bộ khi đang phát triển.
    if st.session_state.vai_tro is None:
        trang_chu()
    elif st.session_state.vai_tro == "giaovien":
        giao_vien()
    elif st.session_state.vai_tro == "hocsinh":
        hoc_sinh()
